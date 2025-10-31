#!/usr/bin/env python3
"""
Reading Speed Measurement Tab for PDF Tools
Measures reading speed in WPM with comprehension testing
"""

import os
import sys
import time
import random
import re
import csv
from datetime import datetime
from typing import List, Dict, Optional, Tuple

import sqlite3

# Import reading position database
try:
    from src.reading_position_db import ReadingPositionDB
except ImportError:
    try:
        from reading_position_db import ReadingPositionDB
    except ImportError:
        ReadingPositionDB = None  # Graceful degradation if module not found
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QFormLayout,
    QPushButton, QLabel, QFileDialog, QSpinBox, QMessageBox,
    QGroupBox, QTextEdit, QLineEdit, QListWidget, QListWidgetItem,
    QProgressBar, QSplitter, QScrollArea, QFrame, QGridLayout,
    QTableWidget, QTableWidgetItem, QHeaderView, QTabWidget,
    QSlider, QComboBox, QDialog, QSizePolicy, QToolButton
)
from PySide6.QtCore import Qt, QTimer, Signal, QSize, QPropertyAnimation, QEasingCurve, QUrl
from PySide6.QtGui import QFont, QPixmap, QIcon, QDesktopServices


try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from reading_position_db import ReadingPositionDB
except ImportError:
    ReadingPositionDB = None


class ReadingSpeedTab(QWidget):
    """Reading Speed Measurement Tab"""

    def __init__(self, localization=None, recent_books_manager=None, parent=None):
        super().__init__(parent)
        self.localization = localization
        self.recent_books_manager = recent_books_manager
        self.parent_window = parent

        # State variables
        self.pdf_path = None
        self.pdf_doc = None
        self.avg_words_per_page = 0
        self.total_pages = 0
        self.start_time = 0
        self.elapsed_time_seconds = 0
        self.reading_timer = QTimer(self)
        self.reading_timer.timeout.connect(self.update_timer_display)
        self.sample_text = ""
        self.comprehension_questions = []

        self.words_per_page_list = []  # Optional per-page word counts from DOCX

        # Reading session state management
        self.is_reading_paused = False
        self.is_reading_stopped = False
        self.current_reading_page = 1
        self.paused_time = 0  # Time when paused
        self.total_paused_duration = 0  # Total time spent paused

        # Database persistence for reading positions
        self.position_db = None
        self.current_document_path = None
        self.current_document_type = None
        self.auto_save_counter = 0
        self.auto_save_interval = 10  # Save every 10 words
        self.total_words_read_session = 0  # Words read in current session

        # Initialize database (with graceful degradation)
        try:
            if ReadingPositionDB is not None:
                self.position_db = ReadingPositionDB()
                print("✅ Reading position database initialized")
            else:
                print("⚠️ Reading position database not available (module not found)")
        except Exception as e:
            print(f"⚠️ Could not initialize reading position database: {e}")
            self.position_db = None

        # Get current theme
        self.current_theme = self.get_current_theme()

        self.init_ui()

    def get_current_theme(self):
        """Get current theme from parent application"""
        try:
            if hasattr(self.parent_window, 'settings'):
                return self.parent_window.settings.get("theme", "light")
            return "light"
        except:
            return "light"

    def get_theme_colors(self):
        """Get theme-appropriate colors"""
        if self.current_theme == "dark":
            return {
                'bg_primary': '#1e1e1e',
                'bg_secondary': '#2d2d2d',
                'bg_tertiary': '#3d3d3d',
                'text_primary': '#ffffff',
                'text_secondary': '#cccccc',
                'text_muted': '#999999',
                'border': '#555555',
                'accent_blue': '#2196F3',
                'accent_green': '#4CAF50',
                'accent_orange': '#FF9800',
                'accent_purple': '#9C27B0',
                'accent_red': '#F44336'
            }
        else:
            return {
                'bg_primary': '#ffffff',
                'bg_secondary': '#fafafa',
                'bg_tertiary': '#f5f5f5',
                'text_primary': '#333333',
                'text_secondary': '#666666',
                'text_muted': '#999999',
                'border': '#e0e0e0',
                'accent_blue': '#2196F3',
                'accent_green': '#4CAF50',
                'accent_orange': '#FF9800',
                'accent_purple': '#9C27B0',
                'accent_red': '#F44336'
            }

    def get_group_box_style(self, accent_color):
        """Get theme-aware group box styling"""
        colors = self.get_theme_colors()
        return f"""
            QGroupBox {{
                font-weight: bold;
                font-size: 16px;
                color: {accent_color};
                border: 2px solid {accent_color};
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: {colors['bg_secondary']};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: {accent_color};
                font-weight: bold;
                background-color: {colors['bg_secondary']};
            }}
        """

    def get_table_style(self, header_color):
        """Get theme-aware table styling"""
        colors = self.get_theme_colors()
        return f"""
            QTableWidget {{
                gridline-color: {colors['border']};
                background-color: {colors['bg_primary']};
                alternate-background-color: {colors['bg_tertiary']};
                selection-background-color: {colors['bg_secondary']};
                border: 1px solid {colors['border']};
                border-radius: 6px;
                font-size: 12px;
                color: {colors['text_primary']};
            }}
            QHeaderView::section {{
                background-color: {header_color};
                color: white;
                padding: 8px;
                border: none;
                font-weight: bold;
                font-size: 12px;
            }}
            QTableWidget::item {{
                padding: 6px;
                border-bottom: 1px solid {colors['border']};
                color: {colors['text_primary']};
            }}
        """

    # Fast Reading Trainer Methods (defined early to avoid AttributeError)
    def select_trainer_file(self, file_type):
        """Unified file selection for trainer"""
        if file_type == "txt":
            self.select_txt_file()
        elif file_type == "docx":
            self.select_docx_file()
        elif file_type == "pdf":
            self.select_pdf_for_trainer()

    def select_txt_file(self):
        """Select a TXT file for fast reading training"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_txt_file") if self.localization else "Select TXT File",
            "",
            "Text Files (*.txt)"
        )

        if file_path:
            self.load_trainer_text_file(file_path)

    def select_docx_file(self):
        """Select a DOCX file for fast reading training"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_docx_file") if self.localization else "Select DOCX File",
            "",
            "Word Documents (*.docx)"
        )

        if file_path:
            self.load_trainer_docx_file(file_path)

    def select_pdf_for_trainer(self):
        """Select a PDF file for fast reading training"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file") if self.localization else "Select PDF File",
            "",
            "PDF Files (*.pdf)"
        )

        if file_path:
            self.load_trainer_pdf_file(file_path)

    def load_trainer_text_file(self, file_path):
        """Load text from a TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            self.prepare_trainer_text(text, file_path)

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error loading TXT file: {str(e)}"
            )

    def load_trainer_docx_file(self, file_path):
        """Load text from a DOCX file with start page support"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    "python-docx library is required for DOCX files. Please install it with: pip install python-docx"
                )
                return

            # Load DOCX document
            doc = Document(file_path)

            # Get start page from spinbox (convert to 0-based index)
            start_page_num = self.start_page_spinbox.value() - 1 if hasattr(self, 'start_page_spinbox') else 0

            print(f"\n📄 Loading DOCX for training:")
            print(f"   File: {os.path.basename(file_path)}")
            print(f"   Total paragraphs: {len(doc.paragraphs)}")
            print(f"   Start page (user setting): {start_page_num + 1}")

            # Extract all text and track words per "page" (approximation)
            # For DOCX, we'll estimate pages based on word count (assume ~250 words per page)
            WORDS_PER_PAGE_ESTIMATE = 250

            all_text = ""
            page_word_counts = []
            current_page_words = 0

            for paragraph in doc.paragraphs:
                para_text_raw = paragraph.text
                # Normalize similar to prepare_trainer_text
                para_text_norm = re.sub(r'\s+', ' ', para_text_raw.strip())
                if para_text_norm:
                    all_text += para_text_norm + " "
                    # Count words in this normalized paragraph
                    para_words = len(para_text_norm.split())
                    current_page_words += para_words
                    # If we've accumulated enough words for a "page", record it
                    if current_page_words >= WORDS_PER_PAGE_ESTIMATE:
                        page_word_counts.append(current_page_words)
                        current_page_words = 0

            # Add remaining words as last page
            if current_page_words > 0:
                page_word_counts.append(current_page_words)

            # Store per-page word counts for later start-page recalculation
            self.trainer_page_word_counts = list(page_word_counts)
            self.trainer_total_pages = len(page_word_counts)


            if not all_text.strip():
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    self.localization.get_text("no_text_found") if self.localization else "No text found in the document"
                )
                return

            # Calculate the word index where the start page begins
            estimated_pages = len(page_word_counts)
            start_page_num = min(start_page_num, estimated_pages - 1)  # Ensure valid range
            start_word_index = sum(page_word_counts[:start_page_num]) if start_page_num > 0 else 0

            print(f"   Estimated pages: {estimated_pages}")
            print(f"   Words before start page: {start_word_index}")
            print(f"   Starting training from word index: {start_word_index}")

            # Prepare the text with start word index
            self.prepare_trainer_text(all_text, file_path, start_word_index=start_word_index, start_page=start_page_num + 1)

            # Update Pages/Minute tab defaults and auto-select mode based on available data
            try:
                if hasattr(self, 'ppm_modes_tabs'):
                    self._ppm_refresh_defaults()
                    self._ppm_autoselect_mode_based_on_data()
            except Exception:
                pass


        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error loading DOCX file: {str(e)}"
            )

    def load_trainer_pdf_file(self, file_path):
        """Load text from a PDF file"""
        if not fitz:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                "PyMuPDF (fitz) is required for PDF processing"
            )
            return

        try:
            # Open PDF
            doc = fitz.open(file_path)

            # Get start page from spinbox (convert to 0-based index)
            start_page_num = self.start_page_spinbox.value() - 1 if hasattr(self, 'start_page_spinbox') else 0
            start_page_num = max(0, min(start_page_num, doc.page_count - 1))  # Ensure valid range

            print(f"\n📄 Loading PDF for training:")
            print(f"   File: {os.path.basename(file_path)}")
            print(f"   Total pages: {doc.page_count}")
            print(f"   Start page (user setting): {start_page_num + 1}")

            # Extract ALL pages to maintain proper page numbering and word count
            all_text = ""
            page_word_counts = []  # Track words per page to calculate start word index

            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                page_text_raw = page.get_text("text") or ""
                # Normalize page text similar to prepare_trainer_text
                page_text_norm = re.sub(r'\s+', ' ', page_text_raw.strip())
                all_text += page_text_norm + " "

                # Count words in this normalized page
                page_words = len(page_text_norm.split())
                page_word_counts.append(page_words)

            # Store per-page word counts for later start-page recalculation
            self.trainer_page_word_counts = list(page_word_counts)
            self.trainer_total_pages = len(page_word_counts)


            doc.close()

            if not all_text.strip():
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    self.localization.get_text("no_extractable_text") if self.localization else "No extractable text found in the PDF"
                )
                return

            # Update Pages/Minute tab defaults and auto-select mode based on available data
            try:
                if hasattr(self, 'ppm_modes_tabs'):
                    self._ppm_refresh_defaults()
                    self._ppm_autoselect_mode_based_on_data()
            except Exception:
                pass


            # Calculate the word index where the start page begins
            start_word_index = sum(page_word_counts[:start_page_num]) if start_page_num > 0 else 0

            print(f"   Words before start page: {start_word_index}")
            print(f"   Starting training from word index: {start_word_index}")

            # Prepare the text and then adjust the starting word index
            self.prepare_trainer_text(all_text, file_path, start_word_index=start_word_index, start_page=start_page_num + 1)

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error loading PDF file: {str(e)}"
            )

    def prepare_trainer_text(self, text, file_path, start_word_index=0, start_page=1):
        """
        Prepare text for fast reading training

        Args:
            text: The full text to train with
            file_path: Path to the source file
            start_word_index: Word index to start training from (default: 0)
            start_page: Page number to start from (for display purposes, default: 1)
        """
        try:
            # Clean and prepare text
            # Remove extra whitespace and normalize
            text = re.sub(r'\s+', ' ', text.strip())

            # Split into words
            words = text.split()

            if len(words) < 10:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    self.localization.get_text("text_too_short") if self.localization else "Text is too short for training (minimum 10 words required)"
                )
                return

            # Validate start_word_index
            if start_word_index >= len(words):
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    f"Start page is beyond the document length. Starting from beginning instead."
                )
                start_word_index = 0
                start_page = 1

            # Store training data
            self.trainer_text = text
            self.trainer_words = words
            self.current_word_index = start_word_index  # Start from specified word index
            self.trainer_file_path = file_path  # Store file path for session saving
            self.session_start_time = datetime.now()  # Track session start time
            self.trainer_start_page = start_page  # Store start page for display

            # Track current document for DB/auto-save
            self.current_document_path = file_path
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                self.current_document_type = 'pdf'
            elif ext in ('.doc', '.docx'):
                self.current_document_type = 'docx'
            else:
                self.current_document_type = 'txt'
            self.current_reading_page = start_page

            print(f"✅ Text prepared for training:")
            print(f"   Total words: {len(words)}")
            print(f"   Starting from word index: {start_word_index}")
            print(f"   Starting from page: {start_page}")
            print(f"   Document type: {self.current_document_type}")

            # Initialize DB row for auto-save if available
            if self.position_db is not None:
                try:
                    init_ok = self.position_db.save_reading_position(
                        self.current_document_path,
                        {
                            'document_type': self.current_document_type,
                            'last_page': start_page if self.current_document_type == 'pdf' else None,
                            'last_word_index': self.current_word_index,
                            'training_mode': self.get_current_training_mode() if hasattr(self, 'get_current_training_mode') else 'standard',
                            'wpm_setting': self.trainer_wpm_spinbox.value() if hasattr(self, 'trainer_wpm_spinbox') else 300,
                            'words_per_glance': self.wpg_spinbox.value() if hasattr(self, 'wpg_spinbox') else 3,
                            'total_words_read': 0,
                        }
                    )
                    print(f"🟢 DB init save: {'Success' if init_ok else 'Skipped/Failed'}")
                except Exception as db_e:
                    print(f"⚠️ DB init save error: {db_e}")

            # Update UI
            filename = os.path.basename(file_path)
            self.trainer_file_display.setText(filename)

            # Enable start button
            self.btn_start_training.setEnabled(True)

            # Update display
            if start_word_index > 0:
                self.word_display.setText(
                    f"{self.localization.get_text('ready_to_train') if self.localization else 'Ready to start training!'}\n"
                    f"Starting from page {start_page}"
                )
            else:
                self.word_display.setText(self.localization.get_text("ready_to_train") if self.localization else "Ready to start training!")

            # Update progress
            self.update_training_progress(value=start_word_index, maximum=len(words), visible=True)
            remaining_words = len(words) - start_word_index
            self.progress_info.setText(
                f"{remaining_words} {self.localization.get_text('words_loaded') if self.localization else 'words remaining'} "
                f"(starting from page {start_page})" if start_word_index > 0
                else f"{len(words)} {self.localization.get_text('words_loaded') if self.localization else 'words loaded'}"
            )

            # Success message
            start_info = f"\n🎯 Starting from page {start_page} (word {start_word_index + 1})" if start_word_index > 0 else ""
            QMessageBox.information(
                self,
                self.localization.get_text("text_loaded") if self.localization else "Text Loaded",
                f"{self.localization.get_text('text_loaded_successfully') if self.localization else 'Text loaded successfully!'}\n\n"
                f"📄 {filename}\n"
                f"📝 {len(words)} {self.localization.get_text('words') if self.localization else 'total words'}"
                f"{start_info}\n\n"
                f"{self.localization.get_text('adjust_settings_and_start') if self.localization else 'Adjust your settings and click Start Training!'}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error preparing text: {str(e)}"
            )

    def start_fast_reading_training(self):
        """
        Start the fast reading training session.

        This method initializes the training session for all modes:
        1. Validates document is loaded
        2. Resets state variables
        3. Calculates timer interval based on WPM and WPG
        4. Starts the training timer
        5. Displays first word/group

        Timer Interval Formula:
        interval_ms = (wpg * 60 * 1000) / wpm

        Example: 3 WPG at 300 WPM = (3 * 60 * 1000) / 300 = 600ms

        Mode-Specific Behavior:
        - Standard: Displays wpg words together
        - RSVP: Displays 1 word (wpg locked to 1)
        - Scrolling: Displays wpg words per line (scrolling pending)
        - Chunking: Displays wpg words as phrase
        - Elimination: Displays wpg words at high speed (400+ WPM)
        """
        try:
            # Validate document loaded
            if not self.trainer_words:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    self.localization.get_text("no_text_loaded") if self.localization else "No text loaded for training"
                )
                return

            # Recalculate start index from requested start page using per-page mapping
            requested_start_page = self.start_page_spinbox.value() if hasattr(self, 'start_page_spinbox') else 1
            start_page_num = max(0, requested_start_page - 1)
            calc_start_index = 0
            if hasattr(self, 'trainer_page_word_counts') and self.trainer_page_word_counts and start_page_num > 0:
                calc_start_index = sum(self.trainer_page_word_counts[:start_page_num])
            # Clamp to available words
            total_words = len(self.trainer_words)
            if calc_start_index >= total_words:
                calc_start_index = max(0, total_words - 1)
            print(f"🔎 Start-page debug: requested page={requested_start_page}, computed start_word_index={calc_start_index}, pre-existing index={getattr(self, 'current_word_index', None)}")
            # Always start at the computed index to honor the user's setting
            self.current_word_index = calc_start_index
            self.current_reading_page = max(1, requested_start_page)

            # Prepare training state (do NOT reset current_word_index so start page works)
            if not hasattr(self, 'current_word_index') or self.current_word_index is None:
                self.current_word_index = 0
            self.training_active = True  # Mark training as active
            print(f"✅ Verify start index: {self.current_word_index} (actual) == {calc_start_index} (computed) -> {self.current_word_index == calc_start_index}")

            self.is_training_paused = False  # Not paused initially
            print(f"🔵 DEBUG: start_fast_reading_training -> starting at word index {self.current_word_index}")

            # Update UI - change unified button to Pause state, enable stop
            self.update_unified_control_button("pause")
            self.btn_unified_control.setEnabled(True)
            self.btn_stop_training.setEnabled(True)

            # Disable settings during training to prevent mid-session changes
            self.wpg_spinbox.setEnabled(False)
            self.trainer_wpm_spinbox.setEnabled(False)
            self.font_size_spinbox.setEnabled(False)

            # Calculate timing based on WPM and WPG
            wpm = self.trainer_wpm_spinbox.value()  # Words per minute
            wpg = self.wpg_spinbox.value()  # Words per glance

            # Calculate interval between word groups
            # WPM = words per minute, so words per second = WPM / 60
            # Time per word group = wpg / (WPM / 60) = (wpg * 60) / WPM
            # Convert to milliseconds by multiplying by 1000
            interval_ms = int((wpg * 60 * 1000) / wpm)

            # Update display styling with selected font
            font_size = self.font_size_spinbox.value()
            font_family = self.font_family_combo.currentText()
            self.word_display.setStyleSheet(f"""
                QLabel {{
                    color: #00FF00;
                    font-size: {font_size}px;
                    font-weight: bold;
                    font-family: '{font_family}', monospace;
                    padding: 20px;
                    background-color: transparent;
                    border: none;
                }}
            """)

            # Start training timer
            self.training_timer.start(interval_ms)

            # Display first word group immediately
            self.display_next_word_group()

            # Update progress info
            self.progress_info.setText(f"{self.localization.get_text('training_active') if self.localization else 'Training active'} - {wpm} WPM, {wpg} WPG")

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error starting training: {str(e)}"
            )

    def is_arabic_text(self, text):
        """Check if text contains Arabic characters"""
        arabic_pattern = r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]'
        return bool(re.search(arabic_pattern, text))

    def display_next_word_group(self):
        """
        Display the next group of words based on current training mode.

        This method is called by the training timer on each tick.
        It handles word display for all 5 training modes:

        1. Standard Mode: Displays wpg words together
        2. RSVP Mode: Displays 1 word (wpg=1) in center
        3. Scrolling Mode: Displays wpg words (scrolling animation pending)
        4. Chunking Mode: Displays wpg words as a phrase (3-7 words)
        5. Elimination Mode: Displays wpg words at high speed

        The method:
        - Checks if training is active and not paused
        - Extracts the next wpg words from trainer_words
        - Detects Arabic text for RTL formatting
        - Highlights focus word (middle word in group)
        - Updates the display
        - Increments current_word_index
        - Updates progress bar
        - Checks for completion
        """
        try:
            # Check training state - exit if not active or paused
            if not self.training_active or self.is_training_paused:
                return

            # Get words per glance setting
            # RSVP mode: wpg=1 (locked)
            # Chunking mode: wpg=3-7
            # Other modes: wpg=1-7 (user choice)
            wpg = self.wpg_spinbox.value()

            # Check if training is complete
            if self.current_word_index >= len(self.trainer_words):
                self.complete_training()
                return

            # Extract next word group from document
            # For RSVP: word_group will contain 1 word
            # For Chunking: word_group will contain 3-7 words (phrase)
            # For others: word_group will contain wpg words
            word_group = []
            for i in range(wpg):
                if self.current_word_index + i < len(self.trainer_words):
                    word_group.append(self.trainer_words[self.current_word_index + i])

            # Update font family and size from settings
            font_family = self.font_family_combo.currentText()
            font_size = self.font_size_spinbox.value()

            # Determine if we have Arabic text in the group
            has_arabic = any(self.is_arabic_text(word) for word in word_group)

            if has_arabic:
                # For Arabic text: Use CSS styling instead of HTML spans to preserve letter connections
                if len(word_group) > 1:
                    # Multi-word group: highlight the entire focus word
                    focus_index = len(word_group) // 2
                    focus_word = word_group[focus_index]

                    # Create display with focus word highlighted via CSS
                    display_parts = []
                    for i, word in enumerate(word_group):
                        if i == focus_index:
                            # Highlight entire focus word to preserve Arabic shaping
                            display_parts.append(f"<span style='color: #FF4444; font-weight: bold; text-decoration: underline;'>{word}</span>")
                        else:
                            display_parts.append(word)
                    display_text = " ".join(display_parts)
                else:
                    # Single Arabic word: highlight entire word to preserve shaping
                    if len(word_group) == 1:
                        word = word_group[0]
                        display_text = f"<span style='color: #FF4444; font-weight: bold; text-decoration: underline;'>{word}</span>"
                    else:
                        display_text = ""

                # Apply styling optimized for Arabic text
                self.word_display.setStyleSheet(f"""
                    QLabel {{
                        color: #00FF00;
                        font-size: {font_size}px;
                        font-weight: bold;
                        font-family: '{font_family}', 'IBM Plex Sans Arabic', 'Tajawal', sans-serif;
                        padding: 20px;
                        background-color: transparent;
                        border: none;
                        direction: rtl;
                    }}
                """)
            else:
                # For English/Latin text: Use character-level ORP highlighting
                if len(word_group) > 1:
                    # Multi-word group: highlight focus word's ORP
                    focus_index = len(word_group) // 2
                    focus_word = word_group[focus_index]

                    # Find the optimal recognition point (ORP) - usually around 1/3 into the word
                    if len(focus_word) > 2:
                        orp_index = max(0, min(len(focus_word) - 1, len(focus_word) // 3))
                        # Create highlighted word with proper HTML formatting
                        highlighted_word = (
                            focus_word[:orp_index] +
                            f"<span style='color: #FF4444; font-weight: bold; text-decoration: underline;'>{focus_word[orp_index]}</span>" +
                            focus_word[orp_index+1:]
                        )
                        word_group[focus_index] = highlighted_word

                    # Join words with proper spacing
                    display_text = " ".join(word_group)
                else:
                    # Single English word: highlight its ORP
                    if len(word_group) == 1:
                        word = word_group[0]
                        if len(word) > 2:
                            orp_index = max(0, min(len(word) - 1, len(word) // 3))
                            display_text = (
                                word[:orp_index] +
                                f"<span style='color: #FF4444; font-weight: bold; text-decoration: underline;'>{word[orp_index]}</span>" +
                                word[orp_index+1:]
                            )
                        else:
                            display_text = word_group[0]
                    else:
                        display_text = ""

                # Apply styling optimized for English text
                self.word_display.setStyleSheet(f"""
                    QLabel {{
                        color: #00FF00;
                        font-size: {font_size}px;
                        font-weight: bold;
                        font-family: '{font_family}', monospace;
                        padding: 20px;
                        background-color: transparent;
                        border: none;
                        direction: ltr;
                    }}
                """)

            # Set the formatted text
            self.word_display.setText(display_text)

            # Update progress
            self.current_word_index += wpg
            self.update_training_progress(value=self.current_word_index)

            # Update progress percentage
            progress_percent = (self.current_word_index / len(self.trainer_words)) * 100
            self.progress_info.setText(f"{progress_percent:.1f}% - {self.current_word_index}/{len(self.trainer_words)} {self.localization.get_text('words') if self.localization else 'words'}")

            # Auto-save reading position every N words (Task 7: Database persistence)
            self.auto_save_counter += wpg
            self.total_words_read_session += wpg

            if self.auto_save_counter >= self.auto_save_interval:
                self._auto_save_position()
                self.auto_save_counter = 0  # Reset counter

        except Exception as e:
            print(f"Error displaying word group: {e}")
            self.stop_fast_reading_training()

    def handle_unified_control_click(self):
        """
        Handle clicks on the unified control button.
        Routes to appropriate action based on current training state.
        """
        try:
            # Add visual feedback animation
            self.animate_button_click(self.btn_unified_control)

            control_state = self.btn_unified_control.property("control_state")

            if control_state == "start":
                # Start training
                self.start_fast_reading_training()
            elif control_state == "pause":
                # Pause training
                self.pause_resume_training()
            elif control_state == "resume":
                # Resume training
                self.pause_resume_training()

        except Exception as e:
            print(f"Error handling unified control click: {e}")

    def animate_button_click(self, button):
        """
        Add visual feedback animation when button is clicked.
        Creates a brief scale animation to indicate the click was registered.
        """
        try:
            # Create a property animation for the button's minimum height
            animation = QPropertyAnimation(button, b"minimumHeight")
            animation.setDuration(150)  # 150ms animation
            animation.setStartValue(button.minimumHeight())
            animation.setKeyValueAt(0.5, button.minimumHeight() - 5)  # Shrink slightly
            animation.setEndValue(button.minimumHeight())  # Return to original
            animation.setEasingCurve(QEasingCurve.InOutQuad)

            # Start the animation
            animation.start()

            # Store reference to prevent garbage collection
            if not hasattr(self, '_button_animations'):
                self._button_animations = []
            self._button_animations.append(animation)

            # Clean up old animations
            if len(self._button_animations) > 5:
                self._button_animations.pop(0)

        except Exception as e:
            print(f"Error animating button click: {e}")

    def update_unified_control_button(self, state):
        """
        Update the unified control button appearance based on training state.
        Icon-only version for compact display.

        States:
        - 'start': Not started (green, ▶️)
        - 'pause': Running (orange, ⏸️)
        - 'resume': Paused (blue, ▶️)
        """
        try:
            colors = self.get_theme_colors()

            if state == "start":
                self.btn_unified_control.setText("▶️")
                self.btn_unified_control.setToolTip(self.localization.get_text("start_training") if self.localization else "Start Training")
                self.btn_unified_control.setStyleSheet(f"""
                    QToolButton {{
                        background-color: {colors['accent_green']};
                        color: white;
                        border: none;
                        border-radius: 8px;
                        font-size: 28px;
                    }}
                    QToolButton:hover {{
                        background-color: #45a049;
                    }}
                    QToolButton:disabled {{
                        background-color: {colors['border']};
                    }}
                    QToolButton:pressed {{
                        background-color: #2E7D32;
                    }}
                """)
                self.btn_unified_control.setProperty("control_state", "start")

            elif state == "pause":
                self.btn_unified_control.setText("⏸️")
                self.btn_unified_control.setToolTip(self.localization.get_text("pause") if self.localization else "Pause Training")
                self.btn_unified_control.setStyleSheet(f"""
                    QToolButton {{
                        background-color: {colors['accent_orange']};
                        color: white;
                        border: none;
                        border-radius: 8px;
                        font-size: 28px;
                    }}
                    QToolButton:hover {{
                        background-color: #F57C00;
                    }}
                    QToolButton:disabled {{
                        background-color: {colors['border']};
                    }}
                    QToolButton:pressed {{
                        background-color: #E65100;
                    }}
                """)
                self.btn_unified_control.setProperty("control_state", "pause")

            elif state == "resume":
                self.btn_unified_control.setText("▶️")
                self.btn_unified_control.setToolTip(self.localization.get_text("resume") if self.localization else "Resume Training")
                self.btn_unified_control.setStyleSheet(f"""
                    QToolButton {{
                        background-color: {colors['accent_blue']};
                        color: white;
                        border: none;
                        border-radius: 8px;
                        font-size: 28px;
                    }}
                    QToolButton:hover {{
                        background-color: #1976D2;
                    }}
                    QToolButton:disabled {{
                        background-color: {colors['border']};
                    }}
                    QToolButton:pressed {{
                        background-color: #0D47A1;
                    }}
                """)
                self.btn_unified_control.setProperty("control_state", "resume")

        except Exception as e:
            print(f"Error updating unified control button: {e}")

    def pause_resume_training(self):
        """Pause or resume the training session"""
        try:
            if not self.training_active:
                return

            if self.is_training_paused:
                # Resume training with updated speed settings
                self.is_training_paused = False

                # Recalculate timer interval with current settings
                wpm = self.trainer_wpm_spinbox.value()
                wpg = self.wpg_spinbox.value()
                interval_ms = int((wpg * 60 * 1000) / wpm)

                # Restart timer with new interval
                self.training_timer.stop()  # Stop first to ensure clean restart
                self.training_timer.start(interval_ms)

                # Update unified control button to show Pause
                self.update_unified_control_button("pause")
                self.progress_info.setText(f"{self.localization.get_text('training_resumed') if self.localization else 'Training resumed'} - {wpm} WPM")
            else:
                # Pause training
                self.is_training_paused = True
                self.training_timer.stop()

                # Update unified control button to show Resume
                self.update_unified_control_button("resume")
                self.progress_info.setText(f"{self.localization.get_text('training_paused') if self.localization else 'Training paused'} - {self.localization.get_text('adjust_speed_if_needed') if self.localization else 'Adjust speed if needed'}")

        except Exception as e:
            print(f"Error pausing/resuming training: {e}")

    def stop_fast_reading_training(self):
        """Stop the training session"""
        try:
            # Save current word index BEFORE resetting (BUG FIX)
            current_word_index_backup = self.current_word_index if hasattr(self, 'current_word_index') else 0

            # Ask to save session if any document is loaded
            if hasattr(self, 'trainer_words') and len(self.trainer_words) > 0:

                # Ask user if they want to save the session
                reply = QMessageBox.question(
                    self,
                    self.localization.get_text("save_session") if self.localization else "Save Session",
                    f"{self.localization.get_text('save_training_progress') if self.localization else 'Save your training progress?'}\n"
                    f"{self.localization.get_text('current_progress') if self.localization else 'Current progress'}: {(current_word_index_backup / len(self.trainer_words)) * 100:.1f}%",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )

                if reply == QMessageBox.Yes:
                    # Save the session with current position
                    self.save_training_session()

                    # Also save to database for persistence
                    print(f"🟡 DB save debug: preparing to save reading position (word_index={current_word_index_backup}, total_words={len(self.trainer_words)})")
                    if self.position_db and hasattr(self, 'trainer_file_path'):
                        try:
                            # Extra diagnostics
                            db_path = getattr(self.position_db, 'db_path', None)
                            print(f"🗄️ DB path: {db_path} | exists={os.path.exists(db_path) if db_path else None} | writable={os.access(db_path, os.W_OK) if db_path and os.path.exists(db_path) else None}")
                            # Build position payload and save using ReadingPositionDB API
                            doc_type = getattr(self, 'current_document_type', None)
                            if not doc_type:
                                ext = os.path.splitext(self.trainer_file_path)[1].lower()
                                if ext == '.pdf':
                                    doc_type = 'pdf'
                                elif ext in ('.doc', '.docx'):
                                    doc_type = 'docx'
                                else:
                                    doc_type = 'txt'
                            last_page = getattr(self, 'current_reading_page', None)
                            pos_data = {
                                'document_type': doc_type,
                                'last_page': last_page if doc_type == 'pdf' else None,
                                'last_word_index': current_word_index_backup,
                                'training_mode': self.get_current_training_mode() if hasattr(self, 'get_current_training_mode') else 'standard',
                                'wpm_setting': self.trainer_wpm_spinbox.value(),
                                'words_per_glance': self.wpg_spinbox.value(),
                                'total_words_read': current_word_index_backup,
                            }
                            print(f"📝 DB payload: {pos_data}")
                            # Schema/connection sanity check
                            if hasattr(self, '_debug_log_db_status_and_schema'):
                                self._debug_log_db_status_and_schema()
                            ok = self.position_db.save_reading_position(self.trainer_file_path, pos_data)
                            print(f"✅ Position saved to database: word {current_word_index_backup}/{len(self.trainer_words)} (ok={ok})")
                        except Exception as db_error:
                            print(f"⚠️ Failed to save position to database: {db_error}")
                    else:
                        print(f"🟠 DB save skipped: position_db={bool(self.position_db)}, trainer_file_path present={hasattr(self, 'trainer_file_path')}")

            # Stop timer
            self.training_timer.stop()

            # Reset training state AFTER saving
            self.training_active = False
            self.is_training_paused = False
            self.current_word_index = 0

            # Update UI - reset unified button to Start state, disable stop
            self.update_unified_control_button("start")
            self.btn_unified_control.setEnabled(True)
            self.btn_stop_training.setEnabled(False)

            # Re-enable settings
            self.wpg_spinbox.setEnabled(True)
            self.trainer_wpm_spinbox.setEnabled(True)
            self.font_size_spinbox.setEnabled(True)

            # Reset display
            self.word_display.setText(self.localization.get_text("training_stopped") if self.localization else "Training stopped")

            # Reset progress
            self.update_training_progress(value=0, visible=False)
            self.progress_info.setText(self.localization.get_text("ready_to_start") if self.localization else "Ready to start training")

        except Exception as e:
            print(f"Error stopping training: {e}")

    def _auto_save_position(self):
        """
        Auto-save current reading position to database.
        Called periodically during training (every N words).

        This is a lightweight, non-blocking operation that saves:
        - Current word index
        - Training mode
        - WPM and WPG settings
        - Total words read
        """
        try:
            # Only save if database is available and document is loaded
            if self.position_db is None or self.current_document_path is None:
                return

            # Prepare position data
            position_data = {
                'document_type': self.current_document_type or 'unknown',
                'last_word_index': self.current_word_index,
                'training_mode': self.get_current_training_mode(),
                'wpm_setting': self.trainer_wpm_spinbox.value(),
                'words_per_glance': self.wpg_spinbox.value(),
                'total_words_read': self.total_words_read_session
            }

            # Add page number for PDFs
            if self.current_document_type == 'pdf' and hasattr(self, 'current_reading_page'):
                position_data['last_page'] = self.current_reading_page

            # Save to database (fast operation, typically <5ms)
            self.position_db.update_reading_position(
                self.current_document_path,
                self.current_word_index,
                self.total_words_read_session
            )

        except Exception as e:
            # Silent failure - don't interrupt training for database errors
            print(f"Auto-save error (non-critical): {e}")

    def complete_training(self):
        """Complete the training session"""
        try:
            # Stop timer
            self.training_timer.stop()

            # Calculate training statistics
            total_words = len(self.trainer_words)
            wpm = self.trainer_wpm_spinbox.value()
            wpg = self.wpg_spinbox.value()

            # Estimate training time
            estimated_time_minutes = total_words / wpm

            # Update display
            self.word_display.setText("🎉 " + (self.localization.get_text("training_complete") if self.localization else "Training Complete!"))

            # Show completion message
            QMessageBox.information(
                self,
                self.localization.get_text("training_complete") if self.localization else "Training Complete",
                f"{self.localization.get_text('congratulations') if self.localization else 'Congratulations! You have completed the training session.'}\n\n"
                f"📊 {self.localization.get_text('statistics') if self.localization else 'Statistics'}:\n"
                f"📝 {total_words} {self.localization.get_text('words_read') if self.localization else 'words read'}\n"
                f"⚡ {wpm} WPM {self.localization.get_text('target_speed') if self.localization else 'target speed'}\n"
                f"👁️ {wpg} {self.localization.get_text('words_per_glance') if self.localization else 'words per glance'}\n"
                f"⏱️ ~{estimated_time_minutes:.1f} {self.localization.get_text('minutes') if self.localization else 'minutes'}"
            )

            # Reset for next session
            self.stop_fast_reading_training()

        except Exception as e:
            print(f"Error completing training: {e}")
            self.stop_fast_reading_training()

    def update_effective_speed(self):
        """Update the effective speed calculation display"""
        try:
            wpg = self.wpg_spinbox.value()
            wpm = self.trainer_wpm_spinbox.value()

            # Calculate interval time per word group
            interval_seconds = (wpg * 60) / wpm

            # Calculate effective reading metrics
            words_per_second = wpm / 60
            groups_per_minute = wpm / wpg

            # Create informative display text
            if self.localization:
                text = f"📊 {self.localization.get_text('effective_reading_speed')}: {wpm} WPM\n"
                text += f"⚡ {wpg} {self.localization.get_text('words')} {self.localization.get_text('every')} {interval_seconds:.1f} {self.localization.get_text('seconds')}\n"
                text += f"👁️ {groups_per_minute:.0f} {self.localization.get_text('word_groups_per_minute')}"
            else:
                text = f"📊 Effective Reading Speed: {wpm} WPM\n"
                text += f"⚡ {wpg} words every {interval_seconds:.1f} seconds\n"
                text += f"👁️ {groups_per_minute:.0f} word groups per minute"

            if hasattr(self, 'effective_speed_label'):
                self.effective_speed_label.setText(text)

            # Update current speed display if it exists
            if hasattr(self, 'current_speed_display'):
                current_speed_text = f"⚡ {self.localization.get_text('current_speed') if self.localization else 'Current Speed'}: {wpm} WPM"
                self.current_speed_display.setText(current_speed_text)
        except Exception as e:
            print(f"Error updating effective speed: {e}")


    def _debug_log_db_status_and_schema(self):
        """Print diagnostic info about the DB file and schema."""
        try:
            if not getattr(self, 'position_db', None):
                print("DB DEBUG: position_db is None")
                return
            db_path = getattr(self.position_db, 'db_path', None)
            print(f"DB DEBUG: db_path={db_path}")
            if db_path:
                exists = os.path.exists(db_path)
                writable = os.access(db_path, os.W_OK) if exists else False
                print(f"DB DEBUG: exists={exists} writable={writable}")
                try:
                    conn = sqlite3.connect(db_path)
                    cur = conn.cursor()
                    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='reading_positions'")
                    present = bool(cur.fetchone())
                    print(f"DB DEBUG: table reading_positions present={present}")
                    cur.execute("PRAGMA table_info(reading_positions)")
                    cols = cur.fetchall()
                    print(f"DB DEBUG: columns={[(c[1], c[2]) for c in cols]}")
                    conn.close()
                except Exception as e:
                    print(f"DB DEBUG: error introspecting schema: {e}")
        except Exception as e:
            print(f"DB DEBUG: schema check error: {e}")


    def save_training_session(self):
        """
        Save current training session for later resume.
        Stores: file path, word index, page number (for PDFs), training mode, WPM, WPG settings.
        """
        try:
            print("\n" + "=" * 80)
            print("🔵 DEBUG: save_training_session() CALLED")
            print("=" * 80)

            # Check if recent_books_manager exists
            if not hasattr(self, 'recent_books_manager'):
                print("❌ ERROR: self.recent_books_manager attribute does not exist!")
                print("=" * 80)
                return False

            if self.recent_books_manager is None:
                print("❌ ERROR: self.recent_books_manager is None!")
                print("=" * 80)
                return False

            print(f"✅ recent_books_manager exists: {type(self.recent_books_manager)}")
            print(f"✅ Database path: {self.recent_books_manager.db_path}")

            if not hasattr(self, 'trainer_file_path') or not self.trainer_file_path:
                print("❌ ERROR: No trainer_file_path found")
                print("=" * 80)
                return False

            if not hasattr(self, 'trainer_words') or not self.trainer_words:
                print("❌ ERROR: No trainer_words found")
                print("=" * 80)
                return False

            print(f"✅ File path: {self.trainer_file_path}")
            print(f"✅ Total words: {len(self.trainer_words)}")
            print(f"✅ Current word index: {self.current_word_index}")

            # Determine document type
            file_ext = os.path.splitext(self.trainer_file_path)[1].lower()
            if file_ext == '.pdf':
                doc_type = 'pdf'
            elif file_ext == '.txt':
                doc_type = 'txt'
            elif file_ext in ['.docx', '.doc']:
                doc_type = 'docx'
            else:
                doc_type = 'unknown'

            print(f"✅ Document type: {doc_type}")

            # Prepare session data
            session_data = {
                'file_path': self.trainer_file_path,
                'file_name': os.path.basename(self.trainer_file_path),
                'current_word_index': self.current_word_index,
                'total_words': len(self.trainer_words),
                'progress_percentage': (self.current_word_index / len(self.trainer_words)) * 100,
                'wpm_setting': self.trainer_wpm_spinbox.value(),
                'wpg_setting': self.wpg_spinbox.value(),
                'font_family': self.font_family_combo.currentText(),
                'font_size': self.font_size_spinbox.value(),
                'session_time': getattr(self, 'session_start_time', datetime.now()).isoformat(),
                'last_updated': datetime.now().isoformat()
            }

            print(f"\n📝 Session data prepared:")
            print(f"   File: {session_data['file_name']}")
            print(f"   Word index: {session_data['current_word_index']}/{session_data['total_words']}")
            print(f"   Progress: {session_data['progress_percentage']:.1f}%")
            print(f"   WPM: {session_data['wpm_setting']}, WPG: {session_data['wpg_setting']}")
            print(f"   Font: {session_data['font_family']} {session_data['font_size']}pt")

            print(f"\n🔄 Calling recent_books_manager.save_training_session()...")
            success = self.recent_books_manager.save_training_session(session_data)

            if success:
                print(f"\n✅ SUCCESS: Training session saved to database!")
                print(f"   Database: {self.recent_books_manager.db_path}")
                print("=" * 80 + "\n")
                return True
            else:
                print(f"\n❌ FAILED: save_training_session() returned False")
                print("=" * 80 + "\n")
                return False

        except Exception as e:
            print(f"\n❌ EXCEPTION in save_training_session(): {e}")
            import traceback
            traceback.print_exc()
            print("=" * 80 + "\n")
            return False

    def load_saved_sessions(self):
        """Load all saved training sessions"""
        try:
            print("DEBUG: load_saved_sessions called")
            if self.recent_books_manager:
                sessions = self.recent_books_manager.get_training_sessions()
                print(f"DEBUG: Found {len(sessions)} sessions")
                for i, session in enumerate(sessions):
                    print(f"DEBUG: Session {i}: {session['file_name']} - {session['progress_percentage']:.1f}%")
                return sessions
            else:
                print("DEBUG: No recent_books_manager found")
            return []
        except Exception as e:
            print(f"Error loading saved sessions: {e}")
            return []

    def resume_training_session(self, session_data):
        """Resume a saved training session"""
        try:
            # Load the file
            file_path = session_data['file_path']
            if not os.path.exists(file_path):
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    f"{self.localization.get_text('file_not_found') if self.localization else 'File not found'}: {file_path}"
                )
                return False

            # Determine file type and load
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext == '.txt':
                self.load_trainer_text_file(file_path)
            elif file_ext == '.docx':
                self.load_trainer_docx_file(file_path)
            elif file_ext == '.pdf':
                self.load_trainer_pdf_file(file_path)
            else:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning") if self.localization else "Warning",
                    self.localization.get_text("unsupported_file_type") if self.localization else "Unsupported file type"
                )
                return False

            # Restore settings
            self.trainer_wpm_spinbox.setValue(session_data.get('wpm_setting', 300))
            self.wpg_spinbox.setValue(session_data.get('wpg_setting', 3))
            self.font_family_combo.setCurrentText(session_data.get('font_family', 'IBM Plex Sans Arabic'))
            self.font_size_spinbox.setValue(session_data.get('font_size', 32))

            # Restore progress
            self.current_word_index = session_data.get('current_word_index', 0)

            # Update progress display
            if hasattr(self, 'trainer_words') and len(self.trainer_words) > 0:
                self.update_training_progress(value=self.current_word_index, maximum=len(self.trainer_words), visible=True)
                progress_percent = (self.current_word_index / len(self.trainer_words)) * 100
                self.progress_info.setText(f"{self.localization.get_text('session_resumed') if self.localization else 'Session resumed'}: {progress_percent:.1f}%")

            # Update effective speed display
            self.update_effective_speed()

            QMessageBox.information(
                self,
                self.localization.get_text("success") if self.localization else "Success",
                f"{self.localization.get_text('session_resumed_successfully') if self.localization else 'Session resumed successfully'}\n"
                f"{self.localization.get_text('progress') if self.localization else 'Progress'}: {session_data.get('progress_percentage', 0):.1f}%"
            )

            return True

        except Exception as e:
            print(f"Error resuming training session: {e}")
            QMessageBox.critical(
                self,
                self.localization.get_text("error") if self.localization else "Error",
                f"{self.localization.get_text('error_resuming_session') if self.localization else 'Error resuming session'}: {str(e)}"
            )
            return False

    def show_saved_sessions_dialog(self):
        """Show dialog with saved training sessions"""
        try:
            sessions = self.load_saved_sessions()
            if not sessions:
                QMessageBox.information(
                    self,
                    self.localization.get_text("info") if self.localization else "Information",
                    self.localization.get_text("no_saved_sessions") if self.localization else "No saved training sessions found"
                )
                return

            # Create dialog
            dialog = QDialog(self)
            dialog.setWindowTitle(self.localization.get_text("saved_training_sessions") if self.localization else "Saved Training Sessions")
            dialog.setMinimumSize(600, 400)

            layout = QVBoxLayout(dialog)

            # Instructions
            instructions = QLabel(self.localization.get_text("select_session_to_resume") if self.localization else "Select a session to resume:")
            instructions.setStyleSheet("font-weight: bold; margin-bottom: 10px;")
            layout.addWidget(instructions)

            # Sessions list
            sessions_list = QListWidget()
            sessions_list.setStyleSheet("""
                QListWidget {
                    border: 2px solid #ddd;
                    border-radius: 8px;
                    padding: 5px;
                    background: white;
                }
                QListWidget::item {
                    padding: 10px;
                    border-bottom: 1px solid #eee;
                    margin: 2px;
                    border-radius: 4px;
                }
                QListWidget::item:hover {
                    background: #f0f8ff;
                }
                QListWidget::item:selected {
                    background: #e3f2fd;
                    color: #1976d2;
                }
            """)

            for session in sessions:
                item_text = f"📄 {session['file_name']}\n"
                item_text += f"📊 {session['progress_percentage']:.1f}% {self.localization.get_text('complete') if self.localization else 'complete'}\n"
                item_text += f"⚡ {session['wpm_setting']} WPM, {session['wpg_setting']} WPG\n"
                item_text += f"📅 {session['last_updated'][:16].replace('T', ' ')}"

                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, session)
                sessions_list.addItem(item)

            layout.addWidget(sessions_list)

            # Buttons
            buttons_layout = QHBoxLayout()

            resume_btn = QPushButton("▶️ " + (self.localization.get_text("resume_session") if self.localization else "Resume Session"))
            resume_btn.setStyleSheet("""
                QPushButton {
                    background: #4CAF50;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background: #45a049;
                }
            """)

            delete_btn = QPushButton("🗑️ " + (self.localization.get_text("delete_session") if self.localization else "Delete Session"))
            delete_btn.setStyleSheet("""
                QPushButton {
                    background: #f44336;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background: #da190b;
                }
            """)

            cancel_btn = QPushButton(self.localization.get_text("cancel") if self.localization else "Cancel")
            cancel_btn.setStyleSheet("""
                QPushButton {
                    background: #757575;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 6px;
                }
                QPushButton:hover {
                    background: #616161;
                }
            """)

            def resume_selected():
                current_item = sessions_list.currentItem()
                if current_item:
                    session_data = current_item.data(Qt.UserRole)
                    if self.resume_training_session(session_data):
                        dialog.accept()
                else:
                    QMessageBox.warning(dialog,
                        self.localization.get_text("warning") if self.localization else "Warning",
                        self.localization.get_text("select_session_first") if self.localization else "Please select a session first")

            def delete_selected():
                current_item = sessions_list.currentItem()
                if current_item:
                    session_data = current_item.data(Qt.UserRole)
                    reply = QMessageBox.question(dialog,
                        self.localization.get_text("confirm_delete") if self.localization else "Confirm Delete",
                        f"{self.localization.get_text('delete_session_confirm') if self.localization else 'Delete this training session?'}\n{session_data['file_name']}")

                    if reply == QMessageBox.Yes:
                        if self.recent_books_manager:
                            self.recent_books_manager.delete_training_session(session_data['file_path'], session_data['current_word_index'])
                        sessions_list.takeItem(sessions_list.row(current_item))
                        if sessions_list.count() == 0:
                            dialog.accept()
                else:
                    QMessageBox.warning(dialog,
                        self.localization.get_text("warning") if self.localization else "Warning",
                        self.localization.get_text("select_session_first") if self.localization else "Please select a session first")

            resume_btn.clicked.connect(resume_selected)
            delete_btn.clicked.connect(delete_selected)
            cancel_btn.clicked.connect(dialog.reject)

            buttons_layout.addWidget(resume_btn)
            buttons_layout.addWidget(delete_btn)
            buttons_layout.addStretch()
            buttons_layout.addWidget(cancel_btn)

            layout.addLayout(buttons_layout)

            dialog.exec()

        except Exception as e:
            print(f"Error showing saved sessions dialog: {e}")

    def create_training_modes_selector(self, layout):
        """Different training modes"""

        modes_group = QGroupBox("🎯 " + (self.localization.get_text("training_mode") if self.localization else "Training Mode"))
        colors = self.get_theme_colors()
        modes_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                font-size: 16px;
                color: {colors['accent_purple']};
                border: 2px solid {colors['accent_purple']};
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: {colors['bg_primary']};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: {colors['accent_purple']};
            }}
        """)
        modes_layout = QVBoxLayout(modes_group)

        # Define only 2 training modes (Standard and RSVP) - most effective and popular
        from PySide6.QtWidgets import QRadioButton, QButtonGroup

        self.training_modes = {
            'standard': {
                'name': '📖 ' + (self.localization.get_text("standard_mode") if self.localization else "Standard (Words Per Glance)"),
                'name_ar': '📖 الوضع القياسي (كلمات في النظرة)',
                'description': self.localization.get_text("standard_mode_desc") if self.localization else "Train to read multiple words at once",
                'description_ar': 'تدرب على قراءة عدة كلمات معاً',
                'details': self.localization.get_text("standard_mode_details") if self.localization else "Perfect for beginners. Shows 1-7 words with center focus highlighting.",
                'details_ar': 'مثالي للمبتدئين. يعرض 1-7 كلمات مع تبريز الكلمة المركزية.',
                'icon': '📖'
            },
            'rsvp': {
                'name': '⚡ ' + (self.localization.get_text("rsvp_mode") if self.localization else "RSVP (Rapid Serial Visual Presentation)"),
                'name_ar': '⚡ العرض السريع المتسلسل',
                'description': self.localization.get_text("rsvp_mode_desc") if self.localization else "Single words flashing rapidly in center",
                'description_ar': 'كلمات فردية تظهر بسرعة في المركز',
                'details': self.localization.get_text("rsvp_mode_details") if self.localization else "Most effective for speed. Eliminates eye movement and regression.",
                'details_ar': 'الأكثر فعالية للسرعة. يلغي حركة العين والرجوع للخلف.',
                'icon': '⚡'
            }
        }

        # Create button group for radio buttons
        self.training_mode_button_group = QButtonGroup(self)

        # Horizontal layout for radio buttons and info button (all in one row)
        modes_row_layout = QHBoxLayout()
        modes_row_layout.setSpacing(15)

        # Create radio buttons for each mode
        for mode_key, mode_data in self.training_modes.items():
            # Container for radio button and description
            mode_container = QVBoxLayout()
            mode_container.setSpacing(3)

            # Create radio button with Arabic support
            # Check if current language is Arabic
            is_arabic = self.localization and hasattr(self.localization, 'current_language') and self.localization.current_language == 'ar'
            mode_name = mode_data.get('name_ar', mode_data['name']) if is_arabic else mode_data['name']
            mode_desc = mode_data.get('description_ar', mode_data['description']) if is_arabic else mode_data['description']

            radio_btn = QRadioButton(mode_name)
            radio_btn.setProperty('mode_key', mode_key)
            radio_btn.setStyleSheet(f"""
                QRadioButton {{
                    font-size: 14px;
                    font-weight: bold;
                    color: {colors['text_primary']};
                    padding: 5px;
                }}
                QRadioButton::indicator {{
                    width: 18px;
                    height: 18px;
                }}
                QRadioButton::indicator:unchecked {{
                    border: 2px solid {colors['border']};
                    border-radius: 9px;
                    background-color: {colors['bg_primary']};
                }}
                QRadioButton::indicator:checked {{
                    border: 2px solid {colors['accent_purple']};
                    border-radius: 9px;
                    background-color: {colors['accent_purple']};
                }}
                QRadioButton:hover {{
                    color: {colors['accent_purple']};
                }}
            """)

            # Add to button group
            self.training_mode_button_group.addButton(radio_btn)

            # Create description label
            desc_label = QLabel(mode_desc)
            desc_label.setWordWrap(True)
            desc_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 11px;
                    color: {colors['text_secondary']};
                    padding-left: 25px;
                }}
            """)

            # Add to container
            mode_container.addWidget(radio_btn)
            mode_container.addWidget(desc_label)

            # Add container to horizontal row
            modes_row_layout.addLayout(mode_container)

        # Add stretch to push info button to the right
        modes_row_layout.addStretch()

        # Icon-only info button for mode explanations
        self.btn_mode_help = QToolButton()
        self.btn_mode_help.setText("ℹ️")
        self.btn_mode_help.setToolTip(self.localization.get_text("learn_more_about_modes") if self.localization else "Learn more about training modes")
        self.btn_mode_help.clicked.connect(self.show_mode_help_dialog)
        self.btn_mode_help.setFixedSize(40, 40)
        self.btn_mode_help.setStyleSheet(f"""
            QToolButton {{
                background-color: {colors['accent_purple']};
                color: white;
                border: none;
                border-radius: 6px;
                font-size: 20px;
            }}
            QToolButton:hover {{
                background-color: {colors['accent_blue']};
            }}
            QToolButton:pressed {{
                background-color: {colors['accent_green']};
            }}
        """)
        modes_row_layout.addWidget(self.btn_mode_help)

        modes_layout.addLayout(modes_row_layout)

        # Set default selection (Standard mode)
        first_button = self.training_mode_button_group.buttons()[0]
        first_button.setChecked(True)

        # Connect signal
        self.training_mode_button_group.buttonClicked.connect(self.on_training_mode_changed)

        layout.addWidget(modes_group)

        # Initialize with first mode
        self.on_training_mode_changed()



    def on_training_mode_changed(self):
        """
        Update display and settings based on selected training mode.

        Only 2 modes now:
        - Standard: 1-7 words, user choice (default 3)
        - RSVP: Locked to 1 word for rapid flash display
        """
        # Get selected mode from radio button group
        selected_button = self.training_mode_button_group.checkedButton()
        if not selected_button:
            return

        mode_key = selected_button.property('mode_key')
        if not mode_key or mode_key not in self.training_modes:
            return

        mode = self.training_modes[mode_key]

        # Adjust UI for mode
        if mode_key == 'rsvp':
            # RSVP Mode (Rapid Serial Visual Presentation)
            # Technical: Single word flashing in center to eliminate eye movement
            # WPG locked to 1, high speed recommended (400+ WPM)
            self.wpg_spinbox.setValue(1)
            self.wpg_spinbox.setEnabled(False)
            self.wpg_spinbox.setToolTip(self.localization.get_text("rsvp_wpg_locked") if self.localization else "RSVP mode displays one word at a time")

        else:  # standard
            # Standard Mode (Words Per Glance)
            # Technical: Basic training to expand peripheral vision
            # User can gradually increase WPG from 1 to 7
            self.wpg_spinbox.setEnabled(True)
            self.wpg_spinbox.setRange(1, 7)
            self.wpg_spinbox.setToolTip(self.localization.get_text("standard_wpg") if self.localization else "Number of words to display at once")

    def get_current_training_mode(self):
        """Get the currently selected training mode key"""
        selected_button = self.training_mode_button_group.checkedButton()
        if selected_button:
            return selected_button.property('mode_key')
        return 'standard'  # Default fallback

    def show_mode_help_dialog(self):
        """Show dialog with training mode explanations and comparisons"""
        dialog = QDialog(self)
        dialog.setWindowTitle(self.localization.get_text("training_modes_help") if self.localization else "Training Modes Help")
        dialog.setMinimumSize(700, 600)

        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Title
        title = QLabel("🎯 " + (self.localization.get_text("training_modes_guide") if self.localization else "Training Modes Guide"))
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # Scroll area for content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; }")

        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(15)

        colors = self.get_theme_colors()

        # Mode explanations
        for mode_key, mode_data in self.training_modes.items():
            mode_frame = QFrame()
            mode_frame.setStyleSheet(f"""
                QFrame {{
                    background-color: {colors['bg_secondary']};
                    border: 2px solid {colors['accent_blue']};
                    border-radius: 10px;
                    padding: 15px;
                }}
            """)
            mode_layout = QVBoxLayout(mode_frame)

            # Mode name
            mode_name = QLabel(mode_data['name'])
            mode_name.setFont(QFont("Arial", 14, QFont.Bold))
            mode_name.setStyleSheet(f"color: {colors['accent_blue']};")
            mode_layout.addWidget(mode_name)

            # Mode description
            mode_desc = QLabel(mode_data['description'])
            mode_desc.setWordWrap(True)
            mode_desc.setStyleSheet(f"color: {colors['text_primary']}; font-size: 13px; padding: 5px 0;")
            mode_layout.addWidget(mode_desc)

            # Additional details based on mode
            details_text = self.get_mode_details(mode_key)
            if details_text:
                details = QLabel(details_text)
                details.setWordWrap(True)
                details.setStyleSheet(f"color: {colors['text_secondary']}; font-size: 12px; font-style: italic; padding: 5px 0;")
                mode_layout.addWidget(details)

            content_layout.addWidget(mode_frame)

        # Mode comparison section
        comparison_frame = QFrame()
        comparison_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {colors['bg_tertiary']};
                border: 2px solid {colors['accent_orange']};
                border-radius: 10px;
                padding: 15px;
            }}
        """)
        comparison_layout = QVBoxLayout(comparison_frame)

        comparison_title = QLabel("📊 " + (self.localization.get_text("mode_comparison") if self.localization else "Mode Comparison"))
        comparison_title.setFont(QFont("Arial", 14, QFont.Bold))
        comparison_title.setStyleSheet(f"color: {colors['accent_orange']};")
        comparison_layout.addWidget(comparison_title)

        comparison_content = QLabel()
        comparison_content.setWordWrap(True)
        comparison_content.setTextFormat(Qt.RichText)
        comparison_text = f"""
        <div style='font-size: 13px; line-height: 1.8; color: {colors['text_primary']};'>
        <h3 style='color: {colors['accent_purple']};'>📊 {self.localization.get_text("mode_comparison") if self.localization else "Mode Comparison"}</h3>
        <table width='100%' cellpadding='8' style='border-collapse: collapse; margin: 15px 0;'>
        <tr style='background-color: {colors['bg_secondary']}; font-weight: bold;'>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>{self.localization.get_text("criterion") if self.localization else "Criterion"}</td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>📖 {self.localization.get_text("standard_mode") if self.localization else "Standard"}</td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>⚡ {self.localization.get_text("rsvp_mode") if self.localization else "RSVP"}</td>
        </tr>
        <tr>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'><b>{self.localization.get_text("ease_of_use") if self.localization else "Ease of Use"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>⭐⭐⭐⭐⭐ {self.localization.get_text("very_easy") if self.localization else "Very Easy"}</td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>⭐⭐⭐⭐ {self.localization.get_text("easy") if self.localization else "Easy"}</td>
        </tr>
        <tr style='background-color: {colors['bg_secondary']};'>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'><b>{self.localization.get_text("effectiveness") if self.localization else "Effectiveness"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>⭐⭐⭐⭐ {self.localization.get_text("excellent") if self.localization else "Excellent"}</td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>⭐⭐⭐⭐⭐ {self.localization.get_text("very_excellent") if self.localization else "Very Excellent"}</td>
        </tr>
        <tr>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'><b>{self.localization.get_text("for_beginners") if self.localization else "For Beginners"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px; color: {colors['accent_green']};'><b>✅ {self.localization.get_text("ideal") if self.localization else "Ideal"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px; color: {colors['accent_green']};'><b>✅ {self.localization.get_text("very_good") if self.localization else "Very Good"}</b></td>
        </tr>
        <tr style='background-color: {colors['bg_secondary']};'>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'><b>{self.localization.get_text("for_advanced") if self.localization else "For Advanced"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px; color: {colors['accent_green']};'><b>✅ {self.localization.get_text("excellent") if self.localization else "Excellent"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px; color: {colors['accent_green']};'><b>✅ {self.localization.get_text("excellent") if self.localization else "Excellent"}</b></td>
        </tr>
        <tr>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'><b>{self.localization.get_text("max_speed") if self.localization else "Max Speed"}</b></td>
            <td style='border: 1px solid {colors['border']}; padding: 10px;'>600 WPM</td>
            <td style='border: 1px solid {colors['border']}; padding: 10px; color: {colors['accent_purple']};'><b>1000+ WPM</b></td>
        </tr>
        </table>
        <br>
        <div style='background-color: {colors['bg_secondary']}; padding: 15px; border-radius: 8px; border-left: 4px solid {colors['accent_orange']};'>
        <p style='color: {colors['accent_orange']}; font-weight: bold; margin: 0 0 10px 0;'>
        💡 {self.localization.get_text("recommended_training_plan") if self.localization else "Recommended Training Plan"}
        </p>
        <p style='margin: 5px 0;'><b>{self.localization.get_text("week_1_2") if self.localization else "Week 1-2"}:</b> {self.localization.get_text("start_standard_mode") if self.localization else "Start with Standard mode (2-3 words, 200-250 WPM)"}</p>
        <p style='margin: 5px 0;'><b>{self.localization.get_text("week_3_4") if self.localization else "Week 3-4"}:</b> {self.localization.get_text("mix_both_modes") if self.localization else "Mix both modes (10 min Standard + 5 min RSVP)"}</p>
        <p style='margin: 5px 0;'><b>{self.localization.get_text("week_5_plus") if self.localization else "Week 5+"}:</b> {self.localization.get_text("advanced_training") if self.localization else "RSVP for speed training (400-600 WPM)"}</p>
        </div>
        </div>
        """
        comparison_content.setText(comparison_text)
        comparison_layout.addWidget(comparison_content)

        content_layout.addWidget(comparison_frame)

        scroll.setWidget(content_widget)
        layout.addWidget(scroll)

        # Close button
        close_btn = QPushButton(self.localization.get_text("close") if self.localization else "Close")
        close_btn.clicked.connect(dialog.accept)
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {colors['accent_blue']};
                color: white;
                border: none;
                padding: 10px 30px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background-color: {colors['accent_green']};
            }}
        """)
        layout.addWidget(close_btn, alignment=Qt.AlignCenter)

        dialog.exec()

    def get_mode_details(self, mode_key):
        """Get additional details for the 2 training modes with beginner-friendly explanations"""
        details = {
            'standard': self.localization.get_text("standard_mode_details") if self.localization else
                "📚 <b>Perfect for beginners!</b><br>"
                "• Start with 1-2 words per glance, gradually increase to 5-7<br>"
                "• Trains your peripheral vision to see more words at once<br>"
                "• <b>When to use:</b> Daily practice, building foundation skills<br>"
                "• <b>Recommended speed:</b> 200-400 WPM<br>"
                "• <b>Tip:</b> Focus on the middle word, let your peripheral vision catch the others<br>"
                "• <b>Best for:</b> 90% of users, natural reading improvement",

            'rsvp': self.localization.get_text("rsvp_mode_details") if self.localization else
                "⚡ <b>Advanced speed training!</b><br>"
                "• Words flash one at a time in the center of the screen<br>"
                "• Eliminates eye movement completely - your eyes stay fixed<br>"
                "• <b>When to use:</b> Breaking speed plateaus, building raw speed<br>"
                "• <b>Recommended speed:</b> 400-1000 WPM<br>"
                "• <b>Tip:</b> Don't try to pronounce words mentally, just recognize them visually<br>"
                "• <b>Example:</b> Like watching a slideshow - each word appears briefly then disappears<br>"
                "• <b>Best for:</b> Intensive training, eliminating bad reading habits"
        }
        return details.get(mode_key, "")

    def update_training_progress(self, value=None, maximum=None, visible=None):
        """Update both the section progress bar and bottom progress bar"""
        if hasattr(self, 'training_progress'):
            if maximum is not None:
                self.training_progress.setMaximum(maximum)
            if value is not None:
                self.training_progress.setValue(value)

        if hasattr(self, 'bottom_progress_bar'):
            if maximum is not None:
                self.bottom_progress_bar.setMaximum(maximum)
            if value is not None:
                self.bottom_progress_bar.setValue(value)
            if visible is not None:
                self.bottom_progress_bar.setVisible(visible)

    def init_ui(self):
        """Initialize the reading speed measurement UI with tabbed interface"""
        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(15, 15, 15, 15)

        # Add development banner at the top
        dev_banner = QLabel("⚠️ قيد التطوير - Under Development ⚠️")
        dev_banner.setAlignment(Qt.AlignCenter)
        dev_banner.setStyleSheet("""
            QLabel {
                background-color: #FFA726;
                color: #000000;
                font-size: 14px;
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
                margin: 5px;
            }
        """)
        dev_banner.setMaximumHeight(40)
        layout.addWidget(dev_banner)

        # Header with title and help button
        header_layout = QHBoxLayout()

        # Title
        title = QLabel(self.localization.get_text("reading_speed_meter"))
        title.setFont(QFont("Arial", 18, QFont.Bold))
        title.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        header_layout.addWidget(title)

        # Add stretch to push help button to the right
        header_layout.addStretch()

        # Help/Guide button
        self.btn_help = QPushButton("❓ " + self.localization.get_text("help"))
        self.btn_help.clicked.connect(self.show_reading_speed_guide)
        self.btn_help.setMaximumWidth(100)
        self.btn_help.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        header_layout.addWidget(self.btn_help)

        layout.addLayout(header_layout)

        desc = QLabel(self.localization.get_text("reading_speed_meter_desc"))
        desc.setWordWrap(True)
        desc.setAlignment(Qt.AlignCenter)
        desc.setStyleSheet("color: #666; margin-bottom: 15px;")
        layout.addWidget(desc)

        # Create tabbed interface with theme-aware styling
        self.tab_widget = QTabWidget()
        colors = self.get_theme_colors()
        tab_style = f"""
            QTabWidget::pane {{
                border: 2px solid {colors['border']};
                border-radius: 8px;
                background-color: {colors['bg_secondary']};
            }}
            QTabBar::tab {{
                background-color: {colors['bg_tertiary']};
                border: 2px solid {colors['border']};
                border-bottom-color: {colors['border']};
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
                min-width: 120px;
                padding: 8px 16px;
                margin-right: 2px;
                font-weight: bold;
                font-size: 14px;
                color: {colors['text_primary']};
            }}
            QTabBar::tab:selected {{
                background-color: {colors['accent_blue']};
                color: white;
                border-bottom-color: {colors['accent_blue']};
            }}
            QTabBar::tab:hover {{
                background-color: {colors['bg_secondary']};
                color: {colors['text_primary']};
            }}
        """
        self.tab_widget.setStyleSheet(tab_style)

        # Create tabs in order: Fast Reading Training (1st), Pages/Minute (2nd), Current Book (3rd), Activity (4th)
        self.create_fast_reading_trainer_tab()   # Fast Reading Training tab - FIRST
        self.create_pages_per_minute_tab()       # Pages Per Minute tab - SECOND
        self.create_current_book_tab()           # Current Book tab - THIRD
        self.create_prepared_books_tab()         # Activity tab - FOURTH

        layout.addWidget(self.tab_widget)

        # Thin progress bar at bottom of page - reduced vertical space
        self.bottom_progress_bar = QProgressBar()
        self.bottom_progress_bar.setMaximumHeight(4)  # Reduced from 8 to 4
        self.bottom_progress_bar.setTextVisible(False)
        self.bottom_progress_bar.setStyleSheet(f"""
            QProgressBar {{
                border: none;
                background-color: {colors['bg_tertiary']};
                border-radius: 2px;
                margin: 0px;
                padding: 0px;
            }}
            QProgressBar::chunk {{
                background-color: {colors['accent_green']};
                border-radius: 2px;
            }}
        """)
        self.bottom_progress_bar.setVisible(False)  # Hidden by default
        self.bottom_progress_bar.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.bottom_progress_bar)

    def create_pages_per_minute_tab(self):
        """Create the Pages Per Minute tab with three accuracy modes.
        Modes:
          1) Simple (pages/minute)
          2) Sample-Based Estimation (avg words/page via AI prompt)
          3) Most Accurate (per-page counts via CSV from AI, persisted)
        """
        tab = QWidget()
        layout = QVBoxLayout(tab)


        # Theme colors for styling
        colors = self.get_theme_colors()

        # Common controls (top): Start Page, Timer, Start/Stop
        top_bar = QHBoxLayout()
        self.ppm_start_page_spin = QSpinBox()
        self.ppm_start_page_spin.setMinimum(1)
        # Fallback maximum to allow >99 even when total pages unknown
        self.ppm_start_page_spin.setMaximum(9999)
        self.ppm_start_page_spin.setToolTip(self.localization.get_text("start_page") if self.localization else "Start page")

        self.ppm_timer_label = QLabel("00:00")
        self.ppm_timer_label.setToolTip(self.localization.get_text("elapsed_time") if self.localization else "Elapsed time")
        # Make timer more readable
        self.ppm_timer_label.setStyleSheet("QLabel { font-size: 20px; font-weight: bold; padding: 2px 6px; }")

        self.ppm_start_stop_btn = QPushButton(self.localization.get_text("start") if self.localization else "Start")
        self.ppm_start_stop_btn.clicked.connect(self._ppm_toggle_timer)
        # Make Start/Stop button more prominent
        self.ppm_start_stop_btn.setMinimumHeight(36)
        self.ppm_start_stop_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {colors['accent_green']};
                color: white;
                border-radius: 6px;
                padding: 6px 14px;
                font-weight: bold;
                font-size: 14px;
            }}
            QPushButton:disabled {{
                background-color: {colors['bg_secondary']};
                color: {colors['text_muted']};
            }}
        """)

        top_bar.addWidget(QLabel(self.localization.get_text("start_page") if self.localization else "Start Page:"))


        top_bar.addWidget(self.ppm_start_page_spin)
        top_bar.addStretch(1)
        top_bar.addWidget(self.ppm_timer_label)
        top_bar.addWidget(self.ppm_start_stop_btn)
        layout.addLayout(top_bar)

        # Result display
        self.ppm_result_label = QLabel("")
        layout.addWidget(self.ppm_result_label)

        # Sub-tab widget for the three modes
        self.ppm_modes_tabs = QTabWidget()
        self.ppm_modes_tabs.setTabPosition(QTabWidget.North)

        # Mode 1: Simple
        simple_w = QWidget()
        simple_l = QVBoxLayout(simple_w)
        simple_l.addWidget(QLabel(self.localization.get_text("ppm_simple_desc") if self.localization else "Enter start page, read, then stop and enter end page"))
        self.ppm_modes_tabs.addTab(simple_w, self.localization.get_text("mode_simple") if self.localization else "Simple")

        # Mode 2: Sample-Based Estimation
        sample_w = QWidget()
        sample_l = QVBoxLayout(sample_w)
        row = QHBoxLayout()
        row.addWidget(QLabel(self.localization.get_text("sample_pages") if self.localization else "Sample pages:"))
        self.ppm_sample_pages_spin = QSpinBox()
        self.ppm_sample_pages_spin.setRange(1, 100)
        self.ppm_sample_pages_spin.setValue(10)
        row.addWidget(self.ppm_sample_pages_spin)
        self.ppm_sample_extract_btn = QPushButton(self.localization.get_text("extract_sample") if self.localization else "Extract Sample")
        self.ppm_sample_extract_btn.clicked.connect(self._ppm_extract_sample_pages)
        # Sample-Based: Document selection row
        sample_doc_row = QHBoxLayout()
        self.ppm_sample_select_btn = QPushButton(self.localization.get_text("select_document") if self.localization else "Select Document")
        self.ppm_sample_select_btn.clicked.connect(lambda: self._ppm_select_document("sample"))
        self.ppm_sample_doc_label = QLabel("—")
        self.ppm_sample_doc_label.setStyleSheet("QLabel { font-style: italic; }")
        sample_doc_row.addWidget(self.ppm_sample_select_btn)
        sample_doc_row.addWidget(self.ppm_sample_doc_label, 1)
        sample_l.addLayout(sample_doc_row)

        self.ppm_sample_extract_btn.setEnabled(False)
        # Moved Most Accurate document row to accurate tab section
        # (removed leftover)
        # moved

        # moved

        row.addWidget(self.ppm_sample_extract_btn)
        row.addStretch(1)
        sample_l.addLayout(row)
        # Refresh prompt when sample size changes
        self.ppm_sample_pages_spin.valueChanged.connect(lambda _: self._ppm_refresh_prompts())

        # Prompt (Arabic) + Copy button
        self.ppm_sample_prompt = QTextEdit()
        self.ppm_sample_prompt.setReadOnly(True)
        self.ppm_sample_prompt.setMaximumHeight(120)
        copy_row = QHBoxLayout()
        copy_row.addWidget(QLabel(self.localization.get_text("ai_prompt") if self.localization else "AI Prompt"))
        btn_copy_sample = QPushButton(self.localization.get_text("copy_prompt") if self.localization else "Copy Prompt")
        btn_copy_sample.clicked.connect(lambda: self._ppm_copy_to_clipboard(self.ppm_sample_prompt.toPlainText()))
        copy_row.addWidget(btn_copy_sample)
        copy_row.addStretch(1)
        sample_l.addLayout(copy_row)
        sample_l.addWidget(self.ppm_sample_prompt)

        # Average words/page input
        avg_row = QHBoxLayout()

        # Initialize document UI and enable/disable states
        self._ppm_update_doc_ui()
        self._ppm_update_enabled_states()

        avg_row.addWidget(QLabel(self.localization.get_text("avg_words_per_page") if self.localization else "Average words/page:"))
        self.ppm_avg_wpp_spin = QSpinBox()
        self.ppm_avg_wpp_spin.setRange(0, 5000)
        self.ppm_avg_wpp_spin.setValue(0)
        avg_row.addWidget(self.ppm_avg_wpp_spin)
        avg_row.addStretch(1)
        sample_l.addLayout(avg_row)

        self.ppm_modes_tabs.addTab(sample_w, self.localization.get_text("mode_sample") if self.localization else "Sample-Based")

        # Mode 3: Most Accurate
        acc_w = QWidget()
        acc_l = QVBoxLayout(acc_w)

        self.ppm_pwc_indicator = QLabel("")
        acc_l.addWidget(self.ppm_pwc_indicator)

        # Prompt (English CSV) + Copy button
        self.ppm_csv_prompt = QTextEdit()
        self.ppm_csv_prompt.setReadOnly(True)
        self.ppm_csv_prompt.setMaximumHeight(140)
        copy_row2 = QHBoxLayout()
        copy_row2.addWidget(QLabel(self.localization.get_text("ai_prompt") if self.localization else "AI Prompt"))
        btn_copy_csv = QPushButton(self.localization.get_text("copy_prompt") if self.localization else "Copy Prompt")
        btn_copy_csv.clicked.connect(lambda: self._ppm_copy_to_clipboard(self.ppm_csv_prompt.toPlainText()))
        # Most Accurate: Document selection row
        acc_doc_row = QHBoxLayout()
        self.ppm_acc_select_btn = QPushButton(self.localization.get_text("select_document") if self.localization else "Select Document")
        self.ppm_acc_select_btn.clicked.connect(lambda: self._ppm_select_document("accurate"))
        self.ppm_acc_open_btn = QPushButton(self.localization.get_text("open_pdf") if self.localization else "Open PDF")
        self.ppm_acc_open_btn.clicked.connect(self._ppm_open_pdf)
        self.ppm_acc_open_btn.setEnabled(False)
        self.ppm_acc_doc_label = QLabel("-")
        self.ppm_acc_doc_label.setStyleSheet("QLabel { font-style: italic; }")
        self.ppm_acc_doc_label.setText("—")
        acc_doc_row.addWidget(self.ppm_acc_select_btn)
        acc_doc_row.addWidget(self.ppm_acc_open_btn)
        acc_doc_row.addWidget(self.ppm_acc_doc_label, 1)
        acc_l.addLayout(acc_doc_row)

        copy_row2.addWidget(btn_copy_csv)
        copy_row2.addStretch(1)
        acc_l.addLayout(copy_row2)
        acc_l.addWidget(self.ppm_csv_prompt)

        # CSV input area
        acc_l.addWidget(QLabel(self.localization.get_text("paste_csv_here") if self.localization else "Paste CSV data here:"))
        self.ppm_csv_input = QTextEdit()
        self.ppm_csv_input.setPlaceholderText(self.localization.get_text("csv_placeholder") if self.localization else "page_number,word_count\n1,250\n2,240\n...")
        self.ppm_csv_input.setMinimumHeight(220)
        self.ppm_csv_input.setStyleSheet("QTextEdit { font-family: 'Consolas', 'Courier New', monospace; }")
        self.ppm_csv_input.textChanged.connect(self._ppm_on_csv_changed)
        acc_l.addWidget(self.ppm_csv_input)
        self.ppm_parse_save_btn = QPushButton(self.localization.get_text("parse_and_save") if self.localization else "Parse & Save")
        self.ppm_parse_save_btn.clicked.connect(self._ppm_parse_csv_and_save)
        self.ppm_parse_save_btn.setEnabled(False)
        acc_l.addWidget(self.ppm_parse_save_btn)

        self.ppm_modes_tabs.addTab(acc_w, self.localization.get_text("mode_accurate") if self.localization else "Most Accurate")

        layout.addWidget(self.ppm_modes_tabs)
        self.tab_widget.addTab(tab, self.localization.get_text("pages_per_minute") if self.localization else "Pages/Minute")

        # Timer state
        if not hasattr(self, "ppm_timer"):
            self.ppm_timer = QTimer(self)
            self.ppm_timer.setInterval(1000)
            self.ppm_timer.timeout.connect(self._ppm_on_tick)
        self.ppm_running = False
        self.ppm_elapsed_seconds = 0

        # Initialize prompts and defaults
        self._ppm_refresh_defaults()
        self._ppm_refresh_prompts()
        self._ppm_autoselect_mode_based_on_data()


    def _ppm_select_document(self, source: str = ""):
        """Let user choose a PDF for the Pages/Minute tab and update shared state."""
        try:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                self.localization.get_text("open_pdf") if self.localization else "Open PDF",
                "",
                self.localization.get_text("pdf_files_filter") if self.localization else "PDF Files (*.pdf)"
            )
            if not file_path:
                return
            if not str(file_path).lower().endswith('.pdf'):
                QMessageBox.information(self, self.localization.get_text("info") if self.localization else "Info", self.localization.get_text("open_pdf_first") if self.localization else "Please open a PDF first.")
                return
            # Count pages
            total = None
            if fitz:
                try:
                    d = fitz.open(file_path)
                    total = d.page_count
                    d.close()
                except Exception:
                    pass
            # Persist to shared attributes so all modes see it
            self.ppm_doc_path = file_path
            self.ppm_total_pages = total
            # Also expose as current_document_path so other helpers can find it
            self.current_document_path = file_path
            # Update UI and defaults
            self._ppm_update_doc_ui()
            self._ppm_refresh_defaults()
            self._ppm_refresh_prompts()
            self._ppm_update_enabled_states()
        except Exception as e:
            print(f"PPM: selecting document failed: {e}")

    def _ppm_update_doc_ui(self):
        """Update document labels and enable/disable controls based on loaded PDF."""
        try:
            path, total = self._ppm_get_doc_path_and_pages()
            name = os.path.basename(path) if path else "—"
            pages_txt = (self.localization.get_text("total_pages") if self.localization else "Total pages:")
            info = f"{name}  ({pages_txt} {total})" if path and total else (name if path else "—")
            if hasattr(self, 'ppm_sample_doc_label'):
                self.ppm_sample_doc_label.setText(info)
            if hasattr(self, 'ppm_acc_doc_label'):
                self.ppm_acc_doc_label.setText(info)
            # Enable/disable buttons
            has_pdf = bool(path and str(path).lower().endswith('.pdf'))
            if hasattr(self, 'ppm_sample_extract_btn'):
                self.ppm_sample_extract_btn.setEnabled(has_pdf)
            if hasattr(self, 'ppm_acc_open_btn'):
                self.ppm_acc_open_btn.setEnabled(has_pdf)
            self._ppm_update_enabled_states()
        except Exception as e:
            print(f"PPM: update_doc_ui failed: {e}")

    def _ppm_update_enabled_states(self):
        try:
            path, _ = self._ppm_get_doc_path_and_pages()
            has_pdf = bool(path and str(path).lower().endswith('.pdf'))
            has_csv = bool(getattr(self, 'ppm_csv_input', None) and self.ppm_csv_input.toPlainText().strip())
            if hasattr(self, 'ppm_parse_save_btn'):
                self.ppm_parse_save_btn.setEnabled(has_pdf and has_csv)
            if hasattr(self, 'ppm_sample_extract_btn'):
                self.ppm_sample_extract_btn.setEnabled(has_pdf)
        except Exception as e:
            print(f"PPM: update_enabled_states failed: {e}")

    def _ppm_on_csv_changed(self):
        self._ppm_update_enabled_states()

    def _ppm_open_pdf(self):
        """Open currently loaded PDF."""
        try:
            path, _ = self._ppm_get_doc_path_and_pages()
            if not path or not os.path.exists(path):
                QMessageBox.information(self, self.localization.get_text("info") if self.localization else "Info", self.localization.get_text("open_pdf_first") if self.localization else "Please open a PDF first.")
                return
            # Prefer Qt API if available
            try:
                from PySide6.QtGui import QDesktopServices
                from PySide6.QtCore import QUrl
                QDesktopServices.openUrl(QUrl.fromLocalFile(path))
                return
            except Exception:
                pass
            # Fallback to OS
            try:
                os.startfile(path)
            except Exception:
                try:
                    import webbrowser
                    webbrowser.open(path)
                except Exception:
                    pass
        except Exception as e:
            print(f"PPM: open_pdf failed: {e}")

    def _ppm_refresh_defaults(self):
        # Set default start page to middle if possible
        path, total = self._ppm_get_doc_path_and_pages()
        if total and total > 0:
            self.ppm_start_page_spin.setMaximum(total)
            self.ppm_start_page_spin.setValue(max(1, total // 2))
        else:
            # Ensure generous maximum even when total pages unknown
            self.ppm_start_page_spin.setMaximum(9999)

    def _ppm_refresh_prompts(self):
        path, total = self._ppm_get_doc_path_and_pages()
        # Sample mode Arabic prompt (short, instruct to compute avg words/page on N middle pages)
        n = self.ppm_sample_pages_spin.value() if hasattr(self, 'ppm_sample_pages_spin') else 10
        start = max(1, (total // 2) - (n // 2)) if total else 1
        end = min(total or (start + n - 1), start + n - 1)
        arabic = (
            f"احسب متوسط عدد الكلمات لكل صفحة في هذا الملف من الصفحات {start} إلى {end}. "
            f"أعطني رقمًا واحدًا فقط بدون أي نص إضافي."
        )
        self.ppm_sample_prompt.setPlainText(arabic)

        # Accurate mode English CSV prompt
        eng = (
            "Return a CSV with two columns: page_number,word_count for ALL pages of this PDF. "
            "No extra text. Headers required: page_number,word_count"
        )
        self.ppm_csv_prompt.setPlainText(eng)

    def _ppm_autoselect_mode_based_on_data(self):
        try:
            if self.position_db is None:
                return
            path, _ = self._ppm_get_doc_path_and_pages()
            if not path:
                return
            if self.position_db.has_page_word_counts(path):
                self.ppm_modes_tabs.setCurrentIndex(2)  # Most Accurate
                self.ppm_pwc_indicator.setText(self.localization.get_text("using_stored_counts") if self.localization else "Using stored per-page counts")
        except Exception:
            pass

    def _ppm_get_doc_path_and_pages(self):
        # Try known attributes (prefer PPM-local state first)
        path = getattr(self, 'ppm_doc_path', None)
        total = getattr(self, 'ppm_total_pages', None)
        if not path:
            if getattr(self, 'pdf_path', None):
                path = self.pdf_path
            elif getattr(self, 'trainer_file_path', None):
                path = self.trainer_file_path
            elif getattr(self, 'current_document_path', None):
                path = self.current_document_path
        if total in (None, 0):
            total = getattr(self, 'total_pages', None)
        # If pages unknown but path is PDF, try to open to count
        try:
            if total in (None, 0) and path and str(path).lower().endswith('.pdf') and fitz:
                doc = fitz.open(path)
                total = doc.page_count
                doc.close()
        except Exception:
            pass
        return path, total

    def _ppm_toggle_timer(self):
        if self.ppm_running:
            self._ppm_stop()
        else:
            self._ppm_start()

    def _ppm_start(self):
        # Reset
        self.ppm_elapsed_seconds = 0
        self.ppm_timer_label.setText("00:00")
        self.ppm_running = True
        # Update button text and style for running state
        self.ppm_start_stop_btn.setText(self.localization.get_text("stop") if self.localization else "Stop")
        try:
            colors = self.get_theme_colors()
            self.ppm_start_stop_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {colors['accent_red']};
                    color: white;
                    border-radius: 6px;
                    padding: 6px 14px;
                    font-weight: bold;
                    font-size: 14px;
                }}
                QPushButton:disabled {{
                    background-color: {colors['bg_secondary']};
                    color: {colors['fg_muted']};
                }}
            """)
        except Exception:
            pass
        self.ppm_timer.start()

    def _ppm_stop(self):
        self.ppm_timer.stop()
        self.ppm_running = False
        # Update button text and style for idle state
        self.ppm_start_stop_btn.setText(self.localization.get_text("start") if self.localization else "Start")
        try:
            colors = self.get_theme_colors()
            self.ppm_start_stop_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {colors['accent_green']};
                    color: white;
                    border-radius: 6px;
                    padding: 6px 14px;
                    font-weight: bold;
                    font-size: 14px;
                }}
                QPushButton:disabled {{
                    background-color: {colors['bg_secondary']};
                    color: {colors['fg_muted']};
                }}
            """)
        except Exception:
            pass
        # Ask end page
        try:
            from PySide6.QtWidgets import QInputDialog
            start_page = self.ppm_start_page_spin.value()
            path, total = self._ppm_get_doc_path_and_pages()
            max_page = total or 9999
            end_page, ok = QInputDialog.getInt(self, self.localization.get_text("end_page") if self.localization else "End Page", self.localization.get_text("what_page_reached") if self.localization else "What page did you reach?", start_page, start_page, max_page)
            if ok:
                self._ppm_compute_and_show_results(start_page, end_page)
        except Exception as e:
            print(f"PPM: could not ask for end page: {e}")

    def _ppm_on_tick(self):
        self.ppm_elapsed_seconds += 1
        m = self.ppm_elapsed_seconds // 60
        s = self.ppm_elapsed_seconds % 60
        self.ppm_timer_label.setText(f"{m:02d}:{s:02d}")

    def _ppm_compute_and_show_results(self, start_page: int, end_page: int):
        minutes = max(self.ppm_elapsed_seconds / 60.0, 1e-6)
        pages_read = max(0, end_page - start_page)
        ppm = pages_read / minutes
        ppm_label = self.localization.get_text("ppm_label") if self.localization else "PPM"
        result = f"{ppm_label}: {ppm:.2f}"

        # Try to compute WPM if possible
        wpm = None
        path, _ = self._ppm_get_doc_path_and_pages()
        try:
            if self.position_db and path and self.position_db.has_page_word_counts(path):
                counts = self.position_db.load_page_word_counts(path) or {}
                if counts:
                    # pages are 1-based; sum start..(end-1)
                    total_words = 0
                    for p in range(start_page, end_page):
                        total_words += int(counts.get(p, 0))
                    wpm = total_words / minutes
            elif self.ppm_modes_tabs.currentIndex() == 1 and self.ppm_avg_wpp_spin.value() > 0:
                total_words = pages_read * self.ppm_avg_wpp_spin.value()
                wpm = total_words / minutes
        except Exception as e:
            print(f"PPM: error computing WPM: {e}")

        if wpm is not None:
            wpm_label = self.localization.get_text("wpm_label") if self.localization else "WPM"
            result += f"   |   {wpm_label}: {wpm:.0f}"
        self.ppm_result_label.setText(result)

    def _ppm_copy_to_clipboard(self, text: str):
        try:
            from PySide6.QtWidgets import QApplication
            cb = QApplication.clipboard()
            cb.setText(text or "")
        except Exception as e:
            print(f"PPM: copy to clipboard failed: {e}")

    def _ppm_parse_csv_and_save(self):
        text = self.ppm_csv_input.toPlainText().strip()
        path, _ = self._ppm_get_doc_path_and_pages()
        if not text or not path or self.position_db is None:
            QMessageBox.information(self, self.localization.get_text("info") if self.localization else "Info", self.localization.get_text("missing_data") if self.localization else "Missing CSV or document not loaded.")
            return
        counts = {}
        try:
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            # Detect header
            start_idx = 0
            if lines and ("," in lines[0] and ("page" in lines[0].lower() or "word" in lines[0].lower())):
                start_idx = 1
            for ln in lines[start_idx:]:
                parts = [p.strip() for p in ln.split(',')]
                if len(parts) >= 2:
                    try:
                        p = int(parts[0])
                        w = int(parts[1])
                        counts[p] = w
                    except ValueError:
                        continue
            if counts:
                ok = self.position_db.save_page_word_counts(path, counts)
                if ok:
                    QMessageBox.information(self, self.localization.get_text("saved") if self.localization else "Saved", self.localization.get_text("saved_page_counts") if self.localization else "Per-page word counts saved.")
                    # Reselect mode 3 automatically
                    self._ppm_autoselect_mode_based_on_data()
                else:
                    QMessageBox.warning(self, self.localization.get_text("error") if self.localization else "Error", self.localization.get_text("save_failed") if self.localization else "Could not save page counts.")
            else:
                QMessageBox.warning(self, self.localization.get_text("error") if self.localization else "Error", self.localization.get_text("parse_failed") if self.localization else "Could not parse any rows.")
        except Exception as e:
            print(f"PPM: CSV parse/save failed: {e}")
            QMessageBox.warning(self, self.localization.get_text("error") if self.localization else "Error", str(e))

    def _ppm_extract_sample_pages(self):
        """Extract N pages from the middle of the current PDF into a separate PDF for AI analysis."""
        try:
            if not fitz:
                QMessageBox.warning(self, self.localization.get_text("error") if self.localization else "Error", "PyMuPDF (fitz) required")
                return
            path, total = self._ppm_get_doc_path_and_pages()
            if not path or not str(path).lower().endswith('.pdf'):
                QMessageBox.information(self, self.localization.get_text("info") if self.localization else "Info", self.localization.get_text("open_pdf_first") if self.localization else "Please open a PDF first.")
                return
            n = max(1, int(self.ppm_sample_pages_spin.value()))
            # Choose middle block
            start = max(1, (total // 2) - (n // 2)) if total else 1
            end = min(total or (start + n - 1), start + n - 1)
            src = fitz.open(path)
            out = fitz.open()
            try:
                out.insert_pdf(src, from_page=start - 1, to_page=end - 1)
                base, ext = os.path.splitext(path)
                out_path = f"{base}_sample_{start}-{end}.pdf"
                out.save(out_path)
                QMessageBox.information(self, self.localization.get_text("done") if self.localization else "Done", (self.localization.get_text("sample_saved_to") if self.localization else "Sample saved to:") + f"\n{out_path}")
            finally:
                out.close()
                src.close()
        except Exception as e:
            print(f"PPM: extract sample failed: {e}")
            QMessageBox.warning(self, self.localization.get_text("error") if self.localization else "Error", str(e))

    def create_current_book_tab(self):
        """Create the Current Book tab with existing functionality"""
        current_book_tab = QWidget()
        tab_layout = QVBoxLayout(current_book_tab)
        tab_layout.setSpacing(15)
        tab_layout.setContentsMargins(15, 15, 15, 15)

        # Main content in splitter with responsive design
        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)  # Prevent sections from collapsing

        # Left side - File selection and controls
        left_widget = QWidget()
        left_widget.setMinimumWidth(250)  # Further reduced to eliminate scrolling
        left_layout = QVBoxLayout(left_widget)
        left_layout.setSpacing(15)  # Consistent spacing
        left_layout.setContentsMargins(10, 10, 10, 10)

        # Book card moved to Dashboard tab - no longer displayed here

        # File selection section
        file_group = QGroupBox(self.localization.get_text("select_pdf_for_speed"))
        colors = self.get_theme_colors()
        file_group.setStyleSheet(self.get_group_box_style(colors['accent_blue']))
        file_layout = QVBoxLayout(file_group)
        file_layout.setSpacing(12)
        file_layout.setContentsMargins(15, 15, 15, 15)

        # PDF file selection - responsive layout
        pdf_layout = QVBoxLayout()  # Changed to vertical layout for better responsiveness

        # File path display
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_file"))
        self.pdf_display.setMinimumHeight(35)
        pdf_layout.addWidget(self.pdf_display)

        # Icon-based button layout - horizontal layout with icons
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(15)
        buttons_layout.setContentsMargins(0, 10, 0, 10)

        # Get icon paths
        icon_base_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'new_icons')

        # Select PDF icon button
        self.btn_select_pdf = QToolButton()
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_icon_path = os.path.join(icon_base_path, 'pdf_file.png')
        if os.path.exists(pdf_icon_path):
            self.btn_select_pdf.setIcon(QIcon(pdf_icon_path))
            self.btn_select_pdf.setIconSize(QSize(64, 64))
        else:
            self.btn_select_pdf.setText("📄")
            self.btn_select_pdf.setStyleSheet("font-size: 48px;")
        self.btn_select_pdf.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
        self.btn_select_pdf.setToolTip(self.localization.get_text("select_pdf_tooltip") if self.localization else "Select PDF file for reading speed measurement")
        self.btn_select_pdf.setMinimumSize(80, 80)
        self.btn_select_pdf.setStyleSheet("""
            QToolButton {
                background-color: transparent;
                border: 2px solid #2196F3;
                border-radius: 10px;
                padding: 5px;
            }
            QToolButton:hover {
                background-color: #E3F2FD;
                border: 2px solid #1976D2;
            }
            QToolButton:pressed {
                background-color: #BBDEFB;
            }
        """)
        buttons_layout.addWidget(self.btn_select_pdf)

        # Prepare Book icon button
        self.btn_prepare_book = QToolButton()
        self.btn_prepare_book.clicked.connect(self.prepare_book_analysis)
        self.btn_prepare_book.setEnabled(False)
        self.btn_prepare_book.setText("🔍")
        self.btn_prepare_book.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
        prepare_tooltip = self.localization.get_text("prepare_book_tooltip") if self.localization else "Analyze book to calculate words per page"
        self.btn_prepare_book.setToolTip(prepare_tooltip)
        self.btn_prepare_book.setMinimumSize(80, 80)
        self.btn_prepare_book.setStyleSheet("""
            QToolButton {
                background-color: transparent;
                border: 2px solid #FF9800;
                border-radius: 10px;
                padding: 5px;
                font-size: 36px;
            }
            QToolButton:hover {
                background-color: #FFF3E0;
                border: 2px solid #F57C00;
            }
            QToolButton:pressed {
                background-color: #FFE0B2;
            }
            QToolButton:disabled {
                background-color: #f5f5f5;
                border: 2px solid #cccccc;
                color: #999999;
            }
        """)
        buttons_layout.addWidget(self.btn_prepare_book)

        buttons_layout.addStretch()
        pdf_layout.addLayout(buttons_layout)
        file_layout.addLayout(pdf_layout)

        # Recent books selection
        recent_label = QLabel(self.localization.get_text("or_select_from_recent"))
        recent_label.setStyleSheet("color: #666; font-weight: bold; margin-top: 10px;")
        file_layout.addWidget(recent_label)

        self.recent_books_list = QListWidget()
        self.recent_books_list.setMaximumHeight(120)
        self.recent_books_list.itemDoubleClicked.connect(lambda item: self.select_recent_book(item))
        self.recent_books_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: #fafafa;
                font-size: 12px;
            }
            QListWidget::item {
                padding: 8px;
                border-bottom: 1px solid #eee;
            }
            QListWidget::item:hover {
                background-color: #e3f2fd;
            }
            QListWidget::item:selected {
                background-color: #1976D2;
                color: white;
            }
        """)
        file_layout.addWidget(self.recent_books_list)

        # Load recent books (placeholder - will be populated when tab is accessed)
        pass

        file_group.setLayout(file_layout)
        left_layout.addWidget(file_group)

        # Analysis results section
        self.analysis_group = QGroupBox(self.localization.get_text("word_analysis_complete"))
        self.analysis_group.setVisible(False)  # Initially hidden
        self.analysis_group.setStyleSheet(self.get_group_box_style(colors['accent_green']))
        analysis_layout = QFormLayout(self.analysis_group)
        analysis_layout.setSpacing(12)
        analysis_layout.setContentsMargins(15, 15, 15, 15)
        analysis_layout.setVerticalSpacing(12)
        analysis_layout.setHorizontalSpacing(15)

        # Create styled labels for analysis results
        avg_words_label = QLabel(self.localization.get_text("avg_words_per_page"))
        avg_words_label.setStyleSheet("QLabel { color: #1976D2; font-weight: bold; }")
        self.lbl_avg_words = QLabel("---")
        self.lbl_avg_words.setStyleSheet("QLabel { color: #333; font-size: 14px; }")
        analysis_layout.addRow(avg_words_label, self.lbl_avg_words)

        total_pages_label = QLabel(self.localization.get_text("total_pages"))
        total_pages_label.setStyleSheet("QLabel { color: #1976D2; font-weight: bold; }")
        self.lbl_total_pages = QLabel("---")
        self.lbl_total_pages.setStyleSheet("QLabel { color: #333; font-size: 14px; }")
        analysis_layout.addRow(total_pages_label, self.lbl_total_pages)

        # Preparation method label
        method_label = QLabel(self.localization.get_text("preparation_method") if self.localization else "Method")
        method_label.setStyleSheet("QLabel { color: #1976D2; font-weight: bold; }")
        self.lbl_method = QLabel("---")
        self.lbl_method.setStyleSheet("QLabel { color: #333; font-size: 14px; }")
        analysis_layout.addRow(method_label, self.lbl_method)

        left_layout.addWidget(self.analysis_group)

        # Right side - Reading controls and timer
        right_widget = QWidget()
        right_widget.setMinimumWidth(300)
        right_layout = QVBoxLayout(right_widget)
        right_layout.setSpacing(15)
        right_layout.setContentsMargins(10, 10, 10, 10)

        # Reading controls section
        self.reading_group = QGroupBox(self.localization.get_text("start_reading_session"))
        self.reading_group.setVisible(False)  # Initially hidden until book is prepared
        self.reading_group.setStyleSheet(self.get_group_box_style(colors['accent_green']))
        reading_layout = QVBoxLayout(self.reading_group)
        reading_layout.setSpacing(15)
        reading_layout.setContentsMargins(15, 15, 15, 15)

        # Start reading button (will toggle between Start/Pause/Resume)
        self.btn_start_reading = QPushButton(self.localization.get_text("start_reading_session"))
        self.btn_start_reading.clicked.connect(self.toggle_reading_session)
        self.btn_start_reading.setEnabled(False)
        self.btn_start_reading.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 16px;
                min-height: 50px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        reading_layout.addWidget(self.btn_start_reading)

        # Timer display
        self.lbl_timer = QLabel("00:00")
        self.lbl_timer.setAlignment(Qt.AlignCenter)
        self.lbl_timer.setStyleSheet("""
            QLabel {
                font-size: 48px;
                font-weight: bold;
                color: #E91E63;
                padding: 20px;
                background-color: #fce4ec;
                border: 3px solid #e91e63;
                border-radius: 15px;
                min-height: 80px;
                margin: 15px 0;
            }
        """)
        reading_layout.addWidget(self.lbl_timer)

        # Finish reading controls (initially hidden)
        self.finish_widget = QWidget()
        self.finish_widget.setVisible(False)
        finish_layout = QVBoxLayout(self.finish_widget)

        # Pages read input
        pages_layout = QHBoxLayout()
        pages_layout.addWidget(QLabel(self.localization.get_text("pages_read")))
        self.spin_pages_read = QSpinBox()
        self.spin_pages_read.setMinimum(1)
        self.spin_pages_read.setMaximum(9999)
        self.spin_pages_read.setValue(1)
        self.spin_pages_read.setStyleSheet("QSpinBox { font-size: 14px; padding: 5px; }")
        pages_layout.addWidget(self.spin_pages_read)
        pages_layout.addStretch()
        finish_layout.addLayout(pages_layout)

        # Finish reading button
        self.btn_finish_reading = QPushButton(self.localization.get_text("finish_reading"))
        self.btn_finish_reading.clicked.connect(self.finish_reading_session)
        self.btn_finish_reading.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        finish_layout.addWidget(self.btn_finish_reading)

        reading_layout.addWidget(self.finish_widget)
        self.reading_group.setLayout(reading_layout)
        right_layout.addWidget(self.reading_group)

        # Results section
        self.results_group = QGroupBox("🎯 " + self.localization.get_text("reading_speed_results"))
        self.results_group.setVisible(False)  # Initially hidden
        results_layout = QVBoxLayout(self.results_group)
        results_layout.setSpacing(15)

        # WPM display
        wpm_frame = QFrame()
        wpm_frame.setFrameStyle(QFrame.Box)
        wpm_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #1976D2;
                border-radius: 10px;
                background-color: #e3f2fd;
                padding: 15px;
            }
        """)
        wpm_layout = QVBoxLayout(wpm_frame)

        wpm_title = QLabel("📚 " + self.localization.get_text("your_reading_speed"))
        wpm_title.setAlignment(Qt.AlignCenter)
        wpm_title.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; color: #1976D2; }")
        wpm_layout.addWidget(wpm_title)

        self.lbl_wpm = QLabel("--- " + self.localization.get_text("wpm"))
        self.lbl_wpm.setAlignment(Qt.AlignCenter)
        self.lbl_wpm.setStyleSheet("""
            QLabel {
                font-size: 32px;
                font-weight: bold;
                color: #1976D2;
                margin: 10px 0;
            }
        """)
        wpm_layout.addWidget(self.lbl_wpm)
        results_layout.addWidget(wpm_frame)

        # Additional stats
        stats_layout = QHBoxLayout()

        # Time taken
        time_frame = QFrame()
        time_frame.setFrameStyle(QFrame.Box)
        time_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ddd;
                border-radius: 8px;
                background-color: #f9f9f9;
                padding: 10px;
            }
        """)
        time_layout = QVBoxLayout(time_frame)
        time_title = QLabel(self.localization.get_text("total_time_taken"))
        time_title.setStyleSheet("QLabel { font-weight: bold; color: #666; }")
        time_layout.addWidget(time_title)
        self.lbl_total_time = QLabel("---")
        self.lbl_total_time.setStyleSheet("QLabel { font-size: 18px; color: #333; }")
        time_layout.addWidget(self.lbl_total_time)
        stats_layout.addWidget(time_frame)

        # Words read
        words_frame = QFrame()
        words_frame.setFrameStyle(QFrame.Box)
        words_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ddd;
                border-radius: 8px;
                background-color: #f9f9f9;
                padding: 10px;
            }
        """)
        words_layout = QVBoxLayout(words_frame)
        words_title = QLabel(self.localization.get_text("total_words_read") if self.localization else "Words Read")
        words_title.setStyleSheet("QLabel { font-weight: bold; color: #666; }")
        words_layout.addWidget(words_title)
        self.lbl_words_read = QLabel("---")
        self.lbl_words_read.setStyleSheet("QLabel { font-size: 18px; color: #333; }")
        words_layout.addWidget(self.lbl_words_read)
        stats_layout.addWidget(words_frame)

        results_layout.addLayout(stats_layout)

        # Recommendation
        self.lbl_recommendation = QLabel("...")
        self.lbl_recommendation.setWordWrap(True)
        self.lbl_recommendation.setStyleSheet("""
            QLabel {
                background-color: #fff3e0;
                border: 1px solid #ff9800;
                border-radius: 8px;
                padding: 15px;
                font-size: 14px;
                color: #e65100;
                margin-top: 10px;
            }
        """)
        results_layout.addWidget(self.lbl_recommendation)

        # New session button
        self.btn_new_session = QPushButton(self.localization.get_text("new_reading_session"))
        self.btn_new_session.clicked.connect(self.reset_session)
        self.btn_new_session.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
                margin-top: 10px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        results_layout.addWidget(self.btn_new_session)

        right_layout.addWidget(self.results_group)
        right_layout.addStretch()

        # Add widgets to splitter
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)  # Left side gets 1 part
        splitter.setStretchFactor(1, 1)  # Right side gets 1 part

        tab_layout.addWidget(splitter)

        # Add the Current Book tab
        current_book_text = self.localization.get_text('current_book') if self.localization else 'Current Book'
        self.tab_widget.addTab(current_book_tab, f"📖 {current_book_text}")



    def update_analysis_group_theme(self):
        """Update analysis group theme based on current theme"""
        if hasattr(self, 'parent_window') and self.parent_window:
            # Get current theme from parent window
            current_theme = getattr(self.parent_window, 'current_theme', 'light')
            if hasattr(self.parent_window, 'settings'):
                current_theme = self.parent_window.settings.get("theme", "light")
        else:
            current_theme = "light"

        if current_theme == "dark":
            # Dark theme styling
            self.analysis_group.setStyleSheet("""
                QGroupBox {
                    font-weight: bold;
                    font-size: 16px;
                    color: #BB86FC;
                    border: 2px solid #BB86FC;
                    border-radius: 10px;
                    margin: 10px 0;
                    padding: 15px;
                    background-color: #2d2d2d;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 15px;
                    padding: 0 10px 0 10px;
                    color: #BB86FC;
                    font-weight: bold;
                    background-color: #2d2d2d;
                }
            """)
        else:
            # Light theme styling
            self.analysis_group.setStyleSheet("""
                QGroupBox {
                    font-weight: bold;
                    font-size: 16px;
                    color: #7B1FA2;
                    border: 2px solid #9C27B0;
                    border-radius: 10px;
                    margin: 10px 0;
                    padding: 15px;
                    background-color: #fafafa;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 15px;
                    padding: 0 10px 0 10px;
                    color: #7B1FA2;
                    font-weight: bold;
                    background-color: #fafafa;
                }
            """)

    def update_sort_combo_theme(self):
        """Update sort combo theme based on current theme"""
        if hasattr(self, 'parent_window') and self.parent_window:
            current_theme = getattr(self.parent_window, 'current_theme', 'light')
            if hasattr(self.parent_window, 'settings'):
                current_theme = self.parent_window.settings.get("theme", "light")
        else:
            current_theme = "light"

        if current_theme == "dark":
            self.sort_combo.setStyleSheet("""
                QComboBox {
                    background-color: #2d2d2d;
                    color: #ffffff;
                    border: 1px solid #555555;
                    border-radius: 4px;
                    padding: 6px 12px;
                    min-width: 120px;
                }
                QComboBox:hover {
                    border-color: #BB86FC;
                }
                QComboBox::drop-down {
                    border: none;
                    background-color: #2d2d2d;
                }
                QComboBox QAbstractItemView {
                    background-color: #2d2d2d;
                    color: #ffffff;
                    border: 1px solid #555555;
                    selection-background-color: #BB86FC;
                }
            """)
        else:
            self.sort_combo.setStyleSheet("""
                QComboBox {
                    background-color: white;
                    color: #000000;
                    border: 1px solid #ddd;
                    border-radius: 4px;
                    padding: 6px 12px;
                    min-width: 120px;
                }
                QComboBox:hover {
                    border-color: #2196F3;
                }
            """)

    def update_table_theme(self, table_widget):
        """Update table theme based on current theme"""
        if hasattr(self, 'parent_window') and self.parent_window:
            current_theme = getattr(self.parent_window, 'current_theme', 'light')
            if hasattr(self.parent_window, 'settings'):
                current_theme = self.parent_window.settings.get("theme", "light")
        else:
            current_theme = "light"

        if current_theme == "dark":
            table_widget.setStyleSheet("""
                QTableWidget {
                    gridline-color: #555555;
                    background-color: #2d2d2d;
                    alternate-background-color: #3d3d3d;
                    color: #ffffff;
                    selection-background-color: #BB86FC;
                    border: 1px solid #555555;
                }
                QHeaderView::section {
                    background-color: #1e1e1e;
                    color: #ffffff;
                    padding: 8px;
                    border: 1px solid #555555;
                    font-weight: bold;
                }
                QTableWidget::item {
                    padding: 8px;
                    border-bottom: 1px solid #555555;
                }
                QTableWidget::item:selected {
                    background-color: #BB86FC;
                    color: #000000;
                }
            """)
        else:
            table_widget.setStyleSheet("""
                QTableWidget {
                    gridline-color: #e0e0e0;
                    background-color: white;
                    alternate-background-color: #f5f5f5;
                    color: #000000;
                    selection-background-color: #e8f5e8;
                    border: 1px solid #e0e0e0;
                }
                QHeaderView::section {
                    background-color: #f0f0f0;
                    color: #000000;
                    padding: 8px;
                    border: 1px solid #e0e0e0;
                    font-weight: bold;
                }
                QTableWidget::item {
                    padding: 8px;
                    border-bottom: 1px solid #e0e0e0;
                }
            """)

    def show_reading_speed_guide(self):
        """Show comprehensive reading speed measurement guide"""
        guide_dialog = QDialog(self)
        guide_dialog.setWindowTitle(self.localization.get_text("reading_speed_guide"))
        guide_dialog.setWindowIcon(self.parent_window.get_app_icon() if hasattr(self.parent_window, 'get_app_icon') else None)
        guide_dialog.resize(800, 600)

        # Create layout
        layout = QVBoxLayout(guide_dialog)

        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)

        # Create content widget
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)

        # Create text widget
        text_widget = QTextEdit()
        text_widget.setReadOnly(True)

        if self.localization and self.localization.current_language == "ar":
            guide_text = self._get_arabic_reading_speed_guide()
        else:
            guide_text = self._get_english_reading_speed_guide()

        text_widget.setHtml(guide_text)
        content_layout.addWidget(text_widget)

        # Set content widget to scroll area
        scroll_area.setWidget(content_widget)
        layout.addWidget(scroll_area)

        # Add close button
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        close_btn = QPushButton(self.localization.get_text("close") if self.localization else "Close")
        close_btn.clicked.connect(guide_dialog.close)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border-radius: 5px;
                border: none;
            }
            QPushButton:hover {
                background-color: #1565C0;
            }
        """)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)

        # Apply theme
        if hasattr(self, 'parent_window') and self.parent_window:
            current_theme = getattr(self.parent_window, 'current_theme', 'light')
            if hasattr(self.parent_window, 'settings'):
                current_theme = self.parent_window.settings.get("theme", "light")
        else:
            current_theme = "light"

        if current_theme == "dark":
            guide_dialog.setStyleSheet("""
                QDialog {
                    background-color: #1e1e1e;
                    color: #ffffff;
                }
                QTextEdit {
                    background-color: #2d2d2d;
                    color: #ffffff;
                    border: 1px solid #555555;
                }
                QScrollArea {
                    background-color: #1e1e1e;
                    border: none;
                }
            """)

        guide_dialog.exec()

    def _get_arabic_reading_speed_guide(self):
        """Get comprehensive Arabic reading speed guide content"""
        return """
        <div style="direction: rtl; text-align: right;">
        <h2 style="color: #1976D2; text-align: center;">📖 دليل قياس سرعة القراءة</h2>

        <h3 style="color: #1976D2;">🎯 مرحباً بك في أداة قياس سرعة القراءة!</h3>
        <p><b>هذه الأداة تساعدك على قياس سرعة قراءتك بدقة مع اختبار الفهم والاستيعاب.</b></p>

        <h3 style="color: #1976D2;">📋 خطوات الاستخدام</h3>

        <h4>1️⃣ اختيار ملف PDF</h4>
        <ul>
        <li>انقر على زر "اختيار ملف" لتحديد ملف PDF المراد قراءته</li>
        <li>تأكد من أن الملف يحتوي على نص قابل للقراءة (وليس صور فقط)</li>
        <li>يُفضل اختيار ملف بحجم مناسب (10-50 صفحة)</li>
        </ul>

        <h4>2️⃣ تحضير الكتاب للقياس</h4>
        <ul>
        <li>انقر على زر "🔍 تحضير الكتاب" بعد اختيار الملف</li>
        <li>ستقوم الأداة بتحليل محتوى الكتاب وحساب متوسط الكلمات في الصفحة</li>
        <li>هذه الخطوة ضرورية للحصول على قياس دقيق</li>
        </ul>

        <h4>3️⃣ بدء القراءة</h4>
        <ul>
        <li>انقر على زر "بدء القراءة" عندما تكون مستعداً</li>
        <li>ستبدأ المؤقت في العد</li>
        <li>اقرأ النص المعروض بتركيز وفهم</li>
        <li>لا تتسرع - الهدف هو القراءة بفهم وليس السرعة فقط</li>
        </ul>

        <h4>4️⃣ إنهاء القراءة</h4>
        <ul>
        <li>عند الانتهاء من القراءة، انقر على "انتهيت من القراءة"</li>
        <li>سيتوقف المؤقت وسيتم حساب الوقت المستغرق</li>
        </ul>

        <h4>5️⃣ اختبار الفهم</h4>
        <ul>
        <li>ستظهر أسئلة حول النص الذي قرأته</li>
        <li>أجب على الأسئلة بصدق وتركيز</li>
        <li>هذا الاختبار يقيس مدى فهمك للنص</li>
        </ul>

        <h4>6️⃣ عرض النتائج</h4>
        <ul>
        <li>ستحصل على تقرير شامل يتضمن:</li>
        <li><b>سرعة القراءة:</b> عدد الكلمات في الدقيقة (WPM)</li>
        <li><b>نسبة الفهم:</b> النسبة المئوية للإجابات الصحيحة</li>
        <li><b>الوقت المستغرق:</b> إجمالي وقت القراءة</li>
        <li><b>عدد الصفحات:</b> الصفحات التي تمت قراءتها</li>
        </ul>

        <h3 style="color: #1976D2;">💡 نصائح للحصول على قياس دقيق</h3>

        <h4>📚 اختيار النص المناسب</h4>
        <ul>
        <li>اختر نصاً بمستوى صعوبة مناسب لك</li>
        <li>تجنب النصوص التقنية المعقدة في البداية</li>
        <li>استخدم نصوص باللغة التي تتقنها</li>
        </ul>

        <h4>🎯 بيئة القراءة</h4>
        <ul>
        <li>اقرأ في مكان هادئ وخالٍ من المشتتات</li>
        <li>تأكد من الإضاءة الجيدة</li>
        <li>اجلس في وضعية مريحة</li>
        <li>أغلق الهاتف والإشعارات</li>
        </ul>

        <h4>⏰ التوقيت</h4>
        <ul>
        <li>اختر وقتاً تكون فيه نشيطاً ومتيقظاً</li>
        <li>تجنب القياس عند التعب أو النعاس</li>
        <li>خذ استراحة قبل البدء إذا كنت متعباً</li>
        </ul>

        <h3 style="color: #1976D2;">📊 فهم النتائج</h3>

        <h4>⚡ سرعة القراءة (WPM)</h4>
        <ul>
        <li><b>أقل من 200 كلمة/دقيقة:</b> سرعة بطيئة - يمكن تحسينها</li>
        <li><b>200-300 كلمة/دقيقة:</b> سرعة متوسطة - جيدة للمبتدئين</li>
        <li><b>300-400 كلمة/دقيقة:</b> سرعة جيدة - فوق المتوسط</li>
        <li><b>400-500 كلمة/دقيقة:</b> سرعة ممتازة - مستوى متقدم</li>
        <li><b>أكثر من 500 كلمة/دقيقة:</b> سرعة استثنائية</li>
        </ul>

        <h4>🎯 نسبة الفهم</h4>
        <ul>
        <li><b>90-100%:</b> فهم ممتاز - توازن مثالي بين السرعة والفهم</li>
        <li><b>80-89%:</b> فهم جيد جداً - مستوى مقبول</li>
        <li><b>70-79%:</b> فهم جيد - يمكن التحسين</li>
        <li><b>أقل من 70%:</b> يُنصح بتقليل السرعة والتركيز على الفهم</li>
        </ul>

        <h3 style="color: #1976D2;">🚀 نصائح لتحسين سرعة القراءة</h3>

        <ul>
        <li><b>التدرب المنتظم:</b> اقرأ يومياً لمدة 15-30 دقيقة</li>
        <li><b>تجنب القراءة الصوتية:</b> لا تنطق الكلمات في ذهنك</li>
        <li><b>توسيع مجال الرؤية:</b> حاول قراءة عدة كلمات في نظرة واحدة</li>
        <li><b>تقليل الرجوع:</b> تجنب إعادة قراءة الجمل</li>
        <li><b>استخدام المؤشر:</b> استخدم إصبعك أو قلماً لتوجيه العين</li>
        <li><b>القراءة المتدرجة:</b> ابدأ بنصوص سهلة ثم انتقل للأصعب</li>
        </ul>

        <h3 style="color: #1976D2;">❓ الأسئلة الشائعة</h3>

        <h4>س: هل يمكنني استخدام أي نوع من ملفات PDF؟</h4>
        <p>ج: نعم، لكن تأكد من أن الملف يحتوي على نص قابل للتحديد وليس مجرد صور.</p>

        <h4>س: كم مرة يجب أن أقيس سرعة قراءتي؟</h4>
        <p>ج: يُنصح بالقياس مرة واحدة أسبوعياً لتتبع التقدم.</p>

        <h4>س: ماذا لو كانت نسبة الفهم منخفضة؟</h4>
        <p>ج: ركز على الفهم أكثر من السرعة. اقرأ ببطء أكثر حتى تتحسن نسبة الفهم.</p>

        <h4>س: هل يمكنني مقارنة نتائجي مع الآخرين؟</h4>
        <p>ج: التركيز على تحسين أدائك الشخصي أهم من المقارنة مع الآخرين.</p>

        <hr>
        <p style="text-align: center; color: #666;"><i>نتمنى لك قراءة ممتعة ومفيدة!</i></p>
        </div>
        """

    def _get_english_reading_speed_guide(self):
        """Get comprehensive English reading speed guide content"""
        return """
        <h2 style="color: #1976D2; text-align: center;">📖 Reading Speed Measurement Guide</h2>

        <h3 style="color: #1976D2;">🎯 Welcome to the Reading Speed Tool!</h3>
        <p><b>This tool helps you accurately measure your reading speed with comprehension testing.</b></p>

        <h3 style="color: #1976D2;">📋 How to Use</h3>

        <h4>1️⃣ Select PDF File</h4>
        <ul>
        <li>Click "Select File" button to choose a PDF file to read</li>
        <li>Ensure the file contains readable text (not just images)</li>
        <li>Prefer files of appropriate size (10-50 pages)</li>
        </ul>

        <h4>2️⃣ Prepare Book for Measurement</h4>
        <ul>
        <li>Click "🔍 Prepare Book" button after selecting the file</li>
        <li>The tool will analyze the book content and calculate average words per page</li>
        <li>This step is essential for accurate measurement</li>
        </ul>

        <h4>3️⃣ Start Reading</h4>
        <ul>
        <li>Click "Start Reading" when you're ready</li>
        <li>The timer will begin counting</li>
        <li>Read the displayed text with focus and understanding</li>
        <li>Don't rush - the goal is reading with comprehension, not just speed</li>
        </ul>

        <h4>4️⃣ Finish Reading</h4>
        <ul>
        <li>When you finish reading, click "Finished Reading"</li>
        <li>The timer will stop and calculate the elapsed time</li>
        </ul>

        <h4>5️⃣ Comprehension Test</h4>
        <ul>
        <li>Questions about the text you read will appear</li>
        <li>Answer the questions honestly and with focus</li>
        <li>This test measures your understanding of the text</li>
        </ul>

        <h4>6️⃣ View Results</h4>
        <ul>
        <li>You'll receive a comprehensive report including:</li>
        <li><b>Reading Speed:</b> Words per minute (WPM)</li>
        <li><b>Comprehension Rate:</b> Percentage of correct answers</li>
        <li><b>Time Taken:</b> Total reading time</li>
        <li><b>Pages Read:</b> Number of pages completed</li>
        </ul>

        <h3 style="color: #1976D2;">💡 Tips for Accurate Measurement</h3>

        <h4>📚 Choosing the Right Text</h4>
        <ul>
        <li>Choose text with appropriate difficulty level for you</li>
        <li>Avoid complex technical texts initially</li>
        <li>Use texts in your proficient language</li>
        </ul>

        <h4>🎯 Reading Environment</h4>
        <ul>
        <li>Read in a quiet place free from distractions</li>
        <li>Ensure good lighting</li>
        <li>Sit in a comfortable position</li>
        <li>Turn off phone and notifications</li>
        </ul>

        <h4>⏰ Timing</h4>
        <ul>
        <li>Choose a time when you're active and alert</li>
        <li>Avoid measuring when tired or sleepy</li>
        <li>Take a break before starting if you're tired</li>
        </ul>

        <h3 style="color: #1976D2;">📊 Understanding Results</h3>

        <h4>⚡ Reading Speed (WPM)</h4>
        <ul>
        <li><b>Less than 200 words/minute:</b> Slow speed - can be improved</li>
        <li><b>200-300 words/minute:</b> Average speed - good for beginners</li>
        <li><b>300-400 words/minute:</b> Good speed - above average</li>
        <li><b>400-500 words/minute:</b> Excellent speed - advanced level</li>
        <li><b>More than 500 words/minute:</b> Exceptional speed</li>
        </ul>

        <h4>🎯 Comprehension Rate</h4>
        <ul>
        <li><b>90-100%:</b> Excellent comprehension - perfect balance of speed and understanding</li>
        <li><b>80-89%:</b> Very good comprehension - acceptable level</li>
        <li><b>70-79%:</b> Good comprehension - can be improved</li>
        <li><b>Less than 70%:</b> Recommended to slow down and focus on understanding</li>
        </ul>

        <h3 style="color: #1976D2;">🚀 Tips to Improve Reading Speed</h3>

        <ul>
        <li><b>Regular Practice:</b> Read daily for 15-30 minutes</li>
        <li><b>Avoid Subvocalization:</b> Don't pronounce words in your mind</li>
        <li><b>Expand Visual Span:</b> Try to read multiple words in one glance</li>
        <li><b>Reduce Regression:</b> Avoid re-reading sentences</li>
        <li><b>Use a Pointer:</b> Use your finger or a pen to guide your eyes</li>
        <li><b>Progressive Reading:</b> Start with easy texts then move to harder ones</li>
        </ul>

        <h3 style="color: #1976D2;">❓ Frequently Asked Questions</h3>

        <h4>Q: Can I use any type of PDF file?</h4>
        <p>A: Yes, but make sure the file contains selectable text and not just images.</p>

        <h4>Q: How often should I measure my reading speed?</h4>
        <p>A: It's recommended to measure once a week to track progress.</p>

        <h4>Q: What if my comprehension rate is low?</h4>
        <p>A: Focus on understanding more than speed. Read slower until comprehension improves.</p>

        <h4>Q: Can I compare my results with others?</h4>
        <p>A: Focusing on improving your personal performance is more important than comparing with others.</p>

        <hr>
        <p style="text-align: center; color: #666;"><i>Happy reading and learning!</i></p>
        """

    def update_theme(self):
        """Update theme for all components in the reading speed tab"""
        # Update current theme
        self.current_theme = self.get_current_theme()

        # Update analysis group theme
        if hasattr(self, 'analysis_group'):
            self.update_analysis_group_theme()

        # Update sort combo theme
        if hasattr(self, 'sort_combo'):
            self.update_sort_combo_theme()

        # Update all table themes
        tables = []
        if hasattr(self, 'sessions_table'):
            tables.append(self.sessions_table)
        if hasattr(self, 'books_table'):
            tables.append(self.books_table)
        if hasattr(self, 'history_table'):
            tables.append(self.history_table)
        if hasattr(self, 'book_sessions_table'):
            tables.append(self.book_sessions_table)

        for table in tables:
            self.update_table_theme(table)

        # Update book card theme if it exists
        if hasattr(self, 'book_card'):
            colors = self.get_theme_colors()
            self.book_card.setStyleSheet(f"""
                QGroupBox {{
                    border: 3px solid {colors['accent_blue']};
                    border-radius: 15px;
                    margin: 15px;
                    padding: 20px;
                    background-color: {colors['bg_primary']};
                    font-weight: bold;
                    font-size: 16px;
                }}
                QGroupBox::title {{
                    subcontrol-origin: margin;
                    left: 15px;
                    padding: 0 15px 0 15px;
                    color: {colors['accent_blue']};
                    font-weight: bold;
                    font-size: 18px;
                }}
            """)

        # Update trends content theme if it exists
        if hasattr(self, 'trends_content'):
            colors = self.get_theme_colors()
            self.trends_content.setStyleSheet(f"""
                QLabel {{
                    font-size: 14px;
                    color: {colors['text_primary']};
                    padding: 15px;
                    background-color: {colors['bg_secondary']};
                    border-radius: 8px;
                    border: 1px solid {colors['border']};
                    line-height: 1.6;
                }}
            """)

        # Update training controls theme if they exist
        self.update_training_controls_theme()

    def update_training_controls_theme(self):
        """Update theme for training controls"""
        colors = self.get_theme_colors()

        # Enhanced spinbox style with theme awareness
        enhanced_spinbox_style = f"""
            QSpinBox {{
                padding: 8px;
                border: 2px solid {colors['border']};
                border-radius: 6px;
                font-size: 14px;
                background: {colors['bg_primary']};
                color: {colors['text_primary']};
                min-height: 20px;
            }}
            QSpinBox:focus {{
                border-color: {colors['accent_orange']};
            }}
            QSpinBox::up-button {{
                subcontrol-origin: border;
                subcontrol-position: top right;
                width: 25px;
                border-left: 1px solid {colors['border']};
                border-bottom: 1px solid {colors['border']};
                border-top-right-radius: 4px;
                background: {colors['bg_tertiary']};
            }}
            QSpinBox::up-button:hover {{
                background: {colors['bg_secondary']};
            }}
            QSpinBox::up-button:pressed {{
                background: {colors['border']};
            }}
            QSpinBox::up-arrow {{
                image: none;
                border: 4px solid transparent;
                border-bottom: 6px solid {colors['text_muted']};
                width: 0px;
                height: 0px;
            }}
            QSpinBox::down-button {{
                subcontrol-origin: border;
                subcontrol-position: bottom right;
                width: 25px;
                border-left: 1px solid {colors['border']};
                border-top: 1px solid {colors['border']};
                border-bottom-right-radius: 4px;
                background: {colors['bg_tertiary']};
            }}
            QSpinBox::down-button:hover {{
                background: {colors['bg_secondary']};
            }}
            QSpinBox::down-button:pressed {{
                background: {colors['border']};
            }}
            QSpinBox::down-arrow {{
                image: none;
                border: 4px solid transparent;
                border-top: 6px solid {colors['text_muted']};
                width: 0px;
                height: 0px;
            }}
        """

        # Update spinboxes if they exist
        if hasattr(self, 'wpg_spinbox'):
            self.wpg_spinbox.setStyleSheet(enhanced_spinbox_style)
        if hasattr(self, 'trainer_wpm_spinbox'):
            self.trainer_wpm_spinbox.setStyleSheet(enhanced_spinbox_style)
        if hasattr(self, 'start_page_spinbox'):
            self.start_page_spinbox.setStyleSheet(enhanced_spinbox_style)
        if hasattr(self, 'font_size_spinbox'):
            self.font_size_spinbox.setStyleSheet(enhanced_spinbox_style)

        # Update font family combo if it exists
        if hasattr(self, 'font_family_combo'):
            self.font_family_combo.setStyleSheet(f"""
                QComboBox {{
                    padding: 8px;
                    border: 2px solid {colors['border']};
                    border-radius: 6px;
                    font-size: 14px;
                    background: {colors['bg_primary']};
                    color: {colors['text_primary']};
                    min-height: 20px;
                    combobox-popup: 0;
                }}
                QComboBox:focus {{
                    border-color: #FF9800;
                }}
                QComboBox::drop-down {{
                    border: none;
                    width: 25px;
                    background: {colors['bg_tertiary']};
                    border-top-right-radius: 4px;
                    border-bottom-right-radius: 4px;
                }}
                QComboBox::drop-down:hover {{
                    background: {colors['bg_secondary']};
                }}
                QComboBox::down-arrow {{
                    image: none;
                    border: 4px solid transparent;
                    border-top: 6px solid {colors['text_muted']};
                    width: 0px;
                    height: 0px;
                }}
                QComboBox QAbstractItemView {{
                    background-color: {colors['bg_primary']};
                    color: {colors['text_primary']};
                    selection-background-color: {colors['accent_blue']};
                }}
            """)

        # Analysis results
        self.analysis_group = QGroupBox(self.localization.get_text("word_analysis_complete"))
        # Store reference for theme updates
        self.analysis_group_widget = self.analysis_group
        self.update_analysis_group_theme()  # Apply initial theme
        analysis_layout = QFormLayout(self.analysis_group)
        analysis_layout.setSpacing(10)

        self.lbl_avg_words = QLabel("---")
        self.lbl_avg_words.setStyleSheet("""
            QLabel {
                font-size: 20px;
                font-weight: bold;
                color: #673AB7;
                padding: 8px;
                background-color: #f3e5f5;
                border-radius: 6px;
                border: 1px solid #ce93d8;
            }
        """)

        self.lbl_total_pages = QLabel("---")
        self.lbl_total_pages.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #607D8B;
                padding: 8px;
                background-color: #eceff1;
                border-radius: 6px;
                border: 1px solid #b0bec5;
            }
        """)

        # Create styled labels for form rows
        avg_words_label = QLabel(self.localization.get_text("avg_words_per_page"))
        avg_words_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: #424242;
                padding: 5px;
            }
        """)

        total_pages_label = QLabel(self.localization.get_text("total_pages"))
        total_pages_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: #424242;
                padding: 5px;
            }
        """)

        analysis_layout.addRow(avg_words_label, self.lbl_avg_words)
        analysis_layout.addRow(total_pages_label, self.lbl_total_pages)

        self.analysis_group.setVisible(False)
        left_layout.addWidget(self.analysis_group)

        # Reading controls
        self.reading_group = QGroupBox(self.localization.get_text("start_reading_session"))
        reading_layout = QVBoxLayout(self.reading_group)

        self.btn_start_reading = QPushButton(self.localization.get_text("start_reading_session"))
        self.btn_start_reading.clicked.connect(self.start_reading_session)
        self.btn_start_reading.setEnabled(False)
        self.btn_start_reading.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-size: 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        reading_layout.addWidget(self.btn_start_reading)

        # Timer display
        self.lbl_timer = QLabel("00:00")
        self.lbl_timer.setAlignment(Qt.AlignCenter)
        self.lbl_timer.setStyleSheet("""
            QLabel {
                font-size: 36px;
                font-weight: bold;
                color: #E91E63;
                padding: 20px;
                background-color: #fce4ec;
                border: 3px solid #e91e63;
                border-radius: 15px;
                min-height: 60px;
                margin: 15px 0;
            }
        """)
        reading_layout.addWidget(self.lbl_timer)

        # Finish reading controls
        finish_layout = QHBoxLayout()
        finish_layout.addWidget(QLabel(self.localization.get_text("pages_read")))

        self.spin_pages_read = QSpinBox()
        self.spin_pages_read.setMinimum(1)
        finish_layout.addWidget(self.spin_pages_read)

        self.btn_finish_reading = QPushButton(self.localization.get_text("finish_reading"))
        self.btn_finish_reading.clicked.connect(self.finish_reading_session)
        self.btn_finish_reading.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 10px 15px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        finish_layout.addWidget(self.btn_finish_reading)

        self.finish_widget = QWidget()
        self.finish_widget.setLayout(finish_layout)
        self.finish_widget.setVisible(False)
        reading_layout.addWidget(self.finish_widget)

        self.reading_group.setVisible(False)
        left_layout.addWidget(self.reading_group)

        left_layout.addStretch()
        splitter.addWidget(left_widget)

        # Right side - Results and comprehension test
        right_widget = QWidget()
        right_widget.setMinimumWidth(300)  # Further reduced to eliminate scrolling
        right_layout = QVBoxLayout(right_widget)
        right_layout.setSpacing(15)  # Consistent spacing
        right_layout.setContentsMargins(10, 10, 10, 10)

        # Enhanced Results section with improved visual design
        self.results_group = QGroupBox("🎯 " + self.localization.get_text("reading_speed_results"))
        colors = self.get_theme_colors()
        self.results_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                font-size: 18px;
                color: {colors['accent_blue']};
                border: 3px solid {colors['accent_blue']};
                border-radius: 15px;
                margin: 15px 0;
                padding: 25px;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 {colors['bg_primary']}, stop:1 {colors['bg_secondary']});
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 20px;
                padding: 5px 15px;
                color: {colors['accent_blue']};
                font-weight: bold;
                font-size: 18px;
                background-color: {colors['bg_primary']};
                border-radius: 8px;
                border: 2px solid {colors['accent_blue']};
            }}
        """)
        results_layout = QVBoxLayout(self.results_group)  # Changed to VBoxLayout for better control
        results_layout.setSpacing(20)

        # Create enhanced result cards with better visual hierarchy

        # Main WPM result - most prominent
        wpm_card = QWidget()
        wpm_layout = QVBoxLayout(wpm_card)
        wpm_layout.setSpacing(5)

        wpm_title = QLabel("📚 " + self.localization.get_text("your_reading_speed"))
        wpm_title.setAlignment(Qt.AlignCenter)
        wpm_title.setStyleSheet(f"""
            QLabel {{
                font-size: 16px;
                font-weight: bold;
                color: {colors['text_secondary']};
                margin-bottom: 5px;
            }}
        """)

        self.lbl_wpm = QLabel("--- " + self.localization.get_text("wpm"))
        self.lbl_wpm.setAlignment(Qt.AlignCenter)
        self.lbl_wpm.setStyleSheet(f"""
            QLabel {{
                font-size: 36px;
                font-weight: bold;
                color: {colors['accent_blue']};
                padding: 20px;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #e3f2fd, stop:1 #bbdefb);
                border-radius: 15px;
                border: 3px solid {colors['accent_blue']};
                min-height: 80px;
            }}
        """)

        wpm_layout.addWidget(wpm_title)
        wpm_layout.addWidget(self.lbl_wpm)
        wmp_card_container = QWidget()
        wmp_card_container.setStyleSheet(f"""
            QWidget {{
                background-color: {colors['bg_primary']};
                border-radius: 15px;
                padding: 15px;
                margin: 5px;
            }}
        """)
        wmp_card_container_layout = QVBoxLayout(wmp_card_container)
        wmp_card_container_layout.addWidget(wpm_card)
        results_layout.addWidget(wmp_card_container)

        # Time result - secondary prominence
        time_card = QWidget()
        time_layout = QHBoxLayout(time_card)
        time_layout.setSpacing(15)

        time_icon = QLabel("⏱️")
        time_icon.setStyleSheet("font-size: 24px;")
        time_icon.setAlignment(Qt.AlignCenter)

        time_content = QWidget()
        time_content_layout = QVBoxLayout(time_content)
        time_content_layout.setSpacing(2)

        time_title = QLabel(self.localization.get_text("total_time_taken"))
        time_title.setStyleSheet(f"""
            QLabel {{
                font-size: 14px;
                font-weight: bold;
                color: {colors['text_secondary']};
            }}
        """)

        self.lbl_total_time = QLabel("---")
        self.lbl_total_time.setStyleSheet(f"""
            QLabel {{
                font-size: 22px;
                font-weight: bold;
                color: #4CAF50;
            }}
        """)

        time_content_layout.addWidget(time_title)
        time_content_layout.addWidget(self.lbl_total_time)

        time_layout.addWidget(time_icon)
        time_layout.addWidget(time_content)
        time_layout.addStretch()

        time_card.setStyleSheet(f"""
            QWidget {{
                background-color: {colors['bg_secondary']};
                border-radius: 12px;
                padding: 15px;
                border: 2px solid #4CAF50;
                margin: 5px;
            }}
        """)
        results_layout.addWidget(time_card)

        # Recommendation section - enhanced with better styling
        rec_card = QWidget()
        rec_layout = QVBoxLayout(rec_card)
        rec_layout.setSpacing(10)

        rec_title = QLabel("💡 " + self.localization.get_text("recommendation"))
        rec_title.setStyleSheet(f"""
            QLabel {{
                font-size: 16px;
                font-weight: bold;
                color: #FF9800;
                margin-bottom: 5px;
            }}
        """)

        self.lbl_recommendation = QLabel("...")
        self.lbl_recommendation.setWordWrap(True)
        self.lbl_recommendation.setTextFormat(Qt.RichText)
        self.lbl_recommendation.setStyleSheet(f"""
            QLabel {{
                font-size: 15px;
                color: {colors['text_primary']};
                padding: 20px;
                background-color: {colors['bg_primary']};
                border-radius: 12px;
                border: 2px solid #FF9800;
                line-height: 1.6;
                min-height: 100px;
            }}
        """)

        rec_layout.addWidget(rec_title)
        rec_layout.addWidget(self.lbl_recommendation)

        rec_card.setStyleSheet(f"""
            QWidget {{
                background-color: {colors['bg_secondary']};
                border-radius: 15px;
                padding: 15px;
                margin: 5px;
            }}
        """)
        results_layout.addWidget(rec_card)


        self.results_group.setVisible(False)
        right_layout.addWidget(self.results_group)

        # Comprehension test section (will be added later)
        self.comprehension_group = QGroupBox(self.localization.get_text("comprehension_test"))
        self.comprehension_layout = QVBoxLayout(self.comprehension_group)
        self.comprehension_group.setVisible(False)
        right_layout.addWidget(self.comprehension_group)

        # Sessions dashboard moved to Dashboard tab - no longer displayed here

        right_layout.addStretch()

        # Add scroll areas for better responsiveness
        left_scroll = QScrollArea()
        left_scroll.setWidget(left_widget)
        left_scroll.setWidgetResizable(True)
        left_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        left_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        left_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: transparent;
            }
        """)

        right_scroll = QScrollArea()
        right_scroll.setWidget(right_widget)
        right_scroll.setWidgetResizable(True)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: transparent;
            }
        """)

        splitter.addWidget(left_scroll)
        splitter.addWidget(right_scroll)

        # Set balanced splitter proportions (50% left, 50% right) since book info was moved
        splitter.setSizes([500, 500])
        splitter.setStretchFactor(0, 1)  # Equal space for both sides
        splitter.setStretchFactor(1, 1)  # Equal space for both sides

        tab_layout.addWidget(splitter)

        # Add the Current Book tab
        current_book_text = self.localization.get_text('current_book') if self.localization else 'Current Book'
        self.tab_widget.addTab(current_book_tab, f"📖 {current_book_text}")



    def create_prepared_books_tab(self):
        """Create the Activity tab with table-based reading sessions display"""
        activity_tab = QWidget()
        layout = QVBoxLayout(activity_tab)
        layout.setSpacing(10)
        layout.setContentsMargins(15, 15, 15, 15)

        # Header with controls
        header_layout = QHBoxLayout()

        # Title
        activity_text = self.localization.get_text("activity") if self.localization else "Activity"
        title = QLabel(f"📊 {activity_text}")
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setStyleSheet("color: #1976D2; margin-bottom: 5px;")
        header_layout.addWidget(title)

        header_layout.addStretch()

        # Activity type filter
        filter_label = QLabel(self.localization.get_text("show") if self.localization else "Show:")
        filter_label.setStyleSheet("font-weight: bold; margin-right: 5px;")
        header_layout.addWidget(filter_label)

        self.activity_filter = QComboBox()
        self.activity_filter.addItems([
            f"📚 {self.localization.get_text('all_activities') if self.localization else 'All Activities'}",
            f"📖 {self.localization.get_text('reading_measurements') if self.localization else 'Reading Measurements'}",
            f"🎯 {self.localization.get_text('training_sessions') if self.localization else 'Training Sessions'}"
        ])
        self.activity_filter.setStyleSheet("""
            QComboBox {
                padding: 5px 10px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-weight: bold;
                min-width: 150px;
            }
            QComboBox:focus {
                border-color: #1976D2;
            }
        """)
        self.activity_filter.currentTextChanged.connect(self.on_activity_filter_changed)
        header_layout.addWidget(self.activity_filter)

        # Refresh button
        refresh_text = self.localization.get_text("refresh") if self.localization else "Refresh"
        self.refresh_btn = QPushButton(f"🔄 {refresh_text}")
        self.refresh_btn.clicked.connect(self.refresh_sessions_table)
        self.refresh_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        header_layout.addWidget(self.refresh_btn)

        # Export button
        export_text = self.localization.get_text("export_all_data") if self.localization else "Export Data"
        self.export_btn = QPushButton(f"📊 {export_text}")
        self.export_btn.clicked.connect(self.export_sessions_data)
        self.export_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        header_layout.addWidget(self.export_btn)
        layout.addLayout(header_layout)

        # Description
        desc_text = self.localization.get_text("reading_sessions_desc") if self.localization else "Reading speed measurement sessions with detailed statistics:"
        desc = QLabel(desc_text)
        desc.setStyleSheet("color: #666; margin-bottom: 10px; font-size: 12px;")
        layout.addWidget(desc)

        # Create sessions table
        self.sessions_table = QTableWidget()
        self.sessions_table.setColumnCount(8)  # Added one more column for Type

        # Set table headers with localization
        headers = [
            self.localization.get_text("book_title") if self.localization else "Book Title",
            self.localization.get_text("type") if self.localization else "Type",
            self.localization.get_text("date") if self.localization else "Date",
            self.localization.get_text("progress_time") if self.localization else "Progress/Time",
            self.localization.get_text("wpm") if self.localization else "WPM",
            self.localization.get_text("settings_details") if self.localization else "Settings/Details",
            self.localization.get_text("status") if self.localization else "Status",
            self.localization.get_text("actions") if self.localization else "Actions"
        ]
        self.sessions_table.setHorizontalHeaderLabels(headers)

        # Configure table appearance
        self.sessions_table.setAlternatingRowColors(True)
        self.sessions_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.sessions_table.setSelectionMode(QTableWidget.SingleSelection)
        self.sessions_table.setSortingEnabled(True)
        self.sessions_table.verticalHeader().setVisible(False)

        # Increase row height for better readability
        self.sessions_table.verticalHeader().setDefaultSectionSize(45)

        # Set column widths
        header = self.sessions_table.horizontalHeader()
        header.setStretchLastSection(True)
        header.resizeSection(0, 250)  # Book Title (increased width)
        header.resizeSection(1, 100)  # Type
        header.resizeSection(2, 120)  # Date
        header.resizeSection(3, 100)  # Progress/Time
        header.resizeSection(4, 80)   # WPM
        header.resizeSection(5, 120)  # Settings/Details
        header.resizeSection(6, 100)  # Status

        # Apply theme-aware styling
        self.apply_table_theme()

        layout.addWidget(self.sessions_table)

        # Summary statistics panel
        self.create_summary_panel(layout)

        # Add the Activity tab
        activity_text = self.localization.get_text("activity") if self.localization else "Activity"
        self.tab_widget.addTab(activity_tab, f"📊 {activity_text}")

        # Connect tab change to activity update
        self.tab_widget.currentChanged.connect(self.on_tab_changed)

    def create_fast_reading_trainer_tab(self):
        """Create the Fast Reading Trainer tab for WPG (Words Per Glance) training"""
        trainer_tab = QWidget()
        layout = QVBoxLayout(trainer_tab)
        layout.setSpacing(15)
        layout.setContentsMargins(15, 15, 15, 15)

        # Remove header section to save space

        # Main content in splitter
        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)

        # Left side - Controls and settings
        left_widget = QWidget()
        left_widget.setMinimumWidth(300)
        left_layout = QVBoxLayout(left_widget)
        left_layout.setSpacing(15)

        # Document selection section
        doc_group = QGroupBox(self.localization.get_text("select_document") if self.localization else "Select Document")
        doc_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #1976D2;
                border: 2px solid #1976D2;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
            }
        """)
        doc_layout = QVBoxLayout(doc_group)
        doc_layout.setSpacing(12)
        doc_layout.setContentsMargins(15, 15, 15, 15)

        # File path display
        self.trainer_file_display = QLineEdit()
        self.trainer_file_display.setReadOnly(True)
        self.trainer_file_display.setPlaceholderText(self.localization.get_text("select_text_file") if self.localization else "Select text file (.txt, .docx, .pdf)")
        self.trainer_file_display.setMinimumHeight(35)
        doc_layout.addWidget(self.trainer_file_display)

        # Icon-based file selection buttons - horizontal layout
        file_buttons_layout = QHBoxLayout()
        file_buttons_layout.setSpacing(15)
        file_buttons_layout.setContentsMargins(0, 10, 0, 10)

        # Get icon paths
        icon_base_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'new_icons')

        # TXT file icon button
        self.btn_select_txt = QToolButton()
        self.btn_select_txt.clicked.connect(lambda: self.select_trainer_file("txt"))
        txt_icon_path = os.path.join(icon_base_path, 'txt_file.png')
        if os.path.exists(txt_icon_path):
            self.btn_select_txt.setIcon(QIcon(txt_icon_path))
            self.btn_select_txt.setIconSize(QSize(64, 64))
        else:
            self.btn_select_txt.setText("📄")
            self.btn_select_txt.setStyleSheet("font-size: 48px;")
        self.btn_select_txt.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
        txt_tooltip = self.localization.get_text("select_txt_tooltip") if self.localization else "Select TXT file for training"
        self.btn_select_txt.setToolTip(txt_tooltip)
        self.btn_select_txt.setMinimumSize(80, 80)
        self.btn_select_txt.setStyleSheet("""
            QToolButton {
                background-color: transparent;
                border: 2px solid #4CAF50;
                border-radius: 10px;
                padding: 5px;
            }
            QToolButton:hover {
                background-color: #E8F5E9;
                border: 2px solid #45a049;
            }
            QToolButton:pressed {
                background-color: #C8E6C9;
            }
        """)
        file_buttons_layout.addWidget(self.btn_select_txt)

        # DOCX file icon button
        self.btn_select_docx = QToolButton()
        self.btn_select_docx.clicked.connect(lambda: self.select_trainer_file("docx"))
        docx_icon_path = os.path.join(icon_base_path, 'word_file.png')
        if os.path.exists(docx_icon_path):
            self.btn_select_docx.setIcon(QIcon(docx_icon_path))
            self.btn_select_docx.setIconSize(QSize(64, 64))
        else:
            self.btn_select_docx.setText("📝")
            self.btn_select_docx.setStyleSheet("font-size: 48px;")
        self.btn_select_docx.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
        docx_tooltip = self.localization.get_text("select_docx_tooltip") if self.localization else "Select DOCX file for training"
        self.btn_select_docx.setToolTip(docx_tooltip)
        self.btn_select_docx.setMinimumSize(80, 80)
        self.btn_select_docx.setStyleSheet("""
            QToolButton {
                background-color: transparent;
                border: 2px solid #2196F3;
                border-radius: 10px;
                padding: 5px;
            }
            QToolButton:hover {
                background-color: #E3F2FD;
                border: 2px solid #1976D2;
            }
            QToolButton:pressed {
                background-color: #BBDEFB;
            }
        """)
        file_buttons_layout.addWidget(self.btn_select_docx)

        file_buttons_layout.addStretch()
        doc_layout.addLayout(file_buttons_layout)
        left_layout.addWidget(doc_group)

        # Training settings section
        settings_group = QGroupBox(self.localization.get_text("training_settings") if self.localization else "Training Settings")

        # Apply theme-aware styling for training settings
        colors = self.get_theme_colors()
        settings_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                font-size: 16px;
                color: {colors['accent_orange']};
                border: 2px solid {colors['accent_orange']};
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: {colors['bg_primary']};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: {colors['accent_orange']};
            }}
        """)
        settings_layout = QVBoxLayout(settings_group)

        # Enhanced spinbox style with theme awareness
        colors = self.get_theme_colors()
        enhanced_spinbox_style = f"""
            QSpinBox {{
                padding: 8px;
                border: 2px solid {colors['border']};
                border-radius: 6px;
                font-size: 14px;
                background: {colors['bg_primary']};
                color: {colors['text_primary']};
                min-height: 20px;
            }}
            QSpinBox:focus {{
                border-color: {colors['accent_orange']};
            }}
            QSpinBox::up-button {{
                subcontrol-origin: border;
                subcontrol-position: top right;
                width: 25px;
                border-left: 1px solid {colors['border']};
                border-bottom: 1px solid {colors['border']};
                border-top-right-radius: 4px;
                background: {colors['bg_tertiary']};
            }}
            QSpinBox::up-button:hover {{
                background: {colors['bg_secondary']};
            }}
            QSpinBox::up-button:pressed {{
                background: {colors['border']};
            }}
            QSpinBox::up-arrow {{
                image: none;
                border: 4px solid transparent;
                border-bottom: 6px solid {colors['text_muted']};
                width: 0px;
                height: 0px;
            }}
            QSpinBox::down-button {{
                subcontrol-origin: border;
                subcontrol-position: bottom right;
                width: 25px;
                border-left: 1px solid {colors['border']};
                border-top: 1px solid {colors['border']};
                border-bottom-right-radius: 4px;
                background: {colors['bg_tertiary']};
            }}
            QSpinBox::down-button:hover {{
                background: {colors['bg_secondary']};
            }}
            QSpinBox::down-button:pressed {{
                background: {colors['border']};
            }}
            QSpinBox::down-arrow {{
                image: none;
                border: 4px solid transparent;
                border-top: 6px solid {colors['text_muted']};
                width: 0px;
                height: 0px;
            }}
        """

        # Words per glance (label removed for compact UI)
        wpg_layout = QHBoxLayout()
        self.wpg_spinbox = QSpinBox()
        self.wpg_spinbox.setRange(1, 7)
        self.wpg_spinbox.setValue(3)
        self.wpg_spinbox.setSuffix(" " + (self.localization.get_text("words") if self.localization else "words"))
        self.wpg_spinbox.setStyleSheet(enhanced_spinbox_style)
        self.wpg_spinbox.setMinimumWidth(120)
        # Connect to update effective speed calculation
        self.wpg_spinbox.valueChanged.connect(self.update_effective_speed)
        wpg_layout.addWidget(self.wpg_spinbox)
        settings_layout.addLayout(wpg_layout)

        # Target speed with slider (label removed for compact UI)

        # Speed controls layout
        speed_controls_layout = QHBoxLayout()

        # Speed slider
        self.speed_slider = QSlider(Qt.Horizontal)
        self.speed_slider.setRange(100, 2000)
        self.speed_slider.setValue(300)
        self.speed_slider.setTickPosition(QSlider.TicksBelow)
        self.speed_slider.setTickInterval(200)
        self.speed_slider.setStyleSheet("""
            QSlider::groove:horizontal {
                border: 1px solid #999999;
                height: 10px;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #B1B1B1, stop:1 #c4c4c4);
                margin: 2px 0;
                border-radius: 5px;
            }
            QSlider::handle:horizontal {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #FF9800, stop:1 #F57C00);
                border: 2px solid #E65100;
                width: 20px;
                margin: -5px 0;
                border-radius: 10px;
            }
            QSlider::handle:horizontal:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #FFB74D, stop:1 #FF9800);
            }
            QSlider::sub-page:horizontal {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #FF9800, stop:1 #F57C00);
                border: 1px solid #E65100;
                height: 10px;
                border-radius: 5px;
            }
        """)

        # Speed spinbox - reduced width to show full label
        self.trainer_wpm_spinbox = QSpinBox()
        self.trainer_wpm_spinbox.setRange(100, 2000)
        self.trainer_wpm_spinbox.setValue(300)
        # Use shorter suffix to fit in smaller width
        self.trainer_wpm_spinbox.setSuffix(" WPM")
        self.trainer_wpm_spinbox.setStyleSheet(enhanced_spinbox_style)
        self.trainer_wpm_spinbox.setMinimumWidth(100)  # Reduced from 120 to 100
        self.trainer_wpm_spinbox.setMaximumWidth(140)  # Add maximum width constraint

        # Connect slider and spinbox
        self.speed_slider.valueChanged.connect(self.trainer_wpm_spinbox.setValue)
        self.trainer_wpm_spinbox.valueChanged.connect(self.speed_slider.setValue)
        self.trainer_wpm_spinbox.valueChanged.connect(self.update_effective_speed)

        speed_controls_layout.addWidget(self.speed_slider, 3)
        speed_controls_layout.addWidget(self.trainer_wpm_spinbox, 1)
        settings_layout.addLayout(speed_controls_layout)


        # Font family (label removed for compact UI)
        font_family_layout = QHBoxLayout()
        self.font_family_combo = QComboBox()
        self.font_family_combo.addItems([
            "IBM Plex Sans Arabic",
            "Tajawal",
            "Courier New",
            "Arial",
            "Consolas",
            "Times New Roman",
            "Verdana"
        ])
        self.font_family_combo.setCurrentText("IBM Plex Sans Arabic")
        self.font_family_combo.setStyleSheet(f"""
            QComboBox {{
                padding: 8px;
                border: 2px solid {colors['border']};
                border-radius: 6px;
                font-size: 14px;
                background: {colors['bg_primary']};
                color: {colors['text_primary']};
                min-height: 20px;
                combobox-popup: 0;
            }}
            QComboBox:focus {{
                border-color: #FF9800;
            }}
            QComboBox::drop-down {{
                border: none;
                width: 25px;
                background: {colors['bg_tertiary']};
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
            }}
            QComboBox::drop-down:hover {{
                background: {colors['bg_secondary']};
            }}
            QComboBox::down-arrow {{
                image: none;
                border: 4px solid transparent;
                border-top: 6px solid {colors['text_muted']};
                width: 0px;
                height: 0px;
            }}
            QComboBox QAbstractItemView {{
                background-color: {colors['bg_primary']};
                color: {colors['text_primary']};
                selection-background-color: {colors['accent_blue']};
            }}
        """)
        self.font_family_combo.setMinimumWidth(150)
        font_family_layout.addWidget(self.font_family_combo)
        settings_layout.addLayout(font_family_layout)

        # Start page configuration (label removed for compact UI)
        start_page_layout = QHBoxLayout()
        self.start_page_spinbox = QSpinBox()
        self.start_page_spinbox.setRange(1, 9999)
        self.start_page_spinbox.setValue(1)
        self.start_page_spinbox.setStyleSheet(enhanced_spinbox_style)
        self.start_page_spinbox.setMinimumWidth(120)
        self.start_page_spinbox.setToolTip(self.localization.get_text("start_page_tooltip") if self.localization else "Page number to start training from (useful to skip introductions)")
        start_page_layout.addWidget(self.start_page_spinbox)
        settings_layout.addLayout(start_page_layout)

        # Font size (label removed for compact UI)
        font_layout = QHBoxLayout()
        self.font_size_spinbox = QSpinBox()
        self.font_size_spinbox.setRange(16, 72)
        self.font_size_spinbox.setValue(32)
        self.font_size_spinbox.setSuffix("px")
        self.font_size_spinbox.setStyleSheet(enhanced_spinbox_style)
        self.font_size_spinbox.setMinimumWidth(120)
        font_layout.addWidget(self.font_size_spinbox)
        settings_layout.addLayout(font_layout)

        left_layout.addWidget(settings_group)

        # Training Modes Selector Section
        self.create_training_modes_selector(left_layout)

        # Control buttons section - Compact icon-only buttons in one row
        controls_group = QGroupBox(self.localization.get_text("training_controls") if self.localization else "Training Controls")
        colors = self.get_theme_colors()
        controls_group.setStyleSheet(f"""
            QGroupBox {{
                font-weight: bold;
                font-size: 16px;
                color: {colors['accent_green']};
                border: 2px solid {colors['accent_green']};
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: {colors['bg_primary']};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: {colors['accent_green']};
            }}
        """)
        controls_layout = QVBoxLayout(controls_group)

        # Icon buttons row - 3 buttons in horizontal layout
        icon_buttons_layout = QHBoxLayout()
        icon_buttons_layout.setSpacing(10)

        # Continue Previous Session icon button
        self.btn_continue_session = QToolButton()
        self.btn_continue_session.setText("📂")
        self.btn_continue_session.setToolTip(self.localization.get_text("continue_previous_session") if self.localization else "Continue Previous Session")
        self.btn_continue_session.clicked.connect(lambda: self.show_saved_sessions_dialog())
        self.btn_continue_session.setFixedSize(60, 60)
        self.btn_continue_session.setStyleSheet(f"""
            QToolButton {{
                background-color: {colors['accent_blue']};
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 28px;
            }}
            QToolButton:hover {{
                background-color: #1976D2;
            }}
            QToolButton:pressed {{
                background-color: #0D47A1;
            }}
        """)
        icon_buttons_layout.addWidget(self.btn_continue_session)

        # Unified Training Control icon button (Smart Button)
        # This button changes its appearance and function based on training state:
        # - Not Started: ▶️ Start (green)
        # - Running: ⏸️ Pause (orange)
        # - Paused: ▶️ Resume (blue)
        self.btn_unified_control = QToolButton()
        self.btn_unified_control.setText("▶️")
        self.btn_unified_control.setToolTip(self.localization.get_text("start_training") if self.localization else "Start Training")
        self.btn_unified_control.clicked.connect(self.handle_unified_control_click)
        self.btn_unified_control.setEnabled(False)
        self.btn_unified_control.setProperty("control_state", "start")  # Track current state
        self.btn_unified_control.setFixedSize(60, 60)
        self.btn_unified_control.setStyleSheet(f"""
            QToolButton {{
                background-color: {colors['accent_green']};
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 28px;
            }}
            QToolButton:hover {{
                background-color: #45a049;
            }}
            QToolButton:disabled {{
                background-color: {colors['border']};
            }}
            QToolButton:pressed {{
                background-color: #2E7D32;
            }}
        """)
        icon_buttons_layout.addWidget(self.btn_unified_control)

        # Stop training icon button (separate for safety)
        self.btn_stop_training = QToolButton()
        self.btn_stop_training.setText("⏹️")
        self.btn_stop_training.setToolTip(self.localization.get_text("stop") if self.localization else "Stop Training")
        self.btn_stop_training.clicked.connect(lambda: self.stop_fast_reading_training())
        self.btn_stop_training.setEnabled(False)
        self.btn_stop_training.setFixedSize(60, 60)
        self.btn_stop_training.setStyleSheet(f"""
            QToolButton {{
                background-color: {colors['accent_red']};
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 28px;
            }}
            QToolButton:hover {{
                background-color: #d32f2f;
            }}
            QToolButton:disabled {{
                background-color: {colors['border']};
            }}
            QToolButton:pressed {{
                background-color: #B71C1C;
            }}
        """)
        icon_buttons_layout.addWidget(self.btn_stop_training)

        # Add stretch to push buttons to the left
        icon_buttons_layout.addStretch()

        controls_layout.addLayout(icon_buttons_layout)

        # Keep references to old buttons for backward compatibility
        self.btn_start_training = self.btn_unified_control
        self.btn_pause_training = self.btn_unified_control

        left_layout.addWidget(controls_group)

        # Progress section
        progress_group = QGroupBox(self.localization.get_text("progress") if self.localization else "Progress")
        progress_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #9C27B0;
                border: 2px solid #9C27B0;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
            }
        """)
        progress_layout = QVBoxLayout(progress_group)

        # Progress bar
        self.training_progress = QProgressBar()
        self.training_progress.setStyleSheet("""
            QProgressBar {
                border: 2px solid #9C27B0;
                border-radius: 5px;
                text-align: center;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #9C27B0;
                border-radius: 3px;
            }
        """)
        progress_layout.addWidget(self.training_progress)

        # Progress info
        self.progress_info = QLabel(self.localization.get_text("ready_to_start") if self.localization else "Ready to start training")
        self.progress_info.setAlignment(Qt.AlignCenter)
        self.progress_info.setStyleSheet("QLabel { color: #666; font-size: 14px; margin-top: 5px; }")
        progress_layout.addWidget(self.progress_info)

        left_layout.addWidget(progress_group)
        left_layout.addStretch()

        # Right side - Training display
        right_widget = QWidget()
        right_widget.setMinimumWidth(400)
        right_layout = QVBoxLayout(right_widget)
        right_layout.setSpacing(0)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # Training display area
        self.training_display = QFrame()
        self.training_display.setFrameStyle(QFrame.Box)
        self.training_display.setStyleSheet("""
            QFrame {
                border: 3px solid #333;
                border-radius: 15px;
                background-color: #000;
                min-height: 400px;
            }
        """)
        display_layout = QVBoxLayout(self.training_display)
        display_layout.setAlignment(Qt.AlignCenter)

        # Word display label
        self.word_display = QLabel(self.localization.get_text("select_document_to_start") if self.localization else "Select a document to start training")
        self.word_display.setAlignment(Qt.AlignCenter)
        self.word_display.setTextFormat(Qt.RichText)  # Enable rich text for HTML formatting
        self.word_display.setStyleSheet("""
            QLabel {
                color: #00FF00;
                font-size: 32px;
                font-weight: bold;
                font-family: 'Courier New', monospace;
                padding: 20px;
                background-color: transparent;
                border: none;
            }
        """)
        self.word_display.setWordWrap(True)
        display_layout.addWidget(self.word_display)

        right_layout.addWidget(self.training_display)

        # Add widgets to splitter
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)  # Left side gets 1 part
        splitter.setStretchFactor(1, 2)  # Right side gets 2 parts

        layout.addWidget(splitter)

        # Initialize training variables
        self.trainer_text = ""
        self.trainer_words = []
        self.current_word_index = 0
        self.training_timer = QTimer()
        self.training_timer.timeout.connect(lambda: self.display_next_word_group())
        self.is_training_paused = False
        self.training_active = False

        # Add the Fast Reading Trainer tab
        trainer_text = self.localization.get_text("fast_reading_trainer") if self.localization else "Fast Reading Trainer"
        self.tab_widget.addTab(trainer_tab, f"🧠 {trainer_text}")

    def create_summary_panel(self, layout):
        """Create summary statistics panel"""
        summary_frame = QFrame()
        summary_frame.setFrameStyle(QFrame.Box)
        summary_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ddd;
                border-radius: 8px;
                background-color: #f9f9f9;
                padding: 10px;
                margin-top: 10px;
            }
        """)

        summary_layout = QHBoxLayout(summary_frame)

        # Total sessions
        total_text = self.localization.get_text("total_sessions") if self.localization else "Total Sessions"
        self.total_sessions_label = QLabel(f"{total_text}: 0")
        self.total_sessions_label.setStyleSheet("font-weight: bold; color: #1976D2;")
        summary_layout.addWidget(self.total_sessions_label)

        summary_layout.addStretch()

        # Average WPM
        avg_wpm_text = self.localization.get_text("average_wpm") if self.localization else "Average WPM"
        self.avg_wpm_label = QLabel(f"{avg_wpm_text}: 0")
        self.avg_wpm_label.setStyleSheet("font-weight: bold; color: #4CAF50;")
        summary_layout.addWidget(self.avg_wpm_label)

        summary_layout.addStretch()

        # Best WPM
        best_wpm_text = self.localization.get_text("best_wpm") if self.localization else "Best WPM"
        self.best_wpm_label = QLabel(f"{best_wpm_text}: 0")
        self.best_wpm_label.setStyleSheet("font-weight: bold; color: #FF9800;")
        summary_layout.addWidget(self.best_wpm_label)

        summary_layout.addStretch()

        # Total reading time
        total_time_text = self.localization.get_text("total_reading_time") if self.localization else "Total Time"
        self.total_time_label = QLabel(f"{total_time_text}: 0h 0m")
        self.total_time_label.setStyleSheet("font-weight: bold; color: #9C27B0;")
        summary_layout.addWidget(self.total_time_label)

        layout.addWidget(summary_frame)

    def apply_table_theme(self):
        """Apply theme-aware styling to the sessions table"""
        colors = self.get_theme_colors()

        self.sessions_table.setStyleSheet(f"""
            QTableWidget {{
                background-color: {colors.get('card_bg', '#ffffff')};
                alternate-background-color: {colors.get('card_hover_bg', '#f5f5f5')};
                color: {colors.get('text_color', '#333333')};
                gridline-color: {colors.get('border_color', '#e0e0e0')};
                border: 1px solid {colors.get('border_color', '#e0e0e0')};
                border-radius: 8px;
                font-size: 12px;
            }}
            QTableWidget::item {{
                padding: 12px 8px;
                border: none;
                min-height: 25px;
            }}
            QTableWidget::item:selected {{
                background-color: {colors.get('accent_blue', '#1976D2')};
                color: white;
            }}
            QHeaderView::section {{
                background-color: {colors.get('header_bg', '#f0f0f0')};
                color: {colors.get('header_text', '#333333')};
                padding: 8px;
                border: none;
                border-bottom: 2px solid {colors.get('accent_blue', '#1976D2')};
                font-weight: bold;
            }}
        """)

    def refresh_sessions_table(self):
        """Refresh the reading sessions table with latest data"""
        if not self.recent_books_manager:
            return

        try:
            # Get all reading sessions
            sessions = self.recent_books_manager.get_all_reading_sessions()

            # Clear existing data
            self.sessions_table.setRowCount(0)

            if not sessions:
                # Show empty state
                self.sessions_table.setRowCount(1)
                no_data_text = self.localization.get_text("no_sessions_yet") if self.localization else "No reading sessions yet"
                item = QTableWidgetItem(no_data_text)
                item.setTextAlignment(Qt.AlignCenter)
                self.sessions_table.setItem(0, 0, item)
                self.sessions_table.setSpan(0, 0, 1, 7)
                self.update_summary_stats([])
                return

            # Populate table with session data
            self.sessions_table.setRowCount(len(sessions))

            for row, session in enumerate(sessions):
                # Book name/title (column 0 - increased width)
                book_name = os.path.basename(session.get('file_path', 'Unknown'))
                if book_name.endswith('.pdf'):
                    book_name = book_name[:-4]
                # Don't truncate as much since we have more width now
                if len(book_name) > 50:
                    book_name = book_name[:47] + "..."
                self.sessions_table.setItem(row, 0, QTableWidgetItem(book_name))

                # Type (column 1) - Reading or Training
                session_type = session.get('session_type', 'reading')
                if session_type == 'training':
                    type_text = self.localization.get_text('training') if self.localization else 'Training'
                    type_icon = "🎯"
                else:
                    type_text = self.localization.get_text('reading') if self.localization else 'Reading'
                    type_icon = "📖"
                self.sessions_table.setItem(row, 1, QTableWidgetItem(f"{type_icon} {type_text}"))

                # Date (column 2)
                date_str = session.get('timestamp', 'Unknown')
                if isinstance(date_str, str) and 'T' in date_str:
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                        date_str = dt.strftime("%Y-%m-%d %H:%M")
                    except:
                        pass
                self.sessions_table.setItem(row, 2, QTableWidgetItem(str(date_str)))

                # Progress/Time (column 3) - Combined pages and time
                pages = session.get('pages_read', 0)
                time_seconds = session.get('time_spent_seconds', 0)
                time_str = self.format_time_duration(time_seconds)
                progress_text = f"{pages} صفحة / {time_str}" if self.localization and self.localization.get_text('pages') else f"{pages} pages / {time_str}"
                self.sessions_table.setItem(row, 3, QTableWidgetItem(progress_text))

                # WPM (column 4)
                wpm = session.get('wpm', 0)
                self.sessions_table.setItem(row, 4, QTableWidgetItem(f"{wpm:.0f}"))

                # Settings/Details (column 5)
                words_per_page = session.get('avg_words_per_page', 0)
                wpg = session.get('wpg_setting', 3)
                settings_text = f"WPP: {words_per_page:.0f}, WPG: {wpg}"
                self.sessions_table.setItem(row, 5, QTableWidgetItem(settings_text))

                # Status (column 6) - Efficiency
                efficiency = self.calculate_efficiency(wpm)
                self.sessions_table.setItem(row, 6, QTableWidgetItem(efficiency))

                # Actions (column 7) - Resume and Delete buttons
                actions_widget = QWidget()
                actions_layout = QHBoxLayout(actions_widget)
                actions_layout.setContentsMargins(4, 4, 4, 4)
                actions_layout.setSpacing(4)

                # Resume button
                resume_btn = QPushButton("▶️")
                resume_btn.setToolTip(self.localization.get_text('resume') if self.localization else 'Resume')
                resume_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #4CAF50;
                        color: white;
                        border: none;
                        padding: 4px 8px;
                        border-radius: 4px;
                        font-weight: bold;
                        min-width: 30px;
                    }
                    QPushButton:hover {
                        background-color: #45a049;
                    }
                """)
                resume_btn.setProperty('session_data', session)
                resume_btn.clicked.connect(lambda checked, s=session: self.resume_session(s))

                # Delete button
                delete_btn = QPushButton("🗑️")
                delete_btn.setToolTip(self.localization.get_text('delete') if self.localization else 'Delete')
                delete_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #f44336;
                        color: white;
                        border: none;
                        padding: 4px 8px;
                        border-radius: 4px;
                        font-weight: bold;
                        min-width: 30px;
                    }
                    QPushButton:hover {
                        background-color: #d32f2f;
                    }
                """)
                delete_btn.setProperty('session_data', session)
                delete_btn.clicked.connect(lambda checked, s=session: self.delete_session(s))

                actions_layout.addWidget(resume_btn)
                actions_layout.addWidget(delete_btn)
                actions_layout.addStretch()

                self.sessions_table.setCellWidget(row, 7, actions_widget)

            # Update summary statistics
            self.update_summary_stats(sessions)

            # Sort by date (most recent first) - column 2 is now the date column
            self.sessions_table.sortItems(2, Qt.DescendingOrder)

        except Exception as e:
            print(f"Error refreshing sessions table: {e}")

    def delete_session(self, session_data):
        """Delete a reading session from the activity table"""
        try:
            # Show confirmation dialog
            reply = QMessageBox.question(
                self,
                self.localization.get_text("confirm_delete") if self.localization else "Confirm Delete",
                f"{self.localization.get_text('delete_session_confirm') if self.localization else 'Are you sure you want to delete this session?'}\n\n"
                f"{self.localization.get_text('book') if self.localization else 'Book'}: {os.path.basename(session_data.get('file_path', 'Unknown'))}\n"
                f"{self.localization.get_text('date') if self.localization else 'Date'}: {session_data.get('timestamp', 'Unknown')}",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                # Delete the session using the recent books manager
                if self.recent_books_manager:
                    success = self.recent_books_manager.delete_reading_session(
                        session_data.get('file_path', ''),
                        session_data.get('session_id', '')
                    )

                    if success:
                        # Refresh the table to show updated data
                        self.refresh_sessions_table()

                        # Show success message
                        QMessageBox.information(
                            self,
                            self.localization.get_text("success") if self.localization else "Success",
                            self.localization.get_text("session_deleted") if self.localization else "Session deleted successfully."
                        )
                    else:
                        # Show error message
                        QMessageBox.warning(
                            self,
                            self.localization.get_text("error") if self.localization else "Error",
                            self.localization.get_text("delete_session_error") if self.localization else "Failed to delete session."
                        )

        except Exception as e:
            print(f"Error deleting session: {e}")
            QMessageBox.critical(
                self,
                self.localization.get_text("error") if self.localization else "Error",
                f"An error occurred while deleting the session: {str(e)}"
            )

    def on_activity_filter_changed(self):
        """Handle activity filter change"""
        self.refresh_sessions_table()

    def format_time_duration(self, seconds):
        """Format time duration in a readable format"""
        if seconds < 60:
            return f"{seconds}s"
        elif seconds < 3600:
            minutes = seconds // 60
            secs = seconds % 60
            return f"{minutes}m {secs}s"
        else:
            hours = seconds // 3600
            minutes = (seconds % 3600) // 60
            return f"{hours}h {minutes}m"

    def calculate_efficiency(self, wpm):
        """Calculate reading efficiency relative to average"""
        if wpm < 150:
            return "🐌 Slow"
        elif wpm < 250:
            return "📖 Average"
        elif wpm < 350:
            return "🚀 Good"
        elif wpm < 450:
            return "⭐ Excellent"
        else:
            return "🏆 Expert"

    def resume_session(self, session):
        """Resume a reading or training session"""
        try:
            session_type = session.get('session_type', 'reading')
            file_path = session.get('file_path', '')

            if not file_path or not os.path.exists(file_path):
                QMessageBox.warning(
                    self,
                    self.localization.get_text('error') if self.localization else 'Error',
                    self.localization.get_text('file_not_found') if self.localization else 'File not found. It may have been moved or deleted.'
                )
                return

            if session_type == 'training':
                # Resume training session
                self.resume_training_session(session)
            else:
                # Resume reading session
                self.resume_reading_session(session)

        except Exception as e:
            print(f"Error resuming session: {e}")
            QMessageBox.critical(
                self,
                self.localization.get_text('error') if self.localization else 'Error',
                f"Error resuming session: {str(e)}"
            )

    def resume_training_session(self, session):
        """Resume a training session from where it was left off"""
        try:
            # Switch to Fast Reading Trainer tab (index 1)
            self.tab_widget.setCurrentIndex(1)

            # Load the file
            file_path = session.get('file_path', '')
            if file_path and os.path.exists(file_path):
                # Load the PDF text
                self.load_trainer_file(file_path)

                # Restore settings
                if 'wpm_setting' in session:
                    self.trainer_wpm_spinbox.setValue(session['wpm_setting'])
                if 'wpg_setting' in session:
                    self.wpg_spinbox.setValue(session['wpg_setting'])
                if 'font_size' in session:
                    self.font_size_spinbox.setValue(session['font_size'])
                if 'font_family' in session:
                    index = self.font_family_combo.findText(session['font_family'])
                    if index >= 0:
                        self.font_family_combo.setCurrentIndex(index)

                # Restore progress
                if 'current_word_index' in session:
                    self.current_word_index = session['current_word_index']

                QMessageBox.information(
                    self,
                    self.localization.get_text('success') if self.localization else 'Success',
                    self.localization.get_text('session_resumed') if self.localization else 'Training session resumed! Click Start to continue.'
                )
        except Exception as e:
            print(f"Error resuming training session: {e}")
            raise

    def resume_reading_session(self, session):
        """Resume a reading session from the last page read"""
        try:
            # Switch to Current Book tab (index 0)
            self.tab_widget.setCurrentIndex(0)

            # Load the PDF
            file_path = session.get('file_path', '')
            if file_path and os.path.exists(file_path):
                self.load_pdf(file_path)

                # Restore last page if available
                if 'last_page' in session:
                    # Note: This would require PDF viewer integration
                    # For now, just notify the user
                    last_page = session['last_page']
                    QMessageBox.information(
                        self,
                        self.localization.get_text('success') if self.localization else 'Success',
                        f"{self.localization.get_text('session_resumed') if self.localization else 'Reading session resumed!'}\n{self.localization.get_text('last_page') if self.localization else 'Last page'}: {last_page}"
                    )
                else:
                    QMessageBox.information(
                        self,
                        self.localization.get_text('success') if self.localization else 'Success',
                        self.localization.get_text('session_resumed') if self.localization else 'Reading session resumed!'
                    )
        except Exception as e:
            print(f"Error resuming reading session: {e}")
            raise

    def update_summary_stats(self, sessions):
        """Update summary statistics labels"""
        hours_text = self.localization.get_text("hours") if self.localization else "hours"
        minutes_text = self.localization.get_text("minutes") if self.localization else "minutes"

        if not sessions:
            self.total_sessions_label.setText(f"{self.localization.get_text('total_sessions') if self.localization else 'Total Sessions'}: 0")
            self.avg_wpm_label.setText(f"{self.localization.get_text('average_wpm') if self.localization else 'Average WPM'}: 0")
            self.best_wpm_label.setText(f"{self.localization.get_text('best_wpm') if self.localization else 'Best WPM'}: 0")
            self.total_time_label.setText(f"{self.localization.get_text('total_reading_time') if self.localization else 'Total Time'}: 0 {hours_text} 0 {minutes_text}")
            return

        # Calculate statistics
        total_sessions = len(sessions)
        total_time = sum(session.get('time_spent_seconds', 0) for session in sessions)
        wpms = [session.get('wpm', 0) for session in sessions if session.get('wpm', 0) > 0]

        avg_wpm = sum(wpms) / len(wpms) if wpms else 0
        best_wpm = max(wpms) if wpms else 0

        # Format total time
        hours = total_time // 3600
        minutes = (total_time % 3600) // 60

        # Get localized time units
        hours_text = self.localization.get_text('hours') if self.localization else 'hours'
        minutes_text = self.localization.get_text('minutes') if self.localization else 'minutes'

        # Update labels
        self.total_sessions_label.setText(f"{self.localization.get_text('total_sessions') if self.localization else 'Total Sessions'}: {total_sessions}")
        self.avg_wpm_label.setText(f"{self.localization.get_text('average_wpm') if self.localization else 'Average WPM'}: {avg_wpm:.0f}")
        self.best_wpm_label.setText(f"{self.localization.get_text('best_wpm') if self.localization else 'Best WPM'}: {best_wpm:.0f}")
        self.total_time_label.setText(f"{self.localization.get_text('total_reading_time') if self.localization else 'Total Time'}: {hours} {hours_text} {minutes} {minutes_text}")

    def export_sessions_data(self):
        """Export reading sessions data to CSV"""
        if not self.recent_books_manager:
            return

        try:
            from PySide6.QtWidgets import QFileDialog
            import csv

            # Get save location
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                self.localization.get_text("export_sessions") if self.localization else "Export Reading Sessions",
                "reading_sessions.csv",
                "CSV Files (*.csv)"
            )

            if not file_path:
                return

            # Get sessions data
            sessions = self.recent_books_manager.get_all_reading_sessions()

            # Write to CSV
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)

                # Write headers
                headers = [
                    "Book Name", "File Path", "Date", "Pages Read",
                    "Time Spent (seconds)", "WPM", "Words per Page", "Efficiency"
                ]
                writer.writerow(headers)

                # Write data
                for session in sessions:
                    book_name = os.path.basename(session.get('file_path', 'Unknown'))
                    if book_name.endswith('.pdf'):
                        book_name = book_name[:-4]

                    wpm = session.get('wpm', 0)
                    efficiency = self.calculate_efficiency(wpm)

                    row = [
                        book_name,
                        session.get('file_path', ''),
                        session.get('timestamp', ''),
                        session.get('pages_read', 0),
                        session.get('time_spent_seconds', 0),
                        f"{wpm:.0f}",
                        f"{session.get('avg_words_per_page', 0):.0f}",
                        efficiency
                    ]
                    writer.writerow(row)

            # Show success message
            QMessageBox.information(
                self,
                self.localization.get_text("success") if self.localization else "Success",
                f"{self.localization.get_text('data_exported') if self.localization else 'Data exported successfully to'}: {file_path}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error") if self.localization else "Error",
                f"Failed to export data: {str(e)}"
            )
        self.refresh_activity_books()

    def resizeEvent(self, event):
        """Handle resize events to refresh activity grid layout"""
        super().resizeEvent(event)
        # Refresh activity books layout when window is resized
        if hasattr(self, 'activity_grid_layout') and hasattr(self, 'current_sort_method'):
            QTimer.singleShot(100, self.refresh_activity_books)

    def refresh_activity_books(self):
        """Refresh the grid of activity books"""
        if not self.recent_books_manager:
            return

        # Clear existing widgets
        for i in reversed(range(self.activity_grid_layout.count())):
            child = self.activity_grid_layout.itemAt(i).widget()
            if child:
                child.setParent(None)

        try:
            prepared_books = self.recent_books_manager.get_all_prepared_books()

            # Sort books based on current sort method
            sorted_books = self.sort_activity_books(prepared_books)

            if not sorted_books:
                # Show empty state
                empty_widget = self.create_empty_activity_state()
                self.activity_grid_layout.addWidget(empty_widget, 0, 0, 1, 3)  # Span 3 columns
                return

            # Calculate responsive grid layout
            available_width = self.activity_scroll.width() - 40
            card_width = 320  # Maximum card width
            min_columns = 1
            max_columns = 4

            # Calculate optimal number of columns based on available width
            optimal_columns = max(min_columns, min(max_columns, available_width // (card_width + 12)))

            # Create cards for each prepared book
            for i, book_data in enumerate(sorted_books):
                row = i // optimal_columns
                col = i % optimal_columns

                card = self.create_activity_book_card(book_data)
                self.activity_grid_layout.addWidget(card, row, col)

            # Add stretch to remaining space
            last_row = (len(sorted_books) - 1) // optimal_columns if sorted_books else 0
            self.activity_grid_layout.setRowStretch(last_row + 1, 1)
            self.activity_grid_layout.setColumnStretch(optimal_columns, 1)

        except Exception as e:
            print(f"Error refreshing activity books: {e}")
            error_widget = QLabel(f"❌ Error loading activity books: {str(e)}")
            error_widget.setStyleSheet("color: red; padding: 20px; font-size: 14px;")
            self.activity_grid_layout.addWidget(error_widget, 0, 0, 1, 3)

    def sort_activity_books(self, books):
        """Sort books based on current sort method"""
        if not books:
            return books

        try:
            if self.current_sort_method == "Recent Activity":
                # Sort by last session date (most recent first)
                return sorted(books, key=lambda b: b.get('statistics', {}).get('last_session_date', ''), reverse=True)

            elif self.current_sort_method == "Best Speed":
                # Sort by best WPM (highest first)
                return sorted(books, key=lambda b: b.get('statistics', {}).get('best_wpm', 0), reverse=True)

            elif self.current_sort_method == "Most Sessions":
                # Sort by total sessions (most first)
                return sorted(books, key=lambda b: b.get('statistics', {}).get('total_sessions', 0), reverse=True)

            elif self.current_sort_method == "Alphabetical":
                # Sort by display name alphabetically
                return sorted(books, key=lambda b: b.get('display_name', '').lower())

            elif self.current_sort_method == "Date Added":
                # Sort by preparation date (most recent first)
                return sorted(books, key=lambda b: b.get('created_at', ''), reverse=True)

            else:
                return books

        except Exception as e:
            print(f"Error sorting activity books: {e}")
            return books

    def on_sort_changed(self, sort_method):
        """Handle sort method change"""
        self.current_sort_method = sort_method
        self.refresh_activity_books()

    def create_empty_activity_state(self):
        """Create empty state widget for activity tab"""
        empty_widget = QWidget()
        empty_layout = QVBoxLayout(empty_widget)
        empty_layout.setAlignment(Qt.AlignCenter)

        empty_label = QLabel("📭 No Activity Found")
        empty_label.setAlignment(Qt.AlignCenter)
        empty_label.setFont(QFont("Arial", 16, QFont.Bold))
        empty_label.setStyleSheet("color: #999; margin-bottom: 10px;")
        empty_layout.addWidget(empty_label)

        desc_label = QLabel(self.localization.get_text("books_placeholder_text"))
        desc_label.setAlignment(Qt.AlignCenter)
        desc_label.setStyleSheet("color: #666; font-size: 14px; line-height: 1.4;")
        empty_layout.addWidget(desc_label)

        empty_widget.setStyleSheet("""
            QWidget {
                border: 2px dashed #ccc;
                border-radius: 12px;
                background-color: #f9f9f9;
                padding: 40px;
                margin: 20px;
            }
        """)

        return empty_widget

    def create_activity_book_card(self, book_data):
        """Create a card widget for an activity book (similar to Recent Books design)"""
        card = QFrame()
        card.setObjectName("ActivityBookCard")
        card.setMinimumSize(280, 420)
        card.setMaximumSize(320, 460)
        card.setStyleSheet(self.get_activity_card_style())

        # Remove card-level click handler - use buttons instead
        # card.mousePressEvent = lambda event: self.open_activity_book(book_data['file_path'])
        # card.setCursor(Qt.PointingHandCursor)

        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        # Cover image section (dominant element)
        cover_container = QWidget()
        cover_layout = QHBoxLayout(cover_container)
        cover_layout.setContentsMargins(8, 4, 8, 4)

        cover_label = QLabel()
        cover_label.setAlignment(Qt.AlignCenter)
        cover_label.setFixedSize(200, 280)  # Similar to Recent Books medium size
        cover_label.setStyleSheet("""
            QLabel {
                background-color: #2A2A2A;
                border: 1px solid #3A3A3A;
                border-radius: 8px;
            }
        """)

        # Load cover image
        cover_pixmap = self.load_activity_book_cover(book_data['file_path'])
        if cover_pixmap:
            cover_label.setPixmap(cover_pixmap)
        else:
            # Default cover with book icon
            cover_label.setText("📖")
            cover_label.setFont(QFont("Arial", 48))
            cover_label.setStyleSheet(cover_label.styleSheet() + "color: #666;")

        cover_layout.addWidget(cover_label, alignment=Qt.AlignCenter)
        layout.addWidget(cover_container)

        # Book title
        title = book_data.get('display_name', os.path.basename(book_data['file_path']).replace('.pdf', ''))
        title_label = QLabel(title)
        title_label.setWordWrap(True)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setFont(QFont("Arial", 12, QFont.Bold))
        title_label.setStyleSheet("color: #FFFFFF; margin: 5px 0;")
        title_label.setMaximumHeight(40)
        layout.addWidget(title_label)

        # Essential metrics in compact format
        metrics_widget = QWidget()
        metrics_layout = QVBoxLayout(metrics_widget)
        metrics_layout.setSpacing(3)
        metrics_layout.setContentsMargins(5, 0, 5, 0)

        # Configuration info
        config_text = f"📄 {book_data['total_pages']} pages • 📝 {book_data['avg_words_per_page']:.0f} WPM"
        config_label = QLabel(config_text)
        config_label.setStyleSheet("color: #CCCCCC; font-size: 10px;")
        config_label.setAlignment(Qt.AlignCenter)
        metrics_layout.addWidget(config_label)

        # Session statistics
        stats = book_data.get('statistics', {})
        if stats and stats.get('total_sessions', 0) > 0:
            sessions_text = f"📊 {stats['total_sessions']} sessions"
            sessions_label = QLabel(sessions_text)
            sessions_label.setStyleSheet("color: #6FB1FF; font-size: 11px; font-weight: bold;")
            sessions_label.setAlignment(Qt.AlignCenter)
            metrics_layout.addWidget(sessions_label)

            # Last session info
            last_session_date = stats.get('last_session_date', '')
            if last_session_date:
                try:
                    from datetime import datetime
                    date_obj = datetime.fromisoformat(last_session_date.replace('Z', '+00:00'))
                    date_str = date_obj.strftime('%m/%d')
                    last_text = f"🕒 Last: {date_str}"
                except:
                    last_text = "🕒 Recent activity"
            else:
                last_text = "🕒 No recent activity"

            last_label = QLabel(last_text)
            last_label.setStyleSheet("color: #9AD27C; font-size: 10px;")
            last_label.setAlignment(Qt.AlignCenter)
            metrics_layout.addWidget(last_label)

            # Speed metrics
            last_wpm = stats.get('last_session_wpm', 0)
            best_wpm = stats.get('best_wpm', 0)

            if last_wpm > 0:
                speed_text = f"⚡ {last_wpm:.0f} WPM"
                if best_wpm > last_wpm:
                    speed_text += f" • 🏆 {best_wpm:.0f}"
            elif best_wpm > 0:
                speed_text = f"🏆 Best: {best_wpm:.0f} WPM"
            else:
                speed_text = "⚡ No speed data"

            speed_label = QLabel(speed_text)
            speed_label.setStyleSheet("color: #FFD700; font-size: 10px; font-weight: bold;")
            speed_label.setAlignment(Qt.AlignCenter)
            metrics_layout.addWidget(speed_label)
        else:
            no_sessions_label = QLabel("📊 No sessions yet")
            no_sessions_label.setStyleSheet("color: #999; font-style: italic; font-size: 11px;")
            no_sessions_label.setAlignment(Qt.AlignCenter)
            metrics_layout.addWidget(no_sessions_label)

        layout.addWidget(metrics_widget)

        # Action buttons container
        buttons_widget = QWidget()
        buttons_layout = QHBoxLayout(buttons_widget)
        buttons_layout.setSpacing(6)
        buttons_layout.setContentsMargins(0, 5, 0, 0)

        # Open in Current Book tab button
        open_btn = QPushButton("📖 Read")
        open_btn.clicked.connect(lambda: self.open_activity_book_in_current_tab(book_data['file_path']))
        open_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 6px 10px;
                border-radius: 4px;
                font-size: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        buttons_layout.addWidget(open_btn)

        # Continue Training button
        train_btn = QPushButton("⚡ Train")
        train_btn.clicked.connect(lambda: self.open_activity_book_in_training_tab(book_data['file_path']))
        train_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 6px 10px;
                border-radius: 4px;
                font-size: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        buttons_layout.addWidget(train_btn)

        layout.addWidget(buttons_widget)

        # Show All Sessions button
        sessions_btn = QPushButton("📊 Sessions")
        sessions_btn.clicked.connect(lambda: self.show_book_session_details(book_data))
        sessions_btn.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 10px;
                font-weight: bold;
                margin-top: 2px;
            }
            QPushButton:hover {
                background-color: #1565C0;
            }
        """)
        layout.addWidget(sessions_btn)

        return card

    def get_activity_card_style(self):
        """Get styling for activity book cards (matching Recent Books theme)"""
        return """
            QFrame#ActivityBookCard {
                background-color: #1E1E1E;
                border: 1px solid #3A3A3A;
                border-radius: 12px;
                margin: 2px;
                padding: 0px;
            }
            QFrame#ActivityBookCard:hover {
                border: 2px solid #1976D2;
                background-color: #252525;
                transform: translateY(-2px);
            }
            QLabel {
                background-color: transparent;
            }
            QPushButton {
                border-radius: 4px;
                font-weight: bold;
                padding: 6px 12px;
            }
        """

    def load_activity_book_cover(self, file_path):
        """Load cover image for activity book card"""
        try:
            if not os.path.exists(file_path):
                return None

            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            if len(doc) > 0:
                page = doc[0]
                # Get page as pixmap
                mat = fitz.Matrix(1.0, 1.0)  # No scaling
                pix = page.get_pixmap(matrix=mat)

                # Convert to QPixmap
                img_data = pix.tobytes("ppm")
                pixmap = QPixmap()
                pixmap.loadFromData(img_data)

                # Scale to fit cover size
                scaled_pixmap = pixmap.scaled(200, 280, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                doc.close()
                return scaled_pixmap

            doc.close()
            return None

        except Exception as e:
            print(f"Error loading cover for {file_path}: {e}")
            return None

    def open_activity_book(self, file_path):
        """Open an activity book for reading speed measurement (legacy method - redirects to current tab)"""
        self.open_activity_book_in_current_tab(file_path)

    def open_activity_book_in_current_tab(self, file_path):
        """Open an activity book in Current Book tab for reading speed measurement"""
        if os.path.exists(file_path):
            # Switch to Current Book tab and load the book
            # Tab order: 0=Fast Reading Training, 1=Current Book, 2=Activity
            self.tab_widget.setCurrentIndex(1)  # Current Book tab (now at index 1)
            self.load_pdf(file_path)
        else:
            QMessageBox.warning(
                self,
                self.localization.get_text("error") if self.localization else "File Not Found",
                f"The file '{file_path}' could not be found. It may have been moved or deleted."
            )

    def open_activity_book_in_training_tab(self, file_path):
        """Open an activity book in Training Session tab and continue from last position"""
        if not os.path.exists(file_path):
            QMessageBox.warning(
                self,
                self.localization.get_text("error") if self.localization else "File Not Found",
                f"The file '{file_path}' could not be found. It may have been moved or deleted."
            )
            return

        try:
            # Switch to Fast Reading Trainer tab (index 1)
            self.tab_widget.setCurrentIndex(1)

            # Check if there's a saved training session for this book
            if self.recent_books_manager:
                sessions = self.recent_books_manager.get_training_sessions(file_path)

                if sessions and len(sessions) > 0:
                    # Get the most recent session
                    latest_session = sessions[0]

                    # Ask user if they want to continue from last position
                    reply = QMessageBox.question(
                        self,
                        self.localization.get_text("continue_session") if self.localization else "Continue Session",
                        f"{self.localization.get_text('continue_from_last_position') if self.localization else 'Continue from last position'}\n\n"
                        f"{self.localization.get_text('progress') if self.localization else 'Progress'}: {latest_session.get('progress_percentage', 0):.1f}%\n"
                        f"{self.localization.get_text('last_session') if self.localization else 'Last session'}: {latest_session.get('session_date', 'Unknown')}",
                        QMessageBox.Yes | QMessageBox.No,
                        QMessageBox.Yes
                    )

                    if reply == QMessageBox.Yes:
                        # Resume the training session
                        self.resume_training_session(latest_session)
                        return

            # If no session or user chose not to continue, load fresh
            self.load_pdf_for_training(file_path)

        except Exception as e:
            print(f"Error opening activity book in training tab: {e}")
            QMessageBox.critical(
                self,
                self.localization.get_text("error") if self.localization else "Error",
                f"Error opening book: {str(e)}"
            )

    def show_book_session_details(self, book_data):
        """Show detailed session information for a book"""
        dialog = BookSessionDetailsDialog(book_data, self.recent_books_manager, self)
        if dialog.exec() == QDialog.Accepted:
            # Refresh activity if any changes were made
            self.refresh_activity_books()

    def delete_prepared_book_config(self, file_path):
        """Delete the reading speed configuration for a book"""
        reply = QMessageBox.question(
            self,
            "Remove Configuration",
            f"Are you sure you want to remove the reading speed configuration for this book?\n\n"
            f"File: {os.path.basename(file_path)}\n\n"
            f"This will not delete the PDF file, only the reading speed settings.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            if self.recent_books_manager.delete_reading_speed_config(file_path):
                QMessageBox.information(
                    self,
                    "Configuration Removed",
                    "The reading speed configuration has been removed successfully."
                )
                self.refresh_activity_books()
            else:
                QMessageBox.warning(
                    self,
                    "Error",
                    "Failed to remove the reading speed configuration."
                )

    def create_dashboard_content(self, dashboard_layout):
        """Create compact, card-based dashboard content"""
        # Create scroll area for dashboard content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        dashboard_widget = QWidget()
        content_layout = QVBoxLayout(dashboard_widget)
        content_layout.setSpacing(15)
        content_layout.setContentsMargins(10, 10, 10, 10)

        # General Statistics Summary Cards
        self.create_general_stats_cards(content_layout)

        # Current Book Information Section (if a book is loaded)
        if hasattr(self, 'pdf_path') and self.pdf_path:
            self.book_card = self.create_book_card()
            content_layout.addWidget(self.book_card)

            # Current Book Sessions Dashboard
            self.sessions_dashboard = self.create_sessions_dashboard()
            content_layout.addWidget(self.sessions_dashboard)

        content_layout.addStretch()
        scroll_area.setWidget(dashboard_widget)
        dashboard_layout.addWidget(scroll_area)

    def create_general_stats_cards(self, layout):
        """Create general statistics summary in card format"""
        # Container for stats cards
        stats_container = QWidget()
        stats_layout = QVBoxLayout(stats_container)
        stats_layout.setSpacing(15)

        # Title
        title = QLabel("📊 Reading Speed Overview")
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setStyleSheet("color: #1976D2; margin-bottom: 10px;")
        stats_layout.addWidget(title)

        # Create grid layout for cards
        cards_widget = QWidget()
        cards_layout = QGridLayout(cards_widget)
        cards_layout.setSpacing(15)

        # Get overall statistics
        overall_stats = self.get_overall_statistics()

        # Create stat cards
        cards_data = [
            ("📚", "Total Books", str(overall_stats.get('total_books', 0)), "Books configured for speed reading"),
            ("📖", "Books Read", str(overall_stats.get('books_with_sessions', 0)), "Books with reading sessions"),
            ("📄", "Pages Read", f"{overall_stats.get('total_pages_read', 0):,}", "Total pages across all sessions"),
            ("⚡", "Average Speed", f"{overall_stats.get('overall_avg_wpm', 0):.0f} WPM", "Overall average reading speed"),
            ("🏆", "Best Speed", f"{overall_stats.get('best_wpm', 0):.0f} WPM", "Highest speed achieved"),
            ("⏱️", "Total Time", self.format_reading_time(overall_stats.get('total_time_seconds', 0)), "Total reading time")
        ]

        # Create cards in 3x2 grid
        for i, (icon, title, value, description) in enumerate(cards_data):
            row = i // 3
            col = i % 3
            card = self.create_stat_card(icon, title, value, description)
            cards_layout.addWidget(card, row, col)

        stats_layout.addWidget(cards_widget)
        layout.addWidget(stats_container)

    def create_stat_card(self, icon, title, value, description):
        """Create a single statistics card"""
        card = QFrame()
        card.setFrameStyle(QFrame.Box)
        card.setStyleSheet("""
            QFrame {
                border: 1px solid #3A3A3A;
                border-radius: 8px;
                background-color: #1E1E1E;
                padding: 15px;
                margin: 2px;
            }
            QFrame:hover {
                border-color: #2196F3;
                background-color: #252525;
            }
            QLabel {
                background-color: transparent;
            }
        """)
        card.setMinimumHeight(100)
        card.setMaximumHeight(120)

        layout = QVBoxLayout(card)
        layout.setSpacing(5)
        layout.setContentsMargins(10, 10, 10, 10)

        # Icon and title row
        header_layout = QHBoxLayout()

        icon_label = QLabel(icon)
        icon_label.setFont(QFont("Arial", 20))
        header_layout.addWidget(icon_label)

        title_label = QLabel(title)
        title_label.setFont(QFont("Arial", 11, QFont.Bold))
        title_label.setStyleSheet("color: #CCCCCC;")
        header_layout.addWidget(title_label)

        header_layout.addStretch()
        layout.addLayout(header_layout)

        # Value
        value_label = QLabel(value)
        value_label.setFont(QFont("Arial", 18, QFont.Bold))
        value_label.setStyleSheet("color: #1976D2;")
        layout.addWidget(value_label)

        # Description
        desc_label = QLabel(description)
        desc_label.setFont(QFont("Arial", 9))
        desc_label.setStyleSheet("color: #999;")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)

        return card

    def get_overall_statistics(self):
        """Get overall reading statistics across all books"""
        if not self.recent_books_manager:
            return {}

        try:
            # Get all prepared books
            prepared_books = self.recent_books_manager.get_all_prepared_books()

            stats = {
                'total_books': len(prepared_books),
                'books_with_sessions': 0,
                'total_pages_read': 0,
                'total_time_seconds': 0,
                'all_wpm_values': [],
                'overall_avg_wpm': 0,
                'best_wpm': 0
            }

            for book in prepared_books:
                book_stats = book.get('statistics', {})
                if book_stats and book_stats.get('total_sessions', 0) > 0:
                    stats['books_with_sessions'] += 1
                    stats['total_pages_read'] += book_stats.get('total_pages_read', 0)
                    stats['total_time_seconds'] += book_stats.get('total_time_seconds', 0)

                    # Collect WPM values for overall average
                    sessions = book_stats.get('sessions', [])
                    for session in sessions:
                        wpm = session.get('words_per_minute', 0)
                        if wpm > 0:
                            stats['all_wpm_values'].append(wpm)

                    # Track best WPM
                    book_best = book_stats.get('best_wpm', 0)
                    if book_best > stats['best_wpm']:
                        stats['best_wpm'] = book_best

            # Calculate overall average WPM
            if stats['all_wpm_values']:
                stats['overall_avg_wpm'] = sum(stats['all_wpm_values']) / len(stats['all_wpm_values'])

            return stats

        except Exception as e:
            print(f"Error getting overall statistics: {e}")
            return {}

    def format_reading_time(self, total_seconds):
        """Format reading time in a human-readable format"""
        seconds_text = self.localization.get_text("seconds") if self.localization else "seconds"
        minutes_text = self.localization.get_text("minutes") if self.localization else "minutes"
        hours_text = self.localization.get_text("hours") if self.localization else "hours"

        if total_seconds < 60:
            return f"{total_seconds:.0f} {seconds_text}"
        elif total_seconds < 3600:
            minutes = total_seconds // 60
            return f"{minutes:.0f} {minutes_text}"
        else:
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            if minutes > 0:
                return f"{hours:.0f} {hours_text} {minutes:.0f} {minutes_text}"
            else:
                return f"{hours:.0f} {hours_text}"

    def start_reading_session(self):
        """Start a reading session by opening the PDF"""
        if not self.pdf_path:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                "Please select a PDF file first."
            )
            return

        try:
            # Open PDF in external viewer
            if os.name == 'nt':  # Windows
                os.startfile(self.pdf_path)
            elif os.name == 'posix':  # macOS and Linux
                os.system(f'open "{self.pdf_path}"' if sys.platform == 'darwin' else f'xdg-open "{self.pdf_path}"')

            # Start timer
            self.start_time = time.time()
            self.reading_timer.start(1000)  # Update every second

            # Update UI
            self.btn_start_reading.setVisible(False)
            self.finish_widget.setVisible(True)
            self.results_group.setVisible(False)
            self.comprehension_group.setVisible(False)

        except Exception as e:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                f"Could not open PDF: {str(e)}"
            )

    def update_timer_display(self):
        """Update the timer display every second"""
        self.elapsed_time_seconds = int(time.time() - self.start_time)
        minutes = self.elapsed_time_seconds // 60
        seconds = self.elapsed_time_seconds % 60
        self.lbl_timer.setText(f"{minutes:02d}:{seconds:02d}")

    def finish_reading_session(self):
        """Stop timer and calculate reading speed"""
        self.reading_timer.stop()

        pages_read = self.spin_pages_read.value()

        if self.elapsed_time_seconds == 0:
            wpm = 0
        else:
            total_words_read = pages_read * self.avg_words_per_page
            wpm = (total_words_read / self.elapsed_time_seconds) * 60

        # Update results display
        self.lbl_wpm.setText(f"{wpm:.0f} {self.localization.get_text('wpm')}")
        self.lbl_total_time.setText(f"{self.elapsed_time_seconds // 60}:{self.elapsed_time_seconds % 60:02d}")
        self.lbl_recommendation.setText(self.get_recommendation(wpm))

        # Show results
        self.results_group.setVisible(True)
        self.finish_widget.setVisible(False)
        self.btn_start_reading.setVisible(True)
        self.lbl_timer.setText("00:00")

        # Save reading session
        self.save_reading_session(wpm, pages_read)

        # Generate comprehension questions if we have sample text
        if self.sample_text and len(self.sample_text.strip()) > 100:
            self.generate_comprehension_questions()

    def get_recommendation(self, wpm):
        """Get reading speed recommendation based on WPM"""
        if wpm < 200:
            return self.localization.get_text("slow_reader")
        elif wpm < 300:
            return self.localization.get_text("average_reader")
        elif wpm < 400:
            return self.localization.get_text("good_reader")
        elif wpm < 500:
            return self.localization.get_text("excellent_reader")
        else:
            return self.localization.get_text("speed_reader")

    def select_pdf_file(self):
        """Open file dialog to select a PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            "PDF Files (*.pdf)"
        )
        if file_path:
            self.load_pdf(file_path)

    def prepare_book_analysis(self):
        """Comprehensive PDF preparation workflow"""
        if not self.pdf_path or not self.pdf_doc:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("select_file_first")
            )
            return

        try:
            # Check if we already have configuration for this book
            if self.recent_books_manager:
                existing_config = self.recent_books_manager.get_reading_speed_config(self.pdf_path)
                if existing_config:
                    self.load_existing_configuration(existing_config)
                    return

            # Start fresh analysis
            self.analyze_pdf_content()

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to prepare book: {str(e)}"
            )

    def analyze_pdf_content(self):
        """Analyze PDF content to determine words per page"""
        try:
            # Show progress
            self.btn_prepare_book.setText("🔍 " + (self.localization.get_text("analyzing_pdf") if self.localization else "Analyzing..."))
            self.btn_prepare_book.setEnabled(False)

            # Sample pages for analysis (up to 5 pages)
            total_pages = self.pdf_doc.page_count
            if total_pages > 5:
                # Sample from beginning, middle, and end
                sample_pages = [0, total_pages//4, total_pages//2, 3*total_pages//4, total_pages-1]
            else:
                sample_pages = list(range(total_pages))

            total_words = 0
            valid_pages = 0
            sample_texts = []

            for page_num in sample_pages:
                try:
                    page = self.pdf_doc.load_page(page_num)
                    text = page.get_text("text").strip()

                    if text and len(text) > 50:  # Minimum text threshold
                        words = len(text.split())
                        if words > 20:  # Reasonable word count
                            total_words += words
                            valid_pages += 1
                            sample_texts.append(text[:200])  # Store sample for display
                except Exception as e:
                    print(f"Error analyzing page {page_num}: {e}")
                    continue

            # Determine preparation method and words per page
            if valid_pages >= 2:
                # Option A: Text-based PDF (preferred)
                avg_words = total_words / valid_pages
                self.show_analysis_results(avg_words, "automatic", sample_texts, total_pages)
            else:
                # Option B/C: OCR or default estimate
                self.show_preparation_options(total_pages)

        except Exception as e:
            self.reset_prepare_button()
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error analyzing PDF: {str(e)}"
            )

    def show_analysis_results(self, avg_words, method, sample_texts, total_pages):
        """Show automatic analysis results and ask for confirmation"""
        dialog = QDialog(self)
        dialog.setWindowTitle(self.localization.get_text("analysis_results") if self.localization else "Analysis Results")
        dialog.setModal(True)
        dialog.resize(600, 500)

        layout = QVBoxLayout(dialog)

        # Title
        title = QLabel("📊 " + (self.localization.get_text("word_analysis_complete") if self.localization else "Word Analysis Complete"))
        title.setStyleSheet("QLabel { font-size: 18px; font-weight: bold; color: #1976D2; margin-bottom: 15px; }")
        layout.addWidget(title)

        # Results summary
        summary_frame = QFrame()
        summary_frame.setFrameStyle(QFrame.Box)
        summary_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #4CAF50;
                border-radius: 10px;
                background-color: #e8f5e8;
                padding: 15px;
                margin-bottom: 15px;
            }
        """)
        summary_layout = QVBoxLayout(summary_frame)

        # Book info
        book_name = os.path.basename(self.pdf_path)
        book_label = QLabel(f"📖 {book_name}")
        book_label.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; }")
        summary_layout.addWidget(book_label)

        # Analysis results
        results_text = f"""
        ✅ {self.localization.get_text('total_pages') if self.localization else 'Total Pages'}: {total_pages}
        ✅ {self.localization.get_text('avg_words_per_page') if self.localization else 'Average Words per Page'}: {avg_words:.0f}
        ✅ {self.localization.get_text('preparation_method') if self.localization else 'Method'}: {self.localization.get_text('automatic_analysis') if self.localization else 'Automatic Analysis'}
        """
        results_label = QLabel(results_text)
        results_label.setStyleSheet("QLabel { font-size: 14px; line-height: 1.5; }")
        summary_layout.addWidget(results_label)

        layout.addWidget(summary_frame)

        # Sample text preview
        if sample_texts:
            preview_label = QLabel(self.localization.get_text("sample_text_preview") if self.localization else "Sample Text Preview:")
            preview_label.setStyleSheet("QLabel { font-weight: bold; margin-top: 10px; }")
            layout.addWidget(preview_label)

            preview_text = QTextEdit()
            preview_text.setPlainText(sample_texts[0])
            preview_text.setMaximumHeight(100)
            preview_text.setReadOnly(True)
            layout.addWidget(preview_text)

        # Manual adjustment option
        adjustment_frame = QFrame()
        adjustment_frame.setFrameStyle(QFrame.Box)
        adjustment_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ddd;
                border-radius: 8px;
                background-color: #f9f9f9;
                padding: 10px;
                margin-top: 15px;
            }
        """)
        adjustment_layout = QHBoxLayout(adjustment_frame)

        adjustment_layout.addWidget(QLabel(self.localization.get_text("adjust_words_per_page") if self.localization else "Adjust words per page:"))

        self.words_spinbox = QSpinBox()
        self.words_spinbox.setRange(50, 1000)
        self.words_spinbox.setValue(int(avg_words))
        self.words_spinbox.setSuffix(" " + (self.localization.get_text("words") if self.localization else "words"))
        adjustment_layout.addWidget(self.words_spinbox)
        adjustment_layout.addStretch()

        layout.addWidget(adjustment_frame)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel") if self.localization else "Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        confirm_btn = QPushButton(self.localization.get_text("confirm_preparation") if self.localization else "Confirm & Prepare")
        confirm_btn.clicked.connect(dialog.accept)
        confirm_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        button_layout.addWidget(confirm_btn)

        layout.addLayout(button_layout)

        # Show dialog and handle result
        if dialog.exec() == QDialog.Accepted:
            final_words = self.words_spinbox.value()
            self.finalize_preparation(final_words, "automatic", total_pages)
        else:
            self.reset_prepare_button()

    def show_preparation_options(self, total_pages):
        """Show preparation options for OCR or manual entry"""
        dialog = QDialog(self)
        dialog.setWindowTitle(self.localization.get_text("preparation_options") if self.localization else "Preparation Options")
        dialog.setModal(True)
        dialog.resize(500, 400)

        layout = QVBoxLayout(dialog)

        # Title
        title = QLabel("⚙️ " + (self.localization.get_text("choose_preparation_method") if self.localization else "Choose Preparation Method"))
        title.setStyleSheet("QLabel { font-size: 18px; font-weight: bold; color: #FF9800; margin-bottom: 15px; }")
        layout.addWidget(title)

        # Info message
        info_text = self.localization.get_text("pdf_analysis_limited") if self.localization else "This PDF has limited extractable text. Please choose a preparation method:"
        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setStyleSheet("QLabel { color: #666; margin-bottom: 20px; }")
        layout.addWidget(info_label)

        # Option B: Default estimate
        default_frame = QFrame()
        default_frame.setFrameStyle(QFrame.Box)
        default_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #2196F3;
                border-radius: 10px;
                background-color: #e3f2fd;
                padding: 15px;
                margin-bottom: 15px;
            }
        """)
        default_layout = QVBoxLayout(default_frame)

        default_title = QLabel("📊 " + (self.localization.get_text("use_default_estimate") if self.localization else "Use Default Estimate"))
        default_title.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; color: #1976D2; }")
        default_layout.addWidget(default_title)

        default_desc = QLabel(self.localization.get_text("default_estimate_desc") if self.localization else "Use a typical range of 120-150 words per page for standard documents.")
        default_desc.setWordWrap(True)
        default_layout.addWidget(default_desc)

        # Words per page selection for default
        words_layout = QHBoxLayout()
        words_layout.addWidget(QLabel(self.localization.get_text("words_per_page") if self.localization else "Words per page:"))
        self.default_words_spinbox = QSpinBox()
        self.default_words_spinbox.setRange(80, 300)
        self.default_words_spinbox.setValue(135)  # Default middle value
        words_layout.addWidget(self.default_words_spinbox)
        words_layout.addStretch()
        default_layout.addLayout(words_layout)

        layout.addWidget(default_frame)

        # Option C: OCR conversion
        ocr_frame = QFrame()
        ocr_frame.setFrameStyle(QFrame.Box)
        ocr_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #9C27B0;
                border-radius: 10px;
                background-color: #f3e5f5;
                padding: 15px;
                margin-bottom: 15px;
            }
        """)
        ocr_layout = QVBoxLayout(ocr_frame)

        ocr_title = QLabel("🔍 " + (self.localization.get_text("use_ocr_conversion") if self.localization else "Use OCR Conversion"))
        ocr_title.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; color: #7B1FA2; }")
        ocr_layout.addWidget(ocr_title)

        ocr_desc = QLabel(self.localization.get_text("ocr_conversion_desc") if self.localization else "Convert PDF to Word format using online OCR for accurate word count.")
        ocr_desc.setWordWrap(True)
        ocr_layout.addWidget(ocr_desc)

        layout.addWidget(ocr_frame)

        # Option D: Manual entry
        manual_frame = QFrame()
        manual_frame.setFrameStyle(QFrame.Box)
        manual_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #FF5722;
                border-radius: 10px;
                background-color: #fbe9e7;
                padding: 15px;
                margin-bottom: 15px;
            }
        """)
        manual_layout = QVBoxLayout(manual_frame)

        manual_title = QLabel("✏️ " + (self.localization.get_text("manual_entry") if self.localization else "Manual Entry"))
        manual_title.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; color: #D84315; }")
        manual_layout.addWidget(manual_title)

        manual_desc = QLabel(self.localization.get_text("manual_entry_desc") if self.localization else "Count words on a sample page and enter manually.")
        manual_desc.setWordWrap(True)
        manual_layout.addWidget(manual_desc)

        # Manual words input
        manual_words_layout = QHBoxLayout()
        manual_words_layout.addWidget(QLabel(self.localization.get_text("words_per_page") if self.localization else "Words per page:"))
        self.manual_words_spinbox = QSpinBox()
        self.manual_words_spinbox.setRange(50, 1000)
        self.manual_words_spinbox.setValue(200)
        manual_words_layout.addWidget(self.manual_words_spinbox)
        manual_words_layout.addStretch()
        manual_layout.addLayout(manual_words_layout)

        layout.addWidget(manual_frame)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel") if self.localization else "Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        default_btn = QPushButton(self.localization.get_text("use_default") if self.localization else "Use Default")
        default_btn.clicked.connect(lambda: self.handle_preparation_choice(dialog, "default"))
        default_btn.setStyleSheet("QPushButton { background-color: #2196F3; color: white; padding: 8px 16px; border-radius: 4px; }")
        button_layout.addWidget(default_btn)

        ocr_btn = QPushButton(self.localization.get_text("use_ocr") if self.localization else "Use OCR")
        ocr_btn.clicked.connect(lambda: self.handle_preparation_choice(dialog, "ocr"))
        ocr_btn.setStyleSheet("QPushButton { background-color: #9C27B0; color: white; padding: 8px 16px; border-radius: 4px; }")
        button_layout.addWidget(ocr_btn)

        manual_btn = QPushButton(self.localization.get_text("manual_entry") if self.localization else "Manual Entry")
        manual_btn.clicked.connect(lambda: self.handle_preparation_choice(dialog, "manual"))
        manual_btn.setStyleSheet("QPushButton { background-color: #FF5722; color: white; padding: 8px 16px; border-radius: 4px; }")
        button_layout.addWidget(manual_btn)

        layout.addLayout(button_layout)

        dialog.exec()

    def handle_preparation_choice(self, dialog, choice):
        """Handle the user's preparation method choice"""
        total_pages = self.pdf_doc.page_count if self.pdf_doc else 1

        if choice == "default":
            words_per_page = self.default_words_spinbox.value()
            self.finalize_preparation(words_per_page, "default_estimate", total_pages)
        elif choice == "ocr":
            self.show_ocr_guidance()
        elif choice == "manual":
            words_per_page = self.manual_words_spinbox.value()
            self.finalize_preparation(words_per_page, "manual_entry", total_pages)

        dialog.accept()

    def show_ocr_guidance(self):
        """Show OCR conversion guidance"""
        dialog = QDialog(self)
        dialog.setWindowTitle(self.localization.get_text("ocr_guidance") if self.localization else "OCR Conversion Guide")
        dialog.setModal(True)
        dialog.resize(600, 500)

        layout = QVBoxLayout(dialog)

        # Title
        title = QLabel("🔍 " + (self.localization.get_text("ocr_conversion_guide") if self.localization else "OCR Conversion Guide"))
        title.setStyleSheet("QLabel { font-size: 18px; font-weight: bold; color: #9C27B0; margin-bottom: 15px; }")
        layout.addWidget(title)

        # Instructions
        instructions = QTextEdit()
        instructions_text = self.localization.get_text("ocr_instructions") if self.localization else """
        Follow these steps to convert your PDF using online OCR:

        1. Click the "Open OCR Tool" button below to open the online converter
        2. Upload your PDF file to the OCR service
        3. Wait for the conversion to complete
        4. Download the converted Word document
        5. Open the Word document and count words on a sample page
        6. Return here and enter the words per page manually

        This method provides the most accurate word count for scanned PDFs.
        """
        instructions.setPlainText(instructions_text)
        instructions.setMaximumHeight(200)
        instructions.setReadOnly(True)
        layout.addWidget(instructions)

        # OCR tool button
        ocr_btn = QPushButton("🌐 " + (self.localization.get_text("open_ocr_tool") if self.localization else "Open OCR Tool"))
        ocr_btn.clicked.connect(self.open_ocr_tool)
        ocr_btn.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 16px;
                margin: 20px 0;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        layout.addWidget(ocr_btn)

        # Manual entry after OCR
        manual_frame = QFrame()
        manual_frame.setFrameStyle(QFrame.Box)
        manual_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #FF5722;
                border-radius: 10px;
                background-color: #fbe9e7;
                padding: 15px;
            }
        """)
        manual_layout = QVBoxLayout(manual_frame)

        manual_title = QLabel("✏️ " + (self.localization.get_text("enter_ocr_results") if self.localization else "Enter OCR Results"))
        manual_title.setStyleSheet("QLabel { font-size: 16px; font-weight: bold; color: #D84315; }")
        manual_layout.addWidget(manual_title)

        manual_desc = QLabel(self.localization.get_text("ocr_results_desc") if self.localization else "After OCR conversion, count words on a sample page and enter below:")
        manual_desc.setWordWrap(True)
        manual_layout.addWidget(manual_desc)

        # Words input
        words_layout = QHBoxLayout()
        words_layout.addWidget(QLabel(self.localization.get_text("words_per_page") if self.localization else "Words per page:"))
        self.ocr_words_spinbox = QSpinBox()
        self.ocr_words_spinbox.setRange(50, 1000)
        self.ocr_words_spinbox.setValue(200)
        words_layout.addWidget(self.ocr_words_spinbox)
        words_layout.addStretch()
        manual_layout.addLayout(words_layout)

        layout.addWidget(manual_frame)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel") if self.localization else "Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        confirm_btn = QPushButton(self.localization.get_text("confirm_ocr_results") if self.localization else "Confirm OCR Results")
        confirm_btn.clicked.connect(lambda: self.finalize_ocr_preparation(dialog))
        confirm_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF5722;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #D84315;
            }
        """)
        button_layout.addWidget(confirm_btn)

        layout.addLayout(button_layout)

        if dialog.exec() == QDialog.Rejected:
            self.reset_prepare_button()

    def finalize_ocr_preparation(self, dialog):
        """Finalize preparation with OCR results"""
        words_per_page = self.ocr_words_spinbox.value()
        total_pages = self.pdf_doc.page_count if self.pdf_doc else 1
        self.finalize_preparation(words_per_page, "ocr_conversion", total_pages)
        dialog.accept()

    def open_ocr_tool(self):
        """Open the online OCR tool"""
        ocr_url = "https://colab.research.google.com/drive/your-ocr-notebook"  # Replace with actual URL
        try:
            import webbrowser
            webbrowser.open(ocr_url)
        except Exception as e:
            QMessageBox.information(
                self,
                self.localization.get_text("information"),
                f"Please open this URL manually: {ocr_url}"
            )

    def finalize_preparation(self, words_per_page, method, total_pages):
        """Finalize book preparation and enable reading session"""
        try:
            # Store preparation data
            self.words_per_page = words_per_page
            self.preparation_method = method
            self.total_pages = total_pages

            # Update UI with results
            self.lbl_avg_words.setText(f"{words_per_page:.0f}")
            self.lbl_total_pages.setText(str(total_pages))

            method_display = {
                "automatic": self.localization.get_text("automatic_analysis") if self.localization else "Automatic Analysis",
                "default_estimate": self.localization.get_text("default_estimate") if self.localization else "Default Estimate",
                "manual_entry": self.localization.get_text("manual_entry") if self.localization else "Manual Entry",
                "ocr_conversion": self.localization.get_text("ocr_conversion") if self.localization else "OCR Conversion"
            }
            self.lbl_method.setText(method_display.get(method, method))

            # Show analysis results
            self.analysis_group.setVisible(True)

            # Enable reading session
            self.reading_group.setVisible(True)
            self.btn_start_reading.setEnabled(True)

            # Reset prepare button
            self.reset_prepare_button()

            # Save configuration for future use
            if self.recent_books_manager:
                self.recent_books_manager.save_reading_speed_config(
                    file_path=self.pdf_path,
                    total_pages=total_pages,
                    avg_words_per_page=words_per_page,
                    preparation_method=method
                )

            # Show success message
            QMessageBox.information(
                self,
                self.localization.get_text("book_prepared") if self.localization else "Book Prepared",
                f"{self.localization.get_text('book_ready') if self.localization else 'Book is ready for reading speed measurement!'}\n\n"
                f"📖 {os.path.basename(self.pdf_path)}\n"
                f"📊 {words_per_page:.0f} {self.localization.get_text('words_per_page') if self.localization else 'words per page'}\n"
                f"📄 {total_pages} {self.localization.get_text('total_pages') if self.localization else 'pages'}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error finalizing preparation: {str(e)}"
            )

    def reset_prepare_button(self):
        """Reset the prepare button to its original state"""
        prepare_text = self.localization.get_text("prepare_book") if self.localization else "Prepare Book"
        self.btn_prepare_book.setText("🔍 " + prepare_text)
        self.btn_prepare_book.setEnabled(True)

    def load_existing_configuration(self, config):
        """Load existing book configuration"""
        try:
            self.words_per_page = config.get('words_per_page', 200)
            self.preparation_method = config.get('method', 'saved')
            self.total_pages = config.get('total_pages', 1)

            # Update UI
            self.lbl_avg_words.setText(f"{self.words_per_page:.0f}")
            self.lbl_total_pages.setText(str(self.total_pages))
            self.lbl_method.setText(self.localization.get_text("saved_configuration") if self.localization else "Saved Configuration")

            # Show sections
            self.analysis_group.setVisible(True)
            self.reading_group.setVisible(True)
            self.btn_start_reading.setEnabled(True)

            # Reset prepare button
            self.reset_prepare_button()

        except Exception as e:
            print(f"Error loading existing configuration: {e}")
            # Fall back to fresh analysis
            self.analyze_pdf_content()

    def start_reading_session(self):
        """Start or resume a reading session with timer"""
        try:
            # Check if resuming from a paused/stopped state
            if self.is_reading_paused or self.is_reading_stopped:
                # Resume from where we left off
                if self.is_reading_paused:
                    # Calculate time spent paused and add to total paused duration
                    self.total_paused_duration += time.time() - self.paused_time

                # Adjust start time to account for paused duration
                self.start_time = time.time() - self.elapsed_time_seconds

                # Reset pause/stop flags
                self.is_reading_paused = False
                self.is_reading_stopped = False

                # Update button text
                self.btn_start_reading.setText(self.localization.get_text("pause_reading") if self.localization else "Pause Reading")
            else:
                # Starting fresh session
                self.start_time = time.time()
                self.elapsed_time_seconds = 0
                self.total_paused_duration = 0
                self.current_reading_page = 1

                # Update button text
                self.btn_start_reading.setText(self.localization.get_text("pause_reading") if self.localization else "Pause Reading")

            # Create and start timer
            if hasattr(self, 'reading_timer'):
                self.reading_timer.stop()

            self.reading_timer = QTimer()
            self.reading_timer.timeout.connect(self.update_timer_display)
            self.reading_timer.start(1000)  # Update every second

            # Update UI
            self.btn_start_reading.setEnabled(True)
            self.finish_widget.setVisible(True)

            # Set default pages to read (user can adjust) only if starting fresh
            if self.current_reading_page == 1:
                self.spin_pages_read.setValue(min(5, self.total_pages))

            # Open PDF in external viewer only if starting fresh
            if self.current_reading_page == 1:
                self.open_pdf_for_reading()

            # Show timer prominently
            self.update_timer_display()

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error starting reading session: {str(e)}"
            )

    def toggle_reading_session(self):
        """Toggle between start, pause, and resume states"""
        try:
            # Check current state
            if not hasattr(self, 'reading_timer') or not self.reading_timer.isActive():
                # Timer is not running - either starting fresh or resuming
                self.start_reading_session()
            else:
                # Timer is running - pause it
                self.pause_reading_session()
        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error toggling reading session: {str(e)}"
            )

    def pause_reading_session(self):
        """Pause the current reading session"""
        try:
            # Stop the timer
            if hasattr(self, 'reading_timer'):
                self.reading_timer.stop()

            # Save current state
            self.is_reading_paused = True
            self.paused_time = time.time()
            self.elapsed_time_seconds = int(time.time() - self.start_time)

            # Update current page from spinbox
            if hasattr(self, 'spin_pages_read'):
                self.current_reading_page = self.spin_pages_read.value()

            # Update button text to show "Resume"
            self.btn_start_reading.setText(self.localization.get_text("resume_reading") if self.localization else "Resume Reading")

        except Exception as e:
            print(f"Error pausing reading session: {e}")

    def stop_reading_session(self):
        """Stop the current reading session (can be resumed later)"""
        try:
            # Stop the timer
            if hasattr(self, 'reading_timer'):
                self.reading_timer.stop()

            # Save current state
            self.is_reading_stopped = True
            self.elapsed_time_seconds = int(time.time() - self.start_time)

            # Update current page from spinbox
            if hasattr(self, 'spin_pages_read'):
                self.current_reading_page = self.spin_pages_read.value()

            # Update button text to show "Resume"
            self.btn_start_reading.setText(self.localization.get_text("resume_reading") if self.localization else "Resume Reading")

        except Exception as e:
            print(f"Error stopping reading session: {e}")

    def update_timer_display(self):
        """Update the timer display every second"""
        try:
            self.elapsed_time_seconds = int(time.time() - self.start_time)
            minutes = self.elapsed_time_seconds // 60
            seconds = self.elapsed_time_seconds % 60
            self.lbl_timer.setText(f"{minutes:02d}:{seconds:02d}")
        except Exception as e:
            print(f"Error updating timer: {e}")

    def open_pdf_for_reading(self):
        """Open PDF in external viewer for reading"""
        try:
            if self.pdf_path and os.path.exists(self.pdf_path):
                import subprocess
                import platform

                system = platform.system()
                if system == "Windows":
                    os.startfile(self.pdf_path)
                elif system == "Darwin":  # macOS
                    subprocess.run(["open", self.pdf_path])
                else:  # Linux
                    subprocess.run(["xdg-open", self.pdf_path])

        except Exception as e:
            print(f"Error opening PDF: {e}")
            QMessageBox.information(
                self,
                self.localization.get_text("information"),
                f"{self.localization.get_text('open_pdf_manually') if self.localization else 'Please open the PDF manually:'}\n{self.pdf_path}"
            )

    def finish_reading_session(self):
        """Finish the reading session and calculate results"""
        try:
            # Stop timer
            if hasattr(self, 'reading_timer'):
                self.reading_timer.stop()

            # Get session data
            pages_read = self.spin_pages_read.value()
            total_time_seconds = self.elapsed_time_seconds
            words_read = pages_read * self.avg_words_per_page

            # Calculate WPM
            if total_time_seconds > 0:
                wpm = (words_read / total_time_seconds) * 60
            else:
                wpm = 0

            # Save session to database
            self.save_reading_session_data(pages_read, total_time_seconds, wpm, words_read)

            # Show results
            self.display_reading_results(pages_read, total_time_seconds, wpm, words_read)

            # Reset session state
            self.is_reading_paused = False
            self.is_reading_stopped = False
            self.current_reading_page = 1
            self.total_paused_duration = 0

            # Reset button text
            self.btn_start_reading.setText(self.localization.get_text("start_reading_session") if self.localization else "Start Reading Session")

            # Update UI state
            self.finish_widget.setVisible(False)
            self.results_group.setVisible(True)

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error finishing reading session: {str(e)}"
            )

    def save_reading_session_data(self, pages_read, total_time_seconds, wpm, words_read):
        """Save reading session data to database"""
        try:
            if self.recent_books_manager:
                session_data = {
                    'file_path': self.pdf_path,
                    'book_name': os.path.basename(self.pdf_path),
                    'session_date': datetime.now().isoformat(),
                    'pages_read': pages_read,
                    'time_seconds': total_time_seconds,
                    'words_per_page': self.words_per_page,
                    'total_words': words_read,
                    'wpm': wpm,
                    'preparation_method': self.preparation_method
                }

                # Save to database
                self.recent_books_manager.save_reading_session(session_data)

                # Refresh activity table if it exists
                if hasattr(self, 'refresh_sessions_table'):
                    self.refresh_sessions_table()

        except Exception as e:
            print(f"Error saving reading session: {e}")

    def display_reading_results(self, pages_read, total_time_seconds, wpm, words_read):
        """Display reading session results"""
        try:
            # Format time
            minutes = total_time_seconds // 60
            seconds = total_time_seconds % 60
            time_str = f"{minutes}m {seconds}s" if minutes > 0 else f"{seconds}s"

            # Update result labels
            self.lbl_wpm.setText(f"{wpm:.0f} " + (self.localization.get_text("wpm") if self.localization else "WPM"))
            self.lbl_total_time.setText(time_str)
            self.lbl_words_read.setText(f"{words_read:.0f}")

            # Generate recommendation
            recommendation = self.generate_reading_recommendation(wpm)
            self.lbl_recommendation.setText(recommendation)

            # Apply theme to results
            self.apply_results_theme()

        except Exception as e:
            print(f"Error displaying results: {e}")

    def generate_reading_recommendation(self, wpm):
        """Generate reading speed recommendation"""
        try:
            if wpm < 150:
                level = self.localization.get_text("beginner_reader") if self.localization else "Beginner Reader"
                advice = self.localization.get_text("beginner_advice") if self.localization else "Practice reading regularly to improve your speed. Focus on reducing subvocalization."
                emoji = "🐌"
            elif wpm < 250:
                level = self.localization.get_text("average_reader") if self.localization else "Average Reader"
                advice = self.localization.get_text("average_advice") if self.localization else "Good reading speed! Try speed reading techniques to reach the next level."
                emoji = "📖"
            elif wpm < 350:
                level = self.localization.get_text("good_reader") if self.localization else "Good Reader"
                advice = self.localization.get_text("good_advice") if self.localization else "Excellent reading speed! You're above average. Keep practicing to maintain this level."
                emoji = "🚀"
            else:
                level = self.localization.get_text("expert_reader") if self.localization else "Expert Reader"
                advice = self.localization.get_text("expert_advice") if self.localization else "Outstanding reading speed! You're in the top tier of readers."
                emoji = "🏆"

            return f"{emoji} {level}\n\n{advice}"

        except Exception as e:
            return "Keep practicing to improve your reading speed!"

    def apply_results_theme(self):
        """Apply theme-aware styling to results"""
        try:
            # Get current theme
            is_dark = self.is_dark_theme()

            if is_dark:
                # Dark theme colors
                wmp_bg = "#2d2d2d"
                wpm_border = "#BB86FC"
                stats_bg = "#1e1e1e"
                recommendation_bg = "#3d2f00"
                recommendation_border = "#ffb74d"
            else:
                # Light theme colors
                wpm_bg = "#e3f2fd"
                wpm_border = "#1976D2"
                stats_bg = "#f9f9f9"
                recommendation_bg = "#fff3e0"
                recommendation_border = "#ff9800"

            # Update WPM frame styling
            wpm_frame = self.lbl_wpm.parent()
            if wpm_frame:
                wpm_frame.setStyleSheet(f"""
                    QFrame {{
                        border: 2px solid {wpm_border};
                        border-radius: 10px;
                        background-color: {wpm_bg};
                        padding: 15px;
                    }}
                """)

            # Update stats frames
            for frame in [self.lbl_total_time.parent(), self.lbl_words_read.parent()]:
                if frame:
                    frame.setStyleSheet(f"""
                        QFrame {{
                            border: 1px solid #ddd;
                            border-radius: 8px;
                            background-color: {stats_bg};
                            padding: 10px;
                        }}
                    """)

            # Update recommendation styling
            self.lbl_recommendation.setStyleSheet(f"""
                QLabel {{
                    background-color: {recommendation_bg};
                    border: 1px solid {recommendation_border};
                    border-radius: 8px;
                    padding: 15px;
                    font-size: 14px;
                    color: #e65100;
                    margin-top: 10px;
                }}
            """)

        except Exception as e:
            print(f"Error applying results theme: {e}")

    def reset_session(self):
        """Reset for a new reading session"""
        try:
            # Hide results and finish controls
            self.results_group.setVisible(False)
            self.finish_widget.setVisible(False)

            # Reset timer
            self.lbl_timer.setText("00:00")
            if hasattr(self, 'reading_timer'):
                self.reading_timer.stop()

            # Re-enable start button
            self.btn_start_reading.setEnabled(True)

            # Reset session variables
            self.start_time = None
            self.elapsed_time_seconds = 0

        except Exception as e:
            print(f"Error resetting session: {e}")

    def is_dark_theme(self):
        """Check if current theme is dark"""
        try:
            if hasattr(self, 'parent_window') and self.parent_window:
                if hasattr(self.parent_window, 'settings'):
                    return self.parent_window.settings.get("theme", "light") == "dark"
                return getattr(self.parent_window, 'current_theme', 'light') == 'dark'

            # Fallback: detect from system palette
            from PySide6.QtGui import QPalette
            from PySide6.QtWidgets import QApplication
            palette = QApplication.palette()
            return palette.color(QPalette.Window).lightness() < 128
        except Exception as e:
            return False

    def load_pdf(self, file_path):
        """Load and initialize PDF file for reading speed analysis"""
        if not fitz:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                "PyMuPDF (fitz) is required for PDF processing"
            )
            return

        try:
            # Show loading message
            self.pdf_display.setText(self.localization.get_text("analyzing_pdf") if self.localization else "Loading...")

            # Open PDF document
            self.pdf_doc = fitz.open(file_path)
            self.pdf_path = file_path
            self.total_pages = self.pdf_doc.page_count

            # Update display
            self.pdf_display.setText(os.path.basename(file_path))
            self.btn_prepare_book.setEnabled(True)

            # Reset UI state
            self.analysis_group.setVisible(False)
            self.reading_group.setVisible(False)
            self.results_group.setVisible(False)

            # Add book to recent books manager for data persistence
            if self.recent_books_manager:
                try:
                    # Create book entry if it doesn't exist - pass just the file path
                    self.recent_books_manager.add_book(file_path)
                except Exception as e:
                    print(f"Error adding book to recent books: {e}")

            # Check for existing configuration
            if self.recent_books_manager:
                existing_config = self.recent_books_manager.get_reading_speed_config(file_path)
                if existing_config:
                    # Ask user if they want to use existing configuration
                    reply = QMessageBox.question(
                        self,
                        self.localization.get_text("existing_configuration") if self.localization else "Existing Configuration",
                        self.localization.get_text("use_existing_config") if self.localization else "This book has been prepared before. Use existing configuration?",
                        QMessageBox.Yes | QMessageBox.No,
                        QMessageBox.Yes
                    )

                    if reply == QMessageBox.Yes:
                        self.load_existing_configuration()
                        return

            # Success message
            QMessageBox.information(
                self,
                self.localization.get_text("pdf_loaded") if self.localization else "PDF Loaded",
                f"{self.localization.get_text('pdf_loaded_successfully') if self.localization else 'PDF loaded successfully!'}\n\n"
                f"📖 {os.path.basename(file_path)}\n"
                f"📄 {self.total_pages} {self.localization.get_text('pages') if self.localization else 'pages'}\n\n"
                f"{self.localization.get_text('click_prepare_to_start') if self.localization else 'Click \"Prepare Book\" to start analysis.'}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Error loading PDF: {str(e)}"
            )
            self.pdf_path = None
            self.pdf_doc = None
            self.btn_prepare_book.setEnabled(False)

    def on_tab_changed(self, index):
        """Handle tab change events"""
        if index == 2:  # Activity tab (index 2 - now 3rd tab)
            self.refresh_sessions_table()

    def save_reading_session(self, wpm, pages_read):
        """Save reading session to database"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            # Save the session
            self.recent_books_manager.save_reading_session(
                file_path=self.pdf_path,
                pages_read=pages_read,
                time_spent_seconds=self.elapsed_time_seconds,
                wpm=wpm,
                avg_words_per_page=self.avg_words_per_page
            )

            # Update the activity tab if it's visible
            if hasattr(self, 'tab_widget') and self.tab_widget.currentIndex() == 2:
                self.refresh_sessions_table()

        except Exception as e:
            print(f"Error saving reading session: {e}")

    def update_dashboard(self):
        """Update dashboard with latest statistics"""
        try:
            self.refresh_general_stats_cards()
        except Exception as e:
            print(f"Error updating dashboard: {e}")

    def refresh_general_stats_cards(self):
        """Refresh the general statistics cards"""
        try:
            # The dashboard cards are created dynamically in create_general_stats_cards
            # This method is called when the dashboard tab is activated
            print("Dashboard refresh requested - cards are created dynamically")

        except Exception as e:
            print(f"Error refreshing stats cards: {e}")

    def cache_analysis_results(self, preparation_method="auto"):
        """Cache analysis results for future use"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            # Save the reading speed configuration
            self.recent_books_manager.save_reading_speed_config(
                file_path=self.pdf_path,
                total_pages=self.total_pages,
                avg_words_per_page=self.avg_words_per_page,
                preparation_method=preparation_method,
                sample_text=self.sample_text[:500] if self.sample_text else ""  # Limit sample text
            )

            # Refresh activity tab if visible
            if hasattr(self, 'tab_widget') and self.tab_widget.currentIndex() == 2:
                self.refresh_activity_books()

        except Exception as e:
            print(f"Error caching analysis results: {e}")

    def analyze_word_count(self):
        """Analyze word count from PDF"""
        if not self.pdf_path:
            return

        try:
            # Open PDF and analyze
            if fitz:
                self.pdf_doc = fitz.open(self.pdf_path)
                self.total_pages = len(self.pdf_doc)

                # Sample 5 random pages for word count analysis
                sample_pages = min(5, self.total_pages)
                page_indices = random.sample(range(self.total_pages), sample_pages)

                total_words = 0
                sample_texts = []

                for page_num in page_indices:
                    page = self.pdf_doc[page_num]
                    text = page.get_text()
                    if text.strip():
                        words = len(text.split())
                        total_words += words
                        sample_texts.append(text[:200])  # Store sample for comprehension

                if total_words > 0:
                    self.avg_words_per_page = total_words / sample_pages
                    self.sample_text = " ".join(sample_texts)
                else:
                    # No text found - might be scanned PDF
                    self.handle_ocr_needed_scenario()
                    return

                # Update UI
                self.lbl_total_pages.setText(str(self.total_pages))
                self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f}")

                # Show analysis results
                self.analysis_group.setVisible(True)
                self.reading_group.setVisible(True)
                self.btn_start_reading.setEnabled(True)

                # Cache results
                self.cache_analysis_results("auto")

            else:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("error"),
                    "PyMuPDF not available. Please install it to analyze PDFs."
                )

        except Exception as e:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                f"Error analyzing PDF: {str(e)}"
            )

    def load_existing_configuration(self):
        """Load existing reading speed configuration for the current book"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            config = self.recent_books_manager.get_reading_speed_config(self.pdf_path)
            if config:
                self.total_pages = config.get('total_pages', 0)
                self.avg_words_per_page = config.get('avg_words_per_page', 0)
                self.sample_text = config.get('sample_text', '')

                # Update UI
                self.lbl_total_pages.setText(str(self.total_pages))
                self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f}")

                # Show sections
                self.analysis_group.setVisible(True)
                self.reading_group.setVisible(True)
                self.btn_start_reading.setEnabled(True)

                print(f"Loaded existing configuration for {os.path.basename(self.pdf_path)}")

        except Exception as e:
            print(f"Error loading existing configuration: {e}")

    def show_quick_setup_dialog(self):
        """Show quick setup dialog for new books"""
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QSpinBox

        dialog = QDialog(self)
        dialog.setWindowTitle("Quick Book Setup")
        dialog.setModal(True)
        dialog.resize(400, 300)

        layout = QVBoxLayout(dialog)

        # Title
        title = QLabel("📚 Setup Reading Speed Measurement")
        title.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(title)

        # Description
        desc = QLabel(f"Setting up: {os.path.basename(self.pdf_path)}")
        desc.setWordWrap(True)
        desc.setStyleSheet("color: #666; margin-bottom: 20px;")
        layout.addWidget(desc)

        # Options
        option1_btn = QPushButton("🚀 Use Default (150 words/page)")
        option1_btn.clicked.connect(lambda: self.setup_with_default(dialog))
        option1_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                margin: 5px 0;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        layout.addWidget(option1_btn)

        option2_btn = QPushButton("✏️ Manual Input")
        option2_btn.clicked.connect(lambda: self.setup_with_manual(dialog))
        option2_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                margin: 5px 0;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        layout.addWidget(option2_btn)

        option3_btn = QPushButton("🔍 Detailed Analysis")
        option3_btn.clicked.connect(lambda: self.setup_with_analysis(dialog))
        option3_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                margin: 5px 0;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        layout.addWidget(option3_btn)

        dialog.exec()

    def setup_with_default(self, dialog):
        """Setup with default 150 words per page"""
        self.avg_words_per_page = 150
        self.total_pages = 1  # Will be updated when PDF is opened
        self.sample_text = ""

        # Update UI
        self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (default)")
        self.analysis_group.setVisible(True)
        self.reading_group.setVisible(True)
        self.btn_start_reading.setEnabled(True)

        # Cache the configuration
        self.cache_analysis_results("default")

        dialog.accept()

    def setup_with_manual(self, dialog):
        """Setup with manual word count input"""
        from PySide6.QtWidgets import QInputDialog

        words_per_page, ok = QInputDialog.getInt(
            self,
            "Manual Word Count",
            "Enter average words per page:",
            150,  # default value
            50,   # minimum
            2000  # maximum
        )

        if ok:
            self.avg_words_per_page = words_per_page
            self.total_pages = 1  # Will be updated when PDF is opened
            self.sample_text = ""

            # Update UI
            self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (manual)")
            self.analysis_group.setVisible(True)
            self.reading_group.setVisible(True)
            self.btn_start_reading.setEnabled(True)

            # Cache the configuration
            self.cache_analysis_results("manual")

        dialog.accept()

    def setup_with_analysis(self, dialog):
        """Setup with detailed PDF analysis"""
        dialog.accept()
        self.analyze_word_count()

    def handle_ocr_needed_scenario(self):
        """Handle scenario where OCR is needed for text extraction"""
        # Update UI to show OCR is needed
        self.lbl_avg_words.setText("0 (OCR needed)")
        self.analysis_group.setVisible(True)

        # Offer online OCR (localized)
        reply = QMessageBox.question(
            self,
            self.localization.get_text("ocr_processing_required"),
            self.localization.get_text("ocr_processing_question"),
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )

        if reply == QMessageBox.Yes:
            self.open_online_ocr_instructions()
        else:
            # User declined online OCR: use a sensible default and enable reading
            default_wpp = 150
            self.avg_words_per_page = default_wpp
            self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (default)")
            self.analysis_group.setVisible(True)
            self.reading_group.setVisible(True)
            self.btn_start_reading.setEnabled(True)
            # Cache default value for persistence
            self.cache_analysis_results("default")

    def open_online_ocr_instructions(self):
        """Show localized steps to use the online Colab OCR and offer to open it"""
        url = "https://colab.research.google.com/github/ieasybooks/tahweel/blob/main/colab_notebook.ipynb#scrollTo=Uf8p_j7vXHTg"
        steps = self.localization.get_text("online_ocr_instructions")
        title = self.localization.get_text("online_ocr_instructions_title")
        reply = QMessageBox.question(
            self,
            title,
            steps + "\n\n" + self.localization.get_text("open_colab") + "?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply == QMessageBox.Yes:
            QDesktopServices.openUrl(QUrl(url))

    def import_ocr_results_for_current_pdf(self):
        """Find TXT/DOCX for current PDF (Downloads or manual), move to preferred folder, and update analysis."""
        if not self.pdf_path:
            QMessageBox.information(self, self.localization.get_text("information"), self.localization.get_text("select_file_first"))
            return

        import os, shutil
        from pathlib import Path
        from PySide6.QtCore import QSettings

        base = os.path.splitext(os.path.basename(self.pdf_path))[0]
        downloads = Path.home() / "Downloads"
        txt_src = downloads / f"{base}.txt"
        docx_src = downloads / f"{base}.docx"

        # Allow manual pick if not found
        if not txt_src.exists() or not docx_src.exists():
            QMessageBox.information(self, self.localization.get_text("information"), self.localization.get_text("ocr_results_not_found"))
            txt_path, _ = QFileDialog.getOpenFileName(self, self.localization.get_text("select_txt_file"), str(downloads), "Text (*.txt)")
            docx_path, _ = QFileDialog.getOpenFileName(self, self.localization.get_text("select_docx_file"), str(downloads), "DOCX (*.docx)")
            txt_src = Path(txt_path) if txt_path else None
            docx_src = Path(docx_path) if docx_path else None
            if not txt_src or not docx_src:
                return
            if txt_src and txt_src.stem != base:
                QMessageBox.warning(self, self.localization.get_text("warning"), self.localization.get_text("file_name_mismatch"))
                return
            if docx_src and docx_src.stem != base:
                QMessageBox.warning(self, self.localization.get_text("warning"), self.localization.get_text("file_name_mismatch"))
                return

        # Destination: save OCR results alongside the original PDF
        dest_dir = Path(self.pdf_path).parent
        dest_dir.mkdir(parents=True, exist_ok=True)

        moved_txt = dest_dir / (txt_src.name if txt_src else f"{base}.txt")
        moved_docx = dest_dir / (docx_src.name if docx_src else f"{base}.docx")

        # Move/copy files
        for src, dst in ((txt_src, moved_txt), (docx_src, moved_docx)):
            try:
                if src and src.exists():
                    if src.resolve() != dst.resolve():
                        shutil.move(str(src), str(dst))
                else:
                    pass
            except Exception:
                try:
                    if src and src.exists():
                        shutil.copy2(str(src), str(dst))
                except Exception:
                    pass

        # Try to compute per-page word counts from DOCX if python-docx is available
        self.words_per_page_list = []
        try:
            counts = self.compute_words_per_page_from_docx(moved_docx)
            if counts and isinstance(counts, list):
                self.words_per_page_list = counts
        except Exception:
            self.words_per_page_list = []

        # Update analysis: prefer DOCX per-page counts; fallback to TXT average
        try:
            if self.words_per_page_list:
                if self.total_pages > 0 and len(self.words_per_page_list) != self.total_pages:
                    if len(self.words_per_page_list) > self.total_pages:
                        self.words_per_page_list = self.words_per_page_list[: self.total_pages]
                    else:
                        last = self.words_per_page_list[-1] if self.words_per_page_list else 0
                        self.words_per_page_list += [last] * (self.total_pages - len(self.words_per_page_list))
                if self.words_per_page_list:
                    self.avg_words_per_page = max(1, int(sum(self.words_per_page_list) / len(self.words_per_page_list)))
            else:
                with open(moved_txt, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                total_words = len([w for w in re.findall(r"\S+", text)])
                if self.total_pages > 0:
                    self.avg_words_per_page = max(1, total_words // max(1, self.total_pages))

            self.lbl_avg_words.setText(str(self.avg_words_per_page))
            self.cache_analysis_results("ocr")
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("ocr_results_moved"))
        except Exception as e:
            QMessageBox.warning(self, self.localization.get_text("warning"), str(e))

    def compute_words_per_page_from_docx(self, docx_path):
        """Attempt to compute per-page word counts by detecting page breaks in DOCX.
        Returns a list of word counts per page, or None if unavailable.
        """
        try:
            import docx  # python-docx (optional dependency)
        except Exception:
            return None
        try:
            document = docx.Document(str(docx_path))
            counts = []
            current = 0
            # Heuristic: count words paragraph by paragraph; split when a page break run appears
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    current += len([w for w in text.split() if w])
                for run in para.runs:
                    # Page break detection
                    try:
                        for br in run._element.xpath('.//w:br'):
                            br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                            if br_type == 'page':
                                counts.append(current)
                                current = 0
                    except Exception:
                        pass
            counts.append(current)
            # Remove trailing empty pages if any
            while len(counts) > 1 and counts[-1] == 0:
                counts.pop()
            return counts
        except Exception:
            return None

    def set_default_ocr_dir(self):
        from PySide6.QtCore import QSettings
        directory = QFileDialog.getExistingDirectory(self, self.localization.get_text("set_default_ocr_dir"))
        if directory:
            settings = QSettings("PDFToolsApp", "PDF Tools")
            settings.setValue("ocr/default_output_dir", directory)
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("preferred_ocr_dir_saved"))

    def generate_comprehension_questions(self):
        """Generate comprehension questions from sample text"""
        # This method can be implemented later if needed
        pass

    def check_comprehension(self):
        """Check comprehension answers"""
        # This method can be implemented later if needed
        pass

    def reset_ui(self):
        """Reset the UI for a new reading session while preserving historical data"""
        # Store current book path to preserve session history
        current_book_path = self.pdf_path

        # Reset current session state only
        self.pdf_path = None
        if self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None
        self.avg_words_per_page = 0
        self.total_pages = 0
        self.start_time = 0
        self.elapsed_time_seconds = 0
        self.sample_text = ""
        self.comprehension_questions = []
        self.words_per_page_list = []

        # Stop timer
        if hasattr(self, 'reading_timer'):
            self.reading_timer.stop()

        # Reset UI elements for current session
        self.pdf_display.clear()
        self.lbl_avg_words.setText("---")
        self.lbl_total_pages.setText("---")
        self.lbl_wpm.setText(f"--- {self.localization.get_text('wpm')}")
        self.lbl_total_time.setText("---")
        self.lbl_recommendation.setText("...")
        self.lbl_timer.setText("00:00")
        self.spin_pages_read.setValue(1)

        # Hide current session sections
        self.analysis_group.setVisible(False)
        self.reading_group.setVisible(False)
        self.results_group.setVisible(False)
        self.comprehension_group.setVisible(False)
        self.finish_widget.setVisible(False)
        self.btn_start_reading.setVisible(True)
        self.btn_start_reading.setEnabled(False)
        self.btn_prepare_book.setEnabled(False)

        # Hide book card and sessions dashboard
        if hasattr(self, 'book_card'):
            self.book_card.setVisible(False)
        if hasattr(self, 'sessions_dashboard'):
            self.sessions_dashboard.setVisible(False)

        # Clear current comprehension layout
        for i in reversed(range(self.comprehension_layout.count())):
            child = self.comprehension_layout.itemAt(i).widget()
            if child:
                child.setParent(None)

        # Refresh recent books list (preserves historical data)
        if hasattr(self, 'load_recent_books'):
            self.load_recent_books()


class BookSessionDetailsDialog(QDialog):
    """Dialog to show detailed session information for a book"""

    def __init__(self, book_data, recent_books_manager, parent=None):
        super().__init__(parent)
        self.book_data = book_data
        self.recent_books_manager = recent_books_manager
        self.setWindowTitle(f"Session Details - {book_data.get('display_name', 'Unknown Book')}")
        self.setMinimumSize(800, 600)
        self.setModal(True)

        self.setup_ui()
        self.load_session_data()

    def setup_ui(self):
        """Setup the dialog UI"""
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header with book info
        header_widget = self.create_header_section()
        layout.addWidget(header_widget)

        # Tabs for different views
        tab_widget = QTabWidget()

        # Session History tab
        sessions_tab = self.create_sessions_tab()
        tab_widget.addTab(sessions_tab, "📊 Session History")

        # Book Details tab
        details_tab = self.create_details_tab()
        tab_widget.addTab(details_tab, "📖 Book Details")

        # Statistics tab
        stats_tab = self.create_statistics_tab()
        tab_widget.addTab(stats_tab, "📈 Statistics")

        layout.addWidget(tab_widget)

        # Buttons
        button_layout = QHBoxLayout()

        # Remove Configuration button
        remove_btn = QPushButton("🗑️ Remove Configuration")
        remove_btn.clicked.connect(self.remove_configuration)
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        button_layout.addWidget(remove_btn)

        button_layout.addStretch()

        # Close button
        close_btn = QPushButton(self.localization.get_text("close_button"))
        close_btn.clicked.connect(self.accept)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)

    def create_header_section(self):
        """Create header section with book overview"""
        header_widget = QFrame()
        header_widget.setStyleSheet("""
            QFrame {
                background-color: #f5f5f5;
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 15px;
            }
        """)

        layout = QHBoxLayout(header_widget)

        # Book title and basic info
        info_layout = QVBoxLayout()

        title = QLabel(self.book_data.get('display_name', 'Unknown Book'))
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setStyleSheet("color: #333;")
        info_layout.addWidget(title)

        file_path = self.book_data['file_path']
        path_label = QLabel(f"📁 {file_path}")
        path_label.setStyleSheet("color: #666; font-size: 11px;")
        path_label.setWordWrap(True)
        info_layout.addWidget(path_label)

        layout.addLayout(info_layout)
        layout.addStretch()

        # Quick stats
        stats = self.book_data.get('statistics', {})
        if stats:
            stats_layout = QVBoxLayout()

            sessions_label = QLabel(f"📊 {stats.get('total_sessions', 0)} Sessions")
            sessions_label.setStyleSheet("color: #2196F3; font-weight: bold;")
            stats_layout.addWidget(sessions_label)

            if stats.get('best_wpm', 0) > 0:
                best_label = QLabel(f"🏆 Best: {stats['best_wpm']:.0f} WPM")
                best_label.setStyleSheet("color: #FF9800; font-weight: bold;")
                stats_layout.addWidget(best_label)

            layout.addLayout(stats_layout)

        return header_widget

    def create_sessions_tab(self):
        """Create sessions history tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Sessions table
        self.sessions_table = QTableWidget()
        # Apply theme-aware styling
        self.update_table_theme(self.sessions_table)
        self.sessions_table.setAlternatingRowColors(True)
        self.sessions_table.setSelectionBehavior(QTableWidget.SelectRows)

        layout.addWidget(self.sessions_table)
        return widget

    def create_details_tab(self):
        """Create book details tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Details in a form-like layout
        details_widget = QWidget()
        details_layout = QFormLayout(details_widget)
        details_layout.setSpacing(10)

        # Book configuration details
        details_layout.addRow("📄 Total Pages:", QLabel(str(self.book_data.get('total_pages', 'N/A'))))
        details_layout.addRow("📝 Words per Page:", QLabel(f"{self.book_data.get('avg_words_per_page', 0):.1f}"))
        details_layout.addRow("🔧 Preparation Method:", QLabel(self.book_data.get('preparation_method', 'Unknown').title()))
        details_layout.addRow("📅 Date Added:", QLabel(self.book_data.get('created_at', 'Unknown')))
        details_layout.addRow("🔄 Last Updated:", QLabel(self.book_data.get('updated_at', 'Unknown')))

        # File information
        file_path = self.book_data['file_path']
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            size_mb = file_size / (1024 * 1024)
            details_layout.addRow("💾 File Size:", QLabel(f"{size_mb:.1f} MB"))
        else:
            details_layout.addRow("💾 File Status:", QLabel("❌ File not found"))

        # Sample text if available
        sample_text = self.book_data.get('sample_text', '')
        if sample_text:
            sample_label = QLabel("📝 Sample Text:")
            sample_text_widget = QTextEdit()
            sample_text_widget.setPlainText(sample_text)
            sample_text_widget.setMaximumHeight(100)
            sample_text_widget.setReadOnly(True)
            details_layout.addRow(sample_label, sample_text_widget)

        layout.addWidget(details_widget)
        layout.addStretch()
        return widget

    def create_statistics_tab(self):
        """Create statistics tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        stats = self.book_data.get('statistics', {})

        if not stats or stats.get('total_sessions', 0) == 0:
            no_stats_label = QLabel("📊 No session statistics available yet.\n\nComplete some reading sessions to see detailed statistics here.")
            no_stats_label.setAlignment(Qt.AlignCenter)
            no_stats_label.setStyleSheet("color: #666; font-size: 14px; padding: 40px;")
            layout.addWidget(no_stats_label)
            return widget

        # Statistics cards
        stats_scroll = QScrollArea()
        stats_scroll.setWidgetResizable(True)

        stats_widget = QWidget()
        stats_layout = QVBoxLayout(stats_widget)

        # Overall performance
        perf_group = QGroupBox("📈 Overall Performance")
        perf_layout = QFormLayout(perf_group)

        perf_layout.addRow("📊 Total Sessions:", QLabel(str(stats.get('total_sessions', 0))))
        perf_layout.addRow("📄 Total Pages Read:", QLabel(str(stats.get('total_pages_read', 0))))
        perf_layout.addRow("⏱️ Total Reading Time:", QLabel(self.format_time_duration(stats.get('total_time_seconds', 0))))
        perf_layout.addRow("⚡ Average Speed:", QLabel(f"{stats.get('average_wpm', 0):.1f} WPM"))
        perf_layout.addRow("🏆 Best Speed:", QLabel(f"{stats.get('best_wpm', 0):.1f} WPM"))

        if stats.get('last_session_date'):
            perf_layout.addRow("🕒 Last Session:", QLabel(stats['last_session_date']))

        stats_layout.addWidget(perf_group)

        # Progress tracking
        if stats.get('sessions'):
            progress_group = QGroupBox("📊 Progress Tracking")
            progress_layout = QVBoxLayout(progress_group)

            sessions = stats['sessions']
            if len(sessions) >= 2:
                first_session = sessions[0]
                last_session = sessions[-1]

                improvement = last_session.get('words_per_minute', 0) - first_session.get('words_per_minute', 0)
                improvement_text = f"{improvement:+.1f} WPM"
                if improvement > 0:
                    improvement_text += " 📈"
                elif improvement < 0:
                    improvement_text += " 📉"
                else:
                    improvement_text += " ➡️"

                progress_info = QLabel(f"Improvement from first to last session: {improvement_text}")
                progress_info.setStyleSheet("font-weight: bold; color: #2196F3;")
                progress_layout.addWidget(progress_info)

            stats_layout.addWidget(progress_group)

        stats_scroll.setWidget(stats_widget)
        layout.addWidget(stats_scroll)

        return widget

    def format_time_duration(self, seconds):
        """Format time duration in human readable format"""
        if seconds < 60:
            return f"{seconds:.0f} seconds"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f} minutes"
        else:
            hours = seconds / 3600
            return f"{hours:.1f} hours"

    def load_session_data(self):
        """Load and display session data in the table"""
        stats = self.book_data.get('statistics', {})
        sessions = stats.get('sessions', [])

        if not sessions:
            self.sessions_table.setRowCount(1)
            self.sessions_table.setColumnCount(1)
            self.sessions_table.setHorizontalHeaderLabels(["Status"])

            no_data_item = QTableWidgetItem("No session data available")
            no_data_item.setTextAlignment(Qt.AlignCenter)
            self.sessions_table.setItem(0, 0, no_data_item)
            return

        # Setup table
        self.sessions_table.setRowCount(len(sessions))
        self.sessions_table.setColumnCount(6)
        self.sessions_table.setHorizontalHeaderLabels([
            "Date", "Pages Read", "Duration", "WPM", "Accuracy", "Notes"
        ])

        # Populate table
        for row, session in enumerate(sessions):
            # Date
            date_item = QTableWidgetItem(session.get('session_date', 'Unknown'))
            self.sessions_table.setItem(row, 0, date_item)

            # Pages read
            pages_item = QTableWidgetItem(str(session.get('pages_read', 0)))
            pages_item.setTextAlignment(Qt.AlignCenter)
            self.sessions_table.setItem(row, 1, pages_item)

            # Duration
            duration = session.get('session_duration', 0)
            duration_text = self.format_time_duration(duration)
            duration_item = QTableWidgetItem(duration_text)
            duration_item.setTextAlignment(Qt.AlignCenter)
            self.sessions_table.setItem(row, 2, duration_item)

            # WPM
            wpm = session.get('words_per_minute', 0)
            wpm_item = QTableWidgetItem(f"{wpm:.1f}")
            wpm_item.setTextAlignment(Qt.AlignCenter)
            self.sessions_table.setItem(row, 3, wpm_item)

            # Accuracy (if available)
            accuracy = session.get('accuracy', 'N/A')
            accuracy_item = QTableWidgetItem(str(accuracy))
            accuracy_item.setTextAlignment(Qt.AlignCenter)
            self.sessions_table.setItem(row, 4, accuracy_item)

            # Notes
            notes = session.get('notes', '')
            notes_item = QTableWidgetItem(notes)
            self.sessions_table.setItem(row, 5, notes_item)

        # Resize columns to content
        self.sessions_table.resizeColumnsToContents()

    def remove_configuration(self):
        """Remove the reading speed configuration for this book"""
        reply = QMessageBox.question(
            self,
            "Remove Configuration",
            f"Are you sure you want to remove the reading speed configuration for this book?\n\n"
            f"Book: {self.book_data.get('display_name', 'Unknown')}\n\n"
            f"This will delete all session data and configuration settings.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            if self.recent_books_manager.delete_reading_speed_config(self.book_data['file_path']):
                QMessageBox.information(
                    self,
                    "Configuration Removed",
                    "The reading speed configuration has been removed successfully."
                )
                self.accept()  # Close dialog
            else:
                QMessageBox.warning(
                    self,
                    "Error",
                    "Failed to remove the reading speed configuration."
                )

    def create_overall_stats_section(self, layout):
        """Create overall reading statistics section"""
        stats_group = QGroupBox("📈 Overall Reading Statistics")
        stats_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #1976D2;
                border: 2px solid #2196F3;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: #fafafa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: #1976D2;
                font-weight: bold;
                background-color: #fafafa;
            }
        """)

        stats_layout = QGridLayout(stats_group)

        # Create statistics labels
        self.total_books_label = QLabel("0")
        wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"
        min_text = self.localization.get_text("minutes") if self.localization else "min"

        self.total_sessions_label = QLabel("0")
        self.total_time_label = QLabel(f"0 {min_text}")
        self.avg_speed_label = QLabel(f"0 {wpm_text}")
        self.best_speed_label = QLabel(f"0 {wpm_text}")
        self.total_pages_label = QLabel("0")

        # Style the statistics labels
        stat_style = """
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #2196F3;
                background-color: #e3f2fd;
                padding: 10px;
                border-radius: 8px;
                border: 1px solid #bbdefb;
                text-align: center;
            }
        """

        for label in [self.total_books_label, self.total_sessions_label, self.total_time_label,
                     self.avg_speed_label, self.best_speed_label, self.total_pages_label]:
            label.setStyleSheet(stat_style)
            label.setAlignment(Qt.AlignCenter)

        # Add to grid
        stats_layout.addWidget(QLabel("📚 Total Books:"), 0, 0)
        stats_layout.addWidget(self.total_books_label, 0, 1)
        stats_layout.addWidget(QLabel("📖 Total Sessions:"), 0, 2)
        stats_layout.addWidget(self.total_sessions_label, 0, 3)

        stats_layout.addWidget(QLabel("⏱️ Total Time:"), 1, 0)
        stats_layout.addWidget(self.total_time_label, 1, 1)
        stats_layout.addWidget(QLabel("📊 Average Speed:"), 1, 2)
        stats_layout.addWidget(self.avg_speed_label, 1, 3)

        stats_layout.addWidget(QLabel("🏆 Best Speed:"), 2, 0)
        stats_layout.addWidget(self.best_speed_label, 2, 1)
        stats_layout.addWidget(QLabel("📄 Total Pages:"), 2, 2)
        stats_layout.addWidget(self.total_pages_label, 2, 3)

        layout.addWidget(stats_group)

    def create_books_performance_section(self, layout):
        """Create books performance comparison section"""
        books_group = QGroupBox("📚 Books Performance Comparison")
        books_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #388E3C;
                border: 2px solid #4CAF50;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: #fafafa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: #388E3C;
                font-weight: bold;
                background-color: #fafafa;
            }
        """)

        books_layout = QVBoxLayout(books_group)

        # Books performance table
        self.books_table = QTableWidget()
        self.books_table.setColumnCount(6)
        self.books_table.setHorizontalHeaderLabels([
            "📖 Book", "📊 Sessions", "⚡ Avg WPM", "🏆 Best WPM", "⏱️ Total Time", "📄 Pages Read"
        ])

        # Apply theme-aware styling
        self.update_table_theme(self.books_table)

        self.books_table.setAlternatingRowColors(True)
        self.books_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.books_table.horizontalHeader().setStretchLastSection(True)
        self.books_table.setMaximumHeight(300)  # Increased for better visibility

        books_layout.addWidget(self.books_table)
        layout.addWidget(books_group)

    def create_reading_history_section(self, layout):
        """Create comprehensive reading history section"""
        history_group = QGroupBox("📅 Reading History")
        history_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #F57C00;
                border: 2px solid #FF9800;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: #fafafa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: #F57C00;
                font-weight: bold;
                background-color: #fafafa;
            }
        """)

        history_layout = QVBoxLayout(history_group)

        # History table
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(6)
        self.history_table.setHorizontalHeaderLabels([
            "📅 Date", "📖 Book", "📄 Pages", "⏱️ Duration", "⚡ WPM", "🎯 Score"
        ])

        # Apply theme-aware styling
        self.update_table_theme(self.history_table)

        self.history_table.setAlternatingRowColors(True)
        self.history_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.history_table.horizontalHeader().setStretchLastSection(True)
        self.history_table.setMaximumHeight(350)  # Increased for better visibility

        history_layout.addWidget(self.history_table)
        layout.addWidget(history_group)

    def create_reading_trends_section(self, layout):
        """Create reading trends and insights section"""
        trends_group = QGroupBox("📈 Reading Trends & Insights")
        trends_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 16px;
                color: #7B1FA2;
                border: 2px solid #9C27B0;
                border-radius: 10px;
                margin: 10px 0;
                padding: 15px;
                background-color: #fafafa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px 0 10px;
                color: #7B1FA2;
                font-weight: bold;
                background-color: #fafafa;
            }
        """)

        trends_layout = QVBoxLayout(trends_group)

        # Trends content
        self.trends_content = QLabel(self.localization.get_text("loading_trends"))
        self.trends_content.setWordWrap(True)
        # Apply theme-aware styling for trends content
        colors = self.get_theme_colors()
        self.trends_content.setStyleSheet(f"""
            QLabel {{
                font-size: 14px;
                color: {colors['text_primary']};
                padding: 15px;
                background-color: {colors['bg_secondary']};
                border-radius: 8px;
                border: 1px solid {colors['border']};
                line-height: 1.6;
            }}
        """)

        trends_layout.addWidget(self.trends_content)
        layout.addWidget(trends_group)

    def create_book_card(self):
        """Create a comprehensive book information display with three organized sections"""
        # Main container with responsive width
        card = QGroupBox()
        card.setVisible(False)  # Initially hidden
        card.setMaximumWidth(900)  # Maximum width for large screens
        card.setMinimumWidth(400)  # Minimum width for usability
        # Apply theme-aware styling for book card
        colors = self.get_theme_colors()
        card.setStyleSheet(f"""
            QGroupBox {{
                border: 3px solid {colors['accent_blue']};
                border-radius: 15px;
                margin: 15px;
                padding: 20px;
                background-color: {colors['bg_primary']};
                font-weight: bold;
                font-size: 16px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 15px 0 15px;
                color: {colors['accent_blue']};
                font-weight: bold;
                font-size: 18px;
            }}
        """)

        main_layout = QVBoxLayout(card)

        # Section 1 - Book Overview
        overview_section = QFrame()
        overview_section.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 2px solid #e9ecef;
                border-radius: 10px;
                padding: 15px;
                margin: 5px 0;
            }
        """)
        overview_layout = QHBoxLayout(overview_section)

        # Book cover
        self.book_cover = QLabel()
        self.book_cover.setFixedSize(100, 130)
        self.book_cover.setStyleSheet("""
            QLabel {
                border: 2px solid #ddd;
                border-radius: 8px;
                background-color: #e3f2fd;
                color: #666;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        self.book_cover.setAlignment(Qt.AlignCenter)
        self.book_cover.setText("📖\nCover")
        overview_layout.addWidget(self.book_cover)

        # Book basic info
        basic_info_layout = QVBoxLayout()

        # Title
        self.book_title = QLabel(self.localization.get_text("book_title_placeholder"))
        self.book_title.setStyleSheet("""
            QLabel {
                font-size: 20px;
                font-weight: bold;
                color: #1976D2;
                margin-bottom: 8px;
            }
        """)
        self.book_title.setWordWrap(True)
        basic_info_layout.addWidget(self.book_title)

        # Filename
        self.book_filename = QLabel(self.localization.get_text("filename_placeholder"))
        self.book_filename.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #666;
                margin-bottom: 15px;
                font-style: italic;
            }
        """)
        basic_info_layout.addWidget(self.book_filename)

        # Overview metrics grid
        overview_grid = QGridLayout()

        # Total pages
        pages_label = QLabel("📄 Total Pages:")
        pages_label.setStyleSheet("font-size: 14px; color: #424242; font-weight: bold;")
        self.book_pages = QLabel("---")
        self.book_pages.setStyleSheet("font-size: 14px; color: #2196F3; font-weight: bold;")
        overview_grid.addWidget(pages_label, 0, 0)
        overview_grid.addWidget(self.book_pages, 0, 1)

        # File size
        size_label = QLabel("💾 File Size:")
        size_label.setStyleSheet("font-size: 14px; color: #424242; font-weight: bold;")
        self.book_size = QLabel("---")
        self.book_size.setStyleSheet("font-size: 14px; color: #4CAF50; font-weight: bold;")
        overview_grid.addWidget(size_label, 1, 0)
        overview_grid.addWidget(self.book_size, 1, 1)

        # Reading progress
        progress_label = QLabel("📊 Progress:")
        progress_label.setStyleSheet("font-size: 14px; color: #424242; font-weight: bold;")
        self.book_progress = QLabel("---")
        self.book_progress.setStyleSheet("font-size: 14px; color: #FF9800; font-weight: bold;")
        overview_grid.addWidget(progress_label, 2, 0)
        overview_grid.addWidget(self.book_progress, 2, 1)

        basic_info_layout.addLayout(overview_grid)
        basic_info_layout.addStretch()

        overview_layout.addLayout(basic_info_layout)
        overview_layout.addStretch()

        main_layout.addWidget(overview_section)

        # Section 2 - Reading Analytics
        analytics_section = QFrame()
        analytics_section.setStyleSheet("""
            QFrame {
                background-color: #e8f5e8;
                border: 2px solid #c8e6c9;
                border-radius: 10px;
                padding: 15px;
                margin: 5px 0;
            }
        """)
        analytics_layout = QVBoxLayout(analytics_section)

        analytics_title = QLabel("📈 Reading Analytics")
        analytics_title.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #2E7D32;
                margin-bottom: 10px;
            }
        """)
        analytics_layout.addWidget(analytics_title)

        analytics_grid = QGridLayout()

        # Average words per page
        avg_words_label = QLabel("📝 Avg Words/Page:")
        avg_words_label.setStyleSheet("font-size: 13px; color: #424242; font-weight: bold;")
        self.analytics_avg_words = QLabel("---")
        self.analytics_avg_words.setStyleSheet("font-size: 13px; color: #2E7D32; font-weight: bold;")
        analytics_grid.addWidget(avg_words_label, 0, 0)
        analytics_grid.addWidget(self.analytics_avg_words, 0, 1)

        # Last session speed
        wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"

        last_speed_label = QLabel("⚡ Last Session:")
        last_speed_label.setStyleSheet("font-size: 13px; color: #424242; font-weight: bold;")
        self.analytics_last_speed = QLabel(f"--- {wpm_text}")
        self.analytics_last_speed.setStyleSheet("font-size: 13px; color: #1976D2; font-weight: bold;")
        analytics_grid.addWidget(last_speed_label, 0, 2)
        analytics_grid.addWidget(self.analytics_last_speed, 0, 3)

        # Best performance
        best_speed_label = QLabel("🏆 Best Speed:")
        best_speed_label.setStyleSheet("font-size: 13px; color: #424242; font-weight: bold;")
        self.analytics_best_speed = QLabel(f"--- {wpm_text}")
        self.analytics_best_speed.setStyleSheet("font-size: 13px; color: #FF9800; font-weight: bold;")
        analytics_grid.addWidget(best_speed_label, 1, 0)
        analytics_grid.addWidget(self.analytics_best_speed, 1, 1)

        # Total reading time
        total_time_label = QLabel("⏱️ Total Time:")
        total_time_label.setStyleSheet("font-size: 13px; color: #424242; font-weight: bold;")
        self.analytics_total_time = QLabel("--- min")
        self.analytics_total_time.setStyleSheet("font-size: 13px; color: #9C27B0; font-weight: bold;")
        analytics_grid.addWidget(total_time_label, 1, 2)
        analytics_grid.addWidget(self.analytics_total_time, 1, 3)

        analytics_layout.addLayout(analytics_grid)
        main_layout.addWidget(analytics_section)

        # Section 3 - Sessions History (will be created separately)
        self.create_sessions_history_section(main_layout)

        return card

    def create_sessions_history_section(self, main_layout):
        """Create Section 3 - Sessions History"""
        history_section = QFrame()
        history_section.setStyleSheet("""
            QFrame {
                background-color: #fff3e0;
                border: 2px solid #ffcc02;
                border-radius: 10px;
                padding: 15px;
                margin: 5px 0;
            }
        """)
        history_layout = QVBoxLayout(history_section)

        history_title = QLabel("📚 Reading Sessions History")
        history_title.setStyleSheet("""
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #E65100;
                margin-bottom: 10px;
            }
        """)
        history_layout.addWidget(history_title)

        # Sessions table
        self.book_sessions_table = QTableWidget()
        self.book_sessions_table.setColumnCount(5)
        self.book_sessions_table.setHorizontalHeaderLabels([
            "📅 Date", "📖 Pages", "⏱️ Time", "⚡ WPM", "🎯 Score"
        ])

        # Apply theme-aware styling
        self.update_table_theme(self.book_sessions_table)

        self.book_sessions_table.setAlternatingRowColors(True)
        self.book_sessions_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.book_sessions_table.horizontalHeader().setStretchLastSection(True)
        self.book_sessions_table.setMaximumHeight(150)
        self.book_sessions_table.setMinimumHeight(100)

        # Set column widths
        header = self.book_sessions_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)  # Date
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)  # Pages
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)  # Time
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)  # WPM
        header.setSectionResizeMode(4, QHeaderView.Stretch)           # Score

        history_layout.addWidget(self.book_sessions_table)
        main_layout.addWidget(history_section)

    def update_book_card(self, file_path):
        """Update the comprehensive book card with all three sections"""
        try:
            # Get basic file info
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)

            # Format file size
            if file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / (1024 * 1024):.1f} MB"

            # Section 1 - Book Overview Updates
            self.book_title.setText(filename.replace('.pdf', ''))
            self.book_filename.setText(filename)
            self.book_pages.setText(str(self.total_pages))
            self.book_size.setText(size_str)

            # Try to get reading progress from recent books manager
            progress_text = "0%"
            if self.recent_books_manager:
                try:
                    book = self.recent_books_manager.get_book_by_path(file_path)
                    if book:
                        progress_text = f"{book.reading_percentage:.1f}%"
                        # Update card title with book display name if available
                        if hasattr(book, 'display_name') and book.display_name:
                            self.book_title.setText(book.display_name)
                        elif hasattr(book, 'title') and book.title:
                            self.book_title.setText(book.title)
                except Exception as e:
                    print(f"Error getting book progress: {e}")

            self.book_progress.setText(progress_text)

            # Section 2 - Reading Analytics Updates
            self.analytics_avg_words.setText(f"{self.avg_words_per_page:.0f}" if self.avg_words_per_page > 0 else "---")

            # Get reading statistics for analytics
            if self.recent_books_manager:
                try:
                    wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"
                    min_text = self.localization.get_text("minutes") if self.localization else "min"
                    hours_text = self.localization.get_text("hours") if self.localization else "h"

                    stats = self.recent_books_manager.get_reading_statistics(file_path)
                    if stats and stats.get('total_sessions', 0) > 0:
                        # Get last session data
                        sessions = stats.get('sessions', [])
                        if sessions:
                            last_session = sessions[0]  # Most recent session
                            self.analytics_last_speed.setText(f"{last_session['words_per_minute']:.0f} {wpm_text}")
                        else:
                            self.analytics_last_speed.setText(f"--- {wpm_text}")

                        # Best speed
                        best_wpm = stats.get('best_wpm', 0)
                        self.analytics_best_speed.setText(f"{best_wpm:.0f} {wpm_text}" if best_wpm > 0 else f"--- {wpm_text}")

                        # Total time
                        total_seconds = stats.get('total_time_seconds', 0)
                        if total_seconds > 0:
                            total_minutes = total_seconds // 60
                            if total_minutes < 60:
                                time_str = f"{total_minutes} {min_text}"
                            else:
                                hours = total_minutes // 60
                                minutes = total_minutes % 60
                                time_str = f"{hours} {hours_text} {minutes} {min_text}"
                            self.analytics_total_time.setText(time_str)
                        else:
                            self.analytics_total_time.setText(f"--- {min_text}")

                        # Section 3 - Sessions History Updates
                        self.update_sessions_history_table(sessions)
                    else:
                        # No sessions yet
                        self.analytics_last_speed.setText(f"--- {wpm_text}")
                        self.analytics_best_speed.setText(f"--- {wpm_text}")
                        self.analytics_total_time.setText("--- min")
                        self.book_sessions_table.setRowCount(0)

                except Exception as e:
                    print(f"Error getting reading statistics: {e}")

            # Try to load thumbnail if available
            self.load_book_thumbnail(file_path)

            # Show the book card
            self.book_card.setVisible(True)
            current_book_text = self.localization.get_text('current_book') if self.localization else 'Current Book'
            self.book_card.setTitle(f"📖 {current_book_text}")

        except Exception as e:
            print(f"Error updating book card: {e}")

    def update_sessions_history_table(self, sessions):
        """Update the sessions history table in Section 3"""
        try:
            self.book_sessions_table.setRowCount(len(sessions))

            for row, session in enumerate(sessions):
                # Format date
                try:
                    from datetime import datetime
                    date_obj = datetime.fromisoformat(session['session_date'])
                    date_str = date_obj.strftime("%m/%d %H:%M")
                except:
                    date_str = session['session_date'][:10]

                # Format time
                duration = session['session_duration']
                if duration < 60:
                    time_str = f"{duration}s"
                else:
                    minutes = duration // 60
                    seconds = duration % 60
                    time_str = f"{minutes}:{seconds:02d}"

                # Add items to table
                self.book_sessions_table.setItem(row, 0, QTableWidgetItem(date_str))
                self.book_sessions_table.setItem(row, 1, QTableWidgetItem(str(session['pages_read'])))
                self.book_sessions_table.setItem(row, 2, QTableWidgetItem(time_str))
                self.book_sessions_table.setItem(row, 3, QTableWidgetItem(f"{session['words_per_minute']:.0f}"))
                self.book_sessions_table.setItem(row, 4, QTableWidgetItem(f"{session['comprehension_score']:.0f}%"))

                # Style the items
                for col in range(5):
                    item = self.book_sessions_table.item(row, col)
                    if item:
                        item.setTextAlignment(Qt.AlignCenter)

        except Exception as e:
            print(f"Error updating sessions history table: {e}")

    def update_dashboard(self):
        """Update dashboard with current statistics"""
        if not self.recent_books_manager:
            return

        try:
            # Refresh the general statistics cards
            self.refresh_general_stats_cards()

            # Update current book information if available
            if hasattr(self, 'book_card') and self.book_card and hasattr(self, 'pdf_path') and self.pdf_path:
                self.update_book_card()

            # Update sessions dashboard if available
            if hasattr(self, 'sessions_dashboard') and self.sessions_dashboard:
                self.update_sessions_dashboard()

        except Exception as e:
            print(f"Error updating dashboard: {e}")

    def refresh_general_stats_cards(self):
        """Refresh the general statistics cards with current data"""
        try:
            # Get updated statistics
            overall_stats = self.get_overall_statistics()

            # Update card values (this would require storing references to the value labels)
            # For now, we'll recreate the dashboard content when needed
            # This is a simple approach that ensures data is always current
            pass

        except Exception as e:
            print(f"Error refreshing general stats cards: {e}")

    def update_overall_stats(self, all_sessions, all_books):
        """Update overall statistics section"""
        try:
            total_books = len(all_books)
            total_sessions = len(all_sessions)

            if total_sessions > 0:
                total_time_seconds = sum(s['session_duration'] for s in all_sessions)
                total_time_minutes = total_time_seconds // 60

                if total_time_minutes < 60:
                    time_str = f"{total_time_minutes} min"
                else:
                    hours = total_time_minutes // 60
                    minutes = total_time_minutes % 60
                    time_str = f"{hours}h {minutes}m"

                avg_wpm = sum(s['words_per_minute'] for s in all_sessions) / total_sessions
                best_wpm = max(s['words_per_minute'] for s in all_sessions)
                total_pages = sum(s['pages_read'] for s in all_sessions)

                wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"
                min_text = self.localization.get_text("minutes") if self.localization else "min"

                self.total_books_label.setText(str(total_books))
                self.total_sessions_label.setText(str(total_sessions))
                self.total_time_label.setText(time_str)
                self.avg_speed_label.setText(f"{avg_wpm:.0f} {wpm_text}")
                self.best_speed_label.setText(f"{best_wpm:.0f} {wpm_text}")
                self.total_pages_label.setText(str(total_pages))
            else:
                # No sessions yet
                wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"
                min_text = self.localization.get_text("minutes") if self.localization else "min"

                self.total_books_label.setText("0")
                self.total_sessions_label.setText("0")
                self.total_time_label.setText(f"0 {min_text}")
                self.avg_speed_label.setText(f"0 {wpm_text}")
                self.best_speed_label.setText(f"0 {wpm_text}")
                self.total_pages_label.setText("0")

        except Exception as e:
            print(f"Error updating overall stats: {e}")

    def update_books_performance_table(self, all_books):
        """Update books performance comparison table"""
        try:
            self.books_table.setRowCount(len(all_books))

            for row, (book_path, book_data) in enumerate(all_books.items()):
                # Format time
                total_seconds = book_data['total_time']
                if total_seconds < 60:
                    time_str = f"{total_seconds}s"
                else:
                    minutes = total_seconds // 60
                    if minutes < 60:
                        time_str = f"{minutes}m"
                    else:
                        hours = minutes // 60
                        mins = minutes % 60
                        time_str = f"{hours}h {mins}m"

                # Add items to table
                self.books_table.setItem(row, 0, QTableWidgetItem(book_data['title']))
                self.books_table.setItem(row, 1, QTableWidgetItem(str(book_data['sessions'])))
                self.books_table.setItem(row, 2, QTableWidgetItem(f"{book_data['avg_wpm']:.0f}"))
                self.books_table.setItem(row, 3, QTableWidgetItem(f"{book_data['best_wpm']:.0f}"))
                self.books_table.setItem(row, 4, QTableWidgetItem(time_str))
                self.books_table.setItem(row, 5, QTableWidgetItem(str(book_data['total_pages'])))

                # Center align all items
                for col in range(6):
                    item = self.books_table.item(row, col)
                    if item:
                        item.setTextAlignment(Qt.AlignCenter)

        except Exception as e:
            print(f"Error updating books performance table: {e}")

    def update_reading_history_table(self, all_sessions):
        """Update comprehensive reading history table"""
        try:
            # Sort sessions by date (most recent first)
            sorted_sessions = sorted(all_sessions, key=lambda x: x['session_date'], reverse=True)

            # Limit to last 50 sessions for performance
            display_sessions = sorted_sessions[:50]

            self.history_table.setRowCount(len(display_sessions))

            for row, session in enumerate(display_sessions):
                # Format date
                try:
                    from datetime import datetime
                    date_obj = datetime.fromisoformat(session['session_date'])
                    date_str = date_obj.strftime("%m/%d %H:%M")
                except:
                    date_str = session['session_date'][:10]

                # Get book name
                book_name = os.path.basename(session['book_path']).replace('.pdf', '')
                if len(book_name) > 20:
                    book_name = book_name[:17] + "..."

                # Format duration
                duration = session['session_duration']
                if duration < 60:
                    time_str = f"{duration}s"
                else:
                    minutes = duration // 60
                    seconds = duration % 60
                    time_str = f"{minutes}:{seconds:02d}"

                # Add items to table
                self.history_table.setItem(row, 0, QTableWidgetItem(date_str))
                self.history_table.setItem(row, 1, QTableWidgetItem(book_name))
                self.history_table.setItem(row, 2, QTableWidgetItem(str(session['pages_read'])))
                self.history_table.setItem(row, 3, QTableWidgetItem(time_str))
                self.history_table.setItem(row, 4, QTableWidgetItem(f"{session['words_per_minute']:.0f}"))
                self.history_table.setItem(row, 5, QTableWidgetItem(f"{session['comprehension_score']:.0f}%"))

                # Center align all items
                for col in range(6):
                    item = self.history_table.item(row, col)
                    if item:
                        item.setTextAlignment(Qt.AlignCenter)

        except Exception as e:
            print(f"Error updating reading history table: {e}")

    def update_trends_analysis(self, all_sessions):
        """Update reading trends and insights"""
        try:
            if not all_sessions:
                self.trends_content.setText("📊 No reading data available yet. Start reading to see your trends!")
                return

            # Calculate trends
            total_sessions = len(all_sessions)
            avg_wpm = sum(s['words_per_minute'] for s in all_sessions) / total_sessions
            best_wpm = max(s['words_per_minute'] for s in all_sessions)
            avg_comprehension = sum(s['comprehension_score'] for s in all_sessions) / total_sessions

            # Speed category
            if avg_wpm < 150:
                speed_category = "Developing Reader"
                speed_color = "#FF5722"
            elif avg_wpm < 250:
                speed_category = "Average Reader"
                speed_color = "#FF9800"
            elif avg_wpm < 350:
                speed_category = "Above Average Reader"
                speed_color = "#4CAF50"
            else:
                speed_category = "Speed Reader"
                speed_color = "#2196F3"

            # Recent performance (last 5 sessions)
            recent_sessions = sorted(all_sessions, key=lambda x: x['session_date'], reverse=True)[:5]
            if len(recent_sessions) >= 2:
                recent_avg = sum(s['words_per_minute'] for s in recent_sessions) / len(recent_sessions)
                older_sessions = all_sessions[:-len(recent_sessions)] if len(all_sessions) > len(recent_sessions) else []

                if older_sessions:
                    older_avg = sum(s['words_per_minute'] for s in older_sessions) / len(older_sessions)
                    improvement = ((recent_avg - older_avg) / older_avg) * 100

                    if improvement > 5:
                        trend_text = f"📈 <span style='color: #4CAF50;'>Improving (+{improvement:.1f}%)</span>"
                    elif improvement < -5:
                        trend_text = f"📉 <span style='color: #FF5722;'>Declining ({improvement:.1f}%)</span>"
                    else:
                        trend_text = f"📊 <span style='color: #FF9800;'>Stable ({improvement:+.1f}%)</span>"
                else:
                    trend_text = "📊 <span style='color: #666;'>Building baseline...</span>"
            else:
                trend_text = "📊 <span style='color: #666;'>Need more sessions for trend analysis</span>"

            # Create comprehensive analysis
            analysis_html = f"""
            <div style='line-height: 1.8;'>
                <h3 style='color: #1976D2; margin-bottom: 15px;'>📊 Your Reading Profile</h3>

                <p><strong>Reading Level:</strong> <span style='color: {speed_color}; font-weight: bold;'>{speed_category}</span></p>
                <p><strong>Average Speed:</strong> {avg_wpm:.0f} WPM</p>
                <p><strong>Best Performance:</strong> {best_wpm:.0f} WPM</p>
                <p><strong>Comprehension Average:</strong> {avg_comprehension:.1f}%</p>

                <h4 style='color: #1976D2; margin: 15px 0 10px 0;'>📈 Recent Trend</h4>
                <p>{trend_text}</p>

                <h4 style='color: #1976D2; margin: 15px 0 10px 0;'>💡 Insights</h4>
                <p>• You've completed <strong>{total_sessions}</strong> reading sessions</p>
                <p>• Your reading speed puts you in the <strong>{speed_category}</strong> category</p>
                {"<p>• Great job maintaining consistent performance! 🎉</p>" if -5 <= (improvement if 'improvement' in locals() else 0) <= 5 else ""}
                {"<p>• Keep up the excellent improvement trend! 🚀</p>" if 'improvement' in locals() and improvement > 5 else ""}
            </div>
            """

            self.trends_content.setText(analysis_html)

        except Exception as e:
            print(f"Error updating trends analysis: {e}")
            self.trends_content.setText("📊 Error loading trends analysis. Please try again.")

    def on_tab_changed(self, index):
        """Handle tab change events"""
        if index == 1:  # Activity tab (now index 1 since dashboard was removed)
            self.refresh_activity_books()

    def load_book_thumbnail(self, file_path):
        """Load book thumbnail if available"""
        try:
            if self.recent_books_manager:
                book = self.recent_books_manager.get_book_by_path(file_path)
                if book and hasattr(book, 'thumbnail_path') and book.thumbnail_path:
                    if os.path.exists(book.thumbnail_path):
                        pixmap = QPixmap(book.thumbnail_path)
                        if not pixmap.isNull():
                            # Scale pixmap to fit the cover label
                            scaled_pixmap = pixmap.scaled(
                                self.book_cover.size(),
                                Qt.KeepAspectRatio,
                                Qt.SmoothTransformation
                            )
                            self.book_cover.setPixmap(scaled_pixmap)
                            self.book_cover.setText("")  # Clear text when image is loaded
                            return

            # Fallback to default cover
            self.book_cover.clear()
            self.book_cover.setText("📖\nCover")

        except Exception as e:
            print(f"Error loading thumbnail: {e}")
            self.book_cover.clear()
            self.book_cover.setText("📖\nCover")

    def create_sessions_dashboard(self):
        """Create the reading sessions dashboard"""
        sessions_text = self.localization.get_text("reading_sessions") if self.localization else "Reading Sessions"
        dashboard = QGroupBox(sessions_text)
        dashboard.setVisible(False)  # Initially hidden
        dashboard.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #4CAF50;
                border-radius: 8px;
                margin: 10px 0;
                padding: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px 0 10px;
                color: #4CAF50;
                font-weight: bold;
            }
        """)

        layout = QVBoxLayout(dashboard)

        # Statistics summary
        stats_frame = QFrame()
        stats_frame.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                padding: 10px;
                margin: 5px 0;
            }
        """)
        stats_layout = QGridLayout(stats_frame)

        # Statistics labels
        wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"

        self.lbl_total_sessions = QLabel("0")
        self.lbl_total_sessions.setStyleSheet("font-size: 18px; font-weight: bold; color: #2196F3;")

        self.lbl_avg_speed = QLabel(f"0 {wpm_text}")
        self.lbl_avg_speed.setStyleSheet("font-size: 18px; font-weight: bold; color: #4CAF50;")

        self.lbl_best_speed = QLabel(f"0 {wpm_text}")
        self.lbl_best_speed.setStyleSheet("font-size: 16px; font-weight: bold; color: #FF9800;")

        self.lbl_total_time = QLabel("0 min")
        self.lbl_total_time.setStyleSheet("font-size: 16px; font-weight: bold; color: #9C27B0;")

        # Add to grid
        stats_layout.addWidget(QLabel(self.localization.get_text("total_sessions_label")), 0, 0)
        stats_layout.addWidget(self.lbl_total_sessions, 0, 1)
        stats_layout.addWidget(QLabel(self.localization.get_text("average_speed_label")), 0, 2)
        stats_layout.addWidget(self.lbl_avg_speed, 0, 3)

        stats_layout.addWidget(QLabel(self.localization.get_text("best_speed_label")), 1, 0)
        stats_layout.addWidget(self.lbl_best_speed, 1, 1)
        stats_layout.addWidget(QLabel(self.localization.get_text("total_time_label")), 1, 2)
        stats_layout.addWidget(self.lbl_total_time, 1, 3)

        layout.addWidget(stats_frame)

        # Sessions history table
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView

        self.sessions_table = QTableWidget()
        self.sessions_table.setColumnCount(5)
        self.sessions_table.setHorizontalHeaderLabels([
            "Date", "Pages", "Time", "WPM", "Score"
        ])

        # Apply theme-aware styling
        self.update_table_theme(self.sessions_table)

        self.sessions_table.setAlternatingRowColors(True)
        self.sessions_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.sessions_table.horizontalHeader().setStretchLastSection(True)
        self.sessions_table.setMaximumHeight(300)  # Increased for better visibility

        layout.addWidget(self.sessions_table)

        return dashboard

    def update_sessions_dashboard(self):
        """Update the sessions dashboard with latest data"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            stats = self.recent_books_manager.get_reading_statistics(self.pdf_path)
            if not stats:
                return

            # Update statistics labels
            wpm_text = self.localization.get_text("wpm") if self.localization else "WPM"

            # Get localized time units
            hours_text = self.localization.get_text('hours') if self.localization else 'hours'
            minutes_text = self.localization.get_text('minutes') if self.localization else 'minutes'

            self.lbl_total_sessions.setText(str(stats.get('total_sessions', 0)))
            self.lbl_avg_speed.setText(f"{stats.get('average_wpm', 0):.0f} {wpm_text}")
            self.lbl_best_speed.setText(f"{stats.get('best_wpm', 0):.0f} {wpm_text}")

            # Format total time with localized units
            total_seconds = stats.get('total_time_seconds', 0)
            total_minutes = total_seconds // 60
            if total_minutes < 60:
                time_str = f"{total_minutes} {minutes_text}"
            else:
                hours = total_minutes // 60
                minutes = total_minutes % 60
                time_str = f"{hours} {hours_text} {minutes} {minutes_text}"
            self.lbl_total_time.setText(time_str)

            # Update sessions table
            sessions = stats.get('sessions', [])
            self.sessions_table.setRowCount(len(sessions))

            for row, session in enumerate(sessions):
                # Format date
                try:
                    from datetime import datetime
                    date_obj = datetime.fromisoformat(session['session_date'])
                    date_str = date_obj.strftime("%m/%d %H:%M")
                except:
                    date_str = session['session_date'][:10]

                # Format time
                duration = session['session_duration']
                if duration < 60:
                    time_str = f"{duration}s"
                else:
                    minutes = duration // 60
                    seconds = duration % 60
                    time_str = f"{minutes}:{seconds:02d}"

                # Add items to table
                self.sessions_table.setItem(row, 0, QTableWidgetItem(date_str))
                self.sessions_table.setItem(row, 1, QTableWidgetItem(str(session['pages_read'])))
                self.sessions_table.setItem(row, 2, QTableWidgetItem(time_str))
                self.sessions_table.setItem(row, 3, QTableWidgetItem(f"{session['words_per_minute']:.0f}"))
                self.sessions_table.setItem(row, 4, QTableWidgetItem(f"{session['comprehension_score']:.0f}%"))

            # Show the dashboard
            self.sessions_dashboard.setVisible(True)

        except Exception as e:
            print(f"Error updating sessions dashboard: {e}")

    def load_recent_books_delayed(self):
        """Load recent books with delay to ensure all methods are available"""
        if hasattr(self, 'recent_books_list'):
            self.load_recent_books()

    def load_recent_books_list(self):
        """Load recent books list - alias for load_recent_books"""
        self.load_recent_books()

    def load_recent_books(self):
        """Load recent books from the manager"""
        self.recent_books_list.clear()

        if not self.recent_books_manager:
            item = QListWidgetItem(self.localization.get_text("no_recent_books"))
            item.setFlags(Qt.NoItemFlags)
            self.recent_books_list.addItem(item)
            return

        try:
            books = self.recent_books_manager.get_books(limit=10)
            if not books:
                item = QListWidgetItem(self.localization.get_text("no_recent_books"))
                item.setFlags(Qt.NoItemFlags)
                self.recent_books_list.addItem(item)
                return

            for book in books:
                display_name = book.display_name if hasattr(book, 'display_name') else os.path.basename(book.file_path)
                item = QListWidgetItem(f"📖 {display_name}")
                item.setData(Qt.UserRole, book.file_path)
                self.recent_books_list.addItem(item)

        except Exception as e:
            print(f"Error loading recent books: {e}")
            item = QListWidgetItem(self.localization.get_text("no_recent_books"))
            item.setFlags(Qt.NoItemFlags)
            self.recent_books_list.addItem(item)

    def select_pdf_file(self):
        """Open file dialog to select PDF"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            "PDF Files (*.pdf)"
        )

        if file_path:
            self.load_pdf(file_path)

    def load_recent_book(self, item):
        """Load a book from recent books list - double click handler"""
        self.select_recent_book(item)

    def select_recent_book(self, item):
        """Select a book from recent books list"""
        file_path = item.data(Qt.UserRole)
        if file_path and os.path.exists(file_path):
            # Load the PDF file
            self.load_pdf(file_path)
        else:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("file_not_found")
            )

    def load_existing_configuration(self):
        """Load existing reading speed configuration for the current book"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            config = self.recent_books_manager.get_reading_speed_config(self.pdf_path)
            if config:
                # Load configuration data
                self.avg_words_per_page = config['avg_words_per_page']
                self.words_per_page_list = config.get('words_per_page_list', [])
                self.sample_text = config.get('sample_text', '')

                # Update UI to show book is ready
                self.pdf_display.setText(f"{os.path.basename(self.pdf_path)} (Prepared)")
                self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} ({config['preparation_method']})")

                # Show analysis and reading sections
                self.analysis_group.setVisible(True)
                self.reading_group.setVisible(True)
                self.btn_start_reading.setEnabled(True)

                # Generate comprehension questions if we have sample text
                if self.sample_text and len(self.sample_text) > 200:
                    self.generate_comprehension_questions()
                    self.comprehension_group.setVisible(True)

                # Show success message
                QMessageBox.information(
                    self,
                    self.localization.get_text("book_ready"),
                    f"This book is already prepared for reading speed measurement!\n\n"
                    f"Configuration:\n"
                    f"• Pages: {self.total_pages}\n"
                    f"• Words per page: {self.avg_words_per_page:.0f}\n"
                    f"• Method: {config['preparation_method'].title()}\n\n"
                    f"You can start reading immediately or reconfigure if needed."
                )

                print(f"Loaded existing configuration: {self.avg_words_per_page:.0f} WPM, method: {config['preparation_method']}")
            else:
                # No configuration found, proceed with analysis
                self.analyze_word_count()

        except Exception as e:
            print(f"Error loading existing configuration: {e}")
            # Fallback to normal analysis
            self.analyze_word_count()

    def show_quick_setup_dialog(self):
        """Show quick setup dialog for first-time books"""
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QSpinBox, QGroupBox, QFormLayout

        dialog = QDialog(self)
        dialog.setWindowTitle("Quick Setup - Reading Speed Measurement")
        dialog.setModal(True)
        dialog.resize(500, 400)

        layout = QVBoxLayout(dialog)

        # Title
        title = QLabel("📚 Setup Reading Speed Measurement")
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(title)

        # Description
        desc = QLabel(
            f"This book ({os.path.basename(self.pdf_path)}) is not yet configured for reading speed measurement.\n"
            f"Choose how you'd like to proceed:"
        )
        desc.setWordWrap(True)
        desc.setStyleSheet("margin-bottom: 15px;")
        layout.addWidget(desc)

        # Quick options group
        quick_group = QGroupBox("⚡ Quick Options")
        quick_layout = QVBoxLayout(quick_group)

        # Default option
        default_btn = QPushButton("🚀 Use Default (150 words/page)")
        default_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 12px;
                border-radius: 6px;
                font-weight: bold;
                text-align: left;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        default_btn.clicked.connect(lambda: self.setup_with_default(dialog))
        quick_layout.addWidget(default_btn)

        # Manual input option
        manual_group = QGroupBox("✏️ Manual Input")
        manual_layout = QFormLayout(manual_group)

        self.quick_words_spin = QSpinBox()
        self.quick_words_spin.setRange(50, 2000)
        self.quick_words_spin.setValue(150)
        manual_layout.addRow("Words per page:", self.quick_words_spin)

        manual_btn = QPushButton(self.localization.get_text("use_manual_value"))
        manual_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        manual_btn.clicked.connect(lambda: self.setup_with_manual(dialog))
        manual_layout.addRow(manual_btn)

        quick_layout.addWidget(manual_group)
        layout.addWidget(quick_group)

        # Advanced option
        advanced_btn = QPushButton("🔍 Detailed Analysis (Recommended)")
        advanced_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                padding: 12px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        advanced_btn.clicked.connect(lambda: self.setup_with_analysis(dialog))
        layout.addWidget(advanced_btn)

        # Cancel button
        cancel_btn = QPushButton(self.localization.get_text("cancel_button"))
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #757575;
                color: white;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #616161;
            }
        """)
        cancel_btn.clicked.connect(dialog.reject)
        layout.addWidget(cancel_btn)

        return dialog.exec() == QDialog.Accepted

    def setup_with_default(self, dialog):
        """Setup with default 150 words per page"""
        self.avg_words_per_page = 150
        self.sample_text = ""
        self.words_per_page_list = []

        # Update UI
        self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (default)")
        self.analysis_group.setVisible(True)
        self.reading_group.setVisible(True)
        self.btn_start_reading.setEnabled(True)

        # Cache the configuration
        self.cache_analysis_results("default")

        dialog.accept()

    def setup_with_manual(self, dialog):
        """Setup with manual words per page value"""
        words_per_page = self.quick_words_spin.value()
        self.avg_words_per_page = words_per_page
        self.sample_text = ""
        self.words_per_page_list = []

        # Update UI
        self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (manual)")
        self.analysis_group.setVisible(True)
        self.reading_group.setVisible(True)
        self.btn_start_reading.setEnabled(True)

        # Cache the configuration
        self.cache_analysis_results("manual")

        dialog.accept()

    def setup_with_analysis(self, dialog):
        """Setup with detailed analysis"""
        self._perform_detailed_analysis = True
        dialog.accept()
        # Continue with the normal analysis process
        # This will be handled by the analyze_word_count method

    def analyze_word_count(self):
        """Analyze word count from 5 random pages with caching support"""
        try:
            # Show quick setup dialog first for new books
            if not self.show_quick_setup_dialog():
                return  # User cancelled

            # If user chose detailed analysis, continue with the analysis
            # (setup_with_analysis sets this flag)
            if hasattr(self, '_perform_detailed_analysis') and self._perform_detailed_analysis:
                self._perform_detailed_analysis = False  # Reset flag

                # Perform fresh analysis
                # Select 5 random pages for analysis
                if self.total_pages > 5:
                    sample_pages = random.sample(range(self.total_pages), 5)
                else:
                    sample_pages = list(range(self.total_pages))

                total_words = 0
                sample_texts = []

                for page_num in sample_pages:
                    page = self.pdf_doc.load_page(page_num)
                    text = page.get_text("text")

                    # Count words (split by whitespace and filter empty strings)
                    words = [word for word in text.split() if word.strip()]
                    total_words += len(words)

                    # Store sample text for comprehension questions
                    if len(text.strip()) > 100:  # Only store substantial text
                        sample_texts.append(text.strip())

                # Calculate average
                if len(sample_pages) > 0:
                    self.avg_words_per_page = total_words / len(sample_pages)
                else:
                    self.avg_words_per_page = 0

                # Store sample text for questions
                if sample_texts:
                    self.sample_text = random.choice(sample_texts)

                # Check if OCR is needed (no extractable text found)
                if self.avg_words_per_page == 0 or total_words == 0:
                    self.handle_ocr_needed_scenario()
                    return

                # Cache the analysis results
                self.cache_analysis_results("auto")

                # Update UI
                self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f}")
                self.analysis_group.setVisible(True)
                self.reading_group.setVisible(True)

                # Show word count verification dialog
                self.show_word_count_verification()

                self.btn_start_reading.setEnabled(True)

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to analyze PDF: {str(e)}"
            )



    def cache_analysis_results(self, preparation_method="auto"):
        """Cache the analysis results for future use"""
        if not self.recent_books_manager or not self.pdf_path:
            return

        try:
            # Save to dedicated reading speed configuration table
            success = self.recent_books_manager.save_reading_speed_config(
                file_path=self.pdf_path,
                total_pages=self.total_pages,
                avg_words_per_page=self.avg_words_per_page,
                words_per_page_list=getattr(self, 'words_per_page_list', []),
                sample_text=self.sample_text[:500] if self.sample_text else '',
                preparation_method=preparation_method
            )

            if success:
                print(f"Reading speed configuration saved: {self.avg_words_per_page:.0f} WPM, method: {preparation_method}")
                # Refresh activity tab if it exists
                if hasattr(self, 'activity_grid_layout'):
                    QTimer.singleShot(100, self.refresh_activity_books)
            else:
                print("Failed to save reading speed configuration")

        except Exception as e:
            print(f"Error caching analysis results: {e}")

    def prepare_book_analysis(self):
        """Pre-analyze book and cache results for future use"""
        if not self.pdf_path or not self.pdf_doc:
            return

        try:
            # Force fresh analysis (ignore cache)
            self.pdf_display.setText(f"Analyzing {os.path.basename(self.pdf_path)}...")

            # Select 5 random pages for analysis
            if self.total_pages > 5:
                sample_pages = random.sample(range(self.total_pages), 5)
            else:
                sample_pages = list(range(self.total_pages))

            total_words = 0
            sample_texts = []

            for page_num in sample_pages:
                page = self.pdf_doc.load_page(page_num)
                text = page.get_text("text")

                # Count words (split by whitespace and filter empty strings)
                words = [word for word in text.split() if word.strip()]
                total_words += len(words)

                # Store sample text for comprehension questions
                if len(text.strip()) > 100:  # Only store substantial text
                    sample_texts.append(text.strip())

            # Calculate average
            if len(sample_pages) > 0:
                self.avg_words_per_page = total_words / len(sample_pages)
            else:
                self.avg_words_per_page = 0

            # Store sample text for questions
            if sample_texts:
                self.sample_text = random.choice(sample_texts)

            # Cache the analysis results
            self.cache_analysis_results("prepared")

            # Update UI
            self.pdf_display.setText(os.path.basename(self.pdf_path))
            self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f} (prepared)")
            self.analysis_group.setVisible(True)
            self.reading_group.setVisible(True)
            self.btn_start_reading.setEnabled(True)

            # Show success message
            QMessageBox.information(
                self,
                self.localization.get_text("book_prepared"),
                f"{self.localization.get_text('analysis_cached')}\n\n"
                f"{self.localization.get_text('avg_words_per_page')} {self.avg_words_per_page:.0f}"
            )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to prepare book: {str(e)}"
            )
            self.pdf_display.setText(os.path.basename(self.pdf_path))

    def show_word_count_verification(self):
        """Show dialog to verify and adjust word count if needed"""
        try:
            from PySide6.QtWidgets import QDialog, QVBoxLayout, QLabel, QSpinBox, QPushButton, QHBoxLayout

            dialog = QDialog(self)
            dialog.setWindowTitle(self.localization.get_text("verify_word_count") if self.localization else "Verify Word Count")
            dialog.setModal(True)
            dialog.resize(450, 250)

            layout = QVBoxLayout(dialog)

            # Title
            title = QLabel("📊 " + (self.localization.get_text("word_count_analysis") if self.localization else "Word Count Analysis"))
            title.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
            layout.addWidget(title)

            # Description
            analyzed_desc = self.localization.get_text('analyzed_pages_desc') if self.localization else 'We analyzed 5 random pages from your book.'
            avg_words_text = self.localization.get_text('average_words_found') if self.localization else 'Average words per page:'
            adjust_text = self.localization.get_text('adjust_if_needed') if self.localization else 'You can adjust this value if it seems incorrect:'

            desc_text = f"{analyzed_desc}\n\n{avg_words_text} {self.avg_words_per_page:.0f}\n\n{adjust_text}"
            desc = QLabel(desc_text)
            desc.setWordWrap(True)
            desc.setStyleSheet("margin: 10px 0;")
            layout.addWidget(desc)

            # Spinbox for adjustment
            spinbox_layout = QHBoxLayout()
            spinbox_label = QLabel(self.localization.get_text("words_per_page") if self.localization else "Words per page:")
            spinbox_layout.addWidget(spinbox_label)

            words_spinbox = QSpinBox()
            words_spinbox.setRange(50, 2000)
            words_spinbox.setValue(int(self.avg_words_per_page))
            words_spinbox.setStyleSheet("font-size: 14px; padding: 5px;")
            spinbox_layout.addWidget(words_spinbox)
            spinbox_layout.addStretch()

            layout.addLayout(spinbox_layout)

            # Buttons
            button_layout = QHBoxLayout()
            button_layout.addStretch()

            accept_btn = QPushButton(self.localization.get_text("accept") if self.localization else "✓ Accept")
            accept_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    padding: 8px 20px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #45a049;
                }
            """)
            accept_btn.clicked.connect(dialog.accept)
            button_layout.addWidget(accept_btn)

            layout.addLayout(button_layout)

            # Show dialog
            if dialog.exec() == QDialog.Accepted:
                # Update with user's value
                self.avg_words_per_page = words_spinbox.value()
                self.lbl_avg_words.setText(f"{self.avg_words_per_page:.0f}")

                # Re-cache with updated value
                self.cache_analysis_results("auto_verified")

        except Exception as e:
            print(f"Error showing word count verification dialog: {e}")
