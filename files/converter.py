"""
PDF → DOCX / TXT converter using Google Drive OCR API.
Replicates what the `tahweel` library does but with user OAuth2
instead of a service-account JSON, so end-users just sign in with
their Google account — no manual credential setup needed.

Dependencies (add to requirements.txt):
    google-auth-oauthlib>=1.2.0
    google-api-python-client>=2.120.0
    google-auth-httplib2>=0.2.0
    pdf2image>=1.17.0
    python-docx>=1.1.0
    Pillow>=10.0.0

System dependency:
    poppler-utils  (apt / brew / choco)
"""

from __future__ import annotations

import io
import json
import os
import tempfile
import threading
import time
from pathlib import Path
from typing import Callable

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from pdf2image import convert_from_path

# ──────────────────────────────────────────────────────────────────────────────
# CONFIG
# ──────────────────────────────────────────────────────────────────────────────

# The only Drive scope we need — create / read / delete our own files.
SCOPES = ["https://www.googleapis.com/auth/drive.file"]

# Where to store the user's token between sessions.
TOKEN_PATH = Path.home() / ".tahweel_token.json"

# Path to YOUR OAuth2 client secrets file.
# Download it from Google Cloud Console → APIs & Services → Credentials
# (Application type: "Desktop app")
CLIENT_SECRETS_PATH = Path(__file__).parent / "client_secrets.json"


# ──────────────────────────────────────────────────────────────────────────────
# AUTH
# ──────────────────────────────────────────────────────────────────────────────

def get_credentials() -> Credentials:
    """
    Return valid Google credentials, refreshing or launching the OAuth
    browser flow as needed.
    """
    creds: Credentials | None = None

    if TOKEN_PATH.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not CLIENT_SECRETS_PATH.exists():
                raise FileNotFoundError(
                    f"client_secrets.json not found at {CLIENT_SECRETS_PATH}. "
                    "Download it from Google Cloud Console."
                )
            flow = InstalledAppFlow.from_client_secrets_file(
                str(CLIENT_SECRETS_PATH), SCOPES
            )
            # Opens the system browser; falls back to localhost redirect.
            creds = flow.run_local_server(port=0, open_browser=True)

        # Persist token for next run.
        TOKEN_PATH.write_text(creds.to_json())

    return creds


def is_authenticated() -> bool:
    """Quick check — is a valid (non-expired) token already saved?"""
    if not TOKEN_PATH.exists():
        return False
    try:
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
        if creds.valid:
            return True
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            TOKEN_PATH.write_text(creds.to_json())
            return True
    except Exception:
        pass
    return False


def sign_out() -> None:
    """Delete the saved token, effectively signing the user out."""
    if TOKEN_PATH.exists():
        TOKEN_PATH.unlink()


def get_user_email() -> str | None:
    """Return the email address of the signed-in user, or None."""
    if not is_authenticated():
        return None
    try:
        creds = get_credentials()
        service = build("oauth2", "v2", credentials=creds)
        info = service.userinfo().get().execute()
        return info.get("email")
    except Exception:
        return None


# ──────────────────────────────────────────────────────────────────────────────
# CORE CONVERSION
# ──────────────────────────────────────────────────────────────────────────────

def _upload_image_as_google_doc(
    drive_service,
    image_path: Path,
    folder_id: str,
) -> str:
    """
    Upload a PNG image to Drive with 'importAsGoogleDoc' + OCR enabled,
    then return the resulting Google Doc file ID.
    """
    file_metadata = {
        "name": image_path.stem,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [folder_id],
    }
    media = MediaFileUpload(str(image_path), mimetype="image/png", resumable=True)
    uploaded = (
        drive_service.files()
        .create(
            body=file_metadata,
            media_body=media,
            fields="id",
            # This is the key that triggers Google Docs OCR ↓
            ocrLanguage="ar",   # Arabic; change if needed
            supportsAllDrives=True,
        )
        .execute()
    )
    return uploaded["id"]


def _export_doc_as(
    drive_service,
    file_id: str,
    mime_type: str,
) -> bytes:
    """Export a Google Doc as a given MIME type and return raw bytes."""
    request = drive_service.files().export_media(fileId=file_id, mimeType=mime_type)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def _delete_drive_file(drive_service, file_id: str) -> None:
    """Permanently delete a Drive file (clean up after conversion)."""
    try:
        drive_service.files().delete(fileId=file_id).execute()
    except HttpError:
        pass


def _ensure_temp_folder(drive_service) -> str:
    """
    Create (or reuse) a Drive folder called 'tahweel_temp' for scratch
    files; returns its ID.
    """
    query = (
        "name = 'tahweel_temp' "
        "and mimeType = 'application/vnd.google-apps.folder' "
        "and trashed = false"
    )
    result = drive_service.files().list(q=query, fields="files(id)").execute()
    files = result.get("files", [])
    if files:
        return files[0]["id"]

    folder_meta = {
        "name": "tahweel_temp",
        "mimeType": "application/vnd.google-apps.folder",
    }
    folder = drive_service.files().create(body=folder_meta, fields="id").execute()
    return folder["id"]


# ──────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ──────────────────────────────────────────────────────────────────────────────

class ConversionProgress:
    """Thread-safe progress holder passed to the callback."""
    def __init__(self, total_pages: int):
        self.total_pages = total_pages
        self.current_page = 0
        self.status = "starting"          # starting | converting | done | error
        self.error: str | None = None

    @property
    def percent(self) -> int:
        if self.total_pages == 0:
            return 0
        return int(self.current_page / self.total_pages * 100)


def convert_pdf(
    pdf_path: str | Path,
    output_dir: str | Path | None = None,
    remove_newlines: bool = True,
    on_progress: Callable[[ConversionProgress], None] | None = None,
) -> dict:
    """
    Convert a PDF file to DOCX and TXT using Google Drive OCR.

    Parameters
    ----------
    pdf_path      : Path to the input PDF.
    output_dir    : Where to save output files (defaults to same dir as PDF).
    remove_newlines: Strip trailing newlines from DOCX paragraphs.
    on_progress   : Optional callback called after each page conversion.

    Returns
    -------
    dict with keys: docx_path, txt_path, pages
    """
    pdf_path = Path(pdf_path)
    if output_dir is None:
        output_dir = pdf_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    creds = get_credentials()
    drive_service = build("drive", "v3", credentials=creds)

    # Step 1 — Rasterise PDF pages to PNG images
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        images = convert_from_path(str(pdf_path), dpi=300, fmt="png")
        total = len(images)
        progress = ConversionProgress(total_pages=total)

        if on_progress:
            on_progress(progress)

        folder_id = _ensure_temp_folder(drive_service)

        all_text_pages: list[str] = []
        docx_texts: list[str] = []

        progress.status = "converting"

        for i, img in enumerate(images):
            img_path = tmp_path / f"page_{i:04d}.png"
            img.save(str(img_path), "PNG")

            # Step 2 — Upload image → Google Doc (OCR happens here)
            doc_id = _upload_image_as_google_doc(drive_service, img_path, folder_id)

            try:
                # Step 3 — Export as plain text for TXT output
                txt_bytes = _export_doc_as(
                    drive_service, doc_id,
                    "text/plain"
                )
                page_text = txt_bytes.decode("utf-8", errors="replace")

                # Export as DOCX for Word output
                docx_bytes = _export_doc_as(
                    drive_service, doc_id,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                all_text_pages.append(page_text)
                docx_texts.append((i, docx_bytes))

            finally:
                # Step 4 — Clean up Drive scratch file
                _delete_drive_file(drive_service, doc_id)

            progress.current_page = i + 1
            if on_progress:
                on_progress(progress)

        # Step 5 — Merge TXT pages
        txt_path = output_dir / pdf_path.with_suffix(".txt").name
        txt_path.write_text(
            "\n\n--- صفحة جديدة ---\n\n".join(all_text_pages),
            encoding="utf-8"
        )

        # Step 6 — Merge DOCX pages into a single file
        _merge_docx_pages(
            docx_texts,
            output_dir / pdf_path.with_suffix(".docx").name,
            remove_newlines=remove_newlines,
        )

        # Clean up the temp Drive folder
        try:
            drive_service.files().delete(fileId=folder_id).execute()
        except HttpError:
            pass

        progress.status = "done"
        if on_progress:
            on_progress(progress)

        return {
            "docx_path": str(output_dir / pdf_path.with_suffix(".docx").name),
            "txt_path": str(txt_path),
            "pages": total,
        }


def _merge_docx_pages(
    docx_page_bytes: list[tuple[int, bytes]],
    output_path: Path,
    remove_newlines: bool = True,
) -> None:
    """Merge per-page DOCX bytes into a single DOCX file."""
    from docx import Document
    from docx.oxml.ns import qn
    from copy import deepcopy
    import lxml.etree as etree

    merged = Document()

    for page_idx, (_, raw) in enumerate(sorted(docx_page_bytes, key=lambda x: x[0])):
        buf = io.BytesIO(raw)
        page_doc = Document(buf)

        for para in page_doc.paragraphs:
            text = para.text
            if remove_newlines:
                text = text.rstrip("\n")
            if not text.strip():
                continue
            new_para = merged.add_paragraph()
            # Preserve alignment (RTL docs)
            new_para.alignment = para.alignment
            run = new_para.add_run(text)
            # Copy font size / bold if present in source
            if para.runs:
                src_run = para.runs[0]
                run.bold = src_run.bold
                if src_run.font.size:
                    run.font.size = src_run.font.size

        if page_idx < len(docx_page_bytes) - 1:
            merged.add_page_break()

    merged.save(str(output_path))
