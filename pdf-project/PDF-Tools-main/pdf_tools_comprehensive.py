#!/usr/bin/env python3
"""
PDF Tools Comprehensive - Single File Solution
Complete PDF manipulation suite with bookmarks and page operations
"""

import sys
import os
import re
import json
import threading
import time
import fitz  # PyMuPDF
from typing import List, Optional, Tuple, Dict, Any
from dataclasses import dataclass
from datetime import datetime
from functools import partial

# Import app paths utility
try:
    from app_paths import get_settings_path, migrate_old_data
    APP_PATHS_AVAILABLE = True
except ImportError:
    APP_PATHS_AVAILABLE = False
    def get_settings_path(settings_name):
        return settings_name
    def migrate_old_data():
        return []

# Import version and update system
try:
    from version import VERSION, get_app_info, get_version_info, CREDITS
    from update_checker import UpdateChecker
    VERSION_SYSTEM_AVAILABLE = True
except ImportError:
    VERSION = "1.0.0"
    VERSION_SYSTEM_AVAILABLE = False
    print("Warning: Version system not available")
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGroupBox, QLabel, QLineEdit, QPushButton, QTableWidget,
    QTableWidgetItem, QSpinBox, QFileDialog, QMessageBox, QStatusBar,
    QSplitter, QTabWidget, QTextEdit, QListWidget, QListWidgetItem,
    QInputDialog, QProgressBar, QCheckBox, QComboBox, QScrollArea,
    QToolBar, QSlider, QColorDialog, QButtonGroup, QRadioButton,
    QFrame, QGraphicsView, QGraphicsScene, QGraphicsPixmapItem,
    QGraphicsTextItem, QGraphicsRectItem, QGraphicsEllipseItem,
    QGraphicsLineItem, QGraphicsPathItem, QSizePolicy, QTreeWidget, QTreeWidgetItem
)
from PySide6.QtCore import Qt, QThread, Signal, QTimer, QPointF, QRectF, QSizeF
from PySide6.QtGui import (QPixmap, QFont, QPainter, QPen, QBrush, QColor,
                          QTransform, QPainterPath, QPolygonF, QAction, QIcon,
                          QMouseEvent, QFontDatabase)
from PySide6.QtWidgets import QProgressDialog, QDialog, QMenu


# Font configuration
def load_custom_fonts():
    """Load IBM Plex Sans Arabic fonts"""
    try:
        # Get the directory where the script/executable is located
        if getattr(sys, 'frozen', False):
            # Running as compiled executable
            application_path = sys._MEIPASS
        else:
            # Running as script
            application_path = os.path.dirname(os.path.abspath(__file__))

        # Font paths - try multiple locations
        font_paths = [
            # For PyInstaller bundled app
            os.path.join(application_path, "fonts", "IBMPlexSansArabic-Regular.ttf"),
            os.path.join(application_path, "fonts", "IBMPlexSansArabic-Bold.ttf"),
            # For development
            os.path.join(application_path, "..", "website", "src", "fonts", "IBMPlexSansArabic-Regular.ttf"),
            os.path.join(application_path, "..", "website", "src", "fonts", "IBMPlexSansArabic-Bold.ttf")
        ]

        font_ids = []
        for font_path in font_paths:
            if os.path.exists(font_path):
                font_id = QFontDatabase.addApplicationFont(font_path)
                if font_id != -1:
                    font_ids.append(font_id)
                    font_families = QFontDatabase.applicationFontFamilies(font_id)
                    print(f"✅ Loaded font: {font_families} from {font_path}")
                else:
                    print(f"❌ Failed to load font: {font_path}")
            else:
                print(f"⚠️ Font file not found: {font_path}")

        if not font_ids:
            print("⚠️ No IBM Plex Sans Arabic fonts loaded. Using system fallback fonts.")

        return font_ids
    except Exception as e:
        print(f"❌ Error loading fonts: {e}")
        return []

def get_arabic_font(size=12, bold=False):
    """Get the appropriate Arabic font"""
    # Try IBM Plex Sans Arabic first
    font_families = QFontDatabase.families()
    if "IBM Plex Sans Arabic" in font_families:
        font = QFont("IBM Plex Sans Arabic", size)
    elif "Tajawal" in font_families:
        font = QFont("Tajawal", size)
    else:
        # Fallback to system Arabic font
        font = QFont("Arial", size)

    if bold:
        font.setBold(True)

    return font

def get_font_family_css():
    """Get the font family CSS string for consistent use across stylesheets"""
    return 'font-family: "IBM Plex Sans Arabic", "Tajawal", "Arial", sans-serif;'

def apply_global_font_style():
    """Apply global font styling to the application"""
    app = QApplication.instance()
    if app:
        # Set default font for the application
        arabic_font = get_arabic_font()
        app.setFont(arabic_font)

        # Note: Global stylesheet will be applied by theme system
        # This function now just sets the default font


@dataclass
class Bookmark:
    """Simple bookmark entry"""
    title: str
    page: int
    level: int = 1


@dataclass
class HistoryEntry:
    """History entry for tracking operations"""
    timestamp: str
    operation: str
    input_files: List[str]
    output_file: str
    status: str
    details: str


class ProgressDialog(QProgressDialog):
    """Custom progress dialog with better styling"""

    def __init__(self, title, message, parent=None):
        super().__init__(message, "Cancel", 0, 0, parent)
        self.setWindowTitle(title)
        self.setModal(True)
        self.setMinimumDuration(0)
        self.setAutoClose(False)
        self.setAutoReset(False)
        self.setValue(0)

        # Style the progress dialog
        self.setStyleSheet("""
            QProgressDialog {
                background-color: white;
                border: 2px solid #1976D2;
                border-radius: 8px;
            }
            QProgressBar {
                border: 2px solid #E0E0E0;
                border-radius: 5px;
                text-align: center;
                background-color: #F5F5F5;
            }
            QProgressBar::chunk {
                background-color: #1976D2;
                border-radius: 3px;
            }
            QLabel {
                color: #333;
                font-weight: bold;
            }
            QPushButton {
                background-color: #F44336;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #D32F2F;
            }
        """)

    def update_message(self, message):
        """Update the progress message"""
        self.setLabelText(message)


class Localization:
    """Localization system for Arabic and English"""

    def __init__(self):
        self.current_language = "ar"  # Arabic as default
        self.translations = {
            "ar": {
                # Main window
                "app_title": "عدة القارئ للملفات الرقمية - PDF Toolkits",
                "ready_status": "جاهز - اختر علامة تبويب للبدء في العمل مع ملفات PDF",

                # Tab names
                "bookmark_manager": "📖 مدير الفهارس",
                "bookmark_extractor": "📤 مستخرج الفهارس",
                "chapter_weight_analyzer": "📊 إعداد خطة المذاكرة",
                "page_operations": "📄 عمليات الصفحات",
                "watermark": "💧 العلامة المائية",
                "extract_images": "🖼️ استخراج الصور",
                "extract_text": "📝 استخراج النص",
                "merge_pdfs": "📚 دمج ملفات PDF",
                "split_pdfs": "✂️ تقسيم PDF",
                "compress": "🗜️ ضغط",
                "page_editing": "✂️ تحرير الصفحات",
                "pdf_viewer": "عارض ومحرر PDF",
                "settings": "⚙️ الإعدادات",
                "history": "📋 السجل",
                "security_removal": "🔓 إزالة الحماية",

                # Common buttons
                "browse": "📁 تصفح",
                "cancel": "إلغاء",
                "ok": "موافق",
                "save": "💾 حفظ",
                "load": "🔄 تحميل",
                "extract": "📤 استخراج",
                "delete": "🗑️ حذف",
                "insert": "✚ إدراج",
                "merge": "📚 دمج",
                "compress_btn": "🗜️ ضغط",
                "rotate": "🔄 تدوير",
                "add_margins": "📏 إضافة هوامش",

                # Common labels
                "pdf_file": "📄 ملف PDF:",
                "page_range": "📄 نطاق الصفحات:",
                "pages": "📄 الصفحات:",
                "position": "📍 الموضع:",
                "output": "📁 الإخراج:",
                "status": "الحالة:",
                "success": "نجح!",
                "error": "خطأ",
                "warning": "تحذير",
                "information": "معلومات",

                # Bookmark Manager
                "bookmark_manager_desc": "🎯 <b>إدارة الإشارات المرجعية:</b> إدراج إشارات مرجعية من ملفات نصية في ملفات PDF مع كشف الإزاحة التلقائي",
                "file_selection": "📁 اختيار الملفات",
                "bookmark_file": "📖 ملف الإشارات المرجعية:",
                "select_pdf_placeholder": "اختر ملف PDF...",
                "select_bookmark_placeholder": "اختر ملف نصي بتنسيق 'العنوان - الصفحة'...",
                "load_files_preview": "🔄 تحميل الملفات والمعاينة",
                "page_adjustment": "📊 تعديل الصفحة",
                "page_offset": "إزاحة الصفحة:",
                "offset_help": "(+/- للتعديل للصفحات الإضافية/المفقودة)",
                "smart_offset_detection": "🎯 كشف الإزاحة الذكي",
                "offset_info": "💡 انقر على أي إشارة مرجعية لمعاينة صفحتها، ثم استخدم الكشف التلقائي إذا لم تكن متوافقة",
                "auto_detect_offset": "🔍 كشف تلقائي للإزاحة من الإشارة المحددة",
                "apply_correction": "✅ تطبيق التصحيح",
                "correction_applied": "تم تطبيق التصحيح",
                "manual_correction": "🎯 تصحيح يدوي",
                "set_correct_page": "تعيين الصفحة الصحيحة",
                "current_viewing": "تعرض حالياً الصفحة",
                "bookmark_should_be": "يجب أن تكون الإشارة المرجعية في الصفحة",
                "confirm_correction": "تأكيد التصحيح",
                "bookmarks_table": "📋 الإشارات المرجعية",
                "title": "العنوان",
                "original": "الأصلي",
                "adjusted": "المعدل",
                "level": "المستوى",
                "preview": "معاينة",
                "insert_bookmarks": "📚 إدراج الإشارات المرجعية في PDF",
                "pdf_preview": "📄 معاينة PDF",
                "prev": "◀ السابق",
                "next": "التالي ▶",
                "no_pdf_loaded": "لم يتم تحميل PDF",
                "pdf_preview_placeholder": "ستظهر معاينة PDF هنا",
                "click_browse_to_load": "انقر على 'تصفح' لتحميل ملف PDF",
                "navigation": "التنقل",
                "hide_sidebar": "إخفاء الشريط الجانبي",
                "show_sidebar": "إظهار الشريط الجانبي",
                "page_thumbnails": "صور مصغرة للصفحات",
                "no_bookmarks": "لا توجد إشارات مرجعية",
                "no_file_loaded": "لم يتم تحميل ملف",
                "page": "صفحة",
                "load_new_pdf": "تحميل ملف آخر",
                "load_different_pdf": "تحميل ملف PDF مختلف",
                "show_controls": "إظهار عناصر التحكم",
                "hide_controls": "إخفاء عناصر التحكم",
                "tools": "الأدوات",
                "zoom": "التكبير",
                "style": "النمط",
                "select_tool": "أداة التحديد",
                "annotation_size": "حجم التعليق",
                "load_new": "تحميل ملف جديد",

                # Navigation tooltips
                "first_page": "الصفحة الأولى",
                "previous_page": "الصفحة السابقة",
                "next_page": "الصفحة التالية",
                "last_page": "الصفحة الأخيرة",

                # Additional labels
                "page_label": "صفحة:",
                "no_offset_needed": "لا حاجة للإزاحة - الإشارة المرجعية في الموضع الصحيح",
                "error_loading_page": "خطأ في تحميل الصفحة",
                "failed_to_render": "فشل في عرض الصفحة",
                "text_format_tooltip": "TXT: نص عادي، JSON: بيانات منظمة، DOCX: مستند Microsoft Word",
                "show_bookmark_analytics": "عرض توزيع أوزان الإشارات المرجعية",
                "close_analytics": "إغلاق التحليلات",
                "error_title": "خطأ",
                "failed_to_load_pdf": "فشل في تحميل PDF",

                # PDF Viewer specific
                "bookmarks": "الإشارات المرجعية",
                "hide_bookmarks": "إخفاء الإشارات المرجعية",
                "show_bookmarks": "عرض الإشارات المرجعية",
                "no_bookmarks": "لم يتم العثور على إشارات مرجعية",
                "bookmark_weight_distribution": "توزيع أوزان الإشارات المرجعية",
                "click_to_go_to_page": "انقر للانتقال إلى الصفحة",
                "go_to_page": "الانتقال إلى الصفحة",
                "enter_page_number": "أدخل رقم الصفحة",
                "resumed_reading": "تم استئناف القراءة",
                "resumed_from_page": "تم الاستئناف من الصفحة",
                "open_in_weight_analyzer": "فتح في محلل أوزان الفصول",
                "pdf_loaded_in_analyzer": "تم تحميل الملف في محلل الأوزان",
                "weight_analyzer_not_available": "محلل أوزان الفصول غير متاح",

                # Page Operations
                "page_operations_desc": "📄 <b>عمليات الصفحات:</b> استخراج وحذف وإدراج الصفحات، دمج الملفات، وتدويرها وقصّها وإضافة الهوامش",
                "extract_pages": "📤 استخراج الصفحات",
                "delete_pages": "🗑️ حذف الصفحات",
                "insert_blank_pages": "📄 إدراج صفحات فارغة",
                "pages_to_extract": "📄 الصفحات المراد استخراجها:",
                "pages_to_delete": "📄 الصفحات المراد حذفها:",
                "insert_at_position": "📍 الإدراج في الموضع:",
                "number_of_pages": "📄 عدد الصفحات:",
                "extract_pages_btn": "📤 استخراج الصفحات",
                "delete_pages_btn": "🗑️ حذف الصفحات",
                "insert_blank_btn": "📄 إدراج صفحات فارغة",

                # Watermark
                "watermark_desc": "💧 <b>أداة العلامة المائية:</b> إضافة علامات مائية نصية إلى صفحات PDF أو إزالة العلامات المائية الموجودة",
                "watermark_text": "💧 نص العلامة المائية:",
                "watermark_position": "📍 الموضع:",
                "font_size": "📏 حجم الخط:",
                "add_watermark": "💧 إضافة علامة مائية",
                "remove_watermark": "🗑️ إزالة العلامة المائية",
                "add_watermark_section": "✚ إضافة علامة مائية",
                "remove_watermark_section": "🗑️ إزالة علامة مائية",
                "watermark_removed_successfully": "تم إزالة العلامة المائية بنجاح",
                "watermark_removal_info": "هذه الميزة المحسنة تحاول إزالة العلامات المائية المختلفة من ملف PDF بما في ذلك: الشعارات، النصوص، الروابط (مثل UPDF، وغيرها). قد تحتاج لمراجعة النتيجة للتأكد من عدم إزالة محتوى مهم.",
                "removal_options": "خيارات الإزالة",
                "aggressive_mode": "الوضع القوي",
                "aggressive_mode_tooltip": "يزيل المزيد من العناصر المشبوهة، لكن قد يزيل محتوى مهم",
                "remove_urls": "إزالة الروابط والمواقع",
                "center": "الوسط",
                "top_left": "أعلى اليسار",
                "top_right": "أعلى اليمين",
                "bottom_left": "أسفل اليسار",
                "bottom_right": "أسفل اليمين",

                # Image Extraction
                "image_extraction_desc": "🖼️ <b>استخراج الصور:</b> استخراج جميع الصور من ملفات PDF مع كشف محسن",
                "extraction_method": "🔧 الطريقة:",
                "quality": "📊 الجودة:",
                "both_methods": "كلاهما (مدمج + مُعرض)",
                "embedded_only": "الصور المدمجة فقط",
                "rendered_only": "الصفحات المُعرضة فقط",
                "high_quality": "عالية (2x)",
                "medium_quality": "متوسطة (1.5x)",
                "standard_quality": "قياسية (1x)",
                "extract_all_images": "🖼️ استخراج جميع الصور",

                # Text Extraction
                "text_extraction_desc": "📝 <b>استخراج النص:</b> استخراج محتوى النص من ملفات PDF بتنسيقات مختلفة مع دعم المعالجة المجمعة",
                "output_format": "📋 تنسيق الإخراج:",
                "extract_text_btn": "📝 استخراج النص",
                "processing_mode": "وضع المعالجة:",
                "single_file": "ملف واحد",
                "batch_processing": "معالجة مجمعة",
                "select_multiple_pdfs": "📁 اختيار عدة ملفات PDF",
                "clear_selection": "🗑️ مسح التحديد",
                "output_directory": "مجلد الإخراج:",
                "select_output_directory": "اختيار مجلد الإخراج",
                "batch_extract_btn": "📝 استخراج النص (مجمع)",
                "docx_format": "مستند Word (DOCX)",
                "txt_format": "نص عادي (TXT)",
                "json_format": "بيانات منظمة (JSON)",

                # Merge PDFs
                "merge_desc": "📚 <b>دمج ملفات PDF:</b> دمج عدة ملفات PDF في مستند واحد",
                "select_pdfs_merge": "📚 اختيار ملفات PDF للدمج",
                "add_pdfs": "✚ إضافة ملفات PDF",
                "remove": "✖ إزالة",
                "merge_pdfs_btn": "📚 دمج ملفات PDF",
                "merge_info": "ℹ️ معلومات الدمج",
                "how_to_use": "<b>كيفية الاستخدام:</b>",
                "merge_instructions": "1. انقر على 'إضافة ملفات PDF' لاختيار عدة ملفات PDF<br>2. سيتم دمج الملفات بالترتيب الذي تظهر به في القائمة<br>3. استخدم 'إزالة' لحذف الملفات المحددة من القائمة<br>4. انقر على 'دمج ملفات PDF' لدمج جميع الملفات في ملف واحد",

                # Split PDF
                "split_pdfs": "✂️ تقسيم PDF",
                "split_desc": "✂️ <b>تقسيم PDF:</b> تقسيم ملف PDF إلى عدة ملفات عند أرقام صفحات محددة",
                "select_pdf_to_split": "📄 اختيار ملف PDF للتقسيم",
                "split_points": "✂️ نقاط التقسيم",
                "split_points_label": "أرقام الصفحات للتقسيم عندها (مفصولة بفواصل):",
                "split_points_placeholder": "مثال: 5, 10, 15",
                "preview_split_points": "👁️ معاينة نقاط التقسيم",
                "split_pdf_btn": "✂️ تقسيم PDF",
                "split_info": "ℹ️ معلومات التقسيم",
                "split_instructions": "1. اختر ملف PDF للتقسيم<br>2. أدخل أرقام الصفحات حيث تريد تقسيم الملف (مفصولة بفواصل)<br>3. انقر على 'معاينة نقاط التقسيم' لرؤية الصفحات التي سيتم التقسيم عندها<br>4. انقر على 'تقسيم PDF' لإنشاء ملفات منفصلة",
                "split_preview_title": "معاينة نقاط التقسيم",
                "page_preview": "معاينة الصفحة {0}",
                "split_at_page": "التقسيم عند الصفحة {0}",
                "output_files_will_be": "سيتم إنشاء {0} ملفات:",
                "file_part": "الجزء {0}",
                "pages_range": "الصفحات {0} - {1}",
                "select_output_directory": "اختيار مجلد الإخراج",
                "save_split_pdfs": "حفظ ملفات PDF المقسمة",
                "pdfs_split_successfully": "تم تقسيم PDF بنجاح",
                "split_complete": "تم إنشاء {0} ملفات بنجاح",
                "invalid_split_points": "نقاط تقسيم غير صالحة",
                "enter_valid_page_numbers": "يرجى إدخال أرقام صفحات صالحة مفصولة بفواصل",
                "page_numbers_out_of_range": "بعض أرقام الصفحات خارج النطاق الصالح (1-{0})",
                "no_pdf_selected_for_split": "لم يتم اختيار ملف PDF للتقسيم",
                "select_pdf_file_first": "يرجى اختيار ملف PDF أولاً",
                "total_pages": "إجمالي الصفحات: {0}",

                # Compress
                "compress_desc": "🗜️ <b>ضغط PDF:</b> تقليل حجم ملف PDF مع الحفاظ على الجودة",
                "compression_options": "⚙️ خيارات الضغط",
                "compression_features": "<b>ميزات الضغط:</b><br>• إزالة الكائنات والمراجع غير المستخدمة<br>• تحسين ضغط الصور<br>• تنظيف هيكل المستند<br>• تقليل حجم الملف مع الحفاظ على الجودة",
                "compress_pdf_btn": "🗜️ ضغط PDF",
                "compression_results": "📊 نتائج الضغط",

                # PDF Viewer & Editor
                "pdf_viewer_desc": "عرض ملفات PDF مع أدوات التحرير والتعليق التوضيحي المتقدمة",
                "load_pdf": "📄 تحميل PDF",
                "pdf_navigation": "🧭 التنقل في PDF",
                "zoom_controls": "🔍 التحكم في التكبير",
                "zoom_in": "🔍+ تكبير",
                "zoom_out": "🔍- تصغير",
                "fit_page": "📄 ملائمة الصفحة",
                "fit_width": "↔️ ملائمة العرض",
                "previous_page": "⬅️ الصفحة السابقة",
                "next_page": "➡️ الصفحة التالية",
                "go_to_page": "🎯 الذهاب إلى الصفحة",
                "annotation_tools": "✏️ أدوات التعليق التوضيحي",
                "text_annotation": "📝 تعليق نصي",
                "rich_text_annotation": "تعليق نصي منسق",
                "highlight_text": "🖍️ تمييز النص",
                "underline_text": "📏 تسطير النص",
                "draw_rectangle": "⬜ رسم مستطيل",
                "draw_circle": "⭕ رسم دائرة",
                "draw_arrow": "➡️ رسم سهم",
                "sticky_note": "📌 ملاحظة لاصقة",
                "freehand_draw": "✏️ رسم حر",
                "annotation_color": "🎨 لون التعليق",
                "annotation_size": "📏 حجم التعليق",
                "ok": "موافق",
                "save_annotations": "💾 حفظ التعليقات",
                "clear_annotations": "🗑️ مسح التعليقات",
                "page_info": "الصفحة {current} من {total}",
                "invalid_page": "صفحة غير صالحة",
                "page_out_of_range": "الصفحة {page} خارج النطاق (1-{total})",

                # Page Editing
                "page_editing_desc": "✂️ <b>تحرير الصفحات:</b> تدوير الصفحات وقصها وإضافة هوامش بيضاء",
                "rotate_pages": "🔄 تدوير الصفحات",
                "rotation": "🔄 التدوير:",
                "rotate_90": "90° في اتجاه عقارب الساعة",
                "rotate_180": "180°",
                "rotate_270": "270° في اتجاه عقارب الساعة",
                "rotate_pages_btn": "🔄 تدوير الصفحات",
                "add_white_margins": "📏 إضافة هوامش بيضاء",
                "margin_size": "📏 حجم الهامش:",
                "points": " نقطة",
                "margin_help": "يضيف مساحة بيضاء حول الصفحات. مفيد لملفات PDF بدون هوامش أو محتوى ضيق.",
                "add_margins_btn": "📏 إضافة هوامش",

                "crop_pages": "✂️ قص الصفحات",
                "crop_box": "منطقة القص:",
                "crop_pages_btn": "✂️ قص الصفحات",
                "save_cropped_pdf": "حفظ PDF المقصوص",
                "invalid_crop_box": "منطقة القص غير صالحة",
                "pages_cropped_successfully": "تم قص الصفحات بنجاح",
                "pages_range_label": "نطاق الصفحات:",

                # Common messages
                "select_file_first": "يرجى اختيار ملف أولاً",
                "no_pages_specified": "يرجى تحديد الصفحات",
                "invalid_page_format": "تنسيق أرقام الصفحات غير صحيح",
                "operation_successful": "تمت العملية بنجاح!",
                "operation_failed": "فشلت العملية",
                "confirm_deletion": "تأكيد الحذف",
                "are_you_sure": "هل أنت متأكد؟",
                "pages_will_be_deleted": "سيتم حذف الصفحات المحددة",
                "missing_files": "ملفات مفقودة",
                "select_both_files": "يرجى اختيار كلا الملفين",
                "no_images_found": "لم يتم العثور على صور",
                "images_found": "تم العثور على صور",
                "text_extracted": "تم استخراج النص",
                "compression_completed": "اكتمل الضغط",
                "size_reduction": "تقليل الحجم",
                "space_saved": "المساحة المحفوظة",
                "original_size": "الحجم الأصلي",
                "compressed_size": "الحجم المضغوط",

                # Bookmark Extractor
                "bookmark_extractor_desc": "📤 <b>مستخرج الفهارس:</b> استخراج الفهارس الموجودة من ملفات PDF وحفظها في ملفات نصية",
                "extract_bookmarks": "📤 استخراج الفهارس",
                "extracted_bookmarks": "📋 الفهارس المستخرجة",
                "save_to_text": "💾 حفظ في ملف نصي",
                "no_bookmarks_found": "لم يتم العثور على فهارس",
                "bookmarks_extracted": "تم استخراج الفهارس",
                "bookmarks_saved": "تم حفظ الفهارس",

                # Chapter Weight Analyzer
                "chapter_weight_desc": "تحليل توزيع الفصول وإنشاء خطط قراءة ذكية بناءً على أوزان الفصول",
                "load_bookmarks_btn": "📖 تحميل الفهارس",
                "analyze_weights_btn": "📊 تحليل الأوزان",
                "generate_plan_btn": "📅 إنشاء خطة القراءة",
                "analysis_options": "⚙️ خيارات التحليل",
                "include_levels": "تضمين المستويات:",
                "level_1_only": "المستوى 1 فقط",
                "levels_1_2": "المستويات 1 و 2",
                "all_levels": "جميع المستويات",
                "chapter_weights_table": "📊 أوزان الفصول",
                "chapter_title": "عنوان الفصل",
                "start_page": "صفحة البداية",
                "end_page": "صفحة النهاية",
                "page_count": "عدد الصفحات",
                "weight_percent": "الوزن %",
                "reading_plan_section": "📅 إعداد خطة المذاكرة",
                "plan_duration": "مدة الخطة (أيام):",
                "start_date": "تاريخ البدء",
                "end_date": "تاريخ الانتهاء",
                "skip_weekends": "تخطي عطلة نهاية الأسبوع",
                "algorithm": "الخوارزمية",
                "algorithm_weight_based": "على أساس الوزن (دمج ذكي)",
                "algorithm_direct_pages": "صفحات مباشرة (تناسبي)",
                "reading_plan_table": "📅 جدول المذاكرة",
                "assigned_days": "الأيام المخصصة",
                "date_range": "نطاق التاريخ",
                "statistics_insights": "📈 الإحصائيات والرؤى",
                "total_chapters": "إجمالي الفصول:",
                "total_pages": "إجمالي الصفحات:",
                "avg_chapter_length": "متوسط طول الفصل:",
                "longest_chapter": "أطول فصل:",
                "shortest_chapter": "أقصر فصل:",
                "export_options": "💾 خيارات التصدير",
                "export_weights_csv": "تصدير الأوزان (CSV)",
                "export_weights_json": "تصدير الأوزان (JSON)",
                "export_plan_csv": "تصدير الخطة (CSV)",
                "export_plan_json": "تصدير الخطة (JSON)",
                "export_plan_text": "تصدير الخطة (نص)",
                "export_weights_excel": "تصدير الأوزان (Excel)",
                "export_plan_excel": "تصدير الخطة (Excel)",
                "export_weights_markdown": "تصدير الأوزان (Markdown)",
                "export_plan_markdown": "تصدير الخطة (Markdown)",
                "export_plan_obsidian": "تصدير الخطة (Obsidian)",
                "export_all_formats": "📦 تصدير جميع الصيغ",
                "export_pie_chart": "تصدير مخطط دائري",
                "export_bar_chart": "تصدير مخطط شريطي",
                "export_weight_chart": "تصدير مخطط المقارنة",
                "weights_exported": "تم تصدير الأوزان بنجاح",
                "plan_exported": "تم تصدير الخطة بنجاح",
                "chart_exported": "تم تصدير المخطط بنجاح",
                "all_formats_exported": "تم تصدير جميع الصيغ بنجاح",
                "no_weights_to_export": "لا توجد أوزان للتصدير. قم بتحليل الفهارس أولاً.",
                "select_chapters": "اختيار الفصول",
                "select_chapters_instruction": "اختر الفصول التي تريد تضمينها في خطة المذاكرة.\nقم بإلغاء تحديد الفصول مثل الفهرس والمقدمات وما إلى ذلك التي تريد تخطيها.",
                "select_all": "تحديد الكل",
                "deselect_all": "إلغاء تحديد الكل",
                "no_chapters_selected": "لم يتم تحديد أي فصول. يرجى تحديد فصل واحد على الأقل.",
                "no_plan_to_export": "لا توجد خطة للتصدير. قم بإنشاء خطة قراءة أولاً.",

                # Report export labels (Arabic)
                "report_reading_plan": "خطة القراءة",
                "report_total_target_duration": "المدة المستهدفة الإجمالية",
                "report_actual_assigned_duration": "المدة المخصصة الفعلية",
                "report_start_date": "تاريخ البدء",
                "report_average_daily_pages": "متوسط الصفحات اليومية (الأيام النشطة)",
                "report_block": "المجموعة",
                "report_days": "الأيام",
                "report_day": "اليوم",
                "report_chapter": "الفصل",
                "report_includes": "يشمل",
                "report_pages": "الصفحات",
                "report_duration": "المدة",
                "report_weight": "الوزن",
                "report_plan_created": "تم إنشاء الخطة",
                "report_to": "إلى",
                "report_days_unit": "يوم(أيام)",
                "plan_summary": "ملخص الخطة",
                "reading_schedule": "جدول القراءة",
                "total_reading_blocks": "إجمالي مجموعات القراءة",
                "daily_progress": "التقدم اليومي",
                "reading_days": "أيام القراءة",
                "quick_reference": "مرجع سريع",
                "reading_block": "مجموعة القراءة",
                "generated": "تم الإنشاء",
                "dates": "التواريخ",
                "copy_plan_to_clipboard": "نسخ الخطة إلى الحافظة",
                "plan_copied_to_clipboard": "تم نسخ الخطة إلى الحافظة بنجاح",

                # Export tooltips (Arabic)
                "tooltip_export_plan_excel": "تصدير خطة القراءة إلى ملف Excel مع تنسيق احترافي وجداول منظمة",
                "tooltip_export_plan_json": "تصدير خطة القراءة إلى تنسيق JSON للاستخدام البرمجي أو التكامل مع التطبيقات الأخرى",
                "tooltip_export_plan_obsidian": "تصدير خطة القراءة إلى تنسيق Markdown متوافق مع Obsidian مع مربعات اختيار للمتابعة اليومية",
                "tooltip_export_plan_text": "تصدير خطة القراءة إلى ملف نصي بسيط وسهل القراءة",
                "tooltip_export_plan_markdown": "تصدير خطة القراءة إلى تنسيق Markdown مع جداول ومربعات اختيار",
                "tooltip_copy_plan": "نسخ خطة القراءة المنسقة إلى الحافظة للصقها في أي تطبيق",
                "export_calendar_visualization": "تصدير التقويم المرئي",
                "tooltip_calendar_visualization": "إنشاء تقويم مرئي يوضح أيام المذاكرة مع ملاحظات إحصائية أسفل التقويم",
                "calendar_exported": "تم تصدير التقويم بنجاح",
                "total_study_days": "إجمالي أيام المذاكرة",
                "study_day": "يوم مذاكرة",
                "rest_day": "يوم راحة",
                "matplotlib_not_available": "مكتبة matplotlib غير متاحة. لا يمكن إنشاء المخططات البيانية.\n\nلتثبيت المكتبة، استخدم الأمر:\npip install matplotlib",
                "openpyxl_not_available": "مكتبة openpyxl غير متاحة. لا يمكن تصدير ملفات Excel.\n\nلتثبيت المكتبة، استخدم الأمر:\npip install openpyxl",
                "pdf_no_bookmarks_error": "هذا الملف لا يحتوي على إشارات مرجعية (Bookmarks).\nلا يمكن تحليل توزيع الأوزان بدون إشارات مرجعية.",
                "export_failed": "فشل التصدير",
                "check_file_permissions": "تحقق من صلاحيات الملف وتأكد من عدم فتحه في برنامج آخر.",
                "select_export_folder": "اختر مجلد التصدير",

                # File dialogs and messages
                "select_pdf_file": "اختيار ملف PDF",
                "select_bookmark_file": "اختيار ملف الفهارس",
                "save_extracted_text": "حفظ النص المستخرج",
                "save_compressed_pdf": "حفظ PDF المضغوط",
                "save_merged_pdf": "حفظ PDF المدموج",
                "save_rotated_pdf": "حفظ PDF المدور",
                "save_pdf_with_margins": "حفظ PDF مع الهوامش",
                "save_pdf_with_watermark": "حفظ PDF مع العلامة المائية",
                "save_extracted_pages": "حفظ الصفحات المستخرجة",
                "save_pdf_with_blank_pages": "حفظ PDF مع الصفحات الفارغة",

                # Progress and status messages
                "loading": "جاري التحميل...",
                "extracting": "جاري الاستخراج...",
                "compressing": "جاري الضغط...",
                "merging": "جاري الدمج...",
                "rotating": "جاري التدوير...",
                "adding_watermark": "جاري إضافة العلامة المائية...",
                "inserting_pages": "جاري إدراج الصفحات...",
                "deleting_pages": "جاري حذف الصفحات...",
                "adding_margins": "جاري إضافة الهوامش...",

                # Error messages
                "file_not_selected": "لم يتم اختيار ملف",
                "files_not_selected": "لم يتم اختيار الملفات",
                "invalid_page_range": "نطاق صفحات غير صحيح",
                "operation_cancelled": "تم إلغاء العملية",
                "file_access_error": "خطأ في الوصول للملف",
                "insufficient_permissions": "صلاحيات غير كافية",
                "disk_space_error": "مساحة القرص غير كافية",
                "corrupted_pdf": "ملف PDF تالف",

                # Success messages
                "files_loaded_successfully": "تم تحميل الملفات بنجاح",
                "bookmarks_inserted_successfully": "تم إدراج الفهارس بنجاح",
                "bookmarks_loaded_successfully": "تم تحميل الإشارات المرجعية بنجاح!",
                "found_bookmarks": "تم العثور على {count} إشارة مرجعية.",
                "preview_and_insert_instruction": "يمكنك الآن معاينة وإدراج الإشارات المرجعية في الخطوة 4 أدناه.",
                "no_selection": "لا يوجد تحديد",
                "select_bookmark_to_fix": "يرجى تحديد إشارة مرجعية لإصلاحها.",
                "text_extracted_successfully": "تم استخراج النص بنجاح",
                "images_extracted_successfully": "تم استخراج الصور بنجاح",
                "pdf_compressed_successfully": "تم ضغط PDF بنجاح",
                "pdfs_merged_successfully": "تم دمج ملفات PDF بنجاح",
                "pages_rotated_successfully": "تم تدوير الصفحات بنجاح",
                "watermark_added_successfully": "تم إضافة العلامة المائية بنجاح",
                "margins_added_successfully": "تم إضافة الهوامش بنجاح",
                "pages_extracted_successfully": "تم استخراج الصفحات بنجاح",
                "pages_deleted_successfully": "تم حذف الصفحات بنجاح",
                "blank_pages_inserted_successfully": "تم إدراج الصفحات الفارغة بنجاح",

                # File types and formats
                "pdf_files": "ملفات PDF",
                "text_files": "ملفات نصية",
                "json_files": "ملفات JSON",
                "all_files": "جميع الملفات",
                "image_files": "ملفات الصور",

                # Units and measurements
                "pages_count": "عدد الصفحات",
                "file_size": "حجم الملف",
                "megabytes": "ميجابايت",
                "kilobytes": "كيلوبايت",
                "bytes": "بايت",
                "percentage": "نسبة مئوية",
                "degrees": "درجة",

                # Interface elements
                "browse_button": "تصفح...",
                "select_button": "اختيار",
                "apply_button": "تطبيق",
                "reset_button": "إعادة تعيين",
                "clear_button": "مسح",
                "refresh_button": "تحديث",
                "close_button": "إغلاق",
                "help_button": "مساعدة",
                "about_button": "حول",

                # Help and instructions
                "help_title": "💡 مساعدة",
                "instructions": "التعليمات",
                "examples": "أمثلة",
                "tips": "نصائح",
                "page_format_examples": "<b>أمثلة على تنسيق الصفحات:</b>",
                "single_pages": "• صفحات منفردة: 1,3,5",
                "page_ranges": "• نطاقات الصفحات: 1-5,10-15",
                "mixed_format": "• تنسيق مختلط: 1,3,5-10,15",

                # Navigation and Preview
                "next": "التالي ▶",
                "prev": "◀ السابق",
                "pdf_preview": "📄 معاينة PDF",
                "no_pdf_loaded": "لم يتم تحميل PDF",
                "pdf_preview_placeholder": "ستظهر معاينة PDF هنا",
                "click_browse_to_load": "انقر على 'تصفح' لتحميل ملف PDF",

                # Page Adjustment Section
                "page_adjustment": "📊 تعديل الصفحة",
                "page_offset": "إزاحة الصفحة:",
                "offset_help": "(+/- للتعديل للصفحات الإضافية/المفقودة)",
                "smart_offset_detection": "🎯 كشف الإزاحة الذكي",
                "offset_info": "💡 انقر على أي إشارة مرجعية لمعاينة صفحتها، ثم استخدم الكشف التلقائي إذا لم تكن متوافقة",
                "auto_detect_offset": "🔍 كشف تلقائي للإزاحة من الإشارة المحددة",

                # Table Headers
                "bookmarks_table": "📋 الإشارات المرجعية",
                "title": "العنوان",
                "original": "الأصلي",
                "adjusted": "المعدل",
                "level": "المستوى",
                "preview": "معاينة",

                # File Selection Labels
                "file_selection": "📁 اختيار الملفات",
                "bookmark_file": "📖 ملف الإشارات المرجعية:",
                "select_pdf_placeholder": "اختر ملف PDF...",
                "select_bookmark_placeholder": "اختر ملف نصي بتنسيق 'العنوان - الصفحة'...",
                "load_files_preview": "🔄 تحميل الملفات والمعاينة",

                # Operation Groups
                "select_operation": "🔧 اختيار العملية",
                "file_operations": "📁 عمليات الملفات",
                "page_operations": "📄 عمليات الصفحات",
                "content_operations": "📝 عمليات المحتوى",

                # Watermark Positions
                "center": "الوسط",
                "top_left": "أعلى اليسار",
                "top_right": "أعلى اليمين",
                "bottom_left": "أسفل اليسار",
                "bottom_right": "أسفل اليمين",

                # Quality Options
                "high_quality": "عالية (2x)",
                "medium_quality": "متوسطة (1.5x)",
                "standard_quality": "قياسية (1x)",
                "both_methods": "كلاهما (مدمج + مُعرض)",
                "embedded_only": "الصور المدمجة فقط",
                "rendered_only": "الصفحات المُعرضة فقط",

                # Rotation Options
                "rotate_90": "90° في اتجاه عقارب الساعة",
                "rotate_180": "180°",
                "rotate_270": "270° في اتجاه عقارب الساعة",

                # Page Range and Numbers
                "page_range": "📄 نطاق الصفحات:",
                "all_pages": "جميع الصفحات",
                "page_numbers": "أرقام الصفحات",
                "page_count": "عدد الصفحات",
                "current_page": "الصفحة الحالية",
                "total_pages": "إجمالي الصفحات",

                # Placeholders
                "page_range_placeholder": "الكل أو 1,3,5-10",
                "watermark_text_placeholder": "أدخل نص العلامة المائية...",
                "select_pdf_to_extract": "اختر PDF لاستخراج الصفحات منه...",
                "select_pdf_to_delete": "اختر PDF لحذف الصفحات منه...",
                "select_pdf_for_blank": "اختر PDF لإدراج صفحات فارغة...",
                "select_pdf_for_watermark": "اختر ملف PDF لإضافة علامة مائية...",
                "select_pdf_for_compression": "اختر ملف PDF للضغط...",
                "select_pdf_for_editing": "اختر ملف PDF لتحرير الصفحات...",
                "extracted_bookmarks_placeholder": "ستظهر الإشارات المرجعية المستخرجة هنا...",
                "compression_results_placeholder": "ستظهر نتائج الضغط هنا...",

                # Group Box Titles
                "compression_options": "⚙️ خيارات الضغط",
                "compression_results": "📊 نتائج الضغط",
                "merge_info": "ℹ️ معلومات الدمج",
                "rotation_options": "🔄 خيارات التدوير",
                "margin_options": "📏 خيارات الهوامش",
                "extraction_options": "📤 خيارات الاستخراج",
                "output_options": "📁 خيارات الإخراج",

                # Help and Information
                "how_to_use": "<b>كيفية الاستخدام:</b>",
                "note": "<b>ملاحظة:</b>",
                "important": "<b>مهم:</b>",
                "tip": "<b>نصيحة:</b>",
                "warning": "<b>تحذير:</b>",

                # File Types in Dialogs
                "pdf_files_filter": "ملفات PDF (*.pdf)",
                "text_files_filter": "ملفات نصية (*.txt)",
                "json_files_filter": "ملفات JSON (*.json)",
                "all_files_filter": "جميع الملفات (*)",
                "image_files_filter": "ملفات الصور (*.png *.jpg *.jpeg)",

                # Compression Features
                "compression_features": "<b>ميزات الضغط:</b><br>• إزالة الكائنات والمراجع غير المستخدمة<br>• تحسين ضغط الصور<br>• تنظيف هيكل المستند<br>• تقليل حجم الملف مع الحفاظ على الجودة",

                # Merge Instructions
                "merge_instructions": "1. انقر على 'إضافة ملفات PDF' لاختيار عدة ملفات PDF<br>2. سيتم دمج الملفات بالترتيب الذي تظهر به في القائمة<br>3. استخدم 'إزالة' لحذف الملفات المحددة من القائمة<br>4. انقر على 'دمج ملفات PDF' لدمج جميع الملفات في ملف واحد",

                # Page Editing Help
                "margin_help": "يضيف مساحة بيضاء حول الصفحات. مفيد لملفات PDF بدون هوامش أو محتوى ضيق.",
                "rotation_help": "تدوير الصفحات المحددة بالزاوية المختارة.",
                "page_format_help": "<b>أمثلة على تنسيق الصفحات:</b><br>• صفحات منفردة: 1,3,5<br>• نطاقات الصفحات: 1-5,10-15<br>• تنسيق مختلط: 1,3,5-10,15<br><br><b>موضع الإدراج:</b> مكان إدراج الصفحات في PDF الهدف<br><b>الحذف:</b> إزالة الصفحات المحددة من PDF<br><b>الاستخراج:</b> إنشاء PDF جديد بالصفحات المحددة فقط",

                # Additional Interface Elements
                "enabled": "مُفعل",
                "disabled": "مُعطل",
                "selected": "محدد",
                "unselected": "غير محدد",
                "visible": "مرئي",
                "hidden": "مخفي",
                "expanded": "موسع",
                "collapsed": "مطوي",

                # Menu items
                "file_menu": "ملف",
                "close": "إغلاق",
                "exit": "خروج",
                "help_menu": "مساعدة",
                "about": "حول البرنامج",
                "update_available": "تحديث متوفر",
                "check_for_updates": "فحص التحديثات",
                "download_update": "تحميل التحديث",
                "remind_later": "تذكيرني لاحقاً",
                "skip_version": "تخطي هذا الإصدار",

                # Progress and loading messages
                "please_wait": "يرجى الانتظار...",
                "processing_images": "جاري معالجة الصور...",
                "compressing_pdf": "جاري ضغط PDF...",
                "rotating_pages": "جاري تدوير الصفحات...",
                "adding_margins": "جاري إضافة الهوامش...",
                "adding_watermark": "جاري إضافة العلامة المائية...",
                "extracting_pages": "جاري استخراج الصفحات...",
                "deleting_pages": "جاري حذف الصفحات...",
                "inserting_pages": "جاري إدراج الصفحات...",
                "operation_in_progress": "العملية قيد التنفيذ...",
                "completed": "اكتملت",

                # Page range help
                "page_range_help": "<b>كيفية تحديد الصفحات:</b><br>• <b>جميع الصفحات:</b> اتركه فارغاً أو اكتب 'الكل'<br>• <b>صفحة واحدة:</b> 5<br>• <b>صفحات متعددة:</b> 1,3,5,7<br>• <b>نطاق صفحات:</b> 1-10<br>• <b>نطاقات متعددة:</b> 1-5,8-12,15<br>• <b>مختلط:</b> 1,3,5-10,15,20-25",
                "page_range_examples": "أمثلة: 1,3,5 أو 1-10 أو 1,3,5-10",

                # Multiple file selection
                "select_multiple_files": "اختيار ملفات متعددة",
                "files_selected": "ملف محدد",
                "no_files_selected": "لم يتم اختيار ملفات",

                # Directory selection
                "select_output_directory": "اختيار مجلد الإخراج",
                "output_directory": "مجلد الإخراج:",
                "images_saved_to": "تم حفظ الصور في:",

                # Tab operations
                "extract_tab": "استخراج",
                "delete_tab": "حذف",
                "insert_tab": "إدراج",
                "merge_tab": "دمج",

                # Merge functionality
                "selected_files": "الملفات المحددة",
                "add_files": "إضافة ملفات",
                "remove_selected": "إزالة المحدد",
                "clear_all": "مسح الكل",
                "drag_to_reorder": "اسحب لإعادة الترتيب - سيتم دمج الملفات بالترتيب المعروض",
                "merge_options": "خيارات الدمج",
                "select_at_least_two_files": "يرجى اختيار ملفين على الأقل للدمج",
                "files_merged": "الملفات المدموجة",

                # File menu items
                "open_pdf": "فتح PDF",
                "recent_files": "الملفات الحديثة",
                "user_guide": "دليل المستخدم",

                # Page operations messages
                "no_pages_specified": "يرجى تحديد الصفحات",
                "confirm_deletion": "تأكيد الحذف",
                "confirm_delete_pages": "هل أنت متأكد من حذف",
                "pages": "الصفحات",
                "position": "الموضع",
                "count": "العدد",
                "save_pdf_with_blank": "حفظ PDF مع صفحات فارغة",
                "save_pdf_deleted_pages": "حفظ PDF مع الصفحات المحذوفة",
                "save_extracted_pages": "حفظ الصفحات المستخرجة",
                "blank_pages_inserted_successfully": "تم إدراج الصفحات الفارغة بنجاح",
                "pages_deleted_successfully": "تم حذف الصفحات بنجاح",
                "pages_extracted_successfully": "تم استخراج الصفحات بنجاح",

                # Enhanced menu items
                "quick_tools": "أدوات سريعة",
                "keyboard_shortcuts": "اختصارات لوحة المفاتيح",

                # Settings
                "settings_title": "⚙️ إعدادات التطبيق",
                "appearance": "🎨 المظهر",
                "language": "🌐 اللغة",
                "theme": "🎨 السمة",
                "light_mode": "☀️ الوضع الفاتح",
                "dark_mode": "🌙 الوضع الداكن",
                "arabic": "العربية",
                "english": "English",
                "weekend_settings": "📅 إعدادات عطلة نهاية الأسبوع",
                "weekend_days_label": "أيام عطلة نهاية الأسبوع:",
                "monday": "الاثنين",
                "tuesday": "الثلاثاء",
                "wednesday": "الأربعاء",
                "thursday": "الخميس",
                "friday": "الجمعة",
                "saturday": "السبت",
                "sunday": "الأحد",
                "apply_settings": "✅ تطبيق الإعدادات",
                "restart_required": "إعادة التشغيل مطلوبة",
                "restart_message": "يجب إعادة تشغيل التطبيق لتطبيق التغييرات. هل تريد إعادة التشغيل الآن؟",

                # History
                "history_title": "📋 سجل العمليات",
                "timestamp": "⏰ الوقت",
                "operation": "🔧 العملية",
                "input_files": "📁 الملفات المدخلة",
                "output_file": "📄 ملف الإخراج",
                "details": "📝 التفاصيل",
                "clear_history": "🗑️ مسح السجل",
                "export_history": "📤 تصدير السجل",
                "no_history": "لا يوجد سجل عمليات حتى الآن",

                # PDF Reading Progress
                "reading_progress": "📝 مُتتبع التعليقات والأنشطة",
                "reading_progress_desc": "تتبع تقدم قراءة ملفات PDF وتحليل التعليقات والتعليقات التوضيحية لمعرفة ما قرأته ومدى تقدمك في الدراسة",
                "select_folder": "📂 اختر مجلد",
                "scan_computer": "💻 فحص الكمبيوتر",
                "start_scan": "▶️ بدء الفحص",
                "stop_scan": "⏹️ إيقاف الفحص",
                "refresh_stats": "🔄 تحديث الإحصائيات",
                "total_pdfs": "📚 إجمالي ملفات PDF:",
                "annotated_pdfs": "📝 ملفات PDF مع التعليقات:",
                "total_annotations": "💬 إجمالي التعليقات:",
                "avg_intensity": "🔥 متوسط الكثافة:",
                "search_placeholder": "البحث بالاسم أو المسار...",
                "filter_all": "جميع ملفات PDF",
                "filter_annotated": "مع التعليقات فقط",
                "filter_recent": "المعدلة حديثاً (30 يوم)",
                "filter_high_activity": "نشاط عالي (10+ تعليقات)",
                "filter_low_activity": "نشاط منخفض (1-5 تعليقات)",
                "filter_no_annotations": "بدون تعليقات",
                "open_pdf": "📖 فتح PDF",
                "export_annotations": "📝 تصدير التعليقات",
                "study_timeline": "📅 الجدول الزمني للدراسة",
                "backup_data": "💾 نسخ احتياطي للبيانات",
                "restore_data": "📥 استعادة البيانات",
                "clear_all_data": "🗑️ مسح جميع البيانات",
                "export_list": "📤 تصدير القائمة",

                # Messages
                "file_not_found": "الملف غير موجود",
                "operation_completed": "تمت العملية بنجاح",
                "operation_failed": "فشلت العملية",
                "select_file": "يرجى اختيار ملف",
                "invalid_format": "تنسيق غير صحيح",
                "processing": "جاري المعالجة...",

                # Recent Books section
                "recent_books": "الكتب الحديثة",
                "recent_books_desc": "إدارة مكتبة الكتب الشخصية وتتبع تقدم القراءة",
                "add_book": "إضافة كتاب",
                "select_pdf_book": "اختر ملف PDF",
                "book_progress": "تقدم القراءة",
                "pages_read": "الصفحات المقروءة",
                "total_pages": "إجمالي الصفحات",
                "last_opened": "آخر فتح",
                "update_progress": "تحديث التقدم",
                "remove_book": "إزالة الكتاب",
                "book_details": "تفاصيل الكتاب",
                "reading_percentage": "نسبة القراءة",
                "file_size": "حجم الملف",
                "date_added": "تاريخ الإضافة",
                "no_books": "لا توجد كتب مضافة بعد",
                "book_added": "تم إضافة الكتاب بنجاح",
                "book_removed": "تم إزالة الكتاب",
                "progress_updated": "تم تحديث التقدم",
                "opened_at_last_page": "تم فتح: {book} في الصفحة {page}",

                # Enhanced Reading Statistics
                "reading_statistics_dashboard": "لوحة إحصائيات القراءة",
                "total_books_read": "إجمالي الكتب المقروءة",
                "total_pages_read": "إجمالي الصفحات المقروءة",
                "reading_streak_days": "سلسلة القراءة (أيام)",
                "time_range": "النطاق الزمني",
                "7_days": "7 أيام",
                "30_days": "30 يوم",
                "90_days": "90 يوم",
                "1_year": "سنة واحدة",
                "export_charts": "تصدير الرسوم البيانية",
                "charts": "الرسوم البيانية",
                "reading_progress_chart": "رسم بياني لتقدم القراءة عبر الزمن",
                "pages_read": "الصفحات المقروءة",
                "category_distribution_chart": "توزيع القراءة حسب الفئة",
                "weekly_reading_reports": "تقارير القراءة الأسبوعية",
                "generate_new_report": "إنشاء تقرير جديد",
                "weekly_reports": "التقارير الأسبوعية",
                "detailed_analytics": "التحليلات التفصيلية",
                "date": "التاريخ",
                "book": "الكتاب",
                "pages": "الصفحات",
                "category": "الفئة",
                "duration": "المدة",
                "refresh": "تحديث",
                "export_all_data": "تصدير جميع البيانات",
                "books_completed": "الكتب المكتملة",
                "total_pages": "إجمالي الصفحات",
                "reading_time": "وقت القراءة",
                "categories": "الفئات",
                "goals": "الأهداف",
                "weekly_report": "التقرير الأسبوعي",
                "save_weekly_report": "حفظ التقرير الأسبوعي",
                "reading_statistics": "إحصائيات القراءة",
                "no_data_available": "لا توجد بيانات متاحة",
                "close": "إغلاق",
                "feature_unavailable": "الميزة غير متاحة",
                "statistics_dependencies_missing": "الإحصائيات المحسنة تتطلب تبعيات إضافية. يرجى تثبيت matplotlib.",

                # Book Information Dialog - Arabic translations
                "book_information": "معلومات الكتاب",
                "reading_preparation": "إعداد القراءة",
                "words_per_page": "الكلمات في الصفحة",
                "estimated_reading_time": "وقت القراءة المقدر",
                "reading_difficulty": "صعوبة القراءة",
                "recommended_session": "الجلسة الموصى بها",
                "easy": "سهل",
                "medium": "متوسط",
                "hard": "صعب",
                "very_hard": "صعب جداً",
                "unknown": "غير معروف",
                "no_data": "لا توجد بيانات",
                "first_added": "تاريخ الإضافة",
                "file_information": "معلومات الملف",
                "file_path": "مسار الملف",
                "bookmarks": "الإشارات المرجعية",
                "seconds": "ثواني",
                "minutes": "دقائق",
                "hours": "ساعات",

                # Recent Books additional translations
                "reading_status": "حالة القراءة",
                "reading": "قيد القراءة",
                "to_read": "للقراءة",
                "completed": "مكتمل",
                "category": "الفئة",
                "uncategorized": "غير مصنف",
                "add_books": "إضافة كتب",
                "search_books": "البحث في الكتب...",
                "all": "الكل",
                "title": "العنوان",
                "progress": "التقدم",
                "rating": "التقييم",
                "recent": "الحديث",
                "clear_filters": "مسح المرشحات",
                "one_book": "كتاب واحد",
                "books": "كتب",
                "small_grid": "شبكة صغيرة",
                "medium_grid": "شبكة متوسطة",
                "large_grid": "شبكة كبيرة",
                "sort": "ترتيب",
                "open": "فتح",
                "preview": "معاينة",
                "rename": "إعادة تسمية",
                "edit_category": "تحرير الفئة",
                "toggle_star": "تبديل النجمة",
                "priority": "الأولوية",
                "normal": "عادي",
                "high": "عالي",
                "urgent": "عاجل",
                "remove": "إزالة",
                "measure_reading_speed": "قياس سرعة القراءة",
                "add_favorite": "إضافة للمفضلة",
                "remove_favorite": "إزالة من المفضلة",
                "open_book": "فتح الكتاب",

                # Home page grid
                "home_title": "عدة القارئ للملفات الرقمية",
                "home_subtitle": "اختر أداة للبدء في العمل مع ملفات PDF",
                "back_to_home": "العودة للرئيسية",

                # Security removal feature
                "security_removal_title": "إزالة حماية PDF",
                "security_removal_desc": "إزالة قيود الحماية والتشفير من ملفات PDF",
                "remove_security": "إزالة الحماية",
                "pdf_password": "كلمة مرور PDF",
                "password_optional": "كلمة المرور (اختيارية)",
                "security_removed": "تم إزالة الحماية بنجاح",
                "security_removal_failed": "فشل في إزالة الحماية",
                "processing": "جاري المعالجة...",
                "select_file": "اختيار الملف",
                "select_pdf_file": "اختر ملف PDF",
                "pdf_file": "ملف PDF",

                # Reading Speed Measurement
                "reading_speed_meter": "قياس سرعة القراءة",
                "reading_speed_meter_desc": "قس سرعة قراءتك بالكلمات في الدقيقة مع اختبار الفهم والحصول على توصيات لتحسين الأداء",
                "select_pdf_for_speed": "اختر ملف PDF لقياس السرعة",
                "or_select_from_recent": "أو اختر من الكتب الحديثة",
                "recent_books_for_speed": "الكتب الحديثة",
                "no_recent_books": "لا توجد كتب حديثة",
                "analyzing_pdf": "جارٍ تحليل PDF...",

                # OCR PDF to Text
                "ocr_pdf_to_text": "🔍 تحويل PDF إلى نص",
                "ocr_pdf_to_text_desc": "تحويل ملفات PDF المسحوبة ضوئياً إلى نص قابل للتحرير باستخدام تقنية التعرف الضوئي على الحروف المتقدمة",
                "select_pdf_file_ocr": "اختر ملف PDF للتحويل بالـ OCR",
                "ocr_file_selection": "📁 اختيار ملف PDF",
                "ocr_settings": "⚙️ إعدادات OCR",
                "processing_controls": "🚀 أدوات التحكم في المعالجة",
                "processing_status": "📊 حالة المعالجة",
                "results_preview": "📄 معاينة النتائج",
                "select_pdf_placeholder_ocr": "اختر ملف PDF لتحويل OCR...",
                "no_file_selected": "لم يتم اختيار ملف",
                "browse": "تصفح",
                "remove_newlines_docx": "إزالة أسطر جديدة من مخرجات DOCX",
                "remove_newlines_tooltip": "مفيد لجعل عدد صفحات DOCX يطابق عدد صفحات PDF",
                "output_directory": "مجلد الإخراج",
                "same_as_input": "نفس مجلد الملف المدخل",
                "start_ocr_conversion": "🔍 بدء تحويل OCR",
                "cancel_processing": "❌ إلغاء المعالجة",
                "processing_progress": "تقدم المعالجة",
                "current_status": "الحالة الحالية",
                "ready_to_start": "جاهز للبدء",
                "extracted_text_preview": "معاينة النص المستخرج",
                "no_text_extracted": "لم يتم استخراج نص بعد",
                "output_files": "ملفات الإخراج",
                "open_txt_file": "📝 فتح ملف TXT",
                "open_docx_file": "📝 فتح ملف DOCX",
                "open_folder": "📁 فتح المجلد",
                "ocr_processing_required": "معالجة OCR مطلوبة",
                "ocr_processing_question": "يبدو أن هذا PDF يحتوي على صور ممسوحة ضوئياً بدلاً من نص قابل للاستخراج.\n\nهل تريد استخدام OCR (التعرف الضوئي على الحروف) لاستخراج النص؟\n\nملاحظة: معالجة OCR قد تستغرق عدة دقائق حسب حجم PDF.",
                "convert_pdf_online": "🌐 تحويل PDF عبر الإنترنت",
                "convert_pdf_online_tooltip": "افتح Colab لتحويل ملفات PDF الممسوحة ضوئيًا إلى نص (TXT/DOCX)",
                "online_ocr_instructions_title": "تحويل PDF عبر الإنترنت (Colab)",
                "online_ocr_instructions": "استخدم أداة OCR عبر الإنترنت (Colab) لتحويل ملفات PDF الممسوحة إلى TXT/DOCX:\n\n1) افتح الدفتر (Notebook).\n2) شغّل خلية الإعداد (تثبيت المتطلبات).\n3) ارفع ملف PDF (من لوحة الملفات اليسرى) أو قم بتركيب Google Drive.\n4) حدّد مسار ملف الإدخال.\n5) شغّل خلية التحويل.\n6) قم بتنزيل ملفات .txt و .docx من مجلد الإخراج.\n\nبعد التنزيل: استخدم التطبيق لتحليل ملفات TXT/DOCX أو استيراد النتائج تلقائياً.",
                "open_colab": "فتح Colab",
                "import_ocr_results_btn": "📥 استيراد نتائج OCR",
                "searching_downloads": "جاري البحث في مجلد التنزيلات عن نتائج OCR...",
                "ocr_results_moved": "تم العثور على ملفات OCR ونقلها إلى المجلد المفضل",
                "ocr_results_not_found": "لم يتم العثور على ملفات OCR المطابقة. يمكنك اختيارها يدوياً.",
                "choose_ocr_results": "اختر ملفات نتائج OCR",
                "select_txt_file": "اختر ملف TXT",
                "select_docx_file": "اختر ملف DOCX",
                "file_name_mismatch": "اسم الملف لا يطابق اسم ملف PDF الأصلي",
                "default_ocr_dir": "المجلد الافتراضي لنتائج OCR",
                "set_default_ocr_dir": "تعيين مجلد افتراضي لنتائج OCR",
                "change_default_ocr_dir": "تغيير المجلد الافتراضي لنتائج OCR",
                "preferred_ocr_dir_saved": "تم حفظ المجلد الافتراضي لنتائج OCR",
                "image_pdf_detected_use_ocr_title": "ملف PDF يعتمد على الصور",
                "image_pdf_detected_use_ocr_message": "لا يحتوي هذا الملف على نص قابل للاستخراج.\n\nيمكنك استخدام 'تحويل PDF عبر الإنترنت' لإنتاج TXT/DOCX ثم استيرادها للتحليل",
                "ocr_processing_dialog": "معالجة PDF بـ OCR...\nقد يستغرق هذا عدة دقائق.",
                "ocr_completed": "اكتمل OCR",
                "ocr_success_message": "اكتملت معالجة OCR بنجاح!\n\nتم استخراج {total_words:,} كلمة من {total_pages} صفحة.\nالمتوسط: {avg_words:.0f} كلمة لكل صفحة.\n\nالملفات المحفوظة:\n• {txt_file}\n• {docx_file}",
                "ocr_warning": "تحذير OCR",
                "ocr_no_text_warning": "اكتملت معالجة OCR ولكن لم يتم استخراج نص.\nقد يحتوي PDF على صور بدون نص قابل للقراءة.",
                "ocr_processing_failed": "فشلت معالجة OCR",
                "ocr_error_message": "فشلت معالجة OCR:\n\n{error_message}\n\nيمكنك تجربة إدخال عدد الكلمات يدوياً بدلاً من ذلك.",
                "manual_word_count_input": "إدخال عدد الكلمات يدوياً",
                "manual_word_count_prompt": "نظراً لفشل استخراج النص التلقائي، يرجى إدخال\nمتوسط عدد الكلمات لكل صفحة يدوياً:\n\n(يمكنك عد الكلمات في صفحة نموذجية وإدخال هذا الرقم)",
                "manual_input_accepted": "تم قبول الإدخال اليدوي",
                "using_manual_words": "استخدام {words_per_page} كلمة لكل صفحة لحساب سرعة القراءة.",
                "initializing_ocr": "تهيئة معالجة OCR...",
                "preparing_pdf_ocr": "تحضير PDF لـ OCR...",
                "running_ocr_conversion": "تشغيل تحويل OCR...",
                "ocr_conversion_completed": "اكتمل تحويل OCR بنجاح!",
                "poppler_not_found": "لم يتم العثور على Poppler-utils. يرجى تثبيت poppler-utils:\n\nخيارات Windows:\n1. تحميل من: https://github.com/oschwartz10612/poppler-windows/releases/\n   استخراج وإضافة إلى PATH\n\n2. تثبيت عبر conda: conda install poppler\n\n3. تثبيت عبر pip: pip install poppler-utils\n\nبعد التثبيت، أعد تشغيل التطبيق.",
                "tahweel_not_found": "لم يتم العثور على حزمة Tahweel. يرجى تثبيتها أولاً.",
                "google_drive_auth_title": "إعداد مصادقة Google Drive",
                "google_drive_auth_message": "يتطلب Tahweel الوصول إلى Google Drive API لمعالجة OCR.\n\nلديك خياران:\n\n1. استخدام حساب الخدمة (موصى به للأتمتة):\n   - إنشاء مشروع Google Cloud\n   - تفعيل Google Drive API\n   - إنشاء بيانات اعتماد حساب الخدمة\n   - تحميل ملف JSON وحفظه كـ 'service_account.json'\n\n2. استخدام المصادقة التفاعلية:\n   - تشغيل أمر tahweel يدوياً أولاً\n   - إكمال تدفق OAuth في المتصفح\n   - سيتم حفظ بيانات الاعتماد للاستخدام المستقبلي\n\nهل تريد المتابعة بدون بيانات اعتماد (قد يطلب المصادقة)؟",
                "word_analysis_complete": "تم تحليل الكلمات",
                "avg_words_per_page": "متوسط الكلمات في الصفحة:",
                "total_pages": "إجمالي الصفحات:",
                "start_reading_session": "ابدأ جلسة القراءة",
                "reading_timer": "مؤقت القراءة:",
                "pages_read": "الصفحات المقروءة:",
                "finish_reading": "انهِ القراءة واحسب السرعة",
                "reading_speed_results": "نتائج سرعة القراءة",
                "your_reading_speed": "سرعة قراءتك:",
                "wpm": "كلمة/دقيقة",
                "total_time_taken": "إجمالي الوقت المستغرق:",
                "minutes": "دقيقة",
                "pdf_options": "خيارات PDF",
                "pdf_reading_tips": "💡 نصائح لقياس سرعة القراءة بدقة:\n\n• اقرأ بطريقتك الطبيعية المعتادة\n• تجنب القراءة بصوت عالٍ أو تحريك الشفاه\n• ركز على فهم المحتوى وليس فقط السرعة\n• اختر نصوص مناسبة لمستواك\n• تأكد من الإضاءة الجيدة وراحة العينين",

                # Bookmark Copy Feature
                "copy_bookmarks_title": "نسخ الفهارس بين ملفات PDF",
                "copy_bookmarks_desc": "انسخ فهرس المحتويات من ملف PDF الأصلي إلى النسخة المحسنة بـ OCR مع نفس أرقام الصفحات",
                "original_pdf": "ملف PDF الأصلي:",
                "enhanced_pdf": "ملف PDF المحسن:",
                "select_original_pdf": "اختر ملف PDF الأصلي (مصدر الفهارس)",
                "select_enhanced_pdf": "اختر ملف PDF المحسن (هدف الفهارس)",
                "copy_bookmarks_button": "نسخ الفهارس",
                "select_both_files": "يرجى اختيار كلا الملفين أولاً",
                "no_bookmarks_found": "لم يتم العثور على فهارس في الملف الأصلي",
                "page_count_mismatch": "عدم تطابق عدد الصفحات",
                "page_count_warning": "الملف المحسن ({enhanced} صفحة) يحتوي على صفحات أقل من الأصلي ({original} صفحة). هل تريد المتابعة؟",
                "bookmarks_copied_success": "تم نسخ {count} فهرس بنجاح!\n\nتم الحفظ باسم: {filename}",
                "open_directory": "فتح المجلد",
                "open_directory_question": "هل تريد فتح مجلد الملف المحفوظ؟",
                "bookmark_copy_failed": "فشل في نسخ الفهارس:\n{error}",
                "copy_bookmarks_utility": "نسخ الفهارس",
                "copy_bookmarks_tooltip": "نسخ فهرس المحتويات من ملف PDF إلى آخر",
                "file_selection": "اختيار الملفات",
                "seconds": "ثانية",
                "recommendation": "التوصية:",
                "comprehension_test": "اختبار الفهم",
                "answer_questions": "أجب على الأسئلة التالية لاختبار فهمك:",
                "question": "السؤال",
                "your_answer": "إجابتك:",
                "submit_answers": "إرسال الإجابات",
                "comprehension_score": "نتيجة الفهم:",
                "correct_answers": "إجابات صحيحة",
                "out_of": "من",
                "excellent_comprehension": "فهم ممتاز! تقرأ بسرعة جيدة مع فهم عالي.",
                "good_comprehension": "فهم جيد. حاول تحسين التوازن بين السرعة والفهم.",
                "needs_improvement": "يحتاج تحسين. ركز أكثر على الفهم أثناء القراءة.",
                "slow_but_thorough": "قراءة بطيئة ولكن دقيقة. هذا جيد للمواد المعقدة.",
                "average_speed": "سرعة متوسطة مع فهم جيد. استمر في الممارسة.",
                "fast_reader": "قارئ سريع! تأكد من عدم التضحية بالفهم.",
                "speed_reader": "قراءة سريعة جداً. مناسبة للمراجعة السريعة.",
                "reset_test": "إعادة تعيين الاختبار",
                "new_reading_session": "جلسة قراءة جديدة",
                "measure_reading_speed": "قياس سرعة القراءة",
                "prepare_book": "تحضير الكتاب",
                "current_book": "الكتاب الحالي",
                "activity": "النشاط",
                "word_analysis_complete": "تحليل الكلمات مكتمل",
                "comprehension_test": "اختبار الفهم",
                "reading_sessions": "جلسات القراءة",
                "prepare_book_desc": "تحليل الكتاب مسبقاً لتوفير الوقت في جلسات القراءة المستقبلية",
                "book_prepared": "تم تحضير الكتاب بنجاح",
                "analysis_cached": "تم حفظ تحليل الكلمات للاستخدام المستقبلي",

                # Additional Reading Speed Localization
                "reading_speed_guide": "دليل قياس سرعة القراءة",
                "word_analysis_complete": "تحليل الكلمات مكتمل",
                "start_reading_session": "بدء جلسة القراءة",
                "book_ready": "الكتاب جاهز",
                "information": "معلومات",
                "select_file_first": "يرجى اختيار ملف أولاً",
                "ocr_results_not_found": "لم يتم العثور على نتائج OCR",
                "select_txt_file": "اختر ملف TXT",
                "select_docx_file": "اختر ملف DOCX",
                "file_name_mismatch": "أسماء الملفات غير متطابقة",
                "ocr_results_moved": "تم نقل نتائج OCR بنجاح",
                "convert_pdf_online": "تحويل PDF عبر الإنترنت",
                "convert_pdf_online_tooltip": "فتح Colab لتحويل ملفات PDF المسحوبة ضوئياً إلى نص",
                "import_ocr_results_btn": "استيراد نتائج OCR",
                "choose_ocr_results": "اختر نتائج OCR",
                "set_default_ocr_dir": "تعيين مجلد OCR الافتراضي",
                "change_default_ocr_dir": "تغيير مجلد OCR الافتراضي",
                "preferred_ocr_dir_saved": "تم حفظ مجلد OCR المفضل",
                "ocr_processing_required": "معالجة OCR مطلوبة",
                "ocr_processing_question": "هذا الملف يحتاج إلى معالجة OCR. هل تريد المتابعة؟",
                "online_ocr_instructions": "تعليمات OCR عبر الإنترنت",
                "online_ocr_instructions_title": "تعليمات OCR عبر الإنترنت",
                "open_colab": "فتح Colab",
                "file_not_found": "الملف غير موجود",

                # Reading Sessions Table
                "reading_sessions_desc": "جلسات قياس سرعة القراءة مع إحصائيات مفصلة:",
                "time_spent": "الوقت المستغرق",
                "efficiency": "الكفاءة",
                "total_sessions": "إجمالي الجلسات",
                "average_wpm": "متوسط الكلمات/دقيقة",
                "best_wpm": "أفضل كلمات/دقيقة",
                "no_sessions_yet": "لا توجد جلسات قراءة بعد",
                "export_sessions": "تصدير الجلسات",
                "data_exported": "تم تصدير البيانات بنجاح إلى",

            # Enhanced Reading Speed Workflow
            "analyzing_pdf": "تحليل ملف PDF...",
            "analysis_results": "نتائج التحليل",
            "sample_text_preview": "معاينة النص النموذجي:",
            "adjust_words_per_page": "تعديل عدد الكلمات في الصفحة:",
            "words": "كلمات",
            "confirm_preparation": "تأكيد والتحضير",
            "preparation_options": "خيارات التحضير",
            "choose_preparation_method": "اختر طريقة التحضير",
            "pdf_analysis_limited": "هذا الملف يحتوي على نص محدود قابل للاستخراج. يرجى اختيار طريقة التحضير:",
            "use_default_estimate": "استخدام التقدير الافتراضي",
            "default_estimate_desc": "استخدم نطاق نموذجي من 120-150 كلمة في الصفحة للمستندات العادية.",
            "use_ocr_conversion": "استخدام تحويل OCR",
            "ocr_conversion_desc": "تحويل PDF إلى تنسيق Word باستخدام OCR عبر الإنترنت للحصول على عدد دقيق للكلمات.",
            "manual_entry_desc": "عد الكلمات في صفحة نموذجية وأدخلها يدوياً.",
            "use_default": "استخدام الافتراضي",
            "use_ocr": "استخدام OCR",
            "ocr_guidance": "دليل تحويل OCR",
            "ocr_conversion_guide": "دليل تحويل OCR",
            "ocr_instructions": """
اتبع هذه الخطوات لتحويل ملف PDF باستخدام OCR عبر الإنترنت:

1. انقر على زر "فتح أداة OCR" أدناه لفتح المحول عبر الإنترنت
2. ارفع ملف PDF إلى خدمة OCR
3. انتظر حتى يكتمل التحويل
4. حمل مستند Word المحول
5. افتح مستند Word وعد الكلمات في صفحة نموذجية
6. ارجع هنا وأدخل عدد الكلمات في الصفحة يدوياً

هذه الطريقة توفر أدق عدد للكلمات لملفات PDF الممسوحة ضوئياً.
            """,
            "open_ocr_tool": "فتح أداة OCR",
            "enter_ocr_results": "إدخال نتائج OCR",
            "ocr_results_desc": "بعد تحويل OCR، عد الكلمات في صفحة نموذجية وأدخلها أدناه:",
            "confirm_ocr_results": "تأكيد نتائج OCR",
            "automatic_analysis": "التحليل التلقائي",
            "default_estimate": "التقدير الافتراضي",
            "ocr_conversion": "تحويل OCR",
            "saved_configuration": "الإعدادات المحفوظة",
            "preparation_method": "طريقة التحضير",
            "open_pdf_manually": "يرجى فتح ملف PDF يدوياً:",
            "reading_speed_results": "نتائج سرعة القراءة",
            "your_reading_speed": "سرعة قراءتك",
            "total_time_taken": "إجمالي الوقت المستغرق",
            "total_words_read": "إجمالي الكلمات المقروءة",
            "new_reading_session": "جلسة قراءة جديدة",
            "pages_read": "الصفحات المقروءة:",
            "finish_reading": "إنهاء القراءة",
            "hours": "ساعة",
            "minutes": "دقيقة",
            "seconds": "ثانية",
            "total_books": "إجمالي الكتب",
            "total_sessions": "إجمالي الجلسات",
            "total_pages_read": "إجمالي الصفحات المقروءة",
            "average_speed": "متوسط السرعة",
            "best_speed": "أفضل سرعة",
            "total_pages": "إجمالي الصفحات",
            "start_page": "صفحة البداية",
            "resume_reading": "استئناف القراءة",
            "pause_reading": "إيقاف مؤقت",
            "stop_reading": "إيقاف القراءة",


                # Training mode and Pages/Minute tab
                "training_mode": "وضع التدريب",
                "elapsed_time": "الوقت المنقضي",
                "start": "ابدأ",
                "end_page": "صفحة النهاية",
                "what_page_reached": "إلى أي صفحة وصلت؟",
                "ppm_simple_desc": "أدخل صفحة البداية، ابدأ المؤقت، ثم عند الإيقاف أدخل صفحة النهاية.",
                "mode_simple": "بسيط",
                "sample_pages": "صفحات العينة:",
                "extract_sample": "استخراج العينة",
                "ai_prompt": "نص موجه للذكاء الاصطناعي",
                "copy_prompt": "نسخ الموجه",
                "avg_words_per_page": "متوسط الكلمات/صفحة:",
                "mode_sample": "تقدير بالعينة",
                "using_stored_counts": "يتم استخدام عدد الكلمات المخزن لكل صفحة",
                "mode_accurate": "الأدق",
                "pages_per_minute": "صفحات/دقيقة",
                "parse_and_save": "تحليل وحفظ",
                "info": "معلومات",
                "missing_data": "البيانات مفقودة أو لم يتم تحميل الملف.",
                "saved": "تم الحفظ",
                "saved_page_counts": "تم حفظ عدد الكلمات لكل صفحة.",
                "save_failed": "تعذر الحفظ",
                "parse_failed": "تعذر التحليل",
                "done": "تم",
                "sample_saved_to": "تم حفظ العينة في:",
                "open_pdf_first": "يرجى فتح ملف PDF أولاً.",
                "paste_csv_here": "الصق بيانات CSV هنا:",
                "csv_placeholder": "page_number,word_count\n1,250\n2,240\n...",
                "ppm_label": "صفحات/دقيقة",
                "wpm_label": "كلمة/دقيقة",


            # Reading Speed Levels and Advice
            "beginner_reader": "قارئ مبتدئ",
            "beginner_advice": "مارس القراءة بانتظام لتحسين سرعتك. ركز على تقليل النطق الداخلي.",
            "average_reader": "قارئ متوسط",
            "average_advice": "سرعة قراءة جيدة! جرب تقنيات القراءة السريعة للوصول إلى المستوى التالي.",
            "good_reader": "قارئ جيد",
            "good_advice": "سرعة قراءة ممتازة! أنت فوق المتوسط. استمر في الممارسة للحفاظ على هذا المستوى.",
            "expert_reader": "قارئ خبير",
            "expert_advice": "سرعة قراءة متميزة! أنت في المستوى الأعلى من القراء.",

            # PDF Loading Messages
            "existing_configuration": "الإعدادات الموجودة",
            "use_existing_config": "تم تحضير هذا الكتاب من قبل. استخدام الإعدادات الموجودة؟",
            "pdf_loaded": "تم تحميل PDF",
            "pdf_loaded_successfully": "تم تحميل PDF بنجاح!",
            "pages": "صفحات",
            "click_prepare_to_start": "انقر على \"تحضير الكتاب\" لبدء التحليل.",

            # Fast Reading Trainer
            "fast_reading_trainer": "مدرب القراءة السريعة",
            "fast_reading_desc": "تدرب على سرعة القراءة من خلال عرض الكلمات في مجموعات بوتيرة محكومة. حسن قدرتك على الكلمات في النظرة الواحدة (WPG).",
            "select_document": "اختيار المستند",
            "select_text_file": "اختر ملف نصي (.txt, .docx, .pdf)",
            "select_txt": "اختيار ملف TXT",
            "select_docx": "اختيار ملف DOCX",
            "select_docx_file": "اختر ملف DOCX",
            "select_pdf_trainer": "اختيار PDF للتدريب",
            "training_settings": "إعدادات التدريب",
            "words_per_glance": "الكلمات في النظرة:",
            "target_speed": "السرعة المستهدفة:",
            "font_size": "حجم الخط:",
            "training_controls": "أدوات التحكم في التدريب",
            "start_training": "بدء التدريب",
            "pause": "إيقاف مؤقت",
            "resume": "استئناف",
            "stop": "إيقاف",
            "progress": "التقدم",
            "ready_to_start": "جاهز لبدء التدريب",
            "select_document_to_start": "اختر مستنداً لبدء التدريب",
            "warning": "تحذير",
            "no_text_found": "لم يتم العثور على نص في المستند",
            "no_extractable_text": "لم يتم العثور على نص قابل للاستخراج في ملف PDF",
            "text_too_short": "النص قصير جداً للتدريب (الحد الأدنى 10 كلمات مطلوب)",
            "ready_to_train": "جاهز لبدء التدريب!",
            "words_loaded": "كلمة محملة",
            "text_loaded": "تم تحميل النص",
            "text_loaded_successfully": "تم تحميل النص بنجاح!",
            "adjust_settings_and_start": "اضبط إعداداتك وانقر على بدء التدريب!",
            "no_text_loaded": "لم يتم تحميل نص للتدريب",
            "training_active": "التدريب نشط",
            "training_resumed": "تم استئناف التدريب",
            "training_paused": "تم إيقاف التدريب مؤقتاً",
            "training_stopped": "تم إيقاف التدريب",
            "training_complete": "اكتمل التدريب!",
            "congratulations": "تهانينا! لقد أكملت جلسة التدريب.",
            "statistics": "الإحصائيات",
            "words_read": "كلمة مقروءة",
            "words": "كلمات",
            "minutes": "دقائق",
            "font_family": "عائلة الخط",
            "effective_reading_speed": "سرعة القراءة الفعلية",
            "every": "كل",
            "seconds": "ثانية",
            "word_groups_per_minute": "مجموعة كلمات في الدقيقة",
            "adjust_speed_if_needed": "اضبط السرعة إذا لزم الأمر",
            "continue_previous_session": "متابعة الجلسة السابقة",
            "save_session": "حفظ الجلسة",
            "save_training_progress": "حفظ تقدم التدريب؟",
            "current_progress": "التقدم الحالي",
            "session_resumed": "تم استئناف الجلسة",
            "session_resumed_successfully": "تم استئناف الجلسة بنجاح",
            "progress": "التقدم",
            "error_resuming_session": "خطأ في استئناف الجلسة",
            "no_saved_sessions": "لا توجد جلسات تدريب محفوظة",
            "saved_training_sessions": "جلسات التدريب المحفوظة",
            "select_session_to_resume": "اختر جلسة للاستئناف:",
            "complete": "مكتمل",
            "resume_session": "استئناف الجلسة",
            "delete_session": "حذف الجلسة",
            "select_session_first": "يرجى اختيار جلسة أولاً",
            "confirm_delete": "تأكيد الحذف",
            "delete_session_confirm": "حذف جلسة التدريب هذه؟",
            "file_not_found": "الملف غير موجود",
            "unsupported_file_type": "نوع ملف غير مدعوم",
            "session_saved_successfully": "تم حفظ جلسة التدريب بنجاح!",
            "all_activities": "جميع الأنشطة",
            "reading_measurements": "قياسات القراءة",
            "training_sessions": "جلسات التدريب",
            "reading_measurement": "قياس القراءة",
            "training_session": "جلسة التدريب",
            "book_title": "عنوان الكتاب",
            "training": "تدريب",
            "resume": "استئناف",
            "session_resumed": "تم استئناف الجلسة!",
            "last_page": "آخر صفحة",
            "file_not_found": "الملف غير موجود. ربما تم نقله أو حذفه.",
            "type": "النوع",
            "file": "الملف",
            "progress_time": "التقدم/الوقت",
            "settings_details": "الإعدادات/التفاصيل",
            "actions": "الإجراءات",
            "no_activities_yet": "لا توجد أنشطة بعد",
            "show": "عرض",
            "completed": "مكتمل",
            "complete": "مكتمل",
                "current_book": "الكتاب الحالي",
                "reading_sessions": "جلسات القراءة",
                "output_settings": "إعدادات الإخراج",
                "select_output_location": "اختر مكان حفظ الملف",
                "output_file": "ملف الإخراج",

                # Section descriptions
                "pdf_viewer_desc": "عرض وتعليق ملفات PDF بأدوات متقدمة",
                "bookmark_manager_desc": "إدارة وتنظيم إشارات PDF المرجعية",
                "bookmark_extractor_desc": "استخراج الإشارات المرجعية من ملفات PDF",
                "reading_progress_desc": "تتبع تقدم القراءة والتعليقات التوضيحية",

                # Statistics Dashboard
                "statistics_dashboard": "لوحة الإحصائيات",
                "statistics_dashboard_desc": "📈 تحليلات شاملة للقراءة والإحصائيات مع الرسوم البيانية",
                "key_metrics": "📊 المقاييس الرئيسية",
                "overview_statistics": "📊 إحصائيات عامة",
                "visual_analytics": "📈 التحليلات المرئية",
                "advanced_analytics": "🔬 التحليلات المتقدمة",
                "total_books": "إجمالي الكتب",
                "read_today": "قُرئ اليوم",
                "pages_today": "صفحات اليوم",
                "reading_streak": "سلسلة القراءة",
                "avg_pages_day": "متوسط الصفحات/يوم",
                "total_time": "إجمالي الوقت",
                "completion_rate": "معدل الإنجاز",
                "favorite_category": "الفئة المفضلة",
                "categories": "الفئات",
                "reading_velocity": "سرعة القراءة",
                "monthly_goal": "الهدف الشهري",
                "productive_time": "وقت الذروة",
                "consistency_score": "درجة الاستمرارية",
                "books_by_category": "الكتب حسب الفئة",
                "daily_reading_progress": "تقدم القراءة اليومي",
                "export": "📤 تصدير",
                "refresh": "🔄 تحديث",

                "page_operations_desc": "استخراج وحذف وإدراج الصفحات، دمج الملفات، وتدويرها وقصّها وإضافة الهوامش",
                "watermark_desc": "إضافة علامات مائية لملفات PDF",
                "extract_images_desc": "استخراج الصور من ملفات PDF",
                "extract_text_desc": "استخراج محتوى النص من ملفات PDF",
                "merge_pdfs_desc": "دمج عدة ملفات PDF في ملف واحد",
                "split_pdfs_desc": "تقسيم ملف PDF إلى عدة ملفات",
                "compress_desc": "تقليل حجم ملف PDF",
                "page_editing_desc": "تدوير وقص وتحرير الصفحات",
                "settings_desc": "تكوين إعدادات التطبيق",
                "history_desc": "عرض سجل العمليات",
                "navigation": "التنقل",
                "refresh": "تحديث",
                "clear_history": "مسح السجل",

                # PDF Comments section
                "show_hide_help": "إظهار/إخفاء المساعدة",
                "ready": "جاهز",
                "no_directory_selected": "لم يتم اختيار مجلد",
                "include_subdirectories": "تضمين المجلدات الفرعية",
                "statistics": "📈 الإحصائيات",
                "search_placeholder": "البحث بالاسم أو المسار...",
                "filter_label": "📂 تصفية:",
                "sort_label": "ترتيب حسب:",
                "all_pdfs": "جميع ملفات PDF",
                "with_annotations_only": "مع التعليقات فقط",
                "recently_modified": "معدل حديثاً (30 يوم)",
                "high_activity": "نشاط عالي (10+ تعليقات)",
                "low_activity": "نشاط منخفض (1-5 تعليقات)",
                "no_annotations": "بدون تعليقات",
                "last_scanned": "آخر فحص",
                "file_name": "اسم الملف",
                "last_modified": "آخر تعديل",
                "annotations_count": "عدد التعليقات",
                "reading_intensity": "كثافة القراءة",
                "file_name_col": "اسم الملف",
                "path_col": "المسار",
                "pages_col": "الصفحات",
                "annotations_col": "التعليقات",
                "intensity_col": "الكثافة",
                "last_modified_col": "آخر تعديل",
                "last_scanned_col": "آخر فحص",
                "open_pdf": "📖 فتح PDF",
                "export_annotations": "📝 تصدير التعليقات",
                "study_timeline": "📅 جدول الدراسة",
                "backup_data": "💾 نسخ احتياطي",
                "restore_data": "📥 استعادة البيانات",
                "clear_all_data": "🗑️ مسح جميع البيانات",

                # Recent Books additional keys
                "export": "تصدير",
                "grid_size": "حجم الشبكة",
                "small": "صغير",
                "medium": "متوسط",
                "large": "كبير",
                "total": "المجموع",
                "starred": "المفضلة",
                "priority": "الأولوية",
                "avg_progress": "متوسط التقدم",
                "no_cover": "لا يوجد غلاف",
                "available": "متاح",

                # Context menu items
                "edit_name": "تعديل الاسم",
                "edit_category": "تعديل الفئة",
                "reading_status": "حالة القراءة",
                "cover_management": "إدارة الغلاف",
                "add_star": "إضافة نجمة",
                "remove_star": "إزالة النجمة",
                "set_priority": "تحديد الأولوية",
                "quick_update": "تحديث سريع",
                "update_progress": "تحديث التقدم",
                "open_external": "فتح بتطبيق خارجي (افتراضي)",
                "open_internal": "فتح بالعارض الداخلي",
                "reading_analytics": "تحليلات القراءة",
                "remove_book": "إزالة الكتاب",
                "upload_custom_cover": "رفع غلاف مخصص",
                "remove_custom_cover": "إزالة الغلاف المخصص",
                "currently_reading": "قيد القراءة حالياً",
                "normal": "عادي",
                "high": "عالي",
                "urgent": "عاجل",

                # Filter options
                "all_books": "جميع الكتب",
                "reading": "قيد القراءة",
                "to_read": "للقراءة",
                "completed": "مكتملة",
                "priority_books": "كتب الأولوية",
                "sort_by": "ترتيب حسب",

                # Sort options
                "last_opened": "آخر فتح",
                "title": "العنوان",
                "progress": "التقدم",
                "date_added": "تاريخ الإضافة",

                # Category options
                "category": "الفئة",
                "all_categories": "جميع الفئات",
                "uncategorized": "غير مصنف",

                # TOC Preparation
                "prepare_toc": "تحضير فهرس المحتويات",
                "toc_preparation_title": "تحضير فهرس المحتويات من PDF",
                "step1_extract_pages": "الخطوة 1: استخراج صفحات الفهرس",
                "toc_page_range": "نطاق صفحات الفهرس:",
                "to": "إلى",
                "extract_toc_pages": "استخراج صفحات الفهرس",
                "step2_ai_extraction": "الخطوة 2: استخراج النص بالذكاء الاصطناعي",
                "ai_studio_instructions": "1. انقر على 'فتح AI Studio' أدناه\n2. انقر على 'نسخ النص التوجيهي' لنسخ التعليمات\n3. الصق التعليمات في AI Studio\n4. ارفع ملف PDF المستخرج\n5. انسخ النتيجة والصقها في الخطوة 3",
                "open_ai_studio": "فتح AI Studio",
                "copy_ai_prompt": "نسخ النص التوجيهي",
                "step3_format_toc": "الخطوة 3: تنسيق الفهرس",
                "paste_ai_result": "الصق نتيجة AI Studio هنا:",
                "toc_paste_placeholder": "الصق النص المستخرج من AI Studio هنا...\n\nمثال:\nالْمُقَدِّمَةُ - ٥\nالنُّسَخُ الْمُعْتَمَدَةُ فِي التَّحْقِيقِ - ٨\nكِتَابُ الطَّهَارَةِ - ١٥\nبَابُ الْمِيَاهِ - ١٥",
                "format_toc": "تنسيق الفهرس",
                "formatted_result": "النتيجة المنسقة",
                "use_bookmarks": "استخدام العلامات المرجعية",
                "ai_toc_prompt": "استخراج فهرس محتويات من ملف PDF بتنسيق مرقم للإشارات المرجعية (Bookmarks)\n\nاستخرج فهرس المحتويات من صفحات ملف PDF المرفق.\n\nالتعليمات:\n\n1. استخرج العناوين وأرقام الصفحات.\n2. استخدم مستويين من العناوين:\n   * المستوى الأول: العناوين الرئيسية (مثل الفصول أو الأجزاء).\n   * المستوى الثاني: العناوين الفرعية داخل كل قسم.\n3. تنسيق المستوى الأول:\n   * اكتب العنوان الرئيسي متبوعًا برقم الصفحة.\n   * الصيغة: العنوان - رقم الصفحة\n4. تنسيق المستوى الثاني:\n   * أضف مسافة بادئة (indentation).\n   * استخدم ترقيمًا بالتنسيق X.Y حيث X هو رقم القسم الرئيسي و Y هو رقم العنوان الفرعي المتسلسل داخله.\n   * يبدأ ترقيم الأقسام الرئيسية من 1.\n   * الصيغة النهائية: X.Y العنوان - رقم الصفحة\n5. افصل بين كل قسم رئيسي وما يتبعه من العناوين الفرعية بسطر فارغ كما في المثال.\n6. استخدم نفس لغة النص في PDF.\n7. لا تضف أي نص آخر غير الفهرس المنسق.\n\nمثال للناتج المطلوب:\n\nالْمُقَدِّمَةُ - ٥\n1.1 مَفَاهِيمُ الْمَشْرُوعِ وَأَهْدَافُهُ - ٨\n1.2 مَنْهَجِيَّةُ الْعَمَلِ وَخُطُواتُ التَّطْبِيقِ - ١٢\n\nفَصْلُ النِّظَامِ الدَّاخِلِيِّ - ١٥\n2.1 بَيَانُ التَّعَارِيفِ وَالْمُصْطَلَحَاتِ - ١٧\n2.2 تَنْظِيمُ الْعَمَلِ الْيَوْمِيِّ - ٢١\n2.3 إِدَارَةُ الْمَوَارِدِ وَالتَّقَارِيرِ - ٢٤\n2.4 نَمَاذِجُ الْوَثَائِقِ وَالإِجْرَاءَاتِ - ٢٨\n2.5 التَّقْيِيمُ وَالتَّحْسِينُ الْمُسْتَمِرُّ - ٣١",

                # TOC Preparation Steps
                "toc_preparation_steps": "📝 خطوات تحضير فهرس المحتويات:",
                "toc_step1": "الخطوة 1: اختر ملف PDF أعلاه",
                "toc_step2": "الخطوة 2: حدد نطاق الصفحات لفهرس المحتويات",
                "toc_step3": "الخطوة 3: استخرج صفحات فهرس المحتويات إلى PDF منفصل",
                "toc_step4": "الخطوة 4: استخدم AI Studio لاستخراج النص من PDF فهرس المحتويات",
                "toc_step4_final": "الخطوة 4: الصق النص المستخرج وقم بتنسيقه إلى إشارات مرجعية",

                # TOC Range Dialog
                "select_toc_range": "تحديد نطاق فهرس المحتويات",
                "select_toc_range_title": "📋 اختر صفحات فهرس المحتويات",
                "toc_range_selection": "تحديد النطاق",
                "toc_range_instructions": "استخدم المعاينة للتنقل عبر PDF وتحديد صفحات فهرس المحتويات. يمكنك استخدام الأزرار السريعة لتعيين الصفحة الحالية كبداية أو نهاية.",
                "start_page": "صفحة البداية:",
                "end_page": "صفحة النهاية:",
                "set_as_start": "تعيين كبداية",
                "set_as_end": "تعيين كنهاية",
                "selected_range": "النطاق المحدد",
                "total_pages_selected": "إجمالي الصفحات المحددة",
                "extract_selected_pages": "استخراج الصفحات المحددة",
                "toc_range_hint": "💡 استخدم الأزرار أعلاه لاستخراج صفحات فهرس المحتويات بعد تحديد النطاق المناسب",

                # TOC Page Selection Dialog
                "select_toc_pages": "📋 تحديد صفحات فهرس المحتويات",
                "set_toc_range": "تعيين نطاق فهرس المحتويات",
                "toc_page_selection": "اختيار صفحات فهرس المحتويات",
                "toc_page_selection_desc": "استخدم هذا الحوار لتصفح PDF وتحديد الصفحات التي تحتوي على فهرس المحتويات بصرياً",
                "navigate_and_select": "تصفح وحدد",
                "toc_range_selected": "تم تحديد نطاق فهرس المحتويات",

                # Improved Workflow Steps
                "step1_load_pdf": "الخطوة 1: تحميل الملف الذي تريد إضافة فهرس له",
                "step2_select_toc": "الخطوة 2: تحديد صفحات فهرس الكتاب",
                "step3_extract_toc": "الخطوة 3: استخراج صفحات الفهرس",
                "step4_format_process": "الخطوة 4: استخدام AI Studio وتحميل الإشارات المرجعية",
                "step4_verify_insert": "الخطوة 4: التحقق من الإشارات المرجعية وإدراجها",

                "load_pdf_primary": "📄 تحميل ملف PDF",
                "select_toc_pages_step": "📋 تحديد صفحات فهرس الكتاب",
                "format_and_load": "🔄 تنسيق الفهرس وتحميل الملفات",
                "already_have_bookmarks": "إذا كان لديك فهرس جاهز في ملف نصي",
                "upload_to_ai_studio": "ارفع الصفحات المستخرجة إلى AI Studio واستخدم الأمر المنسوخ",
                "paste_ai_result_here": "الصق نتيجة AI Studio هنا وسيتم التنسيق تلقائياً",
                "verify_bookmarks": "تحقق من الإشارات المرجعية باستخدام زر 'معاينة'",
                "fix_page_mismatch": "إصلاح عدم تطابق الصفحات",
                "navigate_to_correct_page": "انتقل للصفحة الصحيحة واضغط 'إصلاح'",
                "enable_mismatch_detection": "تفعيل كشف عدم تطابق الصفحات",
                "fix": "إصلاح",
                "auto_calculate": "حساب تلقائي",
                "confirm": "تأكيد",
                "actions": "الإجراءات",
                "delete": "حذف",
                "delete_bookmark": "حذف الإشارة المرجعية",
                "confirm_delete": "هل أنت متأكد من حذف هذه الإشارة المرجعية؟",
                "split_by_bookmarks": "تقسيم PDF حسب الإشارات المرجعية",
                "split_by_bookmarks_desc": "تقسيم PDF إلى ملفات منفصلة بناءً على هيكل الإشارات المرجعية",
                "split_level_1_only": "تقسيم حسب المستوى الأول فقط (الفصول الرئيسية)",
                "split_all_levels": "تقسيم حسب جميع المستويات",
                "output_directory": "مجلد الإخراج",
                "file_naming": "تسمية الملفات",
                "use_bookmark_titles": "استخدام عناوين الإشارات المرجعية كأسماء ملفات",
                "use_sequential_numbers": "استخدام أرقام متسلسلة (001، 002، إلخ)",
                "start_splitting": "بدء التقسيم",
                "bookmark_levels": "مستويات الإشارات المرجعية",
                "additional_options": "خيارات إضافية",
                "include_original_bookmarks": "تضمين الإشارات المرجعية الأصلية في الملفات المقسمة",
                "create_index_file": "إنشاء ملف فهرس يحتوي على قائمة الملفات المقسمة",
                "insert_bookmarks_pdf": "إدراج الإشارات المرجعية في PDF",
                "split_pdf_bookmarks": "تقسيم PDF حسب الإشارات المرجعية",
                "split_options": "خيارات التقسيم",
                "level1_bookmarks": "الإشارات المرجعية من المستوى الأول (الفصول الرئيسية)",
                "level2_bookmarks": "الإشارات المرجعية من المستوى الثاني (الفصول الفرعية)",
                "all_levels": "جميع المستويات",
                "output_directory": "مجلد الإخراج",
                "select_output_dir": "اختر مجلد الإخراج:",
                "choose_output_dir": "اختر مجلد الإخراج...",
                "split_pdf": "تقسيم PDF",
                "cancel": "إلغاء",
                "load_bookmarks_preview": "تحميل الإشارات المرجعية للمعاينة",

                # New Redesigned Bookmark Workflow
                "step2_choose_method": "الخطوة 2: اختيار طريقة مصدر الإشارات المرجعية",
                "choose_bookmark_method_desc": "اختر الطريقة التي تريد استخدامها لإنشاء الإشارات المرجعية:",
                "method_toc_pages": "تحديد صفحات فهارس - لاستخراج فهرس المحتويات من صفحات PDF",
                "method_text_file": "اختيار ملف نص + إمكانية لصق النص - لاستخدام إشارات مرجعية خارجية",
                "step2_bookmark_extraction": "الخطوة 2: لصق الإشارات المرجعية",
                "step2_guidance_text": "الصق المحتوى هنا:",
                "step2_help_text": "إذا لم يكن لديك إشارات مرجعية - ",
                "prepare_bookmarks_button": "تحضير الإشارات المرجعية",
                "step3_generate_bookmarks": "الخطوة 3: إنشاء الإشارات المرجعية",
                "text_file_desc": "اختر ملف نصي يحتوي على الإشارات المرجعية بتنسيق 'العنوان - الصفحة':",
                "or_separator": "أو",
                "paste_text_desc": "الصق النص المحتوي على الإشارات المرجعية مباشرة:",
                "format_bookmarks": "تنسيق الإشارات المرجعية",
                "no_toc_text": "يرجى لصق النص أو اختيار ملف نصي أولاً",
                "no_bookmark_source": "يرجى اختيار ملف نصي أو لصق النص أولاً",
                "step4_verify_insert": "الخطوة 4: التحقق من الإشارات المرجعية والإدراج",
                "verify_bookmarks": "تحقق من الإشارات المرجعية باستخدام المعاينة، ثم اضغط على 'إدراج الإشارات المرجعية' لحفظها في PDF.",

                # Success messages
                "bookmark_source_ready": "مصدر الإشارات المرجعية جاهز للتنسيق!",
                "bookmark_source_ready_desc": "يمكنك الآن الضغط على 'تحميل الإشارات المرجعية' للمتابعة.",
                "toc_text_ready": "نص فهرس المحتويات جاهز للتنسيق!",

                # Message box titles and content
                "success": "نجح",
                "pdf_loaded_success": "تم تحميل PDF بنجاح!",
                "total_pages": "إجمالي الصفحات:",
                "use_preview_instruction": "استخدم المعاينة لتحديد صفحات فهرس المحتويات، ثم انقر على 'استخراج صفحات فهرس المحتويات' لتحديد النطاق.",
                "toc_range_selected": "تم تحديد نطاق فهرس المحتويات",
                "selected_range": "النطاق المحدد:",
                "pages": "صفحات",
                "can_extract_now": "يمكنك الآن استخراج فهرس المحتويات لهذا النطاق.",
                "toc_extracted_success": "تم استخراج صفحات فهرس المحتويات بنجاح!",
                "file": "الملف:",
                "folder_will_open": "سيتم فتح المجلد المحتوي تلقائياً.",
                "prompt_copied": "تم نسخ النص التوجيهي",
                "ai_prompt_copied": "تم نسخ النص التوجيهي للذكاء الاصطناعي إلى الحافظة!",
                "paste_in_ai_studio": "الصقه في AI Studio وارفع صفحات فهرس المحتويات المستخرجة.",

                # Recent Books - Arabic translations
                "add_books": "إضافة كتب",
                "small_grid": "صغير",
                "medium_grid": "متوسط",
                "large_grid": "كبير",
                "search_books": "البحث في الكتب...",
                "status": "الحالة:",
                "category": "الفئة:",
                "sort": "الترتيب:",
                "all": "الكل",
                "reading": "قيد القراءة",
                "to_read": "للقراءة",
                "completed": "مكتمل",
                "recent": "الأحدث",
                "title": "العنوان",
                "progress": "التقدم",
                "rating": "التقييم",
                "clear_filters": "مسح المرشحات",
                "no_books": "لا توجد كتب",
                "one_book": "كتاب واحد",
                "books": "كتب",
                "uncategorized": "غير مصنف",

                # Context Menu - Arabic translations
                "open": "فتح",
                "preview": "معاينة",
                "rename": "إعادة تسمية",
                "edit_category": "تحرير الفئة",
                "update_progress": "تحديث التقدم",
                "reading_status": "حالة القراءة",
                "toggle_star": "تبديل النجمة",
                "priority": "الأولوية",
                "normal": "عادي",
                "high": "عالي",
                "urgent": "عاجل",
                "remove": "إزالة",
                "resumed_reading": "استئناف القراءة",
                "resumed_from_page": "تم الاستئناف من الصفحة",

                # Reading Speed Tab
                "books_placeholder_text": "ستظهر الكتب هنا بعد تكوينها\nلقياس سرعة القراءة وإكمال الجلسات.",
                "close_button": "إغلاق",
                "loading_trends": "جاري تحميل تحليل الاتجاهات...",
                "book_title_placeholder": "عنوان الكتاب",
                "filename_placeholder": "اسم_الملف.pdf",
                "total_sessions_label": "إجمالي الجلسات:",
                "average_speed_label": "متوسط السرعة:",
                "best_speed_label": "أفضل سرعة:",
                "total_time_label": "إجمالي الوقت:",
                "use_manual_value": "استخدام القيمة اليدوية",
                "cancel_button": "إلغاء",

                # Bug Reporting System
                "help": "مساعدة",
                "report_bug": "🐛💡 الإبلاغ عن خطأ أو اقتراح",
                "bug_report_title": "تقرير الخطأ / اقتراح الميزة",
                "bug_report_desc": "ساعدنا في التحسين من خلال الإبلاغ عن الأخطاء أو اقتراح ميزات جديدة. ملاحظاتك قيمة!",
                "report_type": "نوع التقرير",
                "bug_report": "تقرير خطأ",
                "feature_suggestion": "اقتراح ميزة",
                "enhancement": "تحسين",
                "question": "سؤال",
                "bug_information": "معلومات الخطأ",
                "bug_title": "عنوان الخطأ",
                "bug_title_placeholder": "وصف موجز للمشكلة...",
                "bug_description": "الوصف",
                "bug_description_placeholder": "وصف تفصيلي للخطأ...\n\nخطوات إعادة الإنتاج:\n1. \n2. \n3. \n\nالسلوك المتوقع:\n\nالسلوك الفعلي:",
                "username_optional": "اسم المستخدم (اختياري)",
                "username_placeholder": "اسمك أو اسم المستخدم...",
                "email_optional": "البريد الإلكتروني (اختياري)",
                "email_placeholder": "your.email@example.com",
                "severity": "الخطورة",
                "severity_low": "منخفض",
                "severity_medium": "متوسط",
                "severity_high": "عالي",
                "severity_critical": "حرج",
                "category": "الفئة",
                "category_ui": "مشكلة واجهة",
                "category_functionality": "خطأ وظيفي",
                "category_performance": "أداء",
                "category_crash": "انهيار",
                "category_other": "أخرى",
                "contact_info": "معلومات الاتصال (اختيارية)",
                "system_info_note": "سيتم تضمين معلومات النظام تلقائياً للمساعدة في تشخيص المشكلة.",
                "submit_bug": "إرسال تقرير الخطأ",
                "submitting": "جاري الإرسال...",
                "bug_submitted_success": "تم إرسال تقرير الخطأ بنجاح! شكراً لملاحظاتك.",
                "bug_submission_failed": "فشل إرسال تقرير الخطأ. يرجى المحاولة مرة أخرى.",
                "validation_error": "خطأ في التحقق",
                "bug_title_required": "يرجى إدخال عنوان الخطأ.",
                "bug_description_required": "يرجى إدخال وصف الخطأ.",
                "page_name": "الصفحة",
            },
            "en": {
                # Main window
                "app_title": "عدة القارئ للملفات الرقمية - PDF Toolkits",
                "ready_status": "Ready - Select a tab to begin working with PDFs",

                # Tab names
                "bookmark_manager": "📖 Bookmark Manager",
                "bookmark_extractor": "📤 Bookmark Extractor",
                "chapter_weight_analyzer": "📊 Study Plan Preparation",
                "page_operations": "📄 Page Operations",
                "watermark": "💧 Watermark",
                "extract_images": "🖼️ Extract Images",
                "extract_text": "📝 Extract Text",
                "merge_pdfs": "📚 Merge PDFs",
                "split_pdfs": "✂️ Split PDF",
                "compress": "🗜️ Compress",
                "page_editing": "✂️ Page Editing",
                "pdf_viewer": "PDF Viewer & Annotator",
                "settings": "⚙️ Settings",
                "history": "📋 History",

                # Common buttons
                "browse": "📁 Browse",
                "cancel": "Cancel",
                "ok": "OK",
                "save": "💾 Save",
                "load": "🔄 Load",
                "extract": "📤 Extract",
                "delete": "🗑️ Delete",
                "insert": "✚ Insert",
                "merge": "📚 Merge",
                "compress_btn": "🗜️ Compress",
                "rotate": "🔄 Rotate",
                "add_margins": "📏 Add Margins",

                # Common labels
                "pdf_file": "📄 PDF File:",
                "page_range": "📄 Page Range:",
                "pages": "📄 Pages:",
                "position": "📍 Position:",
                "output": "📁 Output:",
                "status": "Status:",
                "success": "Success!",
                "error": "Error",
                "warning": "Warning",
                "information": "Information",

                # Bookmark Manager
                "bookmark_manager_desc": "🎯 <b>Bookmark Management:</b> Insert bookmarks from text files into PDFs with automatic offset detection",
                "file_selection": "📁 File Selection",
                "bookmark_file": "📖 Bookmark File:",
                "select_pdf_placeholder": "Select PDF file...",
                "select_bookmark_placeholder": "Select text file with 'Title - Page' format...",
                "load_files_preview": "🔄 Load Files & Preview",
                "page_adjustment": "📊 Page Adjustment",
                "page_offset": "Page Offset:",
                "offset_help": "(+/- to adjust for extra/missing pages)",
                "smart_offset_detection": "🎯 Smart Offset Detection",
                "offset_info": "💡 Click any bookmark to preview its page, then use auto-detect if misaligned",
                "auto_detect_offset": "🔍 Auto-Detect Offset from Selected Bookmark",
                "apply_correction": "✅ Apply Correction",
                "correction_applied": "Correction Applied",
                "manual_correction": "🎯 Manual Correction",
                "set_correct_page": "Set Correct Page",
                "current_viewing": "Currently viewing page",
                "bookmark_should_be": "Bookmark should be on page",
                "confirm_correction": "Confirm Correction",
                "bookmarks_table": "📋 Bookmarks",
                "title": "Title",
                "original": "Original",
                "adjusted": "Adjusted",
                "level": "Level",
                "preview": "Preview",
                "insert_bookmarks": "📚 Insert Bookmarks into PDF",
                "pdf_preview": "📄 PDF Preview",
                "prev": "◀ Prev",
                "next": "Next ▶",
                "no_pdf_loaded": "No PDF loaded",
                "pdf_preview_placeholder": "PDF preview will appear here",
                "click_browse_to_load": "Click 'Browse' to load a PDF file",
                "navigation": "Navigation",
                "hide_sidebar": "Hide Sidebar",
                "show_sidebar": "Show Sidebar",
                "page_thumbnails": "Page Thumbnails",
                "no_bookmarks": "No Bookmarks",
                "no_file_loaded": "No File Loaded",
                "page": "Page",
                "load_new_pdf": "Load Different File",
                "load_different_pdf": "Load Different PDF",
                "show_controls": "Show Controls",
                "hide_controls": "Hide Controls",
                "tools": "Tools",
                "zoom": "Zoom",
                "style": "Style",
                "select_tool": "Select Tool",
                "annotation_size": "Annotation Size",
                "load_new": "Load New",

                # Navigation tooltips
                "first_page": "First Page",
                "previous_page": "Previous Page",
                "next_page": "Next Page",
                "last_page": "Last Page",

                # Additional labels
                "page_label": "Page:",
                "no_offset_needed": "No offset needed - bookmark is correctly positioned",
                "error_loading_page": "Error loading page",
                "failed_to_render": "Failed to render page",
                "text_format_tooltip": "TXT: Plain text, JSON: Structured data, DOCX: Microsoft Word document",
                "show_bookmark_analytics": "Show Bookmark Weight Distribution",
                "close_analytics": "Close Analytics",
                "error_title": "Error",
                "failed_to_load_pdf": "Failed to load PDF",

                # Page Operations
                "page_operations_desc": "📄 <b>Page Operations:</b> Extract/delete/insert pages, merge files, rotate, crop, and add margins",
                "extract_pages": "📤 Extract Pages",
                "delete_pages": "🗑️ Delete Pages",
                "insert_blank_pages": "📄 Insert Blank Pages",
                "pages_to_extract": "📄 Pages to extract:",
                "pages_to_delete": "📄 Pages to delete:",
                "insert_at_position": "📍 Insert at position:",
                "number_of_pages": "📄 Number of pages:",
                "extract_pages_btn": "📤 Extract Pages",
                "delete_pages_btn": "🗑️ Delete Pages",
                "insert_blank_btn": "📄 Insert Blank Pages",

                # Watermark
                "watermark_desc": "💧 <b>Watermark Tool:</b> Add text watermarks to PDF pages or remove existing watermarks",
                "watermark_text": "💧 Watermark Text:",
                "watermark_position": "📍 Position:",
                "font_size": "📏 Font Size:",
                "add_watermark": "💧 Add Watermark",
                "remove_watermark": "🗑️ Remove Watermark",
                "add_watermark_section": "✚ Add Watermark",
                "remove_watermark_section": "🗑️ Remove Watermark",
                "watermark_removed_successfully": "Watermark removed successfully",
                "watermark_removal_info": "This enhanced feature attempts to remove various types of watermarks from PDF files including: logos, text, URLs (like UPDF, etc.). Please review the result to ensure no important content was removed.",
                "removal_options": "Removal Options",
                "aggressive_mode": "Aggressive Mode",
                "aggressive_mode_tooltip": "Removes more suspicious elements, but may remove important content",
                "remove_urls": "Remove URLs and websites",
                "center": "Center",
                "top_left": "Top Left",
                "top_right": "Top Right",
                "bottom_left": "Bottom Left",
                "bottom_right": "Bottom Right",

                # Image Extraction
                "image_extraction_desc": "🖼️ <b>Image Extraction:</b> Extract all images from PDF files with enhanced detection",
                "extraction_method": "🔧 Method:",
                "quality": "📊 Quality:",
                "both_methods": "Both (Embedded + Rendered)",
                "embedded_only": "Embedded Images Only",
                "rendered_only": "Rendered Pages Only",
                "high_quality": "High (2x)",
                "medium_quality": "Medium (1.5x)",
                "standard_quality": "Standard (1x)",
                "extract_all_images": "🖼️ Extract All Images",

                # Text Extraction
                "text_extraction_desc": "📝 <b>Text Extraction:</b> Extract text content from PDF files in various formats with batch processing support",
                "output_format": "📋 Output Format:",
                "extract_text_btn": "📝 Extract Text",
                "processing_mode": "Processing Mode:",
                "single_file": "Single File",
                "batch_processing": "Batch Processing",
                "select_multiple_pdfs": "📁 Select Multiple PDFs",
                "clear_selection": "🗑️ Clear Selection",
                "output_directory": "Output Directory:",
                "select_output_directory": "Select Output Directory",
                "batch_extract_btn": "📝 Extract Text (Batch)",
                "docx_format": "Word Document (DOCX)",
                "txt_format": "Plain Text (TXT)",
                "json_format": "Structured Data (JSON)",

                # Merge PDFs
                "merge_desc": "📚 <b>PDF Merge:</b> Combine multiple PDF files into a single document",
                "select_pdfs_merge": "📚 Select PDFs to Merge",
                "add_pdfs": "✚ Add PDFs",
                "remove": "✖ Remove",
                "merge_pdfs_btn": "📚 Merge PDFs",
                "merge_info": "ℹ️ Merge Information",
                "how_to_use": "<b>How to use:</b>",
                "merge_instructions": "1. Click 'Add PDFs' to select multiple PDF files<br>2. Files will be merged in the order they appear in the list<br>3. Use 'Remove' to delete selected files from the list<br>4. Click 'Merge PDFs' to combine all files into one",

                # Split PDF
                "split_pdfs": "✂️ Split PDF",
                "split_desc": "✂️ <b>PDF Split:</b> Divide a PDF file into multiple files at specified page numbers",
                "select_pdf_to_split": "📄 Select PDF to Split",
                "split_points": "✂️ Split Points",
                "split_points_label": "Page numbers to split at (comma-separated):",
                "split_points_placeholder": "Example: 5, 10, 15",
                "preview_split_points": "👁️ Preview Split Points",
                "split_pdf_btn": "✂️ Split PDF",
                "split_info": "ℹ️ Split Information",
                "split_instructions": "1. Select a PDF file to split<br>2. Enter page numbers where you want to split the file (comma-separated)<br>3. Click 'Preview Split Points' to see the pages where splitting will occur<br>4. Click 'Split PDF' to create separate files",
                "split_preview_title": "Split Points Preview",
                "page_preview": "Page {0} Preview",
                "split_at_page": "Split at Page {0}",
                "output_files_will_be": "{0} files will be created:",
                "file_part": "Part {0}",
                "pages_range": "Pages {0} - {1}",
                "select_output_directory": "Select Output Directory",
                "save_split_pdfs": "Save Split PDFs",
                "pdfs_split_successfully": "PDF split successfully",
                "split_complete": "Successfully created {0} files",
                "invalid_split_points": "Invalid split points",
                "enter_valid_page_numbers": "Please enter valid page numbers separated by commas",
                "page_numbers_out_of_range": "Some page numbers are out of valid range (1-{0})",
                "no_pdf_selected_for_split": "No PDF selected for splitting",
                "select_pdf_file_first": "Please select a PDF file first",
                "total_pages": "Total Pages: {0}",

                # Compress
                "compress_desc": "🗜️ <b>PDF Compression:</b> Reduce PDF file size while maintaining quality",
                "compression_options": "⚙️ Compression Options",
                "compression_features": "<b>Compression Features:</b><br>• Remove unused objects and references<br>• Optimize image compression<br>• Clean up document structure<br>• Reduce file size while maintaining quality",
                "compress_pdf_btn": "🗜️ Compress PDF",
                "compression_results": "📊 Compression Results",

                # PDF Viewer & Editor
                "pdf_viewer_desc": "View PDF files with advanced editing and annotation tools",
                "load_pdf": "📄 Load PDF",
                "pdf_navigation": "🧭 PDF Navigation",
                "zoom_controls": "🔍 Zoom Controls",
                "zoom_in": "🔍+ Zoom In",
                "zoom_out": "🔍- Zoom Out",
                "fit_page": "📄 Fit Page",
                "fit_width": "↔️ Fit Width",
                "previous_page": "⬅️ Previous Page",
                "next_page": "➡️ Next Page",
                "go_to_page": "🎯 Go to Page",
                "annotation_tools": "✏️ Annotation Tools",
                "text_annotation": "📝 Text Annotation",
                "rich_text_annotation": "Rich Text Annotation",
                "highlight_text": "🖍️ Highlight Text",
                "underline_text": "📏 Underline Text",
                "draw_rectangle": "⬜ Draw Rectangle",
                "draw_circle": "⭕ Draw Circle",
                "draw_arrow": "➡️ Draw Arrow",
                "sticky_note": "📌 Sticky Note",
                "freehand_draw": "✏️ Freehand Draw",
                "annotation_color": "🎨 Annotation Color",
                "annotation_size": "📏 Annotation Size",
                "ok": "OK",
                "save_annotations": "💾 Save Annotations",
                "clear_annotations": "🗑️ Clear Annotations",
                "page_info": "Page {current} of {total}",
                "invalid_page": "Invalid Page",
                "page_out_of_range": "Page {page} is out of range (1-{total})",

                # Page Editing
                "page_editing_desc": "✂️ <b>Page Editing:</b> Rotate pages, crop content, and add white space margins",
                "rotate_pages": "🔄 Rotate Pages",
                "rotation": "🔄 Rotation:",
                "rotate_90": "90° Clockwise",
                "rotate_180": "180°",
                "crop_pages": "✂️ Crop Pages",
                "crop_box": "Crop Box:",
                "crop_pages_btn": "✂️ Crop Pages",
                "save_cropped_pdf": "Save Cropped PDF",
                "invalid_crop_box": "Invalid crop box",
                "pages_cropped_successfully": "Pages cropped successfully",
                "pages_range_label": "Pages:",

                "rotate_270": "270° Clockwise",
                "rotate_pages_btn": "🔄 Rotate Pages",
                "add_white_margins": "📏 Add White Space Margins",
                "margin_size": "📏 Margin Size:",
                "points": " points",
                "margin_help": "Adds white space around pages. Useful for PDFs with no margins or tight content.",
                "add_margins_btn": "📏 Add Margins",

                # Common messages
                "select_file_first": "Please select a file first",
                "no_pages_specified": "Please specify pages",
                "invalid_page_format": "Invalid page number format",
                "operation_successful": "Operation completed successfully!",
                "operation_failed": "Operation failed",
                "confirm_deletion": "Confirm Deletion",
                "are_you_sure": "Are you sure?",
                "pages_will_be_deleted": "The specified pages will be deleted",
                "missing_files": "Missing Files",
                "select_both_files": "Please select both files",
                "no_images_found": "No images found",
                "images_found": "Images found",
                "text_extracted": "Text extracted",
                "compression_completed": "Compression completed",
                "size_reduction": "Size reduction",
                "space_saved": "Space saved",
                "original_size": "Original size",
                "compressed_size": "Compressed size",

                # Bookmark Extractor
                "bookmark_extractor_desc": "📤 <b>Bookmark Extractor:</b> Extract existing bookmarks from PDFs and save to text files",
                "extract_bookmarks": "📤 Extract Bookmarks",
                "extracted_bookmarks": "📋 Extracted Bookmarks",
                "save_to_text": "💾 Save to Text File",
                "no_bookmarks_found": "No bookmarks found",
                "bookmarks_extracted": "Bookmarks extracted",
                "bookmarks_saved": "Bookmarks saved",

                # Chapter Weight Analyzer
                "chapter_weight_desc": "Analyze chapter distribution and create smart reading plans based on chapter weights",
                "load_bookmarks_btn": "📖 Load Bookmarks",
                "analyze_weights_btn": "📊 Analyze Weights",
                "generate_plan_btn": "📅 Generate Reading Plan",
                "analysis_options": "⚙️ Analysis Options",
                "include_levels": "Include Levels:",
                "level_1_only": "Level 1 Only",
                "levels_1_2": "Levels 1 & 2",
                "all_levels": "All Levels",
                "chapter_weights_table": "📊 Chapter Weights",
                "chapter_title": "Chapter Title",
                "start_page": "Start Page",
                "end_page": "End Page",
                "page_count": "Page Count",
                "weight_percent": "Weight %",
                "reading_plan_section": "📅 Study Plan Preparation",
                "algorithm": "Algorithm",
                "algorithm_weight_based": "Weight-based (smart merge)",
                "algorithm_direct_pages": "Direct pages (proportional)",
                "plan_duration": "Plan Duration (days):",
                "start_date": "Start Date",
                "end_date": "End Date",
                "skip_weekends": "Skip Weekends",
                "reading_plan_table": "📅 Study Schedule",
                "assigned_days": "Assigned Days",
                "date_range": "Date Range",
                "statistics_insights": "📈 Statistics & Insights",
                "total_chapters": "Total Chapters:",
                "total_pages": "Total Pages:",
                "avg_chapter_length": "Avg Chapter Length:",
                "longest_chapter": "Longest Chapter:",
                "shortest_chapter": "Shortest Chapter:",
                "export_options": "💾 Export Options",
                "export_weights_csv": "Export Weights (CSV)",
                "export_weights_json": "Export Weights (JSON)",
                "export_plan_csv": "Export Plan (CSV)",
                "export_plan_json": "Export Plan (JSON)",
                "export_plan_text": "Export Plan (Text)",
                "export_weights_excel": "Export Weights (Excel)",
                "export_plan_excel": "Export Plan (Excel)",
                "export_weights_markdown": "Export Weights (Markdown)",
                "export_plan_markdown": "Export Plan (Markdown)",
                "export_plan_obsidian": "Export Plan (Obsidian)",
                "export_all_formats": "📦 Export All Formats",
                "export_pie_chart": "Export Pie Chart",
                "export_bar_chart": "Export Bar Chart",
                "export_weight_chart": "Export Comparison Chart",
                "weights_exported": "Weights exported successfully",
                "plan_exported": "Plan exported successfully",
                "chart_exported": "Chart exported successfully",
                "all_formats_exported": "All formats exported successfully",
                "no_weights_to_export": "No weights to export. Analyze bookmarks first.",
                "select_chapters": "Select Chapters",
                "select_chapters_instruction": "Select the chapters you want to include in your study plan.\nUncheck chapters like TOC, prefaces, etc. that you want to skip.",
                "select_all": "Select All",
                "deselect_all": "Deselect All",
                "no_chapters_selected": "No chapters selected. Please select at least one chapter.",
                "no_plan_to_export": "No plan to export. Generate a reading plan first.",

                # Report export labels (English)
                "report_reading_plan": "READING PLAN",
                "report_total_target_duration": "Total Target Duration",
                "report_actual_assigned_duration": "Actual Assigned Duration",
                "report_start_date": "Start Date",
                "report_average_daily_pages": "Average Daily Pages (Active Days)",
                "report_block": "Block",
                "report_days": "Days",
                "report_day": "Day",
                "report_chapter": "Chapter",
                "report_includes": "Includes",
                "report_pages": "Pages",
                "report_duration": "Duration",
                "report_weight": "Weight",
                "report_plan_created": "Plan created",
                "report_to": "to",
                "report_days_unit": "day(s)",
                "plan_summary": "Plan Summary",
                "reading_schedule": "Reading Schedule",
                "total_reading_blocks": "Total Reading Blocks",
                "daily_progress": "Daily Progress",
                "reading_days": "Reading Days",
                "quick_reference": "Quick Reference",
                "reading_block": "Reading Block",
                "generated": "Generated",
                "dates": "Dates",
                "copy_plan_to_clipboard": "Copy Plan to Clipboard",
                "plan_copied_to_clipboard": "Plan copied to clipboard successfully",

                # Export tooltips (English)
                "tooltip_export_plan_excel": "Export reading plan to Excel file with professional formatting and organized tables",
                "tooltip_export_plan_json": "Export reading plan to JSON format for programmatic use or integration with other apps",
                "tooltip_export_plan_obsidian": "Export reading plan to Obsidian-compatible Markdown with daily checkboxes for tracking",
                "tooltip_export_plan_text": "Export reading plan to simple, easy-to-read text file",
                "tooltip_export_plan_markdown": "Export reading plan to Markdown format with tables and checkboxes",
                "tooltip_copy_plan": "Copy formatted reading plan to clipboard for pasting into any application",
                "export_calendar_visualization": "Export Calendar Visualization",
                "tooltip_calendar_visualization": "Generate a visual calendar showing study days with statistical notes below the calendar",
                "calendar_exported": "Calendar exported successfully",
                "total_study_days": "Total Study Days",
                "study_day": "Study Day",
                "rest_day": "Rest Day",
                "matplotlib_not_available": "matplotlib library is not available. Cannot generate charts.\n\nTo install, use:\npip install matplotlib",
                "openpyxl_not_available": "openpyxl library is not available. Cannot export Excel files.\n\nTo install, use:\npip install openpyxl",
                "pdf_no_bookmarks_error": "This PDF file does not contain bookmarks.\nCannot analyze weight distribution without bookmarks.",
                "export_failed": "Export failed",
                "check_file_permissions": "Check file permissions and ensure it's not open in another program.",
                "select_export_folder": "Select Export Folder",

                # File dialogs and messages
                "select_pdf_file": "Select PDF File",
                "select_bookmark_file": "Select Bookmark File",
                "save_extracted_text": "Save Extracted Text",
                "save_compressed_pdf": "Save Compressed PDF",
                "save_merged_pdf": "Save Merged PDF",
                "save_rotated_pdf": "Save Rotated PDF",
                "save_pdf_with_margins": "Save PDF with Margins",
                "save_pdf_with_watermark": "Save PDF with Watermark",
                "save_extracted_pages": "Save Extracted Pages",
                "save_pdf_with_blank_pages": "Save PDF with Blank Pages",

                # Progress and status messages
                "loading": "Loading...",
                "extracting": "Extracting...",
                "compressing": "Compressing...",
                "merging": "Merging...",
                "rotating": "Rotating...",
                "adding_watermark": "Adding watermark...",
                "inserting_pages": "Inserting pages...",
                "deleting_pages": "Deleting pages...",
                "adding_margins": "Adding margins...",

                # Error messages
                "file_not_selected": "File not selected",
                "files_not_selected": "Files not selected",
                "invalid_page_range": "Invalid page range",
                "operation_cancelled": "Operation cancelled",
                "file_access_error": "File access error",
                "insufficient_permissions": "Insufficient permissions",
                "disk_space_error": "Insufficient disk space",
                "corrupted_pdf": "Corrupted PDF file",

                # Success messages
                "files_loaded_successfully": "Files loaded successfully",
                "bookmarks_inserted_successfully": "Bookmarks inserted successfully",
                "bookmarks_loaded_successfully": "Bookmarks loaded successfully!",
                "found_bookmarks": "Found {count} bookmarks.",
                "preview_and_insert_instruction": "You can now preview and insert bookmarks in Step 4 below.",
                "no_selection": "No Selection",
                "select_bookmark_to_fix": "Please select a bookmark to fix.",
                "text_extracted_successfully": "Text extracted successfully",
                "images_extracted_successfully": "Images extracted successfully",
                "pdf_compressed_successfully": "PDF compressed successfully",
                "pdfs_merged_successfully": "PDFs merged successfully",
                "pages_rotated_successfully": "Pages rotated successfully",
                "watermark_added_successfully": "Watermark added successfully",
                "margins_added_successfully": "Margins added successfully",
                "pages_extracted_successfully": "Pages extracted successfully",
                "pages_deleted_successfully": "Pages deleted successfully",
                "blank_pages_inserted_successfully": "Blank pages inserted successfully",

                # File types and formats
                "pdf_files": "PDF Files",
                "text_files": "Text Files",
                "json_files": "JSON Files",
                "all_files": "All Files",
                "image_files": "Image Files",

                # Units and measurements
                "pages_count": "Page count",
                "file_size": "File size",
                "megabytes": "MB",
                "kilobytes": "KB",
                "bytes": "bytes",
                "percentage": "percentage",
                "degrees": "degrees",

                # Interface elements
                "browse_button": "Browse...",
                "select_button": "Select",
                "apply_button": "Apply",
                "reset_button": "Reset",
                "clear_button": "Clear",
                "refresh_button": "Refresh",
                "close_button": "Close",
                "help_button": "Help",
                "about_button": "About",

                # Help and instructions
                "help_title": "💡 Help",
                "instructions": "Instructions",
                "examples": "Examples",
                "tips": "Tips",
                "page_format_examples": "<b>Page Format Examples:</b>",
                "single_pages": "• Single pages: 1,3,5",
                "page_ranges": "• Page ranges: 1-5,10-15",
                "mixed_format": "• Mixed format: 1,3,5-10,15",

                # Navigation and Preview
                "next": "Next ▶",
                "prev": "◀ Prev",
                "pdf_preview": "📄 PDF Preview",
                "no_pdf_loaded": "No PDF loaded",
                "pdf_preview_placeholder": "PDF preview will appear here",
                "click_browse_to_load": "Click 'Browse' to load a PDF file",

                # Page Adjustment Section
                "page_adjustment": "📊 Page Adjustment",
                "page_offset": "Page Offset:",
                "offset_help": "(+/- to adjust for extra/missing pages)",
                "smart_offset_detection": "🎯 Smart Offset Detection",
                "offset_info": "💡 Click any bookmark to preview its page, then use auto-detect if misaligned",
                "auto_detect_offset": "🔍 Auto-Detect Offset from Selected Bookmark",

                # Table Headers
                "bookmarks_table": "📋 Bookmarks",
                "title": "Title",
                "original": "Original",
                "adjusted": "Adjusted",
                "level": "Level",
                "preview": "Preview",

                # File Selection Labels
                "file_selection": "📁 File Selection",
                "bookmark_file": "📖 Bookmark File:",
                "select_pdf_placeholder": "Select PDF file...",
                "select_bookmark_placeholder": "Select text file with 'Title - Page' format...",
                "load_files_preview": "🔄 Load Files & Preview",

                # Operation Groups
                "select_operation": "🔧 Select Operation",
                "file_operations": "📁 File Operations",
                "page_operations": "📄 Page Operations",
                "content_operations": "📝 Content Operations",

                # Watermark Positions
                "center": "Center",
                "top_left": "Top Left",
                "top_right": "Top Right",
                "bottom_left": "Bottom Left",
                "bottom_right": "Bottom Right",

                # Quality Options
                "high_quality": "High (2x)",
                "medium_quality": "Medium (1.5x)",
                "standard_quality": "Standard (1x)",
                "both_methods": "Both (Embedded + Rendered)",
                "embedded_only": "Embedded Images Only",
                "rendered_only": "Rendered Pages Only",

                # Rotation Options
                "rotate_90": "90° Clockwise",
                "rotate_180": "180°",
                "rotate_270": "270° Clockwise",

                # Page Range and Numbers
                "page_range": "📄 Page Range:",
                "all_pages": "All Pages",
                "page_numbers": "Page Numbers",
                "page_count": "Page Count",
                "current_page": "Current Page",
                "total_pages": "Total Pages",

                # Placeholders
                "page_range_placeholder": "all or 1,3,5-10",
                "watermark_text_placeholder": "Enter watermark text...",
                "select_pdf_to_extract": "Select PDF to extract pages from...",
                "select_pdf_to_delete": "Select PDF to delete pages from...",
                "select_pdf_for_blank": "Select PDF to insert blank pages...",
                "select_pdf_for_watermark": "Select PDF file to add watermark...",
                "select_pdf_for_compression": "Select PDF file to compress...",
                "select_pdf_for_editing": "Select PDF file for page editing...",
                "extracted_bookmarks_placeholder": "Extracted bookmarks will appear here...",
                "compression_results_placeholder": "Compression results will appear here...",

                # Group Box Titles
                "compression_options": "⚙️ Compression Options",
                "compression_results": "📊 Compression Results",
                "merge_info": "ℹ️ Merge Information",
                "rotation_options": "🔄 Rotation Options",
                "margin_options": "📏 Margin Options",
                "extraction_options": "📤 Extraction Options",
                "output_options": "📁 Output Options",

                # Help and Information
                "how_to_use": "<b>How to use:</b>",
                "note": "<b>Note:</b>",
                "important": "<b>Important:</b>",
                "tip": "<b>Tip:</b>",
                "warning": "<b>Warning:</b>",

                # File Types in Dialogs
                "pdf_files_filter": "PDF Files (*.pdf)",
                "text_files_filter": "Text Files (*.txt)",
                "json_files_filter": "JSON Files (*.json)",
                "all_files_filter": "All Files (*)",
                "image_files_filter": "Image Files (*.png *.jpg *.jpeg)",

                # Compression Features
                "compression_features": "<b>Compression Features:</b><br>• Remove unused objects and references<br>• Optimize image compression<br>• Clean up document structure<br>• Reduce file size while maintaining quality",

                # Merge Instructions
                "merge_instructions": "1. Click 'Add PDFs' to select multiple PDF files<br>2. Files will be merged in the order they appear in the list<br>3. Use 'Remove' to delete selected files from the list<br>4. Click 'Merge PDFs' to combine all files into one",

                # Page Editing Help
                "margin_help": "Adds white space around pages. Useful for PDFs with no margins or tight content.",
                "rotation_help": "Rotate selected pages by the chosen angle.",
                "page_format_help": "<b>Page Format Examples:</b><br>• Single pages: 1,3,5<br>• Page ranges: 1-5,10-15<br>• Mixed: 1,3,5-10,15<br><br><b>Insert Position:</b> Where to insert pages in target PDF<br><b>Delete:</b> Removes specified pages from PDF<br><b>Extract:</b> Creates new PDF with only specified pages",

                # Reading Speed Tab
                "books_placeholder_text": "Books will appear here after you configure them\nfor reading speed measurement and complete sessions.",
                "close_button": "Close",
                "loading_trends": "Loading trends analysis...",
                "book_title_placeholder": "Book Title",
                "filename_placeholder": "filename.pdf",
                "total_sessions_label": "Total Sessions:",
                "average_speed_label": "Average Speed:",
                "best_speed_label": "Best Speed:",
                "total_time_label": "Total Time:",
                "use_manual_value": "Use Manual Value",
                "cancel_button": "Cancel",

                # Additional Interface Elements
                "enabled": "Enabled",
                "disabled": "Disabled",
                "selected": "Selected",
                "unselected": "Unselected",
                "visible": "Visible",
                "hidden": "Hidden",
                "expanded": "Expanded",
                "collapsed": "Collapsed",

                # Menu items
                "file_menu": "File",
                "close": "Close",
                "exit": "Exit",
                "help_menu": "Help",
                "about": "About",
                "update_available": "Update Available",
                "check_for_updates": "Check for Updates",
                "download_update": "Download Update",
                "remind_later": "Remind Me Later",
                "skip_version": "Skip This Version",

                # Progress and loading messages
                "please_wait": "Please wait...",
                "processing_images": "Processing images...",
                "compressing_pdf": "Compressing PDF...",
                "rotating_pages": "Rotating pages...",
                "adding_margins": "Adding margins...",
                "adding_watermark": "Adding watermark...",
                "extracting_pages": "Extracting pages...",
                "deleting_pages": "Deleting pages...",
                "inserting_pages": "Inserting pages...",
                "operation_in_progress": "Operation in progress...",
                "completed": "Completed",

                # Page range help
                "page_range_help": "<b>How to specify pages:</b><br>• <b>All pages:</b> Leave empty or type 'all'<br>• <b>Single page:</b> 5<br>• <b>Multiple pages:</b> 1,3,5,7<br>• <b>Page range:</b> 1-10<br>• <b>Multiple ranges:</b> 1-5,8-12,15<br>• <b>Mixed:</b> 1,3,5-10,15,20-25",
                "page_range_examples": "Examples: 1,3,5 or 1-10 or 1,3,5-10",

                # Multiple file selection
                "select_multiple_files": "Select Multiple Files",
                "files_selected": "files selected",
                "no_files_selected": "No files selected",

                # Directory selection
                "select_output_directory": "Select Output Directory",
                "output_directory": "Output Directory:",
                "images_saved_to": "Images saved to:",

                # Tab operations
                "extract_tab": "Extract",
                "delete_tab": "Delete",
                "insert_tab": "Insert",
                "merge_tab": "Merge",

                # Merge functionality
                "selected_files": "Selected Files",
                "add_files": "Add Files",
                "remove_selected": "Remove Selected",
                "clear_all": "Clear All",
                "drag_to_reorder": "Drag to reorder - files will be merged in the displayed order",
                "merge_options": "Merge Options",
                "select_at_least_two_files": "Please select at least two files to merge",
                "files_merged": "Files Merged",

                # File menu items
                "open_pdf": "Open PDF",
                "recent_files": "Recent Files",
                "user_guide": "User Guide",

                # Page operations messages
                "no_pages_specified": "Please specify pages",
                "confirm_deletion": "Confirm Deletion",
                "confirm_delete_pages": "Are you sure you want to delete",
                "pages": "Pages",
                "position": "Position",
                "count": "Count",
                "save_pdf_with_blank": "Save PDF with Blank Pages",
                "save_pdf_deleted_pages": "Save PDF with Deleted Pages",
                "save_extracted_pages": "Save Extracted Pages",
                "blank_pages_inserted_successfully": "Blank pages inserted successfully",
                "pages_deleted_successfully": "Pages deleted successfully",
                "pages_extracted_successfully": "Pages extracted successfully",

                # Enhanced menu items
                "quick_tools": "Quick Tools",
                "keyboard_shortcuts": "Keyboard Shortcuts",

                # Settings
                "settings_title": "⚙️ Application Settings",
                "appearance": "🎨 Appearance",
                "language": "🌐 Language",
                "theme": "🎨 Theme",
                "light_mode": "☀️ Light Mode",
                "dark_mode": "🌙 Dark Mode",
                "arabic": "العربية",
                "english": "English",
                "weekend_settings": "📅 Weekend Settings",
                "weekend_days_label": "Weekend Days:",
                "monday": "Monday",
                "tuesday": "Tuesday",
                "wednesday": "Wednesday",
                "thursday": "Thursday",
                "friday": "Friday",
                "saturday": "Saturday",
                "sunday": "Sunday",
                "apply_settings": "✅ Apply Settings",
                "restart_required": "Restart Required",
                "restart_message": "Application needs to restart to apply changes. Restart now?",

                # History
                "history_title": "📋 Operation History",
                "timestamp": "⏰ Time",
                "operation": "🔧 Operation",
                "input_files": "📁 Input Files",
                "output_file": "📄 Output File",
                "details": "📝 Details",
                "clear_history": "🗑️ Clear History",
                "export_history": "📤 Export History",
                "no_history": "No operation history yet",

                # PDF Reading Progress
                "reading_progress": "📝 Comments & Activity Tracker",
                "reading_progress_desc": "Track your PDF reading progress and analyze comments and annotations to know what you've read and how much you've progressed in your studies",
                "select_folder": "📂 Select Folder",
                "scan_computer": "💻 Scan Computer",
                "start_scan": "▶️ Start Scan",
                "stop_scan": "⏹️ Stop Scan",
                "refresh_stats": "🔄 Refresh Stats",
                "total_pdfs": "📚 Total PDFs:",
                "annotated_pdfs": "📝 PDFs with Annotations:",
                "total_annotations": "💬 Total Annotations:",
                "avg_intensity": "🔥 Average Intensity:",
                "search_placeholder": "Search by filename or path...",
                "filter_all": "All PDFs",
                "filter_annotated": "With Annotations Only",
                "filter_recent": "Recently Modified (30 days)",
                "filter_high_activity": "High Activity (10+ annotations)",
                "filter_low_activity": "Low Activity (1-5 annotations)",
                "filter_no_annotations": "No Annotations",
                "open_pdf": "📖 Open PDF",
                "export_annotations": "📝 Export Annotations",
                "study_timeline": "📅 Study Timeline",
                "backup_data": "💾 Backup Data",
                "restore_data": "📥 Restore Data",
                "clear_all_data": "🗑️ Clear All Data",
                "export_list": "📤 Export List",

                # Messages
                "file_not_found": "File not found",
                "operation_completed": "Operation completed successfully",
                "operation_failed": "Operation failed",
                "select_file": "Please select a file",
                "invalid_format": "Invalid format",
                "processing": "Processing...",

                # Recent Books section
                "recent_books": "Recent Books",
                "recent_books_desc": "Manage your personal book library and track reading progress",
                "add_book": "Add Book",
                "select_pdf_book": "Select PDF File",
                "book_progress": "Reading Progress",
                "pages_read": "Pages Read",
                "total_pages": "Total Pages",
                "last_opened": "Last Opened",
                "update_progress": "Update Progress",
                "remove_book": "Remove Book",
                "book_details": "Book Details",
                "reading_percentage": "Reading Percentage",
                "file_size": "File Size",
                "date_added": "Date Added",
                "no_books": "No books added yet",
                "book_added": "Book added successfully",
                "book_removed": "Book removed",
                "progress_updated": "Progress updated",
                "opened_at_last_page": "Opened: {book} at page {page}",

                # Enhanced Reading Statistics
                "reading_statistics_dashboard": "Reading Statistics Dashboard",
                "total_books_read": "Total Books Read",
                "total_pages_read": "Total Pages Read",
                "reading_streak_days": "Reading Streak (Days)",
                "time_range": "Time Range",
                "7_days": "7 Days",
                "30_days": "30 Days",
                "90_days": "90 Days",
                "1_year": "1 Year",
                "export_charts": "Export Charts",
                "charts": "Charts",
                "reading_progress_chart": "Reading Progress Over Time",
                "pages_read": "Pages Read",
                "category_distribution_chart": "Reading by Category",
                "weekly_reading_reports": "Weekly Reading Reports",
                "generate_new_report": "Generate New Report",
                "weekly_reports": "Weekly Reports",
                "detailed_analytics": "Detailed Analytics",
                "date": "Date",
                "book": "Book",
                "pages": "Pages",
                "category": "Category",
                "duration": "Duration",
                "refresh": "Refresh",
                "export_all_data": "Export All Data",
                "books_completed": "Books Completed",
                "total_pages": "Total Pages",
                "reading_time": "Reading Time",
                "categories": "Categories",
                "goals": "Goals",
                "weekly_report": "Weekly Report",
                "save_weekly_report": "Save Weekly Report",
                "reading_statistics": "Reading Statistics",
                "no_data_available": "No data available",
                "close": "Close",
                "feature_unavailable": "Feature Unavailable",
                "statistics_dependencies_missing": "Enhanced statistics require additional dependencies. Please install matplotlib.",

                # Book Information Dialog - English translations
                "book_information": "Book Information",
                "reading_preparation": "Reading Preparation",
                "words_per_page": "Words per Page",
                "estimated_reading_time": "Estimated Reading Time",
                "reading_difficulty": "Reading Difficulty",
                "recommended_session": "Recommended Session",
                "easy": "Easy",
                "medium": "Medium",
                "hard": "Hard",
                "very_hard": "Very Hard",
                "unknown": "Unknown",
                "no_data": "No data",
                "first_added": "First Added",
                "file_information": "File Information",
                "file_path": "File Path",
                "bookmarks": "Bookmarks",
                "seconds": "seconds",
                "minutes": "minutes",
                "hours": "hours",

                # Recent Books additional translations
                "reading_status": "Reading Status",
                "reading": "Reading",
                "to_read": "To Read",
                "completed": "Completed",
                "category": "Category",
                "uncategorized": "Uncategorized",
                "add_books": "Add Books",
                "search_books": "Search books...",
                "all": "All",
                "title": "Title",
                "progress": "Progress",
                "rating": "Rating",
                "recent": "Recent",
                "clear_filters": "Clear Filters",
                "one_book": "1 book",
                "books": "books",
                "small_grid": "Small",
                "medium_grid": "Medium",
                "large_grid": "Large",
                "sort": "Sort",
                "open": "Open",
                "preview": "Preview",
                "rename": "Rename",
                "edit_category": "Edit Category",
                "toggle_star": "Toggle Star",
                "priority": "Priority",
                "normal": "Normal",
                "high": "High",
                "urgent": "Urgent",
                "remove": "Remove",
                "measure_reading_speed": "Measure Reading Speed",
                "add_favorite": "Add to Favorites",
                "remove_favorite": "Remove from Favorites",
                "open_book": "Open Book",

                # Home page grid
                "home_title": "PDF Tools Comprehensive",
                "home_subtitle": "Select a tool to begin working with PDFs",
                "back_to_home": "Back to Home",

                # Security removal feature
                "security_removal_title": "Remove PDF Security",
                "security_removal_desc": "Remove security restrictions and encryption from PDF files",
                "remove_security": "Remove Security",
                "pdf_password": "PDF Password",
                "password_optional": "Password (optional)",
                "security_removed": "Security removed successfully",
                "security_removal_failed": "Failed to remove security",
                "processing": "Processing...",
                "select_file": "Select File",
                "select_pdf_file": "Select PDF File",
                "pdf_file": "PDF File",

                # Reading Speed Measurement
                "reading_speed_meter": "⏱️ Reading Speed Meter",
                "reading_speed_meter_desc": "⏱️ <b>Reading Speed Meter:</b> Measure your reading speed in words per minute with comprehension testing",
                "select_pdf_for_speed": "Select PDF File for Speed Test",
                "or_select_from_recent": "Or select from recent books",
                "recent_books_for_speed": "Recent Books",
                "no_recent_books": "No recent books available",
                "analyzing_pdf": "Analyzing PDF...",

                # OCR PDF to Text
                "ocr_pdf_to_text": "🔍 OCR PDF to Text",
                "ocr_pdf_to_text_desc": "🔍 <b>OCR PDF to Text:</b> Convert scanned PDFs to editable text using advanced Optical Character Recognition technology",
                "select_pdf_file_ocr": "Select PDF File for OCR Conversion",
                "ocr_file_selection": "📁 Select PDF File",
                "ocr_settings": "⚙️ OCR Settings",
                "processing_controls": "🚀 Processing Controls",
                "processing_status": "📊 Processing Status",
                "results_preview": "📄 Results Preview",
                "select_pdf_placeholder_ocr": "Select a PDF file for OCR conversion...",
                "convert_pdf_online": "🌐 Convert PDF Online",
                "convert_pdf_online_tooltip": "Open Colab to convert scanned PDFs to text (TXT/DOCX)",
                "online_ocr_instructions_title": "Convert PDF Online (Colab)",
                "online_ocr_instructions": "Use the online OCR (Colab) to convert scanned PDFs to TXT/DOCX:\n\n1) Open the Colab notebook.\n2) Run the setup cell (installs requirements).\n3) Upload your PDF (left Files panel) or mount Drive.\n4) Set the input file path.\n5) Run the conversion cell.\n6) Download the generated .txt/.docx from the output folder.\n\nAfter download: Use this tab (single or batch) to select the TXT/DOCX or import results automatically.",
                "open_colab": "Open Colab",
                "import_ocr_results_btn": "📥 Import OCR Results",
                "searching_downloads": "Searching Downloads folder for OCR results...",
                "ocr_results_moved": "Found OCR files and moved to preferred folder",
                "ocr_results_not_found": "Matching OCR files were not found. You can select them manually.",
                "choose_ocr_results": "Choose OCR Result Files",
                "select_txt_file": "Select TXT file",
                "select_docx_file": "Select DOCX file",
                "file_name_mismatch": "File name does not match the original PDF name",
                "default_ocr_dir": "Default OCR Results Folder",
                "set_default_ocr_dir": "Set default folder for OCR results",
                "change_default_ocr_dir": "Change default folder for OCR results",
                "preferred_ocr_dir_saved": "Default OCR results folder saved",
                "image_pdf_detected_use_ocr_title": "Image-based PDF Detected",
                "image_pdf_detected_use_ocr_message": "This file contains no extractable text.\n\nUse 'Convert PDF Online' to produce TXT/DOCX, then import for analysis.",
                "no_file_selected": "No file selected",
                "browse": "Browse",
                "remove_newlines_docx": "Remove newlines from DOCX output",
                "remove_newlines_tooltip": "Useful to make DOCX page count match PDF page count",
                "output_directory": "Output Directory",
                "same_as_input": "Same as input file directory",
                "start_ocr_conversion": "🔍 Start OCR Conversion",
                "cancel_processing": "❌ Cancel Processing",
                "processing_progress": "Processing Progress",
                "current_status": "Current Status",
                "ready_to_start": "Ready to start",
                "extracted_text_preview": "Extracted Text Preview",
                "no_text_extracted": "No text extracted yet",
                "output_files": "Output Files",
                "open_txt_file": "📝 Open TXT File",
                "open_docx_file": "📝 Open DOCX File",
                "open_folder": "📁 Open Folder",
                "ocr_processing_required": "OCR Processing Required",
                "ocr_processing_question": "This PDF appears to contain scanned images rather than extractable text.\n\nWould you like to use OCR (Optical Character Recognition) to extract text?\n\nNote: OCR processing may take several minutes depending on the PDF size.",
                "ocr_processing_dialog": "Processing PDF with OCR...\nThis may take several minutes.",
                "ocr_completed": "OCR Completed",
                "ocr_success_message": "OCR processing completed successfully!\n\nExtracted {total_words:,} words from {total_pages} pages.\nAverage: {avg_words:.0f} words per page.\n\nFiles saved:\n• {txt_file}\n• {docx_file}",
                "ocr_warning": "OCR Warning",
                "ocr_no_text_warning": "OCR processing completed but no text was extracted.\nThe PDF may contain images without readable text.",
                "ocr_processing_failed": "OCR Processing Failed",
                "ocr_error_message": "OCR processing failed:\n\n{error_message}\n\nYou can try manual word count input instead.",
                "manual_word_count_input": "Manual Word Count Input",
                "manual_word_count_prompt": "Since automatic text extraction failed, please enter\nthe estimated average words per page manually:\n\n(You can count words on a typical page and enter that number)",
                "manual_input_accepted": "Manual Input Accepted",
                "using_manual_words": "Using {words_per_page} words per page for reading speed calculation.",
                "initializing_ocr": "Initializing OCR processing...",
                "preparing_pdf_ocr": "Preparing PDF for OCR...",
                "running_ocr_conversion": "Running OCR conversion...",
                "ocr_conversion_completed": "OCR conversion completed successfully!",
                "poppler_not_found": "Poppler-utils not found. Please install poppler-utils:\n\nWindows Options:\n1. Download from: https://github.com/oschwartz10612/poppler-windows/releases/\n   Extract and add to PATH\n\n2. Install via conda: conda install poppler\n\n3. Install via pip: pip install poppler-utils\n\nAfter installation, restart the application.",
                "tahweel_not_found": "Tahweel package not found. Please install it first.",
                "google_drive_auth_title": "Google Drive Authentication Setup",
                "google_drive_auth_message": "Tahweel requires Google Drive API access for OCR processing.\n\nYou have two options:\n\n1. Use Service Account (Recommended for automation):\n   - Create a Google Cloud Project\n   - Enable Google Drive API\n   - Create service account credentials\n   - Download JSON file and save as 'service_account.json'\n\n2. Use Interactive Authentication:\n   - Run tahweel command manually first\n   - Complete OAuth flow in browser\n   - Credentials will be saved for future use\n\nWould you like to continue without credentials (may prompt for authentication)?",
                "word_analysis_complete": "Word analysis complete",
                "avg_words_per_page": "Average words per page:",
                "total_pages": "Total pages:",
                "start_reading_session": "Start Reading Session",
                "reading_timer": "Reading Timer:",
                "pages_read": "Pages read:",
                "finish_reading": "Finish Reading & Calculate Speed",
                "reading_speed_results": "Reading Speed Results",
                "your_reading_speed": "Your reading speed:",
                "wpm": "WPM",
                "total_time_taken": "Total time taken:",
                "minutes": "minutes",
                "pdf_options": "PDF Options",
                "pdf_reading_tips": "💡 Tips for accurate reading speed measurement:\n\n• Read in your natural, usual way\n• Avoid reading aloud or moving your lips\n• Focus on understanding content, not just speed\n• Choose texts appropriate for your level\n• Ensure good lighting and eye comfort",

                # Bookmark Copy Feature
                "copy_bookmarks_title": "Copy Bookmarks Between PDFs",
                "copy_bookmarks_desc": "Copy table of contents from original PDF to OCR-enhanced version with same page numbers",
                "original_pdf": "Original PDF:",
                "enhanced_pdf": "Enhanced PDF:",
                "select_original_pdf": "Select original PDF file (source of bookmarks)",
                "select_enhanced_pdf": "Select enhanced PDF file (target for bookmarks)",
                "copy_bookmarks_button": "Copy Bookmarks",
                "select_both_files": "Please select both files first",
                "no_bookmarks_found": "No bookmarks found in the original file",
                "page_count_mismatch": "Page Count Mismatch",
                "page_count_warning": "Enhanced file ({enhanced} pages) has fewer pages than original ({original} pages). Continue anyway?",
                "bookmarks_copied_success": "Successfully copied {count} bookmarks!\n\nSaved as: {filename}",
                "open_directory": "Open Directory",
                "open_directory_question": "Would you like to open the output directory?",
                "bookmark_copy_failed": "Failed to copy bookmarks:\n{error}",
                "copy_bookmarks_utility": "Copy Bookmarks",
                "copy_bookmarks_tooltip": "Copy table of contents from one PDF to another",
                "file_selection": "File Selection",
                "seconds": "seconds",
                "recommendation": "Recommendation:",
                "comprehension_test": "Comprehension Test",
                "answer_questions": "Answer the following questions to test your comprehension:",
                "question": "Question",
                "your_answer": "Your answer:",
                "submit_answers": "Submit Answers",
                "comprehension_score": "Comprehension Score:",
                "correct_answers": "correct answers",
                "out_of": "out of",
                "excellent_comprehension": "Excellent comprehension! You read at a good speed with high understanding.",
                "good_comprehension": "Good comprehension. Try to improve the balance between speed and understanding.",
                "needs_improvement": "Needs improvement. Focus more on comprehension while reading.",
                "slow_but_thorough": "Slow but thorough reading. This is good for complex material.",
                "average_speed": "Average speed with good comprehension. Keep practicing.",
                "fast_reader": "Fast reader! Make sure you're not sacrificing comprehension.",
                "speed_reader": "Very fast reading. Suitable for quick review.",
                "reset_test": "Reset Test",
                "new_reading_session": "New Reading Session",
                "measure_reading_speed": "Measure Reading Speed",
                "prepare_book": "Prepare Book",
                "current_book": "Current Book",
                "activity": "Activity",
                "word_analysis_complete": "Word Analysis Complete",
                "comprehension_test": "Comprehension Test",
                "reading_sessions": "Reading Sessions",
                "prepare_book_desc": "Pre-analyze book to save time in future reading sessions",
                "book_prepared": "Book prepared successfully",
                "analysis_cached": "Word analysis cached for future use",

                # Additional Reading Speed Localization
                "reading_speed_guide": "Reading Speed Guide",
                "word_analysis_complete": "Word Analysis Complete",
                "start_reading_session": "Start Reading Session",
                "book_ready": "Book Ready",
                "information": "Information",
                "select_file_first": "Please select a file first",
                "ocr_results_not_found": "OCR results not found",
                "select_txt_file": "Select TXT File",
                "select_docx_file": "Select DOCX File",
                "file_name_mismatch": "File names do not match",
                "ocr_results_moved": "OCR results moved successfully",
                "convert_pdf_online": "Convert PDF Online",
                "convert_pdf_online_tooltip": "Open Colab to convert scanned PDFs to text",
                "import_ocr_results_btn": "Import OCR Results",
                "choose_ocr_results": "Choose OCR Results",
                "set_default_ocr_dir": "Set Default OCR Directory",
                "change_default_ocr_dir": "Change Default OCR Directory",
                "preferred_ocr_dir_saved": "Preferred OCR directory saved",
                "ocr_processing_required": "OCR Processing Required",
                "ocr_processing_question": "This file requires OCR processing. Do you want to continue?",
                "online_ocr_instructions": "Online OCR Instructions",
                "online_ocr_instructions_title": "Online OCR Instructions",
                "open_colab": "Open Colab",
                "file_not_found": "File not found",

                # Reading Sessions Table
                "reading_sessions_desc": "Reading speed measurement sessions with detailed statistics:",
                "time_spent": "Time Spent",
                "efficiency": "Efficiency",
                "total_sessions": "Total Sessions",
                "average_wpm": "Average WPM",
                "best_wpm": "Best WPM",
                "no_sessions_yet": "No reading sessions yet",
                "export_sessions": "Export Sessions",
                "data_exported": "Data exported successfully to",

            # Enhanced Reading Speed Workflow
            "analyzing_pdf": "Analyzing PDF...",
            "analysis_results": "Analysis Results",
            "sample_text_preview": "Sample Text Preview:",
            "adjust_words_per_page": "Adjust words per page:",
            "words": "words",
            "confirm_preparation": "Confirm & Prepare",
            "preparation_options": "Preparation Options",
            "choose_preparation_method": "Choose Preparation Method",
            "pdf_analysis_limited": "This PDF has limited extractable text. Please choose a preparation method:",
            "use_default_estimate": "Use Default Estimate",
            "default_estimate_desc": "Use a typical range of 120-150 words per page for standard documents.",
            "use_ocr_conversion": "Use OCR Conversion",
            "ocr_conversion_desc": "Convert PDF to Word format using online OCR for accurate word count.",
            "manual_entry_desc": "Count words on a sample page and enter manually.",
            "use_default": "Use Default",
            "use_ocr": "Use OCR",
            "ocr_guidance": "OCR Conversion Guide",
            "ocr_conversion_guide": "OCR Conversion Guide",
            "ocr_instructions": """
Follow these steps to convert your PDF using online OCR:

1. Click the "Open OCR Tool" button below to open the online converter
2. Upload your PDF file to the OCR service
3. Wait for the conversion to complete
4. Download the converted Word document
5. Open the Word document and count words on a sample page
6. Return here and enter the words per page manually

This method provides the most accurate word count for scanned PDFs.
            """,
            "open_ocr_tool": "Open OCR Tool",
            "enter_ocr_results": "Enter OCR Results",
            "ocr_results_desc": "After OCR conversion, count words on a sample page and enter below:",
            "confirm_ocr_results": "Confirm OCR Results",
            "automatic_analysis": "Automatic Analysis",
            "default_estimate": "Default Estimate",
            "ocr_conversion": "OCR Conversion",
            "saved_configuration": "Saved Configuration",
            "preparation_method": "Preparation Method",
            "open_pdf_manually": "Please open the PDF manually:",
            "reading_speed_results": "Reading Speed Results",
            "your_reading_speed": "Your Reading Speed",
            "total_time_taken": "Total Time Taken",
            "total_words_read": "Total Words Read",
            "new_reading_session": "New Reading Session",
            "pages_read": "Pages read:",
            "finish_reading": "Finish Reading",
            "hours": "hours",
            "minutes": "minutes",
            "seconds": "seconds",
            "total_books": "Total Books",
            "total_sessions": "Total Sessions",
            "total_pages_read": "Total Pages Read",
            "average_speed": "Average Speed",
            "best_speed": "Best Speed",
            "total_pages": "Total Pages",
            "start_page": "Start Page",
            "resume_reading": "Resume Reading",
            "pause_reading": "Pause Reading",
            "stop_reading": "Stop Reading",

            # Training mode and Pages/Minute tab
            "training_mode": "Training Mode",
            "elapsed_time": "Elapsed time",
            "start": "Start",
            "end_page": "End Page",
            "what_page_reached": "What page did you reach?",
            "ppm_simple_desc": "Enter start page, start the timer, then enter end page when you stop.",
            "mode_simple": "Simple",
            "sample_pages": "Sample pages:",
            "extract_sample": "Extract Sample",
            "ai_prompt": "AI Prompt",
            "copy_prompt": "Copy Prompt",
            "avg_words_per_page": "Average words/page:",
            "mode_sample": "Sample-Based",
            "using_stored_counts": "Using stored per-page counts",
            "mode_accurate": "Most Accurate",
            "pages_per_minute": "Pages/Minute",
            "parse_and_save": "Parse & Save",
            "info": "Info",
            "missing_data": "Missing CSV or no document loaded.",
            "saved": "Saved",
            "saved_page_counts": "Per-page word counts saved.",
            "paste_csv_here": "Paste CSV data here:",
            "csv_placeholder": "page_number,word_count\n1,250\n2,240\n...",
            "ppm_label": "PPM",
            "wpm_label": "WPM",

            "save_failed": "Could not save page counts.",
            "parse_failed": "Could not parse any rows.",
            "done": "Done",
            "sample_saved_to": "Sample saved to:",
            "open_pdf_first": "Please open a PDF first.",


            # Reading Speed Levels and Advice
            "beginner_reader": "Beginner Reader",
            "beginner_advice": "Practice reading regularly to improve your speed. Focus on reducing subvocalization.",
            "average_reader": "Average Reader",
            "average_advice": "Good reading speed! Try speed reading techniques to reach the next level.",
            "good_reader": "Good Reader",
            "good_advice": "Excellent reading speed! You're above average. Keep practicing to maintain this level.",
            "expert_reader": "Expert Reader",
            "expert_advice": "Outstanding reading speed! You're in the top tier of readers.",

            # PDF Loading Messages
            "existing_configuration": "Existing Configuration",
            "use_existing_config": "This book has been prepared before. Use existing configuration?",
            "pdf_loaded": "PDF Loaded",
            "pdf_loaded_successfully": "PDF loaded successfully!",
            "pages": "pages",
            "click_prepare_to_start": "Click \"Prepare Book\" to start analysis.",

            # Fast Reading Trainer
            "fast_reading_trainer": "Fast Reading Trainer",
            "fast_reading_desc": "Train your reading speed by displaying words in groups at a controlled pace. Improve your Words Per Glance (WPG) ability.",
            "select_document": "Select Document",
            "select_text_file": "Select text file (.txt, .docx, .pdf)",
            "select_txt": "Select TXT File",
            "select_docx": "Select DOCX File",
            "select_docx_file": "Select DOCX File",
            "select_pdf_trainer": "Select PDF for Training",
            "training_settings": "Training Settings",
            "words_per_glance": "Words per glance:",
            "target_speed": "Target speed:",
            "font_size": "Font size:",
            "training_controls": "Training Controls",
            "start_training": "Start Training",
            "pause": "Pause",
            "resume": "Resume",
            "stop": "Stop",
            "progress": "Progress",
            "ready_to_start": "Ready to start training",
            "select_document_to_start": "Select a document to start training",
            "warning": "Warning",
            "no_text_found": "No text found in the document",
            "no_extractable_text": "No extractable text found in the PDF",
            "text_too_short": "Text is too short for training (minimum 10 words required)",
            "ready_to_train": "Ready to start training!",
            "words_loaded": "words loaded",
            "text_loaded": "Text Loaded",
            "text_loaded_successfully": "Text loaded successfully!",
            "adjust_settings_and_start": "Adjust your settings and click Start Training!",
            "no_text_loaded": "No text loaded for training",
            "training_active": "Training active",
            "training_resumed": "Training resumed",
            "training_paused": "Training paused",
            "training_stopped": "Training stopped",
            "training_complete": "Training Complete!",
            "congratulations": "Congratulations! You have completed the training session.",
            "statistics": "Statistics",
            "words_read": "words read",
            "words": "words",
            "minutes": "minutes",
            "font_family": "Font family",
            "effective_reading_speed": "Effective Reading Speed",
            "every": "every",
            "seconds": "seconds",
            "word_groups_per_minute": "word groups per minute",
            "adjust_speed_if_needed": "Adjust speed if needed",
            "continue_previous_session": "Continue Previous Session",
            "save_session": "Save Session",
            "save_training_progress": "Save your training progress?",
            "current_progress": "Current progress",
            "session_resumed": "Session resumed",
            "session_resumed_successfully": "Session resumed successfully",
            "progress": "Progress",
            "error_resuming_session": "Error resuming session",
            "no_saved_sessions": "No saved training sessions found",
            "saved_training_sessions": "Saved Training Sessions",
            "select_session_to_resume": "Select a session to resume:",
            "complete": "complete",
            "resume_session": "Resume Session",
            "delete_session": "Delete Session",
            "select_session_first": "Please select a session first",
            "confirm_delete": "Confirm Delete",
            "delete_session_confirm": "Delete this training session?",
            "file_not_found": "File not found",
            "unsupported_file_type": "Unsupported file type",
            "session_saved_successfully": "Training session saved successfully!",
            "all_activities": "All Activities",
            "reading_measurements": "Reading Measurements",
            "training_sessions": "Training Sessions",
            "reading_measurement": "Reading Measurement",
            "training_session": "Training Session",
            "book_title": "Book Title",
            "training": "Training",
            "resume": "Resume",
            "session_resumed": "Session resumed!",
            "last_page": "Last page",
            "file_not_found": "File not found. It may have been moved or deleted.",
            "type": "Type",
            "file": "File",
            "progress_time": "Progress/Time",
            "settings_details": "Settings/Details",
            "actions": "Actions",
            "no_activities_yet": "No activities yet",
            "show": "Show",
            "completed": "Completed",
            "complete": "Complete",
                "current_book": "Current Book",
                "reading_sessions": "Reading Sessions",
                "output_settings": "Output Settings",
                "select_output_location": "Select output location",
                "output_file": "Output File",

                # Section descriptions
                "pdf_viewer_desc": "View and annotate PDF files with advanced tools",
                "bookmark_manager_desc": "Manage and organize PDF bookmarks",
                "bookmark_extractor_desc": "Extract bookmarks from PDF files",
                "reading_progress_desc": "Track reading progress and annotations",

                # Statistics Dashboard
                "statistics_dashboard": "Statistics Dashboard",
                "statistics_dashboard_desc": "📈 Comprehensive reading analytics and statistics with visual charts",
                "key_metrics": "📊 Key Metrics",
                "overview_statistics": "📊 Overview Statistics",
                "visual_analytics": "📈 Visual Analytics",
                "advanced_analytics": "🔬 Advanced Analytics",
                "total_books": "Total Books",
                "read_today": "Read Today",
                "pages_today": "Pages Today",
                "reading_streak": "Reading Streak",
                "avg_pages_day": "Avg Pages/Day",
                "total_time": "Total Time",
                "completion_rate": "Completion Rate",
                "favorite_category": "Top Category",
                "categories": "Categories",
                "reading_velocity": "Reading Velocity",
                "monthly_goal": "Monthly Goal",
                "productive_time": "Peak Time",
                "consistency_score": "Consistency Score",
                "books_by_category": "Books by Category",
                "daily_reading_progress": "Daily Progress",
                "export": "📤 Export",
                "refresh": "🔄 Refresh",

                "page_operations_desc": "Extract/delete/insert pages, merge files, rotate, crop, and add margins",
                "watermark_desc": "Add watermarks to PDF files",
                "extract_images_desc": "Extract images from PDF files",
                "extract_text_desc": "Extract text content from PDFs",
                "merge_pdfs_desc": "Merge multiple PDF files into one",
                "split_pdfs_desc": "Split a PDF file into multiple files",
                "compress_desc": "Reduce PDF file size",
                "page_editing_desc": "Rotate, crop, and edit pages",
                "settings_desc": "Configure application settings",
                "history_desc": "View operation history",
                "navigation": "Navigation",
                "refresh": "Refresh",
                "clear_history": "Clear History",

                # PDF Comments section
                "show_hide_help": "Show/Hide Help",
                "ready": "Ready",
                "no_directory_selected": "No directory selected",
                "include_subdirectories": "Include subdirectories",
                "statistics": "📈 Statistics",
                "search_placeholder": "Search by filename or path...",
                "filter_label": "📂 Filter:",
                "sort_label": "Sort by:",
                "all_pdfs": "All PDFs",
                "with_annotations_only": "With Annotations Only",
                "recently_modified": "Recently Modified (30 days)",
                "high_activity": "High Activity (10+ annotations)",
                "low_activity": "Low Activity (1-5 annotations)",
                "no_annotations": "No Annotations",
                "last_scanned": "Last Scanned",
                "file_name": "File Name",
                "last_modified": "Last Modified",
                "annotations_count": "Annotations Count",
                "reading_intensity": "Reading Intensity",
                "file_name_col": "File Name",
                "path_col": "Path",
                "pages_col": "Pages",
                "annotations_col": "Annotations",
                "intensity_col": "Intensity",
                "last_modified_col": "Last Modified",
                "last_scanned_col": "Last Scanned",
                "open_pdf": "📖 Open PDF",
                "export_annotations": "📝 Export Annotations",
                "study_timeline": "📅 Study Timeline",
                "backup_data": "💾 Backup Data",
                "restore_data": "📥 Restore Data",
                "clear_all_data": "🗑️ Clear All Data",

                # Recent Books additional keys
                "export": "Export",
                "grid_size": "Grid Size",
                "small": "Small",
                "medium": "Medium",
                "large": "Large",
                "total": "Total",
                "starred": "Starred",
                "priority": "Priority",
                "avg_progress": "Avg Progress",
                "no_cover": "No Cover",
                "available": "Available",

                # Context menu items
                "edit_name": "Edit Name",
                "edit_category": "Edit Category",
                "reading_status": "Reading Status",
                "cover_management": "Cover Management",
                "add_star": "Add Star",
                "remove_star": "Remove Star",
                "set_priority": "Set Priority",
                "quick_update": "Quick Update",
                "update_progress": "Update Progress",
                "open_external": "Open with External App (Default)",
                "open_internal": "Open with Internal Viewer",
                "reading_analytics": "Reading Analytics",
                "remove_book": "Remove Book",
                "upload_custom_cover": "Upload Custom Cover",
                "remove_custom_cover": "Remove Custom Cover",
                "currently_reading": "Currently Reading",
                "normal": "Normal",
                "high": "High",
                "urgent": "Urgent",

                # Filter options
                "all_books": "All Books",
                "reading": "Reading",
                "to_read": "To Read",
                "completed": "Completed",
                "priority_books": "Priority Books",
                "sort_by": "Sort by",

                # Sort options
                "last_opened": "Last Opened",
                "title": "Title",
                "progress": "Progress",
                "date_added": "Date Added",

                # Category options
                "category": "Category",
                "all_categories": "All Categories",
                "uncategorized": "Uncategorized",

                # TOC Preparation
                "prepare_toc": "Prepare Table of Contents",
                "toc_preparation_title": "Prepare Table of Contents from PDF",
                "step1_extract_pages": "Step 1: Extract TOC Pages",
                "toc_page_range": "TOC page range:",
                "to": "to",
                "extract_toc_pages": "Extract TOC Pages",
                "step2_ai_extraction": "Step 2: AI Text Extraction",
                "ai_studio_instructions": "1. Click 'Open AI Studio' below\n2. Click 'Copy AI Prompt' to copy instructions\n3. Paste instructions in AI Studio\n4. Upload the extracted PDF file\n5. Copy the result and paste it in Step 3",
                "open_ai_studio": "Open AI Studio",
                "copy_ai_prompt": "Copy AI Prompt",
                "step3_format_toc": "Step 3: Format TOC",
                "paste_ai_result": "Paste AI Studio result here:",
                "toc_paste_placeholder": "Paste the extracted text from AI Studio here...\n\nExample:\nIntroduction - 5\nReferences - 8\nChapter 1: Getting Started - 15\nSection 1.1: Overview - 15",
                "format_toc": "Format TOC",
                "formatted_result": "Formatted Result",
                "use_bookmarks": "Use Bookmarks",
                "ai_toc_prompt": "Extract table of contents from PDF file with numbered format for Bookmarks\n\nExtract table of contents from the attached PDF pages.\n\nInstructions:\n\n1. Extract titles and page numbers.\n2. Use two levels of headings:\n   * Level 1: Main headings (such as chapters or parts).\n   * Level 2: Sub-headings within each section.\n3. Level 1 formatting:\n   * Write the main heading followed by page number.\n   * Format: Title - Page Number\n4. Level 2 formatting:\n   * Add indentation.\n   * Use numbering in X.Y format where X is the main section number and Y is the sequential sub-heading number within it.\n   * Main section numbering starts from 1.\n   * Final format:   X.Y Title - Page Number\n5. Separate each main section and its sub-headings with an empty line as shown in the example.\n6. Use the same language as the text in the PDF.\n7. Do not add any other text except the formatted table of contents.\n\nExample output:\n\nIntroduction - 5\n1.1 Project Concepts and Objectives - 8\n1.2 Work Methodology and Implementation Steps - 12\n\nInternal System Chapter - 15\n2.1 Definitions and Terminology Statement - 17\n2.2 Daily Work Organization - 21\n2.3 Resource Management and Reports - 24\n2.4 Document Templates and Procedures - 28\n2.5 Evaluation and Continuous Improvement - 31",

                # TOC Preparation Steps
                "toc_preparation_steps": "📝 TOC Preparation Steps:",
                "toc_step1": "Step 1: Select your PDF file above",
                "toc_step2": "Step 2: Set the page range for your Table of Contents",
                "toc_step3": "Step 3: Extract TOC pages to a separate PDF",
                "toc_step4": "Step 4: Use AI Studio to extract text from the TOC PDF",
                "toc_step4_final": "Step 4: Paste the extracted text and format it into bookmarks",

                # TOC Range Dialog
                "select_toc_range": "Select TOC Range",
                "select_toc_range_title": "📋 Choose Table of Contents Pages",
                "toc_range_selection": "Range Selection",
                "toc_range_instructions": "Use the preview to navigate through the PDF and identify the Table of Contents pages. You can use the quick buttons to set the current page as start or end.",
                "start_page": "Start Page:",
                "end_page": "End Page:",
                "set_as_start": "Set as Start",
                "set_as_end": "Set as End",
                "selected_range": "Selected Range",
                "total_pages_selected": "Total Pages Selected",
                "extract_selected_pages": "Extract Selected Pages",
                "toc_range_hint": "💡 Use the buttons above to extract TOC pages after selecting the appropriate range",

                # TOC Page Selection Dialog
                "select_toc_pages": "📋 Select TOC Pages",
                "set_toc_range": "Set TOC Range",
                "toc_page_selection": "TOC Page Selection",
                "toc_page_selection_desc": "Use this dialog to browse the PDF and visually identify the pages containing the Table of Contents",
                "navigate_and_select": "Navigate and Select",
                "toc_range_selected": "TOC Range Selected",

                # Improved Workflow Steps
                "step1_load_pdf": "Step 1: Load PDF file that you need to bookmark",
                "step2_select_toc": "Step 2: Select TOC pages",
                "step3_extract_toc": "Step 3: Extract TOC pages",
                "step4_format_process": "Step 4: Use AI Studio and Load Bookmarks",
                "step4_verify_insert": "Step 4: Verify bookmarks and insert",

                "load_pdf_primary": "📄 Load PDF File",
                "select_toc_pages_step": "📋 Select TOC Pages",
                "format_and_load": "🔄 Format TOC and Load Files",
                "already_have_bookmarks": "If you already have bookmarks in a text file",
                "upload_to_ai_studio": "Upload extracted pages to AI Studio and use the copied prompt",
                "paste_ai_result_here": "Paste AI Studio result here and it will be formatted automatically",
                "verify_bookmarks": "Verify bookmarks using the 'Preview' button",
                "fix_page_mismatch": "Fix page mismatch",
                "navigate_to_correct_page": "Navigate to correct page and press 'Fix'",
                "enable_mismatch_detection": "Enable mismatch detection",
                "fix": "Fix",
                "auto_calculate": "Auto Calculate",
                "confirm": "Confirm",
                "actions": "Actions",
                "delete": "Delete",
                "delete_bookmark": "Delete Bookmark",
                "confirm_delete": "Are you sure you want to delete this bookmark?",
                "split_by_bookmarks": "Split PDF by Bookmarks",
                "split_by_bookmarks_desc": "Split PDF into separate files based on bookmark structure",
                "split_level_1_only": "Split by Level 1 only (Main chapters)",
                "split_all_levels": "Split by all levels",
                "output_directory": "Output Directory",
                "file_naming": "File Naming",
                "use_bookmark_titles": "Use bookmark titles as filenames",
                "use_sequential_numbers": "Use sequential numbers (001, 002, etc.)",
                "start_splitting": "Start Splitting",
                "bookmark_levels": "Bookmark Levels",
                "additional_options": "Additional Options",
                "include_original_bookmarks": "Include original bookmarks in split files",
                "create_index_file": "Create index file with list of split files",
                "insert_bookmarks_pdf": "Insert Bookmarks into PDF",
                "split_pdf_bookmarks": "Split PDF by Bookmarks",
                "split_options": "Split Options",
                "level1_bookmarks": "Level 1 Bookmarks (Main chapters)",
                "level2_bookmarks": "Level 2 Bookmarks (Sub-chapters)",
                "all_levels": "All Levels",
                "output_directory": "Output Directory",
                "select_output_dir": "Select output directory:",
                "choose_output_dir": "Choose output directory...",
                "split_pdf": "Split PDF",
                "cancel": "Cancel",
                "load_bookmarks_preview": "Load Bookmarks to Preview",

                # New Redesigned Bookmark Workflow - English
                "step2_choose_method": "Step 2: Choose Bookmark Source Method",
                "choose_bookmark_method_desc": "Choose the method you want to use to create bookmarks:",
                "method_toc_pages": "Select TOC Pages - for extracting table of contents from PDF pages",
                "method_text_file": "Choose Text File + Paste Text Option - for using external bookmark text",
                "step2_bookmark_extraction": "Step 2: Paste Bookmarks",
                "step2_guidance_text": "Paste content here:",
                "step2_help_text": "If you don't have bookmarks - ",
                "prepare_bookmarks_button": "Prepare Bookmarks",
                "step3_generate_bookmarks": "Step 3: Generate Bookmarks",
                "text_file_desc": "Choose a text file containing bookmarks in 'Title - Page' format:",
                "or_separator": "OR",
                "paste_text_desc": "Paste text containing bookmarks directly:",
                "format_bookmarks": "Format Bookmarks",
                "no_toc_text": "Please paste text or select a text file first",
                "no_bookmark_source": "Please select a text file or paste text first",
                "step4_verify_insert": "Step 4: Verify Bookmarks and Insert",
                "verify_bookmarks": "Verify bookmarks using the preview, then click 'Insert Bookmarks' to save them to the PDF.",

                # Success messages
                "bookmark_source_ready": "Bookmark source is ready for formatting!",
                "bookmark_source_ready_desc": "You can now press 'Load Bookmarks' to continue.",
                "toc_text_ready": "TOC text is ready for formatting!",

                # Message box titles and content
                "success": "Success",
                "pdf_loaded_success": "PDF loaded successfully!",
                "total_pages": "Total pages:",
                "use_preview_instruction": "Use the preview to identify TOC pages, then click 'Extract TOC Pages' to select the range.",
                "toc_range_selected": "TOC Range Selected",
                "selected_range": "Selected range:",
                "pages": "pages",
                "can_extract_now": "You can now extract TOC for this range.",
                "toc_extracted_success": "TOC pages extracted successfully!",
                "file": "File:",
                "folder_will_open": "The containing folder will open automatically.",
                "prompt_copied": "Prompt Copied",
                "ai_prompt_copied": "AI prompt has been copied to clipboard!",
                "paste_in_ai_studio": "Paste it in AI Studio and upload your extracted TOC pages.",

                # Recent Books - English translations
                "add_books": "Add Books",
                "small_grid": "Small",
                "medium_grid": "Medium",
                "large_grid": "Large",
                "search_books": "Search books...",
                "status": "Status:",
                "category": "Category:",
                "sort": "Sort:",
                "all": "All",
                "reading": "Reading",
                "to_read": "To Read",
                "completed": "Completed",
                "recent": "Recent",
                "title": "Title",
                "progress": "Progress",
                "rating": "Rating",
                "clear_filters": "Clear Filters",
                "no_books": "No books",
                "one_book": "1 book",
                "books": "books",
                "uncategorized": "Uncategorized",

                # Context Menu - English translations
                "open": "Open",
                "preview": "Preview",
                "rename": "Rename",
                "edit_category": "Edit Category",
                "update_progress": "Update Progress",
                "reading_status": "Reading Status",
                "toggle_star": "Toggle Star",
                "priority": "Priority",
                "normal": "Normal",
                "high": "High",
                "urgent": "Urgent",
                "remove": "Remove",

                # PDF Viewer specific
                "bookmarks": "Bookmarks",
                "hide_bookmarks": "Hide Bookmarks",
                "show_bookmarks": "Show Bookmarks",
                "no_bookmarks": "No bookmarks found",
                "bookmark_weight_distribution": "Bookmark Weight Distribution",
                "click_to_go_to_page": "Click to go to page",
                "go_to_page": "Go to Page",
                "enter_page_number": "Enter page number",
                "resumed_reading": "Reading Resumed",
                "resumed_from_page": "Resumed from page",
                "open_in_weight_analyzer": "Open in Chapter Weight Analyzer",
                "pdf_loaded_in_analyzer": "PDF loaded in Weight Analyzer",
                "weight_analyzer_not_available": "Chapter Weight Analyzer is not available",

                # Bug Reporting System
                "help": "Help",
                "report_bug": "🐛💡 Report Bug / Suggestion",
                "bug_report_title": "Bug Report / Feature Suggestion",
                "bug_report_desc": "Help us improve by reporting bugs or suggesting new features. Your feedback is valuable!",
                "report_type": "Report Type",
                "bug_report": "Bug Report",
                "feature_suggestion": "Feature Suggestion",
                "enhancement": "Enhancement",
                "question": "Question",
                "bug_information": "Bug Information",
                "bug_title": "Bug Title",
                "bug_title_placeholder": "Brief description of the issue...",
                "bug_description": "Description",
                "bug_description_placeholder": "Detailed description of the bug...\n\nSteps to reproduce:\n1. \n2. \n3. \n\nExpected behavior:\n\nActual behavior:",
                "username_optional": "Username (Optional)",
                "username_placeholder": "Your name or username...",
                "email_optional": "Email (Optional)",
                "email_placeholder": "your.email@example.com",
                "severity": "Severity",
                "severity_low": "Low",
                "severity_medium": "Medium",
                "severity_high": "High",
                "severity_critical": "Critical",
                "category": "Category",
                "category_ui": "UI Issue",
                "category_functionality": "Functionality Bug",
                "category_performance": "Performance",
                "category_crash": "Crash",
                "category_other": "Other",
                "contact_info": "Contact Information (Optional)",
                "system_info_note": "System information will be automatically included to help diagnose the issue.",
                "submit_bug": "Submit Bug Report",
                "submitting": "Submitting...",
                "bug_submitted_success": "Bug report submitted successfully! Thank you for your feedback.",
                "bug_submission_failed": "Failed to submit bug report. Please try again.",
                "validation_error": "Validation Error",
                "bug_title_required": "Please enter a bug title.",
                "bug_description_required": "Please enter a bug description.",
                "page_name": "Page",
            }
        }

    def get_text(self, key: str) -> str:
        """Get localized text"""
        return self.translations.get(self.current_language, {}).get(key, key)

    def set_language(self, language: str):
        """Set current language"""
        if language in self.translations:
            self.current_language = language


class Settings:
    """Application settings manager"""

    def __init__(self):
        self.settings_file = get_settings_path("pdf_tools_settings.json")
        print(f"⚙️ Settings file: {self.settings_file}")
        self.default_settings = {
            "language": "ar",  # Arabic as default
            "theme": "light",
            "window_geometry": [100, 100, 1400, 900],
            "last_directories": {},
            "history_limit": 100,
            "weekend_days": [5, 6]  # Saturday (5) and Sunday (6) by default
        }
        self.settings = self.load_settings()

    def load_settings(self) -> Dict[str, Any]:
        """Load settings from file"""
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                # Merge with defaults for any missing keys
                for key, value in self.default_settings.items():
                    if key not in settings:
                        settings[key] = value
                return settings
        except Exception as e:
            print(f"Error loading settings: {e}")

        return self.default_settings.copy()

    def save_settings(self):
        """Save settings to file"""
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving settings: {e}")

    def get(self, key: str, default=None):
        """Get setting value"""
        return self.settings.get(key, default)

    def set(self, key: str, value):
        """Set setting value"""
        self.settings[key] = value
        self.save_settings()


class HistoryManager:
    """Manage operation history"""

    def __init__(self, settings: Settings):
        self.settings = settings
        self.history_file = get_settings_path("pdf_tools_history.json")
        print(f"📜 History file: {self.history_file}")
        self.history = self.load_history()

    def load_history(self) -> List[HistoryEntry]:
        """Load history from file"""
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return [HistoryEntry(**entry) for entry in data]
        except Exception as e:
            print(f"Error loading history: {e}")

        return []

    def save_history(self):
        """Save history to file"""
        try:
            # Limit history size
            limit = self.settings.get("history_limit", 100)
            if len(self.history) > limit:
                self.history = self.history[-limit:]

            data = [
                {
                    "timestamp": entry.timestamp,
                    "operation": entry.operation,
                    "input_files": entry.input_files,
                    "output_file": entry.output_file,
                    "status": entry.status,
                    "details": entry.details
                }
                for entry in self.history
            ]

            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving history: {e}")

    def add_entry(self, operation: str, input_files: List[str],
                  output_file: str, status: str, details: str = ""):
        """Add new history entry"""
        try:
            entry = HistoryEntry(
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                operation=operation,
                input_files=input_files if input_files else [],
                output_file=output_file if output_file else "",
                status=status,
                details=details
            )
            self.history.append(entry)
            self.save_history()
            print(f"History entry added: {operation} - {status}")
        except Exception as e:
            print(f"Error adding history entry: {e}")

    def clear_history(self):
        """Clear all history"""
        self.history = []
        self.save_history()

    def export_history(self, file_path: str):
        """Export history to file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("PDF Tools Operation History\n")
                f.write("=" * 50 + "\n\n")

                for entry in self.history:
                    f.write(f"Time: {entry.timestamp}\n")
                    f.write(f"Operation: {entry.operation}\n")
                    f.write(f"Input Files: {', '.join(entry.input_files)}\n")
                    f.write(f"Output File: {entry.output_file}\n")
                    f.write(f"Status: {entry.status}\n")
                    if entry.details:
                        f.write(f"Details: {entry.details}\n")
                    f.write("-" * 30 + "\n\n")

            return True
        except Exception as e:
            print(f"Error exporting history: {e}")
            return False


class PDFOperations:
    """Core PDF operations class with comprehensive functionality"""

    def __init__(self, history_manager=None):
        self.history_manager = history_manager

    def load_bookmarks_from_text(self, file_path: str) -> List[Bookmark]:
        """Load bookmarks from text file with 'Title - Page' format"""
        bookmarks = []

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            for line_num, line in enumerate(lines, 1):
                line = line.strip()
                if not line:
                    continue

                # Parse "Title - Page" format
                if ' - ' in line:
                    parts = line.rsplit(' - ', 1)
                    if len(parts) == 2:
                        title = parts[0].strip()
                        try:
                            page = int(parts[1].strip())

                            # Determine level based on title format
                            level = 1
                            if re.match(r'^\d+\.', title):  # Starts with number (1.1, 2.3, etc.)
                                level = 2
                            elif title.startswith('  ') or title.startswith('\t'):  # Indented
                                level = 2

                            bookmarks.append(Bookmark(title=title, page=page, level=level))

                        except ValueError:
                            print(f"Warning: Invalid page number on line {line_num}: {line}")
                else:
                    print(f"Warning: Invalid format on line {line_num}: {line}")

            print(f"Loaded {len(bookmarks)} bookmarks from {file_path}")
            return bookmarks

        except Exception as e:
            print(f"Error loading bookmarks: {e}")
            return []

    def extract_bookmarks_from_pdf(self, pdf_path: str) -> List[Bookmark]:
        """Extract existing bookmarks from PDF"""
        bookmarks = []
        try:
            doc = fitz.open(pdf_path)
            toc = doc.get_toc()
            doc.close()

            for item in toc:
                level, title, page = item
                bookmarks.append(Bookmark(title=title, page=page, level=level))

            print(f"Extracted {len(bookmarks)} bookmarks from PDF")
            return bookmarks

        except Exception as e:
            print(f"Error extracting bookmarks: {e}")
            return []

    def insert_bookmarks_into_pdf(self, pdf_path: str, bookmarks: List[Bookmark],
                                 output_path: str, page_offset: int = 0) -> bool:
        """Insert bookmarks into PDF file"""
        try:
            doc = fitz.open(pdf_path)

            # Convert bookmarks to TOC format
            toc = []
            for bookmark in bookmarks:
                adjusted_page = bookmark.page + page_offset

                # Validate page number
                if adjusted_page < 1:
                    print(f"Warning: Skipping bookmark '{bookmark.title}' - invalid page {adjusted_page}")
                    continue
                elif adjusted_page > doc.page_count:
                    print(f"Warning: Skipping bookmark '{bookmark.title}' - page {adjusted_page} exceeds document length")
                    continue

                toc.append([bookmark.level, bookmark.title, adjusted_page])

            if toc:
                doc.set_toc(toc)
                doc.save(output_path)
                doc.close()
                print(f"Successfully inserted {len(toc)} bookmarks into PDF")

                # Add to history
                if self.history_manager:
                    self.history_manager.add_entry(
                        operation="Insert Bookmarks",
                        input_files=[pdf_path],
                        output_file=output_path,
                        status="Success",
                        details=f"Inserted {len(toc)} bookmarks with offset {page_offset}"
                    )

                return True
            else:
                print("No valid bookmarks to insert")
                doc.close()

                # Add to history
                if self.history_manager:
                    self.history_manager.add_entry(
                        operation="Insert Bookmarks",
                        input_files=[pdf_path],
                        output_file=output_path,
                        status="Failed",
                        details="No valid bookmarks to insert"
                    )

                return False

        except Exception as e:
            print(f"Error inserting bookmarks: {e}")
            return False

    def parse_toc_text(self, toc_text: str) -> List[Bookmark]:
        """Parse TOC text with simple level detection:
        - First line + every line preceded by empty line = Level 1
        - Otherwise = Level 2
        """
        import re

        bookmarks = []
        if not toc_text.strip():
            return bookmarks

        # Split into lines and keep track of empty lines
        all_lines = toc_text.split('\n')

        # Process each line and determine its level
        for i, line in enumerate(all_lines):
            line = line.strip()
            if not line:  # Skip empty lines
                continue

            title, page = self._parse_line_parts(line)
            if not title or not page:
                continue

            # Determine level based on position and preceding lines
            is_level1 = False

            # Check if this is the first non-empty line
            if i == 0 or all(not all_lines[j].strip() for j in range(i)):
                # First non-empty line is always Level 1
                is_level1 = True
            else:
                # Check if immediately preceded by an empty line
                if i > 0 and not all_lines[i-1].strip():
                    is_level1 = True
                else:
                    is_level1 = False

            level = 1 if is_level1 else 2
            bookmark = Bookmark(title=title, page=int(page), level=level)
            bookmarks.append(bookmark)

            print(f"Parsed: '{title}' -> Page {page}, Level {level} (line {i+1}, preceded by empty: {i > 0 and not all_lines[i-1].strip()})")  # Debug

        print(f"Total bookmarks parsed: {len(bookmarks)}")  # Debug
        return bookmarks

    def _parse_line_parts(self, line: str) -> tuple[str, str]:
        """
        Parse a line to separate title from page number.
        Returns tuple of (title, page_number) as strings.
        """
        try:
            # Handle format: "Title - PageNumber"
            if ' - ' in line:
                parts = line.strip().rsplit(' - ', 1)
                title = parts[0].strip()
                page_number = parts[1].strip()
                return title, page_number

            # Handle format: "Title ... PageNumber" or "Title PageNumber"
            # Remove excessive dots and normalize spaces
            cleaned_line = re.sub(r'\.{2,}', ' ', line)
            cleaned_line = re.sub(r'\s+', ' ', cleaned_line).strip()

            # Try to find page number at the end
            match = re.search(r'(.+?)\s+(\d+)$', cleaned_line)
            if match:
                title = match.group(1).strip()
                page_number = match.group(2).strip()
                return title, page_number

            # Try to find page number at the beginning
            match = re.search(r'^(\d+)\s+(.+)$', cleaned_line)
            if match:
                page_number = match.group(1).strip()
                title = match.group(2).strip()
                return title, page_number

            return line.strip(), ""

        except (IndexError, ValueError, AttributeError) as e:
            print(f"Error parsing line '{line}': {e}")
            return line.strip(), ""

    def insert_pages(self, source_pdf: str, target_pdf: str, pages_to_insert: List[int],
                    insert_position: int, output_path: str) -> bool:
        """Insert pages from source PDF into target PDF"""
        try:
            source_doc = fitz.open(source_pdf)
            target_doc = fitz.open(target_pdf)

            # Insert pages at specified position
            for i, page_num in enumerate(pages_to_insert):
                if 1 <= page_num <= source_doc.page_count:
                    # Get page from source (convert to 0-based index)
                    page = source_doc[page_num - 1]
                    # Insert into target at position + i
                    target_doc.insert_pdf(source_doc, from_page=page_num-1, to_page=page_num-1,
                                        start_at=insert_position + i)
                else:
                    print(f"Warning: Page {page_num} out of range in source PDF")

            target_doc.save(output_path)
            source_doc.close()
            target_doc.close()

            print(f"Successfully inserted {len(pages_to_insert)} pages")
            return True

        except Exception as e:
            print(f"Error inserting pages: {e}")
            return False

    def delete_pages(self, pdf_path: str, pages_to_delete: List[int], output_path: str) -> bool:
        """Delete specified pages from PDF"""
        try:
            doc = fitz.open(pdf_path)

            # Sort pages in descending order to avoid index shifting
            pages_to_delete.sort(reverse=True)

            for page_num in pages_to_delete:
                if 1 <= page_num <= doc.page_count:
                    # Delete page (convert to 0-based index)
                    doc.delete_page(page_num - 1)
                else:
                    print(f"Warning: Page {page_num} out of range")

            doc.save(output_path)
            doc.close()

            print(f"Successfully deleted {len(pages_to_delete)} pages")
            return True

        except Exception as e:
            print(f"Error deleting pages: {e}")
            return False

    def extract_pages(self, pdf_path: str, pages_to_extract: List[int], output_path: str) -> bool:
        """Extract specified pages to new PDF"""
        try:
            doc = fitz.open(pdf_path)
            new_doc = fitz.open()

            for page_num in pages_to_extract:
                if 1 <= page_num <= doc.page_count:
                    # Insert page into new document
                    new_doc.insert_pdf(doc, from_page=page_num-1, to_page=page_num-1)
                else:
                    print(f"Warning: Page {page_num} out of range")

            new_doc.save(output_path)
            doc.close()
            new_doc.close()

            print(f"Successfully extracted {len(pages_to_extract)} pages")
            return True

        except Exception as e:
            print(f"Error extracting pages: {e}")
            return False

    def merge_pdfs(self, input_paths: List[str], output_path: str) -> bool:
        """Merge multiple PDFs into one"""
        try:
            merged_doc = fitz.open()

            for pdf_path in input_paths:
                doc = fitz.open(pdf_path)
                merged_doc.insert_pdf(doc)
                doc.close()

            merged_doc.save(output_path)
            merged_doc.close()

            print(f"Successfully merged {len(input_paths)} PDFs")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Merge PDFs",
                    input_files=input_paths,
                    output_file=output_path,
                    status="Success",
                    details=f"Merged {len(input_paths)} PDF files"
                )

            return True

        except Exception as e:
            print(f"Error merging PDFs: {e}")
            return False

    def add_watermark(self, pdf_path: str, output_path: str, watermark_text: str,
                     position: str = "center", opacity: float = 0.5, font_size: int = 50) -> bool:
        """Add text watermark to PDF"""
        try:
            doc = fitz.open(pdf_path)

            for page_num in range(doc.page_count):
                page = doc[page_num]
                rect = page.rect

                # Calculate position
                if position == "center":
                    x, y = rect.width / 2, rect.height / 2
                elif position == "top-left":
                    x, y = 50, 50
                elif position == "top-right":
                    x, y = rect.width - 200, 50
                elif position == "bottom-left":
                    x, y = 50, rect.height - 50
                elif position == "bottom-right":
                    x, y = rect.width - 200, rect.height - 50
                else:
                    x, y = rect.width / 2, rect.height / 2

                # Add watermark text
                page.insert_text(
                    (x, y), watermark_text,
                    fontsize=font_size,
                    color=(0.5, 0.5, 0.5),  # Gray color
                    rotate=45  # Diagonal
                )

            doc.save(output_path)
            doc.close()

            print("Successfully added watermark")
            return True

        except Exception as e:
            print(f"Error adding watermark: {e}")
            return False

    def remove_watermark(self, pdf_path: str, output_path: str, aggressive_mode: bool = False,
                        target_updf: bool = True, target_urls: bool = True) -> bool:
        """Enhanced watermark removal for various types including UPDF and similar watermarks"""
        try:
            doc = fitz.open(pdf_path)
            removed_count = 0

            # Store original annotations to preserve their properties
            original_annotations = {}

            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_removed = 0

                # Store annotations before processing to preserve their properties
                page_annotations = []
                for annot in page.annots():
                    annot_data = {
                        'type': annot.type,
                        'rect': annot.rect,
                        'content': annot.info.get('content', ''),
                        'opacity': getattr(annot, 'opacity', 1.0),
                        'colors': annot.colors,
                        'border': annot.border
                    }
                    page_annotations.append(annot_data)
                original_annotations[page_num] = page_annotations

                # Method 1: Remove images (logos, image watermarks)
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        # Get image rectangle
                        img_rect = page.get_image_rects(img[0])
                        for rect in img_rect:
                            # Check if image is in typical watermark positions (corners, small size)
                            page_rect = page.rect
                            img_width = rect.width
                            img_height = rect.height

                            # Remove small images in corners or top area (likely watermarks)
                            is_corner_watermark = (
                                (rect.x0 < page_rect.width * 0.3 and rect.y0 < page_rect.height * 0.3) or  # Top-left
                                (rect.x1 > page_rect.width * 0.7 and rect.y0 < page_rect.height * 0.3) or  # Top-right
                                (rect.x0 < page_rect.width * 0.3 and rect.y1 > page_rect.height * 0.7) or  # Bottom-left
                                (rect.x1 > page_rect.width * 0.7 and rect.y1 > page_rect.height * 0.7)     # Bottom-right
                            )

                            is_small_image = img_width < 200 and img_height < 200

                            if is_corner_watermark or is_small_image:
                                # Check if this overlaps with any annotation before removing
                                overlaps_annotation = False
                                for annot in page.annots():
                                    if rect.intersects(annot.rect):
                                        overlaps_annotation = True
                                        break

                                # Only remove if it doesn't overlap with annotations
                                if not overlaps_annotation:
                                    # Cover with white rectangle
                                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page_removed += 1
                    except:
                        continue

                # Method 2: Remove text watermarks (including URLs and common watermark text)
                text_instances = page.get_text("dict")
                watermark_keywords = ["watermark", "trial", "demo", "sample", "confidential", "draft", "preview", "evaluation", "unregistered", "free version", "trial version"]

                # Add UPDF-specific keywords if targeting UPDF
                if target_updf:
                    watermark_keywords.extend(["updf", "www.updf.com"])

                # Add URL patterns if targeting URLs
                if target_urls:
                    watermark_keywords.extend(["www.", ".com", ".net", ".org", "http://", "https://"])

                for block in text_instances.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                text = span.get("text", "").strip().lower()

                                # Check if text contains watermark keywords
                                is_watermark_text = any(keyword in text for keyword in watermark_keywords)

                                # Also remove very small text (often watermarks)
                                font_size = span.get("size", 12)
                                is_small_text = font_size < 10 if not aggressive_mode else font_size < 12

                                # Remove short URLs or website references
                                is_url_like = target_urls and ("www." in text or ".com" in text or ".net" in text)

                                # In aggressive mode, remove more suspicious text
                                is_suspicious = False
                                if aggressive_mode:
                                    is_suspicious = (len(text) < 30 and
                                                   (any(char.isdigit() for char in text) or
                                                    text.isupper() or
                                                    len(text.split()) <= 2))

                                if is_watermark_text or (is_small_text and len(text) < 20) or is_url_like or is_suspicious:
                                    bbox = span.get("bbox")
                                    if bbox:
                                        # Expand the rectangle slightly to ensure complete coverage
                                        rect = fitz.Rect(bbox)
                                        rect.x0 -= 2
                                        rect.y0 -= 2
                                        rect.x1 += 2
                                        rect.y1 += 2

                                        # Check if this overlaps with any annotation before removing
                                        overlaps_annotation = False
                                        for annot in page.annots():
                                            if rect.intersects(annot.rect):
                                                overlaps_annotation = True
                                                break

                                        # Only remove if it doesn't overlap with annotations
                                        if not overlaps_annotation:
                                            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                                            page_removed += 1

                # Method 3: Remove drawings/vector graphics that might be watermarks
                # BUT preserve annotations and highlights by checking for overlaps
                try:
                    # Get all annotations on this page to avoid removing them
                    annotations = page.annots()
                    annot_rects = [annot.rect for annot in annotations]

                    drawings = page.get_drawings()
                    for drawing in drawings:
                        # Check if drawing is in watermark-typical locations
                        if 'rect' in drawing:
                            draw_rect = drawing['rect']
                            if (draw_rect.width < 150 and draw_rect.height < 150 and
                                (draw_rect.x0 < 200 or draw_rect.y0 < 200)):  # Small drawings in top-left area

                                # Check if this drawing overlaps with any annotation
                                overlaps_annotation = False
                                for annot_rect in annot_rects:
                                    if draw_rect.intersects(annot_rect):
                                        overlaps_annotation = True
                                        break

                                # Only remove if it doesn't overlap with annotations
                                if not overlaps_annotation:
                                    page.draw_rect(draw_rect, color=(1, 1, 1), fill=(1, 1, 1))
                                    page_removed += 1
                except:
                    pass

                removed_count += page_removed

            # Restore annotation properties after watermark removal
            for page_num, page_annotations in original_annotations.items():
                page = doc[page_num]
                current_annotations = list(page.annots())

                # Match and restore annotation properties
                for i, annot in enumerate(current_annotations):
                    if i < len(page_annotations):
                        original_data = page_annotations[i]
                        try:
                            # Restore opacity if it was changed
                            if 'opacity' in original_data and original_data['opacity'] != 1.0:
                                annot.set_opacity(original_data['opacity'])
                                annot.update()
                        except:
                            pass

            doc.save(output_path)
            doc.close()

            print(f"Watermark removal completed. Removed {removed_count} potential watermark elements.")
            return True

        except Exception as e:
            print(f"Error removing watermark: {e}")
            return False

    def extract_images(self, pdf_path: str, output_dir: str, page_range: str = "all") -> int:
        """Extract images from PDF with enhanced detection"""
        try:
            doc = fitz.open(pdf_path)
            os.makedirs(output_dir, exist_ok=True)

            # Parse page range
            if page_range == "all":
                pages = range(doc.page_count)
            else:
                pages = self.parse_page_numbers(page_range)
                pages = [p-1 for p in pages if 1 <= p <= doc.page_count]  # Convert to 0-based

            image_count = 0

            for page_num in pages:
                page = doc[page_num]

                # Method 1: Extract embedded images
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        # Save image
                        image_filename = f"page_{page_num+1}_embedded_{img_index+1}.{image_ext}"
                        image_path = os.path.join(output_dir, image_filename)

                        with open(image_path, "wb") as img_file:
                            img_file.write(image_bytes)

                        image_count += 1
                        print(f"Extracted embedded image: {image_filename}")

                    except Exception as e:
                        print(f"Error extracting embedded image {img_index}: {e}")
                        continue

                # Method 2: Render page as image (for vector graphics, text as images, etc.)
                try:
                    # High resolution rendering
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)

                    page_image_filename = f"page_{page_num+1}_rendered.png"
                    page_image_path = os.path.join(output_dir, page_image_filename)
                    pix.save(page_image_path)
                    pix = None

                    image_count += 1
                    print(f"Rendered page as image: {page_image_filename}")

                except Exception as e:
                    print(f"Error rendering page {page_num+1}: {e}")

            doc.close()
            print(f"Total extracted: {image_count} images to {output_dir}")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Extract Images",
                    input_files=[pdf_path],
                    output_file=output_dir,
                    status="Success" if image_count > 0 else "Failed",
                    details=f"Extracted {image_count} images from {len(pages)} pages"
                )

            return image_count

        except Exception as e:
            print(f"Error extracting images: {e}")
            return 0

    def extract_text(self, pdf_path: str, output_path: str, page_range: str = "all",
                    format_type: str = "txt") -> bool:
        """Extract text from PDF with support for multiple formats including DOCX"""
        try:
            doc = fitz.open(pdf_path)

            # Parse page range
            print(f"PDF has {doc.page_count} pages")
            print(f"Page range input: '{page_range}'")

            # Handle localized "all pages" text
            if page_range == "all" or page_range == "All Pages" or page_range == "جميع الصفحات" or not page_range.strip():
                pages = range(doc.page_count)
                print(f"Processing all pages: {list(pages)}")
            else:
                parsed_pages = self.parse_page_numbers(page_range)
                print(f"Parsed pages: {parsed_pages}")
                pages = [p-1 for p in parsed_pages if 1 <= p <= doc.page_count]  # Convert to 0-based
                print(f"Valid pages (0-based): {pages}")

            extracted_text = []
            total_text_length = 0

            for page_num in pages:
                page = doc[page_num]
                text = page.get_text()

                # Debug: Print text length for each page
                print(f"Page {page_num + 1}: extracted {len(text)} characters")
                total_text_length += len(text)

                # Try alternative text extraction methods if no text found
                if not text.strip():
                    # Try different text extraction methods
                    text = page.get_text("text")  # Default method
                    if not text.strip():
                        text = page.get_text("dict")  # Dictionary format
                        if isinstance(text, dict) and 'blocks' in text:
                            text_parts = []
                            for block in text['blocks']:
                                if 'lines' in block:
                                    for line in block['lines']:
                                        if 'spans' in line:
                                            for span in line['spans']:
                                                if 'text' in span:
                                                    text_parts.append(span['text'])
                            text = ' '.join(text_parts)

                    print(f"Page {page_num + 1}: after alternative extraction: {len(text)} characters")

                if format_type == "json":
                    extracted_text.append({
                        "page": page_num + 1,
                        "text": text
                    })
                elif format_type == "docx":
                    # For DOCX, we'll store text with page information
                    extracted_text.append({
                        "page": page_num + 1,
                        "text": text
                    })
                else:
                    extracted_text.append(f"=== Page {page_num + 1} ===\n{text}\n")

            print(f"Total text extracted: {total_text_length} characters from {len(pages)} pages")

            # Save extracted text based on format
            if format_type == "json":
                import json
                print(f"Saving JSON with {len(extracted_text)} pages to {output_path}")
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(extracted_text, f, ensure_ascii=False, indent=2)
                print(f"JSON file saved. Size: {os.path.getsize(output_path)} bytes")
            elif format_type == "docx":
                print(f"Saving DOCX with {len(extracted_text)} pages to {output_path}")
                self._save_as_docx(extracted_text, output_path)
                if os.path.exists(output_path):
                    print(f"DOCX file saved. Size: {os.path.getsize(output_path)} bytes")
                else:
                    print("DOCX file was not created!")
            else:
                print(f"Saving TXT with {len(extracted_text)} entries to {output_path}")
                content_to_write = '\n'.join(extracted_text)
                print(f"Content length: {len(content_to_write)} characters")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content_to_write)
                print(f"TXT file saved. Size: {os.path.getsize(output_path)} bytes")

            doc.close()
            print(f"Extracted text from {len(pages)} pages to {format_type} format")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Extract Text",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Extracted text from {len(pages)} pages in {format_type} format"
                )

            return True

        except Exception as e:
            print(f"Error extracting text: {e}")

            # Add to history even on failure
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Extract Text",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Failed",
                    details=f"Error: {str(e)}"
                )

            return False

    def _save_as_docx(self, extracted_text: list, output_path: str):
        """Save extracted text as DOCX with page breaks"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_BREAK

            # Create new document
            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add content
            for i, page_data in enumerate(extracted_text):
                if isinstance(page_data, dict):
                    page_num = page_data['page']
                    text = page_data['text']
                else:
                    # Fallback for simple text format
                    page_num = i + 1
                    text = str(page_data)

                # Add page header
                if i > 0:
                    # Add page break before new page (except for first page)
                    doc.add_page_break()

                # Add page number as heading
                heading = doc.add_heading(f'Page {page_num}', level=2)

                # Add page content
                if text.strip():
                    # Split text into paragraphs
                    paragraphs = text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            # Set paragraph spacing
                            para.paragraph_format.space_after = Inches(0.1)
                else:
                    doc.add_paragraph("[No text content on this page]")

            # Save document
            doc.save(output_path)

        except ImportError:
            raise Exception("python-docx library is required for DOCX export. Please install it using: pip install python-docx")
        except Exception as e:
            raise Exception(f"Error creating DOCX file: {e}")

    def extract_text_batch(self, pdf_files: list, output_dir: str, page_range: str = "all",
                          format_type: str = "txt", progress_callback=None) -> tuple:
        """Extract text from multiple PDF files"""
        successful = []
        failed = []

        for i, pdf_path in enumerate(pdf_files):
            if progress_callback:
                progress_callback(i + 1, len(pdf_files), os.path.basename(pdf_path))

            try:
                # Generate output filename
                pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
                output_filename = f"{pdf_name}_text.{format_type}"
                output_path = os.path.join(output_dir, output_filename)

                # Extract text
                success = self.extract_text(pdf_path, output_path, page_range, format_type)

                if success:
                    successful.append({
                        'input': pdf_path,
                        'output': output_path,
                        'filename': output_filename
                    })
                else:
                    failed.append({
                        'input': pdf_path,
                        'error': 'Extraction failed'
                    })

            except Exception as e:
                failed.append({
                    'input': pdf_path,
                    'error': str(e)
                })

        return successful, failed

    def rotate_pages(self, pdf_path: str, output_path: str, page_range: str, rotation: int) -> bool:
        """Rotate specific pages in PDF"""
        try:
            doc = fitz.open(pdf_path)

            # Parse page range
            pages = self.parse_page_numbers(page_range)
            pages = [p-1 for p in pages if 1 <= p <= doc.page_count]  # Convert to 0-based

            for page_num in pages:
                page = doc[page_num]
                page.set_rotation(rotation)

            doc.save(output_path)
            doc.close()

            print(f"Rotated {len(pages)} pages by {rotation} degrees")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Rotate Pages",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Rotated {len(pages)} pages by {rotation} degrees"
                )

            return True

        except Exception as e:
            print(f"Error rotating pages: {e}")
            return False

    def compress_pdf(self, pdf_path: str, output_path: str) -> bool:
        """Compress PDF file"""
        try:
            doc = fitz.open(pdf_path)
            doc.save(output_path, garbage=4, deflate=True, clean=True)
            doc.close()

            # Check compression ratio
            original_size = os.path.getsize(pdf_path)
            compressed_size = os.path.getsize(output_path)
            ratio = (1 - compressed_size / original_size) * 100

            print(f"PDF compressed by {ratio:.1f}%")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Compress PDF",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Compressed PDF by {ratio:.1f}%"
                )

            return True

        except Exception as e:
            print(f"Error compressing PDF: {e}")
            return False

    def remove_pdf_security(self, pdf_path: str, output_path: str, password: str = "") -> bool:
        """Remove security restrictions from PDF to enable editing"""
        try:
            # Try to open the PDF with password if provided
            doc = fitz.open(pdf_path)

            # Check if PDF is encrypted
            if doc.needs_pass:
                if password:
                    # Try to authenticate with password
                    if not doc.authenticate(password):
                        print("Invalid password provided")
                        doc.close()
                        return False
                else:
                    print("PDF is password protected. Password required.")
                    doc.close()
                    return False

            # Check if PDF has restrictions
            if doc.is_encrypted:
                print("PDF has security restrictions. Attempting to remove...")

            # Save without restrictions - this removes most security features
            # The key is to save with specific flags that remove restrictions
            doc.save(output_path,
                    garbage=4,          # Remove unused objects
                    deflate=True,       # Compress
                    clean=True,         # Clean up
                    encryption=fitz.PDF_ENCRYPT_NONE,  # Remove encryption
                    permissions=-1      # Grant all permissions
                    )

            doc.close()

            # Verify the output file was created
            if os.path.exists(output_path):
                # Test if the new file can be opened and edited
                test_doc = fitz.open(output_path)
                can_edit = not test_doc.is_encrypted and not test_doc.needs_pass
                test_doc.close()

                if can_edit:
                    print("Successfully removed PDF security restrictions")

                    # Add to history
                    if self.history_manager:
                        self.history_manager.add_entry(
                            operation="Remove PDF Security",
                            input_files=[pdf_path],
                            output_file=output_path,
                            status="Success",
                            details="Removed security restrictions to enable editing"
                        )

                    return True
                else:
                    print("Warning: Some restrictions may still remain")
                    return True  # Still return True as file was processed
            else:
                print("Failed to create output file")
                return False

        except Exception as e:
            print(f"Error removing PDF security: {e}")
            return False

    def add_pdf_annotations(self, pdf_path: str, output_path: str, annotations: List[Dict]) -> bool:
        """Add annotations to PDF file"""
        try:
            doc = fitz.open(pdf_path)

            for annotation in annotations:
                page_num = annotation.get('page', 0)
                if page_num >= doc.page_count:
                    continue

                page = doc[page_num]
                annot_type = annotation.get('type', 'text')

                # Get position
                x1, y1, x2, y2 = annotation.get('rect', [100, 100, 300, 150])
                rect = fitz.Rect(x1, y1, x2, y2)

                if annot_type == 'text':
                    # Add text annotation (sticky note)
                    text = annotation.get('text', 'Comment')
                    annot = page.add_text_annot(rect.tl, text)

                elif annot_type == 'freetext':
                    # Add free text annotation (visible text box)
                    text = annotation.get('text', 'Free text comment')
                    fontsize = annotation.get('fontsize', 12)
                    fontname = annotation.get('fontname', 'helv')
                    text_color = annotation.get('text_color', (1, 0, 0))  # Red by default

                    page.add_freetext_annot(rect, text,
                                          fontsize=fontsize,
                                          fontname=fontname,
                                          text_color=text_color)

                elif annot_type == 'highlight':
                    # Add highlight annotation
                    color = annotation.get('color', (1, 1, 0))  # Yellow by default
                    annot = page.add_highlight_annot(rect)
                    annot.set_colors(stroke=color)
                    annot.update()

            # Save the annotated PDF
            doc.save(output_path)
            doc.close()

            print(f"Successfully added {len(annotations)} annotations")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Add PDF Annotations",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Added {len(annotations)} annotations"
                )

            return True

        except Exception as e:
            print(f"Error adding annotations: {e}")
            return False

    def crop_pages(self, pdf_path: str, output_path: str, page_range: str,
                   crop_box: Tuple[float, float, float, float]) -> bool:
        """Crop pages in PDF"""
        try:
            doc = fitz.open(pdf_path)

            # Parse page range
            pages = self.parse_page_numbers(page_range)
            pages = [p-1 for p in pages if 1 <= p <= doc.page_count]  # Convert to 0-based

            x0, y0, x1, y1 = crop_box

            for page_num in pages:
                page = doc[page_num]
                # Set crop box
                page.set_cropbox(fitz.Rect(x0, y0, x1, y1))

            doc.save(output_path)
            doc.close()

            print(f"Cropped {len(pages)} pages")
            return True

        except Exception as e:
            print(f"Error cropping pages: {e}")
            return False

    def add_margins(self, pdf_path: str, output_path: str, page_range: str,
                   margin_size: float) -> bool:
        """Add white space margins around pages"""
        try:
            doc = fitz.open(pdf_path)

            # Parse page range
            pages = self.parse_page_numbers(page_range)
            pages = [p-1 for p in pages if 1 <= p <= doc.page_count]  # Convert to 0-based

            for page_num in pages:
                page = doc[page_num]
                original_rect = page.rect

                # Create new page with margins
                new_width = original_rect.width + (2 * margin_size)
                new_height = original_rect.height + (2 * margin_size)

                # Insert new page after current
                new_page = doc.new_page(page_num + 1, width=new_width, height=new_height)

                # Copy content to new page with offset
                new_page.show_pdf_page(
                    fitz.Rect(margin_size, margin_size,
                             margin_size + original_rect.width,
                             margin_size + original_rect.height),
                    doc, page_num
                )

                # Delete original page
                doc.delete_page(page_num)

            doc.save(output_path)
            doc.close()

            print(f"Added margins to {len(pages)} pages")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Add Margins",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Added {margin_size}pt margins to {len(pages)} pages"
                )

            return True

        except Exception as e:
            print(f"Error adding margins: {e}")
            return False

    def parse_page_numbers(self, page_string: str) -> List[int]:
        """Parse page number string like '1,3,5-10' into list of integers"""
        pages = []

        try:
            parts = page_string.split(',')
            for part in parts:
                part = part.strip()
                if '-' in part:
                    # Handle range like "5-10"
                    start, end = part.split('-')
                    start, end = int(start.strip()), int(end.strip())
                    pages.extend(range(start, end + 1))
                else:
                    # Handle single page
                    pages.append(int(part))

            return sorted(list(set(pages)))  # Remove duplicates and sort

        except ValueError:
            return []

    def insert_blank_pages(self, pdf_path: str, output_path: str, position: int, count: int) -> bool:
        """Insert blank pages into PDF"""
        try:
            doc = fitz.open(pdf_path)

            # Validate position
            if position < 1 or position > doc.page_count + 1:
                print(f"Invalid position: {position}")
                return False

            # Get page size from first page
            if doc.page_count > 0:
                first_page = doc[0]
                page_rect = first_page.rect
            else:
                # Default A4 size
                page_rect = fitz.Rect(0, 0, 595, 842)

            # Insert blank pages
            for i in range(count):
                # Insert at position (convert to 0-based index)
                insert_pos = position - 1 + i
                new_page = doc.new_page(insert_pos, width=page_rect.width, height=page_rect.height)

            doc.save(output_path)
            doc.close()

            print(f"Successfully inserted {count} blank pages at position {position}")

            # Add to history
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="Insert Blank Pages",
                    input_files=[pdf_path],
                    output_file=output_path,
                    status="Success",
                    details=f"Inserted {count} blank pages at position {position}"
                )

            return True

        except Exception as e:
            print(f"Error inserting blank pages: {e}")
            return False


class BookmarkTab(QWidget):
    """Bookmark management tab"""

    def __init__(self, history_manager=None, localization=None, main_window=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.main_window = main_window
        self.localization = localization or Localization()
        self.bookmarks = []
        self.pdf_path = ""
        self.bookmark_file_path = ""
        self.pdf_doc = None
        self.current_page = 0
        self.total_pages = 0
        self.selected_bookmark_row = -1
        self.selected_bookmark = None

        # Track cumulative offset applied to bookmarks
        self.cumulative_offset = 0
        self.last_applied_offset = 0

        self.init_ui()

    def create_arabic_message_box(self, title, text, icon=QMessageBox.Information):
        """Create a message box with proper Arabic RTL alignment"""
        msg_box = QMessageBox(self)
        msg_box.setIcon(icon)
        msg_box.setWindowTitle(title)
        msg_box.setText(text)
        msg_box.setTextFormat(Qt.RichText)
        msg_box.setLayoutDirection(Qt.RightToLeft)
        msg_box.setStyleSheet("""
            QMessageBox {
                text-align: right;
                direction: rtl;
            }
            QMessageBox QLabel {
                text-align: right;
                qproperty-alignment: AlignRight;
            }
        """)
        return msg_box

    def init_ui(self):
        """Initialize redesigned bookmark tab UI with clear step-by-step workflow"""
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignTop)  # Push all content to top
        layout.setSpacing(10)  # Minimal spacing between steps
        layout.setContentsMargins(10, 10, 10, 10)  # Small margins

        # Add header section with bookmark copy feature
        self.add_header_section(layout)

        # Initialize method selection state
        self.selected_method = None  # None, "toc_pages", or "text_file"

        # Step 1: Load PDF File (Full width)
        self.step1_group = QGroupBox(f"📄 {self.localization.get_text('step1_load_pdf')}")
        self.step1_group.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; color: #1976D2; padding: 8px; }")
        self.step1_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        step1_layout = QVBoxLayout(self.step1_group)
        step1_layout.setContentsMargins(10, 15, 10, 10)  # Add proper margins

        pdf_layout = QHBoxLayout()
        pdf_layout.setSpacing(10)  # Add spacing between elements

        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_placeholder"))
        # Increase height for better visibility
        self.pdf_display.setStyleSheet("""
            QLineEdit {
                padding: 8px 12px;
                font-size: 12px;
                border: 2px solid #ddd;
                border-radius: 6px;
                background-color: #f9f9f9;
                color: #333;  /* Explicit text color for visibility */
                min-height: 20px;
            }
            QLineEdit:focus {
                border-color: #1976D2;
                background-color: #353535;
                color: #ffffff;
            }
        """)
        pdf_layout.addWidget(self.pdf_display, 7)  # Give more space to text field

        self.btn_select_pdf = QPushButton(self.localization.get_text("load_pdf_primary"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        # Reduce button height and improve styling
        self.btn_select_pdf.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                font-weight: bold;
                padding: 8px 12px;
                font-size: 12px;
                border: none;
                border-radius: 6px;
                min-height: 20px;
                max-height: 36px;
                min-width: 110px;
            }
            QPushButton:hover {
                background-color: #1565C0;
            }
            QPushButton:pressed {
                background-color: #0D47A1;
            }
        """)
        pdf_layout.addWidget(self.btn_select_pdf, 1)  # Give less space to button

        step1_layout.addLayout(pdf_layout)

        # Increase height for Step 1 to accommodate larger text field
        self.step1_group.setMaximumHeight(100)
        self.step1_group.setMinimumHeight(100)

        layout.addWidget(self.step1_group)

        # Step 2: Paste Bookmarks (Full width, default visible)
        self.step2_group = QGroupBox(f"📝 {self.localization.get_text('step2_bookmark_extraction')}")
        self.step2_group.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; color: #FF9800; }")
        self.step2_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        step2_layout = QVBoxLayout(self.step2_group)

        # Top row: Guidance text and Prepare Bookmarks button
        top_row_layout = QHBoxLayout()

        guidance_label = QLabel(self.localization.get_text('step2_guidance_text'))
        guidance_label.setStyleSheet("color: #666; font-weight: bold; font-size: 12px;")
        top_row_layout.addWidget(guidance_label)

        top_row_layout.addStretch()  # Push help text and button to the right

        help_label = QLabel(self.localization.get_text('step2_help_text'))
        help_label.setStyleSheet("color: #666; font-size: 11px; font-style: italic;")
        top_row_layout.addWidget(help_label)

        self.btn_prepare_bookmarks = QPushButton(f"📋 {self.localization.get_text('prepare_bookmarks_button')}")
        self.btn_prepare_bookmarks.clicked.connect(self.open_bookmark_preparation_dialog)
        self.btn_prepare_bookmarks.setEnabled(False)  # Enabled when PDF is loaded
        self.btn_prepare_bookmarks.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 8px 15px;
                font-size: 11px;
                border: none;
                border-radius: 4px;
                min-height: 30px;
                max-height: 30px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666;
            }
        """)
        top_row_layout.addWidget(self.btn_prepare_bookmarks)

        step2_layout.addLayout(top_row_layout)

        # Large expanding text area for pasting bookmarks
        self.toc_text = QTextEdit()
        self.toc_text.setPlaceholderText(self.localization.get_text("toc_paste_placeholder"))
        self.toc_text.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        # Set proper text wrapping and size constraints
        self.toc_text.setLineWrapMode(QTextEdit.WidgetWidth)
        self.toc_text.setMinimumHeight(150)  # Ensure minimum readable height
        self.toc_text.setMaximumHeight(400)  # Prevent excessive expansion initially

        # Add styling for better appearance
        self.toc_text.setStyleSheet("""
            QTextEdit {
                border: 2px solid #ddd;
                border-radius: 6px;
                padding: 8px;
                font-size: 12px;
                background-color: #fafafa;
                color: #333;  /* Explicit text color for visibility */
                font-family: 'Segoe UI', Arial, sans-serif;
            }
            QTextEdit:focus {
                border-color: #1976D2;
                background-color: #353535;
                color: #ffffff;
            }
        """)

        self.toc_text.textChanged.connect(self.on_toc_text_changed)
        step2_layout.addWidget(self.toc_text, 1)  # Give stretch factor to expand

        layout.addWidget(self.step2_group)
        # Add stretch to push Step 3 to the bottom
        layout.addStretch()





        # Step 3: Generate Bookmarks (simplified - removed separate format step)
        self.step3_group = QGroupBox(f"✅ {self.localization.get_text('step3_generate_bookmarks')}")
        self.step3_group.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; color: #4CAF50; }")
        self.step3_group.setVisible(False)  # Initially hidden until text is pasted
        self.step3_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        step3_layout = QVBoxLayout(self.step3_group)

        # Generate bookmarks button (merged format and load functionality)
        self.btn_load_bookmarks = QPushButton(self.localization.get_text("load_bookmarks_preview"))
        self.btn_load_bookmarks.clicked.connect(self.format_and_load_files)
        self.btn_load_bookmarks.setEnabled(False)  # Enabled when text is pasted
        self.btn_load_bookmarks.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 12px;
                border-radius: 6px;
                min-height: 40px;
            }
            QPushButton:hover { background-color: #45a049; }
            QPushButton:disabled { background-color: #cccccc; color: #666; }
        """)
        step3_layout.addWidget(self.btn_load_bookmarks)

        layout.addWidget(self.step3_group)

        # Hidden spinboxes for bookmark preparation dialog (not visible in main UI)
        self.toc_start_page = QSpinBox()
        self.toc_start_page.setMinimum(1)
        self.toc_start_page.setMaximum(9999)
        self.toc_start_page.setValue(1)
        self.toc_start_page.setVisible(False)

        self.toc_end_page = QSpinBox()
        self.toc_end_page.setMinimum(1)
        self.toc_end_page.setMaximum(9999)
        self.toc_end_page.setValue(5)
        self.toc_end_page.setVisible(False)

        # Step 4: Verify Bookmarks and Insert (Initially hidden)
        self.step4_group = QGroupBox(f"✅ {self.localization.get_text('step4_verify_insert')}")
        self.step4_group.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; color: #4CAF50; }")
        # Set size policy to allow expansion when needed
        self.step4_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        step4_layout = QVBoxLayout(self.step4_group)

        # Top controls for Step 4
        step4_controls = QHBoxLayout()

        # Verification instructions
        verify_instructions = QLabel(self.localization.get_text("verify_bookmarks"))
        verify_instructions.setWordWrap(True)
        verify_instructions.setStyleSheet("color: #666; font-style: italic; margin-bottom: 10px;")
        step4_controls.addWidget(verify_instructions)

        step4_controls.addStretch()

        # Step 4 Reset button
        self.btn_step4_reset = QPushButton("🔄 " + self.localization.get_text("reset_button"))
        self.btn_step4_reset.clicked.connect(self.reset_step4_workflow)
        self.btn_step4_reset.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        self.btn_step4_reset.setToolTip(
            "إعادة تعيين عملية تصحيح الإشارات المرجعية" if self.localization.current_language == "ar"
            else "Reset bookmark correction process"
        )
        step4_controls.addWidget(self.btn_step4_reset)

        # REMOVED: Show Previous Steps button (no longer needed)

        step4_layout.addLayout(step4_controls)

        # Main content with splitter for table and preview
        splitter = QSplitter(Qt.Horizontal)

        # Left side - Bookmark table (expanded to take more space)
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)

        # Page offset is now handled automatically through navigation
        self.page_offset_value = 0  # Internal tracking

        # Verification tracking
        self.verified_bookmarks = set()  # Track which bookmarks have been verified

        # Bookmark table
        table_group = QGroupBox(self.localization.get_text("bookmarks_table"))
        table_layout = QVBoxLayout(table_group)

        self.bookmark_table = QTableWidget()
        self.bookmark_table.setColumnCount(4)  # Reduced from 5 to 4 (removed Actions column)
        headers = [
            self.localization.get_text("title"),
            self.localization.get_text("original"),
            "الصفحة المعدلة",  # Modified Page in Arabic
            "المستوى"        # Level in Arabic
        ]
        self.bookmark_table.setHorizontalHeaderLabels(headers)
        self.bookmark_table.setColumnWidth(0, 400)  # Wider title column (increased from 300)
        self.bookmark_table.setColumnWidth(1, 100)  # Wider original page column
        self.bookmark_table.setColumnWidth(2, 120)  # Wider for modified page
        self.bookmark_table.setColumnWidth(3, 100)  # Wider for level
        # REMOVED: Actions column (column 4)
        self.bookmark_table.cellClicked.connect(self.on_bookmark_clicked)
        self.bookmark_table.setAlternatingRowColors(True)
        self.bookmark_table.setSelectionBehavior(QTableWidget.SelectRows)
        table_layout.addWidget(self.bookmark_table)
        left_layout.addWidget(table_group)

        # Insert button
        self.btn_insert = QPushButton("إدراج الإشارات المرجعية في الملف")  # Insert bookmarks in the file
        self.btn_insert.clicked.connect(self.insert_bookmarks)
        self.btn_insert.setEnabled(False)

        # Add debug button click handler
        def debug_button_click():
            print("DEBUG: Insert button clicked - connection working!")
            self.insert_bookmarks()

        # Temporarily override connection for debugging
        self.btn_insert.clicked.disconnect()
        self.btn_insert.clicked.connect(debug_button_click)
        self.btn_insert.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                font-size: 16px;
                padding: 15px;
                border-radius: 8px;
                margin: 10px 0;
            }
            QPushButton:hover { background-color: #45a049; }
            QPushButton:disabled { background-color: #cccccc; }
        """)
        left_layout.addWidget(self.btn_insert)

        # REMOVED: Test Connection button (no longer needed)

        splitter.addWidget(left_widget)

        # Right side - PDF preview
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)

        preview_group = QGroupBox(self.localization.get_text("pdf_preview"))
        preview_layout = QVBoxLayout(preview_group)

        # Navigation with offset adjustment
        nav_layout = QHBoxLayout()

        # Previous button
        self.btn_prev = QPushButton("◀ السابق")
        self.btn_prev.clicked.connect(self.prev_page)
        self.btn_prev.setEnabled(False)
        self.btn_prev.setStyleSheet("QPushButton { padding: 8px 12px; }")
        nav_layout.addWidget(self.btn_prev)

        # Page info and offset adjustment
        page_info_layout = QVBoxLayout()

        self.page_label = QLabel("لم يتم تحميل PDF")
        self.page_label.setAlignment(Qt.AlignCenter)
        self.page_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        page_info_layout.addWidget(self.page_label)

        # Offset adjustment controls
        offset_layout = QHBoxLayout()
        offset_layout.addWidget(QLabel("إزاحة الصفحة:"))

        self.offset_spinbox = QSpinBox()
        self.offset_spinbox.setRange(-100, 100)
        self.offset_spinbox.setValue(0)
        self.offset_spinbox.valueChanged.connect(self.on_offset_changed)
        self.offset_spinbox.setStyleSheet("QSpinBox { text-align: center; padding: 4px; }")
        offset_layout.addWidget(self.offset_spinbox)

        # REMOVED: Fix button - automatic navigation updates now handle page adjustments

        # Confirmation button (kept for backward compatibility)
        self.btn_confirm_offset = QPushButton("تأكيد الإزاحة")
        self.btn_confirm_offset.clicked.connect(self.confirm_offset_adjustment)
        self.btn_confirm_offset.setVisible(False)  # Hidden by default
        self.btn_confirm_offset.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border-radius: 4px;
            }
            QPushButton:hover { background-color: #F57C00; }
        """)
        offset_layout.addWidget(self.btn_confirm_offset)

        page_info_layout.addLayout(offset_layout)
        nav_layout.addLayout(page_info_layout)

        # Next button
        self.btn_next = QPushButton("التالي ▶")
        self.btn_next.clicked.connect(self.next_page)
        self.btn_next.setEnabled(False)
        self.btn_next.setStyleSheet("QPushButton { padding: 8px 12px; }")
        nav_layout.addWidget(self.btn_next)

        preview_layout.addLayout(nav_layout)

        # Preview area - Expanded size
        self.pdf_preview = QLabel()
        self.pdf_preview.setAlignment(Qt.AlignCenter)
        self.pdf_preview.setMinimumSize(400, 600)  # Larger minimum size
        self.pdf_preview.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)  # Allow expansion
        self.pdf_preview.setStyleSheet("""
            border: 2px solid #666;
            background-color: white;
            border-radius: 4px;
            padding: 5px;
        """)
        self.pdf_preview.setText("معاينة PDF\nسيتم عرض الصفحات هنا")
        preview_layout.addWidget(self.pdf_preview)

        right_layout.addWidget(preview_group)
        splitter.addWidget(right_widget)

        # Set splitter proportions - give more space to table
        splitter.setSizes([700, 500])  # Table: 700px, Preview: 500px
        splitter.setStretchFactor(0, 1)  # Table stretches
        splitter.setStretchFactor(1, 0)  # Preview has fixed size
        step4_layout.addWidget(splitter)

        # Initially hide Step 4
        self.step4_group.setVisible(False)
        layout.addWidget(self.step4_group, 1)  # Give stretch factor to expand

        # Store reference to enable button later
        self.step4_created = True

        # Main content
        splitter = QSplitter(Qt.Horizontal)

        # Left side - Table and controls
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)

        # Simplified page offset with checkbox
        offset_group = QGroupBox(self.localization.get_text("fix_page_mismatch"))
        offset_layout = QVBoxLayout(offset_group)

        # Automatic offset tracking info
        self.offset_info = QLabel("📍 " + self.localization.get_text("navigate_to_correct_page"))
        self.offset_info.setWordWrap(True)
        self.offset_info.setStyleSheet("color: #666; font-size: 12px; margin-bottom: 10px;")
        offset_layout.addWidget(self.offset_info)

        # Simple offset control
        offset_control_layout = QHBoxLayout()
        offset_control_layout.addWidget(QLabel(self.localization.get_text("page_offset")))

        self.page_offset = QSpinBox()
        self.page_offset.setRange(-100, 100)
        self.page_offset.setValue(0)
        self.page_offset.valueChanged.connect(self.on_offset_changed)
        self.page_offset.setLayoutDirection(Qt.LeftToRight)
        self.page_offset.setStyleSheet("QSpinBox { text-align: center; }")
        offset_control_layout.addWidget(self.page_offset)

        # Auto-calculate offset button
        self.btn_auto_offset = QPushButton("🔄 " + self.localization.get_text("auto_calculate"))
        self.btn_auto_offset.clicked.connect(self.auto_calculate_offset)
        self.btn_auto_offset.setEnabled(False)
        self.btn_auto_offset.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 8px; }")
        offset_control_layout.addWidget(self.btn_auto_offset)

        # Confirm offset button
        self.btn_confirm_offset = QPushButton("✅ " + self.localization.get_text("confirm"))
        self.btn_confirm_offset.clicked.connect(self.confirm_offset)
        self.btn_confirm_offset.setEnabled(False)
        self.btn_confirm_offset.setStyleSheet("QPushButton { background-color: #2196F3; color: white; font-weight: bold; padding: 8px; }")
        offset_control_layout.addWidget(self.btn_confirm_offset)

        offset_control_layout.addStretch()
        offset_layout.addLayout(offset_control_layout)

        # Navigation guidance
        nav_guidance = QLabel(self.localization.get_text("navigate_to_correct_page"))
        nav_guidance.setStyleSheet("color: #666; font-size: 12px; font-style: italic;")
        offset_layout.addWidget(nav_guidance)

        # Initialize tracking variables for automatic offset calculation
        self.initial_preview_page = None
        self.current_preview_page = None
        self.navigation_count = 0  # Track navigation button presses
        self.navigation_direction = 0  # Track net direction (+/- pages)

    def add_header_section(self, layout):
        """Add header section with back to main and bookmark copy utility"""
        # Header container
        header_widget = QWidget()
        header_widget.setMaximumHeight(60)
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(5, 5, 5, 5)
        header_layout.setSpacing(15)

        # Back to Main button (left side)
        self.btn_back_to_main = QPushButton(f"← {self.localization.get_text('back_to_home')}")
        self.btn_back_to_main.clicked.connect(self.navigate_to_home)
        self.btn_back_to_main.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 8px 16px;
                border: none;
                border-radius: 6px;
                min-height: 35px;
                max-width: 150px;
            }
            QPushButton:hover {
                background-color: #1565C0;
            }
        """)
        self.btn_back_to_main.setToolTip(self.localization.get_text('back_to_home'))
        header_layout.addWidget(self.btn_back_to_main)

        # Flexible spacing to push buttons apart
        header_layout.addStretch()

        # Bookmark copy utility button (right side)
        self.btn_bookmark_copy = QPushButton(f"📋 {self.localization.get_text('copy_bookmarks_utility')}")
        self.btn_bookmark_copy.clicked.connect(self.open_bookmark_copy_dialog)
        self.btn_bookmark_copy.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 8px 16px;
                border: none;
                border-radius: 6px;
                min-height: 35px;
                max-width: 200px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        self.btn_bookmark_copy.setToolTip(self.localization.get_text('copy_bookmarks_tooltip'))
        header_layout.addWidget(self.btn_bookmark_copy)

        layout.addWidget(header_widget)

        # Initialize file paths for bookmark copy
        self.original_pdf_path = ""
        self.enhanced_pdf_path = ""

    def navigate_to_home(self):
        """Navigate back to home page"""
        # Get the main window and navigate to home
        main_window = self.parent()
        while main_window and not hasattr(main_window, 'navigation_manager'):
            main_window = main_window.parent()

        if main_window and hasattr(main_window, 'navigation_manager'):
            main_window.navigate_to_home()

    def add_bookmark_copy_section(self, layout):
        """Add bookmark copy feature section for OCR PDFs"""
        # Bookmark Copy Feature (for OCR PDFs)
        copy_group = QGroupBox(f"📋 {self.localization.get_text('copy_bookmarks_title')}")
        copy_group.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; color: #9C27B0; }")
        copy_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        copy_layout = QVBoxLayout(copy_group)
        copy_layout.setContentsMargins(10, 15, 10, 10)

        # Description
        desc_label = QLabel(self.localization.get_text('copy_bookmarks_desc'))
        desc_label.setWordWrap(True)
        desc_label.setStyleSheet("color: #666; font-size: 12px; margin-bottom: 10px;")
        copy_layout.addWidget(desc_label)

        # File selection layout
        files_layout = QVBoxLayout()

        # Original PDF (source)
        original_layout = QHBoxLayout()
        original_layout.addWidget(QLabel(self.localization.get_text('original_pdf')))

        self.original_pdf_display = QLineEdit()
        self.original_pdf_display.setReadOnly(True)
        self.original_pdf_display.setPlaceholderText(self.localization.get_text('select_original_pdf'))
        self.original_pdf_display.setStyleSheet("""
            QLineEdit {
                padding: 6px 10px;
                font-size: 11px;
                border: 2px solid #ddd;
                border-radius: 4px;
                background-color: #f9f9f9;
                color: #333;
            }
        """)
        original_layout.addWidget(self.original_pdf_display, 7)

        self.btn_select_original = QPushButton(self.localization.get_text("browse"))
        self.btn_select_original.clicked.connect(self.select_original_pdf)
        self.btn_select_original.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 6px 10px;
                font-size: 11px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        original_layout.addWidget(self.btn_select_original, 1)
        files_layout.addLayout(original_layout)

        # Enhanced PDF (target)
        enhanced_layout = QHBoxLayout()
        enhanced_layout.addWidget(QLabel(self.localization.get_text('enhanced_pdf')))

        self.enhanced_pdf_display = QLineEdit()
        self.enhanced_pdf_display.setReadOnly(True)
        self.enhanced_pdf_display.setPlaceholderText(self.localization.get_text('select_enhanced_pdf'))
        self.enhanced_pdf_display.setStyleSheet("""
            QLineEdit {
                padding: 6px 10px;
                font-size: 11px;
                border: 2px solid #ddd;
                border-radius: 4px;
                background-color: #f9f9f9;
                color: #333;
            }
        """)
        enhanced_layout.addWidget(self.enhanced_pdf_display, 7)

        self.btn_select_enhanced = QPushButton(self.localization.get_text("browse"))
        self.btn_select_enhanced.clicked.connect(self.select_enhanced_pdf)
        self.btn_select_enhanced.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 6px 10px;
                font-size: 11px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        enhanced_layout.addWidget(self.btn_select_enhanced, 1)
        files_layout.addLayout(enhanced_layout)

        copy_layout.addLayout(files_layout)

        # Copy bookmarks button
        self.btn_copy_bookmarks = QPushButton(f"📋 {self.localization.get_text('copy_bookmarks_button')}")
        self.btn_copy_bookmarks.clicked.connect(self.copy_bookmarks_between_pdfs)
        self.btn_copy_bookmarks.setEnabled(False)
        self.btn_copy_bookmarks.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 10px;
                border-radius: 6px;
                min-height: 35px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666;
            }
        """)
        copy_layout.addWidget(self.btn_copy_bookmarks)

        # Set fixed height for the copy section
        copy_group.setMaximumHeight(200)
        copy_group.setMinimumHeight(200)

        layout.addWidget(copy_group)

        # Initialize file paths
        self.original_pdf_path = ""
        self.enhanced_pdf_path = ""

    def open_bookmark_preparation_dialog(self):
        """Open the first dialog for page extraction and preview"""
        if not hasattr(self, 'pdf_path') or not self.pdf_path or not hasattr(self, 'pdf_doc') or not self.pdf_doc:
            QMessageBox.warning(self, "تحذير", "يرجى تحميل ملف PDF أولاً")
            return

        try:
            # Create the page extraction dialog
            self.open_page_extraction_dialog()
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "خطأ", f"حدث خطأ في فتح نافذة التحضير:\n{str(e)}")

    def open_page_extraction_dialog(self):
        """First dialog: Extract pages with preview and range selection"""
        dialog = QDialog(self)
        dialog.setWindowTitle("استخراج صفحات الفهرس")
        dialog.setModal(True)
        dialog.resize(900, 900)  # Increased height to prevent PDF preview clipping

        layout = QVBoxLayout(dialog)

        # Title
        title_label = QLabel("📄 استخراج صفحات فهرس المحتويات")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #1976D2; margin-bottom: 10px;")
        layout.addWidget(title_label)

        # Info message
        info_label = QLabel(f"ملف PDF محمل: {self.pdf_doc.page_count} صفحة")
        info_label.setStyleSheet("color: #666; margin-bottom: 15px;")
        layout.addWidget(info_label)

        # Page range selection
        range_group = QGroupBox("تحديد نطاق صفحات الفهرس")
        range_group.setStyleSheet("QGroupBox { font-weight: bold; padding: 10px; }")
        range_layout = QHBoxLayout(range_group)

        range_layout.addWidget(QLabel("من صفحة:"))
        start_spin = QSpinBox()
        start_spin.setMinimum(1)
        start_spin.setMaximum(self.pdf_doc.page_count if hasattr(self, 'pdf_doc') and self.pdf_doc else 9999)
        # Auto-detect: start from last 10 pages
        auto_start = max(1, self.pdf_doc.page_count - 9) if hasattr(self, 'pdf_doc') and self.pdf_doc else 1
        start_spin.setValue(auto_start)
        start_spin.setStyleSheet("QSpinBox { padding: 5px; font-size: 12px; }")
        range_layout.addWidget(start_spin)

        range_layout.addWidget(QLabel("إلى صفحة:"))
        end_spin = QSpinBox()
        end_spin.setMinimum(1)
        end_spin.setMaximum(self.pdf_doc.page_count if hasattr(self, 'pdf_doc') and self.pdf_doc else 9999)
        # Auto-detect: end at last page
        auto_end = self.pdf_doc.page_count if hasattr(self, 'pdf_doc') and self.pdf_doc else 5
        end_spin.setValue(auto_end)
        end_spin.setStyleSheet("QSpinBox { padding: 5px; font-size: 12px; }")
        range_layout.addWidget(end_spin)

        layout.addWidget(range_group)

        # PDF Preview area with full functionality
        preview_group = QGroupBox("معاينة الصفحات")
        preview_group.setStyleSheet("QGroupBox { font-weight: bold; padding: 10px; }")
        preview_layout = QVBoxLayout(preview_group)

        # Navigation controls
        nav_layout = QHBoxLayout()

        btn_first = QPushButton("⏮")
        btn_first.setToolTip("الصفحة الأولى")
        btn_first.setStyleSheet("QPushButton { padding: 5px 10px; font-size: 12px; }")
        nav_layout.addWidget(btn_first)

        btn_prev = QPushButton("◀")
        btn_prev.setToolTip("الصفحة السابقة")
        btn_prev.setStyleSheet("QPushButton { padding: 5px 10px; font-size: 12px; }")
        nav_layout.addWidget(btn_prev)

        # Page input
        page_input = QSpinBox()
        page_input.setRange(1, self.pdf_doc.page_count if hasattr(self, 'pdf_doc') and self.pdf_doc else 9999)
        page_input.setValue(auto_start)  # Start at auto-detected page
        page_input.setStyleSheet("QSpinBox { text-align: center; min-width: 60px; padding: 5px; }")
        nav_layout.addWidget(page_input)

        page_label = QLabel(f"/ {self.pdf_doc.page_count if hasattr(self, 'pdf_doc') and self.pdf_doc else '?'}")
        page_label.setAlignment(Qt.AlignCenter)
        nav_layout.addWidget(page_label)

        btn_next = QPushButton("▶")
        btn_next.setToolTip("الصفحة التالية")
        btn_next.setStyleSheet("QPushButton { padding: 5px 10px; font-size: 12px; }")
        nav_layout.addWidget(btn_next)

        btn_last = QPushButton("⏭")
        btn_last.setToolTip("الصفحة الأخيرة")
        btn_last.setStyleSheet("QPushButton { padding: 5px 10px; font-size: 12px; }")
        nav_layout.addWidget(btn_last)

        preview_layout.addLayout(nav_layout)

        # Set Start/End buttons
        set_buttons_layout = QHBoxLayout()
        set_buttons_layout.addStretch()  # Center the buttons

        btn_set_start = QPushButton("📍 تعيين كبداية")
        btn_set_start.setToolTip("تعيين الصفحة الحالية كصفحة البداية")
        btn_set_start.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                margin: 2px;
            }
            QPushButton:hover { background-color: #1565C0; }
        """)
        set_buttons_layout.addWidget(btn_set_start)

        btn_set_end = QPushButton("🏁 تعيين كنهاية")
        btn_set_end.setToolTip("تعيين الصفحة الحالية كصفحة النهاية")
        btn_set_end.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                margin: 2px;
            }
            QPushButton:hover { background-color: #45a049; }
        """)
        set_buttons_layout.addWidget(btn_set_end)

        set_buttons_layout.addStretch()  # Center the buttons
        preview_layout.addLayout(set_buttons_layout)

        # PDF preview display
        pdf_preview = QLabel()
        pdf_preview.setAlignment(Qt.AlignCenter)
        pdf_preview.setMinimumSize(400, 300)
        pdf_preview.setMaximumSize(600, 500)  # Set maximum size to prevent excessive stretching
        pdf_preview.setStyleSheet("border: 2px solid #ddd; background-color: white; border-radius: 6px;")
        # Remove setScaledContents(True) to preserve aspect ratio
        # The pixmap will be scaled properly in the update function
        preview_layout.addWidget(pdf_preview)

        # Page info
        page_info = QLabel("اضغط 'معاينة الصفحات' لرؤية الصفحات المحددة")
        page_info.setAlignment(Qt.AlignCenter)
        page_info.setStyleSheet("color: #666; font-size: 11px; padding: 5px; font-style: italic;")
        preview_layout.addWidget(page_info)

        layout.addWidget(preview_group)

        # Initialize preview variables
        current_preview_page = auto_start - 1  # 0-based index
        preview_start_page = auto_start
        preview_end_page = auto_end

        # Buttons layout
        buttons_layout = QHBoxLayout()

        # Preview button
        preview_btn = QPushButton("👁 معاينة الصفحات")
        preview_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                font-size: 14px;
                border: none;
                border-radius: 6px;
                margin: 5px;
            }
            QPushButton:hover { background-color: #1976D2; }
        """)
        buttons_layout.addWidget(preview_btn)

        # Extract button
        extract_btn = QPushButton("📄 استخراج الصفحات وفتح المجلد")
        extract_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 12px 24px;
                font-size: 14px;
                border: none;
                border-radius: 6px;
                margin: 5px;
            }
            QPushButton:hover { background-color: #45a049; }
        """)
        buttons_layout.addWidget(extract_btn)

        layout.addLayout(buttons_layout)

        # Instructions
        instructions = QLabel("التعليمات:\n1. حدد نطاق صفحات فهرس المحتويات\n2. اضغط 'معاينة الصفحات' للتأكد من الصفحات الصحيحة\n3. اضغط 'استخراج الصفحات' لحفظ الصفحات وفتح المجلد")
        instructions.setStyleSheet("color: #666; font-size: 11px; padding: 15px; background-color: #f0f8ff; border-radius: 6px; border-left: 4px solid #2196F3;")
        instructions.setWordWrap(True)
        layout.addWidget(instructions)

        # Preview functionality methods
        def update_pdf_preview():
            """Update PDF preview with current page"""
            try:
                if not hasattr(self, 'pdf_doc') or not self.pdf_doc:
                    return

                page = self.pdf_doc[current_preview_page]

                # Get page dimensions
                page_rect = page.rect

                # Use fixed preview dimensions for consistent scaling
                max_preview_width = 580  # Maximum width for preview
                max_preview_height = 480  # Maximum height for preview

                # Calculate scale to fit within preview area while maintaining aspect ratio
                scale_x = max_preview_width / page_rect.width if page_rect.width > 0 else 1
                scale_y = max_preview_height / page_rect.height if page_rect.height > 0 else 1
                scale = min(scale_x, scale_y, 2.0)  # Allow up to 2x zoom for better quality

                # Render page with calculated scale
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")

                pixmap = QPixmap()
                pixmap.loadFromData(img_data)

                if not pixmap.isNull():
                    # Scale pixmap to fit preview area while maintaining aspect ratio
                    scaled_pixmap = pixmap.scaled(
                        max_preview_width, max_preview_height,
                        Qt.KeepAspectRatio,  # Maintain aspect ratio
                        Qt.SmoothTransformation  # High quality scaling
                    )
                    pdf_preview.setPixmap(scaled_pixmap)

                # Update page info
                page_num = current_preview_page + 1
                page_status = ""
                if page_num == preview_start_page:
                    page_status = " (بداية النطاق)"
                elif page_num == preview_end_page:
                    page_status = " (نهاية النطاق)"
                elif preview_start_page <= page_num <= preview_end_page:
                    page_status = " (ضمن النطاق المحدد)"

                page_info.setText(f"الصفحة {page_num}{page_status}")
                page_input.setValue(page_num)

                # Update navigation buttons
                btn_first.setEnabled(current_preview_page > 0)
                btn_prev.setEnabled(current_preview_page > 0)
                btn_next.setEnabled(current_preview_page < self.pdf_doc.page_count - 1)
                btn_last.setEnabled(current_preview_page < self.pdf_doc.page_count - 1)

            except Exception as e:
                pdf_preview.setText(f"خطأ في المعاينة: {str(e)}")
                page_info.setText("خطأ في تحميل الصفحة")

        def first_page():
            nonlocal current_preview_page
            current_preview_page = 0
            update_pdf_preview()

        def prev_page():
            nonlocal current_preview_page
            if current_preview_page > 0:
                current_preview_page -= 1
                update_pdf_preview()

        def next_page():
            nonlocal current_preview_page
            if current_preview_page < self.pdf_doc.page_count - 1:
                current_preview_page += 1
                update_pdf_preview()

        def last_page():
            nonlocal current_preview_page
            current_preview_page = self.pdf_doc.page_count - 1
            update_pdf_preview()

        def goto_page(page_num):
            nonlocal current_preview_page
            if 1 <= page_num <= self.pdf_doc.page_count:
                current_preview_page = page_num - 1
                update_pdf_preview()

        # Connect navigation buttons to local functions
        btn_first.clicked.connect(first_page)
        btn_prev.clicked.connect(prev_page)
        btn_next.clicked.connect(next_page)
        btn_last.clicked.connect(last_page)
        page_input.valueChanged.connect(goto_page)

        # Set Start/End button functions
        def set_as_start():
            """Set current preview page as start page"""
            current_page_num = current_preview_page + 1
            start_spin.setValue(current_page_num)
            print(f"DEBUG: Set start page to {current_page_num}")

        def set_as_end():
            """Set current preview page as end page"""
            current_page_num = current_preview_page + 1
            end_spin.setValue(current_page_num)
            print(f"DEBUG: Set end page to {current_page_num}")

        # Connect set start/end buttons
        btn_set_start.clicked.connect(set_as_start)
        btn_set_end.clicked.connect(set_as_end)

        # Connect buttons
        def preview_pages():
            nonlocal current_preview_page, preview_start_page, preview_end_page
            start_page = start_spin.value()
            end_page = end_spin.value()
            if start_page > end_page:
                QMessageBox.warning(dialog, "خطأ", "رقم الصفحة الأولى يجب أن يكون أقل من أو يساوي رقم الصفحة الأخيرة")
                return

            # Update preview range
            preview_start_page = start_page
            preview_end_page = end_page

            # Go to start page of selected range
            current_preview_page = start_page - 1
            update_pdf_preview()

            # Update range info
            range_info = f"النطاق المحدد: من الصفحة {start_page} إلى {end_page}\n"
            range_info += f"إجمالي الصفحات: {end_page - start_page + 1} صفحة"
            page_info.setText(range_info)
            page_info.setStyleSheet("color: #2196F3; font-weight: bold; text-align: center; padding: 10px;")

        def extract_pages():
            start_page = start_spin.value()
            end_page = end_spin.value()
            if start_page > end_page:
                QMessageBox.warning(dialog, "خطأ", "رقم الصفحة الأولى يجب أن يكون أقل من أو يساوي رقم الصفحة الأخيرة")
                return

            try:
                # Extract pages using existing logic
                import os
                import subprocess

                # Create output filename
                base_name = os.path.splitext(os.path.basename(self.pdf_path))[0]
                output_path = os.path.join(os.path.dirname(self.pdf_path), f"{base_name}_TOC_pages_{start_page}-{end_page}.pdf")

                # Extract pages
                output_doc = fitz.open()
                for page_num in range(start_page - 1, end_page):
                    if page_num < self.pdf_doc.page_count:
                        output_doc.insert_pdf(self.pdf_doc, from_page=page_num, to_page=page_num)

                output_doc.save(output_path)
                output_doc.close()

                # Show success message
                QMessageBox.information(dialog, "نجح الاستخراج", f"تم استخراج الصفحات بنجاح:\n{output_path}")

                # Open folder containing the extracted file
                try:
                    # Ensure we have absolute path
                    abs_output_path = os.path.abspath(output_path)

                    if os.name == 'nt':  # Windows
                        # Convert forward slashes to backslashes for Windows
                        windows_path = abs_output_path.replace('/', '\\')
                        subprocess.run(['explorer', '/select,', windows_path], check=False)
                    elif os.name == 'posix':  # macOS and Linux
                        if sys.platform == 'darwin':  # macOS
                            subprocess.run(['open', '-R', abs_output_path], check=False)
                        else:  # Linux
                            subprocess.run(['xdg-open', os.path.dirname(abs_output_path)], check=False)
                except Exception as folder_error:
                    print(f"Error opening folder: {folder_error}")
                    # Fallback: try to open just the directory
                    try:
                        import webbrowser
                        webbrowser.open(os.path.dirname(abs_output_path))
                    except Exception as fallback_error:
                        print(f"Fallback folder opening also failed: {fallback_error}")

                # Close this dialog and open the AI Studio dialog
                dialog.accept()
                self.open_ai_studio_dialog()

            except Exception as e:
                QMessageBox.critical(dialog, "خطأ", f"فشل في استخراج الصفحات:\n{str(e)}")

        # Connect buttons
        preview_btn.clicked.connect(preview_pages)
        extract_btn.clicked.connect(extract_pages)

        # Auto-detect TOC pages and initialize preview
        def auto_detect_toc_pages():
            """Automatically detect potential TOC pages (usually at the end)"""
            if not hasattr(self, 'pdf_doc') or not self.pdf_doc:
                return

            # Check last 15 pages for TOC content
            total_pages = self.pdf_doc.page_count
            search_start = max(0, total_pages - 15)

            toc_keywords = [
                'فهرس', 'محتويات', 'contents', 'index', 'table of contents',
                'الفهرس', 'فهرس المحتويات', 'قائمة المحتويات'
            ]

            potential_toc_pages = []

            for page_num in range(search_start, total_pages):
                try:
                    page = self.pdf_doc[page_num]
                    text = page.get_text().lower()

                    # Check for TOC keywords
                    for keyword in toc_keywords:
                        if keyword in text:
                            potential_toc_pages.append(page_num + 1)  # Convert to 1-based
                            break

                    # Check for typical TOC patterns (dots and page numbers)
                    if '...' in text or '…' in text:
                        # Look for page number patterns
                        import re
                        if re.search(r'\d+\s*$', text, re.MULTILINE):
                            potential_toc_pages.append(page_num + 1)

                except Exception as e:
                    continue

            if potential_toc_pages:
                # Use the range from first detected to last page
                detected_start = min(potential_toc_pages)
                detected_end = total_pages

                # Update the spinboxes with detected range
                start_spin.setValue(detected_start)
                end_spin.setValue(detected_end)

                # Update preview variables
                nonlocal preview_start_page, preview_end_page, current_preview_page
                preview_start_page = detected_start
                preview_end_page = detected_end
                current_preview_page = detected_start - 1

                # Show detection message
                page_info.setText(f"🔍 تم اكتشاف صفحات الفهرس تلقائياً: {detected_start}-{detected_end}")
                page_info.setStyleSheet("color: #4CAF50; font-weight: bold; text-align: center; padding: 10px;")
            else:
                # No TOC detected, show default message
                page_info.setText("اضغط 'معاينة الصفحات' لرؤية الصفحات المحددة")
                page_info.setStyleSheet("color: #666; font-size: 11px; padding: 5px; font-style: italic;")

        # Run auto-detection
        auto_detect_toc_pages()

        # Initialize preview with auto-detected or default range
        update_pdf_preview()

        # Execute dialog
        dialog.exec_()

    def open_ai_studio_dialog(self):
        """Second dialog: AI Studio workflow with copy prompt and open AI Studio"""
        dialog = QDialog(self)
        dialog.setWindowTitle("استخدام Google AI Studio")
        dialog.setModal(True)
        dialog.resize(600, 400)

        layout = QVBoxLayout(dialog)
        layout.setSpacing(20)

        # Title
        title_label = QLabel("🤖 استخدام Google AI Studio لاستخراج الإشارات المرجعية")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #1976D2; margin-bottom: 10px;")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)

        # Instructions
        instructions = QLabel("""تم استخراج صفحات الفهرس بنجاح! الآن اتبع الخطوات التالية:

1. انسخ النص التوجيهي أدناه
2. افتح Google AI Studio
3. الصق النص التوجيهي في AI Studio
4. ارفع ملف PDF المستخرج إلى AI Studio
5. انسخ النتيجة والصقها في الخطوة 2 من التطبيق الرئيسي""")
        instructions.setStyleSheet("""
            color: #333;
            font-size: 12px;
            padding: 20px;
            background-color: #f0f8ff;
            border-radius: 8px;
            border-left: 4px solid #2196F3;
            line-height: 1.5;
        """)
        instructions.setWordWrap(True)
        layout.addWidget(instructions)

        # Copy prompt button (full width)
        copy_btn = QPushButton("📋 نسخ النص التوجيهي")
        copy_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                font-weight: bold;
                padding: 15px;
                font-size: 14px;
                border: none;
                border-radius: 8px;
                min-height: 20px;
            }
            QPushButton:hover {
                background-color: #F57C00;
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                background-color: #E65100;
            }
        """)
        layout.addWidget(copy_btn)

        # Open AI Studio button (full width)
        studio_btn = QPushButton("🌐 فتح Google AI Studio")
        studio_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 15px;
                font-size: 14px;
                border: none;
                border-radius: 8px;
                min-height: 20px;
            }
            QPushButton:hover {
                background-color: #45a049;
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
        """)
        layout.addWidget(studio_btn)

        # Footer note
        footer_note = QLabel("💡 نصيحة: بعد الانتهاء من AI Studio، ارجع إلى التطبيق الرئيسي والصق النتيجة في الخطوة 2")
        footer_note.setStyleSheet("color: #666; font-size: 11px; font-style: italic; text-align: center; margin-top: 10px;")
        footer_note.setAlignment(Qt.AlignCenter)
        footer_note.setWordWrap(True)
        layout.addWidget(footer_note)

        # Connect buttons
        def copy_prompt():
            # Use the existing AI prompt
            prompt_text = self.localization.get_text("ai_toc_prompt")
            QApplication.clipboard().setText(prompt_text)
            QMessageBox.information(dialog, "تم النسخ", "تم نسخ النص التوجيهي إلى الحافظة!")

        def open_ai_studio():
            import webbrowser
            webbrowser.open("https://aistudio.google.com/")
            # Close this dialog after opening AI Studio
            dialog.accept()

        copy_btn.clicked.connect(copy_prompt)
        studio_btn.clicked.connect(open_ai_studio)

        dialog.exec_()




    def update_ui_for_method_selection(self):
        """Update UI elements based on the selected method"""
        if self.selected_method == "toc_pages":
            # Update step titles and descriptions for TOC method
            pass
        elif self.selected_method == "text_file":
            # Update step titles and descriptions for text file method
            pass

    def enable_method_selection(self):
        """Enable bookmark preparation (method selection removed)"""
        # Method selection removed - bookmark preparation is now via dialog
        pass

    def format_bookmarks(self):
        """Format bookmarks based on the selected method"""
        try:
            if not hasattr(self, 'selected_method') or not self.selected_method:
                QMessageBox.warning(self, "Warning", "Please select a bookmark method first.")
                return

            if self.selected_method == "toc_pages":
                # For TOC pages method, format the pasted text
                if hasattr(self, 'toc_text') and self.toc_text and self.toc_text.toPlainText().strip():
                    # Enable the load bookmarks button
                    if hasattr(self, 'btn_load_bookmarks') and self.btn_load_bookmarks:
                        self.btn_load_bookmarks.setEnabled(True)
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        self.localization.get_text("toc_text_ready")
                    )
                else:
                    QMessageBox.warning(self, "Warning", self.localization.get_text("no_toc_text"))
            elif self.selected_method == "text_file":
                # For text file method, check if file is selected or text is pasted
                has_file = hasattr(self, 'bookmark_file_path') and self.bookmark_file_path
                has_text = hasattr(self, 'toc_text') and self.toc_text and self.toc_text.toPlainText().strip()

                if has_file or has_text:
                    # Enable the load bookmarks button
                    if hasattr(self, 'btn_load_bookmarks') and self.btn_load_bookmarks:
                        self.btn_load_bookmarks.setEnabled(True)

                    # FIX: Don't change UI layout on Format button - only enable Load Bookmarks
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('bookmark_source_ready')}\n\n{self.localization.get_text('bookmark_source_ready_desc')}"
                    )
                else:
                    QMessageBox.warning(self, "Warning", self.localization.get_text("no_bookmark_source"))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error formatting bookmarks: {str(e)}")

    def enable_bookmark_controls(self):
        """Enable bookmark-related controls after bookmarks are loaded"""
        print("ENABLE_BOOKMARK_CONTROLS CALLED!")  # Debug
        print(f"Current bookmarks count: {len(self.bookmarks) if hasattr(self, 'bookmarks') and self.bookmarks else 0}")  # Debug

        if hasattr(self, 'btn_insert') and self.btn_insert:
            print(f"Found insert button: {self.btn_insert}")  # Debug
            print(f"Button before enabling: {self.btn_insert.isEnabled()}")  # Debug

            # Force enable the button
            self.btn_insert.setEnabled(True)
            self.btn_insert.setVisible(True)

            print(f"Button after enabling: {self.btn_insert.isEnabled()}")  # Debug
            print(f"Button visible: {self.btn_insert.isVisible()}")  # Debug

            # Also update button style to show it's active
            self.btn_insert.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    font-weight: bold;
                    font-size: 16px;
                    padding: 15px;
                    border-radius: 8px;
                    margin: 10px 0;
                    border: 2px solid #45a049;
                }
                QPushButton:hover {
                    background-color: #45a049;
                    border: 2px solid #4CAF50;
                }
                QPushButton:pressed {
                    background-color: #3d8b40;
                }
            """)
        else:
            print("ERROR: btn_insert not found or is None!")  # Debug

        if hasattr(self, 'btn_prev') and self.btn_prev:
            self.btn_prev.setEnabled(True)
        if hasattr(self, 'btn_next') and self.btn_next:
            self.btn_next.setEnabled(True)

    def force_enable_insert_button(self):
        """Force enable the insert button - called after bookmarks are loaded"""
        print("FORCE_ENABLE_INSERT_BUTTON CALLED!")  # Debug

        if hasattr(self, 'btn_insert') and self.btn_insert:
            print(f"Forcing insert button enable...")  # Debug

            # Force enable
            self.btn_insert.setEnabled(True)
            self.btn_insert.setVisible(True)
            self.btn_insert.show()

            # Update the button text to show it's ready
            self.btn_insert.setText("✅ إدراج الإشارات المرجعية في الملف")

            # Apply active styling
            self.btn_insert.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50 !important;
                    color: white !important;
                    font-weight: bold !important;
                    font-size: 16px !important;
                    padding: 15px !important;
                    border-radius: 8px !important;
                    margin: 10px 0 !important;
                    border: 3px solid #45a049 !important;
                }
                QPushButton:hover {
                    background-color: #45a049 !important;
                    border: 3px solid #4CAF50 !important;
                    transform: scale(1.02);
                }
                QPushButton:pressed {
                    background-color: #3d8b40 !important;
                }
            """)

            print(f"Insert button final state - Enabled: {self.btn_insert.isEnabled()}, Visible: {self.btn_insert.isVisible()}")  # Debug
        else:
            print("ERROR: btn_insert not found in force_enable_insert_button!")  # Debug

    # REMOVED: test_button_connection method (no longer needed)

    def on_bookmark_clicked(self, row, column):
        """Handle bookmark table cell clicks - automatically show preview for any cell"""
        self.selected_bookmark_row = row
        if row < len(self.bookmarks):
            self.selected_bookmark = self.bookmarks[row]
            # Automatically show preview when any cell is clicked
            self.preview_bookmark(row)

    def preview_bookmark(self, row):
        """Preview a specific bookmark using current modified page from table"""
        if row >= len(self.bookmarks) or not self.pdf_doc:
            return

        bookmark = self.bookmarks[row]

        # CRITICAL FIX: Get the current modified page from the table, not recalculate
        modified_page = bookmark.page  # Default fallback
        if hasattr(self, 'bookmark_table') and self.bookmark_table:
            try:
                # Read the current "Modified Page" value from column 2 of the table
                mod_item = self.bookmark_table.item(row, 2)
                if mod_item:
                    modified_page = int(mod_item.text())
                    print(f"DEBUG: Using modified page {modified_page} from table for bookmark '{bookmark.title}'")
                else:
                    print(f"DEBUG: No modified page item found, using original page {bookmark.page}")
            except (ValueError, RuntimeError) as e:
                print(f"DEBUG: Error reading modified page from table: {e}, using original page {bookmark.page}")

        target_page = modified_page - 1  # Convert to 0-based for PDF indexing
        if 0 <= target_page < self.total_pages:
            self.current_page = target_page
            self.selected_bookmark = bookmark
            self.selected_bookmark_row = row
            self.update_preview()

            # Track this bookmark as verified
            self.verified_bookmarks.add(row)

            # Show guidance message for verification
            verified_count = len(self.verified_bookmarks)
            total_bookmarks = len(self.bookmarks)

            # REMOVED: Guidance message - automatic navigation now handles adjustments
            # Just show a brief confirmation that the bookmark was selected
            print(f"DEBUG: Previewing bookmark '{bookmark.title}' at modified page {modified_page}")

            # Highlight the row in the table
            if hasattr(self, 'bookmark_table') and self.bookmark_table:
                self.bookmark_table.selectRow(row)

    def prev_page(self):
        """Go to previous page"""
        if self.pdf_doc and self.current_page > 0:
            self.current_page -= 1
            self.update_preview()
            self.update_page_info()

    def next_page(self):
        """Go to next page"""
        if self.pdf_doc and self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_preview()
            self.update_page_info()

    def update_page_info(self):
        """Update page information display"""
        if hasattr(self, 'page_label') and self.page_label:
            try:
                current_display = self.current_page + 1
                self.page_label.setText(f"الصفحة {current_display} من {self.total_pages}")

                # Check if current page matches any bookmark
                offset = getattr(self, 'page_offset_value', 0)
                if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
                    try:
                        offset = self.offset_spinbox.value()
                    except RuntimeError:
                        offset = getattr(self, 'page_offset_value', 0)

                for i, bookmark in enumerate(self.bookmarks):
                    if bookmark.page + offset == current_display:
                        self.page_label.setText(
                            f"الصفحة {current_display} من {self.total_pages}\n"
                            f"📖 {bookmark.title}"
                        )
                        # Highlight this bookmark in the table
                        if hasattr(self, 'bookmark_table') and self.bookmark_table:
                            try:
                                self.bookmark_table.selectRow(i)
                            except RuntimeError:
                                pass
                        break
            except RuntimeError:
                # Widget has been deleted, ignore
                pass

    def on_level_changed(self, row, level):
        """Handle bookmark level changes"""
        if row < len(self.bookmarks):
            self.bookmarks[row].level = level

    def on_offset_changed(self):
        """Handle page offset changes and update table immediately"""
        if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
            try:
                offset = self.offset_spinbox.value()
                print(f"Offset changed to: {offset}")  # Debug
                self.update_table()
            except RuntimeError:
                pass

    def fix_bookmark_pages(self):
        """Fix bookmark pages with current offset and prompt for verification"""
        if not hasattr(self, 'offset_spinbox') or not self.offset_spinbox:
            return

        offset = self.offset_spinbox.value()

        # Show confirmation with clear explanation
        reply = QMessageBox.question(
            self,
            "🔧 إصلاح الصفحات",
            f"سيتم تطبيق إزاحة {offset:+d} على جميع الإشارات المرجعية.\n\n"
            f"هذا يعني أن كل صفحة ستتغير بمقدار {offset:+d}.\n\n"
            f"هل تريد المتابعة؟",
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            self.page_offset_value = offset
            self.update_table()

            # Show success message and prompt for verification
            verified_count = len(self.verified_bookmarks)
            total_bookmarks = len(self.bookmarks)

            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setWindowTitle("✅ تم الإصلاح")

            if verified_count < 2:
                msg_box.setText(
                    f"تم تطبيق الإزاحة {offset:+d} بنجاح!\n\n"
                    f"📊 تم التحقق من {verified_count} من {total_bookmarks} إشارة مرجعية\n\n"
                    f"🔍 يُنصح بشدة بالتحقق من 1-2 إشارات مرجعية إضافية للتأكد من صحة التعديل.\n\n"
                    f"استخدم أيقونة العين 👁 لمعاينة إشارات مرجعية أخرى."
                )
            else:
                msg_box.setText(
                    f"تم تطبيق الإزاحة {offset:+d} بنجاح!\n\n"
                    f"✅ تم التحقق من {verified_count} إشارة مرجعية - ممتاز!\n\n"
                    f"يمكنك الآن المتابعة لإدراج الإشارات المرجعية في الملف."
                )

            msg_box.setStandardButtons(QMessageBox.Ok)
            msg_box.exec()

    def confirm_offset_adjustment(self):
        """Confirm the offset adjustment with user"""
        if not hasattr(self, 'offset_spinbox') or not self.offset_spinbox:
            return

        offset = self.offset_spinbox.value()
        if offset == 0:
            QMessageBox.information(self, "تأكيد", "لا توجد إزاحة مطبقة.")
            return

        reply = QMessageBox.question(
            self,
            "تأكيد الإزاحة",
            f"هل أنت متأكد من تطبيق إزاحة {offset:+d} على جميع الإشارات المرجعية؟\n\n"
            f"سيتم تحديث جميع أرقام الصفحات وفقاً لهذه الإزاحة.",
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            self.page_offset_value = offset
            self.update_table()
            QMessageBox.information(
                self,
                "تم التأكيد",
                f"تم تطبيق الإزاحة {offset:+d} بنجاح!\n\n"
                f"يمكنك الآن معاينة النتائج والمتابعة لإدراج الإشارات المرجعية."
            )

    def delete_bookmark(self, row):
        """Delete a bookmark"""
        if row < len(self.bookmarks):
            bookmark_title = self.bookmarks[row].title
            reply = QMessageBox.question(
                self,
                "Delete Bookmark",
                f"Are you sure you want to delete bookmark:\n'{bookmark_title}'?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                del self.bookmarks[row]
                self.update_table()

    def hide_previous_steps(self):
        """Hide Steps 1, 2, and 3 when Step 4 is opened to maximize workspace"""
        print("DEBUG: Hiding all previous steps for Step 4 full-page view")

        # Hide the entire top section (Steps 1 and 2)
        if hasattr(self, 'top_section_widget'):
            self.top_section_widget.setVisible(False)
            print("DEBUG: Hidden top section (Steps 1 and 2)")

        # Hide Step 2 (text input areas)
        if hasattr(self, 'step2a_group'):
            self.step2a_group.setVisible(False)
            print("DEBUG: Hidden Step 2A (TOC pages)")
        if hasattr(self, 'step2b_group'):
            self.step2b_group.setVisible(False)
            print("DEBUG: Hidden Step 2B (text paste)")

        # Hide Step 3 (format buttons) - but keep Step 4 visible
        if hasattr(self, 'step3_group'):
            self.step3_group.setVisible(False)
            print("DEBUG: Hidden Step 3 (format buttons)")

        # Ensure Step 4 remains visible (this is the step we want to show)
        if hasattr(self, 'step4_group'):
            self.step4_group.setVisible(True)
            print("DEBUG: Ensured Step 4 is visible")

        # Optimize text area size for Step 4 view
        if hasattr(self, 'toc_text') and self.toc_text:
            # Set reasonable constraints that allow text to be readable but don't take too much space
            self.toc_text.setMaximumHeight(150)  # Increased from 60 to 150 for better readability
            self.toc_text.setMinimumHeight(100)  # Increased from 60 to 100 for better usability
            # Ensure text wrapping is enabled
            self.toc_text.setLineWrapMode(QTextEdit.WidgetWidth)
            print("DEBUG: Optimized text area size for Step 4")

        # Make Step 4 expand to fill available space
        if hasattr(self, 'step4_group'):
            self.step4_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            print("DEBUG: Set Step 4 to expand and fill available space")

        # Force layout update
        self.layout().update()
        self.updateGeometry()

        print("DEBUG: All previous steps hidden, Step 4 now takes full page")

    def show_previous_steps(self):
        """Show Steps 1-3 again (for when user wants to modify settings)"""
        # Show the entire top section (Steps 1 and 2)
        if hasattr(self, 'top_section_widget'):
            self.top_section_widget.setVisible(True)
            print("DEBUG: Shown top section (Steps 1 and 2)")

        # Show appropriate step based on selected method
        if hasattr(self, 'selected_method') and self.selected_method:
            if self.selected_method == "toc_pages" and hasattr(self, 'step3a_group'):
                self.step3a_group.setVisible(True)
            elif self.selected_method == "text_file" and hasattr(self, 'step3b_group'):
                self.step3b_group.setVisible(True)

        if hasattr(self, 'step4_group'):
            self.step4_group.setVisible(True)

        # Restore text area size - remove height constraints to allow full expansion
        if hasattr(self, 'toc_text') and self.toc_text:
            self.toc_text.setMaximumHeight(16777215)  # Remove height constraint (Qt max)
            self.toc_text.setMinimumHeight(0)  # Remove minimum height constraint

        # Reset Step 4 size policy to normal
        if hasattr(self, 'step4_group'):
            self.step4_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
            print("DEBUG: Reset Step 4 size policy to normal")

    def reset_step4_workflow(self):
        """Reset Step 4 bookmark correction workflow"""
        if self.localization.current_language == "ar":
            title = "إعادة تعيين الخطوة 4"
            message = "هل تريد إعادة تعيين عملية تصحيح الإشارات المرجعية والعودة إلى الخطوة 3؟\n\nسيتم الاحتفاظ بالملف والإشارات المرجعية المحملة."
            yes_text = "نعم، إعادة تعيين"
            no_text = "إلغاء"
        else:
            title = "Reset Step 4"
            message = "Do you want to reset the bookmark correction process and return to Step 3?\n\nThe loaded file and bookmarks will be preserved."
            yes_text = "Yes, Reset"
            no_text = "Cancel"

        reply = QMessageBox.question(
            self, title, message,
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            # Hide Step 4
            if hasattr(self, 'step4_group'):
                self.step4_group.setVisible(False)

            # Show previous steps
            self.show_previous_steps()

            # Reset navigation tracking
            self.navigation_count = 0
            self.navigation_direction = 0
            self.page_offset_value = 0

            # Clear bookmarks table and data
            self.bookmarks = []
            if hasattr(self, 'bookmark_table'):
                self.bookmark_table.setRowCount(0)
                self.bookmark_table.clearContents()

            # Clear verified bookmarks
            if hasattr(self, 'verified_bookmarks'):
                self.verified_bookmarks.clear()

            # Reset offset controls
            if hasattr(self, 'page_offset') and self.page_offset:
                self.page_offset.setValue(0)
            if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
                self.offset_spinbox.setValue(0)

            # Reset PDF preview to first page if available
            if hasattr(self, 'pdf_preview') and self.pdf_preview and hasattr(self, 'total_pages'):
                self.current_page = 0
                self.update_preview()

            # Disable insert button
            if hasattr(self, 'btn_insert'):
                self.btn_insert.setEnabled(False)

            # Show success message
            if self.localization.current_language == "ar":
                success_title = "تم إعادة التعيين"
                success_message = "تم إعادة تعيين الخطوة 4 بنجاح. تم مسح جدول الإشارات المرجعية. يمكنك الآن تعديل الإعدادات أو المتابعة من الخطوة 3."
            else:
                success_title = "Reset Complete"
                success_message = "Step 4 has been reset successfully. Bookmarks table cleared. You can now modify settings or continue from Step 3."

            QMessageBox.information(self, success_title, success_message)

    def reset_all_data(self):
        """Reset all bookmark tab data and return to initial state"""
        # Clear all data
        self.pdf_path = ""
        self.bookmarks = []
        self.bookmark_file_path = ""
        self.selected_method = None
        self.navigation_count = 0
        self.navigation_direction = 0
        self.page_offset_value = 0

        # Close PDF document if open
        if hasattr(self, 'pdf_doc') and self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None

        # Reset page tracking
        self.current_page = 0
        self.total_pages = 0

        # Clear verified bookmarks
        if hasattr(self, 'verified_bookmarks'):
            self.verified_bookmarks.clear()

        # Reset UI elements
        if hasattr(self, 'pdf_display'):
            self.pdf_display.clear()

        if hasattr(self, 'toc_text'):
            self.toc_text.clear()

        if hasattr(self, 'bookmark_table'):
            self.bookmark_table.setRowCount(0)

        # Reset prepare bookmarks button
        if hasattr(self, 'btn_prepare_bookmarks'):
            self.btn_prepare_bookmarks.setEnabled(False)

        # Hide all steps except 1 and 2
        if hasattr(self, 'step2a_group'):
            self.step2a_group.setVisible(False)
        if hasattr(self, 'step2b_group'):
            self.step2b_group.setVisible(False)
        if hasattr(self, 'step3_group'):
            self.step3_group.setVisible(False)
        if hasattr(self, 'step4_group'):
            self.step4_group.setVisible(False)

        # Show Steps 1 and 2
        self.show_previous_steps()

        # Clear PDF preview
        if hasattr(self, 'pdf_preview') and self.pdf_preview:
            self.pdf_preview.clear()
            self.pdf_preview.setText("معاينة PDF\nسيتم عرض الصفحات هنا")

        print("DEBUG: All bookmark tab data has been reset")

    def select_pdf_file(self):
        """Select PDF file and enable bookmark preparation"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path

            # Set the display text and ensure it's visible
            self.pdf_display.setText(file_path)
            self.pdf_display.setToolTip(file_path)  # Add tooltip for long paths

            # Force refresh the display to ensure text is visible
            self.pdf_display.repaint()
            self.pdf_display.update()

            print(f"DEBUG: PDF path set to: {file_path}")  # Debug output
            print(f"DEBUG: Display text is: '{self.pdf_display.text()}'")  # Debug output
            print(f"DEBUG: Display widget visible: {self.pdf_display.isVisible()}")  # Debug output

            # Enable prepare bookmarks button FIRST (before any potential errors)
            self.btn_prepare_bookmarks.setEnabled(True)

            # Load PDF preview immediately (with error handling)
            try:
                self.load_pdf_preview()
            except Exception as e:
                # Button should still be enabled even if preview fails
                QMessageBox.warning(self, "تحذير", f"تم تحميل الملف ولكن فشل في عرض المعاينة:\n{str(e)}")

    def open_toc_page_selection_dialog(self):
        """Open TOC page selection dialog"""
        if not self.pdf_path or not self.pdf_doc:
            QMessageBox.warning(self, "No PDF", "Please select and load a PDF file first.")
            return

        dialog = TOCPageSelectionDialog(self.pdf_doc, self.total_pages, self.localization, self.main_window or self)
        if dialog.exec() == QDialog.Accepted:
            start_page, end_page = dialog.get_selected_range()
            # Update the TOC range spinboxes if they exist
            if hasattr(self, 'toc_start_page') and hasattr(self, 'toc_end_page'):
                self.toc_start_page.setValue(start_page)
                self.toc_end_page.setValue(end_page)

            # Show success message with proper localization and icon
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setWindowTitle("📋 " + self.localization.get_text("toc_range_selected"))
            msg_box.setText(
                f"✅ {self.localization.get_text('toc_range_selected')}\n"
                f"📄 {self.localization.get_text('selected_range')} {start_page} {self.localization.get_text('pages')} {end_page}\n\n"
                f"🔄 {self.localization.get_text('can_extract_now')}"
            )
            msg_box.exec()

    # REMOVED: File upload functionality - only text paste is supported now
    # def select_bookmark_file(self):
    #     """Select bookmark file"""
    #     file_path, _ = QFileDialog.getOpenFileName(
    #         self, "Select Bookmark File", "", "Text Files (*.txt);;All Files (*)"
    #     )
    #     if file_path:
    #         self.bookmark_file_path = file_path
    #         self.bookmark_display.setText(file_path)
    #         self.update_load_button()

    def toggle_text_file_section(self):
        """Toggle the expandable text file section"""
        is_visible = self.text_file_section.isVisible()
        self.text_file_section.setVisible(not is_visible)

        # Update button appearance
        if not is_visible:
            self.btn_expand_text_file.setText("📄-")
            self.btn_expand_text_file.setStyleSheet("""
                QPushButton {
                    background-color: #FF9800;
                    color: white;
                    font-weight: bold;
                    padding: 8px;
                    max-width: 50px;
                    border-radius: 4px;
                }
                QPushButton:hover { background-color: #F57C00; }
            """)
        else:
            self.btn_expand_text_file.setText("📄+")
            self.btn_expand_text_file.setStyleSheet("""
                QPushButton {
                    background-color: #666;
                    color: white;
                    font-weight: bold;
                    padding: 8px;
                    max-width: 50px;
                    border-radius: 4px;
                }
                QPushButton:hover { background-color: #888; }
            """)

    def extract_toc_pages_and_open_folder(self):
        """Extract TOC pages and automatically open containing folder"""
        try:
            if not self.pdf_path:
                QMessageBox.warning(self, "No PDF", "Please select a PDF file first.")
                return

            start_page = self.toc_start_page.value()
            end_page = self.toc_end_page.value()

            if start_page > end_page:
                QMessageBox.warning(self, "Invalid Range", "Start page must be less than or equal to end page.")
                return

            # Create output filename with Arabic words for clarity
            base_name = os.path.splitext(os.path.basename(self.pdf_path))[0]
            output_filename = f"{base_name}_فهرس_المحتويات_صفحات_{start_page}-{end_page}.pdf"
            output_path = os.path.join(os.path.dirname(self.pdf_path), output_filename)

            # Extract pages - create list of page numbers
            pages_to_extract = list(range(start_page, end_page + 1))
            success = self.pdf_ops.extract_pages(self.pdf_path, pages_to_extract, output_path)

            if success:
                # Show success message with proper localization and icon
                msg_box = QMessageBox(self)
                msg_box.setIcon(QMessageBox.Information)
                msg_box.setWindowTitle("🎉 " + self.localization.get_text("success"))
                msg_box.setText(
                    f"✅ {self.localization.get_text('toc_extracted_success')}\n\n"
                    f"📁 {self.localization.get_text('file')} {output_filename}\n\n"
                    f"🔄 {self.localization.get_text('folder_will_open')}"
                )
                msg_box.exec()

                # Open containing folder and highlight the file
                self.open_folder_and_highlight(output_path)
            else:
                QMessageBox.critical(self, "Error", "Failed to extract TOC pages.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error extracting TOC pages: {str(e)}")

    def open_folder_and_highlight(self, file_path):
        """Open folder and highlight the specified file"""
        try:
            import subprocess
            import platform

            # Ensure we're using the absolute path
            abs_file_path = os.path.abspath(file_path)

            if platform.system() == "Windows":
                subprocess.run(["explorer", "/select,", abs_file_path], check=False)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", "-R", abs_file_path], check=False)
            else:  # Linux
                subprocess.run(["xdg-open", os.path.dirname(abs_file_path)], check=False)
        except Exception as e:
            print(f"Error opening folder: {e}")
            # Fallback: just open the directory
            try:
                import webbrowser
                webbrowser.open(os.path.dirname(os.path.abspath(file_path)))
            except Exception as fallback_error:
                print(f"Fallback error: {fallback_error}")
                pass

    def on_toc_text_changed(self):
        """Enable load button and show Step 3 when text is pasted"""
        has_text = bool(self.toc_text.toPlainText().strip())
        if hasattr(self, 'btn_load_bookmarks'):
            self.btn_load_bookmarks.setEnabled(has_text)

        # Show Step 3 when text is available
        if hasattr(self, 'step3_group') and has_text:
            self.step3_group.setVisible(True)

    def auto_format_and_load(self):
        """Automatically format TOC and load files when text is pasted (deprecated - now manual)"""
        # This method is kept for compatibility but functionality moved to manual button
        pass

    def format_and_load_files(self):
        """Combined format TOC and load files functionality"""
        try:
            # Check if widgets exist before accessing them
            if not hasattr(self, 'toc_text') or not self.toc_text:
                QMessageBox.warning(self, "Error", "Text input widget not found.")
                return

            # Get text content based on selected method
            toc_text = ""
            if self.selected_method == "text_file":
                # Try to get text from file first, then from text area
                if hasattr(self, 'bookmark_file_path') and self.bookmark_file_path:
                    try:
                        with open(self.bookmark_file_path, 'r', encoding='utf-8') as f:
                            toc_text = f.read().strip()
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Failed to read file: {str(e)}")
                        return
                elif self.toc_text.toPlainText().strip():
                    toc_text = self.toc_text.toPlainText().strip()
            else:
                # For TOC pages method, get text from text area
                toc_text = self.toc_text.toPlainText().strip()

            if not toc_text:
                QMessageBox.warning(self, "No Text", "No text found to process. Please select a file or paste text.")
                return

            # Format the TOC text
            self.bookmarks = self.pdf_ops.parse_toc_text(toc_text)

            if not self.bookmarks:
                QMessageBox.warning(self, "No Bookmarks", "No valid bookmarks found in the text.")
                return

            # Load PDF if not already loaded
            if not self.pdf_doc and self.pdf_path:
                self.load_pdf_preview()

            # Update table and enable controls
            self.update_table()
            self.enable_bookmark_controls()

            # Show Step 4 when bookmarks are loaded and hide previous steps
            if hasattr(self, 'step4_group'):
                print("DEBUG: Showing Step 4 and hiding previous steps")
                self.step4_group.setVisible(True)
                print(f"DEBUG: Step 4 visible: {self.step4_group.isVisible()}")
                # Auto-hide previous sections to maximize screen space
                self.hide_previous_steps()
                print("DEBUG: Step 4 setup complete")

            # Force enable the insert button after everything is set up
            self.force_enable_insert_button()

            QMessageBox.information(
                self,
                self.localization.get_text("success"),
                f"{self.localization.get_text('bookmarks_loaded_successfully')}\n\n"
                f"{self.localization.get_text('found_bookmarks').format(count=len(self.bookmarks))}\n"
                f"{self.localization.get_text('preview_and_insert_instruction')}"
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error formatting TOC: {str(e)}")

    def on_offset_changed(self):
        """Handle offset value changes"""
        self.update_table()

    def update_offset_info(self):
        """Update offset information display"""
        if hasattr(self, 'offset_info') and self.offset_info and self.navigation_count > 0:
            try:
                direction_text = "⬅️" if self.navigation_direction < 0 else "➡️"
                self.offset_info.setText(
                    f"📍 Navigation: {self.navigation_count} presses {direction_text} "
                    f"(Net: {self.navigation_direction:+d} pages)"
                )
            except RuntimeError:
                # Widget has been deleted, ignore
                pass

    def auto_calculate_offset(self):
        """Auto-calculate and populate offset based on navigation tracking"""
        if self.initial_preview_page is not None and self.current_preview_page is not None:
            calculated_offset = self.current_preview_page - self.initial_preview_page

            # Automatically populate the offset spinbox
            self.page_offset.setValue(calculated_offset)

            # Enable confirm button
            if hasattr(self, 'btn_confirm_offset'):
                self.btn_confirm_offset.setEnabled(True)

            # Show detailed confirmation with navigation info
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setWindowTitle("🔄 " + self.localization.get_text("auto_calculate"))
            msg_box.setText(
                f"✅ تم حساب الإزاحة تلقائياً: {calculated_offset:+d}\n\n"
                f"📊 ملخص التنقل:\n"
                f"• عدد الضغطات: {self.navigation_count}\n"
                f"• الاتجاه الصافي: {self.navigation_direction:+d} صفحة\n"
                f"• الصفحة الأولى: {self.initial_preview_page}\n"
                f"• الصفحة الحالية: {self.current_preview_page}\n\n"
                f"📝 تم تعيين الإزاحة تلقائياً في المربع.\n"
                f"اضغط 'تأكيد' لتطبيق هذه الإزاحة على جميع الإشارات المرجعية."
            )
            msg_box.exec()

    def confirm_offset(self):
        """Confirm and apply the calculated offset"""
        current_offset = 0
        if hasattr(self, 'page_offset') and self.page_offset:
            try:
                current_offset = self.page_offset.value()
            except RuntimeError:
                current_offset = 0

        # Show confirmation dialog
        reply = QMessageBox.question(
            self,
            "✅ " + self.localization.get_text("confirm"),
            f"Apply offset of {current_offset:+d} to all bookmarks?\n\n"
            f"This will update all bookmark page numbers.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )

        if reply == QMessageBox.Yes:
            # Apply offset and update table
            self.update_table()

            # Show success message
            QMessageBox.information(
                self,
                "✅ " + self.localization.get_text("success"),
                f"Offset of {current_offset:+d} applied successfully!\n\n"
                f"All bookmark page numbers have been updated."
            )

            # Reset tracking
            self.navigation_count = 0
            self.navigation_direction = 0
            self.btn_confirm_offset.setEnabled(False)

    def on_level_changed(self, row, new_level):
        """Handle level change for a bookmark"""
        if 0 <= row < len(self.bookmarks):
            old_level = self.bookmarks[row].level
            self.bookmarks[row].level = new_level

            # Update title formatting based on level change
            if new_level == 1 and old_level == 2:
                # L2 to L1: Remove numbering
                title = self.bookmarks[row].title
                # Remove hierarchical numbering (e.g., "1.2 Title" -> "Title")
                clean_title = re.sub(r'^\d+\.\d+\s+', '', title)
                self.bookmarks[row].title = clean_title
            elif new_level == 2 and old_level == 1:
                # L1 to L2: Add numbering (find appropriate numbering)
                title = self.bookmarks[row].title
                # Find the last L1 before this item to determine numbering
                l1_counter = 1
                l2_counter = 1

                for i in range(row):
                    if self.bookmarks[i].level == 1:
                        l1_counter = i + 1  # Simple counter for now
                        l2_counter = 1
                    elif self.bookmarks[i].level == 2:
                        l2_counter += 1

                # Add hierarchical numbering
                numbered_title = f"{l1_counter}.{l2_counter} {title}"
                self.bookmarks[row].title = numbered_title

            # Refresh the table to show updated title
            self.update_table()

            # Show confirmation
            QMessageBox.information(
                self,
                "✅ " + self.localization.get_text("success"),
                f"Bookmark level changed from L{old_level} to L{new_level}\n\n"
                f"Title: {self.bookmarks[row].title}"
            )

    def show_bookmark_context_menu(self, position):
        """Show context menu for bookmark table"""
        if self.bookmark_table.itemAt(position) is not None:
            row = self.bookmark_table.rowAt(position.y())
            if 0 <= row < len(self.bookmarks):
                context_menu = QMenu(self)

                # Delete action
                delete_action = QAction("🗑️ " + self.localization.get_text("delete_bookmark"), self)
                delete_action.triggered.connect(lambda: self.delete_bookmark(row))
                context_menu.addAction(delete_action)

                # Show context menu
                context_menu.exec(self.bookmark_table.mapToGlobal(position))

    def delete_bookmark(self, row):
        """Delete a bookmark from the table"""
        if 0 <= row < len(self.bookmarks):
            bookmark = self.bookmarks[row]

            # Confirmation dialog
            reply = QMessageBox.question(
                self,
                "🗑️ " + self.localization.get_text("delete_bookmark"),
                f"{self.localization.get_text('confirm_delete')}\n\n"
                f"📖 {bookmark.title}\n"
                f"📄 Page: {bookmark.page}",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                # Remove bookmark from list
                del self.bookmarks[row]

                # Update table
                self.update_table()

                # Show success message
                QMessageBox.information(
                    self,
                    "✅ " + self.localization.get_text("success"),
                    f"Bookmark deleted successfully!\n\n"
                    f"📖 {bookmark.title}"
                )

    def fix_current_mismatch(self):
        """Fix mismatch for currently selected bookmark"""
        if not self.selected_bookmark:
            QMessageBox.information(
                self,
                self.localization.get_text("no_selection"),
                self.localization.get_text("select_bookmark_to_fix")
            )
            return

        # Apply current offset to selected bookmark
        offset = 0
        if hasattr(self, 'page_offset') and self.page_offset:
            try:
                offset = self.page_offset.value()
            except RuntimeError:
                offset = 0
        if offset != 0:
            self.selected_bookmark['page'] += offset
            self.update_table()
            QMessageBox.information(
                self,
                "Fixed",
                f"Applied offset of {offset} to selected bookmark.\n"
                f"New page: {self.selected_bookmark['page']}"
            )

    def enable_bookmark_controls(self):
        """Enable bookmark-related controls after loading"""
        if hasattr(self, 'btn_fix_mismatch'):
            self.btn_fix_mismatch.setEnabled(bool(self.bookmarks))
        if hasattr(self, 'btn_view_bookmark'):
            self.btn_view_bookmark.setEnabled(bool(self.bookmarks))
        if hasattr(self, 'btn_insert_bookmarks'):
            self.btn_insert_bookmarks.setEnabled(bool(self.bookmarks))

    def update_load_button(self):
        """Update UI state based on loaded files"""
        # Enable controls based on what we have loaded
        has_pdf = bool(self.pdf_path)
        has_bookmarks = bool(self.bookmarks or self.bookmark_file_path)

        # Enable TOC extraction when PDF is selected
        if hasattr(self, 'btn_extract_toc'):
            self.btn_extract_toc.setEnabled(has_pdf)

        # Enable bookmark controls when we have both PDF and bookmarks
        if has_pdf and has_bookmarks:
            self.enable_bookmark_controls()

    def load_files(self):
        """Load files and show preview"""
        try:
            # Check if we have formatted TOC bookmarks from TOC preparation
            if self.bookmarks:
                # Use existing formatted bookmarks from TOC preparation
                pass
            elif self.bookmark_file_path:
                # Load bookmarks from file
                self.bookmarks = self.pdf_ops.load_bookmarks_from_text(self.bookmark_file_path)
            else:
                QMessageBox.warning(self, "No Bookmarks", "Please prepare TOC or select a bookmark file first.")
                return

            if not self.bookmarks:
                QMessageBox.warning(self, "No Bookmarks", "No bookmarks found. Please prepare TOC or select a valid bookmark file.")
                return

            # Load PDF
            if not self.pdf_path:
                QMessageBox.warning(self, "No PDF", "Please select a PDF file first.")
                return

            self.pdf_doc = fitz.open(self.pdf_path)
            self.current_page = 0
            self.total_pages = self.pdf_doc.page_count

            # Update UI
            self.update_table()
            self.update_preview()
            self.btn_insert.setEnabled(True)
            self.btn_prev.setEnabled(True)
            self.btn_next.setEnabled(True)

            QMessageBox.information(
                self, "Files Loaded",
                f"Successfully loaded:\n• {len(self.bookmarks)} bookmarks\n• PDF with {self.total_pages} pages\n\nYou can now preview and insert bookmarks!"
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load files:\n{str(e)}")

    def update_table(self):
        """Update bookmark table with cumulative offset tracking"""
        if not self.bookmarks:
            return

        # Get current offset value
        current_offset = getattr(self, 'page_offset_value', 0)
        if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
            try:
                current_offset = self.offset_spinbox.value()
            except RuntimeError:
                pass

        # Calculate the delta offset (change from last applied offset)
        delta_offset = current_offset - self.last_applied_offset
        print(f"DEBUG: Current offset: {current_offset}, Last applied: {self.last_applied_offset}, Delta: {delta_offset}")

        if hasattr(self, 'bookmark_table') and self.bookmark_table:
            self.bookmark_table.setRowCount(len(self.bookmarks))
        else:
            return

        for i, bookmark in enumerate(self.bookmarks):
            # Title
            title_item = QTableWidgetItem(bookmark.title)
            title_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)  # RTL alignment
            self.bookmark_table.setItem(i, 0, title_item)

            # Original page
            orig_item = QTableWidgetItem(str(bookmark.page))
            orig_item.setTextAlignment(Qt.AlignCenter)
            self.bookmark_table.setItem(i, 1, orig_item)

            # Modified page - use cumulative approach
            # If this is the first time, use original + current offset
            # Otherwise, read current value from table and apply delta
            if delta_offset == current_offset:  # First time setting up table
                modified_page = bookmark.page + current_offset
                print(f"DEBUG: First setup for '{bookmark.title}': {bookmark.page} + {current_offset} = {modified_page}")
            else:
                # Read current modified page from table and apply delta
                try:
                    existing_item = self.bookmark_table.item(i, 2)
                    if existing_item:
                        current_modified = int(existing_item.text())
                        modified_page = current_modified + delta_offset
                        print(f"DEBUG: Delta update for '{bookmark.title}': {current_modified} + {delta_offset} = {modified_page}")
                    else:
                        modified_page = bookmark.page + current_offset
                        print(f"DEBUG: No existing item for '{bookmark.title}', using: {bookmark.page} + {current_offset} = {modified_page}")
                except (ValueError, AttributeError):
                    modified_page = bookmark.page + current_offset
                    print(f"DEBUG: Error reading existing, using: {bookmark.page} + {current_offset} = {modified_page}")
            mod_item = QTableWidgetItem(str(modified_page))
            mod_item.setTextAlignment(Qt.AlignCenter)

            # Highlight invalid pages
            if modified_page <= 0:
                mod_item.setBackground(Qt.red)
                mod_item.setToolTip("صفحة غير صالحة - أقل من 1")
            elif hasattr(self, 'total_pages') and modified_page > self.total_pages:
                mod_item.setBackground(Qt.yellow)
                mod_item.setToolTip(f"صفحة غير صالحة - أكبر من {self.total_pages}")
            else:
                mod_item.setBackground(Qt.white)
                mod_item.setToolTip("")

            self.bookmark_table.setItem(i, 2, mod_item)

            # Level - Improved combo box (only create if doesn't exist)
            existing_combo = self.bookmark_table.cellWidget(i, 3)
            if existing_combo is None or not isinstance(existing_combo, QComboBox):
                level_combo = QComboBox()
                level_combo.addItems(["المستوى 1", "المستوى 2", "المستوى 3"])  # Arabic levels
                level_combo.setCurrentIndex(bookmark.level - 1 if bookmark.level <= 3 else 0)
                level_combo.currentIndexChanged.connect(lambda idx, row=i: self.on_level_changed(row, idx + 1))
                level_combo.setStyleSheet("""
                    QComboBox {
                        text-align: center;
                        padding: 4px;
                        font-weight: bold;
                    }
                """)
                self.bookmark_table.setCellWidget(i, 3, level_combo)
            else:
                # Update existing combo box if level changed
                existing_combo.blockSignals(True)  # Prevent triggering change event
                existing_combo.setCurrentIndex(bookmark.level - 1 if bookmark.level <= 3 else 0)
                existing_combo.blockSignals(False)

            # REMOVED: Action buttons (Preview and Delete) - Now using automatic preview on row selection

        # Update the last applied offset to track cumulative changes
        self.last_applied_offset = current_offset
        print(f"DEBUG: Updated last_applied_offset to {self.last_applied_offset}")

    def update_preview(self):
        """Update PDF preview"""
        if not self.pdf_doc:
            return

        # Check if preview widget exists
        if not hasattr(self, 'pdf_preview') or not self.pdf_preview:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Get preview widget size for optimal scaling
            preview_size = self.pdf_preview.size()
            available_width = max(preview_size.width() - 20, 400)  # Minimum 400px
            available_height = max(preview_size.height() - 20, 600)  # Minimum 600px

            # Get page dimensions
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height

            # Calculate optimal scale to fill the preview area
            scale_x = available_width / page_width
            scale_y = available_height / page_height
            scale = min(scale_x, scale_y, 2.0)  # Max scale of 2.0 for quality

            # Render page with calculated scale
            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            # Convert to QPixmap
            pixmap = QPixmap()
            pixmap.loadFromData(img_data)

            if not pixmap.isNull():
                # The pixmap is already at the right size, just set it
                self.pdf_preview.setPixmap(pixmap)
            else:
                self.pdf_preview.setText("لا يمكن عرض الصفحة")

            # Update page label with clean display (if it exists)
            if hasattr(self, 'page_label') and self.page_label:
                page_text = self.localization.get_text("page_info").format(
                    current=self.current_page + 1,
                    total=self.total_pages
                )

                # Add bookmark title if selected and page matches
                if hasattr(self, 'selected_bookmark') and self.selected_bookmark:
                    if hasattr(self, 'page_offset') and self.page_offset:
                        try:
                            offset = self.page_offset.value()
                            expected_page = self.selected_bookmark.page + offset
                            current_pdf_page = self.current_page + 1

                            if expected_page == current_pdf_page:
                                page_text += f" ✅ ({self.selected_bookmark.title})"
                        except RuntimeError:
                            # Widget deleted
                            pass

                try:
                    self.page_label.setText(page_text)
                except RuntimeError:
                    pass

            # Update navigation buttons
            if hasattr(self, 'btn_prev') and self.btn_prev:
                try:
                    self.btn_prev.setEnabled(self.current_page > 0)
                except RuntimeError:
                    pass
            if hasattr(self, 'btn_next') and self.btn_next:
                try:
                    self.btn_next.setEnabled(self.current_page < self.total_pages - 1)
                except RuntimeError:
                    pass

        except Exception as e:
            if hasattr(self, 'pdf_preview') and self.pdf_preview:
                self.pdf_preview.setText(f"خطأ في المعاينة: {str(e)}")

    def prev_page(self):
        """Go to previous page and update all bookmark offsets"""
        if self.pdf_doc and self.current_page > 0:
            print(f"DEBUG: prev_page called, current_page: {self.current_page}")  # Debug
            self.current_page -= 1
            self.current_preview_page = self.current_page + 1  # Track current page
            print(f"DEBUG: new current_page: {self.current_page}")  # Debug

            # Update offset spinbox to affect all bookmarks
            if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
                try:
                    current_offset = self.offset_spinbox.value()
                    print(f"DEBUG: updating offset from {current_offset} to {current_offset - 1}")  # Debug
                    self.offset_spinbox.setValue(current_offset - 1)
                except RuntimeError:
                    pass
            elif hasattr(self, 'page_offset') and self.page_offset:
                try:
                    current_offset = self.page_offset.value()
                    self.page_offset.setValue(current_offset - 1)
                except RuntimeError:
                    pass

            # Track navigation for automatic offset detection
            if self.initial_preview_page is not None:
                self.navigation_count += 1
                self.navigation_direction -= 1  # Going backwards
                self.update_offset_info()

            print("DEBUG: calling update_preview()")  # Debug
            self.update_preview()

    def next_page(self):
        """Go to next page and update all bookmark offsets"""
        if self.pdf_doc and self.current_page < self.total_pages - 1:
            print(f"DEBUG: next_page called, current_page: {self.current_page}")  # Debug
            self.current_page += 1
            self.current_preview_page = self.current_page + 1  # Track current page
            print(f"DEBUG: new current_page: {self.current_page}")  # Debug

            # Update offset spinbox to affect all bookmarks
            if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
                try:
                    current_offset = self.offset_spinbox.value()
                    print(f"DEBUG: updating offset from {current_offset} to {current_offset + 1}")  # Debug
                    self.offset_spinbox.setValue(current_offset + 1)
                except RuntimeError:
                    pass
            elif hasattr(self, 'page_offset') and self.page_offset:
                try:
                    current_offset = self.page_offset.value()
                    self.page_offset.setValue(current_offset + 1)
                except RuntimeError:
                    pass

            # Track navigation for automatic offset detection
            if self.initial_preview_page is not None:
                self.navigation_count += 1
                self.navigation_direction += 1  # Going forwards
                self.update_offset_info()

            print("DEBUG: calling update_preview()")  # Debug
            self.update_preview()

    def on_bookmark_clicked(self, row, column):
        """Handle bookmark table click and track page navigation"""
        if not self.bookmarks or not self.pdf_doc:
            return

        self.selected_bookmark_row = row
        self.selected_bookmark = self.bookmarks[row]

        # Navigate to bookmark page - safely get offset
        offset = 0
        if hasattr(self, 'page_offset') and self.page_offset:
            try:
                offset = self.page_offset.value()
            except RuntimeError:
                # Widget has been deleted, use default
                offset = 0
        elif hasattr(self, 'offset_spinbox') and self.offset_spinbox:
            try:
                offset = self.offset_spinbox.value()
            except RuntimeError:
                # Widget has been deleted, use default
                offset = 0

        target_page = self.selected_bookmark.page + offset

        if 1 <= target_page <= self.total_pages:
            # Reset and track initial page for offset calculation
            self.initial_preview_page = target_page
            self.navigation_count = 0
            self.navigation_direction = 0

            self.current_page = target_page - 1
            self.current_preview_page = target_page
            self.update_preview()
            self.btn_auto_offset.setEnabled(True)
            self.bookmark_table.selectRow(row)

            # Update offset info display
            if hasattr(self, 'offset_info'):
                self.offset_info.setText("📍 " + self.localization.get_text("navigate_to_correct_page"))
        else:
            QMessageBox.warning(
                self, self.localization.get_text("invalid_page"),
                self.localization.get_text("page_out_of_range").format(
                    page=target_page,
                    total=self.total_pages
                )
            )

    def auto_detect_offset(self):
        """Auto-detect page offset"""
        if not self.selected_bookmark or not self.pdf_doc:
            return

        current_pdf_page = self.current_page + 1
        bookmark_page = self.selected_bookmark.page

        actual_page, ok = QInputDialog.getInt(
            self, "Confirm Actual Page Number",
            f"You're viewing PDF page {current_pdf_page}.\n"
            f"Bookmark '{self.selected_bookmark.title}' expects page {bookmark_page}.\n\n"
            f"What is the ACTUAL page number shown in the PDF content?",
            current_pdf_page, 1, self.total_pages
        )

        if ok:
            correct_offset = actual_page - bookmark_page
            self.page_offset.setValue(correct_offset)
            self.update_table()

            QMessageBox.information(
                self, "Offset Detected!",
                f"Page offset set to {correct_offset:+d}\n\n"
                f"Bookmark page {bookmark_page} → PDF page {actual_page}\n"
                f"All bookmarks adjusted accordingly."
            )

            self.btn_auto_offset.setEnabled(False)
            self.btn_apply_correction.setEnabled(False)

    def apply_manual_correction(self):
        """Apply manual correction for the selected bookmark"""
        if not self.selected_bookmark or not self.pdf_doc:
            return

        current_pdf_page = self.current_page + 1
        bookmark_page = self.selected_bookmark.page

        # Create correction dialog
        dialog = BookmarkCorrectionDialog(
            bookmark_title=self.selected_bookmark.title,
            bookmark_page=bookmark_page,
            current_pdf_page=current_pdf_page,
            total_pages=self.total_pages,
            localization=self.localization,
            parent=self
        )

        if dialog.exec() == QDialog.Accepted:
            correct_page = dialog.get_correct_page()
            if correct_page:
                # Calculate and apply offset
                correct_offset = correct_page - bookmark_page
                self.page_offset.setValue(correct_offset)
                self.update_table()

                QMessageBox.information(
                    self,
                    self.localization.get_text("correction_applied"),
                    f"Page offset set to {correct_offset:+d}\n\n"
                    f"Bookmark '{self.selected_bookmark.title}'\n"
                    f"Original page: {bookmark_page} → Corrected page: {correct_page}\n"
                    f"All bookmarks adjusted accordingly."
                )

                self.btn_auto_offset.setEnabled(False)
                self.btn_apply_correction.setEnabled(False)

    def insert_bookmarks(self):
        """Insert bookmarks into PDF"""
        try:
            print("=" * 50)  # Debug separator
            print("INSERT BOOKMARKS FUNCTION CALLED!")  # Debug
            print(f"Button enabled: {self.btn_insert.isEnabled()}")  # Debug
            print(f"Button visible: {self.btn_insert.isVisible()}")  # Debug

            # Check if we have bookmarks
            print(f"Bookmarks available: {hasattr(self, 'bookmarks')}")  # Debug
            if hasattr(self, 'bookmarks'):
                print(f"Number of bookmarks: {len(self.bookmarks)}")  # Debug
            else:
                print("ERROR: No bookmarks attribute found!")  # Debug
                QMessageBox.critical(self, "خطأ", "لا توجد إشارات مرجعية محملة")
                return

            if not self.bookmarks:
                print("ERROR: Bookmarks list is empty!")  # Debug
                QMessageBox.warning(self, "لا توجد إشارات مرجعية", "لا توجد إشارات مرجعية للإدراج")
                return

            # Check if we have PDF path
            print(f"PDF path available: {hasattr(self, 'pdf_path')}")  # Debug
            if hasattr(self, 'pdf_path'):
                print(f"PDF path: {self.pdf_path}")  # Debug
            else:
                print("ERROR: No pdf_path attribute found!")  # Debug
                QMessageBox.critical(self, "خطأ", "لا يوجد مسار ملف PDF")
                return

            if not self.pdf_path:
                print("ERROR: PDF path is empty!")  # Debug
                QMessageBox.warning(self, "لا يوجد ملف PDF", "لم يتم تحديد ملف PDF")
                return

            # Get current page offset
            print("Getting page offset...")  # Debug
            offset = getattr(self, 'page_offset_value', 0)
            print(f"Default offset from page_offset_value: {offset}")  # Debug

            if hasattr(self, 'offset_spinbox') and self.offset_spinbox:
                try:
                    spinbox_offset = self.offset_spinbox.value()
                    print(f"Spinbox offset: {spinbox_offset}")  # Debug
                    offset = spinbox_offset
                except RuntimeError as e:
                    print(f"RuntimeError accessing spinbox: {e}")  # Debug
                    offset = getattr(self, 'page_offset_value', 0)
            else:
                print("No offset_spinbox found or it's None")  # Debug

            print(f"Final offset to use: {offset}")  # Debug

            # Get output file
            print("Opening file save dialog...")  # Debug
            pdf_name = os.path.basename(self.pdf_path)
            # Use Arabic suffix for localized filename
            suffix = "_مفهرس.pdf" if self.localization.current_language == "ar" else "_with_bookmarks.pdf"
            default_name = pdf_name.replace('.pdf', suffix)
            print(f"Default filename: {default_name}")  # Debug

            output_path, _ = QFileDialog.getSaveFileName(
                self, "حفظ PDF مع الإشارات المرجعية",
                os.path.join(os.path.dirname(self.pdf_path), default_name),
                "PDF Files (*.pdf)"
            )

            if not output_path:
                print("User cancelled file dialog")  # Debug
                return  # User cancelled

            print(f"User selected output path: {output_path}")  # Debug

            # Check if pdf_ops exists
            print("Checking pdf_ops...")  # Debug
            if not hasattr(self, 'pdf_ops'):
                print("ERROR: No pdf_ops attribute found!")  # Debug
                QMessageBox.critical(self, "خطأ", "خطأ في النظام - pdf_ops غير موجود")
                return
            elif not self.pdf_ops:
                print("ERROR: pdf_ops is None!")  # Debug
                QMessageBox.critical(self, "خطأ", "خطأ في النظام - pdf_ops فارغ")
                return
            else:
                print(f"pdf_ops found: {type(self.pdf_ops)}")  # Debug

            # Insert bookmarks with current offset
            print("Calling insert_bookmarks_into_pdf...")  # Debug
            print(f"Parameters: pdf_path={self.pdf_path}, bookmarks_count={len(self.bookmarks)}, output_path={output_path}, offset={offset}")  # Debug

            success = self.pdf_ops.insert_bookmarks_into_pdf(
                self.pdf_path, self.bookmarks, output_path, offset
            )
            print(f"Insert operation result: {success}")  # Debug

            if success:
                QMessageBox.information(
                    self, "نجح!",
                    f"تم إدراج {len(self.bookmarks)} إشارة مرجعية بنجاح!\n\n"
                    f"الإزاحة المطبقة: {offset:+d}\n"
                    f"الملف المحفوظ: {os.path.basename(output_path)}"
                )
            else:
                QMessageBox.critical(self, "خطأ", "فشل في إدراج الإشارات المرجعية")

        except Exception as e:
            print(f"Error in insert_bookmarks: {e}")  # Debug
            import traceback
            traceback.print_exc()  # Print full traceback
            QMessageBox.critical(self, "خطأ", f"فشل في إدراج الإشارات المرجعية:\n{str(e)}")

    # Split by bookmarks functionality moved to separate SplitByBookmarksTab

    def load_pdf_preview(self):
        """Load PDF preview immediately when PDF is selected"""
        if not self.pdf_path:
            return

        try:
            if self.pdf_doc:
                self.pdf_doc.close()

            self.pdf_doc = fitz.open(self.pdf_path)
            self.total_pages = self.pdf_doc.page_count
            self.current_page = 0

            # Update page range for TOC extraction
            if hasattr(self, 'toc_end_page') and self.toc_end_page:
                self.toc_end_page.setMaximum(self.total_pages)

            # Enable navigation buttons if they exist
            if hasattr(self, 'btn_prev') and self.btn_prev:
                self.btn_prev.setEnabled(True)
            if hasattr(self, 'btn_next') and self.btn_next:
                self.btn_next.setEnabled(True)

            # Show first page (only if preview widget exists)
            if hasattr(self, 'pdf_preview') and self.pdf_preview:
                self.update_preview()

            # Enable method selection and set default to text file
            self.enable_method_selection()

            # Show success message with proper localization and icon
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setWindowTitle("📄 " + self.localization.get_text("success"))
            msg_box.setText(
                f"✅ {self.localization.get_text('pdf_loaded_success')}\n"
                f"📊 {self.localization.get_text('total_pages')} {self.total_pages}\n\n"
                f"👁️ {self.localization.get_text('use_preview_instruction')}"
            )
            msg_box.exec()

        except Exception as e:
            print(f"DEBUG: Exception in load_pdf_preview: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to load PDF: {str(e)}")

    def open_page_range_dialog(self):
        """Open dialog for selecting TOC page range with PDF preview"""
        if not self.pdf_path or not self.pdf_doc:
            QMessageBox.warning(self, "No PDF", "Please select and load a PDF file first.")
            return

        dialog = TOCPageRangeDialog(self.pdf_doc, self.total_pages, self.localization, self)
        if dialog.exec() == QDialog.Accepted:
            start_page, end_page = dialog.get_selected_range()
            # Update the spinboxes with selected range
            self.toc_start_page.setValue(start_page)
            self.toc_end_page.setValue(end_page)
            # Extract the pages
            self.extract_toc_pages()

    def extract_toc_pages(self):
        """Extract TOC pages from PDF"""
        if not self.pdf_path:
            QMessageBox.warning(self, "No PDF Selected", "Please select a PDF file first.")
            return

        try:
            start = self.toc_start_page.value()
            end = self.toc_end_page.value()

            if start > end:
                QMessageBox.warning(self, "Invalid Range", "Start page must be less than or equal to end page.")
                return

            if end > self.total_pages:
                QMessageBox.warning(self, "Invalid Range", f"End page cannot exceed total pages ({self.total_pages}).")
                return

            # Create output filename
            pdf_name = os.path.splitext(os.path.basename(self.pdf_path))[0]
            output_dir = os.path.dirname(self.pdf_path)
            output_filename = f"{pdf_name}_TOC_pages_{start}-{end}.pdf"
            output_path = os.path.join(output_dir, output_filename)

            # Extract pages
            doc = fitz.open(self.pdf_path)
            new_doc = fitz.open()

            for page_num in range(start - 1, min(end, doc.page_count)):
                new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)

            new_doc.save(output_path)
            new_doc.close()
            doc.close()

            # Open output directory and highlight the file
            import subprocess
            import platform

            try:
                if platform.system() == "Windows":
                    # Use explorer with /select to highlight the specific file
                    subprocess.run(["explorer", "/select,", output_path.replace('/', '\\')], check=True)
                elif platform.system() == "Darwin":  # macOS
                    # Use open -R to reveal the file in Finder
                    subprocess.run(["open", "-R", output_path], check=True)
                else:  # Linux
                    # For Linux, open the directory (file highlighting varies by file manager)
                    subprocess.run(["xdg-open", output_dir], check=True)
            except subprocess.CalledProcessError:
                # Fallback: just open the directory
                if platform.system() == "Windows":
                    subprocess.run(["explorer", output_dir.replace('/', '\\')], check=False)
                else:
                    subprocess.run(["xdg-open", output_dir], check=False)

            QMessageBox.information(
                self,
                "Pages Extracted",
                f"TOC pages {start}-{end} extracted successfully!\n\nSaved as: {output_filename}\n\nOutput directory opened. Now use AI Studio to extract the TOC text."
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to extract pages: {str(e)}")

    def open_ai_studio(self):
        """Open AI Studio in browser"""
        import webbrowser
        webbrowser.open("https://aistudio.google.com/prompts/new_chat")

    def copy_ai_prompt(self):
        """Copy AI prompt to clipboard"""
        prompt = self.localization.get_text("ai_toc_prompt")

        clipboard = QApplication.clipboard()
        clipboard.setText(prompt)

        # Show success message with proper localization and icon
        msg_box = QMessageBox(self)
        msg_box.setIcon(QMessageBox.Information)
        msg_box.setWindowTitle("📋 " + self.localization.get_text("prompt_copied"))
        msg_box.setText(
            f"✅ {self.localization.get_text('ai_prompt_copied')}\n\n"
            f"🤖 {self.localization.get_text('paste_in_ai_studio')}"
        )
        msg_box.exec()

    def format_toc(self):
        """Format the TOC text using the improved algorithm"""
        raw_text = self.toc_text.toPlainText().strip()

        if not raw_text:
            QMessageBox.warning(self, "No Text", "Please paste the TOC text from AI Studio first.")
            return

        try:
            # Use the improved formatting function
            formatted_bookmarks = self.generate_formatted_toc(raw_text)

            if not formatted_bookmarks:
                QMessageBox.warning(self, "No Valid TOC", "No valid TOC entries found. Please check the format.")
                return

            self.bookmarks = formatted_bookmarks
            self.update_table()
            self.update_load_button()  # Enable load button now that we have bookmarks

            QMessageBox.information(
                self,
                "TOC Formatted",
                f"Successfully formatted {len(formatted_bookmarks)} TOC entries!\n\nBookmarks are now loaded and ready to preview. Click 'Load Files & Preview' to continue."
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to format TOC: {str(e)}")

    def generate_formatted_toc(self, raw_text: str) -> List[Bookmark]:
        """
        Generate formatted TOC from raw text using simplified level detection.
        Level 1 (L1): Every title that appears before an empty line + first title
        Level 2 (L2): All other titles
        """
        def parse_line(line: str) -> tuple[str, str]:
            """Separate title from page number in a line."""
            try:
                parts = line.strip().rsplit(' - ', 1)
                title = parts[0].strip()
                page_number = parts[1].strip()
                return title, page_number
            except (IndexError, ValueError):
                return line.strip(), ""

        def extract_page_number(page_str: str) -> int:
            """Extract page number from string (handles Arabic numerals)"""
            # Arabic to English number mapping
            arabic_to_english = {
                '٠': '0', '١': '1', '٢': '2', '٣': '3', '٤': '4',
                '٥': '5', '٦': '6', '٧': '7', '٨': '8', '٩': '9'
            }

            # Convert Arabic numerals to English
            english_str = page_str
            for arabic, english in arabic_to_english.items():
                english_str = english_str.replace(arabic, english)

            # Extract numbers
            numbers = re.findall(r'\d+', english_str)
            return int(numbers[0]) if numbers else 0

        bookmarks = []
        lines = [line.strip() for line in raw_text.strip().split('\n') if line.strip()]

        if not lines:
            return bookmarks

        # Find lines that are preceded by empty lines in the original text
        original_lines = raw_text.strip().split('\n')
        level1_indices = set()

        # Track which non-empty lines should be Level 1
        non_empty_index = 0
        for i, original_line in enumerate(original_lines):
            if original_line.strip():  # Non-empty line
                # First non-empty line is always Level 1
                if non_empty_index == 0:
                    level1_indices.add(non_empty_index)
                # Check if this line is immediately preceded by an empty line
                elif i > 0 and not original_lines[i - 1].strip():
                    level1_indices.add(non_empty_index)

                non_empty_index += 1

        # Process all lines and assign levels
        for i, line in enumerate(lines):
            title, page_str = parse_line(line)
            page_num = extract_page_number(page_str)

            if page_num > 0:
                # Determine level based on position
                level = 1 if i in level1_indices else 2
                bookmark = Bookmark(title, page_num, level)
                bookmarks.append(bookmark)
                print(f"Generated: '{title}' -> Page {page_num}, Level {level} (index {i})")  # Debug

        return bookmarks

    def select_original_pdf(self):
        """Select the original PDF file (source of bookmarks)"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_original_pdf"),
            "",
            "PDF Files (*.pdf);;All Files (*)"
        )

        if file_path:
            self.original_pdf_path = file_path
            self.original_pdf_display.setText(os.path.basename(file_path))
            self.check_copy_ready()

    def select_enhanced_pdf(self):
        """Select the enhanced PDF file (target for bookmarks)"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_enhanced_pdf"),
            "",
            "PDF Files (*.pdf);;All Files (*)"
        )

        if file_path:
            self.enhanced_pdf_path = file_path
            self.enhanced_pdf_display.setText(os.path.basename(file_path))
            self.check_copy_ready()

    def check_copy_ready(self):
        """Check if both files are selected and enable copy button"""
        if self.original_pdf_path and self.enhanced_pdf_path:
            self.btn_copy_bookmarks.setEnabled(True)
        else:
            self.btn_copy_bookmarks.setEnabled(False)

    def copy_bookmarks_between_pdfs(self):
        """Copy bookmarks from original PDF to enhanced PDF"""
        if not self.original_pdf_path or not self.enhanced_pdf_path:
            QMessageBox.warning(
                self,
                self.localization.get_text("warning"),
                self.localization.get_text("select_both_files")
            )
            return

        try:
            # Open both PDFs
            original_doc = fitz.open(self.original_pdf_path)
            enhanced_doc = fitz.open(self.enhanced_pdf_path)

            # Get bookmarks from original PDF
            original_toc = original_doc.get_toc()

            if not original_toc:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning"),
                    self.localization.get_text("no_bookmarks_found")
                )
                original_doc.close()
                enhanced_doc.close()
                return

            # Verify page counts match (or enhanced has same or more pages)
            original_pages = original_doc.page_count
            enhanced_pages = enhanced_doc.page_count

            if enhanced_pages < original_pages:
                reply = QMessageBox.question(
                    self,
                    self.localization.get_text("page_count_mismatch"),
                    self.localization.get_text("page_count_warning").format(
                        original=original_pages,
                        enhanced=enhanced_pages
                    ),
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    original_doc.close()
                    enhanced_doc.close()
                    return

            # Copy bookmarks to enhanced PDF
            enhanced_doc.set_toc(original_toc)

            # Create output filename
            enhanced_dir = os.path.dirname(self.enhanced_pdf_path)
            enhanced_name = os.path.splitext(os.path.basename(self.enhanced_pdf_path))[0]
            output_filename = f"{enhanced_name}_with_bookmarks.pdf"
            output_path = os.path.join(enhanced_dir, output_filename)

            # Save the enhanced PDF with bookmarks
            enhanced_doc.save(output_path)

            # Close documents
            original_doc.close()
            enhanced_doc.close()

            # Show success message
            success_msg = self.localization.get_text("bookmarks_copied_success").format(
                count=len(original_toc),
                filename=output_filename
            )

            QMessageBox.information(
                self,
                self.localization.get_text("success"),
                success_msg
            )

            # Ask if user wants to open the output directory
            reply = QMessageBox.question(
                self,
                self.localization.get_text("open_directory"),
                self.localization.get_text("open_directory_question"),
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )

            if reply == QMessageBox.Yes:
                import subprocess
                import platform

                try:
                    if platform.system() == "Windows":
                        subprocess.run(["explorer", "/select,", output_path.replace('/', '\\')], check=True)
                    elif platform.system() == "Darwin":  # macOS
                        subprocess.run(["open", "-R", output_path], check=True)
                    else:  # Linux
                        subprocess.run(["xdg-open", enhanced_dir], check=True)
                except subprocess.CalledProcessError:
                    # Fallback: just open the directory
                    if platform.system() == "Windows":
                        subprocess.run(["explorer", enhanced_dir.replace('/', '\\')], check=False)
                    else:
                        subprocess.run(["xdg-open", enhanced_dir], check=False)

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("bookmark_copy_failed").format(error=str(e))
            )

    def open_bookmark_copy_dialog(self):
        """Open bookmark copy dialog"""
        dialog = QDialog(self)
        dialog.setWindowTitle(self.localization.get_text("copy_bookmarks_title"))
        dialog.setModal(True)
        dialog.resize(600, 300)

        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Title and description
        title_label = QLabel(f"📋 {self.localization.get_text('copy_bookmarks_title')}")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #9C27B0; margin-bottom: 10px;")
        layout.addWidget(title_label)

        desc_label = QLabel(self.localization.get_text('copy_bookmarks_desc'))
        desc_label.setWordWrap(True)
        desc_label.setStyleSheet("color: #666; font-size: 12px; margin-bottom: 15px;")
        layout.addWidget(desc_label)

        # File selection section
        files_group = QGroupBox(self.localization.get_text("file_selection"))
        files_layout = QVBoxLayout(files_group)

        # Original PDF
        original_layout = QHBoxLayout()
        original_layout.addWidget(QLabel(self.localization.get_text('original_pdf')))

        self.dialog_original_display = QLineEdit()
        self.dialog_original_display.setReadOnly(True)
        self.dialog_original_display.setPlaceholderText(self.localization.get_text('select_original_pdf'))
        original_layout.addWidget(self.dialog_original_display, 7)

        btn_original = QPushButton(self.localization.get_text("browse"))
        btn_original.clicked.connect(lambda: self.select_file_for_dialog('original'))
        btn_original.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover { background-color: #7B1FA2; }
        """)
        original_layout.addWidget(btn_original, 1)
        files_layout.addLayout(original_layout)

        # Enhanced PDF
        enhanced_layout = QHBoxLayout()
        enhanced_layout.addWidget(QLabel(self.localization.get_text('enhanced_pdf')))

        self.dialog_enhanced_display = QLineEdit()
        self.dialog_enhanced_display.setReadOnly(True)
        self.dialog_enhanced_display.setPlaceholderText(self.localization.get_text('select_enhanced_pdf'))
        enhanced_layout.addWidget(self.dialog_enhanced_display, 7)

        btn_enhanced = QPushButton(self.localization.get_text("browse"))
        btn_enhanced.clicked.connect(lambda: self.select_file_for_dialog('enhanced'))
        btn_enhanced.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover { background-color: #7B1FA2; }
        """)
        enhanced_layout.addWidget(btn_enhanced, 1)
        files_layout.addLayout(enhanced_layout)

        layout.addWidget(files_group)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel"))
        cancel_btn.clicked.connect(dialog.reject)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #757575;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover { background-color: #616161; }
        """)
        button_layout.addWidget(cancel_btn)

        self.dialog_copy_btn = QPushButton(f"📋 {self.localization.get_text('copy_bookmarks_button')}")
        self.dialog_copy_btn.clicked.connect(lambda: self.execute_bookmark_copy(dialog))
        self.dialog_copy_btn.setEnabled(False)
        self.dialog_copy_btn.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover { background-color: #7B1FA2; }
            QPushButton:disabled { background-color: #cccccc; color: #666; }
        """)
        button_layout.addWidget(self.dialog_copy_btn)

        layout.addLayout(button_layout)

        # Initialize dialog file paths
        self.dialog_original_path = ""
        self.dialog_enhanced_path = ""

        dialog.exec()

    def select_file_for_dialog(self, file_type):
        """Select file for bookmark copy dialog"""
        title = self.localization.get_text("select_original_pdf") if file_type == 'original' else self.localization.get_text("select_enhanced_pdf")

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            title,
            "",
            "PDF Files (*.pdf);;All Files (*)"
        )

        if file_path:
            if file_type == 'original':
                self.dialog_original_path = file_path
                self.dialog_original_display.setText(os.path.basename(file_path))
            else:
                self.dialog_enhanced_path = file_path
                self.dialog_enhanced_display.setText(os.path.basename(file_path))

            # Check if both files are selected
            if self.dialog_original_path and self.dialog_enhanced_path:
                self.dialog_copy_btn.setEnabled(True)

    def execute_bookmark_copy(self, dialog):
        """Execute bookmark copy operation from dialog"""
        if not self.dialog_original_path or not self.dialog_enhanced_path:
            return

        # Use the existing copy method with dialog paths
        original_backup = self.original_pdf_path
        enhanced_backup = self.enhanced_pdf_path

        self.original_pdf_path = self.dialog_original_path
        self.enhanced_pdf_path = self.dialog_enhanced_path

        self.copy_bookmarks_between_pdfs()

        # Restore original paths
        self.original_pdf_path = original_backup
        self.enhanced_pdf_path = enhanced_backup

        dialog.accept()


class TOCPageRangeDialog(QDialog):
    """Dialog for selecting TOC page range with PDF preview"""

    def __init__(self, pdf_doc, total_pages, localization, parent=None):
        super().__init__(parent)
        self.pdf_doc = pdf_doc
        self.total_pages = total_pages
        self.localization = localization
        # Start near the end where TOC usually is
        self.current_page = max(0, total_pages - 10)  # Start 10 pages from end
        self.start_page = max(1, total_pages - 9)  # Default start near end
        self.end_page = total_pages  # Default end at last page

        self.setWindowTitle(self.localization.get_text("select_toc_range"))
        self.setModal(True)
        self.resize(1000, 700)  # Larger dialog for better preview

        self.init_ui()
        self.update_preview()

    def init_ui(self):
        """Initialize dialog UI"""
        layout = QVBoxLayout(self)

        # Title
        title = QLabel(self.localization.get_text("select_toc_range_title"))
        title.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(title)

        # Main content
        content_layout = QHBoxLayout()

        # Left side - PDF preview
        preview_group = QGroupBox(self.localization.get_text("pdf_preview"))
        preview_layout = QVBoxLayout(preview_group)

        # Navigation
        nav_layout = QHBoxLayout()

        # First/Previous buttons
        self.btn_first = QPushButton("⏮")
        self.btn_first.setToolTip(self.localization.get_text("first_page"))
        self.btn_first.clicked.connect(self.first_page)
        nav_layout.addWidget(self.btn_first)

        self.btn_prev = QPushButton("◀")
        self.btn_prev.setToolTip(self.localization.get_text("previous_page"))
        self.btn_prev.clicked.connect(self.prev_page)
        nav_layout.addWidget(self.btn_prev)

        # Page input
        self.page_input = QSpinBox()
        self.page_input.setRange(1, self.total_pages)
        self.page_input.setValue(self.current_page + 1)
        self.page_input.valueChanged.connect(self.goto_page)
        self.page_input.setLayoutDirection(Qt.LeftToRight)
        self.page_input.setStyleSheet("QSpinBox { text-align: center; min-width: 60px; }")
        nav_layout.addWidget(self.page_input)

        self.page_label = QLabel()
        self.page_label.setAlignment(Qt.AlignCenter)
        nav_layout.addWidget(self.page_label)

        # Next/Last buttons
        self.btn_next = QPushButton("▶")
        self.btn_next.setToolTip(self.localization.get_text("next_page"))
        self.btn_next.clicked.connect(self.next_page)
        nav_layout.addWidget(self.btn_next)

        self.btn_last = QPushButton("⏭")
        self.btn_last.setToolTip(self.localization.get_text("last_page"))
        self.btn_last.clicked.connect(self.last_page)
        nav_layout.addWidget(self.btn_last)

        preview_layout.addLayout(nav_layout)

        # PDF preview
        self.pdf_preview = QLabel()
        self.pdf_preview.setAlignment(Qt.AlignCenter)
        self.pdf_preview.setMinimumSize(500, 600)
        self.pdf_preview.setStyleSheet("border: 2px solid #ccc; background-color: white; border-radius: 4px;")
        self.pdf_preview.setScaledContents(True)
        preview_layout.addWidget(self.pdf_preview)

        # Current page info
        self.page_info = QLabel()
        self.page_info.setAlignment(Qt.AlignCenter)
        self.page_info.setStyleSheet("color: #666; font-size: 12px; padding: 5px;")
        preview_layout.addWidget(self.page_info)

        content_layout.addWidget(preview_group)

        # Right side - Range selection
        range_group = QGroupBox(self.localization.get_text("toc_range_selection"))
        range_layout = QVBoxLayout(range_group)

        # Instructions
        instructions = QLabel(self.localization.get_text("toc_range_instructions"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #666; margin-bottom: 15px;")
        range_layout.addWidget(instructions)

        # Start page
        start_layout = QHBoxLayout()
        start_layout.addWidget(QLabel(self.localization.get_text("start_page")))
        self.start_spin = QSpinBox()
        self.start_spin.setRange(1, self.total_pages)
        self.start_spin.setValue(self.start_page)
        self.start_spin.valueChanged.connect(self.update_range)
        # Force LTR for numbers in Arabic locale
        self.start_spin.setLayoutDirection(Qt.LeftToRight)
        self.start_spin.setStyleSheet("QSpinBox { text-align: center; }")
        start_layout.addWidget(self.start_spin)
        start_layout.addStretch()
        range_layout.addLayout(start_layout)

        # End page
        end_layout = QHBoxLayout()
        end_layout.addWidget(QLabel(self.localization.get_text("end_page")))
        self.end_spin = QSpinBox()
        self.end_spin.setRange(1, self.total_pages)
        self.end_spin.setValue(self.end_page)
        self.end_spin.valueChanged.connect(self.update_range)
        # Force LTR for numbers in Arabic locale
        self.end_spin.setLayoutDirection(Qt.LeftToRight)
        self.end_spin.setStyleSheet("QSpinBox { text-align: center; }")
        end_layout.addWidget(self.end_spin)
        end_layout.addStretch()
        range_layout.addLayout(end_layout)

        # Quick selection buttons
        quick_layout = QHBoxLayout()
        self.btn_set_start = QPushButton(self.localization.get_text("set_as_start"))
        self.btn_set_start.clicked.connect(self.set_current_as_start)
        quick_layout.addWidget(self.btn_set_start)

        self.btn_set_end = QPushButton(self.localization.get_text("set_as_end"))
        self.btn_set_end.clicked.connect(self.set_current_as_end)
        quick_layout.addWidget(self.btn_set_end)

        range_layout.addLayout(quick_layout)

        # Range info
        self.range_info = QLabel()
        self.range_info.setStyleSheet("background-color: #E3F2FD; padding: 10px; border-radius: 4px; margin-top: 10px;")
        range_layout.addWidget(self.range_info)

        range_layout.addStretch()
        content_layout.addWidget(range_group)

        layout.addLayout(content_layout)

        # Dialog buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel"))
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        ok_btn = QPushButton(self.localization.get_text("extract_selected_pages"))
        ok_btn.clicked.connect(self.accept)
        ok_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 8px 16px; }")
        button_layout.addWidget(ok_btn)

        layout.addLayout(button_layout)

        self.update_range_info()

    def update_preview(self):
        """Update PDF preview with proper A4 aspect ratio"""
        try:
            page = self.pdf_doc[self.current_page]
            # Use proper scaling to maintain A4 proportions
            page_rect = page.rect
            # Calculate scale to fit preview area while maintaining aspect ratio
            preview_width = self.pdf_preview.width() - 20  # Account for padding
            preview_height = self.pdf_preview.height() - 20

            scale_x = preview_width / page_rect.width
            scale_y = preview_height / page_rect.height
            scale = min(scale_x, scale_y, 1.5)  # Limit maximum scale

            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            pixmap = QPixmap()
            pixmap.loadFromData(img_data)

            if not pixmap.isNull():
                self.pdf_preview.setPixmap(pixmap)

            # Update page controls
            self.page_input.setValue(self.current_page + 1)
            self.page_label.setText(f"/ {self.total_pages}")

            # Update page info
            page_status = ""
            if self.current_page + 1 == self.start_page:
                page_status = " (START)"
            elif self.current_page + 1 == self.end_page:
                page_status = " (END)"
            elif self.start_page <= self.current_page + 1 <= self.end_page:
                page_status = " (SELECTED)"

            self.page_info.setText(f"Page {self.current_page + 1}{page_status}")

            # Update navigation buttons
            self.btn_first.setEnabled(self.current_page > 0)
            self.btn_prev.setEnabled(self.current_page > 0)
            self.btn_next.setEnabled(self.current_page < self.total_pages - 1)
            self.btn_last.setEnabled(self.current_page < self.total_pages - 1)

        except Exception as e:
            self.pdf_preview.setText(f"Preview error: {str(e)}")
            self.page_info.setText(self.localization.get_text("error_loading_page"))

    def first_page(self):
        """Go to first page"""
        self.current_page = 0
        self.update_preview()

    def prev_page(self):
        """Go to previous page"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_preview()

    def next_page(self):
        """Go to next page"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_preview()

    def last_page(self):
        """Go to last page"""
        self.current_page = self.total_pages - 1
        self.update_preview()

    def goto_page(self, page_num):
        """Go to specific page"""
        if 1 <= page_num <= self.total_pages:
            self.current_page = page_num - 1
            self.update_preview()

    def set_current_as_start(self):
        """Set current page as start page"""
        self.start_spin.setValue(self.current_page + 1)

    def set_current_as_end(self):
        """Set current page as end page"""
        self.end_spin.setValue(self.current_page + 1)

    def update_range(self):
        """Update range when spinboxes change"""
        self.start_page = self.start_spin.value()
        self.end_page = self.end_spin.value()

        # Ensure start <= end
        if self.start_page > self.end_page:
            if self.sender() == self.start_spin:
                self.end_spin.setValue(self.start_page)
                self.end_page = self.start_page
            else:
                self.start_spin.setValue(self.end_page)
                self.start_page = self.end_page

        self.update_range_info()

    def update_range_info(self):
        """Update range information display"""
        page_count = self.end_page - self.start_page + 1
        self.range_info.setText(
            f"{self.localization.get_text('selected_range')}: {self.start_page}-{self.end_page}\n"
            f"{self.localization.get_text('total_pages_selected')}: {page_count}"
        )

    def get_selected_range(self):
        """Get the selected page range"""
        return self.start_page, self.end_page


class TOCPageSelectionDialog(QDialog):
    """Dialog for selecting TOC page range with enhanced PDF preview for Bookmark Manager"""

    def __init__(self, pdf_doc, total_pages, localization, parent=None):
        super().__init__(parent)
        self.pdf_doc = pdf_doc
        self.total_pages = total_pages
        self.localization = localization
        # Start near the end where TOC usually is
        self.current_page = max(0, total_pages - 10)  # Start 10 pages from end
        self.start_page = max(1, total_pages - 9)  # Default start near end
        self.end_page = total_pages  # Default end at last page

        self.setWindowTitle(self.localization.get_text("toc_page_selection"))
        self.setModal(True)
        self.resize(900, 750)  # Reduced width from 1200 to 900

        # Set app icon if parent has one
        if hasattr(parent, 'get_app_icon'):
            self.setWindowIcon(parent.get_app_icon())

        self.init_ui()
        self.update_preview()

    def init_ui(self):
        """Initialize dialog UI"""
        layout = QVBoxLayout(self)

        # Title and description
        title = QLabel(self.localization.get_text("toc_page_selection"))
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px; color: #9C27B0;")
        layout.addWidget(title)

        description = QLabel(self.localization.get_text("toc_page_selection_desc"))
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; margin-bottom: 15px; font-size: 14px;")
        layout.addWidget(description)

        # Main content
        content_layout = QHBoxLayout()

        # Left side - PDF preview (larger)
        preview_group = QGroupBox(self.localization.get_text("pdf_preview"))
        preview_layout = QVBoxLayout(preview_group)

        # Enhanced navigation controls
        nav_layout = QHBoxLayout()

        # First/Previous buttons
        self.btn_first = QPushButton("⏮")
        self.btn_first.setToolTip(self.localization.get_text("first_page"))
        self.btn_first.clicked.connect(self.first_page)
        self.btn_first.setFixedSize(40, 30)
        nav_layout.addWidget(self.btn_first)

        self.btn_prev = QPushButton("◀")
        self.btn_prev.setToolTip(self.localization.get_text("previous_page"))
        self.btn_prev.clicked.connect(self.prev_page)
        self.btn_prev.setFixedSize(40, 30)
        nav_layout.addWidget(self.btn_prev)

        # Page input and display
        nav_layout.addWidget(QLabel(self.localization.get_text("page_label")))
        self.page_input = QSpinBox()
        self.page_input.setRange(1, self.total_pages)
        self.page_input.setValue(self.current_page + 1)
        self.page_input.valueChanged.connect(self.goto_page)
        self.page_input.setLayoutDirection(Qt.LeftToRight)
        self.page_input.setStyleSheet("QSpinBox { text-align: center; min-width: 70px; }")
        nav_layout.addWidget(self.page_input)

        self.page_label = QLabel()
        self.page_label.setAlignment(Qt.AlignCenter)
        nav_layout.addWidget(self.page_label)

        # Next/Last buttons
        self.btn_next = QPushButton("▶")
        self.btn_next.setToolTip(self.localization.get_text("next_page"))
        self.btn_next.clicked.connect(self.next_page)
        self.btn_next.setFixedSize(40, 30)
        nav_layout.addWidget(self.btn_next)

        self.btn_last = QPushButton("⏭")
        self.btn_last.setToolTip(self.localization.get_text("last_page"))
        self.btn_last.clicked.connect(self.last_page)
        self.btn_last.setFixedSize(40, 30)
        nav_layout.addWidget(self.btn_last)

        nav_layout.addStretch()
        preview_layout.addLayout(nav_layout)

        # PDF preview (reduced width)
        self.pdf_preview = QLabel()
        self.pdf_preview.setAlignment(Qt.AlignCenter)
        self.pdf_preview.setMinimumSize(400, 600)  # Reduced width from 600 to 400
        self.pdf_preview.setMaximumSize(450, 650)  # Added maximum size constraint
        self.pdf_preview.setStyleSheet("border: 2px solid #9C27B0; background-color: white; border-radius: 4px;")
        self.pdf_preview.setScaledContents(True)
        preview_layout.addWidget(self.pdf_preview)

        # Current page info
        self.page_info = QLabel()
        self.page_info.setAlignment(Qt.AlignCenter)
        self.page_info.setStyleSheet("color: #666; font-size: 14px; padding: 5px; font-weight: bold;")
        preview_layout.addWidget(self.page_info)

        content_layout.addWidget(preview_group, 2)  # Give more space to preview

        # Right side - Range selection (smaller)
        range_group = QGroupBox(self.localization.get_text("toc_range_selection"))
        range_layout = QVBoxLayout(range_group)

        # Instructions
        instructions = QLabel(self.localization.get_text("navigate_and_select"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #666; margin-bottom: 15px; font-size: 13px;")
        range_layout.addWidget(instructions)

        # Start page
        start_layout = QHBoxLayout()
        start_layout.addWidget(QLabel(self.localization.get_text("start_page")))
        self.start_spin = QSpinBox()
        self.start_spin.setRange(1, self.total_pages)
        self.start_spin.setValue(self.start_page)
        self.start_spin.valueChanged.connect(self.update_range)
        self.start_spin.setLayoutDirection(Qt.LeftToRight)
        self.start_spin.setStyleSheet("QSpinBox { text-align: center; }")
        start_layout.addWidget(self.start_spin)
        start_layout.addStretch()
        range_layout.addLayout(start_layout)

        # End page
        end_layout = QHBoxLayout()
        end_layout.addWidget(QLabel(self.localization.get_text("end_page")))
        self.end_spin = QSpinBox()
        self.end_spin.setRange(1, self.total_pages)
        self.end_spin.setValue(self.end_page)
        self.end_spin.valueChanged.connect(self.update_range)
        self.end_spin.setLayoutDirection(Qt.LeftToRight)
        self.end_spin.setStyleSheet("QSpinBox { text-align: center; }")
        end_layout.addWidget(self.end_spin)
        end_layout.addStretch()
        range_layout.addLayout(end_layout)

        # Quick selection buttons
        quick_layout = QVBoxLayout()
        self.btn_set_start = QPushButton(self.localization.get_text("set_as_start"))
        self.btn_set_start.clicked.connect(self.set_current_as_start)
        self.btn_set_start.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 8px; }")
        quick_layout.addWidget(self.btn_set_start)

        self.btn_set_end = QPushButton(self.localization.get_text("set_as_end"))
        self.btn_set_end.clicked.connect(self.set_current_as_end)
        self.btn_set_end.setStyleSheet("QPushButton { background-color: #FF9800; color: white; font-weight: bold; padding: 8px; }")
        quick_layout.addWidget(self.btn_set_end)

        range_layout.addLayout(quick_layout)

        # Range info
        self.range_info = QLabel()
        self.range_info.setStyleSheet("background-color: #E8F5E8; padding: 10px; border-radius: 4px; margin-top: 10px; border: 1px solid #4CAF50;")
        range_layout.addWidget(self.range_info)

        range_layout.addStretch()
        content_layout.addWidget(range_group, 1)  # Give less space to controls

        layout.addLayout(content_layout)

        # Dialog buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel"))
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setStyleSheet("QPushButton { padding: 8px 16px; }")
        button_layout.addWidget(cancel_btn)

        ok_btn = QPushButton(self.localization.get_text("ok"))
        ok_btn.clicked.connect(self.accept)
        ok_btn.setStyleSheet("QPushButton { background-color: #9C27B0; color: white; font-weight: bold; padding: 8px 16px; }")
        button_layout.addWidget(ok_btn)

        layout.addLayout(button_layout)

        self.update_range_info()

    def update_preview(self):
        """Update PDF preview with proper A4 aspect ratio"""
        try:
            page = self.pdf_doc[self.current_page]
            # Calculate proper scaling to maintain A4 proportions
            page_rect = page.rect
            preview_width = self.pdf_preview.width() - 20  # Account for padding
            preview_height = self.pdf_preview.height() - 20

            scale_x = preview_width / page_rect.width
            scale_y = preview_height / page_rect.height
            scale = min(scale_x, scale_y, 1.2)  # Limit maximum scale

            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            pixmap = QPixmap()
            pixmap.loadFromData(img_data)

            if not pixmap.isNull():
                self.pdf_preview.setPixmap(pixmap)

            # Update page controls
            self.page_input.setValue(self.current_page + 1)
            self.page_label.setText(f"/ {self.total_pages}")

            # Update page info with status
            page_status = ""
            if self.current_page + 1 == self.start_page:
                page_status = " (START)"
            elif self.current_page + 1 == self.end_page:
                page_status = " (END)"
            elif self.start_page <= self.current_page + 1 <= self.end_page:
                page_status = " (SELECTED)"

            self.page_info.setText(f"Page {self.current_page + 1}{page_status}")

            # Update navigation buttons
            self.btn_first.setEnabled(self.current_page > 0)
            self.btn_prev.setEnabled(self.current_page > 0)
            self.btn_next.setEnabled(self.current_page < self.total_pages - 1)
            self.btn_last.setEnabled(self.current_page < self.total_pages - 1)

        except Exception as e:
            self.pdf_preview.setText(f"Preview error: {str(e)}")
            self.page_info.setText(self.localization.get_text("error_loading_page"))

    def first_page(self):
        """Go to first page"""
        self.current_page = 0
        self.update_preview()

    def prev_page(self):
        """Go to previous page"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_preview()

    def next_page(self):
        """Go to next page"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_preview()

    def last_page(self):
        """Go to last page"""
        self.current_page = self.total_pages - 1
        self.update_preview()

    def goto_page(self, page_num):
        """Go to specific page"""
        if 1 <= page_num <= self.total_pages:
            self.current_page = page_num - 1
            self.update_preview()

    def set_current_as_start(self):
        """Set current page as start page"""
        self.start_spin.setValue(self.current_page + 1)

    def set_current_as_end(self):
        """Set current page as end page"""
        self.end_spin.setValue(self.current_page + 1)

    def update_range(self):
        """Update range when spinboxes change"""
        self.start_page = self.start_spin.value()
        self.end_page = self.end_spin.value()

        # Ensure start <= end
        if self.start_page > self.end_page:
            if self.sender() == self.start_spin:
                self.end_spin.setValue(self.start_page)
                self.end_page = self.start_page
            else:
                self.start_spin.setValue(self.end_page)
                self.start_page = self.end_page

        self.update_range_info()
        self.update_preview()  # Update preview to show new status

    def update_range_info(self):
        """Update range information display"""
        page_count = self.end_page - self.start_page + 1
        self.range_info.setText(
            f"<b>{self.localization.get_text('selected_range')}:</b> {self.start_page}-{self.end_page}<br>"
            f"<b>{self.localization.get_text('total_pages_selected')}:</b> {page_count}"
        )

    def get_selected_range(self):
        """Get the selected page range"""
        return self.start_page, self.end_page


class BookmarkCorrectionDialog(QDialog):
    """Dialog for manual bookmark page correction"""

    def __init__(self, bookmark_title, bookmark_page, current_pdf_page, total_pages, localization, parent=None):
        super().__init__(parent)
        self.bookmark_title = bookmark_title

        # Set app icon if parent has one
        if hasattr(parent, 'get_app_icon'):
            self.setWindowIcon(parent.get_app_icon())
        self.bookmark_page = bookmark_page
        self.current_pdf_page = current_pdf_page
        self.total_pages = total_pages
        self.localization = localization
        self.correct_page = None

        self.setWindowTitle(self.localization.get_text("manual_correction"))
        self.setModal(True)
        self.resize(500, 300)

        self.init_ui()

    def init_ui(self):
        """Initialize dialog UI"""
        layout = QVBoxLayout(self)

        # Title
        title = QLabel(self.localization.get_text("manual_correction"))
        title.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(title)

        # Bookmark info
        info_group = QGroupBox("Bookmark Information")
        info_layout = QVBoxLayout(info_group)

        bookmark_info = QLabel(f"<b>Title:</b> {self.bookmark_title}")
        bookmark_info.setWordWrap(True)
        info_layout.addWidget(bookmark_info)

        page_info = QLabel(f"<b>Original Page:</b> {self.bookmark_page}")
        info_layout.addWidget(page_info)

        current_info = QLabel(f"<b>{self.localization.get_text('current_viewing')}:</b> {self.current_pdf_page}")
        info_layout.addWidget(current_info)

        layout.addWidget(info_group)

        # Correction input
        correction_group = QGroupBox(self.localization.get_text("set_correct_page"))
        correction_layout = QVBoxLayout(correction_group)

        instruction = QLabel(f"{self.localization.get_text('bookmark_should_be')}:")
        correction_layout.addWidget(instruction)

        page_layout = QHBoxLayout()
        self.page_spin = QSpinBox()
        self.page_spin.setRange(1, self.total_pages)
        self.page_spin.setValue(self.current_pdf_page)
        self.page_spin.setLayoutDirection(Qt.LeftToRight)
        self.page_spin.setStyleSheet("QSpinBox { text-align: center; font-size: 14px; padding: 5px; }")
        page_layout.addWidget(QLabel(self.localization.get_text("page_label")))
        page_layout.addWidget(self.page_spin)
        page_layout.addStretch()

        correction_layout.addLayout(page_layout)

        # Quick buttons
        quick_layout = QHBoxLayout()

        btn_current = QPushButton(f"Use Current Page ({self.current_pdf_page})")
        btn_current.clicked.connect(lambda: self.page_spin.setValue(self.current_pdf_page))
        quick_layout.addWidget(btn_current)

        btn_original = QPushButton(f"Keep Original ({self.bookmark_page})")
        btn_original.clicked.connect(lambda: self.page_spin.setValue(self.bookmark_page))
        quick_layout.addWidget(btn_original)

        correction_layout.addLayout(quick_layout)

        # Offset preview
        self.offset_preview = QLabel()
        self.offset_preview.setStyleSheet("background-color: #E3F2FD; padding: 10px; border-radius: 4px; margin-top: 10px;")
        correction_layout.addWidget(self.offset_preview)

        # Connect to update preview
        self.page_spin.valueChanged.connect(self.update_offset_preview)
        self.update_offset_preview()

        layout.addWidget(correction_group)

        # Dialog buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton(self.localization.get_text("cancel"))
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        apply_btn = QPushButton(self.localization.get_text("apply_correction"))
        apply_btn.clicked.connect(self.accept)
        apply_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 8px 16px; }")
        button_layout.addWidget(apply_btn)

        layout.addLayout(button_layout)

    def update_offset_preview(self):
        """Update offset preview"""
        correct_page = self.page_spin.value()
        offset = correct_page - self.bookmark_page

        if offset == 0:
            self.offset_preview.setText(self.localization.get_text("no_offset_needed"))
        else:
            self.offset_preview.setText(
                f"Offset: {offset:+d}\n"
                f"This will adjust ALL bookmarks by {offset:+d} pages"
            )

    def get_correct_page(self):
        """Get the correct page number"""
        return self.page_spin.value()


class PageOperationsTab(QWidget):
    """Page operations tab for inserting/deleting pages"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""
        self.source_pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize page operations tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("page_operations_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #1976D2; padding: 8px; background-color: #E3F2FD; border-radius: 4px;")
        layout.addWidget(instructions)

        # Unified PDF file selection at top
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_placeholder"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        layout.addWidget(file_group)

        # Progress bar for operations
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # Operation tabs
        self.operation_tabs = QTabWidget()

        # Extract tab
        extract_tab = QWidget()
        extract_layout = QVBoxLayout(extract_tab)

        # Pages to extract
        extract_pages_layout = QHBoxLayout()
        extract_pages_layout.addWidget(QLabel(self.localization.get_text("pages_to_extract")))
        self.pages_to_extract = QLineEdit()
        self.pages_to_extract.setPlaceholderText(self.localization.get_text("page_range_examples"))
        extract_pages_layout.addWidget(self.pages_to_extract)
        extract_layout.addLayout(extract_pages_layout)

        # Page range help
        help_text = QLabel(self.localization.get_text("page_range_help"))
        help_text.setWordWrap(True)
        help_text.setStyleSheet("color: #666; font-size: 11px; padding: 5px; background-color: #F5F5F5; border-radius: 3px;")
        extract_layout.addWidget(help_text)

        # Extract button
        self.btn_extract_pages = QPushButton(self.localization.get_text("extract_pages_btn"))
        self.btn_extract_pages.clicked.connect(self.extract_pages)
        self.btn_extract_pages.setEnabled(False)
        self.btn_extract_pages.setStyleSheet("QPushButton { background-color: #2196F3; color: white; font-weight: bold; padding: 10px; }")
        extract_layout.addWidget(self.btn_extract_pages)
        extract_layout.addStretch()

        self.operation_tabs.addTab(extract_tab, self.localization.get_text("extract_tab"))

        # Delete tab
        delete_tab = QWidget()
        delete_layout = QVBoxLayout(delete_tab)

        # Pages to delete
        delete_pages_layout = QHBoxLayout()
        delete_pages_layout.addWidget(QLabel(self.localization.get_text("pages_to_delete")))
        self.pages_to_delete = QLineEdit()
        self.pages_to_delete.setPlaceholderText(self.localization.get_text("page_range_examples"))
        delete_pages_layout.addWidget(self.pages_to_delete)
        delete_layout.addLayout(delete_pages_layout)

        # Page range help
        help_text2 = QLabel(self.localization.get_text("page_range_help"))
        help_text2.setWordWrap(True)
        help_text2.setStyleSheet("color: #666; font-size: 11px; padding: 5px; background-color: #F5F5F5; border-radius: 3px;")
        delete_layout.addWidget(help_text2)

        # Delete button
        self.btn_delete_pages = QPushButton(self.localization.get_text("delete_pages_btn"))
        self.btn_delete_pages.clicked.connect(self.delete_pages)
        self.btn_delete_pages.setEnabled(False)
        self.btn_delete_pages.setStyleSheet("QPushButton { background-color: #F44336; color: white; font-weight: bold; padding: 10px; }")
        delete_layout.addWidget(self.btn_delete_pages)
        delete_layout.addStretch()

        self.operation_tabs.addTab(delete_tab, self.localization.get_text("delete_tab"))

        # Insert blank tab
        insert_tab = QWidget()
        insert_layout = QVBoxLayout(insert_tab)

        # Blank page options
        blank_options_layout = QHBoxLayout()
        blank_options_layout.addWidget(QLabel(self.localization.get_text("insert_at_position")))
        self.blank_position = QSpinBox()
        self.blank_position.setMinimum(1)
        self.blank_position.setValue(1)
        blank_options_layout.addWidget(self.blank_position)

        blank_options_layout.addWidget(QLabel(self.localization.get_text("number_of_pages")))
        self.blank_count = QSpinBox()
        self.blank_count.setMinimum(1)
        self.blank_count.setMaximum(100)
        self.blank_count.setValue(1)
        blank_options_layout.addWidget(self.blank_count)
        blank_options_layout.addStretch()
        insert_layout.addLayout(blank_options_layout)

        # Insert blank pages button
        self.btn_insert_blank = QPushButton(self.localization.get_text("insert_blank_btn"))
        self.btn_insert_blank.clicked.connect(self.insert_blank_pages)
        self.btn_insert_blank.setEnabled(False)
        self.btn_insert_blank.setStyleSheet("QPushButton { background-color: #FF9800; color: white; font-weight: bold; padding: 10px; }")
        insert_layout.addWidget(self.btn_insert_blank)
        insert_layout.addStretch()

        self.operation_tabs.addTab(insert_tab, self.localization.get_text("insert_tab"))

        # Merge PDFs tab
        merge_tab = QWidget()
        merge_layout = QVBoxLayout(merge_tab)

        # Instructions for merge
        merge_instructions = QLabel(self.localization.get_text("merge_pdfs_desc"))
        merge_instructions.setWordWrap(True)
        merge_instructions.setStyleSheet("color: #4CAF50; padding: 8px; background-color: #E8F5E8; border-radius: 4px;")
        merge_layout.addWidget(merge_instructions)

        # File list for merge
        files_group = QGroupBox(self.localization.get_text("selected_files"))
        files_layout = QVBoxLayout(files_group)

        # File list widget with drag and drop
        self.merge_file_list = QListWidget()
        self.merge_file_list.setDragDropMode(QListWidget.InternalMove)
        self.merge_file_list.setDefaultDropAction(Qt.MoveAction)
        self.merge_file_list.setMaximumHeight(200)
        files_layout.addWidget(self.merge_file_list)

        # File management buttons
        file_buttons_layout = QHBoxLayout()

        self.btn_add_files = QPushButton(self.localization.get_text("add_files"))
        self.btn_add_files.clicked.connect(self.add_merge_files)
        self.btn_add_files.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 8px; }")
        file_buttons_layout.addWidget(self.btn_add_files)

        self.btn_remove_file = QPushButton(self.localization.get_text("remove_selected"))
        self.btn_remove_file.clicked.connect(self.remove_selected_file)
        self.btn_remove_file.setStyleSheet("QPushButton { background-color: #F44336; color: white; font-weight: bold; padding: 8px; }")
        file_buttons_layout.addWidget(self.btn_remove_file)

        self.btn_clear_files = QPushButton(self.localization.get_text("clear_all"))
        self.btn_clear_files.clicked.connect(self.clear_merge_files)
        self.btn_clear_files.setStyleSheet("QPushButton { background-color: #FF9800; color: white; font-weight: bold; padding: 8px; }")
        file_buttons_layout.addWidget(self.btn_clear_files)

        file_buttons_layout.addStretch()
        files_layout.addLayout(file_buttons_layout)

        merge_layout.addWidget(files_group)

        # Merge options
        options_group = QGroupBox(self.localization.get_text("merge_options"))
        options_layout = QVBoxLayout(options_group)

        # File order info
        order_info = QLabel(self.localization.get_text("drag_to_reorder"))
        order_info.setWordWrap(True)
        order_info.setStyleSheet("color: #666; font-style: italic; padding: 5px;")
        options_layout.addWidget(order_info)

        merge_layout.addWidget(options_group)

        # Merge button
        self.btn_merge_pdfs = QPushButton(self.localization.get_text("merge_pdfs_btn"))
        self.btn_merge_pdfs.clicked.connect(self.merge_selected_pdfs)
        self.btn_merge_pdfs.setEnabled(False)
        self.btn_merge_pdfs.setStyleSheet("QPushButton { background-color: #9C27B0; color: white; font-weight: bold; padding: 12px; font-size: 14px; }")
        merge_layout.addWidget(self.btn_merge_pdfs)

        merge_layout.addStretch()

        self.operation_tabs.addTab(merge_tab, self.localization.get_text("merge_tab"))


        # Rotate Pages tab (moved from Page Editing)
        rotate_tab = QWidget()
        rotate_layout = QVBoxLayout(rotate_tab)

        # Rotation options
        rot_opts = QHBoxLayout()
        rot_opts.addWidget(QLabel(self.localization.get_text("pages_range_label") if hasattr(self.localization, 'get_text') else "Pages:"))
        self.rotate_page_range = QLineEdit()
        self.rotate_page_range.setPlaceholderText(self.localization.get_text("page_range_examples"))
        rot_opts.addWidget(self.rotate_page_range)

        rot_opts.addWidget(QLabel(self.localization.get_text("rotation")))
        self.rotation_combo_ops = QComboBox()
        self.rotation_combo_ops.addItems([
            self.localization.get_text("rotate_90"),
            self.localization.get_text("rotate_180"),
            self.localization.get_text("rotate_270"),
        ])
        rot_opts.addStretch()
        rotate_layout.addLayout(rot_opts)

        self.btn_rotate_pages_ops = QPushButton(self.localization.get_text("rotate_pages_btn"))
        self.btn_rotate_pages_ops.clicked.connect(self.rotate_pages_ops)
        self.btn_rotate_pages_ops.setEnabled(False)
        rotate_layout.addWidget(self.btn_rotate_pages_ops)
        rotate_layout.addStretch()
        self.operation_tabs.addTab(rotate_tab, self.localization.get_text("rotate_pages"))

        # Crop Pages tab (new UI, uses existing PDFOperations.crop_pages)
        crop_tab = QWidget()
        crop_layout = QVBoxLayout(crop_tab)

        crop_row1 = QHBoxLayout()
        crop_row1.addWidget(QLabel(self.localization.get_text("pages_range_label") if hasattr(self.localization, 'get_text') else "Pages:"))
        self.crop_page_range = QLineEdit()
        self.crop_page_range.setPlaceholderText(self.localization.get_text("page_range_examples"))
        crop_row1.addWidget(self.crop_page_range)
        crop_layout.addLayout(crop_row1)

        # Crop box inputs: x0, y0, x1, y1
        crop_row2 = QHBoxLayout()
        crop_row2.addWidget(QLabel(self.localization.get_text("crop_box") if hasattr(self.localization, 'get_text') else "Crop Box (x0,y0,x1,y1):"))
        self.crop_x0 = QSpinBox(); self.crop_x0.setRange(0, 10000); self.crop_x0.setValue(0)
        self.crop_y0 = QSpinBox(); self.crop_y0.setRange(0, 10000); self.crop_y0.setValue(0)
        self.crop_x1 = QSpinBox(); self.crop_x1.setRange(0, 10000); self.crop_x1.setValue(595)
        self.crop_y1 = QSpinBox(); self.crop_y1.setRange(0, 10000); self.crop_y1.setValue(842)
        for w in (self.crop_x0, self.crop_y0, self.crop_x1, self.crop_y1):
            w.setAccelerated(True)
        crop_row2.addWidget(QLabel("x0")); crop_row2.addWidget(self.crop_x0)
        crop_row2.addWidget(QLabel("y0")); crop_row2.addWidget(self.crop_y0)
        crop_row2.addWidget(QLabel("x1")); crop_row2.addWidget(self.crop_x1)
        crop_row2.addWidget(QLabel("y1")); crop_row2.addWidget(self.crop_y1)
        crop_row2.addStretch()
        crop_layout.addLayout(crop_row2)

        self.btn_crop_pages_ops = QPushButton(self.localization.get_text("crop_pages_btn") if hasattr(self.localization, 'get_text') else "Crop Pages")
        self.btn_crop_pages_ops.clicked.connect(self.crop_pages_ops)
        self.btn_crop_pages_ops.setEnabled(False)
        crop_layout.addWidget(self.btn_crop_pages_ops)
        crop_layout.addStretch()
        self.operation_tabs.addTab(crop_tab, self.localization.get_text("crop_pages") if hasattr(self.localization, 'get_text') else "Crop")

        # Add Margins tab (moved from Page Editing)
        margins_tab = QWidget()
        margins_layout = QVBoxLayout(margins_tab)

        margins_row = QHBoxLayout()
        margins_row.addWidget(QLabel(self.localization.get_text("pages_range_label") if hasattr(self.localization, 'get_text') else "Pages:"))
        self.margins_page_range = QLineEdit()
        self.margins_page_range.setPlaceholderText(self.localization.get_text("page_range_examples"))
        margins_row.addWidget(self.margins_page_range)

        margins_row.addWidget(QLabel(self.localization.get_text("margin_size")))
        self.margin_size_ops = QSpinBox(); self.margin_size_ops.setRange(0, 500); self.margin_size_ops.setValue(50)
        margins_row.addWidget(self.margin_size_ops)
        margins_row.addStretch()
        margins_layout.addLayout(margins_row)

        self.btn_add_margins_ops = QPushButton(self.localization.get_text("add_margins_btn"))
        self.btn_add_margins_ops.clicked.connect(self.add_margins_ops)
        self.btn_add_margins_ops.setEnabled(False)
        margins_layout.addWidget(self.btn_add_margins_ops)
        margins_layout.addStretch()
        self.operation_tabs.addTab(margins_tab, self.localization.get_text("add_white_margins"))

        # Initialize merge file list
        self.merge_files = []

        layout.addWidget(self.operation_tabs)

        # Add stretch to push everything to top
        layout.addStretch()

    def parse_page_numbers(self, page_string: str) -> List[int]:
        """Parse page number string like '1,3,5-10' into list of integers"""
        pages = []

        try:
            parts = page_string.split(',')
            for part in parts:
                part = part.strip()
                if '-' in part:
                    # Handle range like "5-10"
                    start, end = part.split('-')
                    start, end = int(start.strip()), int(end.strip())
                    pages.extend(range(start, end + 1))
                else:
                    # Handle single page
                    pages.append(int(part))

            return sorted(list(set(pages)))  # Remove duplicates and sort

        except ValueError:
            return []

    def select_pdf_file(self):
        """Select PDF file for all operations"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            self.localization.get_text("pdf_files_filter")
        )
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            # Enable all operation buttons
            self.btn_extract_pages.setEnabled(True)
            self.btn_delete_pages.setEnabled(True)
            self.btn_insert_blank.setEnabled(True)
            if hasattr(self, 'btn_rotate_pages_ops'):
                self.btn_rotate_pages_ops.setEnabled(True)
            if hasattr(self, 'btn_crop_pages_ops'):
                self.btn_crop_pages_ops.setEnabled(True)
            if hasattr(self, 'btn_add_margins_ops'):
                self.btn_add_margins_ops.setEnabled(True)

    def insert_blank_pages(self):
        """Insert blank pages into PDF"""
        if not self.pdf_path:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("file_not_selected")
            )
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_with_blank_pages.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_pdf_with_blank"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Show progress
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)  # Indeterminate progress
            QApplication.processEvents()

            position = self.blank_position.value()
            count = self.blank_count.value()

            success = self.pdf_ops.insert_blank_pages(
                self.pdf_path, output_path, position, count
            )

            # Hide progress
            self.progress_bar.setVisible(False)

            if success:
                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    f"{self.localization.get_text('blank_pages_inserted_successfully')}\n\n"
                    f"{self.localization.get_text('position')}: {position}, {self.localization.get_text('count')}: {count}\n"
                    f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                )
            else:
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    self.localization.get_text("operation_failed")
                )

    def delete_pages(self):
        """Delete pages from PDF"""
        if not self.pdf_path:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("file_not_selected")
            )
            return

        pages_text = self.pages_to_delete.text().strip()
        if not pages_text:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("no_pages_specified")
            )
            return

        pages = self.parse_page_numbers(pages_text)
        if not pages:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("invalid_page_range")
            )
            return

        # Confirm deletion
        reply = QMessageBox.question(
            self,
            self.localization.get_text("confirm_deletion"),
            f"{self.localization.get_text('confirm_delete_pages')}: {len(pages)}\n\n"
            f"{self.localization.get_text('pages')}: {', '.join(map(str, pages))}",
            QMessageBox.Yes | QMessageBox.No
        )

        if reply != QMessageBox.Yes:
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_pages_deleted.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_pdf_deleted_pages"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Show progress
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)  # Indeterminate progress
            QApplication.processEvents()

            success = self.pdf_ops.delete_pages(self.pdf_path, pages, output_path)

            # Hide progress
            self.progress_bar.setVisible(False)

            if success:
                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    f"{self.localization.get_text('pages_deleted_successfully')}: {len(pages)}\n\n"
                    f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                )
            else:
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    self.localization.get_text("operation_failed")
                )

    def extract_pages(self):
        """Extract pages to new PDF"""
        if not self.pdf_path:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("file_not_selected")
            )
            return

        pages_text = self.pages_to_extract.text().strip()
        if not pages_text:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("no_pages_specified")
            )
            return

        pages = self.parse_page_numbers(pages_text)
        if not pages:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("invalid_page_range")
            )
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_extracted_pages.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_extracted_pages"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Show progress
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)  # Indeterminate progress
            QApplication.processEvents()

            success = self.pdf_ops.extract_pages(self.pdf_path, pages, output_path)

            # Hide progress
            self.progress_bar.setVisible(False)

            if success:
                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    f"{self.localization.get_text('pages_extracted_successfully')}: {len(pages)}\n\n"
                    f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                )
            else:
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    self.localization.get_text("operation_failed")
                )

    def add_merge_files(self):
        """Add files to merge list"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            self.localization.get_text("select_pdfs_merge"),
            "",
            self.localization.get_text("pdf_files_filter")
        )

        if files:
            for file_path in files:
                if file_path not in self.merge_files:
                    self.merge_files.append(file_path)
                    item = QListWidgetItem(os.path.basename(file_path))
                    item.setData(Qt.UserRole, file_path)
                    item.setToolTip(file_path)
                    self.merge_file_list.addItem(item)

            # Enable merge button if we have files
            self.btn_merge_pdfs.setEnabled(len(self.merge_files) >= 2)

    def remove_selected_file(self):
        """Remove selected file from merge list"""
        current_row = self.merge_file_list.currentRow()
        if current_row >= 0:
            item = self.merge_file_list.takeItem(current_row)
            if item:
                file_path = item.data(Qt.UserRole)
                if file_path in self.merge_files:
                    self.merge_files.remove(file_path)
        # Update merge button state
        self.btn_merge_pdfs.setEnabled(len(self.merge_files) >= 2)

    def rotate_pages_ops(self):
        """Rotate selected pages using PDFOperations.rotate_pages."""
        if not self.pdf_path:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("file_not_selected"))
            return
        pages = self.parse_page_numbers(self.rotate_page_range.text().strip() or "")
        if not pages:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("invalid_page_range"))
            return
        angle_text = self.rotation_combo_ops.currentText()
        angle = 0
        if "90" in angle_text:
            angle = 90
        elif "180" in angle_text:
            angle = 180
        elif "270" in angle_text:
            angle = 270
        else:
            angle = 90

        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_rotated.pdf')
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_rotated_pdf"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )
        if not output_path:
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        QApplication.processEvents()
        success = self.pdf_ops.rotate_pages(self.pdf_path, pages, angle, output_path)
        self.progress_bar.setVisible(False)
        if success:
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("pages_rotated_successfully"))
        else:
            QMessageBox.critical(self, self.localization.get_text("error"), self.localization.get_text("operation_failed"))

    def add_margins_ops(self):
        """Add white margins around selected pages using PDFOperations.add_margins."""
        if not self.pdf_path:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("file_not_selected"))
            return
        pages = self.parse_page_numbers(self.margins_page_range.text().strip() or "")
        if not pages:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("invalid_page_range"))
            return
        margin = self.margin_size_ops.value()
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_margins.pdf')
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_pdf_with_margins"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )
        if not output_path:
            return
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        QApplication.processEvents()
        success = self.pdf_ops.add_margins(self.pdf_path, pages, margin, output_path)
        self.progress_bar.setVisible(False)
        if success:
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("margins_added_successfully"))
        else:
            QMessageBox.critical(self, self.localization.get_text("error"), self.localization.get_text("operation_failed"))

    def crop_pages_ops(self):
        """Crop selected pages using PDFOperations.crop_pages."""
        if not self.pdf_path:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("file_not_selected"))
            return
        pages = self.parse_page_numbers(self.crop_page_range.text().strip() or "")
        if not pages:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("invalid_page_range"))
            return
        x0, y0, x1, y1 = self.crop_x0.value(), self.crop_y0.value(), self.crop_x1.value(), self.crop_y1.value()
        if x1 <= x0 or y1 <= y0:
            QMessageBox.warning(self, self.localization.get_text("error"), self.localization.get_text("invalid_crop_box") if hasattr(self.localization, 'get_text') else "Invalid crop box")
            return
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_cropped.pdf')
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_cropped_pdf") if hasattr(self.localization, 'get_text') else "Save Cropped PDF",
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )
        if not output_path:
            return
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        QApplication.processEvents()
        crop_box = (x0, y0, x1, y1)
        success = self.pdf_ops.crop_pages(self.pdf_path, pages, crop_box, output_path)
        self.progress_bar.setVisible(False)
        if success:
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("pages_cropped_successfully") if hasattr(self.localization, 'get_text') else "Pages cropped successfully")
        else:
            QMessageBox.critical(self, self.localization.get_text("error"), self.localization.get_text("operation_failed"))


    def clear_merge_files(self):
        """Clear all files from merge list"""
        self.merge_file_list.clear()
        self.merge_files.clear()
        self.btn_merge_pdfs.setEnabled(False)

    def merge_selected_pdfs(self):
        """Merge selected PDF files"""
        if len(self.merge_files) < 2:
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                self.localization.get_text("select_at_least_two_files")
            )
            return

        # Get the current order from the list widget
        ordered_files = []
        for i in range(self.merge_file_list.count()):
            item = self.merge_file_list.item(i)
            file_path = item.data(Qt.UserRole)
            ordered_files.append(file_path)

        # Get output file
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_merged_pdf"),
            "merged_document.pdf",
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Show progress
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, len(ordered_files))
            QApplication.processEvents()

            try:
                # Create new PDF document
                merged_doc = fitz.open()

                for i, file_path in enumerate(ordered_files):
                    self.progress_bar.setValue(i)
                    QApplication.processEvents()

                    if not os.path.exists(file_path):
                        continue

                    # Open source PDF
                    source_doc = fitz.open(file_path)

                    # Insert all pages from source
                    merged_doc.insert_pdf(source_doc)
                    source_doc.close()

                # Save merged document
                merged_doc.save(output_path)
                merged_doc.close()

                # Hide progress
                self.progress_bar.setVisible(False)

                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    f"{self.localization.get_text('pdfs_merged_successfully')}\n\n"
                    f"{self.localization.get_text('files_merged')}: {len(ordered_files)}\n"
                    f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                )

            except Exception as e:
                # Hide progress
                self.progress_bar.setVisible(False)

                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}\n\n{str(e)}"
                )


class BookmarkExtractorTab(QWidget):
    """Bookmark extractor tab"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize bookmark extractor tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("bookmark_extractor_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #7B1FA2; padding: 8px; background-color: #F3E5F5; border-radius: 4px;")
        layout.addWidget(instructions)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_placeholder"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        # Extract button
        self.btn_extract = QPushButton(self.localization.get_text("extract_bookmarks"))
        self.btn_extract.clicked.connect(self.extract_bookmarks)
        self.btn_extract.setEnabled(False)
        self.btn_extract.setStyleSheet("QPushButton { background-color: #7B1FA2; color: white; font-weight: bold; padding: 10px; }")
        file_layout.addWidget(self.btn_extract)

        layout.addWidget(file_group)

        # Results area
        results_group = QGroupBox(self.localization.get_text("extracted_bookmarks"))
        results_layout = QVBoxLayout(results_group)

        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setPlaceholderText(self.localization.get_text("extracted_bookmarks_placeholder"))
        results_layout.addWidget(self.results_text)

        # Save button
        self.btn_save = QPushButton(self.localization.get_text("save_to_text"))
        self.btn_save.clicked.connect(self.save_bookmarks)
        self.btn_save.setEnabled(False)
        self.btn_save.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; }")
        results_layout.addWidget(self.btn_save)

        layout.addWidget(results_group)

        # Add stretch
        layout.addStretch()

    def select_pdf_file(self):
        """Select PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            self.localization.get_text("pdf_files_filter")
        )
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.btn_extract.setEnabled(True)

    def extract_bookmarks(self):
        """Extract bookmarks from PDF"""
        if not self.pdf_path:
            return

        try:
            bookmarks = self.pdf_ops.extract_bookmarks_from_pdf(self.pdf_path)

            if bookmarks:
                # Format bookmarks for display
                bookmark_text = []
                for bookmark in bookmarks:
                    indent = "  " * (bookmark.level - 1)
                    bookmark_text.append(f"{indent}{bookmark.title} - {bookmark.page}")

                self.results_text.setPlainText("\n".join(bookmark_text))
                self.btn_save.setEnabled(True)

                QMessageBox.information(
                    self, self.localization.get_text("success"),
                    f"{self.localization.get_text('bookmarks_extracted')}: {len(bookmarks)}"
                )
            else:
                self.results_text.setPlainText(self.localization.get_text("no_bookmarks_found"))
                QMessageBox.information(
                    self,
                    self.localization.get_text("information"),
                    self.localization.get_text("no_bookmarks_found")
                )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
            )

    def save_bookmarks(self):
        """Save extracted bookmarks to file"""
        if not self.results_text.toPlainText():
            return

        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_bookmarks.txt')

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Save Bookmarks",
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            "Text Files (*.txt);;All Files (*)"
        )

        if output_path:
            try:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(self.results_text.toPlainText())

                QMessageBox.information(
                    self, "Success!",
                    f"Bookmarks saved to:\n{os.path.basename(output_path)}"
                )
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save bookmarks:\n{str(e)}")


class WatermarkTab(QWidget):
    """Watermark tab"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""
        self.pdf_path_remove = ""

        self.init_ui()

    def init_ui(self):
        """Initialize watermark tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("watermark_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #00796B; padding: 8px; background-color: #E0F2F1; border-radius: 4px;")
        layout.addWidget(instructions)

        # Create tab widget for add/remove watermark
        self.operation_tabs = QTabWidget()

        # Add watermark tab
        add_tab = QWidget()
        self.create_add_watermark_tab(add_tab)
        self.operation_tabs.addTab(add_tab, self.localization.get_text("add_watermark_section"))

        # Remove watermark tab
        remove_tab = QWidget()
        self.create_remove_watermark_tab(remove_tab)
        self.operation_tabs.addTab(remove_tab, self.localization.get_text("remove_watermark_section"))

        layout.addWidget(self.operation_tabs)

        # Add stretch
        layout.addStretch()

    def create_add_watermark_tab(self, tab):
        """Create the add watermark tab"""
        layout = QVBoxLayout(tab)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_for_watermark"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        layout.addWidget(file_group)

        # Watermark settings
        settings_group = QGroupBox("⚙️ " + self.localization.get_text("watermark"))
        settings_layout = QVBoxLayout(settings_group)

        # Watermark text
        text_layout = QHBoxLayout()
        text_layout.addWidget(QLabel(self.localization.get_text("watermark_text")))
        self.watermark_text = QLineEdit()
        self.watermark_text.setPlaceholderText(self.localization.get_text("watermark_text_placeholder"))
        self.watermark_text.textChanged.connect(self.update_button_state)
        text_layout.addWidget(self.watermark_text)
        settings_layout.addLayout(text_layout)

        # Position and size
        options_layout = QHBoxLayout()

        options_layout.addWidget(QLabel(self.localization.get_text("watermark_position")))
        self.position_combo = QComboBox()
        position_items = [
            self.localization.get_text("center"),
            self.localization.get_text("top_left"),
            self.localization.get_text("top_right"),
            self.localization.get_text("bottom_left"),
            self.localization.get_text("bottom_right")
        ]
        self.position_combo.addItems(position_items)
        options_layout.addWidget(self.position_combo)

        options_layout.addWidget(QLabel(self.localization.get_text("font_size")))
        self.font_size = QSpinBox()
        self.font_size.setRange(10, 200)
        self.font_size.setValue(50)
        options_layout.addWidget(self.font_size)

        settings_layout.addLayout(options_layout)

        # Add watermark button
        self.btn_add_watermark = QPushButton(self.localization.get_text("add_watermark"))
        self.btn_add_watermark.clicked.connect(self.add_watermark)
        self.btn_add_watermark.setEnabled(False)
        self.btn_add_watermark.setStyleSheet("QPushButton { background-color: #00796B; color: white; font-weight: bold; padding: 10px; }")
        settings_layout.addWidget(self.btn_add_watermark)

        layout.addWidget(settings_group)
        layout.addStretch()

    def create_remove_watermark_tab(self, tab):
        """Create the remove watermark tab"""
        layout = QVBoxLayout(tab)

        # File selection for removal
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display_remove = QLineEdit()
        self.pdf_display_remove.setReadOnly(True)
        self.pdf_display_remove.setPlaceholderText(self.localization.get_text("select_pdf_for_watermark"))
        pdf_layout.addWidget(self.pdf_display_remove)

        self.btn_select_pdf_remove = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf_remove.clicked.connect(self.select_pdf_file_remove)
        pdf_layout.addWidget(self.btn_select_pdf_remove)
        file_layout.addLayout(pdf_layout)

        layout.addWidget(file_group)

        # Removal options
        options_group = QGroupBox("⚙️ " + self.localization.get_text("removal_options"))
        options_layout = QVBoxLayout(options_group)

        # Aggressive mode checkbox
        self.aggressive_mode = QCheckBox(self.localization.get_text("aggressive_mode"))
        self.aggressive_mode.setToolTip(self.localization.get_text("aggressive_mode_tooltip"))
        options_layout.addWidget(self.aggressive_mode)

        # Target specific watermarks
        self.target_updf = QCheckBox("Remove UPDF watermarks")
        self.target_updf.setChecked(True)
        options_layout.addWidget(self.target_updf)

        self.target_urls = QCheckBox(self.localization.get_text("remove_urls"))
        self.target_urls.setChecked(True)
        options_layout.addWidget(self.target_urls)

        layout.addWidget(options_group)

        # Remove watermark button
        self.btn_remove_watermark = QPushButton(self.localization.get_text("remove_watermark"))
        self.btn_remove_watermark.clicked.connect(self.remove_watermark)
        self.btn_remove_watermark.setEnabled(False)
        self.btn_remove_watermark.setStyleSheet("QPushButton { background-color: #F44336; color: white; font-weight: bold; padding: 10px; }")
        layout.addWidget(self.btn_remove_watermark)

        # Info text
        info_text = QLabel("ℹ️ " + self.localization.get_text("watermark_removal_info"))
        info_text.setWordWrap(True)
        info_text.setStyleSheet("color: #666; font-style: italic; padding: 10px; background-color: #F5F5F5; border-radius: 4px;")
        layout.addWidget(info_text)

        layout.addStretch()

    def select_pdf_file(self):
        """Select PDF file for adding watermark"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.update_button_state()

    def select_pdf_file_remove(self):
        """Select PDF file for removing watermark"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path_remove = file_path
            self.pdf_display_remove.setText(file_path)
            self.btn_remove_watermark.setEnabled(True)

    def update_button_state(self):
        """Update add watermark button state"""
        has_file = bool(self.pdf_path)
        has_text = bool(self.watermark_text.text().strip())
        self.btn_add_watermark.setEnabled(has_file and has_text)

    def add_watermark(self):
        """Add watermark to PDF"""
        if not self.pdf_path or not self.watermark_text.text().strip():
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_watermarked.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_pdf_with_watermark"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Create progress dialog
            progress = ProgressDialog(
                self.localization.get_text("watermark"),
                self.localization.get_text("adding_watermark"),
                self
            )
            progress.show()
            QApplication.processEvents()

            try:
                success = self.pdf_ops.add_watermark(
                    self.pdf_path, output_path,
                    self.watermark_text.text().strip(),
                    self.position_combo.currentText(),
                    0.5,  # opacity
                    self.font_size.value()
                )

                progress.close()

                if success:
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('watermark_added_successfully')}\n\n"
                        f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                    )
                else:
                    QMessageBox.critical(
                        self,
                        self.localization.get_text("error"),
                        self.localization.get_text("operation_failed")
                    )

            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
                )

    def remove_watermark(self):
        """Remove watermark from PDF"""
        if not self.pdf_path_remove:
            return

        pdf_name = os.path.basename(self.pdf_path_remove)
        default_name = pdf_name.replace('.pdf', '_no_watermark.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF without Watermark",
            os.path.join(os.path.dirname(self.pdf_path_remove), default_name),
            "PDF Files (*.pdf)"
        )

        if output_path:
            # Create progress dialog
            progress = ProgressDialog(
                self.localization.get_text("watermark"),
                "Removing watermark...",
                self
            )
            progress.show()
            QApplication.processEvents()

            try:
                # Get options from checkboxes
                aggressive = self.aggressive_mode.isChecked()
                target_updf = self.target_updf.isChecked()
                target_urls = self.target_urls.isChecked()

                success = self.pdf_ops.remove_watermark(
                    self.pdf_path_remove, output_path,
                    aggressive_mode=aggressive,
                    target_updf=target_updf,
                    target_urls=target_urls
                )

                progress.close()

                if success:
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('watermark_removed_successfully')}\n\n"
                        f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                    )
                else:
                    QMessageBox.critical(
                        self,
                        self.localization.get_text("error"),
                        self.localization.get_text("operation_failed")
                    )

            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
                )


class ImageExtractionTab(QWidget):
    """Image extraction tab"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize image extraction tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("image_extraction_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #E65100; padding: 8px; background-color: #FFF3E0; border-radius: 4px;")
        layout.addWidget(instructions)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_placeholder"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        layout.addWidget(file_group)

        # Extraction options
        options_group = QGroupBox(self.localization.get_text("extraction_options"))
        options_layout = QVBoxLayout(options_group)

        # Page range
        page_layout = QHBoxLayout()
        page_layout.addWidget(QLabel(self.localization.get_text("page_range")))
        self.page_range = QLineEdit()
        self.page_range.setPlaceholderText(self.localization.get_text("page_range_examples"))
        self.page_range.setText(self.localization.get_text("all_pages"))
        page_layout.addWidget(self.page_range)
        options_layout.addLayout(page_layout)

        # Page range help
        help_text = QLabel(self.localization.get_text("page_range_help"))
        help_text.setWordWrap(True)
        help_text.setStyleSheet("color: #666; font-size: 11px; padding: 5px; background-color: #F5F5F5; border-radius: 3px;")
        options_layout.addWidget(help_text)

        # Image extraction options
        image_group = QGroupBox("🖼️ " + self.localization.get_text("extract_images"))
        image_layout = QVBoxLayout(image_group)

        # Extraction method
        method_layout = QHBoxLayout()
        method_layout.addWidget(QLabel(self.localization.get_text("extraction_method")))
        self.extraction_method = QComboBox()
        method_items = [
            self.localization.get_text("both_methods"),
            self.localization.get_text("embedded_only"),
            self.localization.get_text("rendered_only")
        ]
        self.extraction_method.addItems(method_items)
        method_layout.addWidget(self.extraction_method)
        image_layout.addLayout(method_layout)

        # Quality settings
        quality_layout = QHBoxLayout()
        quality_layout.addWidget(QLabel(self.localization.get_text("quality")))
        self.quality_combo = QComboBox()
        quality_items = [
            self.localization.get_text("high_quality"),
            self.localization.get_text("medium_quality"),
            self.localization.get_text("standard_quality")
        ]
        self.quality_combo.addItems(quality_items)
        quality_layout.addWidget(self.quality_combo)
        image_layout.addLayout(quality_layout)

        # Extract button
        self.btn_extract_images = QPushButton(self.localization.get_text("extract_all_images"))
        self.btn_extract_images.clicked.connect(self.extract_images)
        self.btn_extract_images.setEnabled(False)
        self.btn_extract_images.setStyleSheet("QPushButton { background-color: #E65100; color: white; font-weight: bold; padding: 10px; }")
        image_layout.addWidget(self.btn_extract_images)

        options_layout.addWidget(image_group)

        layout.addWidget(options_group)

        # Add stretch
        layout.addStretch()

    def select_pdf_file(self):
        """Select PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.btn_extract_images.setEnabled(True)

    def extract_images(self):
        """Extract images from PDF"""
        if not self.pdf_path:
            return

        # Select output directory
        output_dir = QFileDialog.getExistingDirectory(
            self,
            self.localization.get_text("select_output_directory")
        )
        if not output_dir:
            return

        # Create progress dialog
        progress = ProgressDialog(
            self.localization.get_text("extract_images"),
            self.localization.get_text("processing_images"),
            self
        )
        progress.show()
        QApplication.processEvents()

        try:
            page_range = self.page_range.text().strip() or "all"

            # Update progress
            progress.update_message(self.localization.get_text("processing_images"))
            QApplication.processEvents()

            image_count = self.pdf_ops.extract_images(self.pdf_path, output_dir, page_range)

            progress.close()

            if image_count > 0:
                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    f"{self.localization.get_text('images_extracted_successfully')}\n\n"
                    f"{self.localization.get_text('images_saved_to')} {os.path.basename(output_dir)}\n"
                    f"{image_count} {self.localization.get_text('images_found')}"
                )
            else:
                QMessageBox.information(
                    self,
                    self.localization.get_text("information"),
                    self.localization.get_text("no_images_found")
                )

        except Exception as e:
            progress.close()
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
            )




class TextExtractionTab(QWidget):
    """Enhanced text extraction tab with batch processing and DOCX support"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""
        self.pdf_files = []  # For batch processing
        self.batch_mode = False

        self.init_ui()

    def init_ui(self):
        """Initialize text extraction tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("text_extraction_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #388E3C; padding: 8px; background-color: #E8F5E8; border-radius: 4px;")
        layout.addWidget(instructions)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # Processing mode selection
        mode_layout = QHBoxLayout()
        mode_layout.addWidget(QLabel(self.localization.get_text("processing_mode")))

        self.single_mode_radio = QRadioButton(self.localization.get_text("single_file"))
        self.batch_mode_radio = QRadioButton(self.localization.get_text("batch_processing"))
        self.single_mode_radio.setChecked(True)
        self.single_mode_radio.toggled.connect(self.toggle_processing_mode)

        mode_layout.addWidget(self.single_mode_radio)
        mode_layout.addWidget(self.batch_mode_radio)
        mode_layout.addStretch()
        file_layout.addLayout(mode_layout)

        # Single file selection
        self.single_file_widget = QWidget()
        single_layout = QHBoxLayout(self.single_file_widget)
        single_layout.setContentsMargins(0, 0, 0, 0)

        single_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_placeholder"))
        single_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        single_layout.addWidget(self.btn_select_pdf)
        file_layout.addWidget(self.single_file_widget)

        # Batch file selection
        self.batch_file_widget = QWidget()
        self.batch_file_widget.setVisible(False)
        batch_layout = QVBoxLayout(self.batch_file_widget)
        batch_layout.setContentsMargins(0, 0, 0, 0)

        batch_buttons_layout = QHBoxLayout()
        self.btn_select_multiple = QPushButton(self.localization.get_text("select_multiple_pdfs"))
        self.btn_select_multiple.clicked.connect(self.select_multiple_pdfs)
        batch_buttons_layout.addWidget(self.btn_select_multiple)

        self.btn_clear_selection = QPushButton(self.localization.get_text("clear_selection"))
        self.btn_clear_selection.clicked.connect(self.clear_pdf_selection)
        self.btn_clear_selection.setEnabled(False)
        batch_buttons_layout.addWidget(self.btn_clear_selection)
        batch_buttons_layout.addStretch()
        batch_layout.addLayout(batch_buttons_layout)

        # File list
        self.file_list = QListWidget()
        self.file_list.setMaximumHeight(120)
        batch_layout.addWidget(self.file_list)

        file_layout.addWidget(self.batch_file_widget)
        layout.addWidget(file_group)

        # Extraction options
        options_group = QGroupBox("⚙️ " + self.localization.get_text("extract_text"))
        options_layout = QVBoxLayout(options_group)

        # Page range
        page_layout = QHBoxLayout()
        page_layout.addWidget(QLabel(self.localization.get_text("page_range")))
        self.page_range = QLineEdit()
        self.page_range.setPlaceholderText(self.localization.get_text("page_range_placeholder"))
        self.page_range.setText("all")  # Use "all" keyword instead of localized text
        page_layout.addWidget(self.page_range)
        options_layout.addLayout(page_layout)

        # Format selection
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel(self.localization.get_text("output_format")))
        self.text_format = QComboBox()
        self.text_format.addItems(["txt", "json", "docx"])
        self.text_format.setToolTip(self.localization.get_text("text_format_tooltip"))
        format_layout.addWidget(self.text_format)
        format_layout.addStretch()
        options_layout.addLayout(format_layout)

        # Online OCR via Colab launcher
        self.btn_convert_pdf_online = QPushButton(self.localization.get_text("convert_pdf_online"))
        self.btn_convert_pdf_online.setToolTip(self.localization.get_text("convert_pdf_online_tooltip"))
        self.btn_convert_pdf_online.clicked.connect(self.show_online_ocr_instructions)
        options_layout.addWidget(self.btn_convert_pdf_online)

        # Import OCR results (auto-detect TXT/DOCX and integrate)
        self.btn_import_ocr_results = QPushButton(self.localization.get_text("import_ocr_results_btn"))
        self.btn_import_ocr_results.setToolTip(self.localization.get_text("choose_ocr_results"))
        self.btn_import_ocr_results.clicked.connect(self.import_ocr_results)
        options_layout.addWidget(self.btn_import_ocr_results)

        # Set default OCR results folder
        self.btn_set_default_ocr_dir = QPushButton(self.localization.get_text("set_default_ocr_dir"))
        self.btn_set_default_ocr_dir.setToolTip(self.localization.get_text("change_default_ocr_dir"))
        self.btn_set_default_ocr_dir.clicked.connect(self.set_default_ocr_dir)
        options_layout.addWidget(self.btn_set_default_ocr_dir)

        # Output directory (for batch processing)
        self.output_dir_widget = QWidget()
        self.output_dir_widget.setVisible(False)
        output_dir_layout = QHBoxLayout(self.output_dir_widget)
        output_dir_layout.setContentsMargins(0, 0, 0, 0)

        output_dir_layout.addWidget(QLabel(self.localization.get_text("output_directory")))
        self.output_dir_display = QLineEdit()
        self.output_dir_display.setReadOnly(True)
        self.output_dir_display.setPlaceholderText(self.localization.get_text("select_output_directory"))
        output_dir_layout.addWidget(self.output_dir_display)

        self.btn_select_output_dir = QPushButton(self.localization.get_text("browse"))
        self.btn_select_output_dir.clicked.connect(self.select_output_directory)
        output_dir_layout.addWidget(self.btn_select_output_dir)

        options_layout.addWidget(self.output_dir_widget)

        # Extract button
        self.btn_extract_text = QPushButton(self.localization.get_text("extract_text_btn"))
        self.btn_extract_text.clicked.connect(self.extract_text)
        self.btn_extract_text.setEnabled(False)
        self.btn_extract_text.setStyleSheet("QPushButton { background-color: #388E3C; color: white; font-weight: bold; padding: 10px; }")
        options_layout.addWidget(self.btn_extract_text)

        layout.addWidget(options_group)

        # Progress bar (initially hidden)
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # Status label
        self.status_label = QLabel("")
        self.status_label.setStyleSheet("color: #666; padding: 5px;")
        self.status_label.setVisible(False)
        layout.addWidget(self.status_label)

        # Add stretch
        layout.addStretch()

    def show_online_ocr_instructions(self):
        """Show localized steps to use the online Colab OCR and offer to open it"""
        from PySide6.QtWidgets import QMessageBox
        from PySide6.QtCore import QUrl
        from PySide6.QtGui import QDesktopServices

        url = "https://colab.research.google.com/github/ieasybooks/tahweel/blob/main/colab_notebook.ipynb#scrollTo=Uf8p_j7vXHTg"
        steps = self.localization.get_text("online_ocr_instructions")
        title = self.localization.get_text("online_ocr_instructions_title")

        msg = QMessageBox(self)
        msg.setWindowTitle(title)
        msg.setText(steps)
        msg.setStandardButtons(QMessageBox.Open | QMessageBox.Close)
        open_label = self.localization.get_text("open_colab")
        msg.button(QMessageBox.Open).setText(open_label) if msg.button(QMessageBox.Open) else None
        res = msg.exec()
        if res == QMessageBox.Open:
            QDesktopServices.openUrl(QUrl(url))

    def toggle_processing_mode(self):
        """Toggle between single file and batch processing mode"""
        self.batch_mode = self.batch_mode_radio.isChecked()

        # Show/hide appropriate widgets
        self.single_file_widget.setVisible(not self.batch_mode)
        self.batch_file_widget.setVisible(self.batch_mode)
        self.output_dir_widget.setVisible(self.batch_mode)

        # Update button state
        self.update_extract_button_state()

        # Update button text
        if self.batch_mode:
            self.btn_extract_text.setText(self.localization.get_text("batch_extract_btn"))
        else:
            self.btn_extract_text.setText(self.localization.get_text("extract_text_btn"))

    def select_pdf_file(self):
        """Select single PDF file and intelligently detect if OCR is needed"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            self.localization.get_text("pdf_files_filter")
        )
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.update_extract_button_state()
            # Smart detection: if no extractable text, offer Online OCR
            self.detect_and_offer_ocr(file_path)


    def detect_and_offer_ocr(self, file_path: str, sample_pages: int = 3):
        """Detect if PDF has extractable text; if not, guide to Online OCR (localized)."""
        try:
            import fitz  # PyMuPDF
        except Exception:
            return  # If fitz unavailable, skip detection
        try:
            doc = fitz.open(file_path)
            total = len(doc)
            pages_to_check = min(sample_pages, total)
            words = 0
            for i in range(pages_to_check):
                text = doc.load_page(i).get_text("text").strip()
                words += len(text.split())
            doc.close()
            if words == 0:
                from PySide6.QtWidgets import QMessageBox
                from PySide6.QtCore import QUrl
                from PySide6.QtGui import QDesktopServices
                title = self.localization.get_text("image_pdf_detected_use_ocr_title")
                msg = self.localization.get_text("image_pdf_detected_use_ocr_message")
                box = QMessageBox(self)
                box.setWindowTitle(title)
                box.setText(msg)
                box.setStandardButtons(QMessageBox.Open | QMessageBox.Close)
                open_label = self.localization.get_text("open_colab")
                if box.button(QMessageBox.Open):
                    box.button(QMessageBox.Open).setText(open_label)
                res = box.exec()
                if res == QMessageBox.Open:
                    url = "https://colab.research.google.com/github/ieasybooks/tahweel/blob/main/colab_notebook.ipynb#scrollTo=Uf8p_j7vXHTg"
                    QDesktopServices.openUrl(QUrl(url))
        except Exception:
            pass

    def import_ocr_results(self):
        """Auto-detect TXT/DOCX from Downloads that match current PDF name and move them."""
        if not self.pdf_path:
            return
        import os, shutil
        from pathlib import Path
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from PySide6.QtCore import QSettings

        base = os.path.splitext(os.path.basename(self.pdf_path))[0]
        # Candidate filenames
        candidates = [f"{base}.txt", f"{base}.docx"]

        # Determine Downloads path
        downloads = Path.home() / "Downloads"

        found = {}
        for name in candidates:
            p = downloads / name
            if p.exists():
                found[name] = p

        # If missing, allow manual selection
        if len(found) < 2:
            QMessageBox.information(self, self.localization.get_text("information"), self.localization.get_text("ocr_results_not_found"))
            # Manual TXT
            txt_path, _ = QFileDialog.getOpenFileName(self, self.localization.get_text("select_txt_file"), str(downloads), "Text (*.txt)")
            docx_path, _ = QFileDialog.getOpenFileName(self, self.localization.get_text("select_docx_file"), str(downloads), "DOCX (*.docx)")
            if txt_path:
                p = Path(txt_path)
                if p.stem != base:
                    QMessageBox.warning(self, self.localization.get_text("warning"), self.localization.get_text("file_name_mismatch"))
                    return
                found[f"{base}.txt"] = p
            if docx_path:
                p = Path(docx_path)
                if p.stem != base:
                    QMessageBox.warning(self, self.localization.get_text("warning"), self.localization.get_text("file_name_mismatch"))
                    return
                found[f"{base}.docx"] = p

        # Validate names
        for required in candidates:
            if required not in found:
                return  # still incomplete

        # Preferred destination dir (persisted)
        settings = QSettings("PDFToolsApp", "PDF Tools")
        dest_dir = settings.value("ocr/default_output_dir", str(Path(self.pdf_path).parent))
        dest_dir_path = Path(dest_dir)
        dest_dir_path.mkdir(parents=True, exist_ok=True)

        moved_paths = []
        for name, src_path in found.items():
            dest_path = dest_dir_path / name
            try:
                if src_path.resolve() != dest_path.resolve():
                    shutil.move(str(src_path), str(dest_path))
                moved_paths.append(str(dest_path))
            except Exception:
                # Fallback to copy
                try:
                    shutil.copy2(str(src_path), str(dest_path))
                    moved_paths.append(str(dest_path))
                except Exception:
                    pass

        if moved_paths:
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("ocr_results_moved"))

    def set_default_ocr_dir(self):
        """Let user pick and save preferred OCR output directory (persisted)."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from PySide6.QtCore import QSettings
        directory = QFileDialog.getExistingDirectory(self, self.localization.get_text("set_default_ocr_dir"))
        if directory:
            settings = QSettings("PDFToolsApp", "PDF Tools")
            settings.setValue("ocr/default_output_dir", directory)
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("preferred_ocr_dir_saved"))
            QMessageBox.information(self, self.localization.get_text("success"), self.localization.get_text("ocr_results_moved"))

    def select_multiple_pdfs(self):
        """Select multiple PDF files for batch processing"""
        file_paths, _ = QFileDialog.getOpenFileNames(self, "Select PDF Files", "", "PDF Files (*.pdf)")
        if file_paths:
            self.pdf_files = file_paths
            self.file_list.clear()
            for file_path in file_paths:
                self.file_list.addItem(os.path.basename(file_path))

            self.btn_clear_selection.setEnabled(True)
            self.update_extract_button_state()

    def clear_pdf_selection(self):
        """Clear selected PDF files"""
        self.pdf_files = []
        self.file_list.clear()
        self.btn_clear_selection.setEnabled(False)
        self.update_extract_button_state()

    def select_output_directory(self):
        """Select output directory for batch processing"""
        directory = QFileDialog.getExistingDirectory(self, "Select Output Directory")
        if directory:
            self.output_dir_display.setText(directory)
            self.update_extract_button_state()

    def update_extract_button_state(self):
        """Update extract button enabled state based on current mode"""
        if self.batch_mode:
            # Batch mode: need files and output directory
            has_files = len(self.pdf_files) > 0
            has_output_dir = bool(self.output_dir_display.text().strip())
            self.btn_extract_text.setEnabled(has_files and has_output_dir)
        else:
            # Single mode: need single file
            self.btn_extract_text.setEnabled(bool(self.pdf_path))

    def extract_text(self):
        """Extract text from PDF(s) - handles both single and batch processing"""
        if self.batch_mode:
            self.extract_text_batch()
        else:
            self.extract_text_single()

    def extract_text_single(self):
        """Extract text from single PDF"""
        if not self.pdf_path:
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        format_ext = self.text_format.currentText()
        default_name = pdf_name.replace('.pdf', f'_text.{format_ext}')

        # Set file filter based on format
        if format_ext == "txt":
            file_filter = "Text Files (*.txt)"
        elif format_ext == "json":
            file_filter = "JSON Files (*.json)"
        elif format_ext == "docx":
            file_filter = "Word Documents (*.docx)"
        else:
            file_filter = "All Files (*)"

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Save Extracted Text",
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            file_filter
        )

        if output_path:
            try:
                page_range_text = self.page_range.text().strip()
                # Handle localized text
                if not page_range_text or page_range_text == "All Pages" or page_range_text == "جميع الصفحات":
                    page_range = "all"
                else:
                    page_range = page_range_text

                success = self.pdf_ops.extract_text(
                    self.pdf_path, output_path, page_range, format_ext
                )

                if success:
                    QMessageBox.information(
                        self, "Success!",
                        f"Text extracted successfully!\n\n"
                        f"Output: {os.path.basename(output_path)}\n"
                        f"Format: {format_ext.upper()}"
                    )
                else:
                    QMessageBox.critical(self, "Error", "Failed to extract text")

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to extract text:\n{str(e)}")

    def extract_text_batch(self):
        """Extract text from multiple PDFs"""
        if not self.pdf_files:
            return

        output_dir = self.output_dir_display.text().strip()
        if not output_dir:
            QMessageBox.warning(self, "No Output Directory", "Please select an output directory.")
            return

        # Show progress
        self.progress_bar.setVisible(True)
        self.status_label.setVisible(True)
        self.progress_bar.setMaximum(len(self.pdf_files))
        self.progress_bar.setValue(0)

        # Disable extract button during processing
        self.btn_extract_text.setEnabled(False)

        try:
            page_range_text = self.page_range.text().strip()
            # Handle localized text
            if not page_range_text or page_range_text == "All Pages" or page_range_text == "جميع الصفحات":
                page_range = "all"
            else:
                page_range = page_range_text
            format_ext = self.text_format.currentText()

            def progress_callback(current, total, filename):
                self.progress_bar.setValue(current)
                self.status_label.setText(f"Processing {current}/{total}: {filename}")
                QApplication.processEvents()  # Update UI

            successful, failed = self.pdf_ops.extract_text_batch(
                self.pdf_files, output_dir, page_range, format_ext, progress_callback
            )

            # Hide progress
            self.progress_bar.setVisible(False)
            self.status_label.setVisible(False)

            # Show results
            message = f"Batch extraction completed!\n\n"
            message += f"Successfully processed: {len(successful)} files\n"
            message += f"Failed: {len(failed)} files\n\n"

            if successful:
                message += "Successful files:\n"
                for item in successful[:5]:  # Show first 5
                    message += f"• {item['filename']}\n"
                if len(successful) > 5:
                    message += f"... and {len(successful) - 5} more\n"

            if failed:
                message += f"\nFailed files:\n"
                for item in failed[:3]:  # Show first 3 failures
                    message += f"• {os.path.basename(item['input'])}: {item['error']}\n"
                if len(failed) > 3:
                    message += f"... and {len(failed) - 3} more\n"

            QMessageBox.information(self, "Batch Extraction Results", message)

        except Exception as e:
            QMessageBox.critical(self, "Batch Processing Error", f"Error during batch processing:\n{str(e)}")

        finally:
            # Re-enable extract button
            self.btn_extract_text.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.status_label.setVisible(False)


class MergeTab(QWidget):
    """PDF merge tab"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_files = []

        self.init_ui()

    def init_ui(self):
        """Initialize merge tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("merge_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #5D4037; padding: 8px; background-color: #EFEBE9; border-radius: 4px;")
        layout.addWidget(instructions)

        # Merge PDFs section
        merge_group = QGroupBox(self.localization.get_text("select_pdfs_merge"))
        merge_layout = QVBoxLayout(merge_group)

        # File list
        self.file_list = QListWidget()
        self.file_list.setMaximumHeight(150)
        merge_layout.addWidget(self.file_list)

        # Buttons for merge
        merge_buttons = QHBoxLayout()

        self.btn_add_files = QPushButton(self.localization.get_text("add_pdfs"))
        self.btn_add_files.clicked.connect(self.add_pdf_files)
        merge_buttons.addWidget(self.btn_add_files)

        self.btn_remove_file = QPushButton(self.localization.get_text("remove"))
        self.btn_remove_file.clicked.connect(self.remove_pdf_file)
        merge_buttons.addWidget(self.btn_remove_file)

        self.btn_merge = QPushButton(self.localization.get_text("merge_pdfs_btn"))
        self.btn_merge.clicked.connect(self.merge_pdfs)
        self.btn_merge.setEnabled(False)
        self.btn_merge.setStyleSheet("QPushButton { background-color: #5D4037; color: white; font-weight: bold; }")
        merge_buttons.addWidget(self.btn_merge)

        merge_layout.addLayout(merge_buttons)
        layout.addWidget(merge_group)

        # Merge order info
        info_group = QGroupBox(self.localization.get_text("merge_info"))
        info_layout = QVBoxLayout(info_group)

        info_text = QLabel(self.localization.get_text("merge_instructions"))
        info_text.setWordWrap(True)
        info_text.setStyleSheet("color: #666; padding: 10px;")
        info_layout.addWidget(info_text)

        layout.addWidget(info_group)

        # Add stretch
        layout.addStretch()

    def add_pdf_files(self):
        """Add PDF files for merging"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            self.localization.get_text("select_multiple_files"),
            "",
            self.localization.get_text("pdf_files_filter")
        )

        if files:
            for file_path in files:
                if file_path not in self.pdf_files:
                    self.pdf_files.append(file_path)
                    self.file_list.addItem(os.path.basename(file_path))

            # Update status
            count = len(self.pdf_files)
            if count > 0:
                status_text = f"{count} {self.localization.get_text('files_selected')}"
                self.file_list.setToolTip(status_text)

        self.btn_merge.setEnabled(len(self.pdf_files) >= 2)

    def remove_pdf_file(self):
        """Remove selected PDF file"""
        current_row = self.file_list.currentRow()
        if current_row >= 0:
            self.file_list.takeItem(current_row)
            del self.pdf_files[current_row]
            self.btn_merge.setEnabled(len(self.pdf_files) >= 2)

    def merge_pdfs(self):
        """Merge selected PDFs"""
        if len(self.pdf_files) < 2:
            return

        output_path, _ = QFileDialog.getSaveFileName(
            self, "Save Merged PDF", "merged.pdf", "PDF Files (*.pdf)"
        )

        if output_path:
            try:
                success = self.pdf_ops.merge_pdfs(self.pdf_files, output_path)

                if success:
                    QMessageBox.information(
                        self, "Success!",
                        f"Successfully merged {len(self.pdf_files)} PDFs!\n\n"
                        f"Output: {os.path.basename(output_path)}"
                    )
                else:
                    QMessageBox.critical(self, "Error", "Failed to merge PDFs")

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to merge PDFs:\n{str(e)}")


class CompressTab(QWidget):
    """PDF compression tab"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize compress tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("compress_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #D32F2F; padding: 8px; background-color: #FFEBEE; border-radius: 4px;")
        layout.addWidget(instructions)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_for_compression"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        layout.addWidget(file_group)

        # Compression options
        options_group = QGroupBox(self.localization.get_text("compression_options"))
        options_layout = QVBoxLayout(options_group)

        # Compression level info
        info_text = QLabel(self.localization.get_text("compression_features"))
        info_text.setWordWrap(True)
        info_text.setStyleSheet("color: #666; padding: 10px;")
        options_layout.addWidget(info_text)

        # Compress button
        self.btn_compress = QPushButton(self.localization.get_text("compress_pdf_btn"))
        self.btn_compress.clicked.connect(self.compress_pdf)
        self.btn_compress.setEnabled(False)
        self.btn_compress.setStyleSheet("QPushButton { background-color: #D32F2F; color: white; font-weight: bold; padding: 10px; }")
        options_layout.addWidget(self.btn_compress)

        layout.addWidget(options_group)

        # Results section
        results_group = QGroupBox(self.localization.get_text("compression_results"))
        results_layout = QVBoxLayout(results_group)

        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setMaximumHeight(150)
        self.results_text.setPlaceholderText(self.localization.get_text("compression_results_placeholder"))
        results_layout.addWidget(self.results_text)

        layout.addWidget(results_group)

        # Add stretch
        layout.addStretch()

    def select_pdf_file(self):
        """Select PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.btn_compress.setEnabled(True)

            # Show original file size
            original_size = os.path.getsize(file_path)
            size_mb = original_size / (1024 * 1024)
            self.results_text.setPlainText(f"Original file size: {size_mb:.2f} MB")

    def compress_pdf(self):
        """Compress PDF file"""
        if not self.pdf_path:
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_compressed.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_compressed_pdf"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Create progress dialog
            progress = ProgressDialog(
                self.localization.get_text("compress"),
                self.localization.get_text("compressing_pdf"),
                self
            )
            progress.show()
            QApplication.processEvents()

            try:
                # Show progress in results
                self.results_text.append(f"\n{self.localization.get_text('compressing_pdf')}")
                QApplication.processEvents()

                success = self.pdf_ops.compress_pdf(self.pdf_path, output_path)

                progress.close()

                if success:
                    # Calculate compression results
                    original_size = os.path.getsize(self.pdf_path)
                    compressed_size = os.path.getsize(output_path)
                    ratio = (1 - compressed_size / original_size) * 100

                    original_mb = original_size / (1024 * 1024)
                    compressed_mb = compressed_size / (1024 * 1024)

                    results = (
                        f"✅ {self.localization.get_text('compression_completed')}\n\n"
                        f"{self.localization.get_text('original_size')}: {original_mb:.2f} {self.localization.get_text('megabytes')}\n"
                        f"{self.localization.get_text('compressed_size')}: {compressed_mb:.2f} {self.localization.get_text('megabytes')}\n"
                        f"{self.localization.get_text('size_reduction')}: {ratio:.1f}%\n"
                        f"{self.localization.get_text('space_saved')}: {(original_mb - compressed_mb):.2f} {self.localization.get_text('megabytes')}"
                    )

                    self.results_text.setPlainText(results)

                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('pdf_compressed_successfully')}\n\n"
                        f"{self.localization.get_text('size_reduction')}: {ratio:.1f}%\n"
                        f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                    )
                else:
                    self.results_text.append(f"❌ {self.localization.get_text('operation_failed')}")
                    QMessageBox.critical(
                        self,
                        self.localization.get_text("error"),
                        self.localization.get_text("operation_failed")
                    )

            except Exception as e:
                progress.close()
                self.results_text.append(f"❌ {self.localization.get_text('error')}: {str(e)}")
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
                )


class PageEditingTab(QWidget):
    """Page editing tab for rotate, crop, and add margins"""

    def __init__(self, history_manager=None, localization=None):
        super().__init__()
        self.pdf_ops = PDFOperations(history_manager)
        self.localization = localization or Localization()
        self.pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize page editing tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("page_editing_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #7B1FA2; padding: 8px; background-color: #F3E5F5; border-radius: 4px;")
        layout.addWidget(instructions)

        # File selection
        file_group = QGroupBox(self.localization.get_text("file_selection"))
        file_layout = QVBoxLayout(file_group)

        # PDF file
        pdf_layout = QHBoxLayout()
        pdf_layout.addWidget(QLabel(self.localization.get_text("pdf_file")))
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_for_editing"))
        pdf_layout.addWidget(self.pdf_display)

        self.btn_select_pdf = QPushButton(self.localization.get_text("browse"))
        self.btn_select_pdf.clicked.connect(self.select_pdf_file)
        pdf_layout.addWidget(self.btn_select_pdf)
        file_layout.addLayout(pdf_layout)

        # Page range
        page_layout = QHBoxLayout()
        page_layout.addWidget(QLabel(self.localization.get_text("page_range")))
        self.page_range = QLineEdit()
        self.page_range.setPlaceholderText(self.localization.get_text("page_range_placeholder"))
        self.page_range.setText(self.localization.get_text("all_pages"))
        page_layout.addWidget(self.page_range)
        file_layout.addLayout(page_layout)

        layout.addWidget(file_group)

        # Rotate section
        rotate_group = QGroupBox(self.localization.get_text("rotate_pages"))
        rotate_layout = QVBoxLayout(rotate_group)

        rotate_options = QHBoxLayout()
        rotate_options.addWidget(QLabel(self.localization.get_text("rotation")))
        self.rotation_combo = QComboBox()
        rotation_items = [
            self.localization.get_text("rotate_90"),
            self.localization.get_text("rotate_180"),
            self.localization.get_text("rotate_270")
        ]
        self.rotation_combo.addItems(rotation_items)
        rotate_options.addWidget(self.rotation_combo)
        rotate_options.addStretch()
        rotate_layout.addLayout(rotate_options)

        self.btn_rotate = QPushButton(self.localization.get_text("rotate_pages_btn"))
        self.btn_rotate.clicked.connect(self.rotate_pages)
        self.btn_rotate.setEnabled(False)
        self.btn_rotate.setStyleSheet("QPushButton { background-color: #7B1FA2; color: white; font-weight: bold; }")
        rotate_layout.addWidget(self.btn_rotate)

        layout.addWidget(rotate_group)

        # Add margins section
        margins_group = QGroupBox(self.localization.get_text("add_white_margins"))
        margins_layout = QVBoxLayout(margins_group)

        margin_options = QHBoxLayout()
        margin_options.addWidget(QLabel(self.localization.get_text("margin_size")))
        self.margin_size = QSpinBox()
        self.margin_size.setRange(10, 200)
        self.margin_size.setValue(50)
        self.margin_size.setSuffix(self.localization.get_text("points"))
        margin_options.addWidget(self.margin_size)
        margin_options.addStretch()
        margins_layout.addLayout(margin_options)

        margin_info = QLabel(self.localization.get_text("margin_help"))
        margin_info.setWordWrap(True)
        margin_info.setStyleSheet("color: #666; font-style: italic;")
        margins_layout.addWidget(margin_info)

        self.btn_add_margins = QPushButton(self.localization.get_text("add_margins_btn"))
        self.btn_add_margins.clicked.connect(self.add_margins)
        self.btn_add_margins.setEnabled(False)
        self.btn_add_margins.setStyleSheet("QPushButton { background-color: #7B1FA2; color: white; font-weight: bold; }")
        margins_layout.addWidget(self.btn_add_margins)

        layout.addWidget(margins_group)

        # Add stretch
        layout.addStretch()

    def select_pdf_file(self):
        """Select PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(file_path)
            self.btn_rotate.setEnabled(True)
            self.btn_add_margins.setEnabled(True)

    def rotate_pages(self):
        """Rotate pages in PDF"""
        if not self.pdf_path:
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_rotated.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_rotated_pdf"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Create progress dialog
            progress = ProgressDialog(
                self.localization.get_text("rotate_pages"),
                self.localization.get_text("rotating_pages"),
                self
            )
            progress.show()
            QApplication.processEvents()

            try:
                rotation_text = self.rotation_combo.currentText()
                if "90°" in rotation_text:
                    rotation = 90
                elif "180°" in rotation_text:
                    rotation = 180
                elif "270°" in rotation_text:
                    rotation = 270
                else:
                    rotation = 90

                page_range = self.page_range.text().strip() or "all"
                success = self.pdf_ops.rotate_pages(
                    self.pdf_path, output_path, page_range, rotation
                )

                progress.close()

                if success:
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('pages_rotated_successfully')}\n\n"
                        f"{self.localization.get_text('rotation')}: {rotation_text}\n"
                        f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                    )
                else:
                    QMessageBox.critical(
                        self,
                        self.localization.get_text("error"),
                        self.localization.get_text("operation_failed")
                    )

            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
                )

    def add_margins(self):
        """Add margins to PDF pages"""
        if not self.pdf_path:
            return

        # Get output file
        pdf_name = os.path.basename(self.pdf_path)
        default_name = pdf_name.replace('.pdf', '_with_margins.pdf')

        output_path, _ = QFileDialog.getSaveFileName(
            self,
            self.localization.get_text("save_pdf_with_margins"),
            os.path.join(os.path.dirname(self.pdf_path), default_name),
            self.localization.get_text("pdf_files_filter")
        )

        if output_path:
            # Create progress dialog
            progress = ProgressDialog(
                self.localization.get_text("add_margins"),
                self.localization.get_text("adding_margins"),
                self
            )
            progress.show()
            QApplication.processEvents()

            try:
                page_range = self.page_range.text().strip() or "all"
                margin_size = self.margin_size.value()

                success = self.pdf_ops.add_margins(
                    self.pdf_path, output_path, page_range, margin_size
                )

                progress.close()

                if success:
                    QMessageBox.information(
                        self,
                        self.localization.get_text("success"),
                        f"{self.localization.get_text('margins_added_successfully')}\n\n"
                        f"{self.localization.get_text('margin_size')}: {margin_size}{self.localization.get_text('points')}\n"
                        f"{self.localization.get_text('output')}: {os.path.basename(output_path)}"
                    )
                else:
                    QMessageBox.critical(
                        self,
                        self.localization.get_text("error"),
                        self.localization.get_text("operation_failed")
                    )

            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    f"{self.localization.get_text('operation_failed')}:\n{str(e)}"
                )


class PDFSecurityRemovalTab(QWidget):
    """PDF Security Removal tab for removing PDF restrictions"""

    def __init__(self, history_manager: HistoryManager, localization: Localization = None):
        super().__init__()
        self.history_manager = history_manager
        self.localization = localization or Localization()
        self.pdf_path = ""

        self.init_ui()

    def init_ui(self):
        """Initialize PDF security removal tab UI with proper localization"""
        layout = QVBoxLayout(self)

        # Instructions with localization
        instructions = QLabel(self.localization.get_text("security_removal_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; margin: 10px; padding: 10px;")
        layout.addWidget(instructions)

        # File selection group with localization
        file_group = QGroupBox(self.localization.get_text("select_file"))
        file_layout = QVBoxLayout(file_group)

        # PDF file selection with localization
        pdf_layout = QHBoxLayout()
        self.pdf_path_edit = QLineEdit()
        self.pdf_path_edit.setPlaceholderText(self.localization.get_text("select_pdf_file"))
        self.pdf_path_edit.setReadOnly(True)
        self.pdf_path_edit.setStyleSheet("font-family: 'Tajawal'; font-size: 14px;")
        pdf_browse_btn = QPushButton(self.localization.get_text("browse"))
        pdf_browse_btn.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; font-weight: bold;")
        pdf_browse_btn.clicked.connect(self.browse_pdf_file)
        pdf_label = QLabel(self.localization.get_text("pdf_file") + ":")
        pdf_label.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; font-weight: bold;")
        pdf_layout.addWidget(pdf_label)
        pdf_layout.addWidget(self.pdf_path_edit)
        pdf_layout.addWidget(pdf_browse_btn)
        file_layout.addLayout(pdf_layout)

        # Password input (for encrypted PDFs) with localization
        password_layout = QHBoxLayout()
        self.password_edit = QLineEdit()
        self.password_edit.setPlaceholderText(self.localization.get_text("password_optional"))
        self.password_edit.setEchoMode(QLineEdit.Password)
        self.password_edit.setStyleSheet("font-family: 'Tajawal'; font-size: 14px;")
        password_label = QLabel(self.localization.get_text("pdf_password") + ":")
        password_label.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; font-weight: bold;")
        password_layout.addWidget(password_label)
        password_layout.addWidget(self.password_edit)
        file_layout.addLayout(password_layout)

        layout.addWidget(file_group)

        # Output settings group with localization
        output_group = QGroupBox(self.localization.get_text("output_settings"))
        output_layout = QVBoxLayout(output_group)

        # Output file selection with localization
        output_layout_h = QHBoxLayout()
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setPlaceholderText(self.localization.get_text("select_output_location"))
        self.output_path_edit.setStyleSheet("font-family: 'Tajawal'; font-size: 14px;")
        output_browse_btn = QPushButton(self.localization.get_text("browse"))
        output_browse_btn.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; font-weight: bold;")
        output_browse_btn.clicked.connect(self.browse_output_file)
        output_label = QLabel(self.localization.get_text("output_file") + ":")
        output_label.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; font-weight: bold;")
        output_layout_h.addWidget(output_label)
        output_layout_h.addWidget(self.output_path_edit)
        output_layout_h.addWidget(output_browse_btn)
        output_layout.addLayout(output_layout_h)

        layout.addWidget(output_group)

        # Process button with localization
        self.process_btn = QPushButton(self.localization.get_text("remove_security"))
        self.process_btn.setStyleSheet("font-family: 'Tajawal'; font-size: 16px; font-weight: bold; padding: 10px;")
        self.process_btn.clicked.connect(self.remove_security)
        self.process_btn.setEnabled(False)
        layout.addWidget(self.process_btn)

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("font-family: 'Tajawal'; font-size: 14px;")
        layout.addWidget(self.progress_bar)

        # Status label with localization
        self.status_label = QLabel("")
        self.status_label.setWordWrap(True)
        self.status_label.setStyleSheet("font-family: 'Tajawal'; font-size: 14px; padding: 10px;")
        layout.addWidget(self.status_label)

        layout.addStretch()

    def browse_pdf_file(self):
        """Browse for PDF file with localization"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            "PDF Files (*.pdf)"
        )
        if file_path:
            self.pdf_path = file_path
            self.pdf_path_edit.setText(file_path)

            # Auto-generate output filename
            base_name = os.path.splitext(file_path)[0]
            output_path = f"{base_name}_unlocked.pdf"
            self.output_path_edit.setText(output_path)

            self.check_ready()

    def browse_output_file(self):
        """Browse for output file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "حفظ ملف PDF",
            "",
            "PDF Files (*.pdf)"
        )
        if file_path:
            self.output_path_edit.setText(file_path)
            self.check_ready()

    def check_ready(self):
        """Check if ready to process"""
        ready = (self.pdf_path and
                self.output_path_edit.text().strip() and
                os.path.exists(self.pdf_path))
        self.process_btn.setEnabled(ready)

    def remove_security(self):
        """Remove PDF security restrictions"""
        if not self.pdf_path or not self.output_path_edit.text().strip():
            QMessageBox.warning(self, "تحذير", "يرجى اختيار الملفات أولاً")
            return

        output_path = self.output_path_edit.text().strip()
        password = self.password_edit.text().strip()

        # Show progress with localization
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)  # Indeterminate progress
        self.process_btn.setEnabled(False)
        self.status_label.setText(self.localization.get_text("processing"))

        try:
            # Create PDF tools instance
            pdf_tools = PDFTools(self.history_manager)

            # Remove security
            success = pdf_tools.remove_pdf_security(self.pdf_path, output_path, password)

            if success:
                success_msg = f"✅ {self.localization.get_text('security_removed')}\n{self.localization.get_text('output_file')}: {output_path}"
                self.status_label.setText(success_msg)
                QMessageBox.information(
                    self,
                    self.localization.get_text("success"),
                    self.localization.get_text("security_removed")
                )
            else:
                failure_msg = f"❌ {self.localization.get_text('security_removal_failed')}"
                self.status_label.setText(failure_msg)
                QMessageBox.critical(
                    self,
                    self.localization.get_text("error"),
                    self.localization.get_text("security_removal_failed")
                )

        except Exception as e:
            error_msg = f"{self.localization.get_text('error')}: {str(e)}"
            self.status_label.setText(f"❌ {error_msg}")
            QMessageBox.critical(self, self.localization.get_text("error"), error_msg)

        finally:
            # Hide progress and re-enable button
            self.progress_bar.setVisible(False)
            self.process_btn.setEnabled(True)


class SettingsTab(QWidget):
    """Settings tab for theme and language"""

    def __init__(self, localization: Localization, settings: Settings, main_window=None):
        super().__init__()
        self.localization = localization
        self.settings = settings
        self.main_window = main_window

        self.init_ui()

    def init_ui(self):
        """Initialize settings tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("settings_title"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #1976D2; padding: 8px; background-color: #E3F2FD; border-radius: 4px; font-size: 14px; font-weight: bold;")
        layout.addWidget(instructions)

        # Appearance settings
        appearance_group = QGroupBox(self.localization.get_text("appearance"))
        appearance_layout = QVBoxLayout(appearance_group)

        # Theme selection
        theme_layout = QHBoxLayout()
        theme_layout.addWidget(QLabel(self.localization.get_text("theme")))

        self.theme_combo = QComboBox()
        # Enhanced styling for theme combo
        self.theme_combo.setStyleSheet("""
            QComboBox {
                font-size: 14px;
                padding: 6px 12px;
                min-width: 120px;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                width: 20px;
            }
            QComboBox QAbstractItemView {
                font-size: 14px;
                min-width: 150px;
            }
        """)
        self.theme_combo.addItem(self.localization.get_text("light_mode"), "light")
        self.theme_combo.addItem(self.localization.get_text("dark_mode"), "dark")

        # Set current theme
        current_theme = self.settings.get("theme", "light")
        theme_index = 0 if current_theme == "light" else 1
        self.theme_combo.setCurrentIndex(theme_index)

        theme_layout.addWidget(self.theme_combo)
        theme_layout.addStretch()
        appearance_layout.addLayout(theme_layout)

        # Font size selection
        font_layout = QHBoxLayout()
        font_label_text = "حجم الخط:" if self.localization.current_language == "ar" else "Font Size:"
        font_layout.addWidget(QLabel(font_label_text))

        self.font_size_combo = QComboBox()
        # Enhanced styling for font size combo
        self.font_size_combo.setStyleSheet("""
            QComboBox {
                font-size: 14px;
                padding: 6px 12px;
                min-width: 120px;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                width: 20px;
            }
            QComboBox QAbstractItemView {
                font-size: 14px;
                min-width: 150px;
            }
        """)

        # Font size options
        if self.localization.current_language == "ar":
            self.font_size_combo.addItem("صغير (80%)", "small")
            self.font_size_combo.addItem("متوسط (100%)", "medium")
            self.font_size_combo.addItem("كبير (120%)", "large")
            self.font_size_combo.addItem("كبير جداً (140%)", "extra_large")
        else:
            self.font_size_combo.addItem("Small (80%)", "small")
            self.font_size_combo.addItem("Medium (100%)", "medium")
            self.font_size_combo.addItem("Large (120%)", "large")
            self.font_size_combo.addItem("Extra Large (140%)", "extra_large")

        # Set current font size
        current_font_size = self.settings.get("font_size", "medium")
        font_size_options = ["small", "medium", "large", "extra_large"]
        if current_font_size in font_size_options:
            font_size_index = font_size_options.index(current_font_size)
        else:
            font_size_index = 1  # Default to medium
        self.font_size_combo.setCurrentIndex(font_size_index)

        font_layout.addWidget(self.font_size_combo)
        font_layout.addStretch()
        appearance_layout.addLayout(font_layout)

        layout.addWidget(appearance_group)

        # Language settings
        language_group = QGroupBox(self.localization.get_text("language"))
        language_layout = QVBoxLayout(language_group)

        # Language selection
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel(self.localization.get_text("language")))

        self.language_combo = QComboBox()
        # Enhanced styling for language combo
        self.language_combo.setStyleSheet("""
            QComboBox {
                font-size: 14px;
                padding: 6px 12px;
                min-width: 120px;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                width: 20px;
            }
            QComboBox QAbstractItemView {
                font-size: 14px;
                min-width: 150px;
            }
        """)
        self.language_combo.addItem(self.localization.get_text("arabic"), "ar")
        self.language_combo.addItem(self.localization.get_text("english"), "en")

        # Set current language
        current_lang = self.settings.get("language", "ar")
        lang_index = 0 if current_lang == "ar" else 1
        self.language_combo.setCurrentIndex(lang_index)

        lang_layout.addWidget(self.language_combo)
        lang_layout.addStretch()
        language_layout.addLayout(lang_layout)

        layout.addWidget(language_group)

        # Weekend settings
        weekend_group = QGroupBox(self.localization.get_text("weekend_settings"))
        weekend_layout = QVBoxLayout(weekend_group)

        # Weekend days selection
        weekend_label = QLabel(self.localization.get_text("weekend_days_label"))
        weekend_layout.addWidget(weekend_label)

        # Checkboxes for each day of the week
        days_layout = QHBoxLayout()
        self.day_checkboxes = []
        day_names = ["monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"]

        current_weekend_days = self.settings.get("weekend_days", [5, 6])

        for i, day_name in enumerate(day_names):
            checkbox = QCheckBox(self.localization.get_text(day_name))
            checkbox.setProperty("day_index", i)  # Monday=0, Sunday=6
            if i in current_weekend_days:
                checkbox.setChecked(True)
            checkbox.setStyleSheet("""
                QCheckBox {
                    font-size: 12px;
                    spacing: 5px;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                    border: 2px solid #2196F3;
                    border-radius: 3px;
                    background-color: white;
                }
                QCheckBox::indicator:hover {
                    border: 2px solid #1976D2;
                    background-color: #E3F2FD;
                }
                QCheckBox::indicator:checked {
                    background-color: #2196F3;
                    border: 2px solid #1976D2;
                }
            """)
            self.day_checkboxes.append(checkbox)
            days_layout.addWidget(checkbox)

        days_layout.addStretch()
        weekend_layout.addLayout(days_layout)

        layout.addWidget(weekend_group)

        # Apply button
        self.btn_apply = QPushButton(self.localization.get_text("apply_settings"))
        self.btn_apply.clicked.connect(self.apply_settings)
        self.btn_apply.setStyleSheet("QPushButton { background-color: #1976D2; color: white; font-weight: bold; padding: 10px; }")
        layout.addWidget(self.btn_apply)

        # Settings info
        info_group = QGroupBox("ℹ️ Information")
        info_layout = QVBoxLayout(info_group)

        if self.localization.current_language == "ar":
            info_text = QLabel(
                "<b>معلومات الإعدادات:</b><br>"
                "• السمة: اختر بين الوضع الفاتح والداكن<br>"
                "• اللغة: العربية أو الإنجليزية<br>"
                "• إعادة التشغيل: مطلوبة لتطبيق تغييرات اللغة<br>"
                "• الحفظ التلقائي: يتم حفظ الإعدادات تلقائياً"
            )
        else:
            info_text = QLabel(
                "<b>Settings Information:</b><br>"
                "• Theme: Choose between light and dark mode<br>"
                "• Language: Arabic or English<br>"
                "• Restart: Required for language changes<br>"
                "• Auto-save: Settings are saved automatically"
            )

        info_text.setWordWrap(True)
        info_text.setStyleSheet("color: #666; padding: 10px;")
        info_layout.addWidget(info_text)

        layout.addWidget(info_group)

        # Add stretch
        layout.addStretch()

    def apply_settings(self):
        """Apply settings changes"""
        # Get selected values
        new_theme = self.theme_combo.currentData()
        new_language = self.language_combo.currentData()

        # Get selected weekend days
        weekend_days = []
        for checkbox in self.day_checkboxes:
            if checkbox.isChecked():
                weekend_days.append(checkbox.property("day_index"))

        # Check if language changed
        language_changed = new_language != self.settings.get("language")

        # Save settings
        self.settings.set("theme", new_theme)
        self.settings.set("language", new_language)
        self.settings.set("weekend_days", weekend_days)

        # Apply theme immediately
        if self.main_window:
            self.main_window.apply_theme(new_theme)
        else:
            self.apply_theme(new_theme)

        # Show restart message if language changed
        if language_changed:
            if self.localization.current_language == "ar":
                title = "إعادة التشغيل مطلوبة"
                message = "يجب إعادة تشغيل التطبيق لتطبيق تغييرات اللغة. هل تريد إعادة التشغيل الآن؟"
            else:
                title = "Restart Required"
                message = "Application needs to restart to apply language changes. Restart now?"

            reply = QMessageBox.question(
                self, title, message,
                QMessageBox.Yes | QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                # Restart application
                QApplication.quit()
                os.execl(sys.executable, sys.executable, *sys.argv)
        else:
            # Show success message
            success_msg = "تم تطبيق الإعدادات بنجاح!" if self.localization.current_language == "ar" else "Settings applied successfully!"
            QMessageBox.information(self, self.localization.get_text("success"), success_msg)

    def apply_theme(self, theme: str):
        """Apply theme to application"""
        app = QApplication.instance()
        if theme == "dark":
            # Enhanced Dark theme stylesheet with better text visibility and font integration
            font_family = get_font_family_css()
            dark_style = f"""
            * {{
                {font_family}
            }}
            QMainWindow {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QWidget {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QTabWidget::pane {
                border: 1px solid #555555;
                background-color: #1e1e1e;
                color: #ffffff;
            }
            QTabBar::tab {
                background-color: #2d2d2d;
                color: #ffffff;
                padding: 8px 16px;
                margin: 2px;
                border-radius: 4px;
                border: 1px solid #555555;
            }
            QTabBar::tab:selected {
                background-color: #1976D2;
                color: #ffffff;
            }
            QTabBar::tab:hover {
                background-color: #404040;
                color: #ffffff;
            }
            QGroupBox {
                border: 2px solid #555555;
                border-radius: 5px;
                margin: 10px 0px;
                padding-top: 15px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-weight: bold;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
                color: #ffffff;
                font-weight: bold;
                font-size: 12px;
                background-color: #2d2d2d;
            }
            QLineEdit {
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                selection-background-color: #1976D2;
            }
            QLineEdit:focus {
                border-color: #1976D2;
                background-color: #353535;
            }
            QPushButton {
                border: 2px solid #555555;
                border-radius: 6px;
                padding: 8px 16px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-weight: bold;
                font-size: 11px;
                min-height: 20px;
            }
            QPushButton:hover {
                background-color: #404040;
                border-color: #777777;
                color: #ffffff;
            }
            QPushButton:pressed {
                background-color: #1976D2;
                color: #ffffff;
            }
            QComboBox {
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                min-height: 20px;
            }
            QComboBox:hover {
                border-color: #777777;
                background-color: #353535;
            }
            QComboBox::drop-down {
                border: none;
                background-color: #404040;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #ffffff;
            }
            QComboBox QAbstractItemView {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                selection-background-color: #1976D2;
            }
            QTableWidget {
                background-color: #2d2d2d;
                color: #ffffff;
                gridline-color: #555555;
                border: 1px solid #555555;
                font-size: 11px;
            }
            QTableWidget::item {
                padding: 8px;
                color: #ffffff;
                background-color: #2d2d2d;
            }
            QTableWidget::item:selected {
                background-color: #1976D2;
                color: #ffffff;
            }
            QHeaderView::section {
                background-color: #404040;
                color: #ffffff;
                padding: 8px;
                border: 1px solid #555555;
                font-weight: bold;
                font-size: 11px;
            }
            QTextEdit {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                font-size: 11px;
                padding: 8px;
            }
            QTextEdit:focus {
                border-color: #1976D2;
            }
            QListWidget {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                font-size: 11px;
            }
            QListWidget::item {
                padding: 8px;
                color: #ffffff;
                border-bottom: 1px solid #404040;
            }
            QListWidget::item:selected {
                background-color: #1976D2;
                color: #ffffff;
            }
            QListWidget::item:hover {
                background-color: #404040;
                color: #ffffff;
            }
            QSpinBox {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                font-size: 11px;
                min-height: 20px;
            }
            QSpinBox:focus {
                border-color: #1976D2;
            }
            QSpinBox::up-button, QSpinBox::down-button {
                background-color: #404040;
                border: 1px solid #555555;
                color: #ffffff;
            }
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {
                background-color: #505050;
            }
            QLabel {
                color: #ffffff;
                font-size: 11px;
                font-weight: normal;
                background-color: transparent;
            }
            QStatusBar {
                background-color: #2d2d2d;
                color: #ffffff;
                border-top: 1px solid #555555;
                font-size: 11px;
            }
            QCheckBox {
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 3px;
            }
            QCheckBox::indicator:checked {
                background-color: #1976D2;
                border-color: #1976D2;
            }
            QRadioButton {
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 8px;
            }
            QRadioButton::indicator:checked {
                background-color: #1976D2;
                border-color: #1976D2;
            }
            QScrollBar:vertical {
                background-color: #1e1e1e;
                width: 14px;
                border: none;
                border-radius: 7px;
                margin: 2px;
            }
            QScrollBar::handle:vertical {
                background-color: #4a4a4a;
                border-radius: 6px;
                min-height: 30px;
                margin: 2px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #5a5a5a;
            }
            QScrollBar::handle:vertical:pressed {
                background-color: #1976D2;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background: none;
            }
            QScrollBar:horizontal {
                background-color: #1e1e1e;
                height: 14px;
                border: none;
                border-radius: 7px;
                margin: 2px;
            }
            QScrollBar::handle:horizontal {
                background-color: #4a4a4a;
                border-radius: 6px;
                min-width: 30px;
                margin: 2px;
            }
            QScrollBar::handle:horizontal:hover {
                background-color: #5a5a5a;
            }
            QScrollBar::handle:horizontal:pressed {
                background-color: #1976D2;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {
                background: none;
            }
            QProgressBar {
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 4px;
                text-align: center;
                color: #ffffff;
                font-size: 11px;
            }
            QProgressBar::chunk {
                background-color: #1976D2;
                border-radius: 2px;
            }
            /* Navigation and Home Page Styling */
            SectionCard {
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 12px;
                margin: 8px;
                color: #ffffff;
            }
            SectionCard:hover {
                border-color: #1976D2;
                background-color: #404040;
            }
            BookCard {
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 12px;
                margin: 8px;
                color: #ffffff;
            }
            BookCard:hover {
                border-color: #1976D2;
                background-color: #404040;
            }
            /* Additional UI Components Dark Theme */
            QSpinBox {
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                min-height: 20px;
            }
            QSpinBox:hover {
                border-color: #777777;
                background-color: #353535;
            }
            QSpinBox::up-button, QSpinBox::down-button {
                background-color: #404040;
                border: 1px solid #555555;
                width: 16px;
            }
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {
                background-color: #555555;
            }
            QCheckBox {
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                border: 2px solid #555555;
                border-radius: 3px;
                background-color: #2d2d2d;
            }
            QCheckBox::indicator:checked {
                background-color: #1976D2;
                border-color: #1976D2;
            }
            QRadioButton {
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
                border: 2px solid #555555;
                border-radius: 8px;
                background-color: #2d2d2d;
            }
            QRadioButton::indicator:checked {
                background-color: #1976D2;
                border-color: #1976D2;
            }
            QSlider::groove:horizontal {
                border: 1px solid #555555;
                height: 8px;
                background-color: #2d2d2d;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background-color: #1976D2;
                border: 2px solid #555555;
                width: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }
            QSlider::handle:horizontal:hover {
                background-color: #404040;
            }
            QMessageBox {
                background-color: #2d2d2d;
                color: #ffffff;
            }
            QMessageBox QLabel {
                color: #ffffff;
                font-size: 12px;
            }
            QDialog {
                background-color: #2d2d2d;
                color: #ffffff;
            }
            QFileDialog {
                background-color: #2d2d2d;
                color: #ffffff;
            }
            QMenu {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
            }
            QMenu::item {
                padding: 8px 16px;
                background-color: transparent;
            }
            QMenu::item:selected {
                background-color: #1976D2;
                color: #ffffff;
            }
            QMenuBar {
                background-color: #2d2d2d;
                color: #ffffff;
                border-bottom: 1px solid #555555;
            }
            QMenuBar::item {
                padding: 8px 12px;
                background-color: transparent;
            }
            QMenuBar::item:selected {
                background-color: #404040;
            }
            QToolBar {
                background-color: #2d2d2d;
                border: 1px solid #555555;
                spacing: 3px;
                color: #ffffff;
            }
            QToolButton {
                background-color: #2d2d2d;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 6px;
                color: #ffffff;
            }
            QToolButton:hover {
                background-color: #404040;
                border-color: #777777;
            }
            QToolButton:pressed {
                background-color: #1976D2;
            }
            QStatusBar {
                background-color: #2d2d2d;
                color: #ffffff;
                border-top: 1px solid #555555;
            }
            QSplitter::handle {
                background-color: #555555;
            }
            QSplitter::handle:hover {
                background-color: #777777;
            }
            """
            app.setStyleSheet(dark_style)
        else:
            # Light theme (default)
            app.setStyleSheet("")


class HistoryTab(QWidget):
    """History tab for viewing operation history"""

    def __init__(self, localization: Localization, history_manager: HistoryManager):
        super().__init__()
        self.localization = localization
        self.history_manager = history_manager

        self.init_ui()

    def init_ui(self):
        """Initialize history tab UI"""
        layout = QVBoxLayout(self)

        # Instructions
        instructions = QLabel(self.localization.get_text("history_title"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #795548; padding: 8px; background-color: #EFEBE9; border-radius: 4px; font-size: 14px; font-weight: bold;")
        layout.addWidget(instructions)

        # History controls
        controls_layout = QHBoxLayout()

        self.btn_refresh = QPushButton("🔄 " + ("تحديث" if self.localization.current_language == "ar" else "Refresh"))
        self.btn_refresh.clicked.connect(self.refresh_history)
        controls_layout.addWidget(self.btn_refresh)

        self.btn_clear = QPushButton(self.localization.get_text("clear_history"))
        self.btn_clear.clicked.connect(self.clear_history)
        self.btn_clear.setStyleSheet("QPushButton { background-color: #F44336; color: white; font-weight: bold; }")
        controls_layout.addWidget(self.btn_clear)

        self.btn_export = QPushButton(self.localization.get_text("export_history"))
        self.btn_export.clicked.connect(self.export_history)
        self.btn_export.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; }")
        controls_layout.addWidget(self.btn_export)

        controls_layout.addStretch()
        layout.addLayout(controls_layout)

        # History table
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(6)

        headers = [
            self.localization.get_text("timestamp"),
            self.localization.get_text("operation"),
            self.localization.get_text("input_files"),
            self.localization.get_text("output_file"),
            self.localization.get_text("status"),
            self.localization.get_text("details")
        ]
        self.history_table.setHorizontalHeaderLabels(headers)

        # Set column widths
        self.history_table.setColumnWidth(0, 150)  # Timestamp
        self.history_table.setColumnWidth(1, 150)  # Operation
        self.history_table.setColumnWidth(2, 200)  # Input files
        self.history_table.setColumnWidth(3, 200)  # Output file
        self.history_table.setColumnWidth(4, 100)  # Status
        self.history_table.setColumnWidth(5, 200)  # Details

        layout.addWidget(self.history_table)

        # Control buttons
        buttons_layout = QHBoxLayout()

        # Refresh button
        refresh_btn = QPushButton("🔄 " + self.localization.get_text("refresh"))
        refresh_btn.clicked.connect(self.refresh_history)
        refresh_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        buttons_layout.addWidget(refresh_btn)

        # Clear history button
        clear_btn = QPushButton("🗑️ " + self.localization.get_text("clear_history"))
        clear_btn.clicked.connect(self.clear_history)
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #F44336;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #D32F2F;
            }
        """)
        buttons_layout.addWidget(clear_btn)

        # Add test entry button (for testing)
        test_btn = QPushButton("🧪 Add Test Entry")
        test_btn.clicked.connect(self.add_test_entry)
        test_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        buttons_layout.addWidget(test_btn)

        buttons_layout.addStretch()
        layout.addLayout(buttons_layout)

        # Load history
        self.refresh_history()

    def refresh_history(self):
        """Refresh history display"""
        history = self.history_manager.history
        self.history_table.setRowCount(len(history))

        if not history:
            # Show no history message
            self.history_table.setRowCount(1)
            no_history_item = QTableWidgetItem(self.localization.get_text("no_history"))
            no_history_item.setTextAlignment(Qt.AlignCenter)
            self.history_table.setItem(0, 0, no_history_item)
            self.history_table.setSpan(0, 0, 1, 6)
            return

        for i, entry in enumerate(reversed(history)):  # Show newest first
            # Timestamp
            self.history_table.setItem(i, 0, QTableWidgetItem(entry.timestamp))

            # Operation
            self.history_table.setItem(i, 1, QTableWidgetItem(entry.operation))

            # Input files
            input_files_text = ", ".join([os.path.basename(f) for f in entry.input_files])
            if len(input_files_text) > 50:
                input_files_text = input_files_text[:47] + "..."
            self.history_table.setItem(i, 2, QTableWidgetItem(input_files_text))

            # Output file
            output_file_text = os.path.basename(entry.output_file) if entry.output_file else "-"
            self.history_table.setItem(i, 3, QTableWidgetItem(output_file_text))

            # Status
            status_item = QTableWidgetItem(entry.status)
            if entry.status == "Success" or "نجح" in entry.status:
                status_item.setBackground(Qt.green)
            elif entry.status == "Failed" or "فشل" in entry.status:
                status_item.setBackground(Qt.red)
            self.history_table.setItem(i, 4, status_item)

            # Details
            details_text = entry.details if len(entry.details) <= 50 else entry.details[:47] + "..."
            self.history_table.setItem(i, 5, QTableWidgetItem(details_text))

    def clear_history(self):
        """Clear all history"""
        reply = QMessageBox.question(
            self, "Clear History",
            "Are you sure you want to clear all history?",
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            self.history_manager.clear_history()
            self.refresh_history()
            success_msg = "تم مسح السجل بنجاح!" if self.localization.current_language == "ar" else "History cleared successfully!"
            QMessageBox.information(self, self.localization.get_text("success"), success_msg)

    def add_test_entry(self):
        """Add a test history entry for testing"""
        import random
        operations = ["PDF Merge", "PDF Split", "Extract Images", "Add Watermark", "Compress PDF"]
        statuses = ["Success", "Failed"]

        operation = random.choice(operations)
        status = random.choice(statuses)

        self.history_manager.add_entry(
            operation=operation,
            input_files=["test_file.pdf"],
            output_file="output_file.pdf" if status == "Success" else "",
            status=status,
            details=f"Test entry for {operation}"
        )

        self.refresh_history()
        QMessageBox.information(self, "Test Entry", f"Added test entry: {operation} - {status}")

    def clear_history(self):
        """Clear all history"""
        if self.localization.current_language == "ar":
            title = "مسح السجل"
            message = "هل أنت متأكد من أنك تريد مسح جميع سجلات العمليات؟"
        else:
            title = "Clear History"
            message = "Are you sure you want to clear all operation history?"

        reply = QMessageBox.question(
            self, title, message,
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            self.history_manager.clear_history()
            self.refresh_history()

            success_msg = "تم مسح السجل بنجاح" if self.localization.current_language == "ar" else "History cleared successfully"
            QMessageBox.information(self, self.localization.get_text("success"), success_msg)

    def export_history(self):
        """Export history to file"""
        default_name = "pdf_tools_history.txt"
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.localization.get_text("export_history"),
            default_name,
            "Text Files (*.txt);;All Files (*)"
        )

        if file_path:
            success = self.history_manager.export_history(file_path)
            if success:
                success_msg = "تم تصدير السجل بنجاح" if self.localization.current_language == "ar" else "History exported successfully"
                QMessageBox.information(self, self.localization.get_text("success"), success_msg)
            else:
                error_msg = "فشل في تصدير السجل" if self.localization.current_language == "ar" else "Failed to export history"
                QMessageBox.critical(self, self.localization.get_text("error"), error_msg)


class RichTextAnnotationDialog(QDialog):
    """Enhanced dialog for creating rich text annotations with formatting options"""

    def __init__(self, parent=None, localization=None, initial_text=""):
        super().__init__(parent)
        self.localization = localization
        self.setWindowTitle(self.localization.get_text("text_annotation") if localization else "Text Annotation")
        self.setModal(True)
        self.resize(500, 400)

        # Annotation properties
        self.text_content = initial_text
        self.font_family = "Helvetica"
        self.font_size = 12
        self.is_bold = False
        self.is_italic = False
        self.text_color = QColor(0, 0, 0)  # Black
        self.bg_color = QColor(255, 255, 255)  # White
        self.text_alignment = fitz.TEXT_ALIGN_LEFT

        self.setup_ui()

    def setup_ui(self):
        """Setup the dialog UI"""
        layout = QVBoxLayout(self)
        layout.setSpacing(10)

        # Title
        title_label = QLabel("📝 " + (self.localization.get_text("rich_text_annotation") if self.localization else "Rich Text Annotation"))
        title_label.setStyleSheet("font-size: 14px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title_label)

        # Formatting toolbar
        toolbar_frame = QFrame()
        toolbar_frame.setFrameStyle(QFrame.StyledPanel)
        toolbar_layout = QHBoxLayout(toolbar_frame)
        toolbar_layout.setContentsMargins(5, 5, 5, 5)

        # Font family combo
        font_label = QLabel("Font:")
        toolbar_layout.addWidget(font_label)

        self.font_combo = QComboBox()
        self.font_combo.addItems(["Helvetica", "Times", "Courier", "Arial"])
        self.font_combo.setCurrentText(self.font_family)
        self.font_combo.currentTextChanged.connect(self.on_font_changed)
        toolbar_layout.addWidget(self.font_combo)

        # Font size spinner
        size_label = QLabel("Size:")
        toolbar_layout.addWidget(size_label)

        self.size_spinner = QSpinBox()
        self.size_spinner.setRange(8, 72)
        self.size_spinner.setValue(self.font_size)
        self.size_spinner.valueChanged.connect(self.on_size_changed)
        toolbar_layout.addWidget(self.size_spinner)

        toolbar_layout.addSpacing(10)

        # Bold button
        self.bold_btn = QPushButton("B")
        self.bold_btn.setCheckable(True)
        self.bold_btn.setMaximumWidth(30)
        self.bold_btn.setStyleSheet("font-weight: bold;")
        self.bold_btn.toggled.connect(self.on_bold_toggled)
        toolbar_layout.addWidget(self.bold_btn)

        # Italic button
        self.italic_btn = QPushButton("I")
        self.italic_btn.setCheckable(True)
        self.italic_btn.setMaximumWidth(30)
        self.italic_btn.setStyleSheet("font-style: italic;")
        self.italic_btn.toggled.connect(self.on_italic_toggled)
        toolbar_layout.addWidget(self.italic_btn)

        toolbar_layout.addSpacing(10)

        # Text color button
        self.text_color_btn = QPushButton("🎨 Text")
        self.text_color_btn.clicked.connect(self.choose_text_color)
        toolbar_layout.addWidget(self.text_color_btn)

        # Background color button
        self.bg_color_btn = QPushButton("🎨 BG")
        self.bg_color_btn.clicked.connect(self.choose_bg_color)
        toolbar_layout.addWidget(self.bg_color_btn)

        toolbar_layout.addSpacing(10)

        # Alignment buttons
        align_label = QLabel("Align:")
        toolbar_layout.addWidget(align_label)

        self.align_left_btn = QPushButton("⬅")
        self.align_left_btn.setCheckable(True)
        self.align_left_btn.setChecked(True)
        self.align_left_btn.setMaximumWidth(30)
        self.align_left_btn.clicked.connect(lambda: self.set_alignment(fitz.TEXT_ALIGN_LEFT))
        toolbar_layout.addWidget(self.align_left_btn)

        self.align_center_btn = QPushButton("⬌")
        self.align_center_btn.setCheckable(True)
        self.align_center_btn.setMaximumWidth(30)
        self.align_center_btn.clicked.connect(lambda: self.set_alignment(fitz.TEXT_ALIGN_CENTER))
        toolbar_layout.addWidget(self.align_center_btn)

        self.align_right_btn = QPushButton("➡")
        self.align_right_btn.setCheckable(True)
        self.align_right_btn.setMaximumWidth(30)
        self.align_right_btn.clicked.connect(lambda: self.set_alignment(fitz.TEXT_ALIGN_RIGHT))
        toolbar_layout.addWidget(self.align_right_btn)

        toolbar_layout.addStretch()

        layout.addWidget(toolbar_frame)

        # Font presets
        presets_frame = QFrame()
        presets_layout = QHBoxLayout(presets_frame)
        presets_layout.setContentsMargins(5, 5, 5, 5)

        presets_label = QLabel("Quick Presets:")
        presets_layout.addWidget(presets_label)

        title_preset_btn = QPushButton("📰 Title")
        title_preset_btn.setToolTip("18pt Bold")
        title_preset_btn.clicked.connect(lambda: self.apply_preset('title'))
        presets_layout.addWidget(title_preset_btn)

        body_preset_btn = QPushButton("📄 Body")
        body_preset_btn.setToolTip("12pt Regular")
        body_preset_btn.clicked.connect(lambda: self.apply_preset('body'))
        presets_layout.addWidget(body_preset_btn)

        caption_preset_btn = QPushButton("🔤 Caption")
        caption_preset_btn.setToolTip("10pt Italic")
        caption_preset_btn.clicked.connect(lambda: self.apply_preset('caption'))
        presets_layout.addWidget(caption_preset_btn)

        presets_layout.addStretch()

        layout.addWidget(presets_frame)

        # Text input area
        text_label = QLabel("Text Content:")
        layout.addWidget(text_label)

        self.text_edit = QTextEdit()
        self.text_edit.setPlainText(self.text_content)
        self.text_edit.setMinimumHeight(150)
        self.text_edit.textChanged.connect(self.update_preview)

        # Install event filter for keyboard shortcuts
        self.text_edit.installEventFilter(self)

        layout.addWidget(self.text_edit)

        # Preview area
        preview_label = QLabel("Preview:")
        layout.addWidget(preview_label)

        self.preview_label = QLabel()
        self.preview_label.setFrameStyle(QFrame.StyledPanel)
        self.preview_label.setMinimumHeight(60)
        self.preview_label.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.preview_label.setWordWrap(True)
        self.preview_label.setStyleSheet("padding: 5px; background-color: white;")
        layout.addWidget(self.preview_label)

        # Update initial preview
        self.update_preview()

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        ok_btn = QPushButton(self.localization.get_text("ok") if self.localization else "OK")
        ok_btn.clicked.connect(self.accept)
        ok_btn.setMinimumWidth(80)
        button_layout.addWidget(ok_btn)

        cancel_btn = QPushButton(self.localization.get_text("cancel_button") if self.localization else "Cancel")
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setMinimumWidth(80)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)

    def on_font_changed(self, font_name):
        """Handle font family change"""
        self.font_family = font_name
        self.update_preview()

    def on_size_changed(self, size):
        """Handle font size change"""
        self.font_size = size
        self.update_preview()

    def on_bold_toggled(self, checked):
        """Handle bold toggle"""
        self.is_bold = checked
        self.update_preview()

    def on_italic_toggled(self, checked):
        """Handle italic toggle"""
        self.is_italic = checked
        self.update_preview()

    def choose_text_color(self):
        """Choose text color"""
        color = QColorDialog.getColor(self.text_color, self, "Choose Text Color")
        if color.isValid():
            self.text_color = color
            self.text_color_btn.setStyleSheet(f"background-color: {color.name()};")
            self.update_preview()

    def choose_bg_color(self):
        """Choose background color"""
        color = QColorDialog.getColor(self.bg_color, self, "Choose Background Color")
        if color.isValid():
            self.bg_color = color
            self.bg_color_btn.setStyleSheet(f"background-color: {color.name()};")
            self.update_preview()

    def set_alignment(self, alignment):
        """Set text alignment"""
        self.text_alignment = alignment

        # Update button states
        self.align_left_btn.setChecked(alignment == fitz.TEXT_ALIGN_LEFT)
        self.align_center_btn.setChecked(alignment == fitz.TEXT_ALIGN_CENTER)
        self.align_right_btn.setChecked(alignment == fitz.TEXT_ALIGN_RIGHT)

        self.update_preview()

    def apply_preset(self, preset_type):
        """Apply font preset"""
        if preset_type == 'title':
            # Title: 18pt Bold
            self.size_spinner.setValue(18)
            self.bold_btn.setChecked(True)
            self.italic_btn.setChecked(False)
        elif preset_type == 'body':
            # Body: 12pt Regular
            self.size_spinner.setValue(12)
            self.bold_btn.setChecked(False)
            self.italic_btn.setChecked(False)
        elif preset_type == 'caption':
            # Caption: 10pt Italic
            self.size_spinner.setValue(10)
            self.bold_btn.setChecked(False)
            self.italic_btn.setChecked(True)

        self.update_preview()

    def update_preview(self):
        """Update the preview label with current formatting"""
        text = self.text_edit.toPlainText()

        # Build style string
        font_weight = "bold" if self.is_bold else "normal"
        font_style = "italic" if self.is_italic else "normal"

        alignment_map = {
            fitz.TEXT_ALIGN_LEFT: "left",
            fitz.TEXT_ALIGN_CENTER: "center",
            fitz.TEXT_ALIGN_RIGHT: "right"
        }
        text_align = alignment_map.get(self.text_alignment, "left")

        style = f"""
            font-family: {self.font_family};
            font-size: {self.font_size}px;
            font-weight: {font_weight};
            font-style: {font_style};
            color: {self.text_color.name()};
            background-color: {self.bg_color.name()};
            padding: 5px;
            text-align: {text_align};
        """

        self.preview_label.setStyleSheet(style)
        self.preview_label.setText(text if text else "(Preview will appear here)")

    def get_annotation_data(self):
        """Get the annotation data"""
        return {
            'text': self.text_edit.toPlainText(),
            'font_family': self.font_family,
            'font_size': self.font_size,
            'is_bold': self.is_bold,
            'is_italic': self.is_italic,
            'text_color': self.text_color,
            'bg_color': self.bg_color,
            'alignment': self.text_alignment
        }

    def eventFilter(self, obj, event):
        """Handle keyboard shortcuts"""
        if obj == self.text_edit and event.type() == event.Type.KeyPress:
            # Ctrl+B for bold
            if event.key() == Qt.Key_B and event.modifiers() == Qt.ControlModifier:
                self.bold_btn.setChecked(not self.bold_btn.isChecked())
                return True
            # Ctrl+I for italic
            elif event.key() == Qt.Key_I and event.modifiers() == Qt.ControlModifier:
                self.italic_btn.setChecked(not self.italic_btn.isChecked())
                return True

        return super().eventFilter(obj, event)


class InteractivePDFLabel(QLabel):
    """Interactive PDF display label that handles mouse events for annotations"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_viewer = parent
        self.mouse_pressed = False
        self.start_pos = None
        self.current_pos = None
        self.selection_rect = None
        self.freehand_path = []  # Store freehand drawing points
        self.setContextMenuPolicy(Qt.CustomContextMenu)
        self.customContextMenuRequested.connect(self.show_context_menu)

        # Single touch navigation support
        self.single_touch_enabled = True
        self.touch_start_time = 0
        self.touch_threshold_ms = 200  # Maximum time for a tap to be considered a single touch

    def mousePressEvent(self, event):
        """Handle mouse press events"""
        if event.button() == Qt.LeftButton and self.parent_viewer:
            self.mouse_pressed = True
            self.start_pos = event.position().toPoint()

            # Track touch start time for single touch navigation
            import time
            self.touch_start_time = time.time() * 1000  # Convert to milliseconds

            # Check if clicking on existing annotation in select mode
            if self.parent_viewer.current_tool == "select":
                annot = self.get_annotation_at_position(event.position().toPoint())
                if annot:
                    self.edit_annotation(annot)
                    return

            if self.parent_viewer.current_tool == "text":
                self.handle_text_annotation(event.position().toPoint())
                return  # Don't continue with drag for text
            elif self.parent_viewer.current_tool == "sticky_note":
                self.handle_sticky_note(event.position().toPoint())
                return  # Don't continue with drag for sticky note
            elif self.parent_viewer.current_tool == "freehand":
                self.freehand_path = [event.position().toPoint()]

        super().mousePressEvent(event)

    def show_context_menu(self, pos):
        """Show context menu for annotations"""
        if not self.parent_viewer or not self.parent_viewer.pdf_doc:
            return

        # Check if there's an annotation at this position
        annot = self.get_annotation_at_position(pos)
        if not annot:
            return

        menu = QMenu(self)

        # Edit action
        edit_action = menu.addAction("✏️ Edit Annotation")
        edit_action.triggered.connect(lambda: self.edit_annotation(annot))

        # Delete action
        delete_action = menu.addAction("🗑️ Delete Annotation")
        delete_action.triggered.connect(lambda: self.delete_annotation(annot))

        menu.exec(self.mapToGlobal(pos))

    def get_annotation_at_position(self, pos):
        """Get annotation at the given position"""
        if not self.parent_viewer or not self.parent_viewer.pdf_doc:
            return None

        try:
            page = self.parent_viewer.pdf_doc[self.parent_viewer.current_page]

            # Convert screen coordinates to PDF coordinates
            page_rect = page.rect
            display_rect = self.pixmap().rect() if self.pixmap() else page_rect

            scale_x = page_rect.width / (display_rect.width() / self.parent_viewer.zoom_factor)
            scale_y = page_rect.height / (display_rect.height() / self.parent_viewer.zoom_factor)

            pdf_x = pos.x() * scale_x / self.parent_viewer.zoom_factor
            pdf_y = pos.y() * scale_y / self.parent_viewer.zoom_factor

            # Check all annotations on current page
            for annot in page.annots():
                rect = annot.rect
                if rect.contains(fitz.Point(pdf_x, pdf_y)):
                    return annot

            return None
        except Exception as e:
            print(f"Error getting annotation at position: {e}")
            return None

    def edit_annotation(self, annot):
        """Edit an existing annotation"""
        if not annot or annot.type[0] != fitz.PDF_ANNOT_FREE_TEXT:
            return

        # Get current annotation properties
        info = annot.info
        current_text = info.get("content", "")

        # Get annotation formatting
        colors = annot.colors
        text_color = QColor(int(colors.get("stroke", [0,0,0])[0]*255),
                           int(colors.get("stroke", [0,0,0])[1]*255),
                           int(colors.get("stroke", [0,0,0])[2]*255))

        fill_color_rgb = colors.get("fill", [1,1,1])
        bg_color = QColor(int(fill_color_rgb[0]*255),
                         int(fill_color_rgb[1]*255),
                         int(fill_color_rgb[2]*255))

        # Show rich text dialog with current values
        dialog = RichTextAnnotationDialog(
            parent=self,
            localization=self.parent_viewer.localization if hasattr(self.parent_viewer, 'localization') else None,
            initial_text=current_text
        )

        # Pre-fill dialog with current annotation properties
        dialog.text_color = text_color
        dialog.bg_color = bg_color
        dialog.text_color_btn.setStyleSheet(f"background-color: {text_color.name()};")
        dialog.bg_color_btn.setStyleSheet(f"background-color: {bg_color.name()};")
        dialog.update_preview()

        if dialog.exec() == QDialog.Accepted:
            annotation_data = dialog.get_annotation_data()
            if annotation_data['text']:
                # Update annotation
                self.update_annotation(annot, annotation_data)
                self.parent_viewer.display_current_page()

    def update_annotation(self, annot, annotation_data):
        """Update an existing annotation with new data"""
        try:
            # Get annotation rectangle
            rect = annot.rect

            # Delete old annotation
            page = self.parent_viewer.pdf_doc[self.parent_viewer.current_page]
            page.delete_annot(annot)

            # Extract annotation data
            text = annotation_data['text']
            font_family = annotation_data.get('font_family', 'Helvetica')
            font_size = annotation_data.get('font_size', 12)
            is_bold = annotation_data.get('is_bold', False)
            is_italic = annotation_data.get('is_italic', False)
            text_color = annotation_data.get('text_color', QColor(0, 0, 0))
            bg_color = annotation_data.get('bg_color', QColor(255, 255, 255))
            alignment = annotation_data.get('alignment', fitz.TEXT_ALIGN_LEFT)

            # Map font family
            font_map = {'Helvetica': 'helv', 'Times': 'times', 'Courier': 'cour', 'Arial': 'helv'}
            fontname = font_map.get(font_family, 'helv')

            if is_bold and is_italic:
                fontname += '-bi'
            elif is_bold:
                fontname += '-bo'
            elif is_italic:
                fontname += '-it'

            # Convert colors
            text_color_rgb = [text_color.red()/255.0, text_color.green()/255.0, text_color.blue()/255.0]
            bg_color_rgb = [bg_color.red()/255.0, bg_color.green()/255.0, bg_color.blue()/255.0]

            # Create new annotation with updated properties
            new_annot = page.add_freetext_annot(
                rect,
                text,
                fontsize=font_size,
                fontname=fontname,
                text_color=text_color_rgb,
                fill_color=bg_color_rgb,
                align=alignment
            )

            new_annot.set_colors(stroke=text_color_rgb)
            new_annot.set_border(width=1)
            new_annot.update()

        except Exception as e:
            print(f"Error updating annotation: {e}")

    def delete_annotation(self, annot):
        """Delete an annotation"""
        try:
            page = self.parent_viewer.pdf_doc[self.parent_viewer.current_page]
            page.delete_annot(annot)

            # Remove from annotations list
            if hasattr(self.parent_viewer, 'annotations'):
                # Find and remove matching annotation
                self.parent_viewer.annotations = [
                    a for a in self.parent_viewer.annotations
                    if not (a.get('page') == self.parent_viewer.current_page)
                ]

            # Refresh display
            self.parent_viewer.display_current_page()

        except Exception as e:
            print(f"Error deleting annotation: {e}")

    def mouseMoveEvent(self, event):
        """Handle mouse move events"""
        if self.mouse_pressed and self.start_pos:
            self.current_pos = event.position().toPoint()

            if self.parent_viewer.current_tool in ["highlight", "underline", "rectangle", "circle"]:
                self.update()  # Trigger repaint to show selection
            elif self.parent_viewer.current_tool == "freehand":
                self.freehand_path.append(event.position().toPoint())
                self.update()  # Trigger repaint to show drawing
        else:
            # Show tooltip for annotation on hover
            self.show_annotation_tooltip(event.position().toPoint())

        super().mouseMoveEvent(event)

    def show_annotation_tooltip(self, pos):
        """Show tooltip with annotation info on hover"""
        if not self.parent_viewer or not self.parent_viewer.pdf_doc:
            self.setToolTip("")
            return

        annot = self.get_annotation_at_position(pos)
        if annot:
            info = annot.info
            content = info.get("content", "")
            annot_type = annot.type[1]

            tooltip = f"<b>{annot_type}</b>"
            if content:
                # Limit content length for tooltip
                display_content = content[:100] + "..." if len(content) > 100 else content
                tooltip += f"<br>{display_content}"

            self.setToolTip(tooltip)
        else:
            self.setToolTip("")

    def mouseReleaseEvent(self, event):
        """Handle mouse release events"""
        if event.button() == Qt.LeftButton and self.mouse_pressed:
            self.mouse_pressed = False

            # Check for single touch navigation (quick tap without drag)
            import time
            touch_duration = (time.time() * 1000) - self.touch_start_time
            is_quick_tap = touch_duration < self.touch_threshold_ms
            is_minimal_movement = not self.current_pos or (
                self.start_pos and
                abs(self.start_pos.x() - event.position().toPoint().x()) < 10 and
                abs(self.start_pos.y() - event.position().toPoint().y()) < 10
            )

            # Single touch navigation: tap left side = previous page, tap right side = next page
            if (self.single_touch_enabled and is_quick_tap and is_minimal_movement and
                self.parent_viewer and self.parent_viewer.current_tool == "select" and
                self.parent_viewer.pdf_doc):

                # Get tap position relative to widget width
                tap_x = event.position().toPoint().x()
                widget_width = self.width()

                # Left third = previous page, right third = next page, middle third = no action
                if tap_x < widget_width / 3:
                    # Tap on left side - go to previous page
                    if self.parent_viewer.current_page > 0:
                        self.parent_viewer.previous_page()
                        return
                elif tap_x > (widget_width * 2 / 3):
                    # Tap on right side - go to next page
                    if self.parent_viewer.current_page < self.parent_viewer.total_pages - 1:
                        self.parent_viewer.next_page()
                        return

            if self.start_pos and self.current_pos:
                if self.parent_viewer.current_tool == "highlight":
                    self.handle_highlight_annotation()
                elif self.parent_viewer.current_tool == "underline":
                    self.handle_underline_annotation()
                elif self.parent_viewer.current_tool == "rectangle":
                    self.handle_rectangle_annotation()
                elif self.parent_viewer.current_tool == "circle":
                    self.handle_circle_annotation()
                elif self.parent_viewer.current_tool == "arrow":
                    self.handle_arrow_annotation()
                elif self.parent_viewer.current_tool == "freehand":
                    self.handle_freehand_annotation()

            self.start_pos = None
            self.current_pos = None
            self.freehand_path = []
            self.update()

        super().mouseReleaseEvent(event)

    def paintEvent(self, event):
        """Custom paint event to show selection rectangle"""
        super().paintEvent(event)

        if self.mouse_pressed:
            painter = QPainter(self)

            if self.parent_viewer.current_tool == "freehand" and len(self.freehand_path) > 1:
                # Draw freehand path
                pen = QPen(self.parent_viewer.annotation_color, self.parent_viewer.annotation_size)
                painter.setPen(pen)

                for i in range(1, len(self.freehand_path)):
                    painter.drawLine(self.freehand_path[i-1], self.freehand_path[i])

            elif self.start_pos and self.current_pos:
                # Draw selection rectangle
                painter.setPen(QPen(QColor(0, 0, 255, 128), 2, Qt.DashLine))
                painter.setBrush(QBrush(QColor(0, 0, 255, 50)))

                rect = QRectF(self.start_pos, self.current_pos).normalized()
                painter.drawRect(rect)

    def handle_text_annotation(self, pos):
        """Handle text annotation creation with rich text dialog"""
        if not self.parent_viewer:
            return

        # Show rich text annotation dialog
        dialog = RichTextAnnotationDialog(
            parent=self,
            localization=self.parent_viewer.localization if hasattr(self.parent_viewer, 'localization') else None
        )

        if dialog.exec() == QDialog.Accepted:
            annotation_data = dialog.get_annotation_data()
            if annotation_data['text']:
                self.parent_viewer.add_rich_text_annotation(pos.x(), pos.y(), annotation_data)
            # Keep tool active - don't reset to select

    def handle_sticky_note(self, pos):
        """Handle sticky note creation"""
        text, ok = QInputDialog.getText(self, "Sticky Note", "Enter note:")
        if ok and text and self.parent_viewer:
            self.parent_viewer.add_text_annotation(pos.x(), pos.y(), f"Note: {text}")
            # Keep tool active - don't reset to select

    def handle_highlight_annotation(self):
        """Handle highlight annotation creation"""
        if self.parent_viewer:
            rect = QRectF(self.start_pos, self.current_pos).normalized()
            self.parent_viewer.add_highlight_annotation(rect)
            # Keep tool active - don't reset to select

    def handle_underline_annotation(self):
        """Handle underline annotation creation"""
        if self.parent_viewer:
            rect = QRectF(self.start_pos, self.current_pos).normalized()
            self.parent_viewer.add_underline_annotation(rect)
            # Keep tool active - don't reset to select

    def handle_rectangle_annotation(self):
        """Handle rectangle annotation creation"""
        if self.parent_viewer:
            rect = QRectF(self.start_pos, self.current_pos).normalized()
            self.parent_viewer.add_shape_annotation("rectangle", rect)

    def handle_circle_annotation(self):
        """Handle circle annotation creation"""
        if self.parent_viewer:
            rect = QRectF(self.start_pos, self.current_pos).normalized()
            self.parent_viewer.add_shape_annotation("circle", rect)

    def handle_freehand_annotation(self):
        """Handle freehand drawing annotation"""
        if self.parent_viewer and len(self.freehand_path) > 1:
            self.parent_viewer.add_freehand_annotation(self.freehand_path)


class PDFViewerTab(QWidget):
    """PDF Viewer with editing capabilities"""

    def __init__(self, history_manager=None, localization=None, books_manager=None):
        super().__init__()
        self.history_manager = history_manager
        self.localization = localization or Localization()
        self.books_manager = books_manager  # For last page memory
        self.pdf_doc = None
        self.current_page = 0
        self.total_pages = 0
        self.zoom_factor = 1.0
        self.pdf_path = ""
        self.annotations = []  # Store annotations
        self.current_tool = "select"  # Current annotation tool
        self.annotation_color = QColor(255, 255, 0, 128)  # Default yellow highlight
        self.annotation_size = 2

        self.init_ui()

    def __del__(self):
        """Destructor to ensure PDF document is properly closed"""
        self.cleanup_pdf_document()

    def cleanup_pdf_document(self):
        """Close PDF document and free resources"""
        if self.pdf_doc:
            try:
                self.pdf_doc.close()
                self.pdf_doc = None
            except Exception as e:
                print(f"Error closing PDF document: {e}")

    def init_ui(self):
        """Initialize enhanced PDF viewer UI with sidebar and improved layout"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Instructions (collapsible) - initially hidden to maximize viewing space
        self.instructions_widget = QWidget()
        instructions_layout = QVBoxLayout(self.instructions_widget)
        instructions_layout.setContentsMargins(8, 8, 8, 8)

        instructions = QLabel(self.localization.get_text("pdf_viewer_desc"))
        instructions.setWordWrap(True)
        instructions.setStyleSheet("color: #1976D2; padding: 8px; background-color: #E3F2FD; border-radius: 4px;")
        instructions_layout.addWidget(instructions)

        # Hide instructions by default to maximize viewing space
        self.instructions_widget.setVisible(False)

        layout.addWidget(self.instructions_widget)

        # File selection toolbar (compact) - initially hidden when no PDF loaded
        self.file_toolbar = self.create_file_toolbar()
        layout.addWidget(self.file_toolbar)

        # Main content area with bookmarks sidebar
        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Create annotation toolbar with toggle button for UI elements
        toolbar = self.create_annotation_toolbar()
        main_layout.addWidget(toolbar)

        # Add toggle button for instructions/file toolbar
        self.btn_toggle_ui = QPushButton("▼ " + self.localization.get_text("show_controls"))
        self.btn_toggle_ui.setMaximumWidth(150)
        self.btn_toggle_ui.clicked.connect(self.toggle_ui_elements)
        self.btn_toggle_ui.setStyleSheet("""
            QPushButton {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 10px;
            }
            QPushButton:hover {
                background-color: #e9ecef;
            }
        """)
        # Insert toggle button at the top of the toolbar
        toolbar.layout().insertWidget(0, self.btn_toggle_ui)

        # Create horizontal splitter for sidebar and PDF viewer
        self.main_splitter = QSplitter(Qt.Horizontal)
        self.main_splitter.setChildrenCollapsible(True)

        # Create bookmarks sidebar
        self.bookmarks_panel = self.create_bookmarks_panel()
        self.main_splitter.addWidget(self.bookmarks_panel)

        # Create PDF viewer area
        pdf_viewer_widget = QWidget()
        pdf_viewer_layout = QVBoxLayout(pdf_viewer_widget)
        pdf_viewer_layout.setContentsMargins(0, 0, 0, 0)
        self.create_pdf_viewer_area(pdf_viewer_layout)
        self.main_splitter.addWidget(pdf_viewer_widget)

        # Set initial splitter sizes (sidebar: 250px, viewer: rest)
        self.main_splitter.setSizes([250, 800])
        self.main_splitter.setStretchFactor(0, 0)  # Sidebar doesn't stretch
        self.main_splitter.setStretchFactor(1, 1)  # PDF viewer stretches

        main_layout.addWidget(self.main_splitter, 1)

        # Create bookmark analytics card (initially hidden)
        self.bookmark_analytics_card = self.create_bookmark_analytics_card()
        main_layout.addWidget(self.bookmark_analytics_card)

        layout.addWidget(main_widget, 1)

        # Enhanced status/navigation bar at bottom
        status_bar = self.create_enhanced_status_bar()
        layout.addWidget(status_bar)

    def create_annotation_toolbar(self):
        """Create compact annotation toolbar at the top"""
        toolbar_frame = QFrame()
        toolbar_frame.setFrameStyle(QFrame.NoFrame)
        toolbar_frame.setObjectName("AnnotationToolbar")
        toolbar_frame.setMaximumHeight(45)  # More compact
        toolbar_layout = QHBoxLayout(toolbar_frame)
        toolbar_layout.setContentsMargins(5, 3, 5, 3)

        # Annotation Tools Group
        tools_group = QFrame()
        tools_group.setFrameStyle(QFrame.StyledPanel)
        tools_group.setStyleSheet("QFrame { background-color: #f8f9fa; border-radius: 6px; padding: 4px; }")
        tools_layout = QHBoxLayout(tools_group)
        tools_layout.setContentsMargins(6, 4, 6, 4)
        tools_layout.setSpacing(4)

        # Group label
        tools_label = QLabel("🛠️ " + self.localization.get_text("tools"))
        tools_label.setFont(QFont("Arial", 9, QFont.Bold))
        tools_label.setStyleSheet("color: #495057; background: transparent;")
        tools_layout.addWidget(tools_label)

        tools_layout.addWidget(QFrame())  # Spacer

        self.tool_buttons = QButtonGroup()
        self.tool_buttons.setExclusive(True)

        # Enhanced tool buttons with better styling
        tool_style = """
            QPushButton {
                background-color: #ffffff;
                border: 2px solid #dee2e6;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
                padding: 4px;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border-color: #1976D2;
            }
            QPushButton:checked {
                background-color: #1976D2;
                border-color: #1976D2;
                color: white;
            }
        """

        # Select tool
        self.btn_select = QPushButton("🖱️")
        self.btn_select.setCheckable(True)
        self.btn_select.setChecked(True)
        self.btn_select.setToolTip(self.localization.get_text("select_tool"))
        self.btn_select.clicked.connect(lambda checked: self.set_tool("select") if checked else None)
        self.btn_select.setFixedSize(36, 36)
        self.btn_select.setStyleSheet(tool_style)
        self.tool_buttons.addButton(self.btn_select, 0)
        tools_layout.addWidget(self.btn_select)

        # Text annotation
        self.btn_text_annotation = QPushButton("📝")
        self.btn_text_annotation.setCheckable(True)
        self.btn_text_annotation.setToolTip(self.localization.get_text("text_annotation"))
        self.btn_text_annotation.clicked.connect(lambda checked: self.set_tool("text") if checked else None)
        self.btn_text_annotation.setFixedSize(36, 36)
        self.btn_text_annotation.setStyleSheet(tool_style)
        self.tool_buttons.addButton(self.btn_text_annotation, 1)
        tools_layout.addWidget(self.btn_text_annotation)

        # Highlight
        self.btn_highlight = QPushButton("🖍️")
        self.btn_highlight.setCheckable(True)
        self.btn_highlight.setToolTip(self.localization.get_text("highlight_text"))
        self.btn_highlight.clicked.connect(lambda checked: self.set_tool("highlight") if checked else None)
        self.btn_highlight.setFixedSize(36, 36)
        self.btn_highlight.setStyleSheet(tool_style)
        self.tool_buttons.addButton(self.btn_highlight, 2)
        tools_layout.addWidget(self.btn_highlight)

        # Underline
        self.btn_underline = QPushButton("📏")
        self.btn_underline.setCheckable(True)
        self.btn_underline.setToolTip(self.localization.get_text("underline_text"))
        self.btn_underline.clicked.connect(lambda checked: self.set_tool("underline") if checked else None)
        self.btn_underline.setFixedSize(36, 36)
        self.btn_underline.setStyleSheet(tool_style)
        self.tool_buttons.addButton(self.btn_underline, 3)
        tools_layout.addWidget(self.btn_underline)

        toolbar_layout.addWidget(tools_group)

        # Separator
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.VLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #dee2e6;")
        toolbar_layout.addWidget(separator1)

        # Zoom Controls Group
        zoom_group = QFrame()
        zoom_group.setFrameStyle(QFrame.StyledPanel)
        zoom_group.setStyleSheet("QFrame { background-color: #f8f9fa; border-radius: 6px; padding: 4px; }")
        zoom_layout = QHBoxLayout(zoom_group)
        zoom_layout.setContentsMargins(6, 4, 6, 4)
        zoom_layout.setSpacing(4)

        # Group label
        zoom_label_header = QLabel("🔍 " + self.localization.get_text("zoom"))
        zoom_label_header.setFont(QFont("Arial", 9, QFont.Bold))
        zoom_label_header.setStyleSheet("color: #495057; background: transparent;")
        zoom_layout.addWidget(zoom_label_header)

        zoom_layout.addWidget(QFrame())  # Spacer

        # Enhanced zoom button styling with better text centering
        zoom_style = """
            QPushButton {
                background-color: #ffffff;
                border: 2px solid #dee2e6;
                border-radius: 6px;
                font-size: 18px;
                font-weight: bold;
                padding: 0px;
                text-align: center;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border-color: #28a745;
            }
            QPushButton:disabled {
                background-color: #f8f9fa;
                color: #6c757d;
                border-color: #e9ecef;
            }
        """

        self.btn_zoom_out = QPushButton("−")
        self.btn_zoom_out.clicked.connect(self.zoom_out)
        self.btn_zoom_out.setEnabled(False)
        self.btn_zoom_out.setFixedSize(40, 40)
        self.btn_zoom_out.setToolTip(self.localization.get_text("zoom_out"))
        self.btn_zoom_out.setStyleSheet(zoom_style)
        zoom_layout.addWidget(self.btn_zoom_out)

        self.zoom_label = QLabel("100%")
        self.zoom_label.setMinimumWidth(60)
        self.zoom_label.setAlignment(Qt.AlignCenter)
        self.zoom_label.setFont(QFont("Arial", 11, QFont.Bold))
        self.zoom_label.setStyleSheet("color: #495057; background: transparent; padding: 6px;")
        zoom_layout.addWidget(self.zoom_label)

        self.btn_zoom_in = QPushButton("+")
        self.btn_zoom_in.clicked.connect(self.zoom_in)
        self.btn_zoom_in.setEnabled(False)
        self.btn_zoom_in.setFixedSize(40, 40)
        self.btn_zoom_in.setToolTip(self.localization.get_text("zoom_in"))
        self.btn_zoom_in.setStyleSheet(zoom_style)
        zoom_layout.addWidget(self.btn_zoom_in)

        self.btn_fit_width = QPushButton("↔")
        self.btn_fit_width.clicked.connect(self.fit_width)
        self.btn_fit_width.setEnabled(False)
        self.btn_fit_width.setToolTip(self.localization.get_text("fit_width"))
        self.btn_fit_width.setFixedSize(40, 40)
        self.btn_fit_width.setStyleSheet(zoom_style)
        zoom_layout.addWidget(self.btn_fit_width)

        toolbar_layout.addWidget(zoom_group)

        # Separator
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.VLine)
        separator2.setFrameShadow(QFrame.Sunken)
        separator2.setStyleSheet("color: #dee2e6;")
        toolbar_layout.addWidget(separator2)

        # Annotation Properties Group
        props_group = QFrame()
        props_group.setFrameStyle(QFrame.StyledPanel)
        props_group.setStyleSheet("QFrame { background-color: #f8f9fa; border-radius: 6px; padding: 4px; }")
        props_layout = QHBoxLayout(props_group)
        props_layout.setContentsMargins(6, 4, 6, 4)
        props_layout.setSpacing(4)

        # Group label
        props_label = QLabel("🎨 " + self.localization.get_text("style"))
        props_label.setFont(QFont("Arial", 9, QFont.Bold))
        props_label.setStyleSheet("color: #495057; background: transparent;")
        props_layout.addWidget(props_label)

        props_layout.addWidget(QFrame())  # Spacer

        # Color selection
        self.btn_color = QPushButton()
        self.btn_color.setFixedSize(36, 36)
        self.btn_color.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.annotation_color.name()};
                border: 3px solid #ffffff;
                border-radius: 6px;
                outline: 2px solid #dee2e6;
            }}
            QPushButton:hover {{
                outline: 2px solid #1976D2;
            }}
        """)
        self.btn_color.clicked.connect(self.choose_color)
        self.btn_color.setToolTip(self.localization.get_text("annotation_color"))
        props_layout.addWidget(self.btn_color)

        # Size selection
        self.size_spinbox = QSpinBox()
        self.size_spinbox.setMinimum(1)
        self.size_spinbox.setMaximum(20)
        self.size_spinbox.setValue(self.annotation_size)
        self.size_spinbox.valueChanged.connect(self.set_annotation_size)
        self.size_spinbox.setFixedSize(50, 36)
        self.size_spinbox.setStyleSheet("""
            QSpinBox {
                background-color: #ffffff;
                border: 2px solid #dee2e6;
                border-radius: 6px;
                padding: 4px;
                font-weight: bold;
            }
            QSpinBox:hover {
                border-color: #1976D2;
            }
        """)
        self.size_spinbox.setToolTip(self.localization.get_text("annotation_size"))
        props_layout.addWidget(self.size_spinbox)

        toolbar_layout.addWidget(props_group)

        toolbar_layout.addStretch()

        # Actions Group
        actions_group = QFrame()
        actions_group.setFrameStyle(QFrame.StyledPanel)
        actions_group.setStyleSheet("QFrame { background-color: #f8f9fa; border-radius: 6px; padding: 4px; }")
        actions_layout = QHBoxLayout(actions_group)
        actions_layout.setContentsMargins(6, 4, 6, 4)
        actions_layout.setSpacing(6)

        # Enhanced action button styling
        action_style_base = """
            QPushButton {
                border: none;
                border-radius: 6px;
                font-size: 11px;
                font-weight: bold;
                padding: 8px 12px;
                min-width: 80px;
            }
            QPushButton:hover {
                transform: translateY(-1px);
            }
            QPushButton:disabled {
                background-color: #e9ecef;
                color: #6c757d;
            }
        """

        # Load new PDF button (hidden initially)
        self.btn_load_new_pdf = QPushButton("📁 " + self.localization.get_text("load_new"))
        self.btn_load_new_pdf.clicked.connect(self.show_file_selection)
        self.btn_load_new_pdf.setVisible(False)
        self.btn_load_new_pdf.setStyleSheet(action_style_base + """
            QPushButton {
                background-color: #2196F3;
                color: white;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        actions_layout.addWidget(self.btn_load_new_pdf)

        # Save annotations button
        self.btn_save_annotations = QPushButton("💾 " + self.localization.get_text("save_annotations"))
        self.btn_save_annotations.clicked.connect(self.save_annotations)
        self.btn_save_annotations.setEnabled(False)
        self.btn_save_annotations.setStyleSheet(action_style_base + """
            QPushButton {
                background-color: #28a745;
                color: white;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        actions_layout.addWidget(self.btn_save_annotations)

        # Clear annotations button
        self.btn_clear_annotations = QPushButton("🗑️ " + self.localization.get_text("clear_annotations"))
        self.btn_clear_annotations.clicked.connect(self.clear_annotations)
        self.btn_clear_annotations.setEnabled(False)
        self.btn_clear_annotations.setStyleSheet(action_style_base + """
            QPushButton {
                background-color: #dc3545;
                color: white;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        actions_layout.addWidget(self.btn_clear_annotations)

        toolbar_layout.addWidget(actions_group)

        return toolbar_frame

    def create_bookmarks_panel(self):
        """Create collapsible bookmarks sidebar panel"""
        panel = QFrame()
        panel.setFrameStyle(QFrame.StyledPanel)
        panel.setObjectName("BookmarksPanel")
        panel.setMinimumWidth(200)
        panel.setMaximumWidth(400)

        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        # Header with toggle button
        header_layout = QHBoxLayout()

        # Title
        title_label = QLabel("📑 " + self.localization.get_text("bookmarks"))
        title_label.setFont(QFont("Arial", 12, QFont.Bold))
        title_label.setStyleSheet("color: #1976D2; padding: 4px;")
        header_layout.addWidget(title_label)

        header_layout.addStretch()

        # Toggle button
        self.btn_toggle_bookmarks = QPushButton("◀")
        self.btn_toggle_bookmarks.setFixedSize(24, 24)
        self.btn_toggle_bookmarks.setToolTip(self.localization.get_text("hide_bookmarks"))
        self.btn_toggle_bookmarks.clicked.connect(self.toggle_bookmarks_panel)
        self.btn_toggle_bookmarks.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #cccccc;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #1976D2;
            }
        """)
        header_layout.addWidget(self.btn_toggle_bookmarks)

        # Bookmark analytics button
        self.btn_bookmark_analytics = QPushButton("📊")
        self.btn_bookmark_analytics.setFixedSize(24, 24)
        self.btn_bookmark_analytics.setToolTip(self.localization.get_text("show_bookmark_analytics"))
        self.btn_bookmark_analytics.clicked.connect(self.toggle_bookmark_analytics)
        self.btn_bookmark_analytics.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #cccccc;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #28a745;
            }
        """)
        header_layout.addWidget(self.btn_bookmark_analytics)

        # Open in Chapter Weight Analyzer button
        self.btn_open_weight_analyzer = QPushButton("📈")
        self.btn_open_weight_analyzer.setFixedSize(24, 24)
        self.btn_open_weight_analyzer.setToolTip(self.localization.get_text("open_in_weight_analyzer"))
        self.btn_open_weight_analyzer.clicked.connect(self.open_in_weight_analyzer)
        self.btn_open_weight_analyzer.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #cccccc;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #1976D2;
            }
        """)
        header_layout.addWidget(self.btn_open_weight_analyzer)

        layout.addLayout(header_layout)

        # Bookmarks tree widget
        from PySide6.QtWidgets import QTreeWidget, QTreeWidgetItem
        self.bookmarks_tree = QTreeWidget()
        self.bookmarks_tree.setHeaderHidden(True)
        self.bookmarks_tree.setRootIsDecorated(True)
        self.bookmarks_tree.setIndentation(20)
        self.bookmarks_tree.itemClicked.connect(self.on_bookmark_clicked)
        self.bookmarks_tree.setStyleSheet("""
            QTreeWidget {
                border: 1px solid #cccccc;
                border-radius: 4px;
                background-color: #fafafa;
                selection-background-color: #1976D2;
                selection-color: white;
            }
            QTreeWidget::item {
                padding: 4px;
                border: none;
            }
            QTreeWidget::item:hover {
                background-color: #e3f2fd;
            }
            QTreeWidget::item:selected {
                background-color: #1976D2;
                color: white;
            }
        """)
        layout.addWidget(self.bookmarks_tree, 1)

        # Initially hidden - will be shown when PDF is loaded
        panel.setVisible(False)

        return panel

    def create_bookmark_analytics_card(self):
        """Create bookmark weight distribution analytics card"""
        card = QFrame()
        card.setFrameStyle(QFrame.StyledPanel)
        card.setObjectName("BookmarkAnalyticsCard")
        card.setStyleSheet("""
            QFrame#BookmarkAnalyticsCard {
                background-color: #ffffff;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin: 8px;
            }
            QLabel {
                color: #333333;
            }
        """)
        card.setMaximumHeight(300)
        card.setVisible(False)  # Initially hidden

        layout = QVBoxLayout(card)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Header with title and close button
        header_layout = QHBoxLayout()

        title_label = QLabel("📊 " + self.localization.get_text("bookmark_weight_distribution"))
        title_label.setFont(QFont("Arial", 14, QFont.Bold))
        title_label.setStyleSheet("color: #1976D2; padding: 4px;")
        header_layout.addWidget(title_label)

        header_layout.addStretch()

        # Close button
        close_btn = QPushButton("✕")
        close_btn.setFixedSize(24, 24)
        close_btn.setToolTip(self.localization.get_text("close_analytics"))
        close_btn.clicked.connect(self.hide_bookmark_analytics)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 1px solid #cccccc;
                border-radius: 4px;
                font-weight: bold;
                color: #666666;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #dc3545;
                color: #dc3545;
            }
        """)
        header_layout.addWidget(close_btn)

        layout.addLayout(header_layout)

        # Scroll area for bookmark list
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setStyleSheet("""
            QScrollArea {
                border: 1px solid #e0e0e0;
                border-radius: 4px;
                background-color: #fafafa;
            }
        """)

        # Content widget for bookmarks
        self.analytics_content = QWidget()
        self.analytics_layout = QVBoxLayout(self.analytics_content)
        self.analytics_layout.setContentsMargins(8, 8, 8, 8)
        self.analytics_layout.setSpacing(6)

        scroll_area.setWidget(self.analytics_content)
        layout.addWidget(scroll_area, 1)

        return card

    def create_file_toolbar(self):
        """Create compact file selection toolbar"""
        toolbar = QFrame()
        toolbar.setFrameStyle(QFrame.NoFrame)
        toolbar.setObjectName("FileToolbar")
        toolbar.setMaximumHeight(40)

        layout = QHBoxLayout(toolbar)
        layout.setContentsMargins(8, 4, 8, 4)

        # File path display
        self.pdf_display = QLineEdit()
        self.pdf_display.setReadOnly(True)
        self.pdf_display.setPlaceholderText(self.localization.get_text("select_pdf_file"))
        self.pdf_display.setObjectName("FilePathDisplay")
        layout.addWidget(self.pdf_display, 1)

        # Browse button
        self.btn_browse_pdf = QPushButton("📁 " + self.localization.get_text("browse"))
        self.btn_browse_pdf.clicked.connect(self.browse_pdf)
        self.btn_browse_pdf.setObjectName("BrowseButton")
        layout.addWidget(self.btn_browse_pdf)

        # Load new PDF button (hidden initially)
        self.btn_load_new_pdf = QPushButton("📄 " + self.localization.get_text("load_new_pdf"))
        self.btn_load_new_pdf.clicked.connect(self.show_file_selection)
        self.btn_load_new_pdf.setVisible(False)
        self.btn_load_new_pdf.setObjectName("LoadNewPdfButton")
        layout.addWidget(self.btn_load_new_pdf)

        return toolbar



    def create_pdf_viewer_area(self, layout):
        """Create the main PDF viewer area that takes full height"""
        # Add development banner at the top
        dev_banner = QLabel("⚠️ قيد التطوير - Under Development ⚠️")
        dev_banner.setAlignment(Qt.AlignCenter)
        dev_banner.setStyleSheet("""
            QLabel {
                background-color: #FFA726;
                color: #000000;
                font-size: 14px;
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
                margin: 5px;
            }
        """)
        dev_banner.setMaximumHeight(40)
        layout.addWidget(dev_banner)

        # Create scroll area for PDF display
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(False)  # Don't resize widget automatically
        self.scroll_area.setAlignment(Qt.AlignCenter)
        # Styling will be handled by the global theme

        # Enable wheel events for scrolling
        self.scroll_area.wheelEvent = self.scroll_wheel_event

        # PDF display label (interactive) - enlarged for better viewing
        self.pdf_display_label = InteractivePDFLabel(self)
        self.pdf_display_label.setAlignment(Qt.AlignCenter)
        self.pdf_display_label.setStyleSheet("border: none;")  # Remove hardcoded white background
        # Set localized placeholder text
        no_pdf_text = self.localization.get_text("no_pdf_loaded")
        click_browse_text = self.localization.get_text("click_browse_to_load")
        self.pdf_display_label.setText(f"{no_pdf_text}\n{click_browse_text}")
        # Increased minimum size for better PDF viewing experience
        self.pdf_display_label.setMinimumSize(600, 700)

        self.scroll_area.setWidget(self.pdf_display_label)
        layout.addWidget(self.scroll_area, 1)  # Give it stretch factor of 1 to take available space



    def create_enhanced_status_bar(self):
        """Create enhanced status bar with proper dark mode styling"""
        status_bar = QFrame()
        status_bar.setFrameStyle(QFrame.NoFrame)
        status_bar.setObjectName("StatusBar")
        status_bar.setMaximumHeight(40)  # Slightly more compact

        layout = QHBoxLayout(status_bar)
        layout.setContentsMargins(12, 6, 12, 6)

        # Left side - File info
        self.file_info_label = QLabel(self.localization.get_text("no_file_loaded"))
        self.file_info_label.setObjectName("FileInfoLabel")
        layout.addWidget(self.file_info_label)

        layout.addStretch()

        # Center - Page navigation
        nav_widget = QWidget()
        nav_layout = QHBoxLayout(nav_widget)
        nav_layout.setContentsMargins(0, 0, 0, 0)
        nav_layout.setSpacing(8)

        # Previous page
        self.btn_prev_page = QPushButton("◀")
        self.btn_prev_page.clicked.connect(self.previous_page)
        self.btn_prev_page.setEnabled(False)
        self.btn_prev_page.setFixedSize(30, 30)
        self.btn_prev_page.setToolTip(self.localization.get_text("previous_page"))
        self.btn_prev_page.setObjectName("NavButton")
        nav_layout.addWidget(self.btn_prev_page)

        # Page info with input
        page_widget = QWidget()
        page_layout = QHBoxLayout(page_widget)
        page_layout.setContentsMargins(0, 0, 0, 0)
        page_layout.setSpacing(4)

        page_label = QLabel(self.localization.get_text("page"))
        page_label.setAlignment(Qt.AlignCenter)
        page_layout.addWidget(page_label)

        self.page_spinbox = QSpinBox()
        self.page_spinbox.setMinimum(1)
        self.page_spinbox.valueChanged.connect(self.jump_to_page)
        self.page_spinbox.setEnabled(False)
        self.page_spinbox.setFixedSize(70, 30)
        self.page_spinbox.setAlignment(Qt.AlignCenter)
        self.page_spinbox.setObjectName("PageSpinBox")
        page_layout.addWidget(self.page_spinbox)

        self.page_total_label = QLabel("/ 0")
        self.page_total_label.setAlignment(Qt.AlignCenter)
        self.page_total_label.setObjectName("PageTotalLabel")
        # Make page total label clickable for "Go to Page" dialog
        self.page_total_label.mousePressEvent = self.show_go_to_page_dialog
        self.page_total_label.setStyleSheet("QLabel:hover { color: #1976D2; cursor: pointer; }")
        self.page_total_label.setToolTip(self.localization.get_text("click_to_go_to_page"))
        page_layout.addWidget(self.page_total_label)

        nav_layout.addWidget(page_widget)

        # Next page
        self.btn_next_page = QPushButton("▶")
        self.btn_next_page.clicked.connect(self.next_page)
        self.btn_next_page.setEnabled(False)
        self.btn_next_page.setFixedSize(30, 30)
        self.btn_next_page.setToolTip(self.localization.get_text("next_page"))
        self.btn_next_page.setObjectName("NavButton")
        nav_layout.addWidget(self.btn_next_page)

        layout.addWidget(nav_widget)

        layout.addStretch()

        # Right side - Additional view controls (zoom controls are in the annotation toolbar)
        controls_widget = QWidget()
        controls_layout = QHBoxLayout(controls_widget)
        controls_layout.setContentsMargins(0, 0, 0, 0)
        controls_layout.setSpacing(8)

        # Note: Zoom controls are now only in the annotation toolbar to avoid conflicts
        # This prevents duplicate button creation and ensures proper functionality

        layout.addWidget(controls_widget)

        return status_bar



    def show_file_selection(self):
        """Show file selection toolbar and hide load new PDF button"""
        self.btn_load_new_pdf.setVisible(False)
        self.btn_browse_pdf.setVisible(True)
        self.pdf_display.clear()
        self.instructions_widget.setVisible(True)



    def browse_pdf(self):
        """Browse for PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            "PDF files (*.pdf)"
        )

        if file_path:
            self.pdf_path = file_path
            self.pdf_display.setText(os.path.basename(file_path))
            self.load_pdf()

    def load_pdf(self):
        """Load and display PDF with enhanced interface"""
        if not self.pdf_path:
            return

        try:
            # Close previous PDF document if exists
            self.cleanup_pdf_document()

            # Open PDF document
            self.pdf_doc = fitz.open(self.pdf_path)
            self.total_pages = self.pdf_doc.page_count
            self.current_page = 0

            # Enable controls
            self.enable_controls(True)

            # Update enhanced status bar
            self.update_enhanced_status_bar()

            # Set fit to width as default view mode
            self.fit_width()

            # Display first page
            self.display_current_page()

            # Update interface for loaded PDF
            self.btn_browse_pdf.setVisible(False)
            self.btn_load_new_pdf.setVisible(True)

            # Hide upper UI elements to maximize viewing space
            self.instructions_widget.setVisible(False)
            self.file_toolbar.setVisible(False)
            # Update toggle button text
            self.btn_toggle_ui.setText("▼ " + self.localization.get_text("show_controls"))

            # Update file info in status bar
            filename = os.path.basename(self.pdf_path)
            file_size = os.path.getsize(self.pdf_path) / (1024 * 1024)  # MB
            self.file_info_label.setText(f"📄 {filename} ({file_size:.1f} MB, {self.total_pages} pages)")

            # Load bookmarks and show bookmarks panel
            self.load_bookmarks()

            # Show bookmarks panel when PDF is loaded (if bookmarks exist)
            if hasattr(self, 'bookmarks_panel') and self.bookmarks_panel:
                self.bookmarks_panel.setVisible(True)
                # Update toggle button
                if hasattr(self, 'btn_toggle_bookmarks'):
                    self.btn_toggle_bookmarks.setText("◀")
                    self.btn_toggle_bookmarks.setToolTip(self.localization.get_text("hide_bookmarks"))

            # Restore last page if available
            self.restore_last_page()

        except Exception as e:
            QMessageBox.critical(self, self.localization.get_text("error_title"), f"{self.localization.get_text('failed_to_load_pdf')}:\n{str(e)}")

    def update_enhanced_status_bar(self):
        """Update the enhanced status bar with current page info"""
        if self.pdf_doc:
            self.page_spinbox.setMaximum(self.total_pages)
            self.page_spinbox.setValue(self.current_page + 1)
            self.page_total_label.setText(f"/ {self.total_pages}")
            self.zoom_label.setText(f"{int(self.zoom_factor * 100)}%")
        else:
            self.page_spinbox.setMaximum(0)
            self.page_spinbox.setValue(0)
            self.page_total_label.setText("/ 0")
            self.zoom_label.setText("100%")



    def jump_to_page_direct(self, page_num):
        """Jump directly to a specific page number (0-based)"""
        if self.pdf_doc and 0 <= page_num < self.total_pages:
            self.current_page = page_num
            self.display_current_page()
            self.update_enhanced_status_bar()

    def enable_controls(self, enabled):
        """Enable/disable navigation and zoom controls"""
        self.btn_prev_page.setEnabled(enabled)
        self.btn_next_page.setEnabled(enabled)
        self.btn_zoom_in.setEnabled(enabled)
        self.btn_zoom_out.setEnabled(enabled)
        self.btn_fit_width.setEnabled(enabled)
        self.page_spinbox.setEnabled(enabled)

        if enabled:
            self.page_spinbox.setMaximum(self.total_pages)
            self.page_spinbox.setValue(1)

    def update_page_info(self):
        """Update page information display"""
        if self.pdf_doc:
            current_display = self.current_page + 1
            self.page_info_label.setText(f"{current_display} / {self.total_pages}")
            self.page_spinbox.setValue(current_display)

    def display_current_page(self):
        """Display the current page"""
        if not self.pdf_doc or self.current_page >= self.total_pages:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Create transformation matrix for zoom
            mat = fitz.Matrix(self.zoom_factor, self.zoom_factor)

            # Render page to pixmap
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")

            # Convert to QPixmap and display
            pixmap = QPixmap()
            pixmap.loadFromData(img_data)

            if not pixmap.isNull():
                self.pdf_display_label.setPixmap(pixmap)
                self.pdf_display_label.resize(pixmap.size())
            else:
                self.pdf_display_label.setText(self.localization.get_text("failed_to_render"))

        except Exception as e:
            self.pdf_display_label.setText(f"Error rendering page: {str(e)}")

    def previous_page(self):
        """Go to previous page"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_enhanced_status_bar()
            self.display_current_page()
            self.update_page_controls()
            self.save_last_page()

    def next_page(self):
        """Go to next page"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_enhanced_status_bar()
            self.display_current_page()
            self.update_page_controls()
            self.save_last_page()

    def jump_to_page(self, page_num=None):
        """Jump to specific page from spinbox or parameter"""
        if page_num is None:
            page_num = self.page_spinbox.value()

        if 1 <= page_num <= self.total_pages:
            self.current_page = page_num - 1
            self.update_enhanced_status_bar()
            self.display_current_page()
            self.update_page_controls()
            self.save_last_page()

    def zoom_in(self):
        """Zoom in"""
        self.zoom_factor = min(self.zoom_factor * 1.25, 5.0)
        self.update_enhanced_status_bar()
        self.display_current_page()

    def zoom_out(self):
        """Zoom out"""
        self.zoom_factor = max(self.zoom_factor / 1.25, 0.25)
        self.update_enhanced_status_bar()
        self.display_current_page()

    def update_zoom_display(self):
        """Update zoom display (legacy method for compatibility)"""
        zoom_percent = int(self.zoom_factor * 100)
        if hasattr(self, 'zoom_label'):
            self.zoom_label.setText(f"{zoom_percent}%")

    def fit_width(self):
        """Fit page width to viewer"""
        if not self.pdf_doc:
            return

        # Get page dimensions
        page = self.pdf_doc[self.current_page]
        page_rect = page.rect

        # Get available width
        available_width = self.scroll_area.size().width() - 50

        # Calculate zoom factor to fit width
        self.zoom_factor = min(available_width / page_rect.width, 5.0)

        self.update_enhanced_status_bar()
        self.display_current_page()

    def scroll_wheel_event(self, event):
        """Handle mouse wheel events for page navigation and zooming"""
        # Check if Ctrl is pressed for zooming
        if event.modifiers() & Qt.ControlModifier:
            # Zoom with Ctrl+Wheel
            if event.angleDelta().y() > 0:
                self.zoom_in()
            else:
                self.zoom_out()
        else:
            # Get current scroll position
            v_scrollbar = self.scroll_area.verticalScrollBar()
            current_pos = v_scrollbar.value()
            max_pos = v_scrollbar.maximum()

            if event.angleDelta().y() < 0:  # Scrolling down
                # If at bottom of current page, go to next page
                if current_pos >= max_pos - 10:  # Small threshold
                    if self.current_page < self.total_pages - 1:
                        self.next_page()
                        # Scroll to top of new page
                        v_scrollbar.setValue(0)
                        return
                # Otherwise, normal scroll
                v_scrollbar.setValue(current_pos + 50)
            else:  # Scrolling up
                # If at top of current page, go to previous page
                if current_pos <= 10:  # Small threshold
                    if self.current_page > 0:
                        self.previous_page()
                        # Scroll to bottom of new page
                        QTimer.singleShot(50, lambda: v_scrollbar.setValue(v_scrollbar.maximum()))
                        return
                # Otherwise, normal scroll
                v_scrollbar.setValue(current_pos - 50)

    def set_tool(self, tool_name):
        """Set the current annotation tool"""
        self.current_tool = tool_name

        # Update button states
        if tool_name == "select":
            self.btn_select.setChecked(True)
        elif tool_name == "text":
            self.btn_text_annotation.setChecked(True)
        elif tool_name == "highlight":
            self.btn_highlight.setChecked(True)
        elif tool_name == "underline":
            self.btn_underline.setChecked(True)

        # Update cursor or visual feedback based on tool
        if tool_name == "select":
            self.pdf_display_label.setCursor(Qt.ArrowCursor)
        elif tool_name in ["text", "sticky_note"]:
            self.pdf_display_label.setCursor(Qt.IBeamCursor)
        elif tool_name in ["rectangle", "circle", "arrow", "freehand"]:
            self.pdf_display_label.setCursor(Qt.CrossCursor)
        elif tool_name in ["highlight", "underline"]:
            self.pdf_display_label.setCursor(Qt.PointingHandCursor)

    def choose_color(self):
        """Open color dialog to choose annotation color"""
        color = QColorDialog.getColor(self.annotation_color, self)
        if color.isValid():
            self.annotation_color = color
            self.btn_color.setStyleSheet(f"background-color: {color.name()}; border: 1px solid black;")

    def set_annotation_size(self, size):
        """Set annotation size"""
        self.annotation_size = size

    def save_annotations(self):
        """Save annotations to the same PDF file"""
        if not self.pdf_doc or not self.annotations:
            return

        try:
            # Save PDF with annotations to the same file
            self.pdf_doc.save(self.pdf_path, incremental=True)

            QMessageBox.information(
                self,
                "Success",
                f"Annotations saved to:\n{os.path.basename(self.pdf_path)}"
            )

            # Add to history if available
            if self.history_manager:
                self.history_manager.add_entry(
                    operation="PDF Annotation",
                    input_files=[self.pdf_path],
                    output_file=self.pdf_path,
                    status="Success",
                    details=f"Added {len(self.annotations)} annotations"
                )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save annotations:\n{str(e)}")

    def clear_annotations(self):
        """Clear all annotations"""
        if self.annotations:
            reply = QMessageBox.question(
                self,
                "Clear Annotations",
                "Are you sure you want to clear all annotations?",
                QMessageBox.Yes | QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                self.annotations.clear()
                self.display_current_page()  # Refresh display
                self.btn_save_annotations.setEnabled(False)
                self.btn_clear_annotations.setEnabled(False)

    def add_text_annotation(self, x, y, text):
        """Add free text annotation that displays directly on the PDF page (legacy method)"""
        # Convert to rich text annotation with default settings
        annotation_data = {
            'text': text,
            'font_family': 'Helvetica',
            'font_size': 12,
            'is_bold': False,
            'is_italic': False,
            'text_color': self.annotation_color,
            'bg_color': QColor(255, 255, 255),
            'alignment': fitz.TEXT_ALIGN_LEFT
        }
        self.add_rich_text_annotation(x, y, annotation_data)

    def add_rich_text_annotation(self, x, y, annotation_data):
        """Add rich text annotation with formatting options"""
        if not self.pdf_doc:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Convert screen coordinates to PDF coordinates
            page_rect = page.rect
            display_rect = self.pdf_display_label.pixmap().rect() if self.pdf_display_label.pixmap() else page_rect

            # Calculate scale factors
            scale_x = page_rect.width / (display_rect.width() / self.zoom_factor)
            scale_y = page_rect.height / (display_rect.height() / self.zoom_factor)

            pdf_x = x * scale_x / self.zoom_factor
            pdf_y = y * scale_y / self.zoom_factor

            # Extract annotation data
            text = annotation_data['text']
            font_family = annotation_data.get('font_family', 'Helvetica')
            font_size = annotation_data.get('font_size', 12)
            is_bold = annotation_data.get('is_bold', False)
            is_italic = annotation_data.get('is_italic', False)
            text_color = annotation_data.get('text_color', QColor(0, 0, 0))
            bg_color = annotation_data.get('bg_color', QColor(255, 255, 255))
            alignment = annotation_data.get('alignment', fitz.TEXT_ALIGN_LEFT)

            # Map font family to PyMuPDF font names
            font_map = {
                'Helvetica': 'helv',
                'Times': 'times',
                'Courier': 'cour',
                'Arial': 'helv'  # Use Helvetica for Arial
            }
            fontname = font_map.get(font_family, 'helv')

            # Add bold/italic modifiers
            if is_bold and is_italic:
                fontname += '-bi'
            elif is_bold:
                fontname += '-bo'
            elif is_italic:
                fontname += '-it'

            # Calculate text box size based on text content
            # Count lines and estimate width
            lines = text.split('\n')
            max_line_length = max(len(line) for line in lines) if lines else 0
            num_lines = len(lines)

            # Estimate dimensions (adjust based on font size)
            char_width = font_size * 0.6  # Approximate character width
            line_height = font_size * 1.5  # Line height with spacing

            text_width = max(100, max_line_length * char_width)
            text_height = max(30, num_lines * line_height + 10)  # Add padding

            # Create rectangle for text annotation
            text_rect = fitz.Rect(pdf_x, pdf_y, pdf_x + text_width, pdf_y + text_height)

            # Convert QColor to RGB list (0-1 range)
            text_color_rgb = [text_color.red()/255.0, text_color.green()/255.0, text_color.blue()/255.0]
            bg_color_rgb = [bg_color.red()/255.0, bg_color.green()/255.0, bg_color.blue()/255.0]

            # Add free text annotation with rich formatting
            text_annot = page.add_freetext_annot(
                text_rect,
                text,
                fontsize=font_size,
                fontname=fontname,
                text_color=text_color_rgb,
                fill_color=bg_color_rgb,
                align=alignment
            )

            # Set border color (use text color for border)
            text_annot.set_colors(stroke=text_color_rgb)
            text_annot.set_border(width=1)
            text_annot.update()

            # Store annotation info
            self.annotations.append({
                'type': 'freetext',
                'page': self.current_page,
                'position': (pdf_x, pdf_y),
                'content': text,
                'font_family': font_family,
                'font_size': font_size,
                'is_bold': is_bold,
                'is_italic': is_italic,
                'text_color': text_color.name(),
                'bg_color': bg_color.name(),
                'alignment': alignment
            })

            # Enable save/clear buttons
            self.btn_save_annotations.setEnabled(True)
            self.btn_clear_annotations.setEnabled(True)

            # Refresh display
            self.display_current_page()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add text annotation:\n{str(e)}")

    def add_highlight_annotation(self, rect):
        """Add highlight annotation"""
        if not self.pdf_doc:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Convert screen coordinates to PDF coordinates
            # Get the page dimensions
            page_rect = page.rect
            display_rect = self.pdf_display_label.pixmap().rect() if self.pdf_display_label.pixmap() else page_rect

            # Calculate scale factors
            scale_x = page_rect.width / (display_rect.width() / self.zoom_factor)
            scale_y = page_rect.height / (display_rect.height() / self.zoom_factor)

            pdf_rect = fitz.Rect(
                rect.x() * scale_x / self.zoom_factor,
                rect.y() * scale_y / self.zoom_factor,
                (rect.x() + rect.width()) * scale_x / self.zoom_factor,
                (rect.y() + rect.height()) * scale_y / self.zoom_factor
            )

            # Create highlight annotation
            highlight = page.add_highlight_annot(pdf_rect)
            # Set color using RGB values (0-1 range)
            color_rgb = [c/255.0 for c in self.annotation_color.getRgb()[:3]]
            highlight.set_colors(stroke=color_rgb)
            highlight.update()

            # Store annotation info
            self.annotations.append({
                'type': 'highlight',
                'page': self.current_page,
                'rect': (pdf_rect.x0, pdf_rect.y0, pdf_rect.x1, pdf_rect.y1),
                'color': self.annotation_color.name()
            })

            # Enable save/clear buttons
            self.btn_save_annotations.setEnabled(True)
            self.btn_clear_annotations.setEnabled(True)

            # Refresh display
            self.display_current_page()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add highlight:\n{str(e)}")

    def add_underline_annotation(self, rect):
        """Add underline annotation"""
        if not self.pdf_doc:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Convert screen coordinates to PDF coordinates
            # Get the page dimensions
            page_rect = page.rect
            display_rect = self.pdf_display_label.pixmap().rect() if self.pdf_display_label.pixmap() else page_rect

            # Calculate scale factors
            scale_x = page_rect.width / (display_rect.width() / self.zoom_factor)
            scale_y = page_rect.height / (display_rect.height() / self.zoom_factor)

            pdf_rect = fitz.Rect(
                rect.x() * scale_x / self.zoom_factor,
                rect.y() * scale_y / self.zoom_factor,
                (rect.x() + rect.width()) * scale_x / self.zoom_factor,
                (rect.y() + rect.height()) * scale_y / self.zoom_factor
            )

            # Create underline annotation
            underline = page.add_underline_annot(pdf_rect)
            # Set color using RGB values (0-1 range)
            color_rgb = [c/255.0 for c in self.annotation_color.getRgb()[:3]]
            underline.set_colors(stroke=color_rgb)
            underline.update()

            # Store annotation info
            self.annotations.append({
                'type': 'underline',
                'page': self.current_page,
                'rect': (pdf_rect.x0, pdf_rect.y0, pdf_rect.x1, pdf_rect.y1),
                'color': self.annotation_color.name()
            })

            # Enable save/clear buttons
            self.btn_save_annotations.setEnabled(True)
            self.btn_clear_annotations.setEnabled(True)

            # Refresh display
            self.display_current_page()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add underline:\n{str(e)}")

    def add_shape_annotation(self, shape_type, rect):
        """Add shape annotation (rectangle, circle, arrow)"""
        if not self.pdf_doc:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Convert screen coordinates to PDF coordinates
            pdf_rect = fitz.Rect(
                rect.x() / self.zoom_factor,
                rect.y() / self.zoom_factor,
                (rect.x() + rect.width()) / self.zoom_factor,
                (rect.y() + rect.height()) / self.zoom_factor
            )

            if shape_type == "rectangle":
                # Create rectangle annotation
                annot = page.add_rect_annot(pdf_rect)
            elif shape_type == "circle":
                # Create circle annotation
                annot = page.add_circle_annot(pdf_rect)
            else:
                return  # Unsupported shape type

            # Set annotation properties
            annot.set_colors(stroke=self.annotation_color.getRgbF()[:3])
            annot.set_border(width=self.annotation_size)
            annot.update()

            # Store annotation info
            self.annotations.append({
                'type': shape_type,
                'page': self.current_page,
                'rect': (pdf_rect.x0, pdf_rect.y0, pdf_rect.x1, pdf_rect.y1),
                'color': self.annotation_color.name(),
                'size': self.annotation_size
            })

            # Enable save/clear buttons
            self.btn_save_annotations.setEnabled(True)
            self.btn_clear_annotations.setEnabled(True)

            # Refresh display
            self.display_current_page()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add {shape_type}:\n{str(e)}")

    def add_freehand_annotation(self, path_points):
        """Add freehand drawing annotation"""
        if not self.pdf_doc or not path_points:
            return

        try:
            page = self.pdf_doc[self.current_page]

            # Convert screen coordinates to PDF coordinates
            pdf_points = []
            for point in path_points:
                pdf_x = point.x() / self.zoom_factor
                pdf_y = point.y() / self.zoom_factor
                pdf_points.append(fitz.Point(pdf_x, pdf_y))

            # Create ink annotation (freehand drawing)
            ink_annot = page.add_ink_annot([pdf_points])
            ink_annot.set_colors(stroke=self.annotation_color.getRgbF()[:3])
            ink_annot.set_border(width=self.annotation_size)
            ink_annot.update()

            # Store annotation info
            self.annotations.append({
                'type': 'freehand',
                'page': self.current_page,
                'points': [(p.x, p.y) for p in pdf_points],
                'color': self.annotation_color.name(),
                'size': self.annotation_size
            })

            # Enable save/clear buttons
            self.btn_save_annotations.setEnabled(True)
            self.btn_clear_annotations.setEnabled(True)

            # Refresh display
            self.display_current_page()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add freehand drawing:\n{str(e)}")



    def toggle_ui_elements(self):
        """Toggle the visibility of instructions and file toolbar"""
        is_visible = self.instructions_widget.isVisible()
        self.instructions_widget.setVisible(not is_visible)
        self.file_toolbar.setVisible(not is_visible)

        # Update toggle button
        if is_visible:
            self.btn_toggle_ui.setText("▼ " + self.localization.get_text("show_controls"))
        else:
            self.btn_toggle_ui.setText("▲ " + self.localization.get_text("hide_controls"))

    def toggle_bookmarks_panel(self):
        """Toggle the visibility of the bookmarks panel"""
        if hasattr(self, 'bookmarks_panel'):
            is_visible = self.bookmarks_panel.isVisible()
            self.bookmarks_panel.setVisible(not is_visible)

            # Update toggle button
            if is_visible:
                self.btn_toggle_bookmarks.setText("▶")
                self.btn_toggle_bookmarks.setToolTip(self.localization.get_text("show_bookmarks"))
            else:
                self.btn_toggle_bookmarks.setText("◀")
                self.btn_toggle_bookmarks.setToolTip(self.localization.get_text("hide_bookmarks"))

    def on_bookmark_clicked(self, item, column):
        """Handle bookmark item clicks to navigate to page"""
        if item and hasattr(item, 'page_number'):
            page_num = item.page_number
            if 0 <= page_num < self.total_pages:
                self.current_page = page_num
                self.display_current_page()
                self.update_page_controls()

    def load_bookmarks(self):
        """Load bookmarks from the current PDF document"""
        if not self.pdf_doc:
            return

        try:
            # Clear existing bookmarks
            self.bookmarks_tree.clear()

            # Get table of contents from PDF
            toc = self.pdf_doc.get_toc()

            if not toc:
                # No bookmarks found
                no_bookmarks_item = QTreeWidgetItem(["📄 " + self.localization.get_text("no_bookmarks")])
                no_bookmarks_item.setFlags(Qt.ItemIsEnabled)  # Not selectable
                self.bookmarks_tree.addTopLevelItem(no_bookmarks_item)
                return

            # Build bookmark tree
            parent_stack = []

            for level, title, page_num in toc:
                # Adjust page number (PyMuPDF uses 1-based indexing)
                page_index = max(0, page_num - 1)

                # Create bookmark item
                item_text = f"📑 {title}"
                bookmark_item = QTreeWidgetItem([item_text])
                bookmark_item.page_number = page_index
                bookmark_item.setToolTip(0, f"{title} (Page {page_num})")

                # Determine parent based on level
                if level == 1:
                    # Top level bookmark
                    self.bookmarks_tree.addTopLevelItem(bookmark_item)
                    parent_stack = [bookmark_item]
                else:
                    # Child bookmark - find appropriate parent
                    while len(parent_stack) >= level:
                        parent_stack.pop()

                    if parent_stack:
                        parent_stack[-1].addChild(bookmark_item)
                        parent_stack.append(bookmark_item)
                    else:
                        # Fallback to top level
                        self.bookmarks_tree.addTopLevelItem(bookmark_item)
                        parent_stack = [bookmark_item]

            # Expand all items by default
            self.bookmarks_tree.expandAll()

            # Show bookmarks panel if bookmarks were found
            if hasattr(self, 'bookmarks_panel'):
                self.bookmarks_panel.setVisible(True)
                self.btn_toggle_bookmarks.setText("◀")
                self.btn_toggle_bookmarks.setToolTip(self.localization.get_text("hide_bookmarks"))

        except Exception as e:
            print(f"Error loading bookmarks: {e}")

    def show_go_to_page_dialog(self, event):
        """Show 'Go to Page' dialog when page total label is clicked"""
        if not self.pdf_doc:
            return

        page_num, ok = QInputDialog.getInt(
            self,
            self.localization.get_text("go_to_page"),
            self.localization.get_text("enter_page_number"),
            self.current_page + 1,  # Current page (1-based)
            1,  # Minimum
            self.total_pages  # Maximum
        )

        if ok and 1 <= page_num <= self.total_pages:
            self.current_page = page_num - 1  # Convert to 0-based
            self.display_current_page()
            self.update_page_controls()

    def update_page_controls(self):
        """Update page navigation controls"""
        if self.pdf_doc:
            # Update spinbox
            self.page_spinbox.blockSignals(True)
            self.page_spinbox.setValue(self.current_page + 1)
            self.page_spinbox.blockSignals(False)

            # Update navigation buttons
            self.btn_prev_page.setEnabled(self.current_page > 0)
            self.btn_next_page.setEnabled(self.current_page < self.total_pages - 1)

    def save_last_page(self):
        """Save the current page as the last read page"""
        if self.books_manager and self.pdf_path and self.pdf_doc:
            try:
                # Add book to recent books if not already there
                self.books_manager.add_book(self.pdf_path)

                # Update the last read page
                self.books_manager.update_book_progress(
                    self.pdf_path,
                    pages_read=self.current_page + 1,  # Convert to 1-based
                    total_pages=self.total_pages
                )
            except Exception as e:
                print(f"Error saving last page: {e}")

    def restore_last_page(self):
        """Restore the last read page for this PDF"""
        if self.books_manager and self.pdf_path and self.pdf_doc:
            try:
                # Get book data
                book = self.books_manager.get_book(self.pdf_path)
                if book and hasattr(book, 'pages_read') and book.pages_read > 0:
                    # Convert to 0-based page index
                    last_page = min(book.pages_read - 1, self.total_pages - 1)

                    if last_page > 0:  # Only restore if not on first page
                        self.current_page = last_page
                        self.display_current_page()
                        self.update_page_controls()

                        # Show notification
                        self.show_resume_notification(book.pages_read)

            except Exception as e:
                print(f"Error restoring last page: {e}")

    def show_resume_notification(self, page_num):
        """Show a notification that the PDF was resumed from a specific page"""
        try:
            from PySide6.QtWidgets import QMessageBox

            # Create a non-blocking notification
            msg = QMessageBox(self)
            msg.setWindowTitle(self.localization.get_text("resumed_reading"))
            msg.setText(f"{self.localization.get_text('resumed_from_page')} {page_num}")
            msg.setIcon(QMessageBox.Information)
            msg.setStandardButtons(QMessageBox.Ok)
            msg.setDefaultButton(QMessageBox.Ok)

            # Auto-close after 3 seconds
            QTimer.singleShot(3000, msg.close)
            msg.show()

        except Exception as e:
            print(f"Error showing resume notification: {e}")

    def toggle_bookmark_analytics(self):
        """Toggle the visibility of bookmark analytics card"""
        if hasattr(self, 'bookmark_analytics_card'):
            is_visible = self.bookmark_analytics_card.isVisible()
            if is_visible:
                self.bookmark_analytics_card.setVisible(False)
            else:
                self.show_bookmark_analytics()

    def open_in_weight_analyzer(self):
        """Open current PDF in Chapter Weight Analyzer tab"""
        if not self.pdf_path or not self.pdf_doc:
            QMessageBox.warning(
                self,
                self.localization.get_text("warning"),
                self.localization.get_text("no_pdf_loaded")
            )
            return

        # Check if PDF has bookmarks
        toc = self.pdf_doc.get_toc()
        if not toc:
            QMessageBox.warning(
                self,
                self.localization.get_text("warning"),
                f"{self.localization.get_text('no_bookmarks_found')}\n\n"
                f"⚠️ {self.localization.get_text('pdf_no_bookmarks_error')}"
            )
            return

        try:
            # Find the main window and chapter weight visualizer tab
            main_window = self.parent()
            while main_window and not hasattr(main_window, 'chapter_weight_tab'):
                main_window = main_window.parent()

            if main_window and hasattr(main_window, 'chapter_weight_tab'):
                # Load PDF in chapter weight visualizer
                main_window.chapter_weight_tab.load_pdf_from_external(self.pdf_path)

                # Navigate to chapter weight visualizer tab
                if hasattr(main_window, 'navigation_manager'):
                    main_window.navigation_manager.navigate_to_section("chapter_weight")

                # Show success message
                if hasattr(main_window, 'statusBar'):
                    main_window.statusBar().showMessage(
                        f"{self.localization.get_text('pdf_loaded_in_analyzer')}: {os.path.basename(self.pdf_path)}"
                    )
            else:
                QMessageBox.information(
                    self,
                    self.localization.get_text("info"),
                    self.localization.get_text("weight_analyzer_not_available")
                )
        except Exception as e:
            print(f"Error opening in weight analyzer: {e}")
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"{self.localization.get_text('error')}:\n{str(e)}"
            )

    def show_bookmark_analytics(self):
        """Show and populate bookmark analytics card"""
        if not hasattr(self, 'bookmark_analytics_card') or not self.pdf_doc:
            return

        # Clear existing content
        for i in reversed(range(self.analytics_layout.count())):
            child = self.analytics_layout.itemAt(i).widget()
            if child:
                child.setParent(None)

        try:
            # Get table of contents
            toc = self.pdf_doc.get_toc()

            if not toc:
                # No bookmarks message
                no_data_label = QLabel("📄 No bookmarks found in this PDF")
                no_data_label.setAlignment(Qt.AlignCenter)
                no_data_label.setStyleSheet("color: #666666; font-style: italic; padding: 20px;")
                self.analytics_layout.addWidget(no_data_label)
            else:
                # Calculate bookmark weights
                bookmark_data = self.calculate_bookmark_weights(toc)

                # Create visualization
                for bookmark in bookmark_data:
                    bookmark_widget = self.create_bookmark_weight_widget(bookmark)
                    self.analytics_layout.addWidget(bookmark_widget)

                # Add summary
                summary_widget = self.create_summary_widget(bookmark_data)
                self.analytics_layout.addWidget(summary_widget)

            # Show the card
            self.bookmark_analytics_card.setVisible(True)

        except Exception as e:
            print(f"Error showing bookmark analytics: {e}")

    def hide_bookmark_analytics(self):
        """Hide bookmark analytics card"""
        if hasattr(self, 'bookmark_analytics_card'):
            self.bookmark_analytics_card.setVisible(False)

    def calculate_bookmark_weights(self, toc):
        """Calculate the weight/distribution of bookmarks"""
        bookmark_data = []

        for i, (level, title, page_num) in enumerate(toc):
            # Calculate page range for this bookmark
            start_page = page_num

            # Find next bookmark at same or higher level to determine end page
            end_page = self.total_pages
            for j in range(i + 1, len(toc)):
                next_level, _, next_page = toc[j]
                if next_level <= level:
                    end_page = next_page - 1
                    break

            page_count = max(1, end_page - start_page + 1)
            percentage = (page_count / self.total_pages) * 100

            bookmark_data.append({
                'title': title,
                'level': level,
                'start_page': start_page,
                'end_page': end_page,
                'page_count': page_count,
                'percentage': percentage
            })

        return bookmark_data

    def create_bookmark_weight_widget(self, bookmark):
        """Create a widget showing bookmark weight information"""
        widget = QFrame()
        widget.setFrameStyle(QFrame.StyledPanel)
        widget.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 6px;
                margin: 2px;
                padding: 4px;
            }
        """)

        layout = QVBoxLayout(widget)
        layout.setContentsMargins(8, 6, 8, 6)
        layout.setSpacing(4)

        # Title with indentation based on level
        indent = "  " * (bookmark['level'] - 1)
        title_text = f"{indent}📑 {bookmark['title']}"
        title_label = QLabel(title_text)
        title_label.setFont(QFont("Arial", 10, QFont.Bold))
        title_label.setStyleSheet("color: #495057;")
        layout.addWidget(title_label)

        # Progress bar showing relative weight
        progress_layout = QHBoxLayout()

        # Progress bar
        from PySide6.QtWidgets import QProgressBar
        progress_bar = QProgressBar()
        progress_bar.setRange(0, 100)
        progress_bar.setValue(int(bookmark['percentage']))
        progress_bar.setTextVisible(False)
        progress_bar.setFixedHeight(12)
        progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #dee2e6;
                border-radius: 6px;
                background-color: #e9ecef;
            }
            QProgressBar::chunk {
                background-color: #28a745;
                border-radius: 5px;
            }
        """)
        progress_layout.addWidget(progress_bar, 3)

        # Percentage and page info
        info_text = f"{bookmark['percentage']:.1f}% ({bookmark['page_count']} pages)"
        info_label = QLabel(info_text)
        info_label.setFont(QFont("Arial", 9))
        info_label.setStyleSheet("color: #6c757d;")
        info_label.setAlignment(Qt.AlignRight)
        progress_layout.addWidget(info_label, 1)

        layout.addLayout(progress_layout)

        # Page range info
        range_text = f"Pages {bookmark['start_page']}-{bookmark['end_page']}"
        range_label = QLabel(range_text)
        range_label.setFont(QFont("Arial", 8))
        range_label.setStyleSheet("color: #868e96; font-style: italic;")
        layout.addWidget(range_label)

        return widget

    def create_summary_widget(self, bookmark_data):
        """Create summary widget with overall statistics"""
        widget = QFrame()
        widget.setFrameStyle(QFrame.StyledPanel)
        widget.setStyleSheet("""
            QFrame {
                background-color: #e3f2fd;
                border: 2px solid #1976D2;
                border-radius: 8px;
                margin: 4px;
                padding: 8px;
            }
        """)

        layout = QVBoxLayout(widget)
        layout.setContentsMargins(12, 8, 12, 8)
        layout.setSpacing(4)

        # Summary title
        title_label = QLabel("📈 Summary")
        title_label.setFont(QFont("Arial", 11, QFont.Bold))
        title_label.setStyleSheet("color: #1976D2;")
        layout.addWidget(title_label)

        # Statistics
        total_bookmarks = len(bookmark_data)
        avg_pages = sum(b['page_count'] for b in bookmark_data) / total_bookmarks if total_bookmarks > 0 else 0
        largest_section = max(bookmark_data, key=lambda x: x['page_count']) if bookmark_data else None

        stats_text = f"• Total Bookmarks: {total_bookmarks}\n"
        stats_text += f"• Average Section Size: {avg_pages:.1f} pages\n"
        if largest_section:
            stats_text += f"• Largest Section: {largest_section['title']} ({largest_section['page_count']} pages)"

        stats_label = QLabel(stats_text)
        stats_label.setFont(QFont("Arial", 9))
        stats_label.setStyleSheet("color: #1976D2;")
        layout.addWidget(stats_label)

        return widget

    def closeEvent(self, event):
        """Save last page when closing the viewer"""
        self.save_last_page()
        super().closeEvent(event)


class SplitByBookmarksTab(QWidget):
    """Tab for splitting PDF by bookmarks"""

    def __init__(self, history_manager=None, localization=None, main_window=None):
        super().__init__()
        self.history_manager = history_manager
        self.localization = localization or Localization()
        self.main_window = main_window
        self.pdf_path = None
        self.bookmarks = []
        self.init_ui()

    def init_ui(self):
        """Initialize the user interface"""
        layout = QVBoxLayout(self)
        layout.setSpacing(15)  # Increase spacing between sections
        layout.setContentsMargins(20, 10, 20, 20)  # Reduce top margin, increase others

        # Title - Further increased font size with reduced margin
        title = QLabel(self.localization.get_text("split_by_bookmarks"))
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2196F3; margin: 5px 0px;")  # Increased from 21px to 24px
        layout.addWidget(title)

        # Description - Further increased font size with better line height
        desc = QLabel(self.localization.get_text("split_by_bookmarks_desc"))
        desc.setStyleSheet("color: #666; margin: 2px 0px 10px 0px; font-size: 16px; line-height: 1.5;")  # Increased from 14px to 16px
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # File selection - Further increased font sizes with better spacing
        file_group = QGroupBox(self.localization.get_text("select_pdf"))
        file_group.setStyleSheet("QGroupBox { font-size: 17px; font-weight: bold; padding: 15px; margin: 5px; }")  # Increased from 15px to 17px
        file_layout = QHBoxLayout(file_group)
        file_layout.setSpacing(10)  # Add spacing between elements

        self.file_path_edit = QLineEdit()
        self.file_path_edit.setPlaceholderText(self.localization.get_text("select_pdf_file"))
        self.file_path_edit.setReadOnly(True)
        self.file_path_edit.setStyleSheet("QLineEdit { font-size: 15px; padding: 10px; }")  # Increased from 13px to 15px
        file_layout.addWidget(self.file_path_edit)

        self.btn_browse = QPushButton(self.localization.get_text("browse"))
        self.btn_browse.clicked.connect(self.browse_pdf)
        self.btn_browse.setStyleSheet("QPushButton { font-size: 15px; padding: 10px 18px; }")  # Increased from 13px to 15px
        file_layout.addWidget(self.btn_browse)

        layout.addWidget(file_group)

        # Split options - Further increased font sizes with better spacing
        options_group = QGroupBox(self.localization.get_text("split_options"))
        options_group.setStyleSheet("QGroupBox { font-size: 17px; font-weight: bold; padding: 15px; margin: 5px; }")  # Increased from 15px to 17px
        options_layout = QVBoxLayout(options_group)
        options_layout.setSpacing(12)  # Increase spacing between options

        # Level selection
        level_label = QLabel("📊 " + self.localization.get_text("bookmark_levels") + ":")
        level_label.setStyleSheet("font-weight: bold; margin-top: 5px; font-size: 16px;")  # Increased from 14px to 16px
        options_layout.addWidget(level_label)

        self.level1_radio = QRadioButton("📖 " + self.localization.get_text("split_level_1_only"))
        self.level1_radio.setChecked(True)
        self.level1_radio.setStyleSheet("margin-left: 10px; font-size: 15px; padding: 5px;")  # Increased from 13px to 15px
        options_layout.addWidget(self.level1_radio)

        self.all_levels_radio = QRadioButton("📚 " + self.localization.get_text("split_all_levels"))
        self.all_levels_radio.setStyleSheet("margin-left: 10px; font-size: 15px; padding: 5px;")  # Increased from 13px to 15px
        options_layout.addWidget(self.all_levels_radio)

        # Additional options
        options_layout.addWidget(QLabel(""))  # Spacer

        additional_label = QLabel("⚙️ " + self.localization.get_text("additional_options") + ":")
        additional_label.setStyleSheet("font-weight: bold; margin-top: 5px; font-size: 16px;")  # Increased from 14px to 16px
        options_layout.addWidget(additional_label)

        self.include_toc_checkbox = QCheckBox("📋 " + self.localization.get_text("include_original_bookmarks"))
        self.include_toc_checkbox.setChecked(True)
        self.include_toc_checkbox.setStyleSheet("margin-left: 10px; font-size: 15px; padding: 5px;")  # Increased from 13px to 15px
        options_layout.addWidget(self.include_toc_checkbox)

        self.create_index_checkbox = QCheckBox("📑 " + self.localization.get_text("create_index_file"))
        self.create_index_checkbox.setChecked(False)
        self.create_index_checkbox.setStyleSheet("margin-left: 10px; font-size: 15px; padding: 5px;")  # Increased from 13px to 15px
        options_layout.addWidget(self.create_index_checkbox)

        layout.addWidget(options_group)

        # Output settings - Further increased font sizes with better spacing
        output_group = QGroupBox(self.localization.get_text("output_settings"))
        output_group.setStyleSheet("QGroupBox { font-size: 17px; font-weight: bold; padding: 15px; margin: 5px; }")  # Increased from 15px to 17px
        output_layout = QVBoxLayout(output_group)
        output_layout.setSpacing(12)  # Increase spacing between elements

        # Output directory
        dir_layout = QHBoxLayout()
        dir_label = QLabel(self.localization.get_text("output_directory") + ":")
        dir_label.setStyleSheet("font-size: 15px; font-weight: bold;")  # Increased from 13px to 15px
        dir_layout.addWidget(dir_label)

        self.output_dir_edit = QLineEdit()
        self.output_dir_edit.setPlaceholderText(self.localization.get_text("select_output_directory"))
        self.output_dir_edit.setReadOnly(True)
        self.output_dir_edit.setStyleSheet("QLineEdit { font-size: 15px; padding: 10px; }")  # Increased from 13px to 15px
        dir_layout.addWidget(self.output_dir_edit)

        self.btn_browse_output = QPushButton(self.localization.get_text("browse"))
        self.btn_browse_output.clicked.connect(self.browse_output_dir)
        self.btn_browse_output.setStyleSheet("QPushButton { font-size: 15px; padding: 10px 18px; }")  # Increased from 13px to 15px
        dir_layout.addWidget(self.btn_browse_output)

        output_layout.addLayout(dir_layout)

        # File naming options
        naming_label = QLabel(self.localization.get_text("file_naming") + ":")
        naming_label.setStyleSheet("font-size: 16px; font-weight: bold; margin-top: 10px;")  # Increased from 14px to 16px
        output_layout.addWidget(naming_label)

        self.use_titles_radio = QRadioButton(self.localization.get_text("use_bookmark_titles"))
        self.use_titles_radio.setChecked(True)
        self.use_titles_radio.setStyleSheet("font-size: 15px; margin-left: 10px; padding: 5px;")  # Increased from 13px to 15px
        output_layout.addWidget(self.use_titles_radio)

        self.use_numbers_radio = QRadioButton(self.localization.get_text("use_sequential_numbers"))
        self.use_numbers_radio.setStyleSheet("font-size: 15px; margin-left: 10px; padding: 5px;")  # Increased from 13px to 15px
        output_layout.addWidget(self.use_numbers_radio)

        layout.addWidget(output_group)

        # Split button - Further increased font size
        self.btn_split = QPushButton(self.localization.get_text("start_splitting"))
        self.btn_split.clicked.connect(self.start_splitting)
        self.btn_split.setEnabled(False)
        self.btn_split.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                font-size: 17px;
                padding: 15px;
                border-radius: 5px;
                margin: 10px 0px;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        layout.addWidget(self.btn_split)

        # Progress area - Further increased font size
        self.progress_text = QTextEdit()
        self.progress_text.setMaximumHeight(150)
        self.progress_text.setReadOnly(True)
        self.progress_text.setStyleSheet("QTextEdit { font-size: 15px; line-height: 1.4; }")  # Increased from 13px to 15px
        layout.addWidget(self.progress_text)

    def browse_pdf(self):
        """Browse for PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("select_pdf_file"),
            "",
            "PDF files (*.pdf)"
        )

        if file_path:
            self.pdf_path = file_path
            self.file_path_edit.setText(file_path)
            self.load_bookmarks()

    def browse_output_dir(self):
        """Browse for output directory"""
        dir_path = QFileDialog.getExistingDirectory(
            self,
            self.localization.get_text("select_output_directory")
        )

        if dir_path:
            self.output_dir_edit.setText(dir_path)
            self.update_split_button()

    def load_bookmarks(self):
        """Load bookmarks from PDF"""
        if not self.pdf_path:
            return

        try:
            doc = fitz.open(self.pdf_path)
            toc = doc.get_toc()

            if not toc:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning"),
                    "This PDF does not contain bookmarks. Please add bookmarks first using the Bookmark Manager tab."
                )
                return

            # Convert to Bookmark objects
            self.bookmarks = []
            for item in toc:
                level, title, page = item
                self.bookmarks.append(Bookmark(title=title, page=page, level=level))

            self.progress_text.append(f"✅ Loaded {len(self.bookmarks)} bookmarks from PDF")
            self.update_split_button()

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to load bookmarks: {str(e)}"
            )

    def update_split_button(self):
        """Update split button state"""
        has_pdf = bool(self.pdf_path)
        has_output_dir = bool(self.output_dir_edit.text())
        has_bookmarks = bool(self.bookmarks)

        self.btn_split.setEnabled(has_pdf and has_output_dir and has_bookmarks)

    def start_splitting(self):
        """Start the PDF splitting process"""
        if not self.pdf_path or not self.output_dir_edit.text() or not self.bookmarks:
            return

        try:
            output_dir = self.output_dir_edit.text()
            split_level_1_only = self.level1_radio.isChecked()
            use_titles = self.use_titles_radio.isChecked()
            include_bookmarks = self.include_toc_checkbox.isChecked()
            create_index = self.create_index_checkbox.isChecked()

            # Filter bookmarks based on level selection
            if split_level_1_only:
                split_bookmarks = [b for b in self.bookmarks if b.level == 1]
            else:
                # For "all levels", we need to split at every bookmark level
                split_bookmarks = self.bookmarks

            if not split_bookmarks:
                QMessageBox.warning(
                    self,
                    self.localization.get_text("warning"),
                    "No bookmarks found for the selected level."
                )
                return

            self.progress_text.append(f"🚀 Starting split process...")
            self.progress_text.append(f"📁 Output directory: {output_dir}")
            self.progress_text.append(f"📊 Splitting {len(split_bookmarks)} sections")

            # Perform the split
            doc = fitz.open(self.pdf_path)
            created_files = []  # Track created files for index

            for i, bookmark in enumerate(split_bookmarks):
                # Determine end page using improved logic
                end_page = self.calculate_end_page(bookmark, split_bookmarks, i, doc.page_count, split_level_1_only)

                # Create filename
                if use_titles:
                    # Clean title for filename
                    clean_title = re.sub(r'[<>:"/\\|?*]', '_', bookmark.title)
                    clean_title = clean_title.strip()[:50]  # Limit length
                    filename = f"{clean_title}.pdf"
                else:
                    filename = f"{i+1:03d}_{bookmark.title[:20].replace(' ', '_')}.pdf"

                output_path = os.path.join(output_dir, filename)

                # Create new PDF with pages
                new_doc = fitz.open()
                for page_num in range(bookmark.page - 1, min(end_page, doc.page_count)):
                    new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)

                # Add bookmarks to the new PDF if requested
                if include_bookmarks:
                    # Create TOC for this section with improved logic
                    section_toc = self.create_section_toc(bookmark, end_page, split_level_1_only)

                    if section_toc:
                        new_doc.set_toc(section_toc)

                new_doc.save(output_path)
                new_doc.close()

                # Track created file
                created_files.append({
                    'filename': filename,
                    'title': bookmark.title,
                    'pages': f"{bookmark.page}-{end_page}",
                    'page_count': end_page - bookmark.page + 1
                })

                self.progress_text.append(f"✅ Created: {filename} (Pages {bookmark.page}-{end_page})")

            doc.close()

            # Create index file if requested
            if create_index:
                self.create_index_file(output_dir, created_files)

            self.progress_text.append(f"🎉 Split completed! Created {len(split_bookmarks)} files.")

            # Show success message
            QMessageBox.information(
                self,
                self.localization.get_text("success"),
                f"PDF split successfully!\n\n"
                f"Created {len(split_bookmarks)} files in:\n{output_dir}"
            )

        except Exception as e:
            self.progress_text.append(f"❌ Error: {str(e)}")
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to split PDF: {str(e)}"
            )

    def calculate_end_page(self, current_bookmark, split_bookmarks, current_index, total_pages, split_level_1_only):
        """Calculate the end page for a bookmark section"""
        if split_level_1_only:
            # For Level 1 only: find the next Level 1 bookmark
            for j in range(current_index + 1, len(split_bookmarks)):
                if split_bookmarks[j].level == 1:
                    return split_bookmarks[j].page - 1
            return total_pages
        else:
            # For all levels: find the next bookmark at the same or higher level
            current_level = current_bookmark.level
            for j in range(current_index + 1, len(split_bookmarks)):
                next_bookmark = split_bookmarks[j]
                if next_bookmark.level <= current_level:
                    return next_bookmark.page - 1
            return total_pages

    def create_section_toc(self, section_bookmark, end_page, split_level_1_only):
        """Create TOC for a section with proper bookmark preservation and promotion"""
        section_toc = []

        if split_level_1_only:
            # For Level 1 splits: include all sub-bookmarks and promote them
            for b in self.bookmarks:
                if section_bookmark.page <= b.page <= end_page:
                    # Adjust page numbers for the new document
                    adjusted_page = b.page - section_bookmark.page + 1

                    if b.level == 1:
                        # Keep Level 1 bookmarks as Level 1
                        section_toc.append([1, b.title, adjusted_page])
                    elif b.level == 2:
                        # Promote Level 2 to Level 1 in the split file
                        section_toc.append([1, b.title, adjusted_page])
                    elif b.level >= 3:
                        # Promote Level 3+ to Level 2 in the split file
                        section_toc.append([2, b.title, adjusted_page])
        else:
            # For all levels: preserve relative structure but normalize levels
            section_bookmarks = []
            for b in self.bookmarks:
                if section_bookmark.page <= b.page <= end_page:
                    # Adjust page numbers for the new document
                    adjusted_page = b.page - section_bookmark.page + 1
                    section_bookmarks.append([b.level, b.title, adjusted_page])

            # Normalize levels to ensure first bookmark is level 1
            if section_bookmarks:
                min_level = min(bookmark[0] for bookmark in section_bookmarks)
                for bookmark in section_bookmarks:
                    # Adjust level so the minimum level becomes 1
                    normalized_level = bookmark[0] - min_level + 1
                    section_toc.append([normalized_level, bookmark[1], bookmark[2]])

        return section_toc

    def create_index_file(self, output_dir, created_files):
        """Create an index file listing all split files"""
        try:
            index_path = os.path.join(output_dir, "INDEX.txt")

            with open(index_path, 'w', encoding='utf-8') as f:
                f.write("📚 PDF Split Index File\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"📄 Original PDF: {os.path.basename(self.pdf_path)}\n")
                f.write(f"📅 Split Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"📊 Total Files Created: {len(created_files)}\n\n")

                f.write("📋 File List:\n")
                f.write("-" * 30 + "\n")

                for i, file_info in enumerate(created_files, 1):
                    f.write(f"{i:2d}. {file_info['filename']}\n")
                    f.write(f"    📖 Title: {file_info['title']}\n")
                    f.write(f"    📄 Pages: {file_info['pages']} ({file_info['page_count']} pages)\n\n")

                f.write("\n" + "=" * 50 + "\n")
                f.write("Generated by PDF Tools - Split by Bookmarks\n")

            self.progress_text.append(f"📑 Created index file: INDEX.txt")

        except Exception as e:
            self.progress_text.append(f"⚠️ Warning: Could not create index file: {str(e)}")


class LazyTabLoader:
    """Lazy loader for tabs to improve startup performance"""

    def __init__(self, tab_factory, *args, **kwargs):
        """
        Initialize lazy tab loader

        Args:
            tab_factory: Function or class that creates the tab instance
            *args: Positional arguments to pass to tab_factory
            **kwargs: Keyword arguments to pass to tab_factory
        """
        self.tab_factory = tab_factory
        self.args = args
        self.kwargs = kwargs
        self._instance = None
        self._factory_name = getattr(tab_factory, '__name__', str(tab_factory))

    def get_instance(self):
        """Get or create tab instance"""
        if self._instance is None:
            print(f"📦 Lazy loading {self._factory_name}...")
            self._instance = self.tab_factory(*self.args, **self.kwargs)
            print(f"✅ {self._factory_name} loaded")
        return self._instance

    @property
    def instance(self):
        """Property to get instance"""
        return self.get_instance()

    def is_loaded(self):
        """Check if tab has been loaded"""
        return self._instance is not None


class PDFToolsComprehensive(QMainWindow):
    """Main application window"""

    def __init__(self):
        super().__init__()

        # Track startup time
        import time
        self._startup_time = time.time()
        print("🚀 Starting PDF Tools application...")

        # Initialize settings and localization
        print("⚙️ Initializing settings and localization...")
        self.settings = Settings()
        self.localization = Localization()
        self.localization.set_language(self.settings.get("language", "ar"))
        self.history_manager = HistoryManager(self.settings)

        # Setup application icon
        self.setup_app_icon()

        self.init_ui()

        # Apply saved theme
        self.apply_saved_theme()

        # Report startup time
        elapsed = time.time() - self._startup_time
        print(f"✅ Application started in {elapsed:.2f} seconds")

    def setup_app_icon(self):
        """Setup application icon for main window and dialogs"""
        # Try multiple icon paths (ICO first, then PNG)
        icon_paths = [
            os.path.join(os.path.dirname(__file__), '..', 'assets', 'icons', 'logo.ico'),
            os.path.join(os.path.dirname(__file__), '..', 'assets', 'icons', 'logo.png'),
            os.path.join(os.path.dirname(sys.executable), 'assets', 'icons', 'logo.ico'),
            os.path.join(os.path.dirname(sys.executable), 'assets', 'icons', 'logo.png'),
        ]

        icon = None
        for icon_path in icon_paths:
            if os.path.exists(icon_path):
                icon = QIcon(icon_path)
                if not icon.isNull():
                    break

        if not icon or icon.isNull():
            # Fallback: Create a simple icon if file not found
            icon = QIcon()
            pixmap = QPixmap(64, 64)
            pixmap.fill(Qt.transparent)

            painter = QPainter(pixmap)
            painter.setRenderHint(QPainter.Antialiasing)

            # Draw red book icon background
            painter.setBrush(QBrush(QColor("#E53935")))  # Red color
            painter.setPen(QPen(QColor("#C62828"), 2))
            painter.drawRoundedRect(8, 8, 48, 48, 6, 6)

            # Draw book symbol
            painter.setPen(QPen(QColor("white")))
            painter.setFont(QFont("Arial", 12, QFont.Bold))
            painter.drawText(pixmap.rect(), Qt.AlignCenter, "📖")

            painter.end()
            icon.addPixmap(pixmap)

        self.setWindowIcon(icon)

        # Store icon for use in dialogs
        self.app_icon = icon

    def get_app_icon(self):
        """Get the application icon for use in dialogs"""
        return getattr(self, 'app_icon', QIcon())

    def create_menu_bar(self):
        """Create menu bar with File and Help menus"""
        menubar = self.menuBar()

        # Direct back to home action (no submenu)
        self.back_home_action = menubar.addAction("← " + self.localization.get_text("back_to_home"))
        self.back_home_action.triggered.connect(self.navigate_to_home)
        self.back_home_action.setVisible(False)  # Hidden when on home page

        # Global Reset button (only visible in bookmark manager)
        self.global_reset_action = menubar.addAction("🔄 " + self.localization.get_text("reset_button"))
        self.global_reset_action.triggered.connect(self.global_reset)
        self.global_reset_action.setVisible(False)  # Hidden by default

        # File menu
        file_menu = menubar.addMenu("📁 " + self.localization.get_text("file_menu"))

        # New/Open actions
        open_pdf_action = file_menu.addAction("📄 " + self.localization.get_text("open_pdf"))
        open_pdf_action.setShortcut("Ctrl+O")
        open_pdf_action.triggered.connect(self.open_pdf_file)

        # Recent files submenu
        recent_menu = file_menu.addMenu("📋 " + self.localization.get_text("recent_files"))
        recent_menu.setEnabled(False)  # Will be enabled when we have recent files

        file_menu.addSeparator()

        # Quick access to main tools
        tools_menu = file_menu.addMenu("🛠️ " + self.localization.get_text("quick_tools"))

        bookmark_action = tools_menu.addAction(self.localization.get_text("bookmark_manager"))
        bookmark_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("bookmark_manager"))

        # Add Recent Books menu item (check loader, not tab)
        if self._recent_books_tab_loader:
            recent_books_action = tools_menu.addAction(self.localization.get_text("recent_books"))
            recent_books_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("recent_books"))

        # Add Reading Progress menu item (check loader, not tab)
        if self._reading_progress_tab_loader:
            progress_title = self.localization.get_text("reading_progress")
            progress_action = tools_menu.addAction(progress_title)
            progress_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("reading_progress"))

        merge_action = tools_menu.addAction(self.localization.get_text("merge_pdfs"))
        merge_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("merge_pdfs"))

        # Add Split PDF menu item (check loader, not tab)
        if self._split_tab_loader:
            split_action = tools_menu.addAction(self.localization.get_text("split_pdfs"))
            split_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("split_pdfs"))

        compress_action = tools_menu.addAction(self.localization.get_text("compress"))
        compress_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("compress"))

        file_menu.addSeparator()

        # Settings action
        settings_action = file_menu.addAction(self.localization.get_text("settings"))
        settings_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("settings"))

        # History action
        history_action = file_menu.addAction(self.localization.get_text("history"))
        history_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("history"))

        # Statistics Dashboard action (check loader, not tab)
        if self._statistics_dashboard_tab_loader:
            stats_action = file_menu.addAction(self.localization.get_text("statistics_dashboard"))
            stats_action.triggered.connect(lambda: self.navigation_manager.navigate_to_section("statistics_dashboard"))

        file_menu.addSeparator()

        # Close action
        close_action = file_menu.addAction(self.localization.get_text("close"))
        close_action.setShortcut("Ctrl+W")
        close_action.triggered.connect(self.close)

        # Exit action
        exit_action = file_menu.addAction(self.localization.get_text("exit"))
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)

        # Help menu
        help_menu = menubar.addMenu("❓ " + self.localization.get_text("help_menu"))

        # Check for updates action
        if VERSION_SYSTEM_AVAILABLE:
            update_action = help_menu.addAction(self.localization.get_text("check_for_updates"))
            update_action.triggered.connect(self.check_for_updates)
            help_menu.addSeparator()

        # User Guide action
        guide_action = help_menu.addAction(self.localization.get_text("user_guide"))
        guide_action.setShortcut("F1")
        guide_action.triggered.connect(self.show_user_guide)

        # Keyboard shortcuts
        shortcuts_action = help_menu.addAction(self.localization.get_text("keyboard_shortcuts"))
        shortcuts_action.triggered.connect(self.show_shortcuts)

        help_menu.addSeparator()

        # About action
        about_action = help_menu.addAction(self.localization.get_text("about"))
        about_action.triggered.connect(self.show_about)

    def init_ui(self):
        """Initialize main UI"""
        # Set window title with version
        if VERSION_SYSTEM_AVAILABLE:
            version_info = get_version_info()
            title = f"{self.localization.get_text('app_title')} v{version_info['version']}"
        else:
            title = self.localization.get_text("app_title")
        self.setWindowTitle(title)

        # Load saved geometry or use defaults - fit to screen
        screen = QApplication.primaryScreen().availableGeometry()
        default_width = min(1400, int(screen.width() * 0.9))
        default_height = min(900, int(screen.height() * 0.9))
        geometry = self.settings.get("window_geometry", [50, 50, default_width, default_height])
        self.setGeometry(*geometry)
        self.setMinimumSize(1000, 600)

        # Enable window controls (minimize, maximize, close)
        self.setWindowFlags(Qt.Window | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint | Qt.WindowCloseButtonHint)

        # Set layout direction for Arabic
        if self.localization.current_language == "ar":
            self.setLayoutDirection(Qt.RightToLeft)

        # Initialize tabs first
        self.init_tabs()

        # Create menu bar
        self.create_menu_bar()

        # Central widget with navigation
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Create navigation manager
        from navigation_manager import NavigationManager
        self.navigation_manager = NavigationManager(self.localization)

        # Add all sections to navigation
        self.setup_navigation()

        # Initialize update checker
        if VERSION_SYSTEM_AVAILABLE:
            self.update_checker = UpdateChecker(self, self.localization)
            # Check for updates on startup (silent)
            QTimer.singleShot(3000, lambda: self.update_checker.check_for_updates(silent=True))

        layout.addWidget(self.navigation_manager)

        # Status bar
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage(self.localization.get_text("ready_status"))

    def _add_lazy_section(self, section_key: str, tab_loader: LazyTabLoader, title_key: str):
        """
        Add a section with lazy tab loading

        This method adds a section to the navigation manager using a factory function
        that will create the tab only when first accessed.
        """
        def get_tab():
            """Factory function that returns the tab instance"""
            return tab_loader.get_instance()

        # Add section with factory function instead of actual tab
        self.navigation_manager.add_section_lazy(section_key, get_tab, title_key)

    def setup_navigation(self):
        """Setup navigation with all sections using lazy loading"""
        print("📋 Setting up navigation (lazy mode)...")

        # Add all sections to navigation manager with lazy loading
        # Tabs will only be created when user navigates to them

        self._add_lazy_section("pdf_viewer", self._pdf_viewer_tab_loader, self.localization.get_text("pdf_viewer"))

        # Add Recent Books tab if loader exists
        if self._recent_books_tab_loader:
            self._add_lazy_section("recent_books", self._recent_books_tab_loader, self.localization.get_text("recent_books"))

        # Bookmark-related sections (grouped together)
        self._add_lazy_section("bookmark_manager", self._bookmark_tab_loader, self.localization.get_text("bookmark_manager"))
        self._add_lazy_section("bookmark_extractor", self._bookmark_extractor_tab_loader, self.localization.get_text("bookmark_extractor"))
        self._add_lazy_section("split_by_bookmarks", self._split_by_bookmarks_tab_loader, self.localization.get_text("split_by_bookmarks"))

        # Add Chapter Weight Visualizer tab if loader exists
        if self._chapter_weight_tab_loader:
            self._add_lazy_section("chapter_weight", self._chapter_weight_tab_loader, self.localization.get_text("chapter_weight_analyzer"))

        # Add PDF Reading Progress tab if loader exists
        if self._reading_progress_tab_loader:
            self._add_lazy_section("reading_progress", self._reading_progress_tab_loader, self.localization.get_text("reading_progress"))

        # Add Statistics Dashboard tab if loader exists
        if self._statistics_dashboard_tab_loader:
            self._add_lazy_section("statistics_dashboard", self._statistics_dashboard_tab_loader, self.localization.get_text("statistics_dashboard"))

        # Add Reading Speed tab if loader exists
        if self._reading_speed_tab_loader:
            self._add_lazy_section("reading_speed", self._reading_speed_tab_loader, self.localization.get_text("reading_speed_meter"))

        # Other PDF operations
        self._add_lazy_section("page_operations", self._page_ops_tab_loader, self.localization.get_text("page_operations"))
        self._add_lazy_section("watermark", self._watermark_tab_loader, self.localization.get_text("watermark"))
        self._add_lazy_section("extract_images", self._image_extraction_tab_loader, self.localization.get_text("extract_images"))
        self._add_lazy_section("extract_text", self._text_extraction_tab_loader, self.localization.get_text("extract_text"))
        self._add_lazy_section("merge_pdfs", self._merge_tab_loader, self.localization.get_text("merge_pdfs"))

        if self._split_tab_loader:
            self._add_lazy_section("split_pdfs", self._split_tab_loader, self.localization.get_text("split_pdfs"))

        self._add_lazy_section("compress", self._compress_tab_loader, self.localization.get_text("compress"))
        self._add_lazy_section("page_editing", self._page_editing_tab_loader, self.localization.get_text("page_editing"))
        self._add_lazy_section("security_removal", self._security_removal_tab_loader, "إزالة الحماية")

        # Add Settings and History to navigation manager (for File menu access)
        # These won't appear on home page navigation but can be accessed via File menu
        self._add_lazy_section("settings", self._settings_tab_loader, self.localization.get_text("settings"))
        self._add_lazy_section("history", self._history_tab_loader, self.localization.get_text("history"))

        # Connect navigation signals
        self.navigation_manager.home_page.section_selected.connect(self.on_section_selected)

        print("✅ Navigation setup complete (all tabs will load on demand)")

    def global_reset(self):
        """Global reset functionality for bookmark manager"""
        if hasattr(self, 'bookmark_tab') and self.bookmark_tab:
            # Show confirmation dialog
            if self.localization.current_language == "ar":
                title = "إعادة تعيين شاملة"
                message = "هل أنت متأكد من أنك تريد إعادة تعيين جميع البيانات والعودة إلى البداية؟\n\nسيتم فقدان جميع التقدم الحالي."
                yes_text = "نعم، إعادة تعيين"
                no_text = "إلغاء"
            else:
                title = "Global Reset"
                message = "Are you sure you want to reset all data and start over?\n\nAll current progress will be lost."
                yes_text = "Yes, Reset"
                no_text = "Cancel"

            reply = QMessageBox.question(
                self, title, message,
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                self.bookmark_tab.reset_all_data()

                # Show success message
                if self.localization.current_language == "ar":
                    success_title = "تم إعادة التعيين"
                    success_message = "تم إعادة تعيين جميع البيانات بنجاح. يمكنك الآن البدء من جديد."
                else:
                    success_title = "Reset Complete"
                    success_message = "All data has been reset successfully. You can now start over."

                QMessageBox.information(self, success_title, success_message)

    def navigate_to_home(self):
        """Navigate to home page"""
        self.navigation_manager.navigate_to_home()
        self.back_home_action.setVisible(False)
        self.global_reset_action.setVisible(False)

    def on_section_selected(self, section_key: str):
        """Handle section selection from home page"""
        self.navigation_manager.navigate_to_section(section_key)
        self.back_home_action.setVisible(True)

        # Show global reset button only for bookmark manager
        if section_key == "bookmark_manager":
            self.global_reset_action.setVisible(True)
        else:
            self.global_reset_action.setVisible(False)



    def apply_saved_theme(self):
        """Apply saved theme and font size on startup"""
        theme = self.settings.get("theme", "light")
        font_size = self.settings.get("font_size", "medium")

        # Apply theme first, then font size
        self.apply_theme(theme)


    def apply_theme(self, theme: str):
        """Apply theme to application"""
        app = QApplication.instance()

        if theme == "dark":
            # Enhanced Dark theme stylesheet with better text visibility and font integration
            font_family = get_font_family_css()
            dark_style = f"""
            * {{
                {font_family}
            }}
            QMainWindow {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QWidget {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QTabWidget::pane {{
                border: 1px solid #555555;
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QTabBar::tab {{
                background-color: #2d2d2d;
                color: #ffffff;
                padding: 8px 16px;
                margin: 2px;
                border-radius: 4px;
                border: 1px solid #555555;
                {font_family}
                font-weight: bold;
            }}
            QTabBar::tab:selected {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
                font-weight: bold;
            }}
            QTabBar::tab:hover {{
                background-color: #404040;
                color: #ffffff;
                {font_family}
                font-weight: bold;
            }}
            QGroupBox {{
                border: 2px solid #555555;
                border-radius: 5px;
                margin: 10px 0px;
                padding-top: 15px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-weight: bold;
                {font_family}
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
                color: #ffffff;
                font-weight: bold;
                font-size: 12px;
                background-color: #2d2d2d;
                {font_family}
            }}
            QLineEdit {{
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                selection-background-color: #1976D2;
                {font_family}
            }}
            QLineEdit:focus {{
                border-color: #1976D2;
                background-color: #353535;
                {font_family}
            }}
            QPushButton {{
                border: 2px solid #555555;
                border-radius: 6px;
                padding: 8px 16px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-weight: bold;
                font-size: 11px;
                min-height: 20px;
                {font_family}
            }}
            QPushButton:hover {{
                background-color: #404040;
                border-color: #777777;
                color: #ffffff;
                {font_family}
            }}
            QPushButton:pressed {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
            }}
            QComboBox {{
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                min-height: 20px;
                {font_family}
            }}
            QComboBox:hover {{
                border-color: #777777;
                background-color: #353535;
                {font_family}
            }}
            QComboBox::drop-down {{
                border: none;
                background-color: #404040;
            }}
            QComboBox::down-arrow {{
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #ffffff;
            }}
            QComboBox QAbstractItemView {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                selection-background-color: #1976D2;
                {font_family}
            }}
            QTableWidget {{
                background-color: #2d2d2d;
                color: #ffffff;
                gridline-color: #555555;
                border: 1px solid #555555;
                font-size: 11px;
                {font_family}
            }}
            QTableWidget::item {{
                padding: 8px;
                color: #ffffff;
                background-color: #2d2d2d;
                {font_family}
            }}
            QTableWidget::item:selected {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
            }}
            QHeaderView::section {{
                background-color: #404040;
                color: #ffffff;
                padding: 8px;
                border: 1px solid #555555;
                font-weight: bold;
                font-size: 11px;
                {font_family}
            }}
            QTextEdit {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                font-size: 11px;
                padding: 8px;
                {font_family}
            }}
            QTextEdit:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QListWidget {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                font-size: 11px;
                {font_family}
            }}
            QListWidget::item {{
                padding: 8px;
                color: #ffffff;
                border-bottom: 1px solid #404040;
                {font_family}
            }}
            QListWidget::item:selected {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
            }}
            QListWidget::item:hover {{
                background-color: #404040;
                color: #ffffff;
                {font_family}
            }}
            QSpinBox {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                font-size: 11px;
                min-height: 20px;
                {font_family}
            }}
            QSpinBox:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QSpinBox::up-button, QSpinBox::down-button {{
                background-color: #404040;
                border: 1px solid #555555;
                color: #ffffff;
            }}
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {{
                background-color: #505050;
            }}
            QLabel {{
                color: #ffffff;
                font-size: 11px;
                font-weight: normal;
                background-color: transparent;
                {font_family}
            }}
            QStatusBar {{
                background-color: #2d2d2d;
                color: #ffffff;
                border-top: 1px solid #555555;
                font-size: 11px;
                {font_family}
            }}
            QCheckBox {{
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
                {font_family}
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 3px;
            }}
            QCheckBox::indicator:checked {{
                background-color: #1976D2;
                border-color: #1976D2;
            }}
            QRadioButton {{
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
                {font_family}
            }}
            QRadioButton::indicator {{
                width: 16px;
                height: 16px;
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 8px;
            }}
            QRadioButton::indicator:checked {{
                background-color: #1976D2;
                border-color: #1976D2;
            }}
            QScrollBar:vertical {{
                background-color: #2d2d2d;
                width: 16px;
                border: none;
            }}
            QScrollBar::handle:vertical {{
                background-color: #555555;
                border-radius: 8px;
                min-height: 20px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: #777777;
            }}
            QScrollBar:horizontal {{
                background-color: #2d2d2d;
                height: 16px;
                border: none;
            }}
            QScrollBar::handle:horizontal {{
                background-color: #555555;
                border-radius: 8px;
                min-width: 20px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background-color: #777777;
            }}
            QProgressBar {{
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 4px;
                text-align: center;
                color: #ffffff;
                font-size: 11px;
                {font_family}
            }}
            QProgressBar::chunk {{
                background-color: #1976D2;
                border-radius: 2px;
            }}
            /* Fix for white backgrounds in dark mode */
            QFrame {{
                background-color: #1e1e1e;
                color: #ffffff;
                border: none;
                {font_family}
            }}
            QScrollArea {{
                background-color: #1e1e1e;
                border: none;
            }}
            QScrollArea > QWidget > QWidget {{
                background-color: #1e1e1e;
            }}
            QSplitter {{
                background-color: #1e1e1e;
            }}
            QSplitter::handle {{
                background-color: #555555;
            }}
            QSplitter::handle:horizontal {{
                width: 3px;
            }}
            QSplitter::handle:vertical {{
                height: 3px;
            }}
            /* Ensure all dialogs and message boxes are dark */
            QDialog {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QMessageBox {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QMessageBox QPushButton {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                padding: 6px 12px;
                border-radius: 4px;
                {font_family}
            }}
            QMessageBox QPushButton:hover {{
                background-color: #404040;
                {font_family}
            }}
            /* Fix for toolbar and menu areas */
            QToolBar {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                {font_family}
            }}
            QMenuBar {{
                background-color: #2d2d2d;
                color: #ffffff;
                border-bottom: 1px solid #555555;
                {font_family}
            }}
            QMenuBar::item {{
                background-color: transparent;
                color: #ffffff;
                padding: 4px 8px;
                {font_family}
            }}
            QMenuBar::item:selected {{
                background-color: #404040;
                {font_family}
            }}
            QMenu {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                {font_family}
            }}
            QMenu::item {{
                padding: 6px 20px;
                color: #ffffff;
                {font_family}
            }}
            QMenu::item:selected {{
                background-color: #1976D2;
                {font_family}
            }}
            /* Enhanced Status Bar Styling for Dark Mode */
            QFrame#StatusBar {{
                background-color: #2d2d2d;
                border-top: 1px solid #555555;
                border-bottom: none;
                border-left: none;
                border-right: none;
            }}
            QLabel#FileInfoLabel {{
                color: #cccccc;
                font-size: 12px;
                padding: 2px 8px;
                background-color: transparent;
                border: none;
                {font_family}
            }}
            QLabel#PageTotalLabel {{
                color: #cccccc;
                font-size: 12px;
                background-color: transparent;
                border: none;
                {font_family}
            }}
            QPushButton#NavButton {{
                background-color: #555555;
                color: #ffffff;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                font-size: 14px;
                {font_family}
            }}
            QPushButton#NavButton:hover {{
                background-color: #666666;
                {font_family}
            }}
            QPushButton#NavButton:disabled {{
                background-color: #3d3d3d;
                color: #888888;
                {font_family}
            }}
            QPushButton#ZoomButton {{
                background-color: #1976D2;
                color: #ffffff;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                {font_family}
            }}
            QPushButton#ZoomButton:hover {{
                background-color: #1565C0;
                {font_family}
            }}
            QPushButton#ZoomButton:disabled {{
                background-color: #3d3d3d;
                color: #888888;
                {font_family}
            }}
            QSpinBox#PageSpinBox {{
                background-color: #3d3d3d;
                color: #ffffff;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 4px;
                selection-background-color: #1976D2;
                {font_family}
            }}
            QSpinBox#PageSpinBox:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QLabel#ZoomLabel {{
                color: #ffffff;
                font-size: 12px;
                font-weight: bold;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 4px;
                background-color: #3d3d3d;
                {font_family}
            }}
            /* Toolbar Styling for Dark Mode */
            QFrame#FileToolbar {{
                background-color: #2d2d2d;
                border-bottom: 1px solid #555555;
                border-top: none;
                border-left: none;
                border-right: none;
            }}
            QFrame#AnnotationToolbar {{
                background-color: #2d2d2d;
                border-bottom: 1px solid #555555;
                border-top: none;
                border-left: none;
                border-right: none;
            }}
            /* File Toolbar Components */
            QLineEdit#FilePathDisplay {{
                background-color: #3d3d3d;
                color: #ffffff;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 4px 8px;
                selection-background-color: #1976D2;
                {font_family}
            }}
            QLineEdit#FilePathDisplay:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QPushButton#BrowseButton {{
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border: none;
                border-radius: 4px;
                {font_family}
            }}
            QPushButton#BrowseButton:hover {{
                background-color: #45a049;
                {font_family}
            }}
            QPushButton#LoadNewPdfButton {{
                background-color: #2196F3;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border: none;
                border-radius: 4px;
                {font_family}
            }}
            QPushButton#LoadNewPdfButton:hover {{
                background-color: #1976D2;
                {font_family}
            }}
            /* Fix for any remaining white areas */
            * {{
                outline: none;
            }}
            QWidget[class=""] {{
                background-color: #1e1e1e;
                color: #ffffff;
                {font_family}
            }}
            QProgressBar::chunk {{
                background-color: #1976D2;
                border-radius: 2px;
            }}
            QScrollArea {{
                border: 1px solid #555555;
                background-color: #1e1e1e;
            }}
            /* Navigation and Home Page Styling */
            SectionCard {{
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 12px;
                margin: 8px;
                color: #ffffff;
                {font_family}
            }}
            SectionCard:hover {{
                border-color: #1976D2;
                background-color: #404040;
                {font_family}
            }}
            BookCard {{
                background-color: #2d2d2d;
                border: 2px solid #555555;
                border-radius: 12px;
                margin: 8px;
                color: #ffffff;
                {font_family}
            }}
            BookCard:hover {{
                border-color: #1976D2;
                background-color: #404040;
                {font_family}
            }}
            /* Priority-based border colors */
            BookCard[priority="2"] {{
                border-color: #FF5722;
            }}
            BookCard[priority="1"] {{
                border-color: #FF9800;
            }}
            BookCard[priority="0"] {{
                border-color: #555555;
            }}
            /* Recent Books Statistics Frame */
            QFrame#StatsFrame {{
                background-color: #2d2d2d;
                border: 1px solid #555555;
                border-radius: 8px;
                margin: 5px 0;
                color: #ffffff;
                {font_family}
            }}
            /* Statistics Labels */
            QLabel#StatsLabel {{
                font-weight: bold;
                color: #ffffff;
                font-size: 12px;
                background-color: transparent;
                {font_family}
            }}
            /* Small Progress Bar for Book Cards */
            QProgressBar#SmallProgressBar {{
                background-color: #3d3d3d;
                border: none;
                border-radius: 2px;
            }}
            QProgressBar#SmallProgressBar::chunk {{
                background-color: #1976D2;
                border-radius: 2px;
            }}
            /* Books Grid Area */
            QScrollArea#BooksScrollArea {{
                background-color: #1e1e1e;
                border: none;
            }}
            QScrollArea#BooksScrollArea > QWidget {{
                background-color: #1e1e1e;
            }}
            QScrollArea#BooksScrollArea QScrollBar:vertical {{
                background-color: #2d2d2d;
                border: none;
                width: 12px;
            }}
            QScrollArea#BooksScrollArea QScrollBar::handle:vertical {{
                background-color: #555555;
                border-radius: 6px;
                min-height: 20px;
            }}
            QWidget#BooksWidget {{
                background-color: #1e1e1e;
            }}
            QWidget#BooksViewport {{
                background-color: #1e1e1e;
            }}
            /* No Books Message */
            QLabel#NoBooksLabel {{
                color: #cccccc;
                font-size: 16px;
                padding: 50px;
                border: 2px dashed #555555;
                border-radius: 10px;
                background-color: #2d2d2d;
                {font_family}
            }}
            /* Filter Buttons */
            QPushButton#FilterButton {{
                background-color: #2d2d2d;
                color: #1976D2;
                border: 1px solid #1976D2;
                padding: 5px 15px;
                border-radius: 15px;
                font-size: 12px;
                margin: 2px;
                {font_family}
            }}
            QPushButton#FilterButton:hover {{
                background-color: #1976D2;
                color: white;
                {font_family}
            }}
            QPushButton#FilterButton:checked {{
                background-color: #1976D2;
                color: white;
                {font_family}
            }}
            /* Statistics Dashboard */
            QFrame#StatsHeader {{
                background-color: #1976D2;
                border-radius: 10px;
                padding: 15px;
                {font_family}
            }}
            /* Additional UI Components Dark Theme */
            QSpinBox {{
                border: 2px solid #555555;
                border-radius: 4px;
                padding: 8px;
                background-color: #2d2d2d;
                color: #ffffff;
                font-size: 11px;
                min-height: 20px;
                {font_family}
            }}
            QSpinBox:hover {{
                border-color: #777777;
                background-color: #353535;
                {font_family}
            }}
            QSpinBox::up-button, QSpinBox::down-button {{
                background-color: #404040;
                border: 1px solid #555555;
                width: 16px;
            }}
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {{
                background-color: #555555;
            }}
            QCheckBox {{
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
                {font_family}
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
                border: 2px solid #555555;
                border-radius: 3px;
                background-color: #2d2d2d;
            }}
            QCheckBox::indicator:checked {{
                background-color: #1976D2;
                border-color: #1976D2;
            }}
            QRadioButton {{
                color: #ffffff;
                font-size: 11px;
                spacing: 8px;
                {font_family}
            }}
            QRadioButton::indicator {{
                width: 16px;
                height: 16px;
                border: 2px solid #555555;
                border-radius: 8px;
                background-color: #2d2d2d;
            }}
            QRadioButton::indicator:checked {{
                background-color: #1976D2;
                border-color: #1976D2;
            }}
            QSlider::groove:horizontal {{
                border: 1px solid #555555;
                height: 8px;
                background-color: #2d2d2d;
                border-radius: 4px;
            }}
            QSlider::handle:horizontal {{
                background-color: #1976D2;
                border: 2px solid #555555;
                width: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }}
            QSlider::handle:horizontal:hover {{
                background-color: #404040;
            }}
            QMessageBox {{
                background-color: #2d2d2d;
                color: #ffffff;
                {font_family}
            }}
            QMessageBox QLabel {{
                color: #ffffff;
                font-size: 12px;
                {font_family}
            }}
            QDialog {{
                background-color: #2d2d2d;
                color: #ffffff;
                {font_family}
            }}
            QFileDialog {{
                background-color: #2d2d2d;
                color: #ffffff;
                {font_family}
            }}
            QMenu {{
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #555555;
                {font_family}
            }}
            QMenu::item {{
                padding: 8px 16px;
                background-color: transparent;
                {font_family}
            }}
            QMenu::item:selected {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
            }}
            QMenuBar {{
                background-color: #2d2d2d;
                color: #ffffff;
                border-bottom: 1px solid #555555;
                {font_family}
            }}
            QMenuBar::item {{
                padding: 8px 12px;
                background-color: transparent;
                {font_family}
            }}
            QMenuBar::item:selected {{
                background-color: #404040;
                {font_family}
            }}
            QToolBar {{
                background-color: #2d2d2d;
                border: 1px solid #555555;
                spacing: 3px;
                color: #ffffff;
                {font_family}
            }}
            QToolButton {{
                background-color: #2d2d2d;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 6px;
                color: #ffffff;
                {font_family}
            }}
            QToolButton:hover {{
                background-color: #404040;
                border-color: #777777;
                {font_family}
            }}
            QToolButton:pressed {{
                background-color: #1976D2;
                {font_family}
            }}
            QStatusBar {{
                background-color: #2d2d2d;
                color: #ffffff;
                border-top: 1px solid #555555;
                {font_family}
            }}
            QSplitter::handle {{
                background-color: #555555;
            }}
            QSplitter::handle:hover {{
                background-color: #777777;
            }}
            """
            app.setStyleSheet(dark_style)
        else:
            # Light theme - explicitly set light styling with font integration
            font_family = get_font_family_css()
            light_style = f"""
            * {{
                {font_family}
            }}
            QMainWindow {{
                background-color: #ffffff;
                color: #000000;
                {font_family}
            }}
            QWidget {{
                background-color: #ffffff;
                color: #000000;
                {font_family}
            }}
            QTabWidget::pane {{
                border: 1px solid #cccccc;
                background-color: #ffffff;
                {font_family}
            }}
            QTabBar::tab {{
                background-color: #f0f0f0;
                color: #000000;
                padding: 8px 16px;
                margin: 2px;
                border-radius: 4px;
                border: 1px solid #cccccc;
                {font_family}
                font-weight: bold;
            }}
            QTabBar::tab:selected {{
                background-color: #1976D2;
                color: #ffffff;
                {font_family}
                font-weight: bold;
            }}
            QGroupBox {{
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin: 10px 0px;
                padding-top: 10px;
                background-color: #ffffff;
                color: #000000;
                {font_family}
                font-weight: bold;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #000000;
                {font_family}
                font-weight: bold;
            }}
            QLineEdit {{
                border: 2px solid #cccccc;
                border-radius: 4px;
                padding: 5px;
                background-color: #ffffff;
                color: #000000;
                {font_family}
            }}
            QLineEdit:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QPushButton {{
                border: 2px solid #cccccc;
                border-radius: 4px;
                padding: 5px;
                background-color: #f0f0f0;
                color: #000000;
                {font_family}
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: #e0e0e0;
                {font_family}
            }}
            QPushButton:pressed {{
                background-color: #d0d0d0;
                {font_family}
            }}
            QComboBox {{
                border: 2px solid #cccccc;
                border-radius: 4px;
                padding: 5px;
                background-color: #ffffff;
                color: #000000;
                {font_family}
            }}
            QComboBox:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QTableWidget {{
                background-color: #ffffff;
                color: #000000;
                gridline-color: #cccccc;
                border: 1px solid #cccccc;
                {font_family}
            }}
            QHeaderView::section {{
                background-color: #f0f0f0;
                color: #000000;
                padding: 4px;
                border: 1px solid #cccccc;
                {font_family}
            }}
            QTextEdit {{
                background-color: #ffffff;
                color: #000000;
                border: 2px solid #cccccc;
                border-radius: 4px;
                {font_family}
            }}
            QTextEdit:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QListWidget {{
                background-color: #ffffff;
                color: #000000;
                border: 2px solid #cccccc;
                border-radius: 4px;
                {font_family}
            }}
            QSpinBox {{
                background-color: #ffffff;
                color: #000000;
                border: 2px solid #cccccc;
                border-radius: 4px;
                padding: 5px;
                {font_family}
            }}
            QSpinBox:focus {{
                border-color: #1976D2;
                {font_family}
            }}
            QLabel {{
                color: #000000;
                background-color: transparent;
                {font_family}
            }}
            QStatusBar {{
                background-color: #f0f0f0;
                color: #000000;
                border-top: 1px solid #cccccc;
                {font_family}
            }}
            /* Modern scrollbar styling for light theme */
            QScrollBar:vertical {{
                background-color: #f5f5f5;
                width: 14px;
                border: none;
                border-radius: 7px;
                margin: 2px;
            }}
            QScrollBar::handle:vertical {{
                background-color: #c0c0c0;
                border-radius: 6px;
                min-height: 30px;
                margin: 2px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: #a0a0a0;
            }}
            QScrollBar::handle:vertical:pressed {{
                background-color: #1976D2;
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0px;
            }}
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
                background: none;
            }}
            QScrollBar:horizontal {{
                background-color: #f5f5f5;
                height: 14px;
                border: none;
                border-radius: 7px;
                margin: 2px;
            }}
            QScrollBar::handle:horizontal {{
                background-color: #c0c0c0;
                border-radius: 6px;
                min-width: 30px;
                margin: 2px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background-color: #a0a0a0;
            }}
            QScrollBar::handle:horizontal:pressed {{
                background-color: #1976D2;
            }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
                width: 0px;
            }}
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {{
                background: none;
            }}
            /* Light theme for cards */
            SectionCard {{
                background-color: #ffffff;
                border: 2px solid #cccccc;
                border-radius: 12px;
                margin: 8px;
                color: #000000;
                {font_family}
            }}
            SectionCard:hover {{
                border-color: #1976D2;
                background-color: #f5f5f5;
                {font_family}
            }}
            BookCard {{
                background-color: #ffffff;
                border: 2px solid #cccccc;
                border-radius: 12px;
                margin: 8px;
                color: #000000;
                {font_family}
            }}
            BookCard:hover {{
                border-color: #1976D2;
                background-color: #f5f5f5;
                {font_family}
            }}
            """
            app.setStyleSheet(light_style)

        # Update home page theme if navigation manager exists
        if hasattr(self, 'navigation_manager') and self.navigation_manager:
            if hasattr(self.navigation_manager, 'home_page') and self.navigation_manager.home_page:
                self.navigation_manager.home_page.apply_theme(theme == "dark")

        # Update Statistics Dashboard theme if it exists and has been loaded
        if hasattr(self, '_statistics_dashboard_tab_loader') and self._statistics_dashboard_tab_loader:
            if self._statistics_dashboard_tab_loader.is_loaded():
                self.statistics_dashboard_tab.apply_theme(theme == "dark")

    def closeEvent(self, event):
        """Handle application close event"""
        # Save window geometry
        geometry = [self.x(), self.y(), self.width(), self.height()]
        self.settings.set("window_geometry", geometry)
        event.accept()

    def show_about(self):
        """Show about dialog with proper branding and version info"""
        if VERSION_SYSTEM_AVAILABLE:
            app_info = get_app_info()
            version_info = get_version_info()
            app_name = app_info['display_name']
            version = version_info['version']
            author = app_info['author']
            website = app_info['website']
        else:
            app_name = "عدة القارئ للملفات الرقمية - PDF Toolkits"
            version = "2.0.0"
            author = "Zemam Productivity Apps"
            website = "https://github.com/Mohamedabdeltawab86/PDF-Tools"

        if self.localization.current_language == "ar":
            about_text = (
                f"📚 <b>{app_name}</b><br><br>"
                "مجموعة كاملة لمعالجة ملفات PDF مع دعم العربية والإنجليزية.<br><br>"
                "<b>الميزات الرئيسية:</b><br>"
                "• 🔖 إدارة الفهارس - إنشاء وتحرير فهارس PDF<br>"
                "• 🖼️ استخراج الصور - استخراج جميع الصور من PDF<br>"
                "• 📄 عمليات الصفحات - إدراج وحذف ودمج الصفحات<br>"
                "• 💧 العلامة المائية - إضافة نص أو صور كعلامة مائية<br>"
                "• 📚 الكتب الحديثة - تتبع تقدم القراءة والإحصائيات<br>"
                "• ⏱️ قياس سرعة القراءة - قياس الكلمات في الدقيقة مع اختبار الفهم<br>"
                "• 📊 لوحة الإحصائيات - تحليل شامل للاستخدام والأداء<br>"
                "• 📋 تتبع التعليقات - إدارة التعليقات والملاحظات<br>"
                "• 🔍 استخراج النص - تحويل PDF إلى نص قابل للتحرير<br>"
                "• 🗜️ ضغط PDF - تقليل حجم الملفات<br>"
                "• ⚙️ إعدادات متقدمة - السمات المظلمة/الفاتحة واللغات<br><br>"
                f"<b>الإصدار:</b> {version}<br>"
                f"<b>المطور:</b> {author}<br>"
                f"<b>الموقع:</b> <a href='{website}'>{website}</a><br>"
                "<b>الترخيص:</b> MIT License<br><br>"
                "<b>المطور الرئيسي:</b> محمد عبد التواب<br>"
                "<b>مساعد التطوير:</b> Augment Agent<br><br>"
                "<b>مبني بـ:</b> Python, PySide6, PyMuPDF, SQLite, Matplotlib"
            )
        else:
            about_text = (
                f"📚 <b>{app_name}</b><br><br>"
                "A comprehensive PDF manipulation suite with Arabic/English support.<br><br>"
                "<b>Key Features:</b><br>"
                "• 🔖 Bookmark Manager - Create and edit PDF bookmarks<br>"
                "• 🖼️ Image Extraction - Extract all images from PDFs<br>"
                "• 📄 Page Operations - Insert, delete, and merge pages<br>"
                "• 💧 Watermarking - Add text or image watermarks<br>"
                "• 📚 Recent Books - Track reading progress and statistics<br>"
                "• ⏱️ Reading Speed - Measure WPM with comprehension testing<br>"
                "• 📊 Statistics Dashboard - Comprehensive usage analytics<br>"
                "• 📋 Comment Tracking - Manage comments and annotations<br>"
                "• 🔍 Text Extraction - Convert PDFs to editable text<br>"
                "• 🗜️ PDF Compression - Reduce file sizes<br>"
                "• ⚙️ Advanced Settings - Dark/Light themes and languages<br><br>"
                f"<b>Version:</b> {version}<br>"
                f"<b>Developer:</b> {author}<br>"
                f"<b>Website:</b> <a href='{website}'>{website}</a><br>"
                "<b>License:</b> MIT License<br><br>"
                "<b>Lead Developer:</b> Mohamed Abdeltawab<br>"
                "<b>Development Assistant:</b> Augment Agent<br><br>"
                "<b>Built with:</b> Python, PySide6, PyMuPDF, SQLite, Matplotlib"
            )

        QMessageBox.about(self, self.localization.get_text("about"), about_text)

    def init_tabs(self):
        """Initialize all tabs with lazy loading for improved startup performance"""
        print("📦 Initializing tabs with lazy loading...")

        # Initialize books manager FIRST but defer database initialization
        # The database will be initialized when first accessed
        self._books_manager_loader = None
        try:
            def create_books_manager():
                from recent_books_manager import RecentBooksManager
                return RecentBooksManager()
            self._books_manager_loader = LazyTabLoader(create_books_manager)
        except Exception as e:
            print(f"Error setting up RecentBooksManager loader: {e}")
            self._books_manager_loader = None

        # Create lazy loaders for all tabs
        # These tabs will only be created when first accessed

        self._bookmark_tab_loader = LazyTabLoader(
            BookmarkTab, self.history_manager, self.localization, self
        )
        self._page_ops_tab_loader = LazyTabLoader(
            PageOperationsTab, self.history_manager, self.localization
        )
        self._bookmark_extractor_tab_loader = LazyTabLoader(
            BookmarkExtractorTab, self.history_manager, self.localization
        )
        self._split_by_bookmarks_tab_loader = LazyTabLoader(
            SplitByBookmarksTab, self.history_manager, self.localization, self
        )

        # Chapter Weight Visualizer tab (deferred import)
        def create_chapter_weight_tab():
            try:
                from chapter_weight_visualizer_tab import ChapterWeightVisualizerTab
                return ChapterWeightVisualizerTab(self.localization, self.settings, self)
            except ImportError as e:
                print(f"Warning: Could not import Chapter Weight Visualizer tab: {e}")
                return None
        self._chapter_weight_tab_loader = LazyTabLoader(create_chapter_weight_tab)

        self._watermark_tab_loader = LazyTabLoader(
            WatermarkTab, self.history_manager, self.localization
        )
        self._image_extraction_tab_loader = LazyTabLoader(
            ImageExtractionTab, self.history_manager, self.localization
        )
        self._text_extraction_tab_loader = LazyTabLoader(
            TextExtractionTab, self.history_manager, self.localization
        )
        self._merge_tab_loader = LazyTabLoader(
            MergeTab, self.history_manager, self.localization
        )

        # Split PDF tab (deferred import)
        def create_split_tab():
            try:
                from split_pdf_tab import SplitPDFTab
                return SplitPDFTab(self.history_manager, self.localization, self)
            except ImportError as e:
                print(f"Warning: Could not import Split PDF tab: {e}")
                return None
        self._split_tab_loader = LazyTabLoader(create_split_tab)

        self._compress_tab_loader = LazyTabLoader(
            CompressTab, self.history_manager, self.localization
        )
        self._page_editing_tab_loader = LazyTabLoader(
            PageEditingTab, self.history_manager, self.localization
        )
        self._security_removal_tab_loader = LazyTabLoader(
            PDFSecurityRemovalTab, self.history_manager, self.localization
        )

        # PDF Viewer tab needs books_manager, so we use a factory function
        def create_pdf_viewer_tab():
            return PDFViewerTab(self.history_manager, self.localization, self.books_manager)
        self._pdf_viewer_tab_loader = LazyTabLoader(create_pdf_viewer_tab)

        # PDF Reading Progress tab (deferred import)
        def create_reading_progress_tab():
            try:
                from pdf_reading_progress_tab import PDFReadingProgressTab
                return PDFReadingProgressTab(self.history_manager, self.localization)
            except ImportError as e:
                print(f"Warning: Could not import PDF Reading Progress tab: {e}")
                return None
        self._reading_progress_tab_loader = LazyTabLoader(create_reading_progress_tab)

        # Recent Books tab (deferred import)
        def create_recent_books_tab():
            try:
                from recent_books_tab import RecentBooksTab
                if self.books_manager:
                    return RecentBooksTab(
                        manager=self.books_manager,
                        localization=self.localization,
                        on_open_book=self.open_book_from_recent,
                        on_measure_speed=self.open_reading_speed_with_book,
                        parent=self
                    )
                else:
                    return None
            except ImportError as e:
                print(f"Warning: Could not import Recent Books tab: {e}")
                return None
        self._recent_books_tab_loader = LazyTabLoader(create_recent_books_tab)

        # Statistics Dashboard tab (deferred import)
        def create_statistics_dashboard_tab():
            try:
                from statistics_dashboard_tab import StatisticsDashboardTab
                return StatisticsDashboardTab(self.books_manager, self.localization)
            except ImportError as e:
                print(f"Warning: Could not import Statistics Dashboard tab: {e}")
                return None
        self._statistics_dashboard_tab_loader = LazyTabLoader(create_statistics_dashboard_tab)

        # Reading Speed tab (deferred import)
        def create_reading_speed_tab():
            try:
                import sys
                import os
                sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                from reading_speed_tab import ReadingSpeedTab
                return ReadingSpeedTab(
                    localization=self.localization,
                    recent_books_manager=self.books_manager,
                    parent=self
                )
            except ImportError as e:
                print(f"Warning: Could not import Reading Speed tab: {e}")
                return None
        self._reading_speed_tab_loader = LazyTabLoader(create_reading_speed_tab)

        # Settings and History tabs - these are lightweight, create immediately
        self._settings_tab_loader = LazyTabLoader(
            SettingsTab, self.localization, self.settings, self
        )
        self._history_tab_loader = LazyTabLoader(
            HistoryTab, self.localization, self.history_manager
        )

        print("✅ Tab loaders initialized (tabs will be created on first access)")

    # Property accessors for lazy-loaded tabs
    @property
    def books_manager(self):
        """Lazy-loaded books manager"""
        if self._books_manager_loader:
            return self._books_manager_loader.get_instance()
        return None

    @property
    def bookmark_tab(self):
        """Lazy-loaded bookmark tab"""
        return self._bookmark_tab_loader.get_instance()

    @property
    def page_ops_tab(self):
        """Lazy-loaded page operations tab"""
        return self._page_ops_tab_loader.get_instance()

    @property
    def bookmark_extractor_tab(self):
        """Lazy-loaded bookmark extractor tab"""
        return self._bookmark_extractor_tab_loader.get_instance()

    @property
    def split_by_bookmarks_tab(self):
        """Lazy-loaded split by bookmarks tab"""
        return self._split_by_bookmarks_tab_loader.get_instance()

    @property
    def chapter_weight_tab(self):
        """Lazy-loaded chapter weight tab"""
        return self._chapter_weight_tab_loader.get_instance()

    @property
    def watermark_tab(self):
        """Lazy-loaded watermark tab"""
        return self._watermark_tab_loader.get_instance()

    @property
    def image_extraction_tab(self):
        """Lazy-loaded image extraction tab"""
        return self._image_extraction_tab_loader.get_instance()

    @property
    def text_extraction_tab(self):
        """Lazy-loaded text extraction tab"""
        return self._text_extraction_tab_loader.get_instance()

    @property
    def merge_tab(self):
        """Lazy-loaded merge tab"""
        return self._merge_tab_loader.get_instance()

    @property
    def split_tab(self):
        """Lazy-loaded split tab"""
        return self._split_tab_loader.get_instance()

    @property
    def compress_tab(self):
        """Lazy-loaded compress tab"""
        return self._compress_tab_loader.get_instance()

    @property
    def page_editing_tab(self):
        """Lazy-loaded page editing tab"""
        return self._page_editing_tab_loader.get_instance()

    @property
    def security_removal_tab(self):
        """Lazy-loaded security removal tab"""
        return self._security_removal_tab_loader.get_instance()

    @property
    def pdf_viewer_tab(self):
        """Lazy-loaded PDF viewer tab"""
        return self._pdf_viewer_tab_loader.get_instance()

    @property
    def reading_progress_tab(self):
        """Lazy-loaded reading progress tab"""
        return self._reading_progress_tab_loader.get_instance()

    @property
    def recent_books_tab(self):
        """Lazy-loaded recent books tab"""
        return self._recent_books_tab_loader.get_instance()

    @property
    def statistics_dashboard_tab(self):
        """Lazy-loaded statistics dashboard tab"""
        return self._statistics_dashboard_tab_loader.get_instance()

    @property
    def reading_speed_tab(self):
        """Lazy-loaded reading speed tab"""
        return self._reading_speed_tab_loader.get_instance()

    @property
    def settings_tab(self):
        """Lazy-loaded settings tab"""
        return self._settings_tab_loader.get_instance()

    @property
    def history_tab(self):
        """Lazy-loaded history tab"""
        return self._history_tab_loader.get_instance()

    def open_pdf_file(self):
        """Open PDF file dialog"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.localization.get_text("open_pdf"),
            "",
            self.localization.get_text("pdf_files_filter")
        )
        if file_path:
            # Switch to PDF viewer and load the file
            self.navigation_manager.navigate_to_section("pdf_viewer")
            # You can add logic here to load the file in the current tab
            self.statusBar().showMessage(f"{self.localization.get_text('files_loaded_successfully')}: {os.path.basename(file_path)}")

    def open_book_from_recent(self, file_path: str):
        """Open a book from the Recent Books tab"""
        if not os.path.exists(file_path):
            QMessageBox.warning(
                self,
                "File Not Found",
                f"The file '{file_path}' could not be found. It may have been moved or deleted."
            )
            return

        try:
            # Add to recent books if not already there
            if self.books_manager:
                self.books_manager.add_book(file_path)

            # Switch to PDF viewer and load the file
            self.navigation_manager.navigate_to_section("pdf_viewer")

            # Load the file in the PDF viewer if available
            if hasattr(self, 'pdf_viewer_tab') and self.pdf_viewer_tab:
                self.pdf_viewer_tab.pdf_path = file_path
                if hasattr(self.pdf_viewer_tab, 'load_pdf'):
                    self.pdf_viewer_tab.load_pdf()

            self.statusBar().showMessage(f"Opened: {os.path.basename(file_path)}")

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error Opening File",
                f"Failed to open '{os.path.basename(file_path)}':\n{str(e)}"
            )

    def open_reading_speed_with_book(self, file_path: str):
        """Open reading speed measurement with a pre-selected book"""
        if not os.path.exists(file_path):
            QMessageBox.warning(
                self,
                self.localization.get_text("error"),
                f"The file '{file_path}' could not be found. It may have been moved or deleted."
            )
            return

        try:
            # Navigate to reading speed tab
            if hasattr(self, 'reading_speed_tab') and self.reading_speed_tab:
                self.navigation_manager.navigate_to_section("reading_speed")
                # Pre-load the book
                self.reading_speed_tab.load_pdf(file_path)
                self.statusBar().showMessage(f"Ready to measure reading speed for: {os.path.basename(file_path)}")
            else:
                QMessageBox.information(
                    self,
                    self.localization.get_text("info"),
                    "Reading Speed feature is not available."
                )

        except Exception as e:
            QMessageBox.critical(
                self,
                self.localization.get_text("error"),
                f"Failed to load book for speed measurement:\n{str(e)}"
            )

    def check_for_updates(self):
        """Manually check for updates"""
        if VERSION_SYSTEM_AVAILABLE and hasattr(self, 'update_checker'):
            self.update_checker.check_for_updates(silent=False)
        else:
            QMessageBox.information(
                self,
                "Update Check",
                "Update checking is not available in this version."
            )



    def show_user_guide(self):
        """Show comprehensive user guide dialog"""
        # Create a scrollable dialog for the user guide
        guide_dialog = QDialog(self)
        guide_dialog.setWindowTitle(self.localization.get_text("user_guide"))
        guide_dialog.setWindowIcon(self.get_app_icon())
        guide_dialog.resize(900, 700)

        # Create layout
        layout = QVBoxLayout(guide_dialog)

        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)

        # Create content widget
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)

        # Create text widget
        text_widget = QTextEdit()
        text_widget.setReadOnly(True)

        if self.localization.current_language == "ar":
            guide_text = self._get_arabic_user_guide()
        else:
            guide_text = self._get_english_user_guide()

        text_widget.setHtml(guide_text)
        content_layout.addWidget(text_widget)

        # Set content widget to scroll area
        scroll_area.setWidget(content_widget)
        layout.addWidget(scroll_area)

        # Add close button
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        close_btn = QPushButton(self.localization.get_text("close"))
        close_btn.clicked.connect(guide_dialog.close)
        close_btn.setStyleSheet("QPushButton { background-color: #1976D2; color: white; font-weight: bold; padding: 8px 16px; }")
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)

        # Apply theme
        current_theme = self.settings.get("theme", "dark")
        if current_theme == "dark":
            guide_dialog.setStyleSheet("""
                QDialog {
                    background-color: #1e1e1e;
                    color: #ffffff;
                }
                QTextEdit {
                    background-color: #2d2d2d;
                    color: #ffffff;
                    border: 1px solid #555555;
                }
                QScrollArea {
                    background-color: #1e1e1e;
                    border: none;
                }
            """)

        guide_dialog.exec()

    def _get_english_user_guide(self):
        """Get comprehensive English user guide content"""
        return """
        <h2 style="color: #1976D2; text-align: center;">📖 Comprehensive User Guide</h2>
        <h3 style="color: #1976D2;">Welcome to PDF Toolkits!</h3>

        <p><b>This application provides a comprehensive suite of PDF manipulation tools with full Arabic and English support.</b></p>

        <h3 style="color: #1976D2;">🔖 Core Tools</h3>

        <h4>📖 Bookmark Manager</h4>
        <p><b>Function:</b> Create and edit PDF bookmarks with smart offset detection</p>
        <p><b>How to use:</b></p>
        <ul>
        <li>Select a PDF file using "Choose PDF File" button</li>
        <li>Enter bookmark text in the text box (one line = one bookmark)</li>
        <li>Choose bookmark level (1-6) for each entry</li>
        <li>Click "Insert Bookmarks" to add bookmarks to the file</li>
        </ul>

        <h4>📄 Page Operations</h4>
        <p><b>Function:</b> Insert, delete, merge, and rotate pages</p>
        <p><b>Available operations:</b></p>
        <ul>
        <li><b>Extract Pages:</b> Extract specific pages to a new file</li>
        <li><b>Delete Pages:</b> Remove pages from the file</li>
        <li><b>Insert Blank Pages:</b> Add blank pages at specified positions</li>
        <li><b>Merge PDFs:</b> Combine multiple files into one</li>
        <li><b>Rotate Pages:</b> Rotate pages at different angles</li>
        </ul>

        <h4>🖼️ Image Extraction</h4>
        <p><b>Function:</b> Extract all images from PDF files</p>
        <p><b>Steps:</b></p>
        <ul>
        <li>Choose source PDF file</li>
        <li>Select save folder for images</li>
        <li>Choose image format (PNG, JPEG)</li>
        <li>Click "Extract Images" to start the process</li>
        </ul>

        <h4>📝 Text Extraction</h4>
        <p><b>Function:</b> Convert PDF content to editable text</p>
        <p><b>Supported formats:</b></p>
        <ul>
        <li>Plain text (.txt)</li>
        <li>Word document (.docx)</li>
        <li>HTML (.html)</li>
        <li>Markdown (.md)</li>
        </ul>

        <h3 style="color: #1976D2;">📊 Analytics and Reading Tools</h3>

        <h4>⏱️ Reading Speed Measurement</h4>
        <p><b>Function:</b> Measure reading speed in words per minute with comprehension testing</p>
        <p><b>How to use:</b></p>
        <ul>
        <li>Choose a PDF file to read</li>
        <li>Click "Start Reading" when ready</li>
        <li>Read the displayed text</li>
        <li>Click "Finished Reading" when done</li>
        <li>Answer comprehension questions</li>
        <li>Get detailed performance report</li>
        </ul>

        <h4>📊 Statistics Dashboard</h4>
        <p><b>Function:</b> Comprehensive analysis of app usage and performance</p>
        <p><b>Available information:</b></p>
        <ul>
        <li>Reading statistics (books, pages, time)</li>
        <li>Progress charts</li>
        <li>Performance analysis</li>
        <li>Activity log</li>
        </ul>

        <h4>📚 Recent Books</h4>
        <p><b>Function:</b> Track reading progress and maintain reading history</p>
        <p><b>Features:</b></p>
        <ul>
        <li>Save last read page</li>
        <li>Display book thumbnails</li>
        <li>Track reading time</li>
        <li>Detailed statistics for each book</li>
        </ul>

        <h3 style="color: #1976D2;">⚙️ Settings and Customization</h3>

        <h4>🎨 Themes</h4>
        <ul>
        <li><b>Dark Theme:</b> Suitable for low-light reading</li>
        <li><b>Light Theme:</b> Suitable for natural lighting</li>
        </ul>

        <h4>🌐 Languages</h4>
        <ul>
        <li><b>Arabic:</b> Full right-to-left text support</li>
        <li><b>English:</b> Complete English interface</li>
        </ul>

        <h3 style="color: #1976D2;">⌨️ Keyboard Shortcuts</h3>
        <table border="1" cellpadding="5" cellspacing="0" style="width: 100%;">
        <tr style="background-color: #1976D2; color: white;"><th>Shortcut</th><th>Function</th></tr>
        <tr><td><b>Ctrl+O</b></td><td>Open PDF file</td></tr>
        <tr><td><b>Ctrl+W</b></td><td>Close application</td></tr>
        <tr><td><b>Ctrl+Q</b></td><td>Exit application</td></tr>
        <tr><td><b>F1</b></td><td>Show user guide</td></tr>
        <tr><td><b>Tab</b></td><td>Navigate between fields</td></tr>
        <tr><td><b>Enter</b></td><td>Execute operation</td></tr>
        <tr><td><b>Escape</b></td><td>Cancel operation</td></tr>
        </table>

        <h3 style="color: #1976D2;">💡 Tips and Best Practices</h3>

        <h4>📄 Page Range Formats</h4>
        <ul>
        <li><b>Single page:</b> 5</li>
        <li><b>Multiple pages:</b> 1,3,5,7</li>
        <li><b>Range:</b> 1-10</li>
        <li><b>Open range:</b> 5- (from page 5 to end)</li>
        <li><b>Mixed:</b> 1,3,5-10,15</li>
        </ul>

        <h4>🔧 Troubleshooting</h4>
        <ul>
        <li><b>Protected PDF:</b> Ensure the file is not password protected</li>
        <li><b>Large file:</b> Operations may take longer for large files</li>
        <li><b>Disk space:</b> Ensure sufficient space for saving new files</li>
        <li><b>Memory:</b> Close other applications when processing large files</li>
        </ul>

        <h4>📈 Performance Optimization</h4>
        <ul>
        <li>Use PDF compression to reduce file sizes</li>
        <li>Split large files into smaller parts</li>
        <li>Save your work regularly</li>
        <li>Check operation history to verify results</li>
        </ul>

        <h3 style="color: #1976D2;">🆘 Getting Help</h3>
        <p>If you encounter any issues or have suggestions:</p>
        <ul>
        <li><b>GitHub Issues:</b> <a href="https://github.com/Mohamedabdeltawab86/PDF-Tools/issues">Report a problem</a></li>
        <li><b>Email:</b> dr.m.tawab@outlook.com</li>
        <li><b>Website:</b> <a href="https://github.com/Mohamedabdeltawab86/PDF-Tools">Project page</a></li>
        </ul>

        <hr>
        <p style="text-align: center; color: #666;"><i>Developed by Zemam Productivity Apps</i></p>
        """

    def _get_arabic_user_guide(self):
        """Get comprehensive Arabic user guide content"""
        return """
        <div style="direction: rtl; text-align: right;">
        <h2 style="color: #1976D2; text-align: center;">📖 دليل المستخدم الشامل</h2>
        <h3 style="color: #1976D2;">مرحباً بك في عدة القارئ للملفات الرقمية - PDF Toolkits!</h3>

        <p><b>هذا التطبيق يوفر مجموعة شاملة من أدوات معالجة ملفات PDF مع دعم كامل للغة العربية والإنجليزية.</b></p>

        <h3 style="color: #1976D2;">🔖 الأدوات الأساسية</h3>

        <h4>📖 مدير الفهارس</h4>
        <p><b>الوظيفة:</b> إنشاء وتحرير فهارس PDF مع كشف الإزاحة الذكي</p>
        <p><b>كيفية الاستخدام:</b></p>
        <ul>
        <li>اختر ملف PDF من خلال زر "اختيار ملف PDF"</li>
        <li>أدخل نص الفهارس في المربع النصي (كل سطر = إشارة مرجعية)</li>
        <li>اختر مستوى الفهرس (1-6) لكل إشارة</li>
        <li>انقر "إدراج الفهارس" لإضافة الفهارس إلى الملف</li>
        </ul>

        <h4>📄 عمليات الصفحات</h4>
        <p><b>الوظيفة:</b> إدراج وحذف ودمج وتدوير الصفحات</p>
        <p><b>العمليات المتاحة:</b></p>
        <ul>
        <li><b>استخراج الصفحات:</b> استخراج صفحات محددة إلى ملف جديد</li>
        <li><b>حذف الصفحات:</b> حذف صفحات من الملف</li>
        <li><b>إدراج صفحات فارغة:</b> إضافة صفحات فارغة في مواضع محددة</li>
        <li><b>دمج ملفات PDF:</b> دمج عدة ملفات في ملف واحد</li>
        <li><b>تدوير الصفحات:</b> تدوير الصفحات بزوايا مختلفة</li>
        </ul>

        <h4>🖼️ استخراج الصور</h4>
        <p><b>الوظيفة:</b> استخراج جميع الصور من ملفات PDF</p>
        <p><b>الخطوات:</b></p>
        <ul>
        <li>اختر ملف PDF المصدر</li>
        <li>حدد مجلد الحفظ للصور</li>
        <li>اختر تنسيق الصور (PNG, JPEG)</li>
        <li>انقر "استخراج الصور" لبدء العملية</li>
        </ul>

        <h4>📝 استخراج النص</h4>
        <p><b>الوظيفة:</b> تحويل محتوى PDF إلى نص قابل للتحرير</p>
        <p><b>التنسيقات المدعومة:</b></p>
        <ul>
        <li>نص عادي (.txt)</li>
        <li>مستند Word (.docx)</li>
        <li>HTML (.html)</li>
        <li>Markdown (.md)</li>
        </ul>

        <h3 style="color: #1976D2;">📊 أدوات التحليل والقراءة</h3>

        <h4>⏱️ قياس سرعة القراءة</h4>
        <p><b>الوظيفة:</b> قياس سرعة القراءة بالكلمات في الدقيقة مع اختبار الفهم</p>
        <p><b>كيفية الاستخدام:</b></p>
        <ul>
        <li>اختر ملف PDF للقراءة</li>
        <li>انقر "بدء القراءة" عند الاستعداد</li>
        <li>اقرأ النص المعروض</li>
        <li>انقر "انتهيت من القراءة" عند الانتهاء</li>
        <li>أجب على أسئلة الفهم</li>
        <li>احصل على تقرير مفصل عن أدائك</li>
        </ul>

        <h4>📊 لوحة الإحصائيات</h4>
        <p><b>الوظيفة:</b> تحليل شامل لاستخدام التطبيق والأداء</p>
        <p><b>المعلومات المتاحة:</b></p>
        <ul>
        <li>إحصائيات القراءة (عدد الكتب، الصفحات، الوقت)</li>
        <li>رسوم بيانية للتقدم</li>
        <li>تحليل الأداء</li>
        <li>سجل النشاط</li>
        </ul>

        <h4>📚 الكتب الحديثة</h4>
        <p><b>الوظيفة:</b> تتبع تقدم القراءة والاحتفاظ بسجل القراءة</p>
        <p><b>المميزات:</b></p>
        <ul>
        <li>حفظ آخر صفحة مقروءة</li>
        <li>عرض صور مصغرة للكتب</li>
        <li>تتبع وقت القراءة</li>
        <li>إحصائيات مفصلة لكل كتاب</li>
        </ul>

        <h3 style="color: #1976D2;">⚙️ الإعدادات والتخصيص</h3>

        <h4>🎨 السمات</h4>
        <ul>
        <li><b>السمة المظلمة:</b> مناسبة للقراءة في الإضاءة المنخفضة</li>
        <li><b>السمة الفاتحة:</b> مناسبة للقراءة في الإضاءة الطبيعية</li>
        </ul>

        <h4>🌐 اللغات</h4>
        <ul>
        <li><b>العربية:</b> دعم كامل للنصوص من اليمين إلى اليسار</li>
        <li><b>الإنجليزية:</b> واجهة إنجليزية كاملة</li>
        </ul>

        <h3 style="color: #1976D2;">⌨️ اختصارات لوحة المفاتيح</h3>
        <table border="1" cellpadding="5" cellspacing="0" style="width: 100%;">
        <tr style="background-color: #1976D2; color: white;"><th>الاختصار</th><th>الوظيفة</th></tr>
        <tr><td><b>Ctrl+O</b></td><td>فتح ملف PDF</td></tr>
        <tr><td><b>Ctrl+W</b></td><td>إغلاق التطبيق</td></tr>
        <tr><td><b>Ctrl+Q</b></td><td>خروج من التطبيق</td></tr>
        <tr><td><b>F1</b></td><td>عرض دليل المستخدم</td></tr>
        <tr><td><b>Tab</b></td><td>التنقل بين الحقول</td></tr>
        <tr><td><b>Enter</b></td><td>تنفيذ العملية</td></tr>
        <tr><td><b>Escape</b></td><td>إلغاء العملية</td></tr>
        </table>

        <h3 style="color: #1976D2;">💡 نصائح وأفضل الممارسات</h3>

        <h4>📄 تنسيقات نطاقات الصفحات</h4>
        <ul>
        <li><b>صفحة واحدة:</b> 5</li>
        <li><b>صفحات متعددة:</b> 1,3,5,7</li>
        <li><b>نطاق:</b> 1-10</li>
        <li><b>نطاق مفتوح:</b> 5- (من الصفحة 5 حتى النهاية)</li>
        <li><b>مختلط:</b> 1,3,5-10,15</li>
        </ul>

        <h4>🔧 استكشاف الأخطاء وإصلاحها</h4>
        <ul>
        <li><b>ملف PDF محمي:</b> تأكد من أن الملف غير محمي بكلمة مرور</li>
        <li><b>ملف كبير:</b> قد تستغرق العمليات وقتاً أطول للملفات الكبيرة</li>
        <li><b>مساحة القرص:</b> تأكد من وجود مساحة كافية لحفظ الملفات الجديدة</li>
        <li><b>الذاكرة:</b> أغلق التطبيقات الأخرى عند معالجة ملفات كبيرة</li>
        </ul>

        <h4>📈 تحسين الأداء</h4>
        <ul>
        <li>استخدم ضغط PDF لتقليل أحجام الملفات</li>
        <li>قم بتقسيم الملفات الكبيرة إلى أجزاء أصغر</li>
        <li>احفظ عملك بانتظام</li>
        <li>راجع سجل العمليات للتحقق من النتائج</li>
        </ul>

        <h3 style="color: #1976D2;">🆘 الحصول على المساعدة</h3>
        <p>إذا واجهت أي مشاكل أو كان لديك اقتراحات:</p>
        <ul>
        <li><b>GitHub Issues:</b> <a href="https://github.com/Mohamedabdeltawab86/PDF-Tools/issues">الإبلاغ عن مشكلة</a></li>
        <li><b>البريد الإلكتروني:</b> dr.m.tawab@outlook.com</li>
        <li><b>الموقع:</b> <a href="https://github.com/Mohamedabdeltawab86/PDF-Tools">صفحة المشروع</a></li>
        </ul>

        <hr>
        <p style="text-align: center; color: #666;"><i>تم تطوير هذا التطبيق بواسطة Zemam Productivity Apps</i></p>
        </div>
        """

    def toggle_maximize(self):
        """Toggle between maximized and normal window state"""
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def show_shortcuts(self):
        """Show keyboard shortcuts dialog"""
        if self.localization.current_language == "ar":
            shortcuts_text = (
                "<h3>⌨️ اختصارات لوحة المفاتيح</h3>"
                "<table border='1' cellpadding='5' cellspacing='0'>"
                "<tr><th>الاختصار</th><th>الوظيفة</th></tr>"
                "<tr><td><b>Ctrl+O</b></td><td>فتح ملف PDF</td></tr>"
                "<tr><td><b>Ctrl+W</b></td><td>إغلاق التطبيق</td></tr>"
                "<tr><td><b>Ctrl+Q</b></td><td>خروج</td></tr>"
                "<tr><td><b>F1</b></td><td>دليل المستخدم</td></tr>"
                "</table>"
                "<br><p><b>نصائح:</b></p>"
                "<ul>"
                "<li>استخدم Tab للتنقل بين الحقول</li>"
                "<li>استخدم Enter لتنفيذ العمليات</li>"
                "<li>استخدم Escape لإلغاء العمليات</li>"
                "</ul>"
            )
        else:
            shortcuts_text = (
                "<h3>⌨️ Keyboard Shortcuts</h3>"
                "<table border='1' cellpadding='5' cellspacing='0'>"
                "<tr><th>Shortcut</th><th>Function</th></tr>"
                "<tr><td><b>Ctrl+O</b></td><td>Open PDF File</td></tr>"
                "<tr><td><b>Ctrl+W</b></td><td>Close Application</td></tr>"
                "<tr><td><b>Ctrl+Q</b></td><td>Exit</td></tr>"
                "<tr><td><b>F1</b></td><td>User Guide</td></tr>"
                "</table>"
                "<br><p><b>Tips:</b></p>"
                "<ul>"
                "<li>Use Tab to navigate between fields</li>"
                "<li>Use Enter to execute operations</li>"
                "<li>Use Escape to cancel operations</li>"
                "</ul>"
            )

        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(self.localization.get_text("keyboard_shortcuts"))
        msg_box.setText(shortcuts_text)
        msg_box.setTextFormat(Qt.RichText)
        msg_box.exec()


def main():
    """Main function with performance timing"""
    import time

    # Track total startup time
    total_start = time.time()

    print("=" * 60)
    print("🚀 PDF Tools Application - Performance Optimized Launch")
    print("=" * 60)

    # Migrate old data files to app data directory (first run)
    if APP_PATHS_AVAILABLE:
        step_start = time.time()
        print("🔄 Checking for data migration...")
        migrated = migrate_old_data()
        if migrated:
            print(f"✅ Migrated {len(migrated)} files to user data directory")
        print(f"   ⏱️  Migration check: {time.time() - step_start:.3f}s")

    # Create QApplication
    step_start = time.time()
    app = QApplication(sys.argv)
    print(f"✅ QApplication created: {time.time() - step_start:.3f}s")

    # Set application properties
    app.setApplicationName("PDF Toolkits")
    app.setApplicationVersion("2.0.0")

    # Set application icon - try both ICO and PNG
    step_start = time.time()
    icon_paths = [
        os.path.join(os.path.dirname(__file__), '..', 'assets', 'icons', 'logo.ico'),
        os.path.join(os.path.dirname(__file__), '..', 'assets', 'icons', 'logo.png'),
        os.path.join(os.path.dirname(sys.executable), 'assets', 'icons', 'logo.ico'),  # For frozen app
        os.path.join(os.path.dirname(sys.executable), 'assets', 'icons', 'logo.png'),  # For frozen app
    ]

    app_icon = None
    for icon_path in icon_paths:
        if os.path.exists(icon_path):
            app_icon = QIcon(icon_path)
            if not app_icon.isNull():
                break

    if app_icon and not app_icon.isNull():
        app.setWindowIcon(app_icon)
    print(f"✅ Application icon set: {time.time() - step_start:.3f}s")

    # Load custom fonts
    step_start = time.time()
    load_custom_fonts()
    print(f"✅ Custom fonts loaded: {time.time() - step_start:.3f}s")

    # Apply global font styling
    step_start = time.time()
    apply_global_font_style()
    print(f"✅ Global font styling applied: {time.time() - step_start:.3f}s")

    # Create and show main window
    step_start = time.time()
    window = PDFToolsComprehensive()
    print(f"✅ Main window created: {time.time() - step_start:.3f}s")

    step_start = time.time()
    window.show()
    print(f"✅ Main window shown: {time.time() - step_start:.3f}s")

    # Calculate and display total startup time
    total_time = time.time() - total_start
    print("=" * 60)
    print(f"🎉 APPLICATION READY IN {total_time:.2f} SECONDS")
    print("=" * 60)
    print(f"📊 Performance Target: <3 seconds (Current: {total_time:.2f}s)")
    if total_time < 3:
        print("✅ TARGET ACHIEVED!")
    elif total_time < 5:
        print("⚠️  Good progress, but can be improved further")
    else:
        print("❌ Needs more optimization")
    print("=" * 60)

    # Run application
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
