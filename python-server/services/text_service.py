"""Text Extraction Service - Extract text from PDF to TXT/DOCX/MD"""
import os
from typing import List, Dict, Any, Optional
import fitz  # pymupdf


class TextService:
    @staticmethod
    def parse_page_range(page_range: str, total_pages: int) -> List[int]:
        """Parse page range string like '1-10,15,20-25' into list of page indices (0-based)."""
        if not page_range or page_range.lower() == 'all':
            return list(range(total_pages))
        
        pages = set()
        parts = page_range.replace(' ', '').split(',')
        for part in parts:
            if '-' in part:
                start, end = part.split('-', 1)
                start = max(1, int(start))
                end = min(total_pages, int(end))
                for p in range(start, end + 1):
                    pages.add(p - 1)  # Convert to 0-based
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p - 1)
        return sorted(pages)

    @staticmethod
    def extract_text(pdf_path: str, page_range: str = 'all') -> Dict[str, Any]:
        """Extract text from PDF."""
        try:
            doc = fitz.open(pdf_path)
            total_pages = doc.page_count  # Store before closing
            pages = TextService.parse_page_range(page_range, total_pages)
            
            text_parts = []
            for page_idx in pages:
                page = doc[page_idx]
                text_parts.append(page.get_text())
            
            doc.close()  # Close after all operations
            
            full_text = '\n\n'.join(text_parts)
            word_count = len(full_text.split())
            
            return {
                "text": full_text,
                "pages_extracted": len(pages),
                "total_pages": total_pages,
                "word_count": word_count
            }
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")

    @staticmethod
    def get_word_index(pdf_path: str) -> List[Dict[str, Any]]:
        """Get word count for each page in the PDF using PyMuPDF (Text only)."""
        print(f"[TEXT_SERVICE] Indexing PDF: {pdf_path}", flush=True)
        try:
            doc = fitz.open(pdf_path)
            index = []
            for i in range(doc.page_count):
                page = doc[i]
                # get_text() for text-based PDFs
                text = page.get_text()
                # Simple split based on whitespace
                word_count = len(text.split())
                index.append({
                    "page": i + 1,
                    "word_count": word_count
                })
            doc.close()
            print(f"[TEXT_SERVICE] Successfully indexed {len(index)} pages", flush=True)
            return index
        except Exception as e:
            print(f"[TEXT_SERVICE] Indexing failed: {str(e)}", flush=True)
            raise Exception(f"Failed to index PDF: {str(e)}")

    @staticmethod
    def extract_to_file(pdf_path: str, output_path: str, page_range: str = 'all', 
                        format: str = 'txt', merge_lines: bool = True,
                        font_family: str = 'Calibri', font_size: int = 11) -> Dict[str, Any]:
        """Extract text and save to file (TXT, DOCX, or MD)."""
        try:
            import time
            start_time = time.time()
            
            # Open PDF
            doc = fitz.open(pdf_path)
            total_pages = doc.page_count
            page_indices = TextService.parse_page_range(page_range, total_pages)
            
            structured_pages = []
            
            # Analyze each page
            for idx in page_indices:
                page = doc[idx]
                page_content = TextService._extract_structured_content(page)
                structured_pages.append(page_content)
                
            doc.close()

            # Save based on format
            if format.lower() == 'docx':
                TextService._save_as_docx_structured(structured_pages, output_path, font_family, font_size)
            elif format.lower() == 'md':
                TextService._save_as_md_structured(structured_pages, output_path)
            else:  # txt
                # Flatten for TXT
                flat_text = []
                for p in structured_pages:
                    p_text = "\n\n".join([item['text'] for item in p])
                    flat_text.append(p_text)
                final_text = "\n\n".join(flat_text)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(final_text)
            
            end_time = time.time()
            
            return {
                "output_path": output_path,
                "format": format,
                "duration_seconds": round(end_time - start_time, 2),
                "pages_extracted": len(page_indices),
                "total_pages": total_pages
            }
        except Exception as e:
            raise Exception(f"Extract to file failed: {str(e)}")

    @staticmethod
    def _extract_structured_content(page: fitz.Page) -> List[Dict[str, Any]]:
        """
        Extract content with style information (heading/body).
        Returns list of dicts: {'text': str, 'role': 'h1'|'h2'|'body', 'size': float}
        """
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        
        # 1. Collect all lines and font sizes
        all_lines = []
        font_sizes = []
        
        for block in blocks:
            if block['type'] != 0: continue # Skip images
            for line in block['lines']:
                # Get most common font size in line
                if not line['spans']: continue
                
                # Calculate avg/mode size for the line
                sizes = [s['size'] for s in line['spans']]
                avg_size = sum(sizes) / len(sizes)
                
                # Get text
                text = " ".join([s['text'] for s in line['spans']]).strip()
                if not text: continue
                
                all_lines.append({
                    'text': text,
                    'y0': line['bbox'][1],
                    'y1': line['bbox'][3],
                    'height': line['bbox'][3] - line['bbox'][1],
                    'size': avg_size,
                    'font': line['spans'][0]['font']
                })
                font_sizes.append(round(avg_size, 1))

        if not all_lines:
            return []

        # 2. Determine Body Text Size (Mode)
        if font_sizes:
            from collections import Counter
            body_size = Counter(font_sizes).most_common(1)[0][0]
        else:
            body_size = 11.0 # Default

        # Helper to check content
        import re
        def get_role(item_text, item_size):
            # 1. Filter out Symbols (e.g., "~", "...", "-")
            if re.match(r'^[^a-zA-Z0-9\u0600-\u06FF]+$', item_text):
                return 'body'
            
            # 2. Filter out Short Labels (e.g., "A", "A B", "1") usually map labels
            # Match single letter or sequence of single letters "A B C"
            if len(item_text) < 10 and re.match(r'^([A-Z0-9]\s?)+$', item_text, re.IGNORECASE):
                return 'body'

            # 3. Filter out Captions
            if re.match(r'^(Figure|Fig|Table|Chart|Image)\.?\s?\d+', item_text, re.IGNORECASE):
                return 'body'

            # 4. Filter out Page Numbers (single digits or "Page X")
            if re.match(r'^\d+$', item_text) or re.match(r'^Page\s\d+$', item_text, re.IGNORECASE):
                return 'body'

            # Size thresholds (Adjusted slightly to be more conservative)
            if item_size > body_size * 1.6: return 'h1'
            if item_size > body_size * 1.25: return 'h2'
            
            return 'body'

        # 3. Reflow and Assign Styles
        structured_content = []
        
        # Initialize first block
        first_line = all_lines[0]
        current_role = get_role(first_line['text'], first_line['size'])
        
        current_block = {
            'text': [first_line['text']],
            'size': first_line['size'],
            'role': current_role
        }
        
        for i in range(1, len(all_lines)):
            prev = all_lines[i-1]
            curr = all_lines[i]
            
            gap = curr['y0'] - prev['y1']
            ref_height = max(prev['height'], curr['height'])
            
            # Determine if we should split
            is_new_paragraph = gap > (ref_height * 0.6)
            is_size_change = abs(curr['size'] - prev['size']) > 0.5
            
            if is_new_paragraph or is_size_change:
                # Flush current block
                structured_content.append({
                    'text': " ".join(current_block['text']),
                    'role': current_block['role']
                })
                
                # Start new block
                new_role = get_role(curr['text'], curr['size'])
                current_block = {
                    'text': [curr['text']],
                    'size': curr['size'],
                    'role': new_role
                }
            else:
                # Merge into current block
                current_block['text'].append(curr['text'])
                
        # Flush final block
        structured_content.append({
            'text': " ".join(current_block['text']),
            'role': current_block['role']
        })
        
        return structured_content

    @staticmethod
    def _save_as_docx_structured(pages_content: List[List[Dict]], output_path: str,
                                 font_family: str = 'Calibri', font_size: int = 11):
        """Save using python-docx with styles."""
        try:
            from docx import Document
            from docx.shared import Pt
            import re
            
            def sanitize(t): return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', t)
            
            doc = Document()
            
            # Apply font settings to styles
            styles_to_update = ['Normal', 'Heading 1', 'Heading 2']
            for style_name in styles_to_update:
                if style_name in doc.styles:
                    font = doc.styles[style_name].font
                    font.name = font_family
                    # Only set size for Normal, let headings scale naturally or enforce if needed?
                    # User likely wants body text size control. Headings should perhaps remain relative or scaled.
                    # For now, let's only set Body size, but force Family on all.
                    if style_name == 'Normal':
                        font.size = Pt(font_size)
            
            for i, page in enumerate(pages_content):
                for block in page:
                    clean_text = sanitize(block['text'])
                    if not clean_text: continue
                    
                    paragraph = None
                    if block['role'] == 'h1':
                        paragraph = doc.add_heading(clean_text, level=1)
                    elif block['role'] == 'h2':
                        paragraph = doc.add_heading(clean_text, level=2)
                    else:
                        paragraph = doc.add_paragraph(clean_text)
                    
                    # Force font on runs to ensure it overrides any defaults
                    # (Sometimes Style modification isn't enough for specific environments)
                    if paragraph and paragraph.runs:
                        for run in paragraph.runs:
                            run.font.name = font_family
                            # Apply size to normal text runs
                            if block['role'] == 'body':
                                run.font.size = Pt(font_size)
                
                # Page break between pages
                if i < len(pages_content) - 1:
                    doc.add_page_break()
            
            doc.save(output_path)
        except Exception:
            raise

    @staticmethod
    def _save_as_md_structured(pages_content: List[List[Dict]], output_path: str):
        """Save as Markdown."""
        lines = []
        for i, page in enumerate(pages_content):
            for block in page:
                text = block['text']
                if block['role'] == 'h1':
                    lines.append(f"# {text}")
                elif block['role'] == 'h2':
                    lines.append(f"## {text}")
                else:
                    lines.append(text)
                lines.append("") # Empty line after block
            
            if i < len(pages_content) - 1:
                lines.append("---") # Horizontal rule for page break
                lines.append("")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
    @staticmethod
    def batch_extract(pdf_paths: List[str], output_dir: str, format: str = 'txt',
                      page_range: str = 'all') -> Dict[str, Any]:
        """Batch extract text from multiple PDFs."""
        results = {
            "successful": 0,
            "failed": 0,
            "output_files": [],
            "errors": []
        }
        
        for pdf_path in pdf_paths:
            try:
                base_name = os.path.splitext(os.path.basename(pdf_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.{format}")
                TextService.extract_to_file(pdf_path, output_path, page_range, format)
                results["successful"] += 1
                results["output_files"].append(output_path)
            except Exception as e:
                results["failed"] += 1
                results["errors"].append(f"{pdf_path}: {str(e)}")
        
        return results
